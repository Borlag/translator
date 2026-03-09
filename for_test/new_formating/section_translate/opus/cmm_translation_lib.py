"""
CMM Translation Library — shared module for translating aviation CMM documents (EN→RU).

Provides glossary-based terminology, translation functions, XML-level text replacement,
font handling, and the main translation pipeline. All 10 document parts use this library
for consistent terminology and formatting.

Usage:
    from cmm_translation_lib import translate_document
    translate_document(
        src="path/to/original_new_partN.docx",
        dst="path/to/translated_partN.docx",
    )
"""
import copy
import os
import re
import sys
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn


# ══════════════════════════════════════════════════════════════════════════════
#  GLOSSARIES  (EN → RU)
# ══════════════════════════════════════════════════════════════════════════════

# ── Component name glossary ──────────────────────────────────────────────────
COMPONENT_NAMES = {
    "Main Landing Gear Leg": "Стойка основного шасси",
    "Main Landing Gear": "Основное шасси",
    "MLG Shock Strut": "Стойка амортизатора основного шасси",
    "Main Fitting Subassembly": "Сборка корпуса стойки",
    "Main Fitting": "Корпус стойки",
    "Sliding Tube": "Скользящая труба",
    "Lower Torque Link": "Нижний шлиц-шарнир",
    "Upper Torque Link": "Верхний шлиц-шарнир",
    "Upper Slave Link": "Верхнее ведомое звено",
    "Lower Slave Link": "Нижнее ведомое звено",
    "Upper Diaphragm Tube Subassembly": "Сборка верхней диафрагменной трубы",
    "Upper Diaphragm Tube": "Верхняя диафрагменная труба",
    "Pivot Pin": "Штифт вращения",
    "Pivot Bracket": "Поворотный кронштейн",
    "Upper Pivot Bracket": "Верхний поворотный кронштейн",
    "Harness Support Bracket": "Кронштейн крепления жгута",
    "Lock Stay Cardan": "Кардан фиксатора",
    "Transfer Block": "Переходный блок",
    "Transfer block": "Переходный блок",
    "Spherical Bearing": "Сферический подшипник",
    "Pintle pin": "Штифт навеса стойки",
    "Pintle Pin": "Штифт навеса стойки",
    "Forward Pintle Pin": "Передний штифт навеса стойки",
    "Retaining Pin": "Стопорный штифт",
    "Uplock Pin ain Fitting": "Штифт замка убранного положения корпуса стойки",
    "Cylinder": "Цилиндр",
    "Bracket": "Кронштейн",
    "Pin": "Штифт",
    "Spacer": "Проставка",
    "Drag-arm Spacer": "Проставка тяги",
    "Slave Link": "Ведомое звено",
    "Valve Stem": "Шток клапана",
    "Torque Link": "Шлиц-шарнир",
    "Uplock Pin": "Штифт замка убранного положения",
    "Uplock": "Замок убранного положения",
    "Inflation Valve": "Клапан зарядки",
    "Lower Bearing Subassembly": "Сборка нижнего подшипника",
    "Lower Bearing": "Нижний подшипник",
    # Lowercase variants
    "Sliding tube": "Скользящая труба",
    "Pintle pin": "Штифт навеса стойки",
    # Repair procedure descriptions
    "Bush": "Втулка",
    "Bushes": "Втулки",
    "Machining": "Механическая обработка",
    "Installation": "Установка",
    "Machining and Installation": "Механическая обработка и установка",
    "Machining and installation": "Механическая обработка и установка",
    # ── Part 2+ component names ──
    "Installation of Labels": "Установка этикеток",
    "Label": "Этикетка",
    "Sliding Tube Subassembly": "Сборка скользящей трубы",
    "Installation of Bushes": "Установка втулок",
    "Application of Jointing Compound": "Нанесение герметизирующего состава",
    "Assembly of Lower Bearing Subassembly": "Сборка нижнего подшипника",
    "Seal Configuration": "Конфигурация уплотнений",
    "Crimping of the Pin": "Обжимка штифта",
    "Application of Ardrox AV100D to the Upper Diaphragm Tube": "Нанесение Ardrox AV100D на верхнюю диафрагменную трубу",
    "Application of Ardrox AV100D to the Pin": "Нанесение Ardrox AV100D на штифт",
    "Assembly of Damper": "Сборка демпфера",
    "Bolt Subassembly": "Сборка болта",
    "Bracket Subassembly": "Сборка кронштейна",
    "Bracket Assembly": "Сборка кронштейна",
    "Rod End Assembly": "Сборка наконечника тяги",
    "1M Electrical Axle Harness": "Электрический жгут оси 1М",
    "2M Electrical Axle Harness": "Электрический жгут оси 2М",
    "Slave Link Subassembly": "Сборка ведомого звена",
    "Lower Slave Link Subassembly": "Сборка нижнего ведомого звена",
    "Proximity Switch": "Датчик приближения",
    "Proximity Switches": "Датчики приближения",
    "Harness Support Bracket": "Кронштейн крепления жгута",
    "Upper Torque Link Subassembly": "Сборка верхнего шлиц-шарнира",
    "Lower Torque Link Subassembly": "Сборка нижнего шлиц-шарнира",
    "Shock Absorber Subassembly": "Сборка амортизатора",
    "Housing": "Корпус",
    "Damper": "Демпфер",
    "Cardan Assembly": "Сборка кардана",
    "Transfer Block Subassembly": "Сборка переходного блока",
    "Piston": "Поршень",
    "Locking Plate": "Стопорная пластина",
    "Stop Ring": "Стопорное кольцо",
    "Stop ring": "Стопорное кольцо",
    "Cross Bolt": "Поперечный болт",
    "Upper Diaphragm Tube Sub-assembly": "Сборка верхней диафрагменной трубы",
    "Upper diaphragm tube sub-assembly": "Сборка верхней диафрагменной трубы",
    "Locking Nut": "Контргайка",
    "Locking Washer": "Стопорная шайба",
    "Harness Support": "Крепление жгута",
    "Harness support": "Крепление жгута",
    "Electrical Bonding Resistance Tests": "Проверка сопротивления электрического соединения",
    "Proximity Switch connector shell": "Корпус разъёма датчика приближения",
    "Harness support bracket": "Кронштейн крепления жгута",
    "Upper pivot bracket": "Верхний поворотный кронштейн",
    "Static discharge connector": "Штыревой разъём статического разряда",
    "Lock stay cardan": "Кардан фиксатора",
    "Pivot bracket": "Поворотный кронштейн",
    "Hydraulic fluid": "Гидравлическая жидкость",
    "Hydraulic Fluid": "Гидравлическая жидкость",
    "Bung": "Заглушка",
    "Bracket Subassemblies": "Сборки кронштейнов",
    "Sliding Tube Subassembly": "Сборка скользящей трубы",
    "Lower Bearing Subassembly": "Сборка нижнего подшипника",
    "Upper Bearing Housing": "Корпус верхнего подшипника",
    "Liner": "Вкладыш",
    "Liner Installation": "Установка вкладыша",
    "Labels": "Этикетки",
    "wiring diagram plate": "табличка электрической схемы",
    "Wiring Diagram Plate": "Табличка электрической схемы",
    # ── Part 3+ component names ──
    "Wedge": "Клин",
    "wedge": "клин",
    "Transfer Dowel": "Переходный штифт",
    "Transfer dowel": "Переходный штифт",
    "Lock Stay Cardan Subassembly": "Сборка кардана фиксатора",
    "Lock stay cardan subassembly": "Сборка кардана фиксатора",
    # Lowercase variants for textbox labels
    "Upper bearing housing": "Корпус верхнего подшипника",
    "upper bearing housing": "корпус верхнего подшипника",
    "Gland housing": "Корпус сальника",
    "Gland Housing": "Корпус сальника",
    "gland housing": "корпус сальника",
    "Valve support": "Опора клапана",
    "Valve Support": "Опора клапана",
    "valve support": "опора клапана",
    "Lower Bearing Housing Subassembly": "Сборка корпуса нижнего подшипника",
    "Lower bearing housing subassembly": "Сборка корпуса нижнего подшипника",
    "Lower Bearing Housing": "Корпус нижнего подшипника",
    "Lower bearing housing": "Корпус нижнего подшипника",
    "Main fitting": "Корпус стойки",
    "main fitting": "корпус стойки",
    "Sliding tube": "Скользящая труба",
    "sliding tube": "скользящая труба",
    "Gland Housing Subassembly": "Сборка корпуса сальника",
    "Gland housing subassembly": "Сборка корпуса сальника",
    "Spherical bearing": "Сферический подшипник",
    "spherical bearing": "сферический подшипник",
    # ── Part 4 component name variants ──
    "Lower bearing subassembly": "Сборка нижнего подшипника",
    "Pivot pin": "Штифт вращения",
    "Uplock pin": "Штифт замка убранного положения",
    "Main Fitting Repairs": "Ремонты корпуса стойки",
    "Torque Link Repairs": "Ремонты шлиц-шарнира",
    "Sliding Tube Repairs": "Ремонты скользящей трубы",
    # ── Part 5 component name variants ──
    "Upper Diaphragm Tube Repairs": "Ремонты верхней диафрагменной трубы",
    "Cylinder Repairs": "Ремонты цилиндра",
    "Transfer Block Repairs": "Ремонты переходного блока",
    "Harness Support Bracket Repairs": "Ремонты кронштейна крепления жгута",
    "Upper Pivot Bracket Repairs": "Ремонты верхнего поворотного кронштейна",
    "Lower Bearing Subassembly Repairs": "Ремонты сборки нижнего подшипника",
    "Oversize bearing": "Ремонтный подшипник увеличенного размера",
    "Oversize Bearing": "Ремонтный подшипник увеличенного размера",
    "Repair liner": "Ремонтный вкладыш",
    "Repair Liner": "Ремонтный вкладыш",
}

# ── Fixed phrase translations ────────────────────────────────────────────────
FIXED = {
    # Title page
    "MAIN LANDING GEAR LEG": "СТОЙКА ОСНОВНОГО ШАССИ",
    "PART NUMBER": "НОМЕР ДЕТАЛИ",
    "COMPONENT MAINTENANCE MANUAL WITH": "РУКОВОДСТВО ПО ТЕХНИЧЕСКОМУ ОБСЛУЖИВАНИЮ КОМПОНЕНТОВ С",
    "ILLUSTRATED PARTS LIST": "ИЛЛЮСТРИРОВАННЫМ ПЕРЕЧНЕМ ДЕТАЛЕЙ",
    "STATEMENT OF INITIAL CERTIFICATION": "СВИДЕТЕЛЬСТВО О ПЕРВИЧНОЙ СЕРТИФИКАЦИИ",
    "INTENTIONALLY BLANK": "НАМЕРЕННО ОСТАВЛЕНО ПУСТЫМ",

    # Structural headings
    "NEW/REVISED PAGES": "НОВЫЕ/ПЕРЕСМОТРЕННЫЕ СТРАНИЦЫ",
    "REVISION RECORD": "ЗАПИСЬ ИЗМЕНЕНИЙ",
    "RECORD OF REVISIONS": "ЗАПИСЬ ИЗМЕНЕНИЙ",
    "RECORD OF TEMPORARY REVISIONS": "ЗАПИСЬ ВРЕМЕННЫХ ИЗМЕНЕНИЙ",
    "LIST OF SERVICE BULLETINS": "СПИСОК СЕРВИСНЫХ БЮЛЛЕТЕНЕЙ",
    "LIST OF SERVICE BULLETINS (Continued)": "СПИСОК СЕРВИСНЫХ БЮЛЛЕТЕНЕЙ (Продолжение)",
    "LIST OF EFFECTIVE PAGES": "ПЕРЕЧЕНЬ ДЕЙСТВУЮЩИХ СТРАНИЦ",
    "LIST OF EFFECTIVE PAGES (Continued)": "ПЕРЕЧЕНЬ ДЕЙСТВУЮЩИХ СТРАНИЦ (Продолжение)",
    "TABLE OF CONTENTS": "СОДЕРЖАНИЕ",
    "TABLE OF CONTENTS (Continued)": "СОДЕРЖАНИЕ (Продолжение)",
    "ILLUSTRATIONS": "ИЛЛЮСТРАЦИИ",
    "ILLUSTRATIONS (Continued)": "ИЛЛЮСТРАЦИИ (Продолжение)",
    "UNIT IDENTIFICATION CHART": "ТАБЛИЦА ИДЕНТИФИКАЦИИ ИЗДЕЛИЯ",
    "UNIT IDENTIFICATION CHART (Continued)": "ТАБЛИЦА ИДЕНТИФИКАЦИИ ИЗДЕЛИЯ (Продолжение)",
    "TITLE PAGE": "ТИТУЛЬНАЯ СТРАНИЦА",
    "Issued by": "Выпущено",
    "REV NO.": "НОМ. РЕД.",
    "ISSUE DATE": "ДАТА ВЫПУСКА",
    "INSERTION DATE": "ДАТА ВСТАВКИ",
    "BY": "КЕМ",
    "REMOVAL DATE": "ДАТА УДАЛЕНИЯ",

    # TOC section headings
    "Assembly (Including Storage)": "Сборка (включая хранение)",
    "General": "Общие сведения",
    "Procedure": "Процедура",
    "Storage": "Хранение",
    "Fits and Clearances": "Посадки и зазоры",
    "Fits and Clearances Definitions": "Определения посадок и зазоров",
    "Remarks": "Примечания",
    "Torque Data": "Данные моментов затяжки",
    "Special Tools, Fixtures and Equipment": "Специальные инструменты, приспособления и оборудование",
    "Illustrated Parts List": "Иллюстрированный перечень деталей",
    "Introduction": "Введение",
    "How To Use This Illustrated Parts List": "Как пользоваться иллюстрированным перечнем деталей",
    "Vendor Codes, Names and Addresses": "Коды, наименования и адреса поставщиков",
    "Numerical Index": "Числовой указатель",
    "Detailed Parts List": "Подробный список деталей",

    # Figure captions
    "Diagram of Operation": "Схема работы",
    "Electrical Bonding Resistance Test Points": "Точки проверки сопротивления электрического соединения",
    "Protective Treatment": "Защитная обработка",

    # Section/TOC headings (textbox labels)
    "Title Page": "Титульная страница",
    "Record of Temporary": "Запись временных",
    "Revisions": "Изменений",
    "Unit Identification": "Идентификация изделия",
    "Chart": "Таблица",
    "List of Effective": "Перечень действующих",
    "Pages": "Страниц",
    "Pages (Continued)": "Страниц (Продолжение)",
    "Table of Contents": "Содержание",
    "Testing and Fault": "Проверка и поиск неисправностей",
    "Isolation": "Локализация",
    "Subject": "Раздел",
    "Page": "Страница",
    "Date": "Дата",
    "Operation": "Работа",
    "INCORPORATED": "ВКЛЮЧЕНО",
    "Description and": "Описание и",
    "Disassembly": "Разборка",
    "Cleaning": "Очистка",
    "Repair": "Ремонт",
    "Repair Procedure Conditions": "Условия выполнения процедуры ремонта",
    "Approved Repairs": "Утверждённые ремонты",
    "Check": "Проверка",
    "Equipment and Materials": "Оборудование и материалы",
    "Test Conditions": "Условия испытаний",
    "Detailed Inspection": "Детальный осмотр",
    "Special Detailed Inspection": "Специальный детальный осмотр",
    "Fault Isolation": "Поиск неисправностей",
    "Description and Operation": "Описание и работа",
    "Description": "Описание",
    "Data": "Данные",
    "Superseded": "Заменён",
    "Assembly": "Сборка",
    "Assembly (Including": "Сборка (включая",
    "Storage)": "хранение)",
    "Storage) (Continued)": "хранение) (Продолжение)",
    "Tables 101": "Таблицы 101",
    "Tables": "Таблицы",
    "Only": "Только",
    "Conditions": "Условия",
    "Inspection": "Осмотр",
    "Materials": "Материалы",
    "Special Tools, Fixtures": "Специальные инструменты, приспособления",
    "and Equipment": "и оборудование",
    "and Equipment (Continued)": "и оборудование (Продолжение)",
    "Repair (Continued)": "Ремонт (Продолжение)",
    "Cleaning (Continued)": "Очистка (Продолжение)",
    "Check (Continued)": "Проверка (Продолжение)",
    "Disassembly (Continued)": "Разборка (Продолжение)",
    "Assembly (Continued)": "Сборка (Продолжение)",
    "Illustrations": "Иллюстрации",
    "(Continued)": "(Продолжение)",
    "Isolation (Continued)": "Локализация (Продолжение)",
    "REV": "РЕД",

    # Part 2+ section headings
    "INTRODUCTION": "ВВЕДЕНИЕ",
    "DESCRIPTION AND OPERATION": "ОПИСАНИЕ И РАБОТА",
    "TESTING AND FAULT ISOLATION": "ПРОВЕРКА И ПОИСК НЕИСПРАВНОСТЕЙ",
    "DISASSEMBLY": "РАЗБОРКА",
    "CLEANING": "ОЧИСТКА",
    "CHECK": "ПРОВЕРКА",
    "REPAIR": "РЕМОНТ",
    "ASSEMBLY": "СБОРКА",
    "Reference Publications": "Справочные публикации",
    "Equipment": "Оборудование",
    "Not applicable.": "Не применяется.",
    "Not applicable": "Не применяется",
    "Detailed Inspection.": "Детальный осмотр.",
    "Special Detailed Inspection.": "Специальный детальный осмотр.",
    "Leakage Tests": "Испытания на герметичность",
    "Initial Operations": "Начальные операции",
    "Compression": "Сжатие",
    "Recoil": "Обратный ход",
    "Wire Thread Inserts": "Резьбовые спиральные вставки",
    "Special Tools": "Специальные инструменты",
    "Paint Removal": "Удаление лакокрасочного покрытия",
    "Prepare for Transport and Storage": "Подготовка к транспортировке и хранению",
    "Examine Parts Visually": "Визуальный осмотр деталей",
    "Examine Dimensions": "Проверка размеров",
    "Spring Data": "Данные пружины",
    "Corrosion.": "Коррозия.",
    "Distortion and/or cracks.": "Деформация и/или трещины.",
    "Wear or fretting.": "Износ или фреттинг.",
    "Scores, dents or burrs.": "Задиры, вмятины или заусенцы.",
    "Unserviceable screw threads.": "Непригодная резьба.",
    "Deterioration of protective treatment.": "Ухудшение защитного покрытия.",
    "Parts of permanent assemblies that are not correctly attached.": "Неправильно прикреплённые детали неразъёмных сборок.",
    "Procedure": "Процедура",
    "INCORPORATED": "ВКЛЮЧЕНО",
    "PRE SB 201-32-72:": "До SB 201-32-72:",
    "PRE SB 201-32-58:": "До SB 201-32-58:",
    "POST SB 201-32-72:": "После SB 201-32-72:",
    "POST SB 201-32-58:": "После SB 201-32-58:",
    "Component Maintenance Manual": "Руководство по техническому обслуживанию компонентов",
    "Diagram of Operation Figure 2": "Схема работы Рисунок 2",
    "General": "Общие сведения",
    "Materials": "Материалы",
    "These materials are necessary:": "Необходимы следующие материалы:",
    "These special tools are necessary:": "Необходимы следующие специальные инструменты:",
    "NOTE: Alternative equivalents are permitted.": "ПРИМЕЧАНИЕ: Допускаются альтернативные эквиваленты.",
    "Unless instructions are different:": "Если инструкции не указывают иное:",
    "Cleaning": "Очистка",
    # ─── Description and Operation sentences ───
    "The main landing gear leg is a two stage, telescopic shock absorber.":
        "Стойка основного шасси представляет собой двухступенчатый телескопический амортизатор.",
    "To be given subsequently.": "Будет предоставлено дополнительно.",
    "Make sure that the work area, the tools and the equipment are clean.":
        "Убедитесь, что рабочая зона, инструменты и оборудование чистые.",
    "Special Dimension Check:": "Специальная проверка размеров:",
    "Piston (17-200) Leakage Tests": "Испытания на герметичность поршня (17-200)",
    "Proximity Switches (7-40 and 7-230) Adjustments and Tests":
        "Регулировка и проверка датчиков приближения (7-40 и 7-230)",
    "Electrical Bonding Resistance Tests (Refer to Figure 102)":
        "Проверка сопротивления электрического соединения (обратитесь к рисунку 102)",
    "Procedure to Fill and Pressurize the Main Landing Gear Leg (1-1)":
        "Процедура заполнения и создания давления в стойке основного шасси (1-1)",
    "Leakage Tests": "Испытания на герметичность",
    "Only,": "Только,",
    "M-DLNDT3": "M-DLNDT3",
    "Parts 1": "Части 1",
    "Cond H1025": "Конд. H1025",
    "Type 1": "Тип 1",
    "Weight with hydraulic fluid": "Масса с гидравлической жидкостью",
    "Weight without hydraulic fluid": "Масса без гидравлической жидкости",
    "They are a set: keep them together.": "Они являются комплектом: храните их вместе.",
    "threaded insert": "резьбовая вставка",
    "bracket assembly": "сборка кронштейна",
    "Fits and Clearances - Key Diagram": "Посадки и зазоры — Ключевая схема",
    "Proximity Switch (7-230) — Target (7-180) — Adjustment": "Датчик приближения (7-230) — Мишень (7-180) — Регулировка",
    "Proximity Switch (7-40) — Target (6-130) — Adjustment": "Датчик приближения (7-40) — Мишень (6-130) — Регулировка",
    # ─── Part 3 (Repair) section headings and sentences ───
    "Repair Levels": "Уровни ремонта",
    "Surface Damage": "Повреждение поверхности",
    "Identification": "Идентификация",
    "Protective Treatment Repair": "Ремонт защитного покрытия",
    "Cadmium Plated Surfaces": "Кадмированные поверхности",
    "Anodized Surfaces": "Анодированные поверхности",
    "Paint Finish": "Лакокрасочное покрытие",
    "Protective Treatment Replacement": "Замена защитного покрытия",
    "Protective Treatment Table 601": "Таблица защитной обработки 601",
    "Repair of surface damage.": "Ремонт повреждений поверхности.",
    "On a radius.": "На радиусе.",
    "Remove burrs from external screw threads.": "Удалите заусенцы с наружных резьб.",
    "Concession": "Допуск на отклонение",
    "Salvage": "Восстановление",
    "Clean the parts after repair: refer to CLEANING.": "Очистите детали после ремонта: обратитесь к ОЧИСТКЕ.",
    "Figure Deleted Figure 609": "Рисунок удалён Рисунок 609",
    "Figure Deleted Figure 611": "Рисунок удалён Рисунок 611",
    "the 1,5 mm (0.060 in) hole for the Bowden cable(1-45).":
        "отверстие 1,5 мм (0,060 дюйм) для троса Боудена (1-45).",
    "Refer to Figure 623. Apply cadmium":
        "Обратитесь к рисунку 623. Нанесите кадмиевое",
    "Refer to Figure 636. Apply cadmium":
        "Обратитесь к рисунку 636. Нанесите кадмиевое",
    "the bearings, bush bores and flanges": "подшипники, отверстия под втулки и фланцы",
    "Refer to para 3.C. and Figure 618.": "Обратитесь к п. 3.C. и рисунку 618.",
    "to the areas where Sermetel W is applied": "на участки, где нанесён Sermetel W",
    "to the chromium plated areas": "на хромированные участки",
    "to the threaded surfaces.": "на резьбовые поверхности.",
    "to the threaded surfaces": "на резьбовые поверхности",
    # ─── Protective Treatment Table 601 cell paragraphs ───
    "20-410B and 20-420B only": "только 20-410B и 20-420B",
    "20-410D and 20-420D only": "только 20-410D и 20-420D",
    "plate all over including the holes less than diameter 10 mm (0.393 in) but not to the chromium plated areas and areas A: refer to PCS-2100. Make the cadmium plate thickness between 0,010 and 0,020 mm":
        "покрытие повсюду, включая отверстия диаметром менее 10 мм (0,393 дюйм), но не на хромированные участки и участки A: обратитесь к PCS-2100. Толщина кадмиевого покрытия должна быть от 0,010 до 0,020 мм",
    "(0.0004 and 0.0008 in). The cadmium plate must overlap the chromium plate run out. The cadmium plate is optional on the lubrication fitting bores where the lubrication adaptors (20-130),":
        "(0,0004 и 0,0008 дюйм). Кадмиевое покрытие должно перекрывать выход хромового покрытия. Кадмиевое покрытие является необязательным на отверстиях смазочных ниппелей, где устанавливаются смазочные адаптеры (20-130),",
    "(20-160), (20-190) and (20-220) will":
        "(20-160), (20-190) и (20-220) будут",
    "Apply primer paint only to areas B: refer to PCS-2500.":
        "Нанесите грунтовочную краску только на участки B: обратитесь к PCS-2500.",
    "Apply wet primer to PCS-2804 or apply resin to PCS-2802 to the area D.":
        "Нанесите влажную грунтовку по PCS-2804 или нанесите смолу по PCS-2802 на участок D.",
    "Apply paint all over but not on the chromium plated areas, the areas A, B, C and on the lubrication fitting bores where the lubrication adaptors (20-130), (20-160),":
        "Нанесите краску повсюду, кроме хромированных участков, участков A, B, C и отверстий смазочных ниппелей, где устанавливаются смазочные адаптеры (20-130), (20-160),",
    "(20-190) and (20-220) will install: refer to PCS-2500. Paint finish is optional in areas E.":
        "(20-190) и (20-220) будут установлены: обратитесь к PCS-2500. Лакокрасочное покрытие является необязательным на участках E.",
    "(20-190) and (20-220) will install: refer to PCS-2500.":
        "(20-190) и (20-220) будут установлены: обратитесь к PCS-2500.",
    # ─── Drawing annotation paragraphs ───
    "SERMETEL W ON INTERNAL": "SERMETEL W НА ВНУТРЕННЕМ",
    "SERMETEL W TO": "SERMETEL W ДО",
    "14 PLACES\tW": "14 МЕСТ\tW",
    # ─── Part 4 drawing annotation paragraphs ───
    "PRIMER PAINT FACE C": "ГРУНТОВОЧНАЯ КРАСКА ПОВЕРХНОСТЬ C",
    "C AFTER THREAD": "C ПОСЛЕ РЕЗЬБЫ",
    "CADMIUM PLATE AND PAINT TO OVERLAP ON CHROMIUM RADIUS": "КАДМИЕВОЕ ПОКРЫТИЕ И КРАСКА ПЕРЕКРЫВАЮТ ХРОМОВЫЙ РАДИУС",
    "CADMIUM PLATE AND PAINT TO OVERLAP": "КАДМИЕВОЕ ПОКРЫТИЕ И КРАСКА ПЕРЕКРЫВАЮТ",
    "AND PRIMER PAINT": "И ГРУНТОВОЧНАЯ КРАСКА",
    "KNUCKLE BORES": "ОТВЕРСТИЯ КУЛАКА",
    "KNUCKLE TOOLING LUG": "ТЕХНОЛОГИЧЕСКАЯ ПРОУШИНА КУЛАКА",
    "GREASE HOLES": "СМАЗОЧНЫЕ ОТВЕРСТИЯ",
    "EXTERNAL SERMETEL LIMIT": "НАРУЖНЫЙ ПРЕДЕЛ SERMETEL",
    "AXLE NUT CROSS BOLT HOLES": "ОТВЕРСТИЯ ПОПЕРЕЧНОГО БОЛТА ГАЙКИ ОСИ",
    "BRAKE FLANGE": "ТОРМОЗНОЙ ФЛАНЕЦ",
    "RETRACTION BORES": "ОТВЕРСТИЯ УБОРКИ",
    "PINTLE CROSS BORES": "ПОПЕРЕЧНЫЕ ОТВЕРСТИЯ ШТИФТА НАВЕСА",
    "NO CADMIUM PLATE OR PAINT BEYOND THIS LINE": "КАДМИЕВОЕ ПОКРЫТИЕ ИЛИ КРАСКА ЗА ЭТОЙ ЛИНИЕЙ НЕ ДОПУСКАЮТСЯ",
    "UPPER DIAPHRAGM TUBE CROSS BORE": "ПОПЕРЕЧНОЕ ОТВЕРСТИЕ ВЕРХНЕЙ ДИАФРАГМЕННОЙ ТРУБЫ",
    "LOWER CARDAN BORE": "ОТВЕРСТИЕ НИЖНЕГО КАРДАНА",
    "EXTENT OF FINE LIMIT DIA.": "ПРОТЯЖЁННОСТЬ ЧИСТОВОГО ПРЕДЕЛЬНОГО ДИАМ.",
    "DIAMETERS THRU BORES INCLUDING CHAMFERS": "ДИАМЕТРЫ СКВОЗНЫХ ОТВЕРСТИЙ ВКЛЮЧАЯ ФАСКИ",
    "DIAMETER THRU BORE INCLUDING CHAMFERS": "ДИАМЕТР СКВОЗНОГО ОТВЕРСТИЯ ВКЛЮЧАЯ ФАСКИ",
    "PINTLE BORES": "ОТВЕРСТИЯ ШТИФТА НАВЕСА",
    "BARREL": "КОРПУС ЦИЛИНДРА",
    "(BORE AND CHAMFER INCLUDED)": "(ОТВЕРСТИЕ И ФАСКА ВКЛЮЧЕНЫ)",
    "B INCLUDING CHAMFERS": "B ВКЛЮЧАЯ ФАСКИ",
    "B AND E CHAMFERS ONLY TYPICAL 2 PLACES": "B И E ТОЛЬКО ФАСКИ ТИПИЧНО 2 МЕСТА",
    "INCLUDING CHAMFER": "ВКЛЮЧАЯ ФАСКУ",
    "INCLUDING RADIUS": "ВКЛЮЧАЯ РАДИУС",
    "SPOTFACE (REARSIDE ONLY)": "ЗЕНКОВКА (ТОЛЬКО ОБРАТНАЯ СТОРОНА)",
    # ─── Part 4 chrome plating termination annotations ───
    "FULL CHROME\t\t PLATING THICKNESS": "ПОЛНАЯ ТОЛЩИНА\t\t ХРОМОВОГО ПОКРЫТИЯ",
    "FULL CHROME PLATING THICKNESS": "ПОЛНАЯ ТОЛЩИНА ХРОМОВОГО ПОКРЫТИЯ",
    "CHROME PLATING": "ХРОМОВОЕ ПОКРЫТИЕ",
    "CHROME PLATING DEPOSIT": "ХРОМОВОЕ ПОКРЫТИЕ",
    "PLATING THICKNESS": "ТОЛЩИНА ПОКРЫТИЯ",
    "PLATING LIMIT": "ПРЕДЕЛ ПОКРЫТИЯ",
    "PAINT DEPOSIT OVERLAP": "ПЕРЕКРЫТИЕ СЛОЯ КРАСКИ",
    "ZINC-NICKEL DEPOSIT OVERLAP": "ПЕРЕКРЫТИЕ СЛОЯ ЦИНК-НИКЕЛЯ",
    "DIA. AFTER GRINDING CHROME": "ДИАМ. ПОСЛЕ ШЛИФОВАНИЯ ХРОМА",
    "DIA. AFTER CHROME PLATING": "ДИАМ. ПОСЛЕ ХРОМИРОВАНИЯ",
    "A SMOOTH TRANSITION": "ПЛАВНЫЙ ПЕРЕХОД",
    "WAVY OR IRREGULAR LINE PERMISSIBLE": "ВОЛНИСТАЯ ИЛИ НЕРОВНАЯ ЛИНИЯ ДОПУСТИМА",
    "EDGE BLENDED WITH A SMOOTH TRANSITION": "КРОМКА СОПРЯЖЕНА С ПЛАВНЫМ ПЕРЕХОДОМ",
    "NO ZINC-NICKEL OR PAINT DEPOSIT TO REMAIN ON OR PROUD OF WORKING DIA. AFTER GRINDING CHROME":
        "ЦИНК-НИКЕЛЬ ИЛИ КРАСКА НЕ ДОЛЖНЫ ОСТАВАТЬСЯ НА РАБОЧЕМ ДИАМ. ИЛИ ВЫСТУПАТЬ НАД НИМ ПОСЛЕ ШЛИФОВАНИЯ ХРОМА",
    "NO ZINC-NICKEL OR PAINT DEPOSIT TO REMAIN": "ЦИНК-НИКЕЛЬ ИЛИ КРАСКА НЕ ДОЛЖНЫ ОСТАВАТЬСЯ",
    "ON OR PROUD OF WORKING DIA. AFTER GRINDING CHROME":
        "НА РАБОЧЕМ ДИАМ. ИЛИ ВЫСТУПАТЬ НАД НИМ ПОСЛЕ ШЛИФОВАНИЯ ХРОМА",
    "NO ZINC-NICKEL OR PAINT DEPOSIT": "ЦИНК-НИКЕЛЬ ИЛИ КРАСКА НЕ ДОПУСКАЮТСЯ",
    "TO REMAIN ON WORKING DIA. AFTER CHROME PLATING":
        "ОСТАВАТЬСЯ НА РАБОЧЕМ ДИАМ. ПОСЛЕ ХРОМИРОВАНИЯ",
    "EXTERNAL THICK ZINC-NICKEL": "НАРУЖНЫЙ ТОЛСТЫЙ СЛОЙ ЦИНК-НИКЕЛЯ",
    "INTERNAL THICK ZINC-NICKEL": "ВНУТРЕННИЙ ТОЛСТЫЙ СЛОЙ ЦИНК-НИКЕЛЯ",
    "BARREL OUTER DIA. LOWER CHROME TERMINATION": "НАРУЖНЫЙ ДИАМ. КОРПУСА ЦИЛИНДРА НИЖНЕЕ ОКОНЧАНИЕ ХРОМА",
    "BARREL OUTER DIA. UPPER CHROME TERMINATION": "НАРУЖНЫЙ ДИАМ. КОРПУСА ЦИЛИНДРА ВЕРХНЕЕ ОКОНЧАНИЕ ХРОМА",
    "JOURNAL A OUTER CHROME TERMINATION (TYPICAL)": "ШЕЙКА A НАРУЖНОЕ ОКОНЧАНИЕ ХРОМА (ТИПИЧНО)",
    "JOURNAL A,B,C CHROME TERMINATION (TYPICAL)": "ШЕЙКА A,B,C ОКОНЧАНИЕ ХРОМА (ТИПИЧНО)",
    "JOURNAL C INNER CHROME TERMINATION (TYPICAL)": "ШЕЙКА C ВНУТРЕННЕЕ ОКОНЧАНИЕ ХРОМА (ТИПИЧНО)",
    "BREAK FLANGE FACE CHROME TERMINATION (TYPICAL)": "ТОРЕЦ ТОРМОЗНОГО ФЛАНЦА ОКОНЧАНИЕ ХРОМА (ТИПИЧНО)",
    "HPC SEAL ABUTMENT LOWER CHROME TERMINATION": "НИЖНЕЕ ОКОНЧАНИЕ ХРОМА УПОРА УПЛОТНЕНИЯ HPC",
    "HPC SEAL ABUTMENT UPPER CHROME TERMINATION": "ВЕРХНЕЕ ОКОНЧАНИЕ ХРОМА УПОРА УПЛОТНЕНИЯ HPC",
    "CHROME PLATING WILL TERMINATE ANYWHERE ON THE CHAMFER": "ХРОМОВОЕ ПОКРЫТИЕ ЗАКАНЧИВАЕТСЯ В ЛЮБОМ МЕСТЕ НА ФАСКЕ",
    "CHROMIUM PLATED SURFACE": "ХРОМИРОВАННАЯ ПОВЕРХНОСТЬ",
    # ─── Part 5 section headings and sentences ───
    "Specified Damage and Material Specification.": "Указанное повреждение и спецификация материала.",
    "Specified Damage and Material Specification": "Указанное повреждение и спецификация материала",
    "Specified Damage": "Указанное повреждение",
    "Material Specification": "Спецификация материала",
    "Repair Parts": "Ремонтные детали",
    "These repair parts are necessary:": "Необходимы следующие ремонтные детали:",
    "Repair parts are not necessary.": "Ремонтные детали не требуются.",
    "Materials are not necessary.": "Материалы не требуются.",
    "Special tools are not necessary.": "Специальные инструменты не требуются.",
    "Repair loose but undamaged liner:": "Ремонт незакреплённого, но неповреждённого вкладыша:",
    "Repair damaged liner:": "Ремонт повреждённого вкладыша:",
    "Loose or damaged liner.": "Незакреплённый или повреждённый вкладыш.",
    "Damaged or loose liner.": "Повреждённый или незакреплённый вкладыш.",
    "Damage or wear to the diameters A and B.": "Повреждение или износ диаметров A и B.",
    "Damage or wear to diameter A.": "Повреждение или износ диаметра A.",
    "Damage or wear to the diameter A and/or the adjacent inside face.": "Повреждение или износ диаметра A и/или прилегающей внутренней поверхности.",
    "Damage or wear to the diameter A.": "Повреждение или износ диаметра A.",
    "Damage or corrosion to diameter A.": "Повреждение или коррозия диаметра A.",
    "Procedure (Refer to Figure 601)": "Процедура (обратитесь к рисунку 601)",
    "LARGER VIEW AT A": "УВЕЛИЧЕННЫЙ ВИД A",
    "MACHINING": "МЕХАНИЧЕСКАЯ ОБРАБОТКА",
    "LARGER VIEW AT": "УВЕЛИЧЕННЫЙ ВИД",
    "SECTION Y-Y": "СЕЧЕНИЕ Y-Y",
    "BEFORE CHROMIUM PLATE": "ДО ХРОМОВОГО ПОКРЫТИЯ",
    "AFTER CHROMIUM PLATE": "ПОСЛЕ ХРОМОВОГО ПОКРЫТИЯ",
    "AFTER CHROMIUM PLATE AND GRINDING": "ПОСЛЕ ХРОМОВОГО ПОКРЫТИЯ И ШЛИФОВАНИЯ",
    "CHROMIUM PLATE CAN STOP ANYWHERE ON THE CHAMFER.": "ХРОМОВОЕ ПОКРЫТИЕ МОЖЕТ ЗАКАНЧИВАТЬСЯ В ЛЮБОМ МЕСТЕ НА ФАСКЕ.",
    "IT CAN BE FINISHED BY GRINDING": "ОНО МОЖЕТ БЫТЬ ЗАВЕРШЕНО ШЛИФОВАНИЕМ",
    "CHROMIUM PLATE TEMINATION TO M-DLPS1031-6": "ОКОНЧАНИЕ ХРОМОВОГО ПОКРЫТИЯ ПО M-DLPS1031-6",
    "CHROMIUM PLATE TEMINATION TO M-DLPS1031-3": "ОКОНЧАНИЕ ХРОМОВОГО ПОКРЫТИЯ ПО M-DLPS1031-3",
    "CHROMIUM PLATE TEMINATION TO M-DLPS1031-1": "ОКОНЧАНИЕ ХРОМОВОГО ПОКРЫТИЯ ПО M-DLPS1031-1",
    "CHROMIUM PLATE TERMINATION TO M-DLPS1031-6": "ОКОНЧАНИЕ ХРОМОВОГО ПОКРЫТИЯ ПО M-DLPS1031-6",
    "CHROMIUM PLATE TERMINATION TO M-DLPS1031-3": "ОКОНЧАНИЕ ХРОМОВОГО ПОКРЫТИЯ ПО M-DLPS1031-3",
    "CHROMIUM PLATE TERMINATION TO M-DLPS1031-1": "ОКОНЧАНИЕ ХРОМОВОГО ПОКРЫТИЯ ПО M-DLPS1031-1",
    "DETAIL X": "ДЕТАЛЬ X",
    "DETAIL Y": "ДЕТАЛЬ Y",
    "DETAIL Z": "ДЕТАЛЬ Z",
    "AFTER CHROMIUM PLATE IS APPLIED AND BEFORE GRINDING": "ПОСЛЕ НАНЕСЕНИЯ ХРОМОВОГО ПОКРЫТИЯ И ДО ШЛИФОВАНИЯ",
    "BEFORE CHROMIUM PLATE IS APPLIED": "ДО НАНЕСЕНИЯ ХРОМОВОГО ПОКРЫТИЯ",
    "NOT CHROMIUM PLATED": "БЕЗ ХРОМОВОГО ПОКРЫТИЯ",
    "CHROMIUM PLATED LENGTH": "ДЛИНА ХРОМИРОВАННОГО УЧАСТКА",
    "THE CHROMIUM PLATE MUST STOP IN THIS LENGTH. AN IRREGULAR LINE IS PERMITTED": "ХРОМОВОЕ ПОКРЫТИЕ ДОЛЖНО ЗАКАНЧИВАТЬСЯ НА ЭТОЙ ДЛИНЕ. НЕРОВНАЯ ЛИНИЯ ДОПУСКАЕТСЯ",
    "THE CHROMIUM PLATE MUST STOP IN THIS LENGTH. AN IRREGULAR LINE IS PERMITTED.": "ХРОМОВОЕ ПОКРЫТИЕ ДОЛЖНО ЗАКАНЧИВАТЬСЯ НА ЭТОЙ ДЛИНЕ. НЕРОВНАЯ ЛИНИЯ ДОПУСКАЕТСЯ.",
    "LENGTH OF CHROMIUM PLATE": "ДЛИНА ХРОМОВОГО ПОКРЫТИЯ",
    "EDGES SMOOTHED OUT. CHROMIUM PLATE TERMINATION TO M-DLPS1031-6": "КРОМКИ СГЛАЖЕНЫ. ОКОНЧАНИЕ ХРОМОВОГО ПОКРЫТИЯ ПО M-DLPS1031-6",
    "CHROMIUM PLATE TERMINATION TO M-DLPS1031-7 or M-DLPS1031-8": "ОКОНЧАНИЕ ХРОМОВОГО ПОКРЫТИЯ ПО M-DLPS1031-7 или M-DLPS1031-8",
    "EDGES SMOOTHED OUT. CHROMIUM PLATE TERMINATION": "КРОМКИ СГЛАЖЕНЫ. ОКОНЧАНИЕ ХРОМОВОГО ПОКРЫТИЯ",
    "DIAMETER A BEFORE CHROMIUM PLATE": "ДИАМЕТР A ДО ХРОМОВОГО ПОКРЫТИЯ",
    "DIAMETER A AFTER GRINDING": "ДИАМЕТР A ПОСЛЕ ШЛИФОВАНИЯ",
    "DIAMETER BEFORE CHROMIUM PLATE": "ДИАМЕТР ДО ХРОМОВОГО ПОКРЫТИЯ",
    "DIAMETER A AFTER CHROMIUM PLATE AND GRINDING": "ДИАМЕТР A ПОСЛЕ ХРОМОВОГО ПОКРЫТИЯ И ШЛИФОВАНИЯ",
    "PART SECTION Z-Z": "ДЕТАЛЬ СЕЧЕНИЕ Z-Z",
    "SECTION Z-Z": "СЕЧЕНИЕ Z-Z",
    "SECTION Z-Z WITHOUT BEARING": "СЕЧЕНИЕ Z-Z БЕЗ ПОДШИПНИКА",
    "SECTION Z-Z WITH BEARING": "СЕЧЕНИЕ Z-Z С ПОДШИПНИКОМ",
    "DIM. C": "РАЗМЕР C",
    "DIAMETER A": "ДИАМЕТР A",
    "LARGER DETAIL Z": "УВЕЛИЧЕННАЯ ДЕТАЛЬ Z",
    "NOT TO SCALE": "НЕ В МАСШТАБЕ",
    "PIN": "ШТИФТ",
    "SHOT PEEN": "ДРОБЕСТРУЙНАЯ ОБРАБОТКА",
    "SMOOTH BLEND M-DLPS1031-6": "ПЛАВНЫЙ ПЕРЕХОД M-DLPS1031-6",
    "OVERSIZE BEARING - MACHINING": "РЕМОНТНЫЙ ПОДШИПНИК — МЕХАНИЧЕСКАЯ ОБРАБОТКА",
    "LINER DIMENSIONS": "РАЗМЕРЫ ВКЛАДЫША",
    "MAKE EDGES SMOOTH": "СГЛАДИТЬ КРОМКИ",
    "CENTERS TYPICAL": "ЦЕНТРЫ ТИПИЧНО",
    "RADIUS 4 PLACES": "РАДИУС 4 МЕСТА",
    "TYPICAL BOTH REPAIR SLEEVES": "ТИПИЧНО ОБА РЕМОНТНЫХ РУКАВА",
    "MACHINING REPAIR SLEEVE": "МЕХАНИЧЕСКАЯ ОБРАБОТКА РЕМОНТНОГО РУКАВА",
    "MACHINING THE REPAIR BUSH": "МЕХАНИЧЕСКАЯ ОБРАБОТКА РЕМОНТНОЙ ВТУЛКИ",
    "NICKEL PLATE MUST NOT ENTER LUBRICATION HOLES OR CROSS HOLE": "НИКЕЛЕВОЕ ПОКРЫТИЕ НЕ ДОЛЖНО ПОПАДАТЬ В СМАЗОЧНЫЕ ОТВЕРСТИЯ ИЛИ ПОПЕРЕЧНОЕ ОТВЕРСТИЕ",
    "CHROMIUM": "ХРОМ",
    "Remove all of the tape and clean the parts as necessary.": "Удалите всю ленту и очистите детали по необходимости.",
    "Measure the new diameter A.": "Измерьте новый диаметр A.",
    "DEGREES": "ГРАДУСОВ",
    # ─── Part 5 — Examine/inclusion/CAUTION full sentences ───
    "CAUTION: FOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO MESSIER-DOWTY LIMITED: REFER TO GUIDE-CS-001.":
        "ВНИМАНИЕ: ПРИ ПОВРЕЖДЕНИЯХ, ПРЕВЫШАЮЩИХ ДОПУСКИ ДАННОЙ СХЕМЫ РЕМОНТА, ОБРАТИТЕСЬ В MESSIER-DOWTY LIMITED: СМ. GUIDE-CS-001.",
    "CAUTION: FOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO MESSIER-DOWTY LTD: REFER TO GUIDE-CS-001.":
        "ВНИМАНИЕ: ПРИ ПОВРЕЖДЕНИЯХ, ПРЕВЫШАЮЩИХ ДОПУСКИ ДАННОЙ СХЕМЫ РЕМОНТА, ОБРАТИТЕСЬ В MESSIER-DOWTY LTD: СМ. GUIDE-CS-001.",
    "Examine the gland housing for flaws: refer to M-DLNDT8.":
        "Проверьте корпус сальника на наличие дефектов: обратитесь к M-DLNDT8.",
    "Examine the pin for flaws: refer to PCS-3600 and PCS-3100 inclusion class 3.":
        "Проверьте штифт на наличие дефектов: обратитесь к PCS-3600 и PCS-3100, класс включений 3.",
    "Examine the pin for flaws: refer to PCS-3600 and PCS-3100, inclusion class 3.":
        "Проверьте штифт на наличие дефектов: обратитесь к PCS-3600 и PCS-3100, класс включений 3.",
    "Examine the pin for flaws: refer to M-DLNDT3.":
        "Проверьте штифт на наличие дефектов: обратитесь к M-DLNDT3.",
    "Examine the uplock pin for flaws: refer to PCS-3600 and PCS-3100, inclusion class 3.":
        "Проверьте штифт замка убранного положения на наличие дефектов: обратитесь к PCS-3600 и PCS-3100, класс включений 3.",
    "Examine the chromium plated surface for flaws: refer to M-DLNDT3.":
        "Проверьте хромированную поверхность на наличие дефектов: обратитесь к M-DLNDT3.",
    "Examine the pin for flaws: refer to PCS-3600 and PCS-3100, inclusion class 2.":
        "Проверьте штифт на наличие дефектов: обратитесь к PCS-3600 и PCS-3100, класс включений 2.",
    "Examine the pin for flaws: refer to PCS-3600 and PCS-3100, inclusion class 4.":
        "Проверьте штифт на наличие дефектов: обратитесь к PCS-3600 и PCS-3100, класс включений 4.",
    "Examine the ground chromium plate for flaws: refer to PCS-3100, inclusion class 4 and PCS-3002.":
        "Проверьте шлифованное хромовое покрытие на наличие дефектов: обратитесь к PCS-3100, класс включений 4 и PCS-3002.",
    # ─── Part 5 — adhesive/liner procedure sentences ───
    "Apply adhesive PVC tape, Material Ref. Item TBA, around the gland housing to the sides of and touching the repair liner. Make sure that the edges of the adhesive PVC tape, Material Ref. Item TBA, bond tightly to the gland housing.":
        "Наклейте клейкую ПВХ-ленту, мат. ссылка позиция TBA, вокруг корпуса сальника по сторонам и с касанием ремонтного вкладыша. Убедитесь, что края клейкой ПВХ-ленты, мат. ссылка позиция TBA, плотно приклеены к корпусу сальника.",
    "Use a brush to apply a smooth layer of the prepared surface treatment mixture to the contact surface of the gland housing.":
        "Кистью нанесите ровный слой подготовленной смеси для обработки поверхности на контактную поверхность корпуса сальника.",
    "Put the gland housing in the preheated oven for a minimum of 4 minutes and until the applied surface treatment mixture is dry.":
        "Поместите корпус сальника в предварительно нагретую печь минимум на 4 минуты и до высыхания нанесённой смеси для обработки поверхности.",
    "Use a brush to apply Araldite, 2015, Material Ref. Item TBA, to the gland housing.":
        "Кистью нанесите Araldite 2015, мат. ссылка позиция TBA, на корпус сальника.",
    "Assemble the repair liner to the gland housing and use masking tape, Material Ref. Item 08-715, to hold it in that position. Use one layer of masking tape, Material Ref. Item 08-715, at each side of the repair liner. The masking tape must be sufficiently wide to bond to the repair liner and the adhesive PVC tape, Material Ref. Item TBA: make sure the ends touch but do not overlap.":
        "Установите ремонтный вкладыш на корпус сальника и закрепите маскировочной лентой, мат. ссылка позиция 08-715. Используйте один слой маскировочной ленты, мат. ссылка позиция 08-715, с каждой стороны ремонтного вкладыша. Маскировочная лента должна быть достаточно широкой для приклеивания к ремонтному вкладышу и клейкой ПВХ-ленте, мат. ссылка позиция TBA: убедитесь, что концы соприкасаются, но не перекрываются.",
    "Use a brush to apply a flat layer of the prepared mixture to the surfaces of the gland housing made rough.":
        "Кистью нанесите ровный слой подготовленной смеси на зашероховаченные поверхности корпуса сальника.",
    "Machine the diameter and width of the repair liner to the dimensions shown before you remove the adhesive PVC tape, Material Ref. Item TBA.":
        "Обработайте диаметр и ширину ремонтного вкладыша до указанных размеров перед снятием клейкой ПВХ-ленты, мат. ссылка позиция TBA.",
    # ─── Part 5 — remaining procedure sentences ───
    "Spray primer paint lightly on the cadmium plated surface: refer to PCS-2500.":
        "Слегка нанесите грунтовочную краску распылением на кадмированную поверхность: обратитесь к PCS-2500.",
    "Refer to PCS-6000-07 and identify the part with the applicable Messier-Dowty Limited repair number, adjacent to the part number, after painting:":
        "Обратитесь к PCS-6000-07 и идентифицируйте деталь соответствующим номером ремонта Messier-Dowty Limited, рядом с номером детали, после окраски:",
    "Grind diameter A to the dimensions shown with a surface finish of 0,4 micrometers (16 micro-inches). Refer to M-DLPS1031-3 and M-DLPS1031-6 for the chromium plate terminations where shown.":
        "Шлифуйте диаметр A до указанных размеров с шероховатостью поверхности 0,4 микрометра (16 микродюймов). Обратитесь к M-DLPS1031-3 и M-DLPS1031-6 для определения границ хромового покрытия в указанных местах.",
    "Machine the repair bushes to the contour of the pin and the dimensions shown.":
        "Обработайте ремонтные втулки по контуру штифта и указанным размерам.",
    "Machine the diameter A to remove the minimum amount of material necessary to remove the damage or wear, restore the 20,00 mm (0.787 in) radius in two places as shown: refer to M-DLPS1004-4-1 and Figure 601. Do not machine diameter A more than 33,99 mm (1.3383 in). Make the surface finish 1,6 micrometers (63 micro-inches).":
        "Обработайте диаметр A, удаляя минимально необходимое количество материала для устранения повреждения или износа, восстановите радиус 20,00 мм (0,787 дюйм) в двух местах как показано: обратитесь к M-DLPS1004-4-1 и Рисунок 601. Не обрабатывайте диаметр A более 33,99 мм (1,3383 дюйм). Шероховатость поверхности 1,6 микрометра (63 микродюйма).",
    "Apply sulphamate nickel plate to the reworked areas: refer to MIL STD 868A solution 2, PCS-2120 and Figure 601. The sulphamate nickel plate thickness must be sufficient to get the correct diameter after machining. Make sure that the cross hole and the lubrication holes are sufficiently masked: refer to Figure 601.":
        "Нанесите сульфаматное никелевое покрытие на доработанные участки: обратитесь к MIL STD 868A раствор 2, PCS-2120 и Рисунок 601. Толщина сульфаматного никелевого покрытия должна быть достаточной для получения правильного диаметра после обработки. Убедитесь, что поперечное отверстие и смазочные отверстия достаточно замаскированы: обратитесь к Рисунок 601.",
    "NOTE: The above procedure includes de-embrittle for 23 hours at 185oC to 195oC (366oF to 383oF).":
        "ПРИМЕЧАНИЕ: Вышеуказанная процедура включает снятие водородной хрупкости в течение 23 часов при 185°C — 195°C (366°F — 383°F).",
    "NOTE: The above procedure includes de-embrittle for 4 hours at 185oC to 195oC (366oF to 383oF).":
        "ПРИМЕЧАНИЕ: Вышеуказанная процедура включает снятие водородной хрупкости в течение 4 часов при 185°C — 195°C (366°F — 383°F).",
    "Examine the edges of sulphamate nickel plate to make sure they are properly bonded: use 5 or 10X magnification.":
        "Проверьте края сульфаматного никелевого покрытия, чтобы убедиться в их надлежащем сцеплении: используйте 5- или 10-кратное увеличение.",
    "Apply cadmium plate all over the pin except where chromium plated: refer to PCS-2100 and Figure 601. The cadmium plate thickness must be between 0,010 and 0,020 mm (0.0004 and 0.0008 in). Make sure the sulphamate nickel plate is fully encapsulated by cadmium plate.":
        "Нанесите кадмиевое покрытие на штифт повсюду, кроме хромированных участков: обратитесь к PCS-2100 и Рисунок 601. Толщина кадмиевого покрытия должна быть от 0,010 до 0,020 мм (0,0004 и 0,0008 дюйм). Убедитесь, что сульфаматное никелевое покрытие полностью инкапсулировано кадмиевым покрытием.",
    "Apply paint to the reworked areas: refer to REPAIR and PCS-2500.":
        "Нанесите краску на доработанные участки: обратитесь к REPAIR и PCS-2500.",
    "Measure and record the new diameters A.":
        "Измерьте и запишите новые диаметры A.",
    "Shot peen the reworked areas: refer to M-DLPS123 and Figure 601.":
        "Выполните дробеструйную обработку доработанных участков: обратитесь к M-DLPS123 и Рисунок 601.",
    "Locally apply cadmium plate to the reworked areas: refer to PCS-2141.":
        "Локально нанесите кадмиевое покрытие на доработанные участки: обратитесь к PCS-2141.",
    "Locally apply cadmium plate to the reworked areas of the sleeves: refer to PCS-2141.":
        "Локально нанесите кадмиевое покрытие на доработанные участки рукавов: обратитесь к PCS-2141.",
    "Machine the bores of the repair sleeves to a diameter between 12,700 and 12,733 mm (0.5000 and 0.5013 in). Machine the inner and outer ends of both sleeves to inner and outer profile of the pin. Make the radius between 0,5 and 0,75 mm (0.020 and 0.029 in) at inside and outside of both the sleeves: refer to Figure 601.":
        "Обработайте отверстия ремонтных рукавов до диаметра от 12,700 до 12,733 мм (0,5000 — 0,5013 дюйм). Обработайте внутренние и наружные торцы обоих рукавов по внутреннему и наружному профилю штифта. Радиус от 0,5 до 0,75 мм (0,020 — 0,029 дюйм) на внутренней и наружной сторонах обоих рукавов: обратитесь к Рисунок 601.",
    "Measure and make a record of the new diameter A and the thickness of the lug D.":
        "Измерьте и запишите новый диаметр A и толщину проушины D.",
    "Use the Press Pad 460004330/85 and Drift 460004331/7 to install the oversize bearing to the bracket: refer to M-DLPS1011-14. Check line ream the oversize bearing to the dimension shown in Figure 601.":
        "Используйте прижимную подушку 460004330/85 и выколотку 460004331/7 для установки ремонтного подшипника увеличенного размера в кронштейн: обратитесь к M-DLPS1011-14. Проверьте развёртку ремонтного подшипника до размера, указанного на Рисунок 601.",
    "Use the Press Pad 460004330/136 and Drift 460004331/7 to install the oversize bearing to the bracket: refer to M-DLPS1011-14. Check line ream the oversize bearing to the dimension shown in Figure 601.":
        "Используйте прижимную подушку 460004330/136 и выколотку 460004331/7 для установки ремонтного подшипника увеличенного размера в кронштейн: обратитесь к M-DLPS1011-14. Проверьте развёртку ремонтного подшипника до размера, указанного на Рисунок 601.",
    "Shot peen the reworked areas: refer to M-DLPS123 and Figure 601. The shot peen must not enter the bore or thread.":
        "Выполните дробеструйную обработку доработанных участков: обратитесь к M-DLPS123 и Рисунок 601. Дробеструйная обработка не должна попадать в отверстие или резьбу.",
    # Part 5 figure captions
    "Repair to Lower Bearing Subassembly Figure 601": "Ремонт сборки нижнего подшипника Рисунок 601",
    "Repair to Pin Figure 601": "Ремонт штифта Рисунок 601",
    "Repair to Pivot Pin Figure 601": "Ремонт штифта вращения Рисунок 601",
    "Repair to Uplock Pin Figure 601": "Ремонт штифта замка убранного положения Рисунок 601",
    "Repair to Bracket Figure 601": "Ремонт кронштейна Рисунок 601",
    # ─── Part 4 section/figure headings ───
    "Approved Repairs Table 602": "Утверждённые ремонты Таблица 602",
    "Approved Repairs Table 602 (Continued)": "Утверждённые ремонты Таблица 602 (Продолжение)",
    "Landing Systems Repair No.": "Ремонт Safran Landing Systems №",
    "Applicable Part": "Применяемая деталь",
    "TRANSFER BLOCK REFER TO FIGURE 655": "ПЕРЕХОДНЫЙ БЛОК ОБРАТИТЕСЬ К РИСУНКУ 655",
    "UPPER DIAPHRAGM TUBE REFER TO FIGURE 653": "ВЕРХНЯЯ ДИАФРАГМЕННАЯ ТРУБА ОБРАТИТЕСЬ К РИСУНКУ 653",
    "UPPER PIVOT BRACKET REFER TO FIGURE 657": "ВЕРХНИЙ ПОВОРОТНЫЙ КРОНШТЕЙН ОБРАТИТЕСЬ К РИСУНКУ 657",
    "UPPER TORQUE LINK REFER TO FIGURE 651": "ВЕРХНИЙ ШЛИЦ-ШАРНИР ОБРАТИТЕСЬ К РИСУНКУ 651",
    "LOWER TORQUE LINK REFER TO FIGURE 651": "НИЖНИЙ ШЛИЦ-ШАРНИР ОБРАТИТЕСЬ К РИСУНКУ 651",
    "CYLINDER REFER TO FIGURE 654": "ЦИЛИНДР ОБРАТИТЕСЬ К РИСУНКУ 654",
    "HARNESS SUPPORT BRACKET REFER TO FIGURE 656": "КРОНШТЕЙН КРЕПЛЕНИЯ ЖГУТА ОБРАТИТЕСЬ К РИСУНКУ 656",
    "BRAKE FLANGE TYPICAL 12 PLACES": "ТОРМОЗНОЙ ФЛАНЕЦ ТИПИЧНО 12 МЕСТ",
    "12 PLACES INCLUDING": "12 МЕСТ ВКЛЮЧАЯ",
    "TYPICAL 2 TRANSFER BLOCK LUGS": "ТИПИЧНО 2 ПРОУШИНЫ ПЕРЕХОДНОГО БЛОКА",
    "TYPICAL 2 BRAKE MANIFOLD LUGS": "ТИПИЧНО 2 ПРОУШИНЫ ТОРМОЗНОГО КОЛЛЕКТОРА",
    "SECTION L-L CHANGE OVER VALVE HOLES AND LUGS": "СЕЧЕНИЕ L-L ОТВЕРСТИЯ И ПРОУШИНЫ ПЕРЕПУСКНОГО КЛАПАНА",
    "TORQUE LINK AND RETAINING PIN BORES": "ОТВЕРСТИЯ ШЛИЦ-ШАРНИРА И СТОПОРНОГО ШТИФТА",
    "INCLUDING CHAMFER TYPICAL 2 LUGS": "ВКЛЮЧАЯ ФАСКУ ТИПИЧНО 2 ПРОУШИНЫ",
    "DO NOT PAINT": "НЕ КРАСИТЬ",
    "FACE": "ПОВЕРХНОСТЬ",
    "ON FACE": "НА ПОВЕРХНОСТИ",
    "SPOTFACE": "ЗЕНКОВКА",
    "A\tSPOTFACE": "A\tЗЕНКОВКА",
    "C\tSPOTFACE": "C\tЗЕНКОВКА",
    "A\tCHROMIUM PLATED SURFACE": "A\tХРОМИРОВАННАЯ ПОВЕРХНОСТЬ",
    # ─── Part 4 "FOR MAIN FITTING" drawing annotations ───
    "FOR MAIN FITTING (20-410C, 20-420C) ONLY": "ТОЛЬКО ДЛЯ КОРПУСА СТОЙКИ (20-410C, 20-420C)",
    "FOR MAIN FITTING (20-410C AND 20-420C) ONLY": "ТОЛЬКО ДЛЯ КОРПУСА СТОЙКИ (20-410C И 20-420C)",
    "FOR MAIN FITTING (20-410B, 20-420B, 20-410D AND 20-420D) ONLY":
        "ТОЛЬКО ДЛЯ КОРПУСА СТОЙКИ (20-410B, 20-420B, 20-410D И 20-420D)",
    "FOR MAIN FITTING (20-410B, 20-420B,": "ТОЛЬКО ДЛЯ КОРПУСА СТОЙКИ (20-410B, 20-420B,",
    "DRAG ARM HOLES": "ОТВЕРСТИЯ ТЯГИ",
    "RUN OUT BAND": "ЗОНА ВЫХОДА",
    "EDGE BLENDED WITH": "КРОМКА СОПРЯЖЕНА С",
    "ZINC NICKEL PLATE": "ЦИНК-НИКЕЛЕВОЕ ПОКРЫТИЕ",
    "WORKING DIA.": "РАБОЧИЙ ДИАМ.",
    "MAIN FITTING": "КОРПУС СТОЙКИ",
    "SURFACE": "ПОВЕРХНОСТЬ",
    "SLIDING TUBE": "СКОЛЬЗЯЩАЯ ТРУБА",
    "B HOLE": "B ОТВЕРСТИЕ",
    "C\t20-410D AND 20-420D) ONLY": "C\t20-410D И 20-420D) ТОЛЬКО",
    "C\t\t2 LUGS\tC": "C\t\t2 ПРОУШИНЫ\tC",
    # ─── Repair section procedure paragraphs ───
    "Approved repairs are in para 4. The repairs in this CMM have been approved under Airbus\u2019 EASA Design Organisation Approval No. EASA.21J.031.":
        "Утверждённые ремонты указаны в п. 4. Ремонты в данном CMM утверждены в рамках одобрения проектной организации Airbus EASA № EASA.21J.031.",
    "Repair damage to small areas of cadmium plated surfaces: refer to PCS-2141.":
        "Ремонт повреждений небольших участков кадмированных поверхностей: обратитесь к PCS-2141.",
    "Repair damage to small areas of anodized surfaces: refer to PCS-2220.":
        "Ремонт повреждений небольших участков анодированных поверхностей: обратитесь к PCS-2220.",
    "Repair damage to small areas of paint finish: refer to M-DLPS1003-1, use paint to PCS-2500.":
        "Ремонт повреждений небольших участков лакокрасочного покрытия: обратитесь к M-DLPS1003-1, используйте краску по PCS-2500.",
    "Chip damage of less than 10,0 mm2 (0.015 in2) can be restored with Sermetel 249 with Sermetel 273 catalyst: refer to M-DLPS637 (cold rework only).":
        "Сколы площадью менее 10,0 мм² (0,015 дюйм²) могут быть восстановлены с помощью Sermetel 249 с катализатором Sermetel 273: обратитесь к M-DLPS637 (только холодная доработка).",
    "PRIMER PAINT TO PCS-2500 OVER SERMETEL W TO LENGTH 305,00mm (12.000in) for (18-80)":
        "ГРУНТОВОЧНАЯ КРАСКА ПО PCS-2500 ПОВЕРХ SERMETEL W ДЛИНОЙ 305,00 мм (12,000 дюйм) для (18-80)",
    # ─── Protective Treatment Table 601 paragraph-level entries ───
    "PCS-2500. Do not paint:": "PCS-2500. Не окрашивайте:",
    "M-DLPS100-2. Do not include areas that have chromium plate. Paint: refer to M-DLPS1003-1 and":
        "M-DLPS100-2. Не включайте участки с хромовым покрытием. Окраска: обратитесь к M-DLPS1003-1 и",
    "PCS-2100. Do not include areas that have chromium plate. Paint: refer to M-DLPS1003-1 and":
        "PCS-2100. Не включайте участки с хромовым покрытием. Окраска: обратитесь к M-DLPS1003-1 и",
    "AMS5659 condition H1025": "AMS5659 состояние H1025",
    "9-190 Only": "только 9-190",
    "10-160 Only": "только 10-160",
    "MIL-A-8625 Type IB, Class 1.": "MIL-A-8625 тип IB, класс 1.",
    "M-DLPS100-2, M-DLPS137 or DEF STAN 03-19.":
        "M-DLPS100-2, M-DLPS137 или DEF STAN 03-19.",
    "Paint: refer to M-DLPS1003-1 and": "Окраска: обратитесь к M-DLPS1003-1 и",
    # ─── Protective Treatment Table 601 — full paragraph translations ───
    "Apply cadmium plate all over but not to the chromium plated areas: refer to PCS-2101. Make the cadmium plate thickness between 0,010 and 0,015 mm (0.0004 and 0.00059 in).":
        "Нанесите кадмиевое покрытие повсюду, кроме хромированных участков: обратитесь к PCS-2101. Толщина кадмиевого покрытия должна быть от 0,010 до 0,015 мм (0,0004 и 0,00059 дюйм).",
    "Apply cadmium plate all over: refer to M-DLPS100-1. The cadmium plate must be 0,010 to 0,015 mm (0.0004 to 0.0006 in) thick. Apply paint all over: refer to PCS-2500. Do not apply paint to the thread or to the surfaces that enter the transfer block (2-340 and 2-350)":
        "Нанесите кадмиевое покрытие повсюду: обратитесь к M-DLPS100-1. Толщина кадмиевого покрытия должна быть от 0,010 до 0,015 мм (0,0004 — 0,0006 дюйм). Нанесите краску повсюду: обратитесь к PCS-2500. Не наносите краску на резьбу и на поверхности, входящие в переходный блок (2-340 и 2-350)",
    "Apply cadmium plate all over: refer to M-DLPS100-2. Make the cadmium plate thickness between 0,010":
        "Нанесите кадмиевое покрытие повсюду: обратитесь к M-DLPS100-2. Толщина кадмиевого покрытия должна быть от 0,010",
    "Apply cadmium plate all over: refer to M-DLPS100-2S. The cadmium plate must be 0,010 to 0,015 mm (0.0004 to 0.0006 in) thick.":
        "Нанесите кадмиевое покрытие повсюду: обратитесь к M-DLPS100-2S. Толщина кадмиевого покрытия должна быть от 0,010 до 0,015 мм (0,0004 — 0,0006 дюйм).",
    "Apply cadmium plate to PCS-2100. The cadmium plate thickness should be between 0,010 and 0,015 mm (0.0004 and 0.0006 in). Do not apply cadmium plate:":
        "Нанесите кадмиевое покрытие по PCS-2100. Толщина кадмиевого покрытия должна быть от 0,010 до 0,015 мм (0,0004 и 0,0006 дюйм). Не наносите кадмиевое покрытие:",
    "Apply cadmium plate: refer to M-DLPS100-2. Paint: refer to":
        "Нанесите кадмиевое покрытие: обратитесь к M-DLPS100-2. Окраска: обратитесь к",
    "Apply one layer of primer paint only to the areas A: refer to PCS-2500. Apply paint all over as per PCS-2500 but not to:":
        "Нанесите только один слой грунтовочной краски на участки A: обратитесь к PCS-2500. Нанесите краску повсюду по PCS-2500, кроме:",
    "Apply paint all over but not on the chromium plated areas, the areas A, B, C and on the lubrication fitting bores where the lubrication adaptors (20-130), (20-160), (20-":
        "Нанесите краску повсюду, кроме хромированных участков, участков A, B, C и отверстий смазочных ниппелей, где устанавливаются смазочные адаптеры (20-130), (20-160), (20-",
    "Apply paint all over but not to the areas A and B: refer to PCS-2500. Apply primer paint only to areas A. Do not paint areas B.":
        "Нанесите краску повсюду, кроме участков A и B: обратитесь к PCS-2500. Нанесите грунтовочную краску только на участки A. Не красьте участки B.",
    "Apply paint all over but not to the areas A, C and D: refer to PCS-2500.":
        "Нанесите краску повсюду, кроме участков A, C и D: обратитесь к PCS-2500.",
    "Apply paint all over but not to the areas A: refer to PCS-2500. Do not paint areas A.":
        "Нанесите краску повсюду, кроме участков A: обратитесь к PCS-2500. Не красьте участки A.",
    "Apply paint: refer to M-DLPS1003-1 and PCS-2500. Do not paint:":
        "Нанесите краску: обратитесь к M-DLPS1003-1 и PCS-2500. Не красьте:",
    "Apply primer paint only to the areas":
        "Нанесите грунтовочную краску только на участки",
    "Apply primer paint to the areas A: refer to PCS-2500. Apply paint all over: refer to PCS-2500. Do not apply paint to:":
        "Нанесите грунтовочную краску на участки A: обратитесь к PCS-2500. Нанесите краску повсюду: обратитесь к PCS-2500. Не наносите краску на:",
    "Apply sermetel W only to the areas C: refer to IFC 40-860-03MD.":
        "Нанесите Sermetel W только на участки C: обратитесь к IFC 40-860-03MD.",
    "E. Apply paint all over but not to the chromium plated areas A, areas E and F: refer to PCS-2500.":
        "E. Нанесите краску повсюду, кроме хромированных участков A, участков E и F: обратитесь к PCS-2500.",
    "M-DLPS100-2. Do not include the area that has chromium plate. Paint: refer to M-DLPS1003-1 and":
        "M-DLPS100-2. Не включайте участок с хромовым покрытием. Окраска: обратитесь к M-DLPS1003-1 и",
    "M-DLPS100-2. The cadmium plate must be 0,010 to 0,015 mm (0.0004 to 0.0006 in) thick. Paint: refer to":
        "M-DLPS100-2. Толщина кадмиевого покрытия должна быть от 0,010 до 0,015 мм (0,0004 — 0,0006 дюйм). Окраска: обратитесь к",
    "M-DLPS102-1. Apply paint all over but not to the areas A: refer to":
        "M-DLPS102-1. Нанесите краску повсюду, кроме участков A: обратитесь к",
    "M-DLPS131. Apply brush cadmium plate to the areas D: refer to":
        "M-DLPS131. Нанесите кадмиевое покрытие кистью на участки D: обратитесь к",
    "M-DLPS131. Do not include areas that have chromium plate. Paint: refer to M-DLPS1003-1 and":
        "M-DLPS131. Не включайте участки с хромовым покрытием. Окраска: обратитесь к M-DLPS1003-1 и",
    "M-DLPS137. Paint all over externally and internally to areas B but not to the chromium plated areas, the bush and bearing bores, the chamfers, the lubrication fitting bores and areas C and D: refer to PCS-2500.":
        "M-DLPS137. Окрасьте повсюду снаружи и внутрь на участки B, кроме хромированных участков, отверстий под втулки и подшипники, фасок, отверстий смазочных ниппелей и участков C и D: обратитесь к PCS-2500.",
    "PCS-2100. Make the cadmium plate thickness 0,010 to 0,015 mm":
        "PCS-2100. Толщина кадмиевого покрытия должна быть от 0,010 до 0,015 мм",
    "PCS-2500. Apply only primer to face D: refer to PCS-2500.":
        "PCS-2500. Нанесите только грунтовку на поверхность D: обратитесь к PCS-2500.",
    "PCS-2500. Apply paint to areas D: refer to PCS-2500. No bare cadmium permitted.":
        "PCS-2500. Нанесите краску на участки D: обратитесь к PCS-2500. Открытый кадмий не допускается.",
    "PCS-2500. Apply primer paint only to the areas A.":
        "PCS-2500. Нанесите грунтовочную краску только на участки A.",
    "PCS-2500. Apply primer paint only to the contact faces of the flanges. Do not paint:":
        "PCS-2500. Нанесите грунтовочную краску только на контактные поверхности фланцев. Не красьте:",
    "Paint areas A: refer to PCS-2500. Primer paint only on face D and areas E and F including the chamfer. Do not paint:":
        "Нанесите краску на участки A: обратитесь к PCS-2500. Грунтовочная краска только на поверхность D и участки E и F, включая фаску. Не красьте:",
    "Paint external areas only: refer to M-DLPS1003-1 and PCS-2500.":
        "Окрасьте только наружные участки: обратитесь к M-DLPS1003-1 и PCS-2500.",
    "Paint: refer to M-DLPS1003-1 and PCS-2500. Do not paint:":
        "Окраска: обратитесь к M-DLPS1003-1 и PCS-2500. Не красьте:",
    "Paint: refer to PCS-2500. Do not paint:":
        "Окраска: обратитесь к PCS-2500. Не красьте:",
    "Passivate: refer to AMS2700": "Пассивировать: обратитесь к AMS2700",
    "Passivate: refer to AMS2700.": "Пассивировать: обратитесь к AMS2700.",
    "Refer to Figure 610. Apply cadmium plate: refer to M-DLPS131. Do not include areas that have chromium plate. Paint: refer to M-DLPS1003-1 and PCS-2500. Do not paint areas A and areas that have chromium plate.":
        "Обратитесь к рисунку 610. Нанесите кадмиевое покрытие: обратитесь к M-DLPS131. Не включайте участки с хромовым покрытием. Окраска: обратитесь к M-DLPS1003-1 и PCS-2500. Не красьте участки A и участки с хромовым покрытием.",
    "Refer to Figure 613. Apply cadmium plate: refer to M-DLPS100-2. Do not include areas that have chromium plate. Paint area A: refer to":
        "Обратитесь к рисунку 613. Нанесите кадмиевое покрытие: обратитесь к M-DLPS100-2. Не включайте участки с хромовым покрытием. Окрасьте участок A: обратитесь к",
    "Refer to Figure 613. Apply cadmium plate: refer to PCS-2101. Make the cadmium plate thickness between 0,010 and 0,015 mm (0.0004 and 0.0005 in). Do not include areas that have chromium plate. Paint area A: refer to PCS-2500.":
        "Обратитесь к рисунку 613. Нанесите кадмиевое покрытие: обратитесь к PCS-2101. Толщина кадмиевого покрытия должна быть от 0,010 до 0,015 мм (0,0004 и 0,0005 дюйм). Не включайте участки с хромовым покрытием. Окрасьте участок A: обратитесь к PCS-2500.",
    "Refer to Figure 614. Passivate all over: refer to AMS 2700. Apply paint all over externally but not on areas A. Apply primer paint to areas B: refer to PCS-2500.":
        "Обратитесь к рисунку 614. Пассивируйте повсюду: обратитесь к AMS 2700. Нанесите краску повсюду снаружи, кроме участков A. Нанесите грунтовочную краску на участки B: обратитесь к PCS-2500.",
    "Refer to Figure 614. Passivate all over: refer to AMS 2700. Apply paint all over externally but not on areas A. Apply primer paint to areas B: refer to PCS-2500":
        "Обратитесь к рисунку 614. Пассивируйте повсюду: обратитесь к AMS 2700. Нанесите краску повсюду снаружи, кроме участков A. Нанесите грунтовочную краску на участки B: обратитесь к PCS-2500",
    "Refer to Figure 615. Apply cadmium plate: refer to M-DLPS100-2. Do not include areas that have chromium plate. Paint: refer to M-DLPS1003-1 and PCS-2500. Do not paint areas that have chromium plate and areas A and B. Apply primer to areas B.":
        "Обратитесь к рисунку 615. Нанесите кадмиевое покрытие: обратитесь к M-DLPS100-2. Не включайте участки с хромовым покрытием. Окраска: обратитесь к M-DLPS1003-1 и PCS-2500. Не красьте участки с хромовым покрытием и участки A и B. Нанесите грунтовку на участки B.",
    "Refer to Figure 616. Apply cadmium plate internally and externally on area A: refer to M-DLPS100-2. Do not apply cadmium plate to bores B":
        "Обратитесь к рисунку 616. Нанесите кадмиевое покрытие изнутри и снаружи на участок A: обратитесь к M-DLPS100-2. Не наносите кадмиевое покрытие на отверстия B",
    "Refer to Figure 617. Cadmium plate all over to M-DLPS100-2. Make cadmium plate thickness 0,010 to 0,015 mm (0.0004 to 0.0006 in). Do":
        "Обратитесь к рисунку 617. Кадмиевое покрытие повсюду по M-DLPS100-2. Толщина кадмиевого покрытия от 0,010 до 0,015 мм (0,0004 — 0,0006 дюйм). Не",
    "Refer to Figure 619. Apply primer paint only to the areas D: refer to Figure 619 and PCS-2500.":
        "Обратитесь к рисунку 619. Нанесите грунтовочную краску только на участки D: обратитесь к рисунку 619 и PCS-2500.",
    "Refer to Figure 622. Apply cadmium plate all over but not to the chromium plated areas and areas A: refer to":
        "Обратитесь к рисунку 622. Нанесите кадмиевое покрытие повсюду, кроме хромированных участков и участков A: обратитесь к",
    "Refer to Figure 624. Apply cadmium plate all over but not to the areas A: refer to PCS-2101. Cadmium plate is optional on radii and chamfer B. Make the cadmium plate thickness between 0,010 and 0,015 mm":
        "Обратитесь к рисунку 624. Нанесите кадмиевое покрытие повсюду, кроме участков A: обратитесь к PCS-2101. Кадмиевое покрытие не является обязательным на радиусах и фаске B. Толщина кадмиевого покрытия должна быть от 0,010 до 0,015 мм",
    "Refer to Figure 625. Apply cadmium plate all over but not to the chromium plated areas A and areas B and C: refer to PCS-2100. Make the cadmium plate thickness between 0,010 and 0,020 mm (0.0004 and":
        "Обратитесь к рисунку 625. Нанесите кадмиевое покрытие повсюду, кроме хромированных участков A и участков B и C: обратитесь к PCS-2100. Толщина кадмиевого покрытия должна быть от 0,010 до 0,020 мм (0,0004 и",
    "Refer to Figure 626. Chromic acid anodise all over but not the spotfaces A: refer to M-DLPS102-1. Apply Alocrom 1200 to the areas A: refer to M-DLPS114. Apply paint all over but not to the areas A, B, C and D: refer to PCS-2500. Primer paint only on faces D. Apply light coat of primer to area B.":
        "Обратитесь к рисунку 626. Хромовокислотное анодирование повсюду, кроме площадок под крепёж A: обратитесь к M-DLPS102-1. Нанесите Alocrom 1200 на участки A: обратитесь к M-DLPS114. Нанесите краску повсюду, кроме участков A, B, C и D: обратитесь к PCS-2500. Грунтовочная краска только на поверхностях D. Нанесите лёгкий слой грунтовки на участок B.",
    "Refer to Figure 627. Chromic acid anodise all over but not the spotfaces A: refer to MIL-A-8625 Type 1B, Class 1. Apply Alocrom 1200 to the areas A: refer to PCS-2220 Type 2. Apply one coat of primer to the areas B: refer to PCS-2500. Apply primer only to the areas D: refer to":
        "Обратитесь к рисунку 627. Хромовокислотное анодирование повсюду, кроме площадок под крепёж A: обратитесь к MIL-A-8625 тип 1B, класс 1. Нанесите Alocrom 1200 на участки A: обратитесь к PCS-2220 тип 2. Нанесите один слой грунтовки на участки B: обратитесь к PCS-2500. Нанесите грунтовку только на участки D: обратитесь к",
    "Refer to Figure 629. Before installation of bushes: Apply primer paint to the areas A: refer to":
        "Обратитесь к рисунку 629. Перед установкой втулок: Нанесите грунтовочную краску на участки A: обратитесь к",
    "Refer to Figure 630. Before installation of bushes: Apply primer paint to areas A but not to the areas B: refer to PCS-2500.":
        "Обратитесь к рисунку 630. Перед установкой втулок: Нанесите грунтовочную краску на участки A, кроме участков B: обратитесь к PCS-2500.",
    "Refer to Figure 634. Chromic acid anodise all over but not the spotface A: refer to M-DLPS102-1. Apply Alocrom 1200 to the areas A: refer to M-DLPS114. Apply paint all over but not to the areas A, B, C and D: refer to PCS-2500. Apply a light coat of primer to the hole B: refer to":
        "Обратитесь к рисунку 634. Хромовокислотное анодирование повсюду, кроме площадки под крепёж A: обратитесь к M-DLPS102-1. Нанесите Alocrom 1200 на участки A: обратитесь к M-DLPS114. Нанесите краску повсюду, кроме участков A, B, C и D: обратитесь к PCS-2500. Нанесите лёгкий слой грунтовки на отверстие B: обратитесь к",
    "Refer to Figure 638. Apply paint all over externally, but not to areas A: refer to PCS-2500.":
        "Обратитесь к рисунку 638. Нанесите краску повсюду снаружи, кроме участков A: обратитесь к PCS-2500.",
    "Refer to Figure 639. Apply cadmium plate: refer to PCS-2100. Do not include areas that have chromium plate. Make the cadmium plate thickness between 0,010 and 0,020":
        "Обратитесь к рисунку 639. Нанесите кадмиевое покрытие: обратитесь к PCS-2100. Не включайте участки с хромовым покрытием. Толщина кадмиевого покрытия должна быть от 0,010 до 0,020",
    "Refer to Figures 620 and 621. Apply cadmium plate to M-DLPS131, do not apply cadmium plate to areas A. Primer paint all over but not in holes and on areas identified B and the 22,0 mm (0.87 in) diameters C: refer to PCS-2500.":
        "Обратитесь к рисункам 620 и 621. Нанесите кадмиевое покрытие по M-DLPS131, не наносите кадмиевое покрытие на участки A. Грунтовочная краска повсюду, кроме отверстий и участков, обозначенных B, и диаметров 22,0 мм (0,87 дюйм) C: обратитесь к PCS-2500.",
    "Refer to PCS-2500 and apply paint all over externally, but not to:":
        "Обратитесь к PCS-2500 и нанесите краску повсюду снаружи, кроме:",
    "Refer to PCS-2500 and apply paint internally along surface B, but not along surface C.":
        "Обратитесь к PCS-2500 и нанесите краску изнутри вдоль поверхности B, кроме поверхности C.",
    "not paint areas A and B. Apply primer to area A: refer to Figure 604.":
        "не красьте участки A и B. Нанесите грунтовку на участок A: обратитесь к рисунку 604.",
    "not paint the screw threads and the face that touches the wheel bearings.":
        "не красьте резьбу и поверхность, касающуюся подшипников колеса.",
    "the holes (with or without threads)":
        "отверстия (с резьбой или без)",
    "the lubrication fittings and their identification washers":
        "смазочные ниппели и их идентификационные шайбы",
    "the split pin hole.": "отверстие для шплинта.",
    "the two holes through the end. Apply a thin coat of primer paint to the holes through the end: refer to PCS-2500.":
        "два сквозных отверстия на конце. Нанесите тонкий слой грунтовочной краски на сквозные отверстия на конце: обратитесь к PCS-2500.",
    "where identified on Figure 618. Apply Sermetel W where shown: refer to IFC 40-860-03MD. Apply primer as shown and finish paint to PCS-2500. Do not apply paint:":
        "как указано на рисунке 618. Нанесите Sermetel W, как показано: обратитесь к IFC 40-860-03MD. Нанесите грунтовку, как показано, и финишную краску по PCS-2500. Не наносите краску:",
    "15,00 mm (0.591 in) diameter areas around the holes on the inside face of one flange.":
        "участки диаметром 15,00 мм (0,591 дюйм) вокруг отверстий на внутренней поверхности одного фланца.",
    "After installation of bushes: Apply paint to the areas B but not to the bushes: refer to IFC 30-117-05.":
        "После установки втулок: Нанесите краску на участки B, кроме втулок: обратитесь к IFC 30-117-05.",
    "After installation of bushes: Apply paint to the areas C but not to the areas D: refer to PCS-2500.":
        "После установки втулок: Нанесите краску на участки C, кроме участков D: обратитесь к PCS-2500.",
    "plate internally and externally over area A: refer to PCS-2101. Make the cadmium plate thickness between 0,010 and 0,015 mm (0.0004 and":
        "покрытие изнутри и снаружи на участок A: обратитесь к PCS-2101. Толщина кадмиевого покрытия должна быть от 0,010 до 0,015 мм (0,0004 и",
    "Apply only primer paint to areas C including chamfer: refer to":
        "Нанесите только грунтовочную краску на участки C, включая фаску: обратитесь к",
    "190) and (20-220) will install: refer to PCS-2500.":
        "190) и (20-220) будут установлены: обратитесь к PCS-2500.",
    "190) and (20-220) will install: refer to PCS-2500. Paint finish is optional in areas E.":
        "190) и (20-220) будут установлены: обратитесь к PCS-2500. Лакокрасочное покрытие является необязательным на участках E.",
    "M-DLPS1003-1 and PCS-2500. Do":
        "M-DLPS1003-1 и PCS-2500. Не",
    "Do not paint areas F.": "Не красьте участки F.",
    "or Steel, 4340 to AMS6414": "или сталь, 4340 по AMS6414",
    "Apply primer paint only to the contact face. Paint must not go in the bores.":
        "Нанесите грунтовочную краску только на контактную поверхность. Краска не должна попадать в отверстия.",
    "PCS-2500. Do not paint areas A and areas that have chromium plate: refer to Figure 608.":
        "PCS-2500. Не красьте участки A и участки с хромовым покрытием: обратитесь к рисунку 608.",
    "M-DLPS1003-1 and PCS-2500.": "M-DLPS1003-1 и PCS-2500.",
    "where identified on Figure 618": "как указано на рисунке 618",
    "to the chromium plated areas.": "на хромированные участки.",
    "mm (0.0004 and 0.00078 in). The cadmium plate must overlap the chromium plate run out. Bare metal not permitted.":
        "мм (0,0004 и 0,00078 дюйм). Кадмиевое покрытие должно перекрывать выход хромового покрытия. Открытый металл не допускается.",
    "the areas A.": "участки A.",
    "and 0,015 mm (0.0004": "до 0,015 мм (0,0004",
    "and 0.0005 in).": "и 0,0005 дюйм).",
    "Chromium plated areas,": "Хромированные участки,",
    "Areas A and B.": "Участки A и B.",
    "not cadmium plate:": "не наносить кадмиевое покрытие:",
    "chromium plated area C": "хромированный участок C",
    "the 3 holes in face D.": "3 отверстия в поверхности D.",
    "(0.0004 and 0.0006 in).": "(0,0004 и 0,0006 дюйм).",
    "and C": "и C",
    "the area that has chromium plate": "участок с хромовым покрытием",
    "Apply paint all over: refer to M-DLPS100-1.": "Нанесите кадмиевое покрытие повсюду: обратитесь к M-DLPS100-1.",
    "IFC 30-117-05.": "IFC 30-117-05.",
    # ─── Partially translated paragraph sentences ───
    "Repair of wear or damage with an approved Messier-Dowty Limited or Safran Landing Systems repair.":
        "Ремонт износа или повреждения с помощью утверждённого ремонта Messier-Dowty Limited или Safran Landing Systems.",
    "Before you repair a part that is identified with a concession, salvage or repair number, write to Safran Landing Systems for approval. Such numbers are adjacent to the part number, for example:":
        "Перед ремонтом детали, обозначенной номером допуска на отклонение, восстановления или ремонта, обратитесь в Safran Landing Systems за разрешением. Такие номера расположены рядом с номером детали, например:",
    "Before you repair a part that is identified with a concession, salvage or repair number, write to Safran Landing Systems for approval. Such numbers are adjacent to the part number":
        "Перед ремонтом детали, обозначенной номером допуска на отклонение, восстановления или ремонта, обратитесь в Safran Landing Systems за разрешением. Такие номера расположены рядом с номером детали",
    "If the repairs in this manual cannot correct the wear or damage to the part, write to Safran Landing Systems: refer to M-DLPS3002.":
        "Если ремонты, описанные в данном руководстве, не могут устранить износ или повреждение детали, обратитесь в Safran Landing Systems: см. M-DLPS3002.",
    "Identify the parts after repair with the Messier-Dowty Limited or Safran Landing Systems Repair Number: refer to the applicable repair for instructions.":
        "Идентифицируйте детали после ремонта с помощью номера ремонта Messier-Dowty Limited или Safran Landing Systems: обратитесь к соответствующему ремонту за инструкциями.",
    "There are two levels of repair procedure for parts that are found to be unserviceable after inspection: refer to CHECK.":
        "Существует два уровня процедуры ремонта деталей, признанных непригодными после осмотра: обратитесь к ПРОВЕРКЕ.",
    "CAUTION: DO NOT REPAIR A PART WITH A PROCEDURE THAT IS NOT APPROVED.":
        "ВНИМАНИЕ: НЕ РЕМОНТИРУЙТЕ ДЕТАЛЬ ПО НЕУТВЕРЖДЁННОЙ ПРОЦЕДУРЕ.",
    "Unless instructions are different in the approved repair, the applicable tolerances are:":
        "Если в утверждённом ремонте не указано иное, применяются следующие допуски:",
    "Protective treatment replacement procedures and the applicable parts are given in Table 601.":
        "Процедуры замены защитного покрытия и применяемые детали приведены в таблице 601.",
    "CAUTION: YOU MUST COMPLETE THE PROCESSES THAT FOLLOW IN THE SEQUENCE SHOWN. FAILURE TO DO THE PROCESSES IN THE CORRECT SEQUENCE CAN DAMAGE THE SLIDING TUBE (18-80) OR (18-80A) OR":
        "ВНИМАНИЕ: ВЫ ДОЛЖНЫ ВЫПОЛНИТЬ СЛЕДУЮЩИЕ ПРОЦЕССЫ В УКАЗАННОЙ ПОСЛЕДОВАТЕЛЬНОСТИ. НЕВЫПОЛНЕНИЕ ПРОЦЕССОВ В ПРАВИЛЬНОЙ ПОСЛЕДОВАТЕЛЬНОСТИ МОЖЕТ ПОВРЕДИТЬ СКОЛЬЗЯЩУЮ ТРУБУ (18-80) ИЛИ (18-80A) ИЛИ",
    "(18-80B) OR REDUCE THE EFFECT OF THE PROTECTIVE TREATMENTS.":
        "(18-80B) ИЛИ СНИЗИТЬ ЭФФЕКТИВНОСТЬ ЗАЩИТНЫХ ПОКРЫТИЙ.",
    "Chromium plate processes.": "Процессы хромового покрытия.",
    "Cadmium plate processes.": "Процессы кадмиевого покрытия.",
    "Sermetel W processes.": "Процессы Sermetel W.",
    "Paint processes.": "Процессы окраски.",
    "If you apply protective treatment processes that include Sermetel W to any ultra high tensile (UHT) steel part, the sequence of the processes is important. The sequence of the protective treatment processes must be:":
        "Если вы применяете процессы защитной обработки, включающие Sermetel W, к любой детали из сверхвысокопрочной стали (UHT), последовательность процессов важна. Последовательность процессов защитной обработки должна быть:",
    "If you apply protective treatment processes that include Sermetel W to any ultra high tensile (UHT) steel part, the sequence of the processes is important. The sequence of the protective treatment processes must be as follows:":
        "Если вы применяете процессы защитной обработки, включающие Sermetel W, к любой детали из сверхвысокопрочной стали (UHT), последовательность процессов важна. Последовательность процессов защитной обработки должна быть следующей:",
    # ─── Part 3 repair procedure sentences ───
    "Repair isolated external scores, smooth dents and abrasions, that have no cracks and no effect on internal dimensions: refer to para (2). Such damage must not be:":
        "Отремонтируйте изолированные наружные задиры, плавные вмятины и потёртости, не имеющие трещин и не влияющие на внутренние размеры: обратитесь к п. (2). Такие повреждения не должны быть:",
    "More than 19,00 mm (0.750 in) in length": "Более 19,00 мм (0,750 дюйм) в длину",
    "More than 0,76 mm (0.030 in) in depth": "Более 0,76 мм (0,030 дюйм) в глубину",
    "Less than one diameter from a hole and less than 6,35 mm (0.250 in) from a bearing surface":
        "Менее одного диаметра от отверстия и менее 6,35 мм (0,250 дюйм) от опорной поверхности",
    "General tolerance: + or - 0,25 mm (0.010 in)": "Общий допуск: + или - 0,25 мм (0,010 дюйм)",
    "Holes that are drilled or machined: + 0,25 to - 0,05 mm (+ 0.010 to - 0.002 in)":
        "Отверстия сверлёные или обработанные: + 0,25 до - 0,05 мм (+ 0,010 до - 0,002 дюйм)",
    "Angular tolerance: + or - 0,5 degree.": "Угловой допуск: + или - 0,5 градуса.",
    "Remove burrs, corrosion and sharp edges: the area of damage must not be more than 645 mm2 (1.0 in2) for each 6450 mm2 (10.0 in2). Subsequently, remove 0,127 mm (0.0050 in) more of the material and repair the protective treatment.":
        "Удалите заусенцы, коррозию и острые кромки: площадь повреждения не должна превышать 645 мм² (1,0 дюйм²) на каждые 6450 мм² (10,0 дюйм²). Затем удалите дополнительно 0,127 мм (0,0050 дюйм) материала и восстановите защитное покрытие.",
    "In a bore that will not seal, ignore abrasions and small scores that have no burrs. If there are burrs, remove them plus 0,127 mm (0.0050 in) of material from the area. Repair the protective treatment.":
        "В отверстии без уплотнения допускаются потёртости и мелкие задиры без заусенцев. При наличии заусенцев удалите их и дополнительно 0,127 мм (0,0050 дюйм) материала. Восстановите защитное покрытие.",
    "In a bore that will seal, polish scores to remove them. Make sure that the surface finish, concentricity and fits and clearances do not change.":
        "В отверстии с уплотнением отполируйте задиры для их удаления. Убедитесь, что чистота поверхности, концентричность и посадки и зазоры не изменились.",
    # Mixed-case variants for textbox labels
    "Record of Revisions": "Запись изменений",
    "List of Service Bulletins": "Список сервисных бюллетеней",
    "List of Effective Pages": "Перечень действующих страниц",
    "Record of Temporary Revisions": "Запись временных изменений",
    "Testing and Fault Isolation": "Проверка и поиск неисправностей",
    "Unit Identification Chart": "Таблица идентификации изделия",
    "Assembly (Including Storage)": "Сборка (включая хранение)",
    "Special Tools, Fixtures and Equipment": "Специальные инструменты, приспособления и оборудование",
    "Illustrated Parts List": "Иллюстрированный перечень деталей",
    "Fits and Clearances": "Посадки и зазоры",
    # ─── Procedural sentences (FIXED avoids PROCEDURAL_VOCAB ordering issues) ───
    "Examine the unit for damage before you start the tests.":
        "Осмотрите изделие на наличие повреждений перед началом испытаний.",
    "Discard parts that you must not use again. These include:":
        "Утилизируйте детали, непригодные для повторного использования. К ним относятся:",
    "Remove the spherical bearing (4-50) from the rod end (4-60).":
        "Снимите сферический подшипник (4-50) со стержневого наконечника (4-60).",
    "Remove the bracket subassembly (4-330).":
        "Снимите сборку кронштейна (4-330).",
    "Remove the spherical bearing (5-70) from the bracket (5-80): refer to M-DLPS1014-2.":
        "Снимите сферический подшипник (5-70) с кронштейна (5-80): см. M-DLPS1014-2.",
    "Remove the spherical bearing (5-150) from the bracket (5-160): refer to M-DLPS1014-2.":
        "Снимите сферический подшипник (5-150) с кронштейна (5-160): см. M-DLPS1014-2.",
    "Remove the pivot bracket subassembly (7-120).":
        "Снимите сборку поворотного кронштейна (7-120).",
    "Remove the O-ring seals (10-150) from the lubrication shaft subassembly (10-90).":
        "Снимите уплотнительные кольца (10-150) с сборки вала смазки (10-90).",
    "Release the cup washers (13-20).":
        "Отпустите тарельчатые шайбы (13-20).",
    "Remove the sliding tube subassembly (17-240) from the Build Trolley 460007240.":
        "Снимите сборку скользящей трубы (17-240) со сборочной тележки 460007240.",
    "Remove the bearing (20-250 only).":
        "Снимите подшипник (20-250 только).",
    "Remove the bearing (20-370 only).":
        "Снимите подшипник (20-370 только).",
    "Remove drag arm sleeve (20-370A)":
        "Снимите втулку тяги (20-370A)",
    "Remove the split pin (4-20), the nut (4-30), the washer (4-40) and the rod end assembly (4-10).":
        "Снимите шплинт (4-20), гайку (4-30), шайбу (4-40) и сборку стержневого наконечника (4-10).",
    "Remove the joint seal (16-80), the sealing ring (16-90) and the wiper ring (16-100).":
        "Снимите уплотнение стыка (16-80), уплотнительное кольцо (16-90) и грязесъёмное кольцо (16-100).",
    "Remove the joint seal (16A-80), the sealing ring (16A-90) and the wiper ring (16A-100).":
        "Снимите уплотнение стыка (16A-80), уплотнительное кольцо (16A-90) и грязесъёмное кольцо (16A-100).",
    "Examine all parts shown in Tables 501 and 502 to the applicable NDT and information given.":
        "Осмотрите все детали, указанные в таблицах 501 и 502, в соответствии с применимым неразрушающим контролем и указанной информацией.",
    "Parts that are included in Tables 501 and 502 must be fully disassembled to the lowest detail level for inspection.":
        "Детали, включённые в таблицы 501 и 502, должны быть полностью разобраны до наименьшего уровня детализации для осмотра.",
    "Parts that are included in Tables 501 and 502 must be fully disassembled to the lowest detail level for NDT inspection. This includes the removal of all of the bushes.":
        "Детали, включённые в таблицы 501 и 502, должны быть полностью разобраны до наименьшего уровня детализации для неразрушающего контроля. Это включает снятие всех втулок.",
    "Hold the main landing gear leg (1-1), use with 460007281 and 460007282":
        "Удерживайте стойку основного шасси (1-1), используйте с 460007281 и 460007282",
    # ─── TOC / section heading sentences ───
    "Pre SB 201-32-49 or Pre SB 201-32-58 or Pre SB 201-32-60 Lower Bearing (16-150)":
        "До SB 201-32-49 или до SB 201-32-58 или до SB 201-32-60 Нижний подшипник (16-150)",
    "or (16A-150) Grease Groove Dimensions After Installation in the Gland Housing (16-140)":
        "или (16A-150) Размеры смазочных канавок после установки в корпус сальника (16-140)",
    "Post SB-201-32-58 - Lower Bearing Subassembly Machining and Liner Installation .":
        "После SB-201-32-58 — Механическая обработка сборки нижнего подшипника и установка вкладыша .",
    "Rod End Assembly (4-10), Proximity Switch (4-100) and Bracket Subassembly (4-330)":
        "Сборка стержневого наконечника (4-10), датчик приближения (4-100) и сборка кронштейна (4-330)",
    "Bracket Assembly (5-10), Bracket Subassemblies (5-90 and 5-270) and Uplock Pin (5-400)":
        "Сборка кронштейна (5-10), сборки кронштейнов (5-90 и 5-270) и штифт замка убранного положения (5-400)",
    "Slave Link Subassembly (6-190) and Lower Slave Link Subassembly (6-290)":
        "Сборка ведомого звена (6-190) и сборка нижнего ведомого звена (6-290)",
    "Proximity Switches (7-40 and 7-230) and Harness Support Bracket (7-100)":
        "Датчики приближения (7-40 и 7-230) и кронштейн крепления жгута (7-100)",
    "1M Electrical Axle Harness (11-40) and 2M Electrical Axle Harness (11-50)":
        "Электрический жгут оси 1М (11-40) и электрический жгут оси 2М (11-50)",
    "Lower Bearing Subassembly (16-110D or 16A-110E) Post Ref. Code: 2253":
        "Сборка нижнего подшипника (16-110D или 16A-110E) После Код ссылки: 2253",
    "rings (16-50 or 16A-50).":
        "кольца (16-50 или 16A-50).",
    "Remove the joint seal (16-80 or 16A-80), the sealing ring (16-90 or 16A-90) and the wiper ring (16-100 or 16A-100).":
        "Снимите уплотнение стыка (16-80 или 16A-80), уплотнительное кольцо (16-90 или 16A-90) и грязесъёмное кольцо (16-100 или 16A-100).",
    "Spherical Bearing (19-50) and Bung (19-60)":
        "Сферический подшипник (19-50) и заглушка (19-60)",
    # ─── Table cell translations ───
    "Lift the sliding tube subassembly (17-240) and related parts":
        "Поднимите сборку скользящей трубы (17-240) и связанные детали",
    "Hold the sliding tube subassembly (17-240) and related parts":
        "Удерживайте сборку скользящей трубы (17-240) и связанные детали",
    # ─── NOTE/technical sentences ───
    "NOTE:\tThe thread size is M142 x 1.5 pitch - 5h6h to BS3643.":
        "ПРИМЕЧАНИЕ:\tРазмер резьбы M142 x 1,5 шаг — 5h6h по BS3643.",
    # ─── Component name combos for TOC ───
    "1M Electrical Axle Harness (11-40) and 2M Electrical Axle Harness (11-50)-":
        "Электрический жгут оси 1М (11-40) и электрический жгут оси 2М (11-50) —",
    "Labels (20-10, 20-30, 20-40, 20-60 and 20-80) and wiring diagram plate (1-110)":
        "Этикетки (20-10, 20-30, 20-40, 20-60 и 20-80) и табличка электрической схемы (1-110)",
    # ─── Testing section procedural sentences ───
    "The hydraulic test rig must have a hand pump and a power pump. The power pump must have a controlled flow of not less than 4,5 l/min (4.62 in3/sec).":
        "Стенд для гидравлических испытаний должен иметь ручной насос и силовой насос. Силовой насос должен иметь регулируемый расход не менее 4,5 л/мин (4,62 дюйм³/с).",
    "During all hydraulic tests, the unit and the test circuit must be hydraulically full.":
        "Во время всех гидравлических испытаний изделие и испытательный контур должны быть полностью заполнены гидравлической жидкостью.",
    "Use these special tools to install the main landing gear leg (1-1) vertically in the loading press:":
        "Используйте следующие специальные инструменты для установки стойки основного шасси (1-1) вертикально в нагрузочный пресс:",
    "CAUTION: DO NOT PUT AN END LOAD OF MORE THAN 5,08 TONNES (5 TONS) ON THE MAIN LANDING GEAR LEG (1-1).":
        "ВНИМАНИЕ: НЕ ПРИКЛАДЫВАЙТЕ ОСЕВУЮ НАГРУЗКУ БОЛЕЕ 5,08 ТОНН (5 ТОНН) НА СТОЙКУ ОСНОВНОГО ШАССИ (1-1).",
    # ─── Testing section ───
    "During the proximity switch tests the ambient temperature must be between 15 and 25 \uf0b0C (59 and 77 \uf0b0F).":
        "Во время испытаний датчиков приближения температура окружающей среды должна быть между 15 и 25 °C (59 и 77 °F).",
    "Main Landing Gear Leg (1-1) Tests":
        "Стойка основного шасси (1-1) Испытания",
    "Refer to Figure 101 and measure the dimension X: it must be between 483,05 and 487,85 mm (19.017 and 19.207 in).":
        "Обратитесь к рисунку 101 и измерьте размер X: он должен быть между 483,05 и 487,85 мм (19,017 и 19,207 дюйм).",
    "Open the charging valve (13-60) and release the nitrogen pressure. Do not close the charging valve (13-60).":
        "Откройте зарядный клапан (13-60) и сбросьте давление азота. Не закрывайте зарядный клапан (13-60).",
    "Main Landing Gear Leg (1-1) Figure 101":
        "Стойка основного шасси (1-1) Рисунок 101",
    "Write this data on a label and attach it to the unit: THE GEAR MUST BE INFLATED TO THE APPROPRIATE PRESSURES BEFORE BEING PLACED IN SERVICE.":
        "Запишите эти данные на этикетку и прикрепите её к изделию: СТОЙКА ДОЛЖНА БЫТЬ НАКАЧАНА ДО СООТВЕТСТВУЮЩИХ ДАВЛЕНИЙ ПЕРЕД ВВОДОМ В ЭКСПЛУАТАЦИЮ.",
    "Connect the 28 VDC power supply, the Lampbox 460005842 and the main landing gear leg (1-1).":
        "Подсоедините источник питания 28 В пост. тока, контрольную лампу 460005842 и стойку основного шасси (1-1).",
    "Use the loading press to fully extend the main landing gear leg (1-1).":
        "Используйте нагрузочный пресс для полного выдвижения стойки основного шасси (1-1).",
    "Use the loading press to slowly close the main landing gear leg (1-1):":
        "Используйте нагрузочный пресс для медленного закрытия стойки основного шасси (1-1):",
    "The proximity switch (7-230) must operate before the main landing gear leg (1-1) has closed by 26,00 mm (1.0236 in).":
        "Датчик приближения (7-230) должен сработать до того, как стойка основного шасси (1-1) закроется на 26,00 мм (1,0236 дюйм).",
    "The proximity switch (7-40) must operate before the main landing gear leg (1-1) has closed by 29,30 mm (1.1535 in).":
        "Датчик приближения (7-40) должен сработать до того, как стойка основного шасси (1-1) закроется на 29,30 мм (1,1535 дюйм).",
    "Disconnect the 28 VDC supply and the Lampbox 460005842.":
        "Отсоедините источник питания 28 В пост. тока и контрольную лампу 460005842.",
    "Electrical Bonding Resistance Test Points (Tables 101 and 102) Figure 102":
        "Точки проверки сопротивления электрического соединения (Таблицы 101 и 102) Рисунок 102",
    "If necessary, remove the wire thread inserts:":
        "При необходимости снимите резьбовые спиральные вставки:",
    # ─── Disassembly sentences ───
    "Remove the pintle pin (1-60).":
        "Снимите штифт навеса (1-60).",
    "Use the Extractor Pad and Drawbolt 460006415 to remove the bush (2-310) and the Extractor 460006416 to remove the bush (2-320) from the transfer block (2-340 and 2-340A).":
        "Используйте Extractor Pad и Drawbolt 460006415 для снятия втулки (2-310) и Extractor 460006416 для снятия втулки (2-320) из переходного блока (2-340 и 2-340A).",
    "Use the Extractor Pad and Drawbolt 460006415 to remove the bush (2-310) and the Extractor 460006416 to remove the bush (2-320) from the transfer block (2-340B).":
        "Используйте Extractor Pad и Drawbolt 460006415 для снятия втулки (2-310) и Extractor 460006416 для снятия втулки (2-320) из переходного блока (2-340B).",
    "Remove the nuts (6-100), the washers (6-110), the cap screws (6-120), the target (6-130) and the spacers (6-140) or the laminated shim (6-140A).":
        "Снимите гайки (6-100), шайбы (6-110), винты с головкой (6-120), мишень (6-130) и проставки (6-140) или набор прокладок (6-140A).",
    "NOTE: If the calculated gap is in the tolerance, the spacers (6-140) or the laminated shim (6-140A) is not installed.":
        "ПРИМЕЧАНИЕ: Если рассчитанный зазор находится в пределах допуска, проставки (6-140) или набор прокладок (6-140A) не устанавливается.",
    "Remove the nuts (7-10), the washers (7-20), the cap screws (7-30), the proximity switch (7-40) and the spacer (7-50) or the laminated shim (7-50A).":
        "Снимите гайки (7-10), шайбы (7-20), винты с головкой (7-30), датчик приближения (7-40) и проставку (7-50) или набор прокладок (7-50A).",
    "NOTE: If the calculated gap is in the tolerance, the spacer (7-50) or the laminated shim (7-50A) is not installed.":
        "ПРИМЕЧАНИЕ: Если рассчитанный зазор находится в пределах допуска, проставка (7-50) или набор прокладок (7-50A) не устанавливается.",
    "Remove the nuts (7-150), the washers (7-160), the cap screws (7-170), the target (7-180) and the spacer (7-190) or the laminated shim (7-190A).":
        "Снимите гайки (7-150), шайбы (7-160), винты с головкой (7-170), мишень (7-180) и проставку (7-190) или набор прокладок (7-190A).",
    "NOTE: If the calculated gap is in the tolerance, the spacer (7-190) or the laminated shim (7-190A) is not installed.":
        "ПРИМЕЧАНИЕ: Если рассчитанный зазор находится в пределах допуска, проставка (7-190) или набор прокладок (7-190A) не устанавливается.",
    "Remove the nuts (7-200), the washers (7-210), the cap screws (7-220), the proximity switch (7-230) and the spacer (7-240) or the laminated shim (7-240A).":
        "Снимите гайки (7-200), шайбы (7-210), винты с головкой (7-220), датчик приближения (7-230) и проставку (7-240) или набор прокладок (7-240A).",
    "Remove the cap screws (11-10), the washers (11-20) and the harness support (11-30).":
        "Снимите винты с головкой (11-10), шайбы (11-20) и крепление жгута (11-30).",
    "Remove the retaining pins (13-10) and the cup washers (13-20).":
        "Снимите стопорные штифты (13-10) и тарельчатые шайбы (13-20).",
    "Use Pin Spanner 460007279 to remove the upper bearing housing (15-40).":
        "Используйте Pin Spanner 460007279 для снятия корпуса верхнего подшипника (15-40).",
    "Remove the lower bearing subassembly (16-110) and its related parts from the sliding tube subassembly (17-240).":
        "Снимите сборку нижнего подшипника (16-110) и связанные детали с сборки скользящей трубы (17-240).",
    "Remove the lower bearing (16-150) from the lower bearing housing subassembly (16-120).":
        "Снимите нижний подшипник (16-150) из сборки корпуса нижнего подшипника (16-120).",
    "Remove the lower bearing subassembly (16A-110D) and its related parts from the sliding tube subassembly (17-240).":
        "Снимите сборку нижнего подшипника (16A-110D) и связанные детали с сборки скользящей трубы (17-240).",
    "Remove the inner liner (16A-117) from the lower bearing subassembly (16A-110D) and discard it.":
        "Снимите внутренний вкладыш (16A-117) из сборки нижнего подшипника (16A-110D) и утилизируйте его.",
    "Remove the lower bearing subassembly (16-110D or 16A-110E) and its related parts from the sliding tube subassembly (17-240).":
        "Снимите сборку нижнего подшипника (16-110D или 16A-110E) и связанные детали с сборки скользящей трубы (17-240).",
    "Remove the O-ring seal (16-20A or 16A-20A), the backing rings (16-30 or 16A-30), the O-ring seal (16-40A or 16A-40A) and the backing":
        "Снимите уплотнительное кольцо (16-20A или 16A-20A), опорные кольца (16-30 или 16A-30), уплотнительное кольцо (16-40A или 16A-40A) и опорные",
    "Remove the common lower bearing bushes (16-130A or 16A-130B) from the lower bearing housing (16-140B or 16A-140C).":
        "Снимите общие втулки нижнего подшипника (16-130A или 16A-130B) из корпуса нижнего подшипника (16-140B или 16A-140C).",
    "Remove the cap screws (17-30), the washers (17-40) and the valve support (17-50).":
        "Снимите винты с головкой (17-30), шайбы (17-40) и опору клапана (17-50).",
    "Remove the lubrication fittings (18-52) and identification washers (18-54).":
        "Снимите фитинги для смазки (18-52) и идентификационные шайбы (18-54).",
    "Remove the lubrication fittings (20-200) and identification washers (20-210). Remove the lubrication adapters (20-220).":
        "Снимите фитинги для смазки (20-200) и идентификационные шайбы (20-210). Снимите адаптеры для смазки (20-220).",
    "Use the Extractor Plate 460007259/460006151/9 and the Drift 460004331/21 to remove the bushes (20-340 and 20-350).":
        "Используйте Extractor Plate 460007259/460006151/9 и Drift 460004331/21 для снятия втулок (20-340 и 20-350).",
    "WARNING: DO NOT GET PAINT STRIPPER ON YOUR SKIN, IN YOUR EYES OR NEAR A FLAME. DO NOT BREATHE THE FUMES. ONLY USE IN A LOCATION THAT HAS A CONTINUOUS FLOW OF CLEAN AIR. PAINT STRIPPER IS POISONOUS AND FLAMMABLE.":
        "ПРЕДУПРЕЖДЕНИЕ: НЕ ДОПУСКАЙТЕ ПОПАДАНИЯ СМЫВКИ КРАСКИ НА КОЖУ, В ГЛАЗА ИЛИ ВБЛИЗИ ОТКРЫТОГО ОГНЯ. НЕ ВДЫХАЙТЕ ПАРЫ. РАБОТАЙТЕ ТОЛЬКО В ХОРОШО ПРОВЕТРИВАЕМОМ ПОМЕЩЕНИИ. СМЫВКА КРАСКИ ЯДОВИТА И ОГНЕОПАСНА.",
    # ─── WARNING about hydraulic fluid ───
    "WARNING: DO NOT GET HYDRAULIC FLUID ON YOUR SKIN OR IN YOUR EYES. DO NOT BREATHE THE FUMES. ONLY USE IN A LOCATION THAT HAS A CONTINUOUS FLOW OF CLEAN AIR. HYDRAULIC FLUID IS POISONOUS AND DANGEROUS.":
        "ПРЕДУПРЕЖДЕНИЕ: НЕ ДОПУСКАЙТЕ ПОПАДАНИЯ ГИДРАВЛИЧЕСКОЙ ЖИДКОСТИ НА КОЖУ ИЛИ В ГЛАЗА. НЕ ВДЫХАЙТЕ ПАРЫ. РАБОТАЙТЕ ТОЛЬКО В ХОРОШО ПРОВЕТРИВАЕМОМ ПОМЕЩЕНИИ. ГИДРАВЛИЧЕСКАЯ ЖИДКОСТЬ ЯДОВИТА И ОПАСНА.",
    # ─── Trailing "and" sentences (line-broken procedural text) ───
    "Release the tab washers (2-180 and 2-200). Remove the bolts (2-170 and 2-190) and":
        "Ослабьте стопорные шайбы (2-180 и 2-200). Снимите болты (2-170 и 2-190) и",
    "Remove the split pin (3-180), the nut (3-190), the washer (3-200), the bolt (3-210) and":
        "Снимите шплинт (3-180), гайку (3-190), шайбу (3-200), болт (3-210) и",
    "Remove the split pin (4-220), the nut (4-230), the washer (4-240), the bolt (4-250) and":
        "Снимите шплинт (4-220), гайку (4-230), шайбу (4-240), болт (4-250) и",
    "Remove the split pin (5-170), the nut (5-180), the washer (5-190), the bolt (5-200) and":
        "Снимите шплинт (5-170), гайку (5-180), шайбу (5-190), болт (5-200) и",
    "Remove the split pins (8-40), the nuts (8-50), the washers (8-60), the bolts (8-70) and":
        "Снимите шплинты (8-40), гайки (8-50), шайбы (8-60), болты (8-70) и",
    "Remove the split pin (10-40), the nut (10-50), the washers (10-60), the bolt (10-70) and":
        "Снимите шплинт (10-40), гайку (10-50), шайбы (10-60), болт (10-70) и",
    # ─── Sentences with "the lower torque link" continuation ───
    "Remove the pin (11-130), the harness support bracket (11-140), the lower torque link":
        "Снимите штифт (11-130), кронштейн крепления жгута (11-140), нижний шлиц-шарнир",
    # ─── Sentences with tool names ───
    "Use the Torque Reaction Adapter 460007242 to hold the pin (9-70) and use the Torque Adapter T14544 to remove the nut (9-50). Remove the spacer (9-60), the pin (9-70), the spacer (9-80) and the sleeve (9-90).":
        "Используйте Torque Reaction Adapter 460007242 для удержания штифта (9-70) и используйте Torque Adapter T14544 для снятия гайки (9-50). Снимите проставку (9-60), штифт (9-70), проставку (9-80) и втулку (9-90).",
    "Use the Torque Adapter 460007283, the Torque Reactor 460007278, the Holding Blocks 460006406 and the Bench Clamp MT1025 to remove the diaphragm subassembly (15-190), the compression orifice plate (15-220), the clapper seat (15-230) and the baffle (15-240).":
        "Используйте Torque Adapter 460007283, Torque Reactor 460007278, Holding Blocks 460006406 и Bench Clamp MT1025 для снятия сборки диафрагмы (15-190), пластины отверстия сжатия (15-220), седла хлопушки (15-230) и дефлектора (15-240).",
    "Release the lock washer (17-90) and use the Torque Adapter 460006404 to remove the jacking dome (17-80). Remove the lock washer (17-90).":
        "Ослабьте стопорную шайбу (17-90) и используйте Torque Adapter 460006404 для снятия домкратной точки (17-80). Снимите стопорную шайбу (17-90).",
    "Hold the cylinder (17-230) in the Bench Clamp MT1025 and Holding Blocks MT1026/63.":
        "Удерживайте цилиндр (17-230) в Bench Clamp MT1025 и Holding Blocks MT1026/63.",
    "Remove the cylinder (17-230) from the Bench Clamp MT1025 and Holding Blocks MT1026/63.":
        "Снимите цилиндр (17-230) из Bench Clamp MT1025 и Holding Blocks MT1026/63.",
    "Use the Torque Adapter 460007232 to remove the locking nut (19-52). Remove the locking washer (19-54) and the outer race and the ball of the spherical bearing (19-50).":
        "Используйте Torque Adapter 460007232 для снятия контргайки (19-52). Снимите стопорную шайбу (19-54) и наружное кольцо и шарик сферического подшипника (19-50).",
    "Use the Hydraulic-Pneumatic Pump Set 460006497, the Bolt 460006498/7, the Press Pad 460006499/25 and the Extraction Tube 460004680 and remove the forward pintle bush (20-250A).":
        "Используйте Hydraulic-Pneumatic Pump Set 460006497, Bolt 460006498/7, Press Pad 460006499/25 и Extraction Tube 460004680 и снимите переднюю втулку навеса (20-250A).",
    "Use the Extraction Pad 460006263/460006232 and the Extraction Bar 460006262 to remove the bearing (20-300).":
        "Используйте Extraction Pad 460006263/460006232 и Extraction Bar 460006262 для снятия подшипника (20-300).",
    "Use the Extraction Pad 460006261 and the Extraction Bar 460006262 to remove the bearing (20-310).":
        "Используйте Extraction Pad 460006261 и Extraction Bar 460006262 для снятия подшипника (20-310).",
    "Use the Press Pad Assembly 460006267 and remove the drag arm sleeve (20-370A only).":
        "Используйте Press Pad Assembly 460006267 и снимите втулку тяги (20-370A только).",
    # ─── Nitrogen pressure sentences (split across lines) ───
    "Slowly increase the nitrogen pressure to between 9,32 and 10,68 bar (135 and":
        "Медленно увеличьте давление азота до 9,32–10,68 бар (135 и",
    "Slowly increase the nitrogen pressure to between 67,59 and 70,34 bar (980 and1020 lbf/in2).":
        "Медленно увеличьте давление азота до 67,59–70,34 бар (980 и 1020 фунт/дюйм²).",
    # ─── Weight lines (full tab-separated with "approximately") ───
    "Weight with hydraulic fluid\t522 kg (1151 lb) approximately":
        "Масса с гидравлической жидкостью\t522 кг (1151 фунт) приблизительно",
    "Weight without hydraulic fluid\t505 kg (1113 lb) approximately":
        "Масса без гидравлической жидкости\t505 кг (1113 фунт) приблизительно",
    # ─── Torque adapter sentences ───
    "Use the Torque Adapter 460007230 to remove the nut (14-60).":
        "Используйте Torque Adapter 460007230 для снятия гайки (14-60).",
    "Use the Torque Adapter 460006404 to remove the nut (17-20) and the valve stem (17-10). Remove the valve stem (17-10).":
        "Используйте Torque Adapter 460006404 для снятия гайки (17-20) и штока клапана (17-10). Снимите шток клапана (17-10).",
    "Use the Torque Adapter 460006404 and the loading press to connect the valve stem (17-10) to the charging valve (13-60).":
        "Используйте Torque Adapter 460006404 и нагрузочный пресс для соединения штока клапана (17-10) с зарядным клапаном (13-60).",
    "Use the Torque Adapter 460006404 to connect the valve stem (17-20) to the charging valve (13-60). Release the valve stem (17-20).":
        "Используйте Torque Adapter 460006404 для соединения штока клапана (17-20) с зарядным клапаном (13-60). Отпустите шток клапана (17-20).",
    # ─── Sentences with stray "to" ───
    "Use the Charging Adapter 460002502 and the Turner Inflation Equipment T14218: connect the charging valve (17-20) to the nitrogen supply. Open the charging valve (17-20).":
        "Используйте адаптер для зарядки 460002502 и оборудование для заправки Turner T14218: подсоедините зарядный клапан (17-20) к источнику азота. Откройте зарядный клапан (17-20).",
    "Use the Charging Adapter 460002502 to connect the hydraulic test rig to the charging valve (13-60).":
        "Используйте адаптер для зарядки 460002502 для подсоединения стенда для гидравлических испытаний к зарядному клапану (13-60).",
    "Use the Charging Adapter 460002502 and the Turner Inflation Equipment T14218 to connect the nitrogen supply to the charging valve (13-60).":
        "Используйте адаптер для зарядки 460002502 и оборудование для заправки Turner T14218 для подсоединения источника азота к зарядному клапану (13-60).",
    "Use the Charging Adapter 460002502 and the Turner Inflation Equipment T14218 to connect the charging valve (17-20) to the nitrogen supply.":
        "Используйте адаптер для зарядки 460002502 и оборудование для заправки Turner T14218 для подсоединения зарядного клапана (17-20) к источнику азота.",
    "Slowly increase the nitrogen pressure to between 13,11 and 14,48 bar (190 and 210 lbf/in2).":
        "Медленно увеличьте давление азота до 13,11–14,48 бар (190 и 210 фунт/дюйм²).",
    "Close the charging valve (17-20); use the Crowfoot Wrench T14500 to torque it to between 5,7 and 7,9 N m (50 and 70 lbf in).":
        "Закройте зарядный клапан (17-20); используйте рожковый ключ T14500 для затяжки моментом 5,7–7,9 Н·м (50–70 фунт·дюйм).",
    "Use the Turner Inflation Equipment T14218 and the Charging Adapter 460002502 to connect the nitrogen supply to the charging valve (13-60).":
        "Используйте оборудование для заправки Turner T14218 и адаптер для зарядки 460002502 для подсоединения источника азота к зарядному клапану (13-60).",
    "Slowly increase the nitrogen pressure to between 6,90 and 8,27 bar (100 and 120 lbf/in2).":
        "Медленно увеличьте давление азота до 6,90–8,27 бар (100 и 120 фунт/дюйм²).",
    "Close the charging valve (13-60); use the Crowfoot Wrench T14500 to torque it to between 5,7 and 7,9 N m (50 and 70 lbf in).":
        "Закройте зарядный клапан (13-60); используйте рожковый ключ T14500 для затяжки моментом 5,7–7,9 Н·м (50–70 фунт·дюйм).",
    "Open the charging valve (13-60) and reduce the nitrogen pressure to between 3,45 and 4,82 bar (50 and 70 lbf/in2).":
        "Откройте зарядный клапан (13-60) и уменьшите давление азота до 3,45–4,82 бар (50 и 70 фунт/дюйм²).",
    "Open the charging valve (17-20) and reduce the nitrogen pressure to between 3,45 and 4,82 bar (50 and 70 lbf/in2).":
        "Откройте зарядный клапан (17-20) и уменьшите давление азота до 3,45–4,82 бар (50 и 70 фунт/дюйм²).",
    "Close the charging valve (17-20); use the Crowfoot Wrench T14500 to torque it to between 5,7 and 7,9 N m (50 and 70 lbf in).":
        "Закройте зарядный клапан (17-20); используйте рожковый ключ T14500 для затяжки моментом 5,7–7,9 Н·м (50–70 фунт·дюйм).",
}

# ── Table header translations ────────────────────────────────────────────────
TABLE_HEADERS = {
    "Subject Reference": "Ссылка на раздел",
    "Remove and Destroy Pages": "Удалить и уничтожить страницы",
    "Insert New/Revised": "Вставить новые/пересмотренные",
    "Pages": "Страницы",
    "Dated": "Дата",
    "Reason for Change": "Причина изменения",
    "REV.\nNo.": "НОМ.\nРЕД.",
    "ISSUE DATE": "ДАТА ВЫПУСКА",
    "DATE INSERTED": "ДАТА ВСТАВКИ",
    "PAGE NUMBER": "НОМЕР СТРАНИЦЫ",
    "BY": "КЕМ",
    "DATE REMOVED": "ДАТА УДАЛЕНИЯ",
    "SB NUMBER": "НОМЕР SB",
    "SB TITLE": "НАЗВАНИЕ SB",
    "SB REVISION NUMBER": "НОМЕР РЕДАКЦИИ SB",
    "DATE INCORPORATED INTO MANUAL": "ДАТА ВКЛЮЧЕНИЯ В РУКОВОДСТВО",
    "COVER SB NO.": "НОМЕР ОХВАТ. SB",
    "PART NUMBER": "НОМЕР ДЕТАЛИ",
    "DASH NO.": "НОМЕР ТИРЕ",
    "MOD. STRIKE NO.": "НОМЕР СНЯТИЯ МОД.",
    "SAFRAN LANDING SYSTEMS MODIFICATION NUMBER": "НОМЕР МОДИФИКАЦИИ SAFRAN LANDING SYSTEMS",
    "SAFRAN LANDING SYSTEMS\nSERVICE BULLETIN NUMBER": "НОМЕР СЕРВИСНОГО БЮЛЛЕТЕНЯ\nSAFRAN LANDING SYSTEMS",
    "SAFRAN LANDING SYSTEMS SERVICE BULLETIN NUMBER": "НОМЕР СЕРВИСНОГО БЮЛЛЕТЕНЯ SAFRAN LANDING SYSTEMS",
    "Page": "Страница",
    "Date": "Дата",
    "Subject": "Раздел",
    "Initial Issue": "Первоначальный выпуск",
    "No effect": "Без изменений",
    # Part 2+ table headers
    "Part No.": "№ детали",
    "Equipment": "Оборудование",
    "Function": "Функция",
    "Special Tool": "Специальный инструмент",
    "Ref. Item": "Поз. ссылки",
    "Material": "Материал",
    "TEST POINT": "ТОЧКА ПРОВЕРКИ",
    "IPL FIGURE AND ITEM No.": "№ РИСУНКА И ПОЗИЦИИ ИПД",
    "NAME": "НАИМЕНОВАНИЕ",
    "LIMIT VALUE MILLIOHMS": "ПРЕДЕЛЬНОЕ ЗНАЧЕНИЕ, МИЛЛИОМЫ",
    "IPL Fig/Item": "Поз. рис. ИПД",
    "Part Name": "Наименование детали",
    "Fig Item No.": "№ поз. рис.",
    "Fig Item": "Поз. рис.",
    "Name": "Наименование",
    "Material Type": "Тип материала",
    "Spec": "Спецификация",
    # Part 3 (Repair) table headers
    "Item No.": "№ позиции",
    "Specification": "Спецификация",
    # Part 5 (Repair) table headers
    "Material Specification": "Спецификация материала",
    "Repair Part": "Ремонтная деталь",
    "Repair liner": "Ремонтный вкладыш",
    "Oversize bearing": "Ремонтный подшипник",
    "Cleaning tissues": "Чистящие салфетки",
    "Emery cloth, 60-100 grit": "Наждачная ткань, зернистость 60-100",
    "Emery cloth": "Наждачная ткань",
    "Cleaning agent": "Чистящее средство",
    "Masking tape": "Малярная лента",
    "Adhesive PVC tape": "Клейкая ПВХ-лента",
    "Adhesive": "Адгезив",
    "Install the oversize bearing": "Установить ремонтный подшипник",
    "Adhesive (Loctite Grade 601)": "Адгезив (Loctite Grade 601)",
    "Accomet C": "Accomet C",
    "Araldite, 2015": "Araldite, 2015",
    "Alocrom": "Alocrom",
    "Fibreslip B40": "Fibreslip B40",
    "Fibreslip, B40": "Fibreslip, B40",
    "Sealant": "Герметик",
}

# ── Revision table section name translations ─────────────────────────────────
SECTION_NAMES = {
    "Record of Revisions": "Запись изменений",
    "Unit Identification Chart": "Таблица идентификации изделия",
    "List of Effective Pages": "Перечень действующих страниц",
    "Table of Contents": "Содержание",
    "Disassembly": "Разборка",
    "Check": "Проверка",
    "Repair": "Ремонт",
    "Assembly (Including Storage)": "Сборка (включая хранение)",
    "Fits and Clearances": "Посадки и зазоры",
    "Special Tools, Fixtures and Equipment": "Специальные инструменты, приспособления и оборудование",
    "Illustrated Parts List": "Иллюстрированный перечень деталей",
    # List of Effective Pages section names
    "Subject": "Раздел",
    "(Continued)": "(Продолжение)",
    "Testing and Fault": "Проверка и поиск неисправностей",
    "Isolation": "Локализация",
    "Isolation (Continued)": "Локализация (Продолжение)",
    "Cleaning": "Очистка",
    "Repair (Continued)": "Ремонт (Продолжение)",
    "Assembly (Including": "Сборка (включая",
    "Storage)": "хранение)",
    "Storage) (Continued)": "хранение) (Продолжение)",
    "Special Tools, Fixtures": "Специальные инструменты, приспособления",
    "and Equipment": "и оборудование",
}

# ── Reason for change phrase translations ────────────────────────────────────
REASON_PHRASES = {
    "Updated revision status": "Обновлён статус редакции",
    "Updated pages": "Обновлены страницы",
    "Added content": "Добавлено содержание",
    "Updated page numbers": "Обновлены номера страниц",
    "Updated figures": "Обновлены рисунки",
    "Updated figure numbers": "Обновлены номера рисунков",
    "Updated tables": "Обновлены таблицы",
    "Added paras": "Добавлены пункты",
    "Added para": "Добавлен пункт",
    "Updated fig-items": "Обновлены поз. рисунков",
    "Updated fig-item": "Обновлена поз. рисунка",
    "Added fig-item": "Добавлена поз. рисунка",
    "in para": "в пункте",
    "in paras": "в пунктах",
    "and": "и",
    "only": "только",
    "only in para": "только в пункте",
    "in": "в",
    "Updated Messier-Dowty Limited to Safran Landing Systems": "Обновлено Messier-Dowty Limited на Safran Landing Systems",
    "Updated Messier-Dowty Limited to Safran Landing System": "Обновлено Messier-Dowty Limited на Safran Landing Systems",
    "Updated Messier-Dowty Limited to Safran Landing Sy": "Обновлено Messier-Dowty Limited на Safran Landing Sy",
    "Updated Messier-Dowty": "Обновлено Messier-Dowty",
    "Limited to Safran Landing Systems": "Limited на Safran Landing Systems",
    "Added repair no": "Добавлен ремонт №",
    "Added repair no.": "Добавлен ремонт №",
    "Added caution at para": "Добавлено предупреждение в пункте",
    "Updated material specification": "Обновлена спецификация материала",
    "Updated materia": "Обновлены материа",
    "Updated paras": "Обновлены пункты",
    "Updated IPL figs": "Обновлены рисунки ИПД",
    "to include Ref": "для включения ссылок",
    "Updated figure numbers": "Обновлены номера рисунков",
    "Updated figure": "Обновлён рисунок",
    "Updated figures": "Обновлены рисунки",
    "Added figure": "Добавлен рисунок",
    "Added Ref. Codes": "Добавлены коды ссылок",
    "Updated tables": "Обновлены таблицы",
    "Updated table": "Обновлена таблица",
    # Standalone "Updated" at end of line (continuation to next paragraph)
    ". Updated": ". Обновлены",
    "Deleted": "Удалены",
    "titles": "заголовки",
    "conversion value": "значение пересчёта",
    "caution at para": "предупреждение в пункте",
    "caution": "предупреждение",
    "Updated IPL fig": "Обновлён рисунок ИПД",
    "figure": "рисунок",
    "figures": "рисунки",
    "Codes": "Коды",
    "details": "детали",
    "to": "до",
}

# ── SB title translations ───────────────────────────────────────────────────
SB_TITLE_PARTS = {
    # Full SB title translations (sorted longest-first for matching)
    "MLG - Installation of stub bolt subassembly for the forward pintle pin in place of the cross bolt.":
        "ОШ — Установка сборки болта-вставки для переднего штифта навеса стойки взамен болта.",
    "MLG - To allow an increase in aircraft maximum take-off weight to 93 tonne.":
        "ОШ — Для обеспечения увеличения максимальной взлётной массы воздушного судна до 93 тонн.",
    "MLG -To add tracking numbers to parts listed in Airbus Airworthiness Limitations Section (ALS).":
        "ОШ — Добавление учётных номеров к деталям, указанным в разделе ограничений лётной годности Airbus (ALS).",
    "MLG - Installation of a 201585 series MLG Leg and Dressings where a 201387 MLG Leg and Dressings has been installed":
        "ОШ — Установка стойки ОШ серии 201585 и обвязки вместо установленной стойки ОШ серии 201387 и обвязки",
    "MLG -To add tracking numbers to parts listed in Airbus Maintenance Planning Document, Section 9-1. (Torque link apex pin nut)":
        "ОШ — Добавление учётных номеров к деталям, указанным в документе планирования ТО Airbus, раздел 9-1. (Гайка штифта вершины шлиц-шарнира)",
    "MLG -To add tracking numbers to parts listed in Airbus Maintenance Planning Document, Section 9-":
        "ОШ — Добавление учётных номеров к деталям, указанным в документе планирования ТО Airbus, раздел 9-",
    "MLG - Introduction of a new lower bearing subassembly":
        "ОШ — Введение новой сборки нижнего подшипника",
    "MLG - Introduction of new charging labels":
        "ОШ — Введение новых этикеток зарядки",
    "MLG - Introduction of new 1M and 2M Axle harnesses":
        "ОШ — Введение новых жгутов осей 1М и 2М",
    "MLG - Introduction of new 1M and 2M Leg Harness and of new 1M and 2M Axle Harnesses":
        "ОШ — Введение новых жгутов стойки 1М и 2М и новых жгутов осей 1М и 2М",
    "MLG - Introduction of new 1M and 2M Leg Harness an":
        "ОШ — Введение новых жгутов стойки 1М и 2М и",
    "MLG Leg-Introduction of new retaining pins and a new lower bearing subassembly with a new self lubricating liner":
        "Стойка ОШ — Введение новых стопорных штифтов и новой сборки нижнего подшипника с новым самосмазывающимся вкладышем",
    "MLG Leg - Introduction of new retaining pins for the lower bearing subassembly":
        "Стойка ОШ — Введение новых стопорных штифтов для сборки нижнего подшипника",
    "MLG Leg - Introduction of a new lower bearing subassembly with a new low friction inner liner":
        "Стойка ОШ — Введение новой сборки нижнего подшипника с новым низкофрикционным внутренним вкладышем",
    "MLG Leg - Barkhausen Noise Inspection of Main Landing Gear Sliding Tube Axles.":
        "Стойка ОШ — Контроль методом шума Баркгаузена осей скользящей трубы основного шасси.",
    "MLG Leg - Introduction of a new Main Fitting":
        "Стойка ОШ — Введение нового корпуса стойки",
    "MLG Leg - Introduction of a new torque link damper unit":
        "Стойка ОШ — Введение нового демпфера шлиц-шарнира",
    "MLG Leg - Introduction of a new torque link damper":
        "Стойка ОШ — Введение нового демпфера шлиц-шарнира",
    "MLG Leg - Introduction of a new main fitting subassembly and related parts":
        "Стойка ОШ — Введение новой сборки корпуса стойки и связанных деталей",
    "MLG Leg - Introduction of a new main fitting subas":
        "Стойка ОШ — Введение новой сборки корпуса стойки",
    "MLG - Introduction of a new upper pivot bracket":
        "ОШ — Введение нового верхнего поворотного кронштейна",
    "MLG - Introduction of a new changeover valve stem and housing":
        "ОШ — Введение нового штока и корпуса переключающего клапана",
    "MLG - Introduction of a new changeover valve stem ":
        "ОШ — Введение нового штока переключающего клапана",
    "MLG Complete - Modification of the transfer block subassembly":
        "ОШ в сборе — Модификация сборки переходного блока",
    "MLG Complete - Modification of the transfer block ":
        "ОШ в сборе — Модификация переходного блока",
    "MLG - Conversion of low - friction lower - bearing MLG to standard lower - bearing MLG":
        "ОШ — Замена низкофрикционного ОШ с нижним подшипником на стандартный ОШ с нижним подшипником",
    "MLG - Conversion of low - friction lower - bearing":
        "ОШ — Замена низкофрикционного нижнего подшипника",
    "MLG complete - Introduction of a new transfer block subassembly":
        "ОШ в сборе — Введение новой сборки переходного блока",
    "MLG complete - Introduction of a new transfer bloc":
        "ОШ в сборе — Введение нового переходного блока",
}

# ── Part names for table cells (singular → translation) ─────────────────────
PART_NAMES_TABLE = {
    "Split pin": "Шплинт",
    "Split pins": "Шплинты",
    "Tab washer": "Стопорная шайба",
    "Tab washers": "Стопорные шайбы",
    "O-ring seal": "Уплотнительное кольцо",
    "O-ring seals": "Уплотнительные кольца",
    "Backing ring": "Опорное кольцо",
    "Backing rings": "Опорные кольца",
    "Seal": "Уплотнение",
    "Joint seal": "Соединительное уплотнение",
    "Sealing ring": "Уплотнительное кольцо",
    "Wiper ring": "Грязесъёмное кольцо",
    "Inner liner": "Внутренний вкладыш",
    "Lock washer": "Стопорная шайба",
    "Locking plate": "Стопорная пластина",
    "Screws": "Винты",
    "Heat shrink sleeve": "Термоусадочная трубка",
    "Ferrule": "Наконечник",
    "Bowden cable": "Трос в оболочке",
    "Bolt": "Болт",
    "Cross Bolt": "Поперечный болт",
    "Special Bolt": "Специальный болт",
    "Threaded insert": "Резьбовая вставка",
    "Retaining pin": "Стопорный штифт",
    "Inflation valve": "Клапан зарядки",
    "Locking pin": "Стопорный штифт",
    "Retaining ring": "Стопорное кольцо",
    "Recoil orifice plate": "Пластина отверстия обратного хода",
    "Lock plate": "Стопорная пластина",
    "Compression orifice plate": "Пластина отверстия сжатия",
    "Clapper seat": "Седло клапана",
    "Level tube": "Трубка уровня",
    "Upper diaphragm tube": "Верхняя диафрагменная труба",
    "Upper dia- phragm tube": "Верхняя диафрагменная труба",
    "Jacking dome": "Домкратная точка",
    "Washer": "Шайба",
    "Valve stem": "Шток клапана",
    "Sleeve": "Втулка",
    "Upper torque link": "Верхний шлиц-шарнир",
    "Lower torque link": "Нижний шлиц-шарнир",
    "Slave link": "Ведомое звено",
    "Lower slave link": "Нижнее ведомое звено",
    "Harness support bracket": "Кронштейн крепления жгута",
    "Pivot bracket": "Поворотный кронштейн",
    "Lock stay cardan": "Кардан фиксатора",
    "Uplock pin": "Штифт замка убранного положения",
    "Static dis- charge con- nector": "Штыревой разъём статического разряда",
    "Main fitting": "Корпус стойки",
    "Proximity switch connector shell": "Корпус разъёма датчика приближения",
    "Static discharge connector": "Штыревой разъём статического разряда",
    "Nut": "Гайка",
    "Rod end": "Наконечник тяги",
    "Main landing gear leg (1-1) tests": "Испытания стойки основного шасси (1-1)",
    "Electrical bonding resistance tests": "Проверка сопротивления электрического соединения",
    "Proximity switch and target tests": "Испытания датчика приближения и мишени",
    # Part 3 (Repair) Table 601 part names
    "Cross bolt": "Поперечный болт",
    "Cross Bolt": "Поперечный болт",
    "Upper torque link subassembly": "Сборка верхнего шлиц-шарнира",
    "Lower torque link subassembly": "Сборка нижнего шлиц-шарнира",
    "Harness support": "Крепление жгута",
    "Stop ring": "Стопорное кольцо",
    "Upper diaphragm tube sub-assembly": "Сборка верхней диафрагменной трубы",
    "Locking nut": "Контргайка",
    "Locking washer": "Стопорная шайба",
    "Main fitting subassembly": "Сборка корпуса стойки",
}

# ── Tool names for table cells ──────────────────────────────────────────────
TOOL_NAMES_TABLE = {
    "Hydraulic Test Rig": "Стенд для гидравлических испытаний",
    "Nitrogen Supply": "Источник азота",
    "Loading Press": "Нагрузочный пресс",
    "28 VDC Power Supply": "Источник питания 28 В постоянного тока",
    "Turner Inflation Equipment": "Оборудование для заправки Turner",
    "Crowfoot Wrench": "Рожковый ключ",
    "Charging Adapter": "Адаптер для зарядки",
    "Lampbox": "Контрольная лампа",
    "Holding Fixture": "Удерживающее приспособление",
    "Load Cell and Adapter": "Датчик нагрузки и адаптер",
    "Press Adapter": "Адаптер пресса",
    "Offset Adapter": "Смещённый адаптер",
    "Bottom Press Adapter": "Нижний адаптер пресса",
    "Bench Clamp": "Настольный зажим",
    "Holding Blocks": "Удерживающие блоки",
    "Torque Adapter": "Моментный адаптер",
    "Extractor": "Экстрактор",
    "Drift": "Выколотка",
    "Extraction Tube": "Извлекающая труба",
    "Lifting Bar Assembly": "Сборка подъёмной штанги",
    "Lifting Tackle": "Подъёмное приспособление",
    "Extraction Pad": "Извлекающая пластина",
    "Extraction Bar": "Извлекающая штанга",
    "Press Pad Assembly": "Сборка прижимной пластины",
    "Assembly/Extraction Tool": "Инструмент для сборки/извлечения",
    "Extractor Pad and Drawbolt": "Извлекающая пластина и вытяжной болт",
    "Hydraulic-Pneumatic Pump Set": "Гидропневматический насосный комплект",
    "Press Pad": "Прижимная пластина",
    "Build Trolley": "Сборочная тележка",
    "Pin Spanner": "Штифтовый ключ",
    "Torque Reaction Adapter": "Адаптер реакции момента",
    "Torque Reaction": "Реакция момента",
    "Torque Reactor": "Реактор момента",
    "Extractor Plate": "Извлекающая пластина",
    "Pintle Location Assembly": "Сборка позиционирования навеса",
    "Spherical Bearing Locator": "Позиционер сферического подшипника",
    "Location Frame": "Установочная рама",
    "Adapter": "Адаптер",
    "Transport and Build": "Транспортировочная и",
    "Trolley Support Arms Towing Frame": "Тележка, опорные рычаги, буксировочная рама",
    "Jacking Dome Adapter": "Адаптер домкратной точки",
    "Milliohmmeter Megger, Type BT51": "Миллиомметр Megger, тип BT51",
}

# ── Procedural vocabulary for word-level translation ────────────────────────
# Order: longest phrases first, then shorter ones.
# These are applied sequentially via str.replace (no regex).
PROCEDURAL_VOCAB = [
    # ══════════════════════════════════════════════════════════════════════════
    # FULL SENTENCE TRANSLATIONS (must come FIRST before word-level replacements
    # so that "in the", "and", "or", etc. don't break sentence matching)
    # ══════════════════════════════════════════════════════════════════════════

    # ─── Introduction section sentences ───
    ("This manual contains Description, Operation, Maintenance procedures and an Illustrated Parts List (IPL). IPL Figure and Item numbers in parentheses follow the part name to identify them.",
     "Настоящее руководство содержит описание, работу, процедуры технического обслуживания и иллюстрированный перечень деталей (ИПД). Номера рисунков и позиций ИПД в скобках следуют за наименованием детали для их идентификации."),
    ("A Unit Identification Chart is included to show the modification status of the unit. The modification status is related to the unit part number by the last two digits of the dash number.",
     "Таблица идентификации изделия включена для отображения статуса модификации изделия. Статус модификации связан с номером детали изделия двумя последними цифрами номера тире."),
    ("All references in this manual are to the left configuration of the unit unless the instructions tell you differently.",
     "Все ссылки в настоящем руководстве относятся к левой конфигурации изделия, если инструкции не указывают иное."),
    ("All dimensions and quantities in this manual are in SI units with Imperial units in parentheses. A comma shows a decimal part of an SI unit. A full point shows a decimal part of an Imperial unit.",
     "Все размеры и величины в настоящем руководстве указаны в единицах СИ с имперскими единицами в скобках. Запятая обозначает десятичную часть единицы СИ. Точка обозначает десятичную часть имперской единицы."),
    ("This manual refers to Process Specifications (M-DLPS and PCS) and Non-destructive Tests (M-DLNDT). These documents are available from Safran Landing Systems.",
     "В настоящем руководстве ссылаются на спецификации процессов (M-DLPS и PCS) и неразрушающие испытания (M-DLNDT). Эти документы доступны в Safran Landing Systems."),
    ("This manual refers to Process Specifications (M-DLPS and PCS) and Non-destructive Tests (M-DLNDT). These are available within the Safran Landing Systems Technical Publications on-line service.",
     "В настоящем руководстве ссылаются на спецификации процессов (M-DLPS и PCS) и неразрушающие испытания (M-DLNDT). Они доступны через онлайн-службу технических публикаций Safran Landing Systems."),
    ("All the materials in this manual have a Ref. Item identification. This is the reference item number of the material in the Aircraft Manufacturer\u2019s Consumable Materials List.",
     "Все материалы в настоящем руководстве имеют идентификатор позиции ссылки. Это номер позиции ссылки материала в списке расходных материалов производителя воздушного судна."),
    ("All the materials in this manual have a Ref. Item identification. This is the reference item number of the material in the Aircraft Manufacturer's Consumable Materials List.",
     "Все материалы в настоящем руководстве имеют идентификатор позиции ссылки. Это номер позиции ссылки материала в списке расходных материалов производителя воздушного судна."),
    ("The repairs in this CMM have been approved under Airbus\u2019 EASA Design Organisation Approval No. EASA.21J.031.",
     "Ремонты в настоящем руководстве утверждены в рамках одобрения проектной организации Airbus EASA № EASA.21J.031."),
    ("The repairs in this CMM have been approved under Airbus' EASA Design Organisation Approval No. EASA.21J.031.",
     "Ремонты в настоящем руководстве утверждены в рамках одобрения проектной организации Airbus EASA № EASA.21J.031."),
    ("On occasion a REF. CODE can be identified in the NOMENCLATURE column in the DETAILED PARTS LIST. This is a Safran Landing Systems reference code and is used for cross-reference purposes only.",
     "Иногда в столбце НАИМЕНОВАНИЕ ПОДРОБНОГО СПИСКА ДЕТАЛЕЙ может быть указан КОД ССЫЛКИ. Это ссылочный код Safran Landing Systems, используемый только для перекрёстных ссылок."),
    ("On occasion a REF. CODE can be identified in the NOMENCLATURE column of the DETAILED PARTS LIST. This is a Safran Landing Systems drawing reference.",
     "Иногда в столбце НАИМЕНОВАНИЕ ПОДРОБНОГО СПИСКА ДЕТАЛЕЙ может быть указан КОД ССЫЛКИ. Это ссылка на чертёж Safran Landing Systems."),
    # Part 2 variant of P79
    ("A Unit Identification Chart is included to show the modification status of the unit. The modification status is related to the unit part number by the dash number: the dash number is marked on the unit name plate adjacent to the part number.",
     "Таблица идентификации изделия включена для отображения статуса модификации изделия. Статус модификации связан с номером детали изделия номером тире: номер тире указан на табличке изделия рядом с номером детали."),
    # P87 shop verification
    ("The accuracy and the adequacy of the instructions in this CMM have been technically verified by shop verification (performed or simulated) or by similarity with manufacturing instructions or with component maintenance manuals instructions from other programs that have been verified in shop.",
     "Точность и достаточность инструкций настоящего руководства по техническому обслуживанию компонентов технически подтверждены верификацией в цеху (выполненной или смоделированной) или по аналогии с производственными инструкциями или с инструкциями руководств по техническому обслуживанию компонентов других программ, прошедших верификацию в цеху."),

    # ─── Description and Operation sentences ───
    ("The main landing gear leg has a sliding tube subassembly that operates in a main fitting subassembly. The sliding tube subassembly operates through a lower bearing subassembly in the main fitting subassembly.",
     "Стойка основного шасси имеет сборку скользящей трубы, работающую в сборке корпуса стойки. Сборка скользящей трубы перемещается через сборку нижнего подшипника в сборке корпуса стойки."),
    ("An upper torque link subassembly attaches to the main fitting subassembly. A lower torque link subassembly attaches to the sliding tube subassembly. A damper connects the upper and lower torque link subassemblies.",
     "Сборка верхнего шлиц-шарнира крепится к сборке корпуса стойки. Сборка нижнего шлиц-шарнира крепится к сборке скользящей трубы. Демпфер соединяет сборки верхнего и нижнего шлиц-шарниров."),
    ("A slave link subassembly and a lower slave link subassembly attach opposite the upper and lower torque link subassemblies.",
     "Сборка ведомого звена и сборка нижнего ведомого звена крепятся напротив сборок верхнего и нижнего шлиц-шарниров."),
    ("The Upper and Lower Torque Link Subassemblies",
     "Сборки верхнего и нижнего шлиц-шарниров"),
    ("The upper and lower torque link subassemblies prevent the sliding tube subassembly from turning in the main fitting subassembly.",
     "Сборки верхнего и нижнего шлиц-шарниров предотвращают поворот сборки скользящей трубы в сборке корпуса стойки."),
    ("The damper controls the movement of the upper and lower torque link subassemblies.",
     "Демпфер контролирует перемещение сборок верхнего и нижнего шлиц-шарниров."),
    ("Examine the thread form of the diaphragm subassembly (15-190) and diaphragm (15-210A) with shadow graph projection.",
     "Осмотрите профиль резьбы сборки диафрагмы (15-190) и диафрагмы (15-210A) с помощью проекции профилографа."),
    ("Measure all the parts that are in FITS AND CLEARANCES and compare with the dimensions in the table.",
     "Измерьте все детали, указанные в ПОСАДКАХ И ЗАЗОРАХ, и сравните с размерами в таблице."),
    ("Where K = 273", "Где K = 273"),
    ("Grease Groove Dimensions After Installation", "Размеры канавки для смазки после установки"),
    ("Gland Housing", "корпус сальника"),
    ("Lower Bearing Subassembly Machining and Liner Installation",
     "Механическая обработка и установка вкладыша сборки нижнего подшипника"),
    ("Transfer Block Subassembly", "Сборка переходного блока"),
    # Part 2 variant of P128
    ("The main landing gear leg has a sliding tube subassembly that operates in a main fitting subassembly. The sliding tube subassembly operates through a lower bearing subassembly. The lower bearing subassembly also seals the sliding tube subassembly in the main fitting subassembly.",
     "Стойка основного шасси имеет сборку скользящей трубы, работающую в сборке корпуса стойки. Сборка скользящей трубы перемещается через сборку нижнего подшипника. Сборка нижнего подшипника также герметизирует сборку скользящей трубы в сборке корпуса стойки."),
    # Part 2 variant of P129
    ("An upper torque link subassembly attaches to the main fitting subassembly. A lower torque link subassembly attaches to the sliding tube subassembly. A damper attaches to the upper torque link subassembly. A pin installs through the damper and connects the upper and lower torque link subassemblies.",
     "Сборка верхнего шлиц-шарнира крепится к сборке корпуса стойки. Сборка нижнего шлиц-шарнира крепится к сборке скользящей трубы. Демпфер крепится к сборке верхнего шлиц-шарнира. Штифт устанавливается через демпфер и соединяет сборки верхнего и нижнего шлиц-шарниров."),
    # P131
    ("A rod and a cylinder install in the sliding tube subassembly. A piston installs in the cylinder. An upper diaphragm tube subassembly installs in the main fitting subassembly. A baffle, a compression orifice plate and a diaphragm subassembly install in the upper diaphragm tube subassembly. The rod goes through the baffle.",
     "Шток и цилиндр устанавливаются в сборку скользящей трубы. Поршень устанавливается в цилиндр. Сборка верхней диафрагменной трубы устанавливается в сборку корпуса стойки. Дефлектор, пластина отверстия сжатия и сборка диафрагмы устанавливаются в сборку верхней диафрагменной трубы. Шток проходит через дефлектор."),
    # P132
    ("An upper bearing housing installs between the top of the sliding tube subassembly and the main fitting subassembly. A recoil orifice plate operates in the upper bearing housing.",
     "Корпус верхнего подшипника устанавливается между верхней частью сборки скользящей трубы и сборкой корпуса стойки. Пластина отверстия обратного хода работает в корпусе верхнего подшипника."),
    # P133
    ("Operation (Refer to Figure 2)",
     "Работа (обратитесь к рисунку 2)"),
    # P135
    ("The sliding tube subassembly moves into the main fitting subassembly. The subsequent decrease in volume causes hydraulic fluid to flow through the upper bearing housing: the recoil orifice plate moves and slows the flow of hydraulic fluid. The decrease in volume also causes hydraulic fluid to move through the diaphragm and lift the compression orifice plate: the hydraulic fluid flows through the baffle and into the upper diaphragm tube subassembly. This slows the speed of the compression.",
     "Сборка скользящей трубы перемещается в сборку корпуса стойки. Последующее уменьшение объёма заставляет гидравлическую жидкость протекать через корпус верхнего подшипника: пластина отверстия обратного хода перемещается и замедляет поток гидравлической жидкости. Уменьшение объёма также заставляет гидравлическую жидкость перемещаться через диафрагму и поднимать пластину отверстия сжатия: гидравлическая жидкость протекает через дефлектор в сборку верхней диафрагменной трубы. Это замедляет скорость сжатия."),
    # P136
    ("Hydraulic fluid that moves into the upper diaphragm tube compresses the nitrogen in the main fitting subassembly and the upper diaphragm tube subassembly. As the pressure of the nitrogen increases, the hydraulic fluid in the rod moves against the piston. The piston is pushed into the cylinder and compresses the nitrogen in it. This slows the speed of the compression more.",
     "Гидравлическая жидкость, перемещающаяся в верхнюю диафрагменную трубу, сжимает азот в сборке корпуса стойки и сборке верхней диафрагменной трубы. По мере увеличения давления азота гидравлическая жидкость в штоке перемещается к поршню. Поршень вдавливается в цилиндр и сжимает находящийся в нём азот. Это дополнительно замедляет скорость сжатия."),
    # P138
    ("After compression, the nitrogen pressure in the cylinder pushes the piston to the end of the cylinder: hydraulic fluid moves out of the cylinder and into the rod. The nitrogen pressure in the main fitting subassembly and the upper diaphragm subassembly pushes the hydraulic fluid through the baffle: the compression orifice plate is pushed against the diaphragm subassembly and limits the flow of hydraulic fluid through it. This slows the speed of the recoil. The sliding tube subassembly moves out of the main fitting subassembly.",
     "После сжатия давление азота в цилиндре выталкивает поршень к концу цилиндра: гидравлическая жидкость перемещается из цилиндра в шток. Давление азота в сборке корпуса стойки и сборке верхней диафрагменной трубы выталкивает гидравлическую жидкость через дефлектор: пластина отверстия сжатия прижимается к сборке диафрагмы и ограничивает поток гидравлической жидкости через неё. Это замедляет скорость обратного хода. Сборка скользящей трубы выдвигается из сборки корпуса стойки."),

    # ─── Cross-reference CMM lines ───
    ("Safran Landing Systems UK Ltd Component Maintenance Manual, Axle Harness 1M and 2M, 32-12-29.",
     "Safran Landing Systems UK Ltd Руководство по техническому обслуживанию компонентов, Электрический жгут оси 1М и 2М, 32-12-29."),
    ("Safran Landing Systems UK Ltd Component Maintenance Manual, Damper, 32-11-93.",
     "Safran Landing Systems UK Ltd Руководство по техническому обслуживанию компонентов, Демпфер, 32-11-93."),
    ("Safran Landing Systems UK Ltd Component Maintenance Manual, Damper, 32-12-85.",
     "Safran Landing Systems UK Ltd Руководство по техническому обслуживанию компонентов, Демпфер, 32-12-85."),
    # P89 Part 2 CMM reference
    ("Safran Landing Systems UK Ltd Component Maintenance Manual, Main Landing Gear Leg and Dressings, 32-12-21.",
     "Safran Landing Systems UK Ltd Руководство по техническому обслуживанию компонентов, Стойка основного шасси и обвязка, 32-12-21."),

    # ─── Testing and Fault Isolation sentences ───
    ("The temperature of the test fluid must be between 20 and 40 \uf0b0C (68 and 104 \uf0b0F).",
     "Температура испытательной жидкости должна быть между 20 и 40 \uf0b0C (68 и 104 \uf0b0F)."),
    ("The test fluid must be clean: refer to M-DLPS910-1.",
     "Испытательная жидкость должна быть чистой: см. M-DLPS910-1."),
    ("155 lbf/in2). Make a record of the pressure. Close the charging valve (17-20) and hold the nitrogen pressure for 15 minutes.",
     "155 фунт/дюйм²). Запишите давление. Закройте зарядный клапан (17-20) и удерживайте давление азота в течение 15 минут."),
    ("Open the charging valve (17-20) and measure the nitrogen pressure: it must be the same as the record in para (2). Leakage must not occur.",
     "Откройте зарядный клапан (17-20) и измерьте давление азота: оно должно совпадать с записью в п. (2). Утечка не допускается."),
    ("Make sure that all of the nitrogen pressure has been released: remove the charging valve (17-20).",
     "Убедитесь, что всё давление азота было сброшено: снимите зарядный клапан (17-20)."),
    ("Refer to ASSEMBLY: install the charging valve (17-20) and complete the assembly procedure.",
     "См. СБОРКА: установите зарядный клапан (17-20) и завершите процедуру сборки."),
    ("Assemble the Load Cell and Adapter 460006232 and the Offset Adapter 460006234 to the main landing gear leg (1-1).",
     "Соберите датчик нагрузки и Adapter 460006232 и Offset Adapter 460006234 к стойке основного шасси (1-1)."),
    ("Slowly increase the hydraulic pressure to between 13,11 and 14,48 bar (190 and 210 lbf/in2) and let the unit extend fully.",
     "Медленно увеличивайте гидравлическое давление до 13,11 — 14,48 бар (190 — 210 фунт/дюйм²) и дайте изделию полностью выдвинуться."),
    ("Release the hydraulic pressure and fully close the unit.",
     "Сбросьте гидравлическое давление и полностью закройте изделие."),
    ("Do para (c) and (d) until the hydraulic fluid that comes out of the unit does not have air in it.",
     "Выполняйте п. (c) и (d) до тех пор, пока гидравлическая жидкость, выходящая из изделия, не будет без воздуха."),
    ("Fully close the unit and disconnect the hydraulic test rig.",
     "Полностью закройте изделие и отсоедините стенд для гидравлических испытаний."),
    ("CAUTION: DO NOT USE A PRESSURE OF MORE THAN 7,58 BAR (110 LBF/IN2).",
     "ВНИМАНИЕ: НЕ ИСПОЛЬЗУЙТЕ ДАВЛЕНИЕ БОЛЕЕ 7,58 БАР (110 ФУНТ/ДЮЙМ²)."),
    ("Slowly increase the nitrogen pressure until the unit starts to extend.\tHold the pressure and fully extend the unit. The pressure must not be more than 7,58 bar",
     "Медленно увеличивайте давление азота до начала выдвижения изделия.\tУдерживайте давление и полностью выдвиньте изделие. Давление не должно превышать 7,58 бар"),
    ("NOTE: The charging valve (17-20) must be open to let the unit extend fully.",
     "ПРИМЕЧАНИЕ: Зарядный клапан (17-20) должен быть открыт для полного выдвижения изделия."),
    ("NOTE: Nitrogen will be released through the charging valve (13-60) as the piston (17-200) moves.",
     "ПРИМЕЧАНИЕ: Азот будет выходить через зарядный клапан (13-60) по мере перемещения поршня (17-200)."),
    ("Keep the unit in this condition for a minimum of six hours.",
     "Оставьте изделие в этом состоянии минимум на шесть часов."),
    ("Compare the pressures P1A and P2A and compare the pressures P1B and P2B. The pressures P1A and P2A must be the same and the pressures P1B and P2B must be the same, unless:",
     "Сравните давления P1A и P2A и сравните давления P1B и P2B. Давления P1A и P2A должны совпадать, и давления P1B и P2B должны совпадать, если только:"),
    ("If there is a difference between the temperatures T1 and T2, calculate the correct value for the nitrogen pressures (these will be P3A and P3B) and adjust the pressures to the corrected values. Use the formula:",
     "Если имеется разница между температурами T1 и T2, рассчитайте правильное значение давлений азота (это будут P3A и P3B) и доведите давления до скорректированных значений. Используйте формулу:"),
    ("If there is an error because of the gauge capacity:",
     "Если имеется погрешность из-за ёмкости манометра:"),
    ("Release the pressure in the gauge.",
     "Сбросьте давление в манометре."),
    ("Calculate the correct values for the nitrogen pressures (these will be P5A and P5B) and adjust the pressures to the corrected values. Use the formula:",
     "Рассчитайте правильные значения давлений азота (это будут P5A и P5B) и доведите давления до скорректированных значений. Используйте формулу:"),
    ("Complete the torque procedure for the retaining pins (13-10): refer to ASSEMBLY.",
     "Завершите процедуру затяжки удерживающих штифтов (13-10): см. СБОРКА."),
    ("Use the loading press: set the dimension between the pins (10-80 and 11-130) to between 632,80 and 636,95 mm (24.9134 and 25.0767 in).",
     "Используйте нагрузочный пресс: установите размер между штифтами (10-80 и 11-130) в диапазоне 632,80 — 636,95 мм (24.9134 — 25.0767 дюйм)."),
    ("Adjust the spacers (6-140, 7-50, 7-190 and 7-240) or laminated shims (6-140A, 7-50A, 7-90A and 7-240A): refer to ASSEMBLY.",
     "Отрегулируйте проставки (6-140, 7-50, 7-190 и 7-240) или набор прокладок (6-140A, 7-50A, 7-90A и 7-240A): см. СБОРКА."),
    ("NOTE: If the calculated gap is in the tolerance, the spacers (6-140, 7-50, 7-190 and 7-240) or laminated shims (6-140A, 7-50A, 7-90A and 7-240A) are not necessary.",
     "ПРИМЕЧАНИЕ: Если расчётный зазор находится в пределах допуска, проставки (6-140, 7-50, 7-190 и 7-240) или набор прокладок (6-140A, 7-50A, 7-90A и 7-240A) не требуются."),
    ("NOTE: Make sure that the main landing gear leg (1-1) is electrically isolated from the equipment that is used to hold it.",
     "ПРИМЕЧАНИЕ: Убедитесь, что стойка основного шасси (1-1) электрически изолирована от оборудования, используемого для её удержания."),
    ("Use the Milliohmmeter Megger, Type BT51, to measure the electrical bonding resistance.",
     "Используйте миллиомметр Megger, тип BT51, для измерения сопротивления электрического соединения."),
    ("Measure between the bearing (20-250) and the test points given in Table 101. The electrical bonding resistance must not be more than the limit given in Table 101.",
     "Измерьте между подшипником (20-250) и контрольными точками, указанными в таблице 101. Сопротивление электрического соединения не должно превышать предел, указанный в таблице 101."),
    ("Measure between the axle of the sliding tube subassembly (17-240) and the test points given in Table 102. The electrical bonding resistance must not be more than the limit given in Table 102.",
     "Измерьте между осью сборки скользящей трубы (17-240) и контрольными точками, указанными в таблице 102. Сопротивление электрического соединения не должно превышать предел, указанный в таблице 102."),
    ("Electrical Bonding Resistance Tests Table 101 (Refer to Figure 102)",
     "Проверка сопротивления электрического соединения Таблица 101 (обратитесь к рисунку 102)"),
    ("Electrical Bonding Resistance Tests Table 102 (Refer to Figure 102)",
     "Проверка сопротивления электрического соединения Таблица 102 (обратитесь к рисунку 102)"),

    # ─── Disassembly section key sentences ───
    ("NOTE: Refer to TESTING AND FAULT ISOLATION to find the necessary level of disassembly. This will give the condition of the component or the possible cause of its malfunction.",
     "ПРИМЕЧАНИЕ: См. ИСПЫТАНИЯ И ПОИСК НЕИСПРАВНОСТЕЙ для определения необходимого уровня разборки. Это даст состояние компонента или возможную причину его неисправности."),
    ("Bend the outer coil of the wire thread insert to the centre of the hole.",
     "Согните наружный виток резьбовой спиральной вставки к центру отверстия."),
    ("Remove the wire thread insert. Make sure that broken pieces do not stay in the hole.",
     "Снимите резьбовую спиральную вставку. Убедитесь, что обломки не остались в отверстии."),
    ("These special tools are necessary:",
     "Необходимы следующие специальные инструменты:"),
    ("NOTE: Alternative equivalents are permitted.",
     "ПРИМЕЧАНИЕ: Допускаются альтернативные эквиваленты."),
    ("Procedure (Refer to IPL Figures 1 to 20)",
     "Процедура (обратитесь к рисункам ИПД от 1 до 20)"),
    ("Use these special tools as necessary during the procedure to lift and to hold the unit:",
     "Используйте следующие специальные инструменты при необходимости в ходе процедуры для подъёма и удержания изделия:"),
    ("Post SB 201-32-22: cut the Bowden cable (1-45) and remove the cross bolts (1-47 and 1-49).",
     "После SB 201-32-22: разрежьте трос Боудена (1-45) и снимите поперечные болты (1-47 и 1-49)."),
    ("WARNING: RELEASE ALL NITROGEN PRESSURE BEFORE YOU REMOVE THE CHARGING VALVES (13-60 AND 17-20).",
     "ПРЕДУПРЕЖДЕНИЕ: СБРОСЬТЕ ВСЁ ДАВЛЕНИЕ АЗОТА ПЕРЕД СНЯТИЕМ ЗАРЯДНЫХ КЛАПАНОВ (13-60 И 17-20)."),
    ("Slowly open the charging valve (17-20) and release all of the second stage nitrogen pressure.",
     "Медленно откройте зарядный клапан (17-20) и сбросьте всё давление азота второй ступени."),
    ("Slowly open the charging valve (13-60) and release all of the first stage nitrogen pressure.",
     "Медленно откройте зарядный клапан (13-60) и сбросьте всё давление азота первой ступени."),
    ("Use the Lifting Tackle 460006211 and install the sliding tube subassembly (17-240) in the Build Trolley 460007240.",
     "Используйте подъёмное приспособление 460006211 и установите сборку скользящей трубы (17-240) в сборочную тележку 460007240."),
    ("Remove the upper bearing housing (15-40) and related parts as follows:",
     "Снимите корпус верхнего подшипника (15-40) и связанные детали следующим образом:"),
    ("CAUTION: DISCARD THE SCREWS (15-90) AND THE LOCKING PLATES (15-80) WHEN REMOVED.",
     "ВНИМАНИЕ: УТИЛИЗИРУЙТЕ ВИНТЫ (15-90) И СТОПОРНЫЕ ПЛАСТИНЫ (15-80) ПОСЛЕ СНЯТИЯ."),
    ("Remove the two piece stop with inserts (15-130).",
     "Снимите двухчастный упор со вставками (15-130)."),
    ("Remove the locking pins (15-50), the retaining ring (15-60) and the recoil orifice plate (15-70).",
     "Снимите стопорные штифты (15-50), удерживающее кольцо (15-60) и пластину отверстия обратного хода (15-70)."),
    ("Use the Torque Adapter 460007283, the Torque Reactor 460007278, the Holding Blocks 460006406 and the Bench Clamp MT1025 to remove the diaphragm subassembly (15-190), the compression orifice plate (15-220), the clapper seat (15-230) and the baffle (15-240).",
     "Используйте Torque Adapter 460007283, Torque Reactor 460007278, Holding Blocks 460006406 и Bench Clamp MT1025 для снятия сборки диафрагмы (15-190), пластины отверстия сжатия (15-220), седла хлопушки (15-230) и дефлектора (15-240)."),
    ("Remove the lower bearing (16A-150A) from the lower bearing housing subassembly (16A-120B). Discard the machined lower bearing (16A-150A).",
     "Снимите нижний подшипник (16A-150A) из сборки корпуса нижнего подшипника (16A-120B). Утилизируйте обработанный нижний подшипник (16A-150A)."),
    ("Release the lock washer (17-90) and use the Torque Adapter 460006404 to remove the jacking dome (17-80). Remove the lock washer (17-90).",
     "Ослабьте стопорную шайбу (17-90) и используйте Torque Adapter 460006404 для снятия домкратной точки (17-80). Снимите стопорную шайбу (17-90)."),
    ("Hold the cylinder (17-230) in the Bench Clamp MT1025 and Holding Blocks MT1026/63.",
     "Удерживайте цилиндр (17-230) в Bench Clamp MT1025 и Holding Blocks MT1026/63."),
    ("Remove the lock plate (17-120) and use the Pin Spanner 460007284 to remove the nut subassembly (17-130). Remove the rod (17-160) and the washer (17-170).",
     "Снимите стопорную пластину (17-120) и используйте Pin Spanner 460007284 для снятия сборки гайки (17-130). Снимите шток (17-160) и шайбу (17-170)."),
    ("Use the Torque Adapter 460007232 to remove the locking nut (19-52). Remove the locking washer (19-54) and the outer race and the ball of the spherical bearing (19-50).",
     "Используйте Torque Adapter 460007232 для снятия контргайки (19-52). Снимите стопорную шайбу (19-54) и наружное кольцо и шарик сферического подшипника (19-50)."),
    ("NOTE: The outer race and the ball are parts of the spherical bearing (19-550).",
     "ПРИМЕЧАНИЕ: Наружное кольцо и шарик являются деталями сферического подшипника (19-550)."),
    ("Use the Hydraulic-Pneumatic Pump Set 460006497, the Bolt 460006498/7, the Press Pad 460006499/25 and the Extraction Tube 460004680 and remove the forward pintle bush (20-250A).",
     "Используйте Hydraulic-Pneumatic Pump Set 460006497, Bolt 460006498/7, Press Pad 460006499/25 и Extraction Tube 460004680 и снимите переднюю втулку навеса (20-250A)."),
    ("Remove the grooved spherical bearing (6-300) or the self lubricating bearing (6-300A) from the lower slave link (6-310).",
     "Снимите рифлёный сферический подшипник (6-300) или самосмазывающийся подшипник (6-300A) из нижнего ведомого звена (6-310)."),
    ("Use the Torque Reaction Adapter 460007242 to hold the pin (9-70) and use the Torque Adapter T14544 to remove the nut (9-50). Remove the spacer (9-60), the pin (9-70), the spacer (9-80) and the sleeve (9-90).",
     "Используйте Torque Reaction Adapter 460007242 для удержания штифта (9-70) и используйте Torque Adapter T14544 для снятия гайки (9-50). Снимите проставку (9-60), штифт (9-70), проставку (9-80) и втулку (9-90)."),
    ("Use the Crowfoot Wrench T14500 to remove the charging valve (13-60). Remove the O-ring seal (13-67) from the charging valve (13-60).",
     "Используйте Crowfoot Wrench T14500 для снятия зарядного клапана (13-60). Снимите уплотнительное кольцо (13-67) из зарядного клапана (13-60)."),
    ("Use the Crowfoot Wrench T14500 to remove the charging valve (17-20). Remove the O-ring seal (17-27) from the charging valve (17-20).",
     "Используйте Crowfoot Wrench T14500 для снятия зарядного клапана (17-20). Снимите уплотнительное кольцо (17-27) из зарядного клапана (17-20)."),
    ("Use the Assembly/Extraction Tool 460006410 to remove the level tube (15-300) and remove the O-ring seal (15-310).",
     "Используйте Assembly/Extraction Tool 460006410 для снятия трубки уровня (15-300) и снимите уплотнительное кольцо (15-310)."),
    ("Release the lock indentations of the locking washer (19-54).",
     "Ослабьте фиксирующие вдавливания стопорной шайбы (19-54)."),

    # ─── Part 5 Repair procedure sentences ───
    ("Examine the part to make sure that you have obeyed all the repair instructions correctly.",
     "Осмотрите деталь, чтобы убедиться, что все инструкции по ремонту выполнены правильно."),
    ("Examine the part to make sure that you have obeyed the repair instructions correctly.",
     "Осмотрите деталь, чтобы убедиться, что инструкции по ремонту выполнены правильно."),
    ("Special tools are not necessary.", "Специальные инструменты не требуются."),
    ("Materials are not necessary.", "Материалы не требуются."),
    ("Repair parts are not necessary.", "Ремонтные детали не требуются."),
    ("These repair parts are necessary:", "Необходимы следующие ремонтные детали:"),
    ("Repair loose but undamaged liner:", "Ремонт незакреплённого, но неповреждённого вкладыша:"),
    ("Repair damaged liner:", "Ремонт повреждённого вкладыша:"),
    ("CAUTION: FOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO MESSIER-DOWTY LIMITED: REFER TO GUIDE-CS-001.",
     "ВНИМАНИЕ: ПРИ ПОВРЕЖДЕНИЯХ, ПРЕВЫШАЮЩИХ ПРЕДЕЛЫ ДАННОЙ СХЕМЫ РЕМОНТА, ОБРАТИТЕСЬ К MESSIER-DOWTY LIMITED: СМ. GUIDE-CS-001."),
    ("CAUTION: FOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO MESSIER-DOWTY LTD: REFER TO GUIDE-CS-001.",
     "ВНИМАНИЕ: ПРИ ПОВРЕЖДЕНИЯХ, ПРЕВЫШАЮЩИХ ПРЕДЕЛЫ ДАННОЙ СХЕМЫ РЕМОНТА, ОБРАТИТЕСЬ К MESSIER-DOWTY LTD: СМ. GUIDE-CS-001."),

    # ─── Cleaning and Inspection sentences ───
    ("WARNING: DO NOT GET CLEANING AGENTS ON YOUR SKIN, IN YOUR EYES OR NEAR A FLAME. DO NOT BREATHE THE FUMES. ONLY USE IN A LOCATION THAT HAS A CONTINUOUS FLOW OF CLEAN AIR. CLEANING AGENTS ARE POISONOUS AND FLAMMABLE.",
     "ПРЕДУПРЕЖДЕНИЕ: НЕ ДОПУСКАЙТЕ ПОПАДАНИЯ ЧИСТЯЩИХ СРЕДСТВ НА КОЖУ, В ГЛАЗА ИЛИ ВБЛИЗИ ОГНЯ. НЕ ВДЫХАЙТЕ ПАРЫ. ИСПОЛЬЗУЙТЕ ТОЛЬКО В ПОМЕЩЕНИИ С НЕПРЕРЫВНЫМ ПОТОКОМ ЧИСТОГО ВОЗДУХА. ЧИСТЯЩИЕ СРЕДСТВА ЯДОВИТЫ И ОГНЕОПАСНЫ."),
    ("CAUTION: DO NOT USE CHLORINATED SOLVENTS. CHLORINATED SOLVENTS CAN MIX WITH VERY SMALL QUANTITIES OF WATER IN HYDRAULIC SYSTEMS TO MAKE HYDROCHLORIC ACID. HYDROCHLORIC ACID WILL CAUSE CORROSION ON METAL SURFACES.",
     "ВНИМАНИЕ: НЕ ИСПОЛЬЗУЙТЕ ХЛОРИРОВАННЫЕ РАСТВОРИТЕЛИ. ХЛОРИРОВАННЫЕ РАСТВОРИТЕЛИ МОГУТ СМЕШИВАТЬСЯ С ОЧЕНЬ МАЛЫМИ КОЛИЧЕСТВАМИ ВОДЫ В ГИДРАВЛИЧЕСКИХ СИСТЕМАХ С ОБРАЗОВАНИЕМ СОЛЯНОЙ КИСЛОТЫ. СОЛЯНАЯ КИСЛОТА ВЫЗОВЕТ КОРРОЗИЮ МЕТАЛЛИЧЕСКИХ ПОВЕРХНОСТЕЙ."),
    ("Clean all the metal parts with white spirit, Material Ref. Item 11-524. Make sure that you fully remove all sealants, adhesives and jointing compounds.",
     "Очистите все металлические детали уайт-спиритом, поз. ссылки материала 11-524. Убедитесь, что полностью удалены все герметики, адгезивы и соединительные составы."),
    ("Examine the rod (17-160) for the diameter of radial damping holes. The diameter of each hole must be between 5,40 and 5,60 mm (0.213 and 0.220 in).",
     "Осмотрите шток (17-160) на предмет диаметра радиальных демпфирующих отверстий. Диаметр каждого отверстия должен быть между 5,40 и 5,60 мм (0,213 и 0,220 дюйм)."),
    ("Examine the 4 holes in the sliding tube (18-80) where the bracket (8-170) installs, for burrs. If you find burrs contact Safran Landing Systems who will supply an applicable repair.",
     "Осмотрите 4 отверстия в скользящей трубе (18-80), где устанавливается кронштейн (8-170), на предмет заусенцев. При обнаружении заусенцев обратитесь в Safran Landing Systems, которые предоставят соответствующий ремонт."),
    ("NOTE: Use a good light source and 10x magnification to view the area, to look for burrs.",
     "ПРИМЕЧАНИЕ: Используйте хороший источник света и 10-кратное увеличение для осмотра области и поиска заусенцев."),
    ("Unless instructions are different:",
     "Если инструкции не указывают иное:"),
    ("CAUTION: YOU MUST DISASSEMBLE ALL PARTS, THIS WILL INCLUDE THE BUSHES, THEY MUST BE REMOVED AND DISCARDED. YOU MUST APPLY THE NDT INSPECTION TO THE DETAIL LEVEL PART ONLY AS IDENTIFIED IN TABLES 501 AND 502. IF THE BUSHES ARE NOT REMOVED THE INSPECTION IS NOT COMPLETE FOR THE DETAIL PART AND DAMAGE CAN OCCUR.",
     "ВНИМАНИЕ: ВЫ ДОЛЖНЫ РАЗОБРАТЬ ВСЕ ДЕТАЛИ, ВКЛЮЧАЯ ВТУЛКИ, ОНИ ДОЛЖНЫ БЫТЬ СНЯТЫ И УТИЛИЗИРОВАНЫ. ВЫ ДОЛЖНЫ ПРИМЕНЯТЬ НЕРАЗРУШАЮЩИЙ КОНТРОЛЬ ТОЛЬКО К ДЕТАЛИ НИЖНЕГО УРОВНЯ, КАК УКАЗАНО В ТАБЛИЦАХ 501 И 502. ЕСЛИ ВТУЛКИ НЕ СНЯТЫ, КОНТРОЛЬ ДЕТАЛИ НЕ ЯВЛЯЕТСЯ ПОЛНЫМ, И МОЖЕТ ПРОИЗОЙТИ ПОВРЕЖДЕНИЕ."),

    # ─── Section headings and sub-headings ───
    ("Description (Refer to Figures 1 and 2)",
     "Описание (обратитесь к рисункам 1 и 2)"),
    ("Diagram of Operation Figure 2",
     "Схема работы Рисунок 2"),
    ("Special Tools", "Специальные инструменты"),
    ("Initial Operations", "Начальные операции"),
    ("Procedure", "Процедура"),
    ("Cleaning", "Очистка"),
    ("General", "Общие сведения"),
    ("Materials", "Материалы"),
    ("These materials are necessary:", "Необходимы следующие материалы:"),
    ("Discard", "Утилизируйте"),
    ("fully close the unit", "полностью закройте изделие"),
    ("fully extend the unit", "полностью выдвиньте изделие"),

    # ══════════════════════════════════════════════════════════════════════════
    # WORD/PHRASE-LEVEL TRANSLATIONS (applied after full sentences)
    # ══════════════════════════════════════════════════════════════════════════

    # ─── Full standard phrases ───
    ("WARNING: DO NOT GET HYDRAULIC FLUID ON YOUR SKIN OR IN YOUR EYES. DO NOT BREATHE THE FUMES. ONLY USE IN A LOCATION THAT HAS A CONTINUOUS FLOW OF CLEAN AIR.",
     "ПРЕДУПРЕЖДЕНИЕ: НЕ ДОПУСКАЙТЕ ПОПАДАНИЯ ГИДРАВЛИЧЕСКОЙ ЖИДКОСТИ НА КОЖУ ИЛИ В ГЛАЗА. НЕ ВДЫХАЙТЕ ПАРЫ. РАБОТАЙТЕ ТОЛЬКО В ХОРОШО ПРОВЕТРИВАЕМОМ ПОМЕЩЕНИИ."),
    ("WARNING: DO NOT GET CLEANING AGENTS ON YOUR SKIN, IN YOUR EYES OR NEAR A FLAME. DO NOT BREATHE THE FUMES. ONLY USE IN A LOCATION THAT HAS A CONTINUOUS FLOW OF CLEAN AIR.",
     "ПРЕДУПРЕЖДЕНИЕ: НЕ ДОПУСКАЙТЕ ПОПАДАНИЯ ЧИСТЯЩИХ СРЕДСТВ НА КОЖУ, В ГЛАЗА ИЛИ ВБЛИЗИ ОТКРЫТОГО ОГНЯ. НЕ ВДЫХАЙТЕ ПАРЫ. РАБОТАЙТЕ ТОЛЬКО В ХОРОШО ПРОВЕТРИВАЕМОМ ПОМЕЩЕНИИ."),
    ("WARNING: DO NOT GET PAINT STRIPPER ON YOUR SKIN, IN YOUR EYES OR NEAR A FLAME. DO NOT BREATHE THE FUMES. ONLY USE IN A LOCATION THAT HAS A CONTINUOUS FLOW OF CLEAN AIR.",
     "ПРЕДУПРЕЖДЕНИЕ: НЕ ДОПУСКАЙТЕ ПОПАДАНИЯ СМЫВКИ КРАСКИ НА КОЖУ, В ГЛАЗА ИЛИ ВБЛИЗИ ОТКРЫТОГО ОГНЯ. НЕ ВДЫХАЙТЕ ПАРЫ. РАБОТАЙТЕ ТОЛЬКО В ХОРОШО ПРОВЕТРИВАЕМОМ ПОМЕЩЕНИИ."),
    ("WARNING: RELEASE ALL NITROGEN PRESSURE BEFORE YOU REMOVE THE CHARGING VALVES",
     "ПРЕДУПРЕЖДЕНИЕ: СБРОСЬТЕ ВСЁ ДАВЛЕНИЕ АЗОТА ПЕРЕД СНЯТИЕМ ЗАРЯДНЫХ КЛАПАНОВ"),
    ("CAUTION: DO NOT USE CHLORINATED SOLVENTS. CHLORINATED SOLVENTS CAN MIX WITH VERY SMALL QUANTITIES OF WATER IN HYDRAULIC SYSTEMS TO MAKE HYDROCHLORIC ACID.",
     "ВНИМАНИЕ: НЕ ИСПОЛЬЗУЙТЕ ХЛОРИРОВАННЫЕ РАСТВОРИТЕЛИ. ХЛОРИРОВАННЫЕ РАСТВОРИТЕЛИ МОГУТ СМЕШИВАТЬСЯ С ОЧЕНЬ МАЛЫМ КОЛИЧЕСТВОМ ВОДЫ В ГИДРАВЛИЧЕСКИХ СИСТЕМАХ, ОБРАЗУЯ СОЛЯНУЮ КИСЛОТУ."),
    ("CAUTION: DO NOT PUT AN END LOAD OF MORE THAN", "ВНИМАНИЕ: НЕ ПРИКЛАДЫВАЙТЕ ОСЕВУЮ НАГРУЗКУ БОЛЕЕ"),
    ("CAUTION: DO NOT CAUSE DAMAGE TO THE PAINT FINISH.", "ВНИМАНИЕ: НЕ ПОВРЕДИТЕ ЛАКОКРАСОЧНОЕ ПОКРЫТИЕ."),
    ("CAUTION: YOU MUST DISASSEMBLE ALL PARTS, THIS WILL INCLUDE THE BUSHES, THEY MUST BE REMOVED AND DISCARDED. YOU MUST APPLY THE NDT INSPECTION TO THE DETAIL PARTS.",
     "ВНИМАНИЕ: ВЫ ДОЛЖНЫ РАЗОБРАТЬ ВСЕ ДЕТАЛИ, ВКЛЮЧАЯ ВТУЛКИ — ОНИ ДОЛЖНЫ БЫТЬ СНЯТЫ И УТИЛИЗИРОВАНЫ. ВЫ ДОЛЖНЫ ПРИМЕНИТЬ НЕРАЗРУШАЮЩИЙ КОНТРОЛЬ К ДЕТАЛЯМ."),
    ("CAUTION: DISCARD THE SCREWS", "ВНИМАНИЕ: УТИЛИЗИРУЙТЕ ВИНТЫ"),
    ("THE LOCKING PLATES", "СТОПОРНЫЕ ПЛАСТИНЫ"),
    ("WHEN REMOVED.", "ПРИ СНЯТИИ."),
    ("AND THE CHARGING VALVES", "И ЗАРЯДНЫХ КЛАПАНОВ"),
    ("ON THE MAIN LANDING GEAR LEG", "НА СТОЙКУ ОСНОВНОГО ШАССИ"),
    ("THE GEAR MUST BE INFLATED TO THE APPROPRIATE PRESSURES BEFORE BEING PLACED IN SERVICE.",
     "СТОЙКА ДОЛЖНА БЫТЬ НАКАЧАНА ДО СООТВЕТСТВУЮЩИХ ДАВЛЕНИЙ ПЕРЕД ВВОДОМ В ЭКСПЛУАТАЦИЮ."),

    # ─── Instruction starts (verbs) ───
    ("Slowly increase", "Медленно увеличьте"),
    ("Slowly open", "Медленно откройте"),
    ("Make sure that all of", "Убедитесь, что всё"),
    ("Make sure that there is no pressure in", "Убедитесь, что нет давления в"),
    ("Make sure that", "Убедитесь, что"),
    ("Make sure", "Убедитесь"),
    ("Make a record of", "Запишите"),
    ("Do para", "Выполните пункт"),
    ("Refer to ASSEMBLY:", "Обратитесь к СБОРКЕ:"),
    ("Refer to ASSEMBLY.", "Обратитесь к СБОРКЕ."),
    ("Refer to TESTING AND FAULT ISOLATION", "Обратитесь к ПРОВЕРКЕ И ПОИСКУ НЕИСПРАВНОСТЕЙ"),
    ("Refer to REPAIR.", "Обратитесь к РЕМОНТУ."),
    ("Refer to Figure", "Обратитесь к рисунку"),
    ("refer to Figure", "обратитесь к рисунку"),
    ("Refer to M-DLPS", "Обратитесь к M-DLPS"),
    ("refer to M-DLPS", "обратитесь к M-DLPS"),
    ("Refer to PCS-", "Обратитесь к PCS-"),
    ("refer to PCS-", "обратитесь к PCS-"),
    ("Refer to para", "Обратитесь к пункту"),
    ("refer to para", "обратитесь к пункту"),
    ("Refer to ASSEMBLY", "Обратитесь к СБОРКЕ"),
    ("refer to REPAIR", "обратитесь к РЕМОНТУ"),
    ("Refer to ", "Обратитесь к "),
    ("refer to ", "обратитесь к "),

    # ─── Common instructional phrases ───
    ("Remove the damaged paint:", "Удалите повреждённую краску:"),
    ("Clean the part:", "Очистите деталь:"),
    ("Paint the part:", "Покрасьте деталь:"),
    ("Visually examine each part.", "Визуально осмотрите каждую деталь."),
    ("Carefully examine changes of section and areas which contact sealing rings.",
     "Тщательно осмотрите переходы сечений и зоны контакта с уплотнительными кольцами."),
    ("Examine each part for these types of damage:", "Осмотрите каждую деталь на следующие типы повреждений:"),
    ("These special tools are necessary:", "Необходимы следующие специальные инструменты:"),
    ("These materials are necessary:", "Необходимы следующие материалы:"),
    ("This equipment is necessary:", "Необходимо следующее оборудование:"),
    ("NOTE: Alternative equivalents are permitted.", "ПРИМЕЧАНИЕ: Допускается использование эквивалентных аналогов."),
    ("NOTE: If the calculated gap is in the tolerance,", "ПРИМЕЧАНИЕ: Если рассчитанный зазор находится в пределах допуска,"),
    ("is not installed.", "не устанавливается."),
    ("Use approved persons and good aircraft engineering practice for all procedures in this manual.",
     "Все процедуры настоящего руководства должны выполняться допущенным персоналом с соблюдением надлежащей авиационной инженерной практики."),
    ("Discard parts that you must not use again.", "Утилизируйте детали, непригодные для повторного использования."),
    ("The procedure to examine the parts is in two levels:", "Процедура осмотра деталей выполняется на двух уровнях:"),
    ("Unless instructions are different:", "Если не указано иное:"),
    ("Parts that are included in Tables 501 and 502 must be fully disassembled to the lowest detail level for NDT inspection. This includes the removal of all bushes.",
     "Детали, включённые в таблицы 501 и 502, должны быть полностью разобраны до мельчайших деталей для проведения неразрушающего контроля. Это включает снятие всех втулок."),

    # ─── Verb phrases with "the" ───
    ("Remove the", "Снимите"),
    ("Install the", "Установите"),
    ("Use the", "Используйте"),
    ("Connect the", "Подсоедините"),
    ("Disconnect the", "Отсоедините"),
    ("Open the", "Откройте"),
    ("Close the", "Закройте"),
    ("Release the", "Ослабьте"),
    ("Measure the", "Измерьте"),
    ("Hold the", "Удерживайте"),
    ("Complete the", "Завершите"),
    ("Set the", "Установите"),
    ("Attach the", "Прикрепите"),
    ("Apply the", "Нанесите"),
    ("Reduce the", "Уменьшите"),
    ("Increase the", "Увеличьте"),
    ("Keep the", "Оставьте"),
    ("Write this", "Запишите эти"),
    ("Compare the", "Сравните"),
    ("Calculate the", "Рассчитайте"),
    ("Adjust the", "Отрегулируйте"),
    ("Clean the", "Очистите"),
    ("Dry the", "Высушите"),
    ("Bend the", "Согните"),
    ("Remove and discard the", "Снимите и утилизируйте"),
    ("Assemble the", "Соберите"),

    # ─── "to [verb]" patterns (MUST come BEFORE lowercase verb phrases) ───
    ("to remove the", "для снятия"),
    ("to remove", "для снятия"),
    ("to hold the", "для удержания"),
    ("to hold", "для удержания"),
    ("to measure the", "для измерения"),
    ("to measure", "для измерения"),
    ("to lift the", "для подъёма"),
    ("to lift", "для подъёма"),

    # ─── Lowercase verb phrases (after comma, semicolon, colon) ───
    ("remove the", "снимите"),
    ("install the", "установите"),
    ("use the", "используйте"),
    ("let the", "дайте"),
    ("measure the", "измерьте"),
    ("torque it to", "затяните моментом"),
    ("connect the", "подсоедините"),
    ("disconnect the", "отсоедините"),
    ("open the", "откройте"),
    ("close the", "закройте"),
    ("release the", "ослабьте"),
    ("reduce the", "уменьшите"),

    # ─── Part nouns (with "the", plurals first) ───
    ("the split pins", "шплинты"),
    ("the split pin", "шплинт"),
    ("the slotted nuts", "шлицевые гайки"),
    ("the slotted nut", "шлицевую гайку"),
    ("the self-locking nuts", "самоконтрящиеся гайки"),
    ("the self-locking nut", "самоконтрящуюся гайку"),
    ("the locking nuts", "контргайки"),
    ("the locking nut", "контргайку"),
    ("the locking washers", "стопорные шайбы"),
    ("the locking washer", "стопорную шайбу"),
    ("the lock washers", "стопорные шайбы"),
    ("the lock washer", "стопорную шайбу"),
    ("the cup washers", "тарельчатые шайбы"),
    ("the cup washer", "тарельчатую шайбу"),
    ("the tab washers", "стопорные шайбы"),
    ("the tab washer", "стопорную шайбу"),
    ("the cap screws", "винты с головкой"),
    ("the cap screw", "винт с головкой"),
    ("the retaining pins", "стопорные штифты"),
    ("the retaining pin", "стопорный штифт"),
    ("the stop rings", "стопорные кольца"),
    ("the stop ring", "стопорное кольцо"),
    ("the O-ring seals", "уплотнительные кольца"),
    ("the O-ring seal", "уплотнительное кольцо"),
    ("the backing rings", "опорные кольца"),
    ("the backing ring", "опорное кольцо"),
    ("the locking pins", "стопорные штифты"),
    ("the locking pin", "стопорный штифт"),
    ("the locking plates", "стопорные пластины"),
    ("the locking plate", "стопорную пластину"),
    ("the lock plate", "стопорную пластину"),
    ("the lubrication fittings", "смазочные ниппели"),
    ("the lubrication fitting", "смазочный ниппель"),
    ("the identification washers", "идентификационные шайбы"),
    ("the identification washer", "идентификационную шайбу"),
    ("the lubrication adapters", "смазочные адаптеры"),
    ("the lubrication adapter", "смазочный адаптер"),
    ("the lubrication shaft subassembly", "сборку смазочного вала"),
    ("the retaining ring", "стопорное кольцо"),
    ("the recoil orifice plate", "пластину отверстия обратного хода"),
    ("the two piece stop with inserts", "двухсоставной упор со вставками"),
    ("the transfer dowels", "переходные штифты"),
    ("the valve support", "опору клапана"),
    ("the valve stem", "шток клапана"),
    ("the jacking dome", "домкратную точку"),
    ("the outer race and the ball of the spherical bearing",
     "наружное кольцо и шарик сферического подшипника"),
    ("the outer race and the ball", "наружное кольцо и шарик"),
    ("the spherical bearing", "сферический подшипник"),
    ("the grooved spherical bearing", "сферический подшипник с канавкой"),
    ("the self lubricating bearing", "самосмазывающийся подшипник"),
    ("the inflation valve subassembly", "сборку клапана зарядки"),
    ("the inflation valve", "клапан зарядки"),
    ("the charging valves", "зарядные клапаны"),
    ("the charging valve", "зарядный клапан"),
    ("the nuts", "гайки"),
    ("the nut", "гайку"),
    ("the washers", "шайбы"),
    ("the washer", "шайбу"),
    ("the bolts", "болты"),
    ("the bolt", "болт"),
    ("the screws", "винты"),
    ("the screw", "винт"),
    ("the pins", "штифты"),
    ("the pin", "штифт"),
    ("the spacers", "проставки"),
    ("the spacer", "проставку"),
    ("the shims", "прокладки"),
    ("the shim", "прокладку"),
    ("the bearings", "подшипники"),
    ("the bearing", "подшипник"),
    ("the bushes", "втулки"),
    ("the bush", "втулку"),
    ("the seals", "уплотнения"),
    ("the seal", "уплотнение"),
    ("the sleeves", "втулки"),
    ("the sleeve", "втулку"),
    ("the clamp", "хомут"),
    ("the dust cap", "пылезащитную крышку"),
    ("the brackets", "кронштейны"),
    ("the bracket", "кронштейн"),
    ("the retainers", "фиксаторы"),
    ("the retainer", "фиксатор"),
    ("the wedge", "клин"),
    ("the bung", "заглушку"),
    ("the plates", "пластины"),
    ("the plate", "пластину"),
    ("the rod", "шток"),
    ("the piston", "поршень"),
    ("the cylinder", "цилиндр"),
    ("the housing", "корпус"),
    ("the labels", "этикетки"),
    ("the label", "этикетку"),
    ("the joint seal", "соединительное уплотнение"),
    ("the sealing ring", "уплотнительное кольцо"),
    ("the wiper ring", "грязесъёмное кольцо"),
    ("the inner liner", "внутренний вкладыш"),
    ("the bonding cable", "соединительный провод"),
    ("the static discharge connector", "штыревой разъём статического разряда"),
    ("the damper", "демпфер"),
    ("the level tube", "трубку уровня"),
    ("the unit", "изделие"),
    ("the wire thread insert", "резьбовую спиральную вставку"),
    ("the wire thread inserts", "резьбовые спиральные вставки"),
    ("the lock indentations", "фиксирующие вдавливания"),
    ("the wiring diagram plate", "табличку электрической схемы"),

    # ─── Named assemblies/subassemblies ───
    ("the upper torque link subassembly", "сборку верхнего шлиц-шарнира"),
    ("the lower torque link subassembly", "сборку нижнего шлиц-шарнира"),
    ("the slave link subassembly", "сборку ведомого звена"),
    ("the lower slave link subassembly", "сборку нижнего ведомого звена"),
    ("the pivot bracket subassembly", "сборку поворотного кронштейна"),
    ("the bracket subassembly", "сборку кронштейна"),
    ("the bracket assembly", "сборку кронштейна"),
    ("the rod end assembly", "сборку наконечника тяги"),
    ("the bolt subassembly", "сборку болта"),
    ("the cardan assembly", "сборку кардана"),
    ("the transfer block subassembly", "сборку переходного блока"),
    ("the shock absorber subassembly", "сборку амортизатора"),
    ("the sliding tube subassembly", "сборку скользящей трубы"),
    ("the upper diaphragm tube subassembly", "сборку верхней диафрагменной трубы"),
    ("the lower bearing subassembly", "сборку нижнего подшипника"),
    ("the lower bearing housing subassembly", "сборку корпуса нижнего подшипника"),
    ("the main fitting subassembly", "сборку корпуса стойки"),
    ("the washer subassembly", "сборку шайбы"),
    ("the nut subassembly", "сборку гайки"),
    ("the diaphragm subassembly", "сборку диафрагмы"),
    ("the upper bearing housing", "корпус верхнего подшипника"),
    ("the lower bearing", "нижний подшипник"),
    ("the gland housing", "корпус сальника"),
    ("the upper diaphragm tube", "верхнюю диафрагменную трубу"),
    ("the main fitting", "корпус стойки"),
    ("the upper pivot bracket", "верхний поворотный кронштейн"),
    ("the harness support bracket", "кронштейн крепления жгута"),
    ("the harness support", "крепление жгута"),
    ("the proximity switch", "датчик приближения"),
    ("the target", "мишень"),
    ("the uplock pin", "штифт замка убранного положения"),
    ("the ground stud subassembly", "сборку штыря заземления"),
    ("the drag arm sleeve", "втулку тяги"),
    ("the lock stay cardan subassembly", "сборку кардана фиксатора"),
    ("the lock stay cardan", "кардан фиксатора"),
    ("the main landing gear leg", "стойку основного шасси"),

    # ─── Tool names (when used in procedural text) ───
    ("the Crowfoot Wrench T14500", "рожковый ключ T14500"),
    ("the Crowfoot Wrench", "рожковый ключ"),
    ("the Turner Inflation Equipment T14218", "оборудование для заправки Turner T14218"),
    ("the Turner Inflation Equipment", "оборудование для заправки Turner"),
    ("the Charging Adapter 460002502", "адаптер для зарядки 460002502"),
    ("the Charging Adapter", "адаптер для зарядки"),
    ("the Milliohmmeter Megger, Type BT51", "миллиомметр Megger, тип BT51"),
    ("the Lampbox 460005842", "контрольную лампу 460005842"),
    ("the Lampbox", "контрольную лампу"),
    ("the Loading Press", "нагрузочный пресс"),
    ("the loading press", "нагрузочный пресс"),
    ("the Holding Fixture 460006231", "удерживающее приспособление 460006231"),
    ("the Press Adapter 460006233", "адаптер пресса 460006233"),
    ("the Bottom Press Adapter 460007260", "нижний адаптер пресса 460007260"),
    ("the Load Cell and Adapter 460006232", "датчик нагрузки и адаптер 460006232"),
    ("the Offset Adapter 460006234", "смещённый адаптер 460006234"),
    ("the Lifting Bar Assembly 460006208", "сборку подъёмной штанги 460006208"),
    ("the Spherical Bearing Locator 460007282", "позиционер сферического подшипника 460007282"),
    ("the Pintle Location Assembly 460007281", "сборку позиционирования навеса 460007281"),
    ("the Transport and Build Trolley 460006213", "транспортировочно-сборочную тележку 460006213"),
    ("the Support Arms 460006215", "опорные рычаги 460006215"),
    ("the Towing Frame 460006216", "буксировочную раму 460006216"),
    ("the Jacking Dome Adapter 460006223", "адаптер домкратной точки 460006223"),
    ("the Adapter 460006237", "адаптер 460006237"),
    ("the Spacer 460007231", "проставку 460007231"),
    ("the Location Frame 460007234", "установочную раму 460007234"),
    ("the Location Frame 460007235", "установочную раму 460007235"),
    ("the Build Trolley 460007240", "сборочную тележку 460007240"),
    ("the Lifting Tackle 460006211", "подъёмное приспособление 460006211"),
    ("the Bench Clamp MT1025", "настольный зажим MT1025"),
    ("the Holding Blocks MT1026/63", "удерживающие блоки MT1026/63"),
    ("the Holding Blocks 460006406", "удерживающие блоки 460006406"),
    ("Pin Spanner 460007279", "штифтовый ключ 460007279"),
    ("Pin Spanner 460007284", "штифтовый ключ 460007284"),
    ("the Torque Reaction Adapter 460007242", "адаптер реакции момента 460007242"),
    ("the Torque Adapter T14544", "моментный адаптер T14544"),
    ("the Torque Adapter 460007230", "моментный адаптер 460007230"),
    ("the Torque Adapter 460007283", "моментный адаптер 460007283"),
    ("the Torque Adapter 460007232", "моментный адаптер 460007232"),
    ("the Torque Adapter 460006404", "моментный адаптер 460006404"),
    ("the Torque Reactor 460007278", "реактор момента 460007278"),
    ("the Extractor Pad and Drawbolt 460006415", "извлекающую пластину и вытяжной болт 460006415"),
    ("the Extractor 460006416", "экстрактор 460006416"),
    ("the Extractor 460006413", "экстрактор 460006413"),
    ("the Extractor 460001355", "экстрактор 460001355"),
    ("the Extractor 460006253", "экстрактор 460006253"),
    ("the Assembly/Extraction Tool 460006410", "инструмент для сборки/извлечения 460006410"),
    ("the Hydraulic-Pneumatic Pump Set 460006497", "гидропневматический насосный комплект 460006497"),
    ("the Press Pad Assembly 460006267", "сборку прижимной пластины 460006267"),
    ("the 28 VDC power supply", "источник питания 28 В постоянного тока"),
    ("the 28 VDC supply", "источник питания 28 В постоянного тока"),
    ("the hydraulic test rig", "стенд для гидравлических испытаний"),
    ("the nitrogen supply", "источник азота"),
    ("the nitrogen pressure", "давление азота"),
    ("the hydraulic pressure", "гидравлическое давление"),
    ("the test fluid", "испытательная жидкость"),
    ("the inflation equipment", "оборудование для заправки"),
    ("the test circuit", "испытательный контур"),

    # ─── Common prepositions/connectors with context ───
    ("from the", "из"),
    ("in the", "в"),
    ("to the", "к"),
    ("on the", "на"),
    ("for the", "для"),
    ("at the", "у"),
    ("and the", "и"),
    ("or the", "или"),
    ("with the", "с"),
    ("and its related parts", "и связанные детали"),
    ("and its attached parts", "и прикреплённые детали"),
    ("as necessary during the procedure", "при необходимости в ходе процедуры"),
    ("as follows:", "следующим образом:"),
    ("as necessary", "при необходимости"),
    ("must be", "должно быть"),
    ("must not be more than", "не должно превышать"),
    ("must not be", "не должно быть"),
    ("must not occur", "не должна происходить"),
    ("must have", "должен иметь"),
    ("must operate before", "должен сработать до того, как"),
    ("has closed by", "будет закрыта на"),
    (" again.", " снова."),
    ("Leakage", "Утечка"),
    ("until the", "до тех пор, пока"),
    ("until", "до"),
    ("There is a difference between the temperatures", "Имеется разница между температурами"),
    ("There is an error because of the pressure gauge capacity.", "Имеется погрешность из-за ёмкости манометра."),
    ("If there is a difference", "Если имеется разница"),
    ("If there is an error", "Если имеется погрешность"),
    ("calculate the correct value for", "рассчитайте правильное значение"),
    ("the correct values for", "правильные значения"),
    ("adjust the pressures to the corrected values", "доведите давления до скорректированных значений"),
    ("Use the formula:", "Используйте формулу:"),
    ("for temperatures in C", "для температур в °C"),
    ("for temperatures in F", "для температур в °F"),
    ("for temperatures in \uf0b0C", "для температур в °C"),
    ("for temperatures in \uf0b0F", "для температур в °F"),
    ("for pressures in bar", "для давлений в бар"),
    ("for pressures in lbf/in2", "для давлений в фунт/дюйм²"),
    ("not less than", "не менее"),
    ("not more than", "не более"),
    ("approximately", "приблизительно"),
    ("a minimum of", "минимум"),

    # ─── Keywords ───
    ("WARNING:", "ПРЕДУПРЕЖДЕНИЕ:"),
    ("CAUTION:", "ВНИМАНИЕ:"),
    ("NOTE:", "ПРИМЕЧАНИЕ:"),
    ("PRE SB", "До SB"),
    ("POST SB", "После SB"),
    ("Pre SB", "До SB"),
    ("Post SB", "После SB"),
    ("Examination of Magnetic Steel Parts by Non-destructive Testing",
     "Контроль деталей из магнитной стали неразрушающими методами"),
    ("Examination of Non-Magnetic Parts by Non-destructive Testing",
     "Контроль немагнитных деталей неразрушающими методами"),
    ("Examine all parts shown in Tables", "Осмотрите все детали, указанные в таблицах"),
    ("to the applicable NDT and information given", "в соответствии с применимым неразрушающим контролем и указанной информацией"),
    ("The ambient temperature", "Температура окружающей среды"),
    ("the ambient temperature", "температуру окружающей среды"),
    ("The temperature of the test fluid", "Температура испытательной жидкости"),
    ("The hydraulic test rig", "Стенд для гидравлических испытаний"),
    ("The inflation equipment", "Оборудование для заправки"),
    ("The test fluid must be clean:", "Испытательная жидкость должна быть чистой:"),
    ("During all hydraulic tests,", "Во время всех гидравлических испытаний,"),
    ("Examine the unit for damage before you start the tests.",
     "Осмотрите изделие на наличие повреждений перед началом испытаний."),
    ("During the proximity switch tests", "Во время испытаний датчиков приближения"),
    ("AECMA Simplified English to PSC-85-16598 is used in this manual.",
     "В данном руководстве используется упрощённый английский AECMA по PSC-85-16598."),


    # ─── Remaining sentence patterns ───
    ("until the unit starts to extend", "до начала выдвижения изделия"),
    ("Hold the pressure and fully extend the unit.", "Удерживайте давление и полностью выдвиньте изделие."),
    ("The pressure must not be more than", "Давление не должно превышать"),
    ("There is a difference between the temperatures", "Имеется разница между температурами"),
    ("starts to extend", "начнёт выдвигаться"),
    ("horizontally", "горизонтально"),
    ("vertically", "вертикально"),
    ("Adjustment", "Регулировка"),
    ("drag arm sleeve", "втулку тяги"),
    ("forward pintle bush", "переднюю втулку навеса"),
    ("Parts", "Части"),
    ("These include:", "К ним относятся:"),
    ("hydraulic fluid", "гидравлическая жидкость"),
    ("Hydraulic fluid", "Гидравлическая жидкость"),
    ("White spirit", "Уайт-спирит"),
    ("white spirit", "уайт-спирит"),
    ("Material Ref. Item", "Поз. ссылки материала"),
    ("Clean all the metal parts with", "Очистите все металлические детали"),
    ("Dry all the metal parts.", "Высушите все металлические детали."),
    ("Use clean PVC or polythene gloves to prevent corrosion of metal parts.",
     "Используйте чистые перчатки из ПВХ или полиэтилена для предотвращения коррозии металлических деталей."),
    ("Prevent corrosion of the metal parts that you do not immediately use for assembly procedures:",
     "Предотвратите коррозию металлических деталей, которые вы не используете немедленно для процедур сборки:"),
    ("for left configuration units", "для изделий в левой конфигурации"),
    ("for right configuration units", "для изделий в правой конфигурации"),
    ("left configuration", "левая конфигурация"),
    ("right configuration", "правая конфигурация"),
    ("to lift and to hold the unit", "для подъёма и удержания изделия"),
    ("Nitrogen", "Азот"),
    ("Labels", "Этикетки"),
    ("wiring diagram plate", "табличка электрической схемы"),
    ("- Installation", "— Установка"),
    ("Use with", "Используется с"),
    ("The thread size is", "Размер резьбы —"),
    ("pitch", "шаг"),
    ("lbf/in2", "фунт/дюйм²"),

    # ─── Table 101/102 related ───
    ("BEARING", "ПОДШИПНИК"),
    ("AXLE OF SLIDING TUBE SUBASSEMBLY", "ОСЬ СБОРКИ СКОЛЬЗЯЩЕЙ ТРУБЫ"),
    ("Table 101", "Таблица 101"),
    ("Table 102", "Таблица 102"),
    ("Tables 501 and 502", "таблицах 501 и 502"),
    ("Table 501", "Таблица 501"),
    ("Table 502", "Таблица 502"),
    ("Figure 1", "Рисунок 1"),
    ("Figure 2", "Рисунок 2"),
    ("Figure 101", "Рисунок 101"),
    ("Figure 102", "Рисунок 102"),
    ("Main Landing Gear Leg", "Стойка основного шасси"),

    # ─── Drawing annotation terms (for figure textboxes) ───
    ("LENGTH OF CADMIUM PLATE", "ДЛИНА КАДМИЕВОГО ПОКРЫТИЯ"),
    ("NO CADMIUM PLATE", "БЕЗ КАДМИЕВОГО ПОКРЫТИЯ"),
    ("NO PAINT", "БЕЗ КРАСКИ"),
    ("CADMIUM PLATE", "КАДМИЕВОЕ ПОКРЫТИЕ"),
    ("LENGTH OF", "ДЛИНА"),
    ("UP TO BUSH FLANGES", "ДО ФЛАНЦЕВ ВТУЛОК"),
    ("UP TO", "ДО"),
    ("FOR ALL SECTION VIEWS SEE SHEET", "ВСЕ ВИДЫ СЕЧЕНИЙ СМ. ЛИСТ"),
    ("LIMIT OF SERMETEL W TERMINATION FROM CENTER", "ГРАНИЦА ПЕРЕХОДА SERMETEL W ОТ ЦЕНТРА"),
    ("CENTERLINE OF SLIDING TUBE", "ОСЕВАЯ ЛИНИЯ СКОЛЬЗЯЩЕЙ ТРУБЫ"),
    ("SERMETEL W ON INTERNAL", "SERMETEL W НА ВНУТРЕННЕМ"),
    ("SERMETEL W TO IFC", "SERMETEL W ДО IFC"),
    ("SERMETEL W TO", "SERMETEL W ДО"),
    ("SERMETEL W", "SERMETEL W"),
    ("PRIMER PAINT TO", "ГРУНТОВОЧНАЯ КРАСКА ДО"),
    ("PRIMER PAINT", "ГРУНТОВОЧНАЯ КРАСКА"),
    ("CHROMIUM PLATE", "ХРОМОВОЕ ПОКРЫТИЕ"),
    ("PAINT TO", "КРАСКА ДО"),
    ("VIEW ON ARROW", "ВИД ПО СТРЕЛКЕ"),
    ("RADIUS INTERSECTION POINT", "ТОЧКИ ПЕРЕСЕЧЕНИЯ РАДИУСА"),
    ("INTERSECTION POINT", "ТОЧКА ПЕРЕСЕЧЕНИЯ"),
    ("INTERSECTION", "ПЕРЕСЕЧЕНИЕ"),
    ("SECTION", "СЕЧЕНИЕ"),
    ("DETAIL", "ДЕТАЛЬ"),
    ("IN THIS FACE ONLY", "ТОЛЬКО НА ДАННОЙ ПОВЕРХНОСТИ"),
    ("THIS FACE ONLY", "ТОЛЬКО ДАННАЯ ПОВЕРХНОСТЬ"),
    ("FROM OUTSIDE FACE", "ОТ НАРУЖНОЙ ПОВЕРХНОСТИ"),
    ("FROM THIS SURFACE", "ОТ ДАННОЙ ПОВЕРХНОСТИ"),
    ("TO A DEPTH OF", "НА ГЛУБИНУ"),
    ("HOLE TO DEPTH OF", "ОТВЕРСТИЕ НА ГЛУБИНУ"),
    ("OVER LENGTH", "ПО ДЛИНЕ"),
    ("POSITION ONLY", "ТОЛЬКО ПОЛОЖЕНИЕ"),
    ("THROUGH DIA.", "СКВОЗНОЙ ДИАМ."),
    ("INNER DIAMETER", "ВНУТРЕННИЙ ДИАМЕТР"),
    ("DIAMETER", "ДИАМЕТР"),
    ("DIA. SPOTFACE", "ДИАМ. ЗЕНКОВКА"),
    ("DIA.", "ДИАМ."),
    ("BOTH HOLES", "ОБА ОТВЕРСТИЯ"),
    ("BOTH SIDES", "ОБЕ СТОРОНЫ"),
    ("INTERNALLY", "ИЗНУТРИ"),
    ("MAXIMUM", "МАКСИМУМ"),
    ("TYPICAL", "ТИПИЧНО"),
    ("PLACES", "МЕСТ"),
    ("HOLES", "ОТВЕРСТИЙ"),
    ("BORES", "ОТВЕРСТИЙ"),
    ("CHAMFER", "ФАСКА"),
    ("RUNOUT", "БИЕНИЕ"),
    ("CENTERLINE", "ОСЕВАЯ ЛИНИЯ"),
    ("DEEP", "ГЛУБИНА"),
    ("POSITIONS", "ПОЗИЦИЙ"),
    ("Figure Deleted", "Рисунок удалён"),
    ("VIEW", "ВИД"),
    ("FROM CENTER", "ОТ ЦЕНТРА"),
    # ─── Part 4 drawing annotation terms (chrome plating / dimensions) ───
    ("FULL CHROME PLATING THICKNESS", "ПОЛНАЯ ТОЛЩИНА ХРОМОВОГО ПОКРЫТИЯ"),
    ("CHROME PLATING DEPOSIT", "ХРОМОВОЕ ПОКРЫТИЕ"),
    ("CHROME PLATING", "ХРОМОВОЕ ПОКРЫТИЕ"),
    ("PLATING THICKNESS", "ТОЛЩИНА ПОКРЫТИЯ"),
    ("PLATING LIMIT", "ПРЕДЕЛ ПОКРЫТИЯ"),
    ("PAINT DEPOSIT OVERLAP", "ПЕРЕКРЫТИЕ СЛОЯ КРАСКИ"),
    ("ZINC-NICKEL DEPOSIT OVERLAP", "ПЕРЕКРЫТИЕ СЛОЯ ЦИНК-НИКЕЛЯ"),
    ("ZINC-NICKEL", "ЦИНК-НИКЕЛЬ"),
    ("DEPOSIT OVERLAP", "ПЕРЕКРЫТИЕ СЛОЯ"),
    ("DIA. AFTER GRINDING CHROME", "ДИАМ. ПОСЛЕ ШЛИФОВАНИЯ ХРОМА"),
    ("DIA. AFTER CHROME PLATING", "ДИАМ. ПОСЛЕ ХРОМИРОВАНИЯ"),
    ("AFTER GRINDING CHROME", "ПОСЛЕ ШЛИФОВАНИЯ ХРОМА"),
    ("AFTER CHROME PLATING", "ПОСЛЕ ХРОМИРОВАНИЯ"),
    ("AFTER GRINDING", "ПОСЛЕ ШЛИФОВАНИЯ"),
    ("CHROME TERMINATION", "ОКОНЧАНИЕ ХРОМА"),
    ("CHROME RUN OUT BAND", "ЗОНА ВЫХОДА ХРОМА"),
    ("RUN OUT BAND", "ЗОНА ВЫХОДА"),
    ("RUN OUT", "ВЫХОД"),
    ("BAND", "ЗОНА"),
    ("MAX.", "МАКС."),
    ("SLIDING TUBE", "СКОЛЬЗЯЩАЯ ТРУБА"),
    ("MAIN FITTING", "КОРПУС СТОЙКИ"),
    ("CHROME", "ХРОМ"),
    ("GRINDING", "ШЛИФОВАНИЕ"),
    ("PLATING", "ПОКРЫТИЕ"),
    ("DEPOSIT", "СЛОЙ"),
    ("OVERLAP", "ПЕРЕКРЫТИЕ"),
    ("TERMINATION", "ОКОНЧАНИЕ"),
    ("SMOOTH TRANSITION", "ПЛАВНЫЙ ПЕРЕХОД"),
    ("TRANSITION", "ПЕРЕХОД"),
    ("EDGE BLENDED", "КРОМКА СОПРЯЖЕНА"),
    ("WAVY OR IRREGULAR LINE PERMISSIBLE", "ВОЛНИСТАЯ ИЛИ НЕРОВНАЯ ЛИНИЯ ДОПУСТИМА"),
    ("IRREGULAR", "НЕРОВНАЯ"),
    ("WAVY", "ВОЛНИСТАЯ"),
    ("PERMISSIBLE", "ДОПУСТИМА"),
    ("WORKING DIA.", "РАБОЧИЙ ДИАМ."),
    ("PROUD OF", "ВЫСТУПАТЬ НАД"),
    ("REMAIN ON", "ОСТАВАТЬСЯ НА"),
    ("REMAIN", "ОСТАВАТЬСЯ"),
    ("THRU BORES", "СКВОЗНЫХ ОТВЕРСТИЙ"),
    ("THRU BORE", "СКВОЗНОГО ОТВЕРСТИЯ"),
    ("CROSS BORE", "ПОПЕРЕЧНОЕ ОТВЕРСТИЕ"),
    ("CROSS BOLT", "ПОПЕРЕЧНЫЙ БОЛТ"),
    ("CROSS BORES", "ПОПЕРЕЧНЫЕ ОТВЕРСТИЯ"),
    ("BORE", "ОТВЕРСТИЕ"),
    ("SPOTFACE RADIUS", "РАДИУС ЗЕНКОВКИ"),
    ("SPOTFACE", "ЗЕНКОВКА"),
    ("CHAMFERS", "ФАСОК"),
    ("RADIUS", "РАДИУС"),
    ("FLANGE FACE", "ТОРЕЦ ФЛАНЦА"),
    ("FLANGE", "ФЛАНЕЦ"),
    ("LUGS", "ПРОУШИН"),
    ("LUG", "ПРОУШИНА"),
    ("KNUCKLE", "КУЛАК"),
    ("TOOLING", "ТЕХНОЛОГИЧЕСКАЯ"),
    ("GREASE", "СМАЗОЧНЫЕ"),
    ("BARREL OUTER", "НАРУЖНЫЙ КОРПУС ЦИЛИНДРА"),
    ("BARREL", "КОРПУС ЦИЛИНДРА"),
    ("RETRACTION", "УБОРКИ"),
    ("JOURNAL", "ШЕЙКА"),
    ("ABUTMENT", "УПОР"),
    ("SEAL", "УПЛОТНЕНИЕ"),
    ("MANIFOLD", "КОЛЛЕКТОР"),
    ("BRAKE", "ТОРМОЗНОЙ"),
    ("AXLE", "ОСЬ"),
    ("OUTER", "НАРУЖНЫЙ"),
    ("INNER", "ВНУТРЕННИЙ"),
    ("UPPER", "ВЕРХНИЙ"),
    ("LOWER", "НИЖНИЙ"),
    ("EXTERNAL", "НАРУЖНЫЙ"),
    ("INTERNAL", "ВНУТРЕННИЙ"),
    ("LIMIT", "ПРЕДЕЛ"),
    ("EXTENT", "ПРОТЯЖЁННОСТЬ"),
    ("FINE", "ЧИСТОВОЙ"),
    ("ANYWHERE", "В ЛЮБОМ МЕСТЕ"),
    ("BEYOND", "ЗА"),
    ("INCLUDING", "ВКЛЮЧАЯ"),
    ("INCLUDED", "ВКЛЮЧИТЕЛЬНО"),
    ("MIN.", "МИН."),
    ("RAD.", "РАД."),
    ("DIAMETERS", "ДИАМЕТРЫ"),
    ("REFER TO FIGURE", "ОБРАТИТЕСЬ К РИСУНКУ"),
    ("REFER TO", "ОБРАТИТЕСЬ К"),
    ("TRANSFER BLOCK", "ПЕРЕХОДНЫЙ БЛОК"),
    ("TORQUE LINK", "ШЛИЦ-ШАРНИР"),
    ("RETAINING PIN", "СТОПОРНЫЙ ШТИФТ"),
    ("DIAPHRAGM TUBE", "ДИАФРАГМЕННАЯ ТРУБА"),
    ("CARDAN", "КАРДАН"),
    ("PINTLE", "ШТИФТ НАВЕСА"),
    ("CHANGE OVER VALVE", "ПЕРЕПУСКНОЙ КЛАПАН"),
    ("AFTER THREAD", "ПОСЛЕ РЕЗЬБЫ"),
    ("AFTER", "ПОСЛЕ"),
    ("THREAD", "РЕЗЬБА"),
    ("FACE", "ПОВЕРХНОСТЬ"),

    # ─── Protective treatment procedural phrases (for table cells and paragraphs) ───
    ("Apply cadmium plate all over but not to", "Нанесите кадмиевое покрытие повсюду, кроме"),
    ("Apply cadmium plate all over but not on", "Нанесите кадмиевое покрытие повсюду, кроме"),
    ("Apply cadmium plate, but not to the", "Нанесите кадмиевое покрытие, кроме"),
    ("Apply cadmium plate: refer to", "Нанесите кадмиевое покрытие: обратитесь к"),
    ("Apply cadmium plate", "Нанесите кадмиевое покрытие"),
    ("Apply primer paint only to the contact face. Paint must not go in the bores.",
     "Нанесите грунтовочную краску только на контактную поверхность. Краска не должна попадать в отверстия."),
    ("Apply primer paint only to the areas", "Нанесите грунтовочную краску только на участки"),
    ("Apply primer paint only to", "Нанесите грунтовочную краску только на"),
    ("Apply primer paint to the areas", "Нанесите грунтовочную краску на участки"),
    ("Apply primer paint to", "Нанесите грунтовочную краску на"),
    ("Apply primer paint", "Нанесите грунтовочную краску"),
    ("Apply paint all over but not on", "Нанесите краску повсюду, кроме"),
    ("Apply paint all over but not to", "Нанесите краску повсюду, кроме"),
    ("Apply paint to the areas", "Нанесите краску на участки"),
    ("Apply paint to", "Нанесите краску на"),
    ("Apply sermetel W only to the areas", "Нанесите Sermetel W только на участки"),
    ("Apply sermetel W to the areas", "Нанесите Sermetel W на участки"),
    ("Apply only primer paint to areas", "Нанесите только грунтовочную краску на участки"),
    ("Apply only primer paint to", "Нанесите только грунтовочную краску на"),
    ("Do not paint areas", "Не красьте участки"),
    ("Do not paint the screw threads", "Не красьте резьбу"),
    ("not paint areas", "не красьте участки"),
    ("not paint the screw threads", "не красьте резьбу"),
    ("not cadmium plate:", "не кадмировать:"),
    ("Protective treatment is not necessary", "Защитная обработка не требуется"),
    ("Chromic acid anodise all over: refer to", "Хромовокислотное анодирование повсюду: обратитесь к"),
    ("Chromic acid anodise all over", "Хромовокислотное анодирование повсюду"),
    ("Passivate: refer to", "Пассивировать: обратитесь к"),
    ("The cadmium plate must overlap the chromium plate run out. Bare metal not permitted.",
     "Кадмиевое покрытие должно перекрывать выход хромового покрытия. Открытый металл не допускается."),
    ("The cadmium plate must overlap the chromium plate run out.",
     "Кадмиевое покрытие должно перекрывать выход хромового покрытия."),
    ("cadmium plate must overlap", "кадмиевое покрытие должно перекрывать"),
    ("chromium plate run out", "выход хромового покрытия"),
    ("Bare metal not permitted", "Открытый металл не допускается"),
    ("The Sermetel W coating must overlap the chromium plated areas and cadmium plated areas.",
     "Покрытие Sermetel W должно перекрывать хромированные и кадмированные участки."),
    ("Sermetel W coating thickness between", "толщина покрытия Sermetel W между"),
    ("Sermetel W coating must overlap", "покрытие Sermetel W должно перекрывать"),
    ("Sermetel is optional in areas", "Sermetel не является обязательным на участках"),
    ("If sermetel is not applied in areas", "Если Sermetel не нанесён на участки"),
    ("apply cadmium plate to areas", "нанесите кадмиевое покрытие на участки"),
    ("Make the sermetel W coating thickness between", "Выполните толщину покрытия Sermetel W между"),
    ("areas that have chromium plate", "участки с хромовым покрытием"),
    ("area that has chromium plate", "участок с хромовым покрытием"),
    ("chromium plated areas", "хромированные участки"),
    ("chromium plated area", "хромированный участок"),
    ("cadmium plated areas", "кадмированные участки"),
    ("the split pin hole", "отверстие для шплинта"),
    ("the thread and undercut", "резьбу и подрезку"),
    ("the thread", "резьбу"),
    ("the threads", "резьбы"),
    ("the threaded surfaces", "резьбовые поверхности"),
    ("the axial hole and chamfers", "осевое отверстие и фаски"),
    ("the two radial holes", "два радиальных отверстия"),
    ("diameter areas around the holes", "участки диаметром вокруг отверстий"),
    ("on the inside face of one flange", "на внутренней поверхности одного фланца"),
    ("the areas A and D", "участки A и D"),
    ("the areas A and B", "участки A и B"),
    ("the areas A, B, C", "участки A, B, C"),
    ("the areas", "участки"),
    ("the area B", "участок B"),
    ("the area", "участок"),
    ("areas A and B", "участки A и B"),
    ("the bearings, bush bores and flanges", "подшипники, отверстия под втулки и фланцы"),
    ("Before installation of bushes:", "Перед установкой втулок:"),
    ("After installation of bushes:", "После установки втулок:"),
    ("but not to the bushes", "но не на втулки"),
    ("the 3 holes in face D", "3 отверстия в поверхности D"),
    ("including chamfer", "включая фаску"),
    ("where identified on Figure", "как указано на рисунке"),
    ("to the areas where Sermetel W is applied", "на участках, где нанесён Sermetel W"),
    ("to the chromium plated areas", "на хромированных участках"),
    ("to the threaded surfaces", "на резьбовых поверхностях"),
    ("where the lubrication adaptors", "где устанавливаются смазочные адаптеры"),
    ("will install", "будут установлены"),
    ("all over but not to", "повсюду, кроме"),
    ("all over but not on", "повсюду, кроме"),
    ("all over", "повсюду"),
    ("but not to", "кроме"),
    ("but not on", "кроме"),
    ("the face that touches the wheel bearings", "поверхность, касающуюся подшипников колеса"),
    ("the screw threads and", "резьбу и"),

    # ─── Standalone protective treatment terms ───
    # These MUST come BEFORE the bare part names section to ensure
    # "cadmium plate" → "кадмиевое покрытие" runs before "plate" → "пластину"
    ("cadmium plate", "кадмиевое покрытие"),
    ("Cadmium plate", "Кадмиевое покрытие"),
    ("chromium plate", "хромовое покрытие"),
    ("Chromium plate", "Хромовое покрытие"),
    ("primer paint", "грунтовочную краску"),
    ("Primer paint", "Грунтовочную краску"),
    ("Refer to Figure", "Обратитесь к рисунку"),
    ("refer to Figure", "обратитесь к рисунку"),
    ("Refer to", "Обратитесь к"),
    ("refer to", "обратитесь к"),
    ("Apply Alocrom", "Нанесите Alocrom"),
    ("apply Alocrom", "нанесите Alocrom"),
    ("Apply a light coat of", "Нанесите тонкий слой"),
    ("Apply a coat of", "Нанесите слой"),
    ("Do not include", "Не включайте"),
    ("do not include", "не включайте"),
    ("Do not apply", "Не наносите"),
    ("do not apply", "не наносите"),
    ("Anodise", "Анодируйте"),
    ("Anodize", "Анодируйте"),
    ("anodise", "анодируйте"),
    ("anodize", "анодируйте"),
    ("Passivate", "Пассивируйте"),
    ("passivate", "пассивируйте"),
    ("spotfaces", "зенковки"),
    ("Spotfaces", "Зенковки"),
    ("spotface", "зенковку"),
    ("Spotface", "Зенковку"),
    ("externally", "снаружи"),
    ("Externally", "Снаружи"),
    ("internally", "изнутри"),
    ("Internally", "Изнутри"),
    ("thickness", "толщиной"),
    ("Thickness", "Толщиной"),
    ("thick", "толщиной"),
    ("light coat", "тонкий слой"),
    ("primer", "грунтовку"),
    ("paint", "краску"),
    ("permitted", "допускается"),
    ("Permitted", "Допускается"),
    ("optional", "не является обязательным"),
    ("Optional", "Не является обязательным"),
    # Word-level entries for remaining PT text translation
    # IMPORTANT: These MUST come after longer phrase entries above
    ("Aluminium alloy", "Алюминиевый сплав"),
    ("aluminium alloy", "алюминиевый сплав"),
    ("Stainless steel", "Нержавеющая сталь"),
    ("stainless steel", "нержавеющая сталь"),
    ("a thin coat of", "тонкий слой"),
    ("thin coat", "тонкий слой"),
    ("but not the", "кроме"),
    ("but not", "кроме"),
    ("not paint", "не красьте"),
    ("not include", "не включайте"),
    ("Apply only", "Нанесите только"),
    ("apply only", "нанесите только"),
    ("Apply", "Нанесите"),
    ("apply", "нанесите"),
    ("Paint", "Покрасьте"),
    ("only to", "только на"),
    ("only", "только"),
    ("areas", "участки"),
    ("area", "участок"),
    ("holes", "отверстия"),
    ("hole", "отверстие"),
    ("bores", "отверстия"),
    ("bore", "отверстие"),
    (" faces", " поверхности"),
    (" face", " поверхность"),
    ("flanges", "фланцы"),
    ("flange", "фланец"),
    ("Class", "Класс"),
    ("class", "класс"),
    ("Type", "Тип"),
    ("threaded", "резьбовой"),
    (" thread", " резьбу"),
    ("diameter", "диаметр"),
    ("shank", "стержень"),
    ("shanks", "стержни"),
    ("adjacent", "прилегающий"),
    ("head of the", "головку"),
    (" head", " головку"),
    ("below", "ниже"),
    ("above", "выше"),
    ("undercut", "подрез"),
    (" bore ", " отверстие "),
    (" bore.", " отверстие."),
    # Zinc-nickel plating / paint color terms
    ("Zinc-nickel plate", "цинково-никелевое покрытие"),
    ("zinc-nickel plate", "цинково-никелевое покрытие"),
    ("Zinc-nickel", "цинково-никелевое"),
    ("zinc-nickel", "цинково-никелевое"),
    ("extend onto", "распространяться на"),
    ("extend on", "распространяться на"),
    ("chrome surfaces", "хромовые поверхности"),
    ("completely covering", "полностью покрывая"),
    ("base material", "основной материал"),
    ("three coat process", "трёхслойный процесс"),
    ("two coat", "двухслойный"),
    ("three coat", "трёхслойный"),
    ("over layer", "верхний слой"),
    ("red color", "красного цвета"),
    ("green color", "зелёного цвета"),
    ("color", "цвет"),
    ("as given", "как указано"),
    ("where shown", "как показано"),
    ("split line", "линия разъёма"),
    ("per-mitted", "допускается"),
    ("per- mitted", "допускается"),
    (" outer ", " наружный "),
    (" after ", " после "),
    ("split moulds", "разъёмные формы"),
    ("split mould", "разъёмную форму"),
    ("opposite race", "противоположной дорожки"),
    ("inside moulds", "внутри форм"),
    ("inside", "внутри"),
    ("is not", "не"),
    ("moulds", "формы"),
    ("mould", "форму"),
    ("race", "дорожки"),
    ("for lubrication", "для смазки"),

    # ─── Part 5 repair procedural phrases ───
    # Long phrases first (MUST come before word-level)
    ("Identify the part with the Messier-Dowty Limited repair number",
     "Идентифицируйте деталь ремонтным номером Messier-Dowty Limited"),
    ("Identify the part with the Messier-Dowty Ltd repair number",
     "Идентифицируйте деталь ремонтным номером Messier-Dowty Ltd"),
    ("adjacent to the part number", "рядом с номером детали"),
    ("Examine the part to make sure that you have obeyed all the repair instructions correctly.",
     "Осмотрите деталь, чтобы убедиться, что все инструкции по ремонту выполнены правильно."),
    ("Examine the part to make sure that you have obeyed the repair instructions correctly.",
     "Осмотрите деталь, чтобы убедиться, что инструкции по ремонту выполнены правильно."),
    ("make sure that you have obeyed all the repair instructions correctly",
     "убедитесь, что все инструкции по ремонту выполнены правильно"),
    ("make sure that you have obeyed the repair instructions correctly",
     "убедитесь, что инструкции по ремонту выполнены правильно"),
    ("the surface finish must be", "чистота поверхности должна быть"),
    ("The surface finish must be", "Чистота поверхности должна быть"),
    ("with a surface finish of", "с чистотой поверхности"),
    ("a surface finish of", "чистотой поверхности"),
    ("Machine a surface finish of", "Обработайте до чистоты поверхности"),
    ("surface finish", "чистота поверхности"),
    ("micro-inches", "микродюймов"),
    ("micrometers", "микрометров"),
    # Chromium/cadmium plating procedural
    ("Remove the chromium plate from diameter", "Удалите хромовое покрытие с диаметра"),
    ("Remove the chromium plate from diameters", "Удалите хромовое покрытие с диаметров"),
    ("Remove the chromium plate only from the diameter", "Удалите хромовое покрытие только с диаметра"),
    ("Remove the chromium plate from", "Удалите хромовое покрытие с"),
    ("Remove the cadmium plate from", "Удалите кадмиевое покрытие с"),
    ("Apply chromium plate to the diameter", "Нанесите хромовое покрытие на диаметр"),
    ("Apply chromium plate to the diameters", "Нанесите хромовое покрытие на диаметры"),
    ("Apply chromium plate to diameter", "Нанесите хромовое покрытие на диаметр"),
    ("Apply chromium plate to", "Нанесите хромовое покрытие на"),
    ("Apply cadmium plate all over the pin", "Нанесите кадмиевое покрытие на весь штифт"),
    ("Apply cadmium plate to the machined areas but not where chromium plated",
     "Нанесите кадмиевое покрытие на обработанные участки, кроме хромированных"),
    ("Apply cadmium plate to the machined areas", "Нанесите кадмиевое покрытие на обработанные участки"),
    ("Apply cadmium plate to all surfaces of", "Нанесите кадмиевое покрытие на все поверхности"),
    ("Apply cadmium plate to the outside diameter", "Нанесите кадмиевое покрытие на наружный диаметр"),
    ("Locally apply cadmium plate to the reworked areas", "Локально нанесите кадмиевое покрытие на доработанные участки"),
    ("Locally apply cadmium plate", "Локально нанесите кадмиевое покрытие"),
    ("with a minimum chromium plate thickness of", "с минимальной толщиной хромового покрытия"),
    ("The thickness of the cadmium plate must be between", "Толщина кадмиевого покрытия должна быть между"),
    ("chromium plate terminations are smooth", "окончания хромового покрытия гладкие"),
    ("chromium plate termination information", "информацию об окончании хромового покрытия"),
    ("chromium plate termination", "окончание хромового покрытия"),
    ("minimum chromium plate thickness", "минимальная толщина хромового покрытия"),
    ("the chromium plated areas", "хромированные участки"),
    ("the chromium plated surface", "хромированную поверхность"),
    ("to show the parent metal", "до основного металла"),
    ("but not on the chromium plated areas", "кроме хромированных участков"),
    ("but not where chromium plated", "кроме хромированных участков"),
    ("after removal of the chromium plate", "после удаления хромового покрытия"),
    ("Spray primer paint lightly on the cadmium plated surface",
     "Слегка нанесите распылением грунтовочную краску на кадмированную поверхность"),
    ("Paint the pin all over, but not on the threads, the thread undercut and the chromium plated areas",
     "Покрасьте штифт полностью, кроме резьб, подрезов резьбы и хромированных участков"),
    # Machining procedural
    ("Machine the diameter A to remove the minimum amount of material necessary to remove the damage or wear",
     "Обработайте диаметр A, снимая минимально необходимое количество материала для устранения повреждения или износа"),
    ("Machine the diameter A sufficiently to remove the damage or wear",
     "Обработайте диаметр A в достаточной степени для устранения повреждения или износа"),
    ("Machine diameter A sufficiently to remove the damage or wear",
     "Обработайте диаметр A в достаточной степени для устранения повреждения или износа"),
    ("Machine diameter A just sufficiently to remove the damage or corrosion",
     "Обработайте диаметр A ровно настолько, чтобы устранить повреждение или коррозию"),
    ("Machine diameter A to remove the damage or wear after removal of the chromium plate",
     "Обработайте диаметр A для устранения повреждения или износа после удаления хромового покрытия"),
    ("Machine diameter A to remove damage and wear",
     "Обработайте диаметр A для устранения повреждения и износа"),
    ("Machine diameter A sufficiently to remove",
     "Обработайте диаметр A в достаточной степени для устранения"),
    ("Machine the diameter and width of the repair liner to the dimensions",
     "Обработайте диаметр и ширину ремонтного вкладыша до размеров"),
    ("Machine the diameter", "Обработайте диаметр"),
    ("Machine diameter", "Обработайте диаметр"),
    ("Machine the face Q", "Обработайте поверхность Q"),
    ("Machine the chamfer to the dimensions shown", "Обработайте фаску до размеров, показанных"),
    ("Machine the radii to the dimensions as shown", "Обработайте радиусы до размеров, как показано"),
    ("Machine the repair bushes to the dimensions shown and calculated",
     "Обработайте ремонтные втулки до показанных и рассчитанных размеров"),
    ("Machine the oversize bearing to the dimensions shown and calculated",
     "Обработайте ремонтный подшипник до показанных и рассчитанных размеров"),
    ("Machine the repair bushes to the contour of the pin",
     "Обработайте ремонтные втулки по контуру штифта"),
    ("Machine the bores of the repair sleeves", "Обработайте отверстия ремонтных рукавов"),
    ("Machine (do not grind) the sulphamate nickel plate",
     "Обработайте (не шлифуйте) сульфаматное никелевое покрытие"),
    ("Machine the", "Обработайте"),
    ("Machine ", "Обработайте "),
    ("to the dimensions shown and calculated", "до показанных и рассчитанных размеров"),
    ("to the dimensions shown in Figure 601", "до размеров, показанных на рисунке 601"),
    ("to the dimensions shown", "до показанных размеров"),
    ("to the dimensions given in Figure 601", "до размеров, указанных на рисунке 601"),
    ("to the dimensions given in Figure", "до размеров, указанных на рисунке"),
    ("to the dimensions", "до размеров"),
    ("to remove the minimum amount of material necessary", "снимая минимально необходимое количество материала"),
    ("to remove the damage or wear", "для устранения повреждения или износа"),
    ("to remove damage or wear", "для устранения повреждения или износа"),
    ("to remove the damage or corrosion", "для устранения повреждения или коррозии"),
    ("to remove damage and wear", "для устранения повреждения и износа"),
    # Examination/Inspection
    ("Examine the pin for flaws", "Проверьте штифт на наличие дефектов"),
    ("Examine the pivot pin for flaws", "Проверьте штифт вращения на наличие дефектов"),
    ("Examine the uplock pin for flaws", "Проверьте штифт замка убранного положения на наличие дефектов"),
    ("Examine the bracket for flaws", "Проверьте кронштейн на наличие дефектов"),
    ("Examine the part for flaws", "Проверьте деталь на наличие дефектов"),
    ("Examine the ground chromium plate for flaws", "Проверьте шлифованное хромовое покрытие на наличие дефектов"),
    ("Examine the chromium plated surface for flaws", "Проверьте хромированную поверхность на наличие дефектов"),
    ("Examine the machined areas for flaws", "Проверьте обработанные участки на наличие дефектов"),
    ("Examine the repair bushes for flaws", "Проверьте ремонтные втулки на наличие дефектов"),
    ("Examine the edges of sulphamate nickel plate", "Проверьте кромки сульфаматного никелевого покрытия"),
    ("Examine the part", "Осмотрите деталь"),
    ("for flaws", "на наличие дефектов"),
    ("inclusion class", "класс включений"),
    # Shot peen / grind
    ("Shot peen the machined areas", "Выполните дробеструйную обработку обработанных участков"),
    ("Shot peen the machined area", "Выполните дробеструйную обработку обработанного участка"),
    ("Shot peen the reworked areas", "Выполните дробеструйную обработку доработанных участков"),
    ("Shot peen the pin", "Выполните дробеструйную обработку штифта"),
    ("Shot peen diameter", "Выполните дробеструйную обработку диаметра"),
    ("Shot peen and apply chromium plate to diameter",
     "Выполните дробеструйную обработку и нанесите хромовое покрытие на диаметр"),
    ("Shot peen", "Выполните дробеструйную обработку"),
    ("shot peen", "дробеструйную обработку"),
    ("Finish grind the pin", "Выполните чистовую шлифовку штифта"),
    ("Finish grind", "Выполните чистовую шлифовку"),
    ("Grind the diameter", "Шлифуйте диаметр"),
    ("Grind diameter", "Шлифуйте диаметр"),
    ("Grit blast the reworked areas", "Выполните дробеструйную очистку доработанных участков"),
    ("Grit blast the sulphamate nickel area", "Выполните дробеструйную очистку участка сульфаматного никеля"),
    ("Grit blast", "Выполните дробеструйную очистку"),
    # Liner/adhesive repair phrases
    ("Remove all of the adhesive from the external diameter of the gland housing and the internal diameter of the liner",
     "Удалите весь адгезив с наружного диаметра корпуса сальника и внутреннего диаметра вкладыша"),
    ("Remove the used adhesive from the external diameter of the gland housing and the internal diameter of the liner",
     "Удалите использованный адгезив с наружного диаметра корпуса сальника и внутреннего диаметра вкладыша"),
    ("Remove the used adhesive from the external diameter of the gland housing",
     "Удалите использованный адгезив с наружного диаметра корпуса сальника"),
    ("Remove all of the adhesive from the external diameter of the gland housing",
     "Удалите весь адгезив с наружного диаметра корпуса сальника"),
    ("Remove the damaged liner and remove all of the adhesive from the external diameter of the gland housing",
     "Снимите повреждённый вкладыш и удалите весь адгезив с наружного диаметра корпуса сальника"),
    ("Remove the damaged liner and remove the used adhesive from the external diameter of the gland housing",
     "Снимите повреждённый вкладыш и удалите использованный адгезив с наружного диаметра корпуса сальника"),
    ("Temporarily put the liner in position on the gland housing",
     "Временно установите вкладыш на корпус сальника"),
    ("Temporarily put the repair liner in position on the gland housing",
     "Временно установите ремонтный вкладыш на корпус сальника"),
    ("Put the liner in position on the gland housing",
     "Установите вкладыш на корпус сальника"),
    ("Put the repair liner in position on the gland housing",
     "Установите ремонтный вкладыш на корпус сальника"),
    ("Remove the liner from the gland housing", "Снимите вкладыш с корпуса сальника"),
    ("Remove the repair liner from the gland housing", "Снимите ремонтный вкладыш с корпуса сальника"),
    ("Clean the contact surfaces of the liner and the gland housing",
     "Очистите контактные поверхности вкладыша и корпуса сальника"),
    ("Clean the surfaces of the liner and the gland housing that will touch",
     "Очистите поверхности вкладыша и корпуса сальника, которые будут соприкасаться"),
    ("Apply adhesive, Material Ref. Item", "Нанесите адгезив, поз. ссылки материала"),
    ("Apply adhesive PVC tape, Material Ref. Item", "Нанесите клейкую ПВХ-ленту, поз. ссылки материала"),
    ("Assemble the liner to the gland housing", "Соберите вкладыш с корпусом сальника"),
    ("Assemble the repair liner to the gland housing", "Соберите ремонтный вкладыш с корпусом сальника"),
    ("Clamp the repair liner to the gland housing using an applicable tool",
     "Прижмите ремонтный вкладыш к корпусу сальника с помощью подходящего инструмента"),
    ("Put the gland housing in the preheated oven",
     "Поместите корпус сальника в предварительно нагретую печь"),
    ("Put the gland housing in the oven",
     "Поместите корпус сальника в печь"),
    ("Remove the gland housing from the oven",
     "Извлеките корпус сальника из печи"),
    ("Remove all of the tape and clean the parts as necessary",
     "Удалите всю ленту и очистите детали по необходимости"),
    ("If necessary, cut the liner at the scarf joints to adjust its length",
     "При необходимости подрежьте вкладыш в местах стыков для подгонки длины"),
    ("If necessary, cut the repair liner at the scarf joints to adjust its length",
     "При необходимости подрежьте ремонтный вкладыш в местах стыков для подгонки длины"),
    ("Measure the diameter of the gland housing across the contact area",
     "Измерьте диаметр корпуса сальника по контактной зоне"),
    ("to the contact surface of the liner near to the scarf joints",
     "на контактную поверхность вкладыша вблизи стыков"),
    ("to the contact surface of the gland housing",
     "на контактную поверхность корпуса сальника"),
    ("to the cleaned surface of the liner near to the scarf joints",
     "на очищенную поверхность вкладыша вблизи стыков"),
    ("Preheat an oven to between", "Предварительно нагрейте печь до"),
    ("Set the temperature of an oven to between", "Установите температуру печи от"),
    ("Prepare a surface treatment mixture of 1 part by volume of Accomet C",
     "Приготовьте смесь для обработки поверхности из 1 части по объёму Accomet C"),
    ("Prepare a mixture of 1 part by volume of Accomet C",
     "Приготовьте смесь из 1 части по объёму Accomet C"),
    ("parts by volume of clean cold water",
     "частей по объёму чистой холодной воды"),
    ("Use a brush to apply a smooth layer of the prepared surface treatment mixture to the contact surface of the gland housing",
     "Кисточкой нанесите ровный слой приготовленной смеси для обработки поверхности на контактную поверхность корпуса сальника"),
    ("Use a brush to apply a flat layer of the prepared mixture to the surfaces of the gland housing made rough",
     "Кисточкой нанесите ровный слой приготовленной смеси на зашкуренные поверхности корпуса сальника"),
    ("Use a brush to apply Araldite, 2015", "Кисточкой нанесите Araldite, 2015"),
    ("Use a brush to apply a smooth layer of the prepared surface treatment mixture",
     "Кисточкой нанесите ровный слой приготовленной смеси для обработки поверхности"),
    ("Use a brush to apply", "Кисточкой нанесите"),
    ("for a minimum of 4 minutes and until the applied surface treatment mixture is dry",
     "минимум на 4 минуты и до высыхания нанесённой смеси для обработки поверхности"),
    ("for a minimum of 4 minutes and until the applied mixture is dry",
     "минимум на 4 минуты и до высыхания нанесённой смеси"),
    ("Alternative procedure for paragraphs", "Альтернативная процедура для пунктов"),
    ("and allow to cool for a minimum of 30 minutes",
     "и дайте остыть минимум 30 минут"),
    ("allow to cool for a minimum of", "дайте остыть минимум"),
    ("and let its temperature decrease for a minimum of 30 minutes",
     "и дайте температуре снизиться минимум на 30 минут"),
    ("then remove the adhesive PVC tape", "затем удалите клейкую ПВХ-ленту"),
    ("the cleaning tissues, Material Ref. Item", "чистящие салфетки, поз. ссылки материала"),
    ("the cleaning agent, Material Ref. Item", "чистящее средство, поз. ссылки материала"),
    ("cleaning tissues, Material Ref. Item", "чистящие салфетки, поз. ссылки материала"),
    ("cleaning agent, Material Ref. Item", "чистящее средство, поз. ссылки материала"),
    ("Material Ref. Item", "поз. ссылки материала"),
    ("to roughen the surfaces to be bonded", "для зашкуривания поверхностей склеивания"),
    ("to make rough, the surfaces that will bond", "для зашкуривания поверхностей склеивания"),
    ("to clean the roughened surfaces", "для очистки зашкуренных поверхностей"),
    ("to clean the surfaces made rough", "для очистки зашкуренных поверхностей"),
    ("Do not damage the edges of the adhesive PVC tape",
     "Не повредите края клейкой ПВХ-ленты"),
    ("Make sure that the edges of the adhesive PVC tape",
     "Убедитесь, что края клейкой ПВХ-ленты"),
    ("adhesive PVC tape", "клейкая ПВХ-лента"),
    ("bond tightly to the gland housing", "плотно прилегают к корпусу сальника"),
    ("bond tightly", "плотно прилегают"),
    ("around the gland housing to the sides of and touching the repair liner",
     "вокруг корпуса сальника по сторонам и в контакте с ремонтным вкладышем"),
    ("around the gland housing", "вокруг корпуса сальника"),
    ("to hold it in that position", "для фиксации в этом положении"),
    ("Use one layer of masking tape", "Используйте один слой малярной ленты"),
    ("masking tape, Material Ref. Item", "малярную ленту, поз. ссылки материала"),
    ("masking tape", "малярную ленту"),
    ("Use the Emery cloth, 60-100 grit", "Используйте наждачную ткань, зернистость 60-100"),
    ("Emery cloth, 60-100 grit", "наждачную ткань, зернистость 60-100"),
    ("Emery cloth", "наждачную ткань"),
    ("scarf joints", "стыки"),
    ("contact surface", "контактная поверхность"),
    ("contact surfaces", "контактные поверхности"),
    ("contact area", "контактная зона"),
    ("repair liner", "ремонтный вкладыш"),
    ("the liner", "вкладыш"),
    ("liner", "вкладыш"),
    # Bearing/bush installation repair phrases
    ("Use the Press Pad", "Используйте Press Pad"),
    ("and Drift", "и Drift"),
    ("to install the oversize bearing to the bracket",
     "для установки ремонтного подшипника в кронштейн"),
    ("Check line ream the oversize bearing", "Проверьте линейную развёртку ремонтного подшипника"),
    ("line ream the oversize bearing", "линейную развёртку ремонтного подшипника"),
    ("line ream", "линейную развёртку"),
    ("Apply sealant, Material Ref. Item", "Нанесите герметик, поз. ссылки материала"),
    ("Apply adhesive, Material Ref. Item", "Нанесите адгезив, поз. ссылки материала"),
    ("to seal between the ends of the oversize bearing and the bracket",
     "для герметизации между торцами ремонтного подшипника и кронштейном"),
    ("Passivate the oversize bearing", "Пассивируйте ремонтный подшипник"),
    ("Apply cadmium plate to the oversize bearing",
     "Нанесите кадмиевое покрытие на ремонтный подшипник"),
    ("but not on the internal diameter", "кроме внутреннего диаметра"),
    ("Prepare the machined surfaces of the bracket",
     "Подготовьте обработанные поверхности кронштейна"),
    ("anodize the surfaces", "анодируйте поверхности"),
    ("apply Alocrom to the surfaces", "нанесите Alocrom на поверхности"),
    ("apply Alocrom to the contact surface", "нанесите Alocrom на контактную поверхность"),
    ("apply Alocrom to the machined areas", "нанесите Alocrom на обработанные участки"),
    ("Apply Alocrom, Material Ref. Item", "Нанесите Alocrom, поз. ссылки материала"),
    ("to the machined areas", "на обработанные участки"),
    ("Calculate the diameter B and the dimension C of the oversize bearing",
     "Рассчитайте диаметр B и размер C ремонтного подшипника"),
    ("Calculate diameter C of the repair bushes", "Рассчитайте диаметр C ремонтных втулок"),
    ("Calculate diameter B of the repair sleeves", "Рассчитайте диаметр B ремонтных рукавов"),
    ("Measure and make a record of the new diameter A and the thickness of the lug D",
     "Измерьте и запишите новый диаметр A и толщину проушины D"),
    ("Measure and make a record of the new diameter", "Измерьте и запишите новый диаметр"),
    ("Measure and record the new diameters", "Измерьте и запишите новые диаметры"),
    ("Measure the new diameter", "Измерьте новый диаметр"),
    ("Install the repair bushes to diameter", "Установите ремонтные втулки на диаметр"),
    ("Install the repair sleeves to diameters", "Установите ремонтные рукава на диаметры"),
    ("until aligned with the outside diameter of the pin",
     "до совмещения с наружным диаметром штифта"),
    ("to the outside diameter of the repair bushes", "на наружный диаметр ремонтных втулок"),
    ("to the outside diameter of the repair sleeves", "на наружный диаметр ремонтных рукавов"),
    ("the outside diameter and the chamfer of the repair sleeves",
     "наружный диаметр и фаску ремонтных рукавов"),
    ("oversize bearing", "ремонтный подшипник"),
    ("repair bushes", "ремонтные втулки"),
    ("repair bush", "ремонтную втулку"),
    ("repair sleeves", "ремонтные рукава"),
    ("repair sleeve", "ремонтный рукав"),
    ("sulphamate nickel plate", "сульфаматное никелевое покрытие"),
    ("Sulphamate nickel plate", "Сульфаматное никелевое покрытие"),
    ("sulphamate nickel", "сульфаматный никель"),
    ("reworked areas", "доработанные участки"),
    ("machined areas", "обработанные участки"),
    ("machined area", "обработанный участок"),
    ("the parent metal is not damaged or worn", "основной металл не повреждён и не изношен"),
    ("the parent metal is damaged or worn", "основной металл повреждён или изношен"),
    ("the bare metal is not damaged or corroded", "основной металл не повреждён и не подвержен коррозии"),
    ("the bare metal is damaged or corroded", "основной металл повреждён или подвержен коррозии"),
    ("the pin base metal has been machined", "основной металл штифта был обработан"),
    ("parent metal", "основной металл"),
    ("bare metal", "основной металл"),
    ("If the", "Если"),
    ("if the", "если"),
    ("the minimum diameter is", "минимальный диаметр составляет"),
    ("The minimum diameter is", "Минимальный диаметр составляет"),
    ("The diameter must not be less than", "Диаметр не должен быть менее"),
    ("the diameter must not be less than", "диаметр не должен быть менее"),
    ("minimum of", "минимум"),
    ("down to a minimum of", "до минимума"),
    ("is sufficiently masked", "достаточно замаскирован"),
    ("Make sure that the remainder of the pin", "Убедитесь, что остальная часть штифта"),
    ("If there is evidence of delamination, remove the sulphamate nickel plate and do the repair again",
     "При обнаружении расслоения удалите сульфаматное никелевое покрытие и повторите ремонт"),
    ("delamination", "расслоение"),
    ("The sulphamate nickel plate thickness must be sufficient to get the correct diameter after machining",
     "Толщина сульфаматного никелевого покрытия должна быть достаточной для получения правильного диаметра после обработки"),
    ("de-embrittle for", "отпуск от водородного охрупчивания в течение"),
    ("Apply paint to the reworked areas", "Нанесите краску на доработанные участки"),
    ("Apply paint", "Нанесите краску"),
    # Dimension/calculation phrases
    ("as shown in Figure 601", "как показано на рисунке 601"),
    ("as shown", "как показано"),
    ("given in Figure 601", "указанных на рисунке 601"),
    ("shown in Figure 601", "показанных на рисунке 601"),
    ("shown before you remove", "показанных до удаления"),
    ("the dimensions shown", "показанных размеров"),
    ("dimensions shown and calculated", "показанные и рассчитанные размеры"),
    ("the dimensions", "размеры"),
    ("to get the bore diameter", "для получения диаметра отверстия"),
    ("to get the correct diameter after machining", "для получения правильного диаметра после обработки"),
    ("A (as measured)", "A (измеренный)"),
    ("D (as measured)", "D (измеренный)"),
    ("as measured", "измеренный"),
    ("only the chromium plate has been repaired", "отремонтировано только хромовое покрытие"),
    ("chromium plate and parent metal have been repaired", "отремонтированы хромовое покрытие и основной металл"),
    ("after painting", "после окрашивания"),
    ("over lapping adjacent", "с перекрытием прилегающей"),
    ("angled surface by at least", "наклонной поверхности минимум на"),
    ("Almen A intensity to be", "интенсивность по Алмену A должна быть"),
    # Textbox/drawing terms for Part 5
    ("SULPHAMATE NICKEL PLATE", "СУЛЬФАМАТНОЕ НИКЕЛЕВОЕ ПОКРЫТИЕ"),
    ("AFTER CADMIUM PLATE", "ПОСЛЕ КАДМИЕВОГО ПОКРЫТИЯ"),
    ("BEFORE SULPHAMATE NICKEL PLATE", "ДО СУЛЬФАМАТНОГО НИКЕЛЕВОГО ПОКРЫТИЯ"),
    ("AFTER SULPHAMATE NICKEL PLATE", "ПОСЛЕ СУЛЬФАМАТНОГО НИКЕЛЕВОГО ПОКРЫТИЯ"),
    ("REPAIR BUSH", "РЕМОНТНАЯ ВТУЛКА"),
    ("REPAIR SLEEVE", "РЕМОНТНЫЙ РУКАВ"),
    ("DIA. B", "ДИАМ. B"),
    ("MUST NOT ENTER", "НЕ ДОЛЖНО ПОПАДАТЬ В"),
    ("LUBRICATION HOLES", "СМАЗОЧНЫЕ ОТВЕРСТИЯ"),
    ("CROSS HOLE", "ПОПЕРЕЧНОЕ ОТВЕРСТИЕ"),
    ("MINIMUM", "МИНИМУМ"),
    ("LENGTH", "ДЛИНА"),
    ("BEFORE", "ДО"),
    ("STOP", "ЗАКАНЧИВАТЬСЯ"),
    ("MUST", "ДОЛЖНО"),

    # ─── NDT table contents ───
    ("Inclusion class 4 on areas without chromium plate", "Класс включений 4 на участках без хромового покрытия"),
    ("Inclusion class 3 on areas without chromium plate", "Класс включений 3 на участках без хромового покрытия"),
    ("Inclusion class 2 on areas without chromium plate", "Класс включений 2 на участках без хромового покрытия"),
    ("on areas without chromium plate", "на участках без хромового покрытия"),
    ("Chromium plated areas", "Участки с хромовым покрытием"),
    ("- Chromium", "— Хромовые"),
    ("plated areas", "покрытые участки"),
    ("Inclusion class 4", "Класс включений 4"),
    ("Inclusion class 3", "Класс включений 3"),
    ("Inclusion class 2", "Класс включений 2"),
    ("- Inclusion class", "— Класс включений"),
    ("Aluminium Alloy", "Алюминиевый сплав"),
    ("Stainless Steel", "Нержавеющая сталь"),
    ("Steel", "Сталь"),
    ("Parts 1", "Части 1"),

    # ─── Tool listing patterns ───
    ("The Holding Fixture", "Удерживающее приспособление"),
    ("The Press Adapter", "Адаптер пресса"),
    ("The Bottom Press Adapter", "Нижний адаптер пресса"),
    ("The Lifting Bar Assembly", "Сборка подъёмной штанги"),
    ("The Spherical Bearing Locator", "Позиционер сферического подшипника"),
    ("The Pintle Location Assembly", "Сборка позиционирования навеса"),
    ("The Transport and Build Trolley", "Транспортировочно-сборочная тележка"),
    ("The Support Arms", "Опорные рычаги"),
    ("The Towing Frame", "Буксировочная рама"),
    ("The Jacking Dome Adapter", "Адаптер домкратной точки"),
    ("The Adapter", "Адаптер"),
    ("The Spacer", "Проставка"),
    ("The Location Frame", "Установочная рама"),

    # ─── Remaining patterns ───
    ("Proximity switch", "датчик приближения"),
    ("Proximity switches", "датчики приближения"),
    ("proximity switch", "датчик приближения"),
    ("proximity switches", "датчики приближения"),
    ("Grease Groove", "Канавка для смазки"),

    # ─── Bare part names (without "the") for after verb phrases ───
    # These catch part names left after "Remove the" → "Снимите" etc.
    ("split pins", "шплинты"),
    ("split pin", "шплинт"),
    ("tab washers", "стопорные шайбы"),
    ("tab washer", "стопорную шайбу"),
    ("lock washers", "стопорные шайбы"),
    ("lock washer", "стопорную шайбу"),
    ("locking plates", "стопорные пластины"),
    ("locking plate", "стопорную пластину"),
    ("lock plates", "стопорные пластины"),
    ("lock plate", "стопорную пластину"),
    ("nuts", "гайки"),
    ("nut", "гайку"),
    ("bolts", "болты"),
    ("bolt", "болт"),
    ("washers", "шайбы"),
    ("washer", "шайбу"),
    ("screws", "винты"),
    ("screw", "винт"),
    ("bushes", "втулки"),
    ("bush", "втулку"),
    ("bearings", "подшипники"),
    ("bearing", "подшипник"),
    ("seals", "уплотнения"),
    ("seal", "уплотнение"),
    ("sleeves", "втулки"),
    ("sleeve", "втулку"),
    ("spacers", "проставки"),
    ("spacer", "проставку"),
    ("shims", "прокладки"),
    ("shim", "прокладку"),
    ("brackets", "кронштейны"),
    ("bracket", "кронштейн"),
    ("retainers", "фиксаторы"),
    ("retainer", "фиксатор"),
    ("plates", "пластины"),
    ("plate", "пластину"),
    ("clamp", "хомут"),
    ("dust cap", "пылезащитную крышку"),
    ("wedge", "клин"),
    ("bung", "заглушку"),
    ("labels", "этикетки"),
    ("label", "этикетку"),
    ("pins", "штифты"),
    ("pin", "штифт"),
    ("rod", "шток"),
    ("piston", "поршень"),
    ("cylinder", "цилиндр"),
    ("housing", "корпус"),
    ("charging valves", "зарядные клапаны"),
    ("charging valve", "зарядный клапан"),
    ("lubrication adapters", "адаптеры для смазки"),
    ("lubrication adapter", "адаптер для смазки"),
    ("bonding cable", "соединительный провод"),
    ("static discharge connector", "штыревой разъём статического разряда"),
    ("wiper ring", "грязесъёмное кольцо"),
    ("sealing ring", "уплотнительное кольцо"),
    ("inner liner", "внутренний вкладыш"),
    ("joint seal", "соединительное уплотнение"),
    ("level tube", "трубку уровня"),
    ("damper", "демпфер"),
    ("wiring diagram plate", "табличку электрической схемы"),
    ("wire thread inserts", "резьбовые спиральные вставки"),
    ("wire thread insert", "резьбовую спиральную вставку"),
    ("lock indentations", "фиксирующие вдавливания"),
    # Additional bare part names
    ("compression orifice plate", "пластину отверстия сжатия"),
    ("recoil orifice plate", "пластину отверстия обратного хода"),
    ("orifice plate", "пластину отверстия"),
    ("clapper seat", "седло хлопушки"),
    ("baffle", "дефлектор"),
    ("retaining ring", "удерживающее кольцо"),
    ("retaining pins", "удерживающие штифты"),
    ("retaining pin", "удерживающий штифт"),
    ("locking pins", "стопорные штифты"),
    ("locking pin", "стопорный штифт"),
    ("locking nut", "контргайку"),
    ("locking washer", "стопорную шайбу"),
    ("cap screws", "винты с головкой"),
    ("cap screw", "винт с головкой"),
    ("target", "мишень"),
    ("Bowden cable", "трос Боудена"),
    ("cross bolts", "поперечные болты"),
    ("cross bolt", "поперечный болт"),
    ("stop rings", "стопорные кольца"),
    ("stop ring", "стопорное кольцо"),
    ("backing rings", "опорные кольца"),
    ("backing ring", "опорное кольцо"),
    ("O-ring seals", "уплотнительные кольца"),
    ("O-ring seal", "уплотнительное кольцо"),
    ("transfer dowels", "переходные штифты"),
    ("transfer dowel", "переходный штифт"),
    ("valve support", "опору клапана"),
    ("valve stem", "шток клапана"),
    ("nut subassembly", "сборку гайки"),
    ("washer subassembly", "сборку шайбы"),
    ("lock plate", "стопорную пластину"),
    ("two piece stop with inserts", "двухчастный упор со вставками"),
    ("outer race", "наружное кольцо"),
    ("ball", "шарик"),
    ("rod end assembly", "сборку наконечника тяги"),
    ("rod end", "наконечник тяги"),
    ("cardan assembly", "сборку кардана"),
    ("lock stay cardan subassembly", "сборку фиксирующего кардана"),
    ("lock stay cardan", "фиксирующий кардан"),
    ("ground stud subassembly", "сборку клеммы заземления"),
    ("inflation valve subassembly", "сборку клапана заправки"),
    ("inflation valve", "клапан заправки"),
    ("lubrication shaft subassembly", "сборку вала смазки"),
    ("lubrication fittings", "фитинги для смазки"),
    ("lubrication fitting", "фитинг для смазки"),
    ("lubrication adapters", "адаптеры для смазки"),
    ("lubrication adapter", "адаптер для смазки"),
    ("identification washers", "идентификационные шайбы"),
    ("identification washer", "идентификационную шайбу"),
    ("harness support bracket", "кронштейн крепления жгута"),
    ("harness support", "крепление жгута"),
    ("upper pivot bracket", "верхний поворотный кронштейн"),
    ("pivot bracket", "поворотный кронштейн"),
    ("grooved spherical bearing", "рифлёный сферический подшипник"),
    ("self lubricating bearing", "самосмазывающийся подшипник"),
    ("spherical bearing", "сферический подшипник"),
    ("upper bearing housing", "корпус верхнего подшипника"),
    ("lower bearing housing subassembly", "сборку корпуса нижнего подшипника"),
    ("lower bearing housing", "корпус нижнего подшипника"),
    ("lower bearing", "нижний подшипник"),
    ("common lower bearing bushes", "общие втулки нижнего подшипника"),
    ("gland housing", "корпус сальника"),
    ("upper diaphragm tube", "верхнюю диафрагменную трубу"),
    ("sliding tube", "скользящую трубу"),
    ("main fitting", "корпус стойки"),
    ("upper torque link", "верхний шлиц-шарнир"),
    ("lower torque link", "нижний шлиц-шарнир"),
    ("lower slave link", "нижнее ведомое звено"),
    ("slave link", "ведомое звено"),
    ("jacking dome", "домкратную точку"),
    ("electrical axle harness", "электрический жгут оси"),
    ("laminated shims", "набор прокладок"),
    ("laminated shim", "набор прокладок"),
    # Named subassemblies (bare, without "the")
    ("upper torque link subassembly", "сборку верхнего шлиц-шарнира"),
    ("lower torque link subassembly", "сборку нижнего шлиц-шарнира"),
    ("slave link subassembly", "сборку ведомого звена"),
    ("lower slave link subassembly", "сборку нижнего ведомого звена"),
    ("sliding tube subassembly", "сборку скользящей трубы"),
    ("main fitting subassembly", "сборку корпуса стойки"),
    ("lower bearing subassembly", "сборку нижнего подшипника"),
    ("pivot bracket subassembly", "сборку поворотного кронштейна"),
    ("bracket subassembly", "сборку кронштейна"),
    ("main landing gear leg", "стойку основного шасси"),
    ("shock absorber subassembly", "сборку амортизатора"),
    ("transfer block subassembly", "сборку переходного блока"),
    ("bolt subassembly", "сборку болта"),
    ("diaphragm subassembly", "сборку диафрагмы"),
    ("upper diaphragm tube subassembly", "сборку верхней диафрагменной трубы"),

    # ─── Testing/procedural verbs and phrases ───
    ("connect the", "подсоедините"),
    ("Connect the", "Подсоедините"),
    ("disconnect the", "отсоедините"),
    ("Disconnect the", "Отсоедините"),
    ("nitrogen supply", "источник азота"),
    ("nitrogen pressure", "давление азота"),
    ("hydraulic pressure", "гидравлическое давление"),
    ("charging valve", "зарядный клапан"),
    ("charging valves", "зарядные клапаны"),
    ("Charging Adapter", "адаптер для зарядки"),
    ("Turner Inflation Equipment", "оборудование для заправки Turner"),
    ("Crowfoot Wrench", "рожковый ключ"),
    ("loading press", "нагрузочный пресс"),
    ("load cell", "датчик нагрузки"),
    ("hand pump", "ручной насос"),
    ("power pump", "силовой насос"),
    ("test fluid", "испытательная жидкость"),
    ("test circuit", "испытательный контур"),
    ("inflation equipment", "оборудование для заправки"),
    ("hydraulic test rig", "стенд для гидравлических испытаний"),
    ("ambient temperature", "температура окружающей среды"),
    ("controlled flow", "регулируемый расход"),
    ("para ", "п. "),
    ("torque it to", "затяните с моментом"),
    ("torque to", "затяните с моментом"),
    ("refer to", "см."),
    ("Refer to", "См."),

    # ─── Bare residual words (catch orphans after phrase-level matching) ───
    ("Remove drag arm sleeve", "Снимите втулку тяги"),
    ("Remove the locking nut", "Снимите контровочную гайку"),
    ("locking nut", "контровочная гайка"),
    ("locking washer", "контровочная шайба"),
    ("locking pin", "стопорный штифт"),
    ("locking pins", "стопорные штифты"),
    ("locking", "стопорный"),
    ("spherical bearing", "сферический подшипник"),
    ("spherical", "сферический"),
    ("subassemblies", "сборки"),
    ("subassembly", "сборка"),
    ("pivot", "поворотный"),
    ("cup washers", "тарельчатые шайбы"),
    ("cup washer", "тарельчатая шайба"),
    ("rod end", "стержневой наконечник"),
    (" only)", " только)"),
    ("O-ring", "уплотнительное кольцо"),

    # ─── Part 5 remaining word-level entries ───
    # Verb starters for repair procedures
    ("Temporarily put", "Временно установите"),
    ("Temporarily", "Временно"),
    ("Clamp the", "Прижмите"),
    ("Preheat an", "Предварительно нагрейте"),
    ("Preheat", "Предварительно нагрейте"),
    ("Prepare a", "Приготовьте"),
    ("Prepare", "Приготовьте"),
    ("Identify the part", "Идентифицируйте деталь"),
    ("Identify the", "Идентифицируйте"),
    ("Identify", "Идентифицируйте"),
    ("Put the", "Поместите"),
    ("Remove all of the", "Удалите всё"),
    ("remove all of the", "удалите всё"),
    ("Examine the", "Проверьте"),
    ("examine the", "проверьте"),
    # Material names (for table cells and in-text)
    ("cleaning tissues", "чистящие салфетки"),
    ("Cleaning tissues", "Чистящие салфетки"),
    ("cleaning agent", "чистящее средство"),
    ("Cleaning agent", "Чистящее средство"),
    ("adhesive", "адгезив"),
    ("Adhesive", "Адгезив"),
    ("Sealant", "Герметик"),
    ("sealant", "герметик"),
    ("repair number", "ремонтным номером"),
    ("part number", "номером детали"),
    # Repair-specific nouns and adjectives
    ("repair liner", "ремонтный вкладыш"),
    ("gland housing", "корпус сальника"),
    ("the gland", "корпус"),
    ("gland", "корпус"),
    ("the liner", "вкладыш"),
    ("liner", "вкладыш"),
    ("scarf joints", "стыки"),
    ("the scarf", "стык"),
    ("scarf", "стык"),
    ("contact surfaces", "контактные поверхности"),
    ("contact surface", "контактную поверхность"),
    ("contact area", "контактную зону"),
    ("the contact", "контакт"),
    ("reworked areas", "доработанные участки"),
    ("machined areas", "обработанные участки"),
    ("machined area", "обработанный участок"),
    ("machined", "обработанный"),
    ("roughened surfaces", "зашкуренные поверхности"),
    ("roughened", "зашкуренные"),
    ("roughen", "зашкурить"),
    ("cleaned surface", "очищенную поверхность"),
    ("preheated oven", "предварительно нагретую печь"),
    ("preheated", "предварительно нагретый"),
    ("Oversize bearing", "Ремонтный подшипник"),
    ("oversize bearing", "ремонтный подшипник"),
    ("repair bushes", "ремонтные втулки"),
    ("repair bush", "ремонтная втулка"),
    ("repair sleeves", "ремонтные рукава"),
    ("repair sleeve", "ремонтный рукав"),
    ("repair instructions", "инструкции по ремонту"),
    ("repair liner", "ремонтный вкладыш"),
    ("repair", "ремонт"),
    ("parent metal", "основной металл"),
    ("bare metal", "основной металл"),
    ("chromium plate", "хромовое покрытие"),
    ("cadmium plate", "кадмиевое покрытие"),
    ("sulphamate nickel plate", "сульфаматное никелевое покрытие"),
    ("sulphamate nickel", "сульфаматный никель"),
    ("nickel plate", "никелевое покрытие"),
    # Common verbs/phrases for Part 5
    ("this must be between", "значение должно быть между"),
    ("Make sure the", "Убедитесь, что"),
    ("this must be", "значение должно быть"),
    ("to the dimensions", "до размеров"),
    ("to adjust its length", "для подгонки длины"),
    ("to adjust its", "для подгонки"),
    ("in position on", "в положение на"),
    ("in position", "в положение"),
    ("using an applicable tool", "с помощью подходящего инструмента"),
    ("an applicable", "подходящий"),
    ("applicable", "применяемый"),
    ("if necessary", "при необходимости"),
    ("If necessary", "При необходимости"),
    ("as necessary", "по необходимости"),
    ("as shown", "как показано"),
    ("to the dimensions", "до размеров"),
    ("inclusion class", "класс включений"),
    ("correctly", "правильно"),
    ("correctly.", "правильно."),
    ("sufficiently", "в достаточной степени"),
    ("termination information", "информацию об окончании"),
    ("terminations are smooth", "окончания гладкие"),
    ("terminations", "окончания"),
    ("termination", "окончание"),
    ("all over the pin", "на весь штифт"),
    ("all over the", "повсюду"),
    ("all over", "повсюду"),
    # Common words for repair procedures
    (" width ", " ширину "),
    (" width", " ширину"),
    ("width ", "ширину "),
    ("smooth layer", "ровный слой"),
    ("smooth", "гладкий"),
    ("flat layer", "ровный слой"),
    ("layer", "слой"),
    (" dry.", " сухой."),
    (" dry", " сухой"),
    ("mixture", "смесь"),
    ("temperature", "температуру"),
    ("decrease", "снижение"),
    ("damage or wear", "повреждение или износ"),
    ("damage or corrosion", "повреждение или коррозию"),
    ("damage", "повреждение"),
    ("worn", "изношен"),
    ("wear", "износ"),
    ("corroded", "подвержен коррозии"),
    ("corrosion", "коррозию"),
    ("removal of", "удаления"),
    ("removal", "удаление"),
    (" amount ", " количество "),
    (" amount", " количество"),
    ("necessary", "необходимо"),
    (" cut ", " подрежьте "),
    ("minutes", "минут"),
    (" length.", " длины."),
    (" length", " длину"),
    ("dimensions", "размеры"),
    ("dimension", "размер"),
    ("external", "наружный"),
    ("internal", "внутренний"),
    ("outside", "наружный"),
    ("adjacent", "прилегающий"),
    ("a minimum of", "минимум"),
    ("minimum", "минимум"),
    ("maximum", "максимум"),
    ("that will touch", "которые будут соприкасаться"),
    ("touching", "касаясь"),
    ("touch", "соприкасаться"),
    ("tightly", "плотно"),
    ("Bond the", "Приклейте"),
    ("near to the", "вблизи"),
    ("near to", "вблизи"),
    ("near", "рядом"),
    ("an oven", "печь"),
    ("the oven", "печи"),
    (" oven", " печь"),
    (" cool", " остыть"),
    (" used", " использованный"),
    ("damaged", "повреждённый"),
    ("loose", "незакреплённый"),
    ("Loose", "Незакреплённый"),
    ("Damaged", "Повреждённый"),
    ("flaws", "дефекты"),
    # Connectors & short words (MUST be last — after all full sentence translations) ───
    ("from", "из"),
    ("between", "между"),
    (" and ", " и "),
    (" or ", " или "),
]

# ══════════════════════════════════════════════════════════════════════════════
#  TRANSLATION FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def is_only_numbers_or_codes(text: str) -> bool:
    """Return True if text is just numbers, part numbers, dates, page refs."""
    stripped = text.strip()
    if not stripped:
        return True
    # Pure numbers, dates, page refs, codes
    if re.match(r'^[\d\s,\-/\.\t©®]+$', stripped):
        return True
    # Part numbers like 201587001
    if re.match(r'^[\d\s,]+$', stripped):
        return True
    # Dates like "Dec 6/2019", "Mar 18/2025"
    if re.match(r'^(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d+/\d+$', stripped):
        return True
    # Tab-separated page ref + date: "601\tDec 6/2019" or "602\tMar 18/2025"
    if re.match(r'^\d+\t(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d+/\d+$', stripped):
        return True
    # ATA codes like 32-12-22
    if re.match(r'^\d{2}-\d{2}-\d{2}$', stripped):
        return True
    # Single short codes
    if re.match(r'^[A-Z]?\d+[\-\.]\d+', stripped) and len(stripped) < 15:
        return True
    # Standards like "BS EN 4007 (MAT206)", "EN 6049" (pure spec, no sentences)
    if re.match(r'^(BS\s+)?EN\s+\d+[\s\(\)A-Z\d\-]*$', stripped):
        return True
    # Material specs like "MTL-2701", "MAT206" (pure spec, no sentences)
    if re.match(r'^(MTL|MAT)\-?\d+[\s\(\)A-Z\d\-]*$', stripped):
        return True
    # Process/NDT specs like "M-DLNDT3", "M-DLPS1014-2" (pure code, no sentences)
    if re.match(r'^M-DL(NDT|PS)[\d\-]+\.?$', stripped):
        return True
    return False


def translate_toc_entry(text: str) -> str:
    """Translate a TOC entry preserving dot leaders and page numbers."""
    # Handle "Repair No." pattern in TOC (with dot leaders: ". . . .")
    m = re.match(r'^(Repair No\.\s*\d+-\d+)\s+(.*?)(\s*\.\s+\.[\s\.]+.*)', text)
    if m:
        prefix = m.group(1)
        comp_name = m.group(2).strip()
        rest = m.group(3)
        translated_comp = translate_component_name(comp_name)
        prefix_ru = prefix.replace("Repair No.", "Ремонт №")
        rest_ru = rest.replace("Repair No.", "Ремонт №")
        return f"{prefix_ru} {translated_comp}{rest_ru}"

    # Handle "Page Repair No." pattern
    m = re.match(r'^Page\s+(Repair No\.\s*\d+-\d+)\s+(.*?)(\s*\.\s+\.[\s\.]+.*)', text)
    if m:
        prefix = m.group(1)
        comp_name = m.group(2).strip()
        rest = m.group(3)
        translated_comp = translate_component_name(comp_name)
        prefix_ru = prefix.replace("Repair No.", "Ремонт №")
        rest_ru = rest.replace("Repair No.", "Ремонт №")
        return f"Стр. {prefix_ru} {translated_comp}{rest_ru}"

    # Handle figure TOC entries (component name with fig-item refs)
    # Require at least 2 dots (". .") to distinguish from abbreviation dots like "No."
    m = re.match(r'^(.*?)(\s*\.\s+\.[\s\.]+\s*\d*.*)$', text)
    if m:
        name_part = m.group(1).strip()
        rest = m.group(2)
        if name_part in FIXED:
            return f"{FIXED[name_part]}{rest}"
        translated_name = translate_component_name(name_part)
        return f"{translated_name}{rest}"

    # Handle tab-separated entries like "Storage\t798.12"
    m = re.match(r'^(.*?)\t(.*)$', text)
    if m:
        name_part = m.group(1).strip()
        page = m.group(2)
        if name_part in FIXED:
            page_ru = page.replace("Repair No.", "Ремонт №")
            return f"{FIXED[name_part]}\t{page_ru}"
        m_repair = re.match(r'^(Repair No\.\s*\d+-\d+)\s+(.*)$', name_part)
        if m_repair:
            prefix = m_repair.group(1).replace("Repair No.", "Ремонт №")
            comp = translate_component_name(m_repair.group(2).strip())
            page_ru = page.replace("Repair No.", "Ремонт №")
            return f"{prefix} {comp}\t{page_ru}"
        translated_name = translate_component_name(name_part)
        page_ru = page.replace("Repair No.", "Ремонт №")
        return f"{translated_name}\t{page_ru}"

    return text


def _translate_suffix(suffix: str) -> str:
    """Translate common suffixes like 'Protective Treatment (Sheet 1)'."""
    result = suffix
    if "Protective Treatment" in result:
        result = result.replace("Protective Treatment", "Защитная обработка")
    # "Installation of Bushes" must come before "Installation" alone
    if "Installation of Bushes" in result:
        result = result.replace("Installation of Bushes", "Установка втулок")
    if "Sheet" in result:
        result = re.sub(r'Sheet\s+(\d+)\s+of\s+(\d+)', r'Лист \1 из \2', result)
        result = re.sub(r'Sheet\s+(\d+)', r'Лист \1', result)
    if result.strip() == "Withdrawn" or result.strip() == "(Withdrawn)":
        result = result.replace("Withdrawn", "Отозвано")
    if "Superseded" in result:
        result = result.replace("Superseded", "Заменён")
    if "Only" in result:
        result = result.replace("Only", "Только")
    if " only" in result:
        result = result.replace(" only", " только")
    if "Machining and Installation" in result:
        result = result.replace("Machining and Installation", "Механическая обработка и установка")
    elif "Machining and Liner Installation" in result:
        result = result.replace("Machining and Liner Installation", "Механическая обработка и установка вкладыша")
    elif "Liner Installation" in result:
        result = result.replace("Liner Installation", "Установка вкладыша")
    elif "Installation" in result:
        result = result.replace("Installation", "Установка")
    if "Machining" in result and "обработка" not in result:
        result = result.replace("Machining", "Механическая обработка")
    if "Assembly" in result and "Сборка" not in result:
        result = result.replace("Assembly", "Сборка")
    if "Hole Locations" in result:
        result = result.replace("Hole Locations", "Расположение отверстий")
    if "Adjustment" in result:
        result = result.replace("Adjustment", "Регулировка")
    if "Related Parts" in result:
        result = result.replace("Related Parts", "Связанные детали")
    if "Dimensions After" in result:
        result = result.replace("Dimensions After", "Размеры после")
    if "Refer to IPL Figures" in result:
        result = result.replace("Refer to IPL Figures", "Обратитесь к рисункам ИПД")
    if "Refer to Figures" in result:
        result = result.replace("Refer to Figures", "Обратитесь к рисункам")
    if "Refer to Figure" in result:
        result = result.replace("Refer to Figure", "Обратитесь к рисунку")
    if "Grease Groove Dimensions" in result:
        result = result.replace("Grease Groove Dimensions", "Размеры канавки для смазки")
    if "Gland Housing" in result and "корпус" not in result:
        result = result.replace("Gland Housing", "корпусе сальника")
    if "in the" in result:
        result = result.replace("in the", "в")
    if "Tables" in result:
        result = result.replace("Tables", "Таблицы")
    if "Table" in result and "Таблицы" not in result:
        result = result.replace("Table", "Таблица")
    # Component names that may appear in TOC suffixes/refs
    if "Bushes" in result:
        result = result.replace("Bushes", "Втулки")
    if "Bush" in result and "Втулк" not in result:
        result = result.replace("Bush", "Втулка")
    if "Bearing" in result and "подшипник" not in result.lower():
        result = result.replace("Bearing", "Подшипник")
    if "Target" in result:
        result = result.replace("Target", "Мишень")
    if "Ardrox Application" in result:
        result = result.replace("Ardrox Application", "Нанесение Ardrox")
    elif "Application" in result and "Нанесение" not in result:
        result = result.replace("Application", "Нанесение")
    return result


def _lookup_base_name(name: str) -> str:
    """Look up base component name in glossaries."""
    if name in COMPONENT_NAMES:
        return COMPONENT_NAMES[name]
    if name in FIXED:
        return FIXED[name]
    return name


def translate_component_name(name: str) -> str:
    """Translate a component name from the glossary."""
    if name in COMPONENT_NAMES:
        return COMPONENT_NAMES[name]
    if name in FIXED:
        return FIXED[name]

    # ── Pattern: "Name (refs) ... or/and (refs) ... - Suffix" ──
    dash_idx = name.rfind(" - ")
    if dash_idx > 0:
        name_part = name[:dash_idx].strip()
        suffix_part = name[dash_idx + 3:].strip()
        suffix_ru = _translate_suffix(suffix_part)

        paren_idx = name_part.find("(")
        if paren_idx > 0:
            base = name_part[:paren_idx].strip()
            refs = name_part[paren_idx:]
            base_ru = _lookup_base_name(base)
            refs_ru = refs.replace("and", "и").replace("or", "или")
            refs_ru = _translate_suffix(refs_ru)
            return f"{base_ru} {refs_ru} — {suffix_ru}"
        else:
            base_ru = _lookup_base_name(name_part)
            return f"{base_ru} — {suffix_ru}"

    # ── Pattern: "Name (refs)" without suffix ──
    paren_idx = name.find("(")
    if paren_idx > 0:
        base = name[:paren_idx].strip()
        refs = name[paren_idx:]
        base_ru = _lookup_base_name(base)
        refs_ru = _translate_suffix(refs)
        refs_ru = refs_ru.replace(" and ", " и ").replace(" or ", " или ")
        if base_ru != base:
            return f"{base_ru} {refs_ru}"

    # ── Pattern: "(Withdrawn)" suffix ──
    if name.endswith("(Withdrawn)"):
        base = name.replace("(Withdrawn)", "").strip()
        base_ru = translate_component_name(base)
        return f"{base_ru} (Отозвано)"

    return name


def translate_sb_title(text: str) -> str:
    """Translate SB title text."""
    for en, ru in sorted(SB_TITLE_PARTS.items(), key=lambda x: -len(x[0])):
        if text.startswith(en):
            remainder = text[len(en):]
            if remainder:
                return ru + remainder
            return ru
    return text


def translate_reason(text: str) -> str:
    """Translate reason-for-change text in revision tables."""
    result = text
    normalized = re.sub(r'\n', ' ', result)
    # Fix garbled text like "Safran Lиing Systems" -> "Safran Landing Systems"
    normalized = re.sub(r'Safran L[а-яёА-ЯЁ\w]*ing Systems', 'Safran Landing Systems', normalized)
    sorted_phrases = sorted(REASON_PHRASES.keys(), key=len, reverse=True)
    for en in sorted_phrases:
        ru = REASON_PHRASES[en]
        if len(en) <= 4:
            # Word-boundary replacement for short words
            normalized = re.sub(r'\b' + re.escape(en) + r'\b', ru, normalized)
        else:
            normalized = normalized.replace(en, ru)
    return normalized


def translate_repair_description(text: str) -> str:
    """Translate 'Repair to Component — Process' descriptions."""
    result = text
    parts = result.split('\t')
    desc = parts[0]
    rest_parts = '\t'.join(parts[1:]) if len(parts) > 1 else ""

    rest_parts = rest_parts.replace("Repair No.", "Ремонт №")

    desc = desc.replace("Oversize Bushes", "Ремонтные (увеличенные) втулки")
    desc = desc.replace("Oversize Bush(es)", "Ремонтные (увеличенные) втулки")
    desc = desc.replace("Oversize Bush (es)", "Ремонтные (увеличенные) втулки")
    desc = desc.replace("Oversize Bush", "Ремонтная (увеличенная) втулка")
    desc = desc.replace("Oversize Lubrication adapter", "Ремонтный (увеличенный) смазочный адаптер")
    desc = desc.replace("Oversize Transfer Dowel", "Ремонтный (увеличенный) переходной штифт")
    desc = desc.replace("Oversize Thread Insert", "Ремонтная (увеличенная) резьбовая вставка")
    desc = desc.replace("Lower Bearing Subassembly", "Сборка нижнего подшипника")
    desc = desc.replace("Machining and Inner Liner Installation", "Механическая обработка и установка внутреннего вкладыша")
    desc = desc.replace("Machining and Liner Installation", "Механическая обработка и установка вкладыша")
    desc = desc.replace("Repair Threaded Inserts", "Ремонт резьбовых вставок")
    desc = desc.replace("Repair Threaded Insert", "Ремонт резьбовой вставки")
    desc = desc.replace("Repair Bearing", "Ремонт подшипника")
    desc = desc.replace("Repair Sleeves", "Ремонт втулок")
    desc = desc.replace("Repair Sleeve", "Ремонт втулки")
    desc = desc.replace("Repair sleeve", "Ремонт втулки")
    desc = desc.replace("Repair to ", "Ремонт ")
    desc = desc.replace("Repair Bushes", "Ремонт втулок")
    desc = desc.replace("Repair Bush", "Ремонт втулки")

    for en, ru in sorted(COMPONENT_NAMES.items(), key=lambda x: -len(x[0])):
        desc = desc.replace(en, ru)

    desc = desc.replace("Machining and Installation", "Механическая обработка и установка")
    desc = desc.replace("Machining and installation", "Механическая обработка и установка")
    desc = desc.replace("Machining", "Механическая обработка")
    desc = desc.replace("Installation", "Установка")
    desc = desc.replace("Cadmium Plating", "Кадмирование")
    desc = desc.replace("Chromium Plate Termination", "Граница хромового покрытия")
    desc = re.sub(r'Sheet\s+(\d+)', r'Лист \1', desc)
    desc = desc.replace("Rework", "Доработка")
    desc = desc.replace(" — ", " — ")  # already correct
    desc = desc.replace(" and ", " и ")

    if rest_parts:
        return f"{desc}\t{rest_parts}"
    return desc


def translate_long_text(text: str) -> str:
    """Translate longer text blocks (certification, copyright, etc.)."""
    if text.startswith("This manual complies with British Civil Airworthiness"):
        return "Данное руководство соответствует требованиям Британских авиационных правил лётной годности, раздел A, глава A5-3."

    if text.startswith("NOTE: The above certification does not apply"):
        return ("ПРИМЕЧАНИЕ: Вышеуказанная сертификация не распространяется на изменения или "
                "дополнения, внесённые после даты первичной сертификации другими "
                "одобренными организациями. Изменения или дополнения, внесённые другими "
                "одобренными организациями, должны быть отдельно сертифицированы и "
                "зарегистрированы на отдельных учётных листах.")

    if text.startswith("This document and all information contained herein"):
        return ("Настоящий документ и вся содержащаяся в нём информация являются "
                "исключительной собственностью Safran Landing Systems (и/или аффилированных компаний).")

    if text.startswith("No intellectual property rights are granted"):
        return ("Передача данного документа или раскрытие его содержания не предоставляет "
                "никаких прав на интеллектуальную собственность. Воспроизведение данного "
                "документа для третьих лиц запрещено без письменного согласия Safran Landing Systems "
                "(и/или соответствующей аффилированной компании).")

    if text.startswith("Record the issue date and insertion date"):
        return ("Запишите дату выпуска и дату вставки данной редакции в Запись изменений "
                "и сохраните данное Письмо о рассылке.")

    if text.startswith("The technical data in this document"):
        return ("Технические данные в настоящем документе (или файле) могут содержать данные США "
                "и подлежать экспортному контролю в соответствии с")

    return text


def translate_header_line(text: str) -> str:
    """Translate header/footer running text."""
    m = re.match(r'^PART\s+No\.\s+([\d\s]+)AND\s+([\d]+)\s+COMPONENT MAINTENANCE MANUAL\s+MAIN LANDING GEAR LEG$', text)
    if m:
        return f"ДЕТАЛЬ № {m.group(1)}И {m.group(2)} РУКОВОДСТВО ПО ТЕХНИЧЕСКОМУ ОБСЛУЖИВАНИЮ КОМПОНЕНТОВ СТОЙКА ОСНОВНОГО ШАССИ"

    if "SAFRAN LANDING SYSTEMS" in text and "SUBSEQUENT REVISION PAGE DATES" in text:
        m2 = re.match(r'[©®\s]*SAFRAN LANDING SYSTEMS\s+(\d+)\s+\(AND SUBSEQUENT REVISION PAGE DATES\)', text)
        if m2:
            return f"© SAFRAN LANDING SYSTEMS {m2.group(1)} (И ПОСЛЕДУЮЩИЕ ДАТЫ ПЕРЕСМОТРА СТРАНИЦ)"
        return text.replace("AND SUBSEQUENT REVISION PAGE DATES", "И ПОСЛЕДУЮЩИЕ ДАТЫ ПЕРЕСМОТРА СТРАНИЦ")

    return text


def translate_hf_text(text: str) -> str:
    """Translate header/footer textbox text.

    Handles patterns specific to headers/footers:
    - Section labels (RECORD OF REVISIONS, LIST OF EFFECTIVE PAGES, etc.)
    - "Page X of Y" references
    - "Page X Mar 18/2025" combined page+date
    - "Letter of Transmittal No. N"
    - "TITLE PAGE Sep 16/2016"
    - "COMPONENT MAINTENANCE MANUAL 32-12-22 MAIN LANDING GEAR LEG"
    - Running header "PART No. ... COMPONENT MAINTENANCE MANUAL ..."
    Falls back to translate_text() for anything else.
    """
    stripped = text.strip()
    if not stripped:
        return text

    # Skip already-Russian content
    if re.search(r'[А-Яа-яЁё]{3,}', stripped):
        return text

    # Skip pure dates/numbers/codes
    if is_only_numbers_or_codes(stripped):
        return text

    leading = text[:len(text) - len(text.lstrip())]
    trailing = text[len(text.rstrip()):]
    core = stripped

    # 1. Exact match in FIXED (handles most section labels)
    if core in FIXED:
        return leading + FIXED[core] + trailing

    # 2. Running header: "PART No. ... COMPONENT MAINTENANCE MANUAL MAIN LANDING GEAR LEG"
    result = translate_header_line(core)
    if result != core:
        return leading + result + trailing

    # 3. "COMPONENT MAINTENANCE MANUAL 32-12-22 MAIN LANDING GEAR LEG"
    m = re.match(r'^COMPONENT\s+MAINTENANCE\s+MANUAL\s+(\d{2}-\d{2}-\d{2})\s+MAIN LANDING GEAR LEG$', core)
    if m:
        return leading + f"РУКОВОДСТВО ПО ТЕХНИЧЕСКОМУ ОБСЛУЖИВАНИЮ КОМПОНЕНТОВ {m.group(1)} СТОЙКА ОСНОВНОГО ШАССИ" + trailing

    # 4. "TITLE PAGE Sep 16/2016" — section label + date
    m = re.match(r'^(TITLE PAGE|RECORD OF REVISIONS|RECORD OF TEMPORARY REVISIONS|'
                 r'LIST OF SERVICE BULLETINS|LIST OF EFFECTIVE PAGES|'
                 r'UNIT IDENTIFICATION CHART|TABLE OF CONTENTS|ILLUSTRATIONS)'
                 r'(?:\s+\(Continued\))?\s+'
                 r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d+/\d+)$', core)
    if m:
        label = m.group(1)
        # Check for (Continued) in the original
        has_cont = "(Continued)" in core
        date = m.group(2)
        label_full = label + (" (Continued)" if has_cont else "")
        label_ru = FIXED.get(label_full, FIXED.get(label, label))
        return leading + f"{label_ru} {date}" + trailing

    # 5. "Letter of Transmittal No. N"
    m = re.match(r'^Letter of Transmittal No\.\s*(\d+)$', core)
    if m:
        return leading + f"Письмо о рассылке № {m.group(1)}" + trailing

    # 6. "Page X of Y" (page count in footer)
    m = re.match(r'^Page\s+(\d+)\s+of\s+(\d+)$', core)
    if m:
        return leading + f"Стр. {m.group(1)} из {m.group(2)}" + trailing

    # 7. "Page X Mar 18/2025" — page + date combined
    m = re.match(r'^Page\s+(\d+)\s+((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d+/\d+)$', core)
    if m:
        return leading + f"Стр. {m.group(1)} {m.group(2)}" + trailing

    # 8. Standalone "Page X"
    m = re.match(r'^Page\s+(\d+)$', core)
    if m:
        return leading + f"Стр. {m.group(1)}" + trailing

    # 9. "32-12-22 Mar 18/2025" — ATA code + date (keep as-is)
    m = re.match(r'^\d{2}-\d{2}-\d{2}\s+', core)
    if m:
        return text

    # 10. Fallback: use main translate_text()
    return translate_text(text)


def _translation_quality_ok(original: str, translated: str) -> bool:
    """Check if translation result is acceptable (not garbled mixed text).

    Returns False if the result still has too many English words,
    indicating partial word-level replacement created garbage.
    """
    if translated == original:
        return False
    # Count remaining English words (3+ letters, excluding codes/specs)
    en_words = re.findall(r'[A-Za-z]{3,}', translated)
    # Filter out known codes/specs that should stay English
    code_patterns = re.compile(
        r'^(SB|IPL|CMM|PCS|AMS|NDT|NLG|MLG|DPI|FPI|SDS|VDC|PVC|REF|IPL|'
        r'Safran|Landing|Systems|Airbus|EASA|Messier|Ardrox|Turner|Megger|'
        r'Ltd|Ref|Item|Fig|lbf|min|bar|PSC|CODE|CAGE|MAF|EDES|'
        r'Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|'
        r'Extractor|Drift|Drawbolt|Spanner|Wrench|Crowfoot|Milliohmmeter|'
        r'Adapter|Reactor|Clamp|Tackle|Trolley|Fixture|Locator|'
        r'Extraction|Pneumatic|Hydraulic|Pump|Press|Bench|Holding|Blocks|'
        r'Lifting|Towing|Transport|Build|Torque|Offset|Assembly|'
        r'BT51|MT1025|MT1026|Pin|Set|Pad|Tube|Bolt|Bar|Frame|Arms|Plate|'
        r'DLPS|DLNDT|AECMA|MTL|MAT|NCD|NCT|CND|CNU|Material|Pre|Post|'
        r'Sermetel|IFC|MIL|DIA|SPOTFACE|Qty|CON|UHT|'
        r'Alocrom|Araldite|Accomet|Ardrox|Fibreslip|Loctite|Grade|'
        r'Dowty|Limited|Ltd|Guide|STD|RAD|MIN|MAX|'
        r'Type|PVC|TBA|grit|oC|oF)$'
    )
    en_words_real = [w for w in en_words if not code_patterns.match(w)
                     and not re.match(r'^[A-Z]\d', w)]
    ru_words = re.findall(r'[А-Яа-яЁё]{3,}', translated)

    total = len(en_words_real) + len(ru_words)
    if total < 2:
        return True  # Too few words to judge

    en_ratio = len(en_words_real) / total
    # If more than 25% of words are still English, translation is garbled
    return en_ratio <= 0.25


def translate_procedural_text(text: str) -> str:
    """Translate procedural/instructional text using word/phrase-level replacement.

    Applied as a fallback when exact-match translation fails.
    Uses PROCEDURAL_VOCAB for ordered phrase replacement.
    Returns original text if translation quality is poor (mixed garbage).
    """
    result = text
    for en, ru in PROCEDURAL_VOCAB:
        result = result.replace(en, ru)
    # Quality gate: if result is garbled mixed text, return original
    if not _translation_quality_ok(text, result):
        return text
    return result


def translate_text(text: str) -> str:
    """Main translation function. Returns translated text or original if untranslatable."""
    stripped = text.strip()
    if not stripped:
        return text

    if is_only_numbers_or_codes(stripped):
        return text

    # Skip already-Russian content
    if re.search(r'[А-Яа-яЁё]{3,}', stripped):
        return text

    # Preserve leading/trailing whitespace
    leading = text[:len(text) - len(text.lstrip())]
    trailing = text[len(text.rstrip()):]
    core = text.strip()

    # 1. Exact match in FIXED
    if core in FIXED:
        return leading + FIXED[core] + trailing

    # 2. Long text translations
    for start_phrase in [
        "This manual complies",
        "NOTE: The above certification",
        "This document and all information",
        "No intellectual property rights",
        "Record the issue date",
        "The technical data in this document",
    ]:
        if core.startswith(start_phrase):
            return leading + translate_long_text(core) + trailing

    # 3. Header line
    if "COMPONENT MAINTENANCE MANUAL" in core or "SUBSEQUENT REVISION PAGE DATES" in core:
        result = translate_header_line(core)
        if result != core:
            return leading + result + trailing

    # 4. Protective Treatment patterns (check BEFORE dot leaders)
    # Skip digit+tab entries which should be handled by the tab handler
    if "Protective Treatment" in core and not re.match(r'^\d+\t', core):
        m_cont = re.match(r'^-\s*(Protective Treatment.*?)(\s*\.[\s\.]+.*)$', core)
        if m_cont:
            suffix_ru = _translate_suffix(m_cont.group(1))
            return leading + f"— {suffix_ru}{m_cont.group(2)}" + trailing
        m_pt = re.match(r'^(Protective Treatment\s*-?\s*.*?)(\s*\.[\s\.]+.*)$', core)
        if m_pt:
            suffix_ru = _translate_suffix(m_pt.group(1))
            return leading + f"{suffix_ru}{m_pt.group(2)}" + trailing
        result = translate_component_name(core)
        if result != core:
            return leading + result + trailing

    # 4b. Tab-separated entries
    if "\t" in core:
        m_lep = re.match(r'^(Repair No\.\s*\d+-\d+)\t(.*)$', core)
        if m_lep:
            prefix = m_lep.group(1).replace("Repair No.", "Ремонт №")
            return leading + f"{prefix}\t{m_lep.group(2)}" + trailing

        m_blank = re.match(r'^(\d+)\tBlank$', core)
        if m_blank:
            return leading + f"{m_blank.group(1)}\tПусто" + trailing

        if core.startswith("Fig."):
            return leading + core.replace("Fig.", "Рис.").replace("Page", "Страница") + trailing

        m_fig = re.match(r'^(\d+)\t(.*)$', core)
        if m_fig:
            fig_num = m_fig.group(1)
            rest = m_fig.group(2)
            if ". . ." in rest:
                m_dots = re.match(r'^(.*?)(\s*\.[\s\.]+\s*)(\t?\d+.*)$', rest)
                if m_dots:
                    comp_part = m_dots.group(1).strip()
                    dots = m_dots.group(2)
                    page = m_dots.group(3)
                    translated_comp = translate_component_name(comp_part)
                    return leading + f"{fig_num}\t{translated_comp}{dots}{page}" + trailing
                m_dots2 = re.match(r'^(.*?)(\s*\.[\s\.]+\s*)$', rest)
                if m_dots2:
                    comp_part = m_dots2.group(1).strip()
                    dots = m_dots2.group(2)
                    m_trailing_dash = re.match(r'^(.*?)\s*-\s*$', comp_part)
                    if m_trailing_dash:
                        inner = m_trailing_dash.group(1).strip()
                        translated_inner = translate_component_name(inner)
                        return leading + f"{fig_num}\t{translated_inner} —{dots}" + trailing
                    translated_comp = translate_component_name(comp_part)
                    return leading + f"{fig_num}\t{translated_comp}{dots}" + trailing
            else:
                translated_rest = translate_component_name(rest.strip())
                return leading + f"{fig_num}\t{translated_rest}" + trailing

        if not re.match(r'^\d', core):
            toc_result = translate_toc_entry(core)
            if toc_result != core:
                return leading + toc_result + trailing
            # Fall through to procedural text for tab-separated non-TOC entries

        # else: digit-prefixed tab entries — fall through to procedural text
        # (previously returned early, but textbox dimension+annotation text needs translation)

    # 4c. "(Sheet X of Y)" standalone continuation lines
    m_sheet = re.match(r'^\(Sheet\s+(\d+)\s+of\s+(\d+)\)\s*(.*)', core)
    if m_sheet:
        sheet_ru = f"(Лист {m_sheet.group(1)} из {m_sheet.group(2)})"
        rest = m_sheet.group(3)
        return leading + sheet_ru + (" " + rest if rest else "") + trailing

    # 4d. "Post SB-..." / "Pre SB-..." headings with component name
    m_sb = re.match(r'^(Post|Pre)\s+(SB[-\d]+)\s*[—–-]\s*(.+?)(\s*\.[\s\.]+.*)?$', core)
    if m_sb:
        sb_prefix = "После" if m_sb.group(1) == "Post" else "До"
        sb_num = m_sb.group(2)
        comp_name = m_sb.group(3).strip()
        dots_rest = m_sb.group(4) or ""
        comp_ru = translate_component_name(comp_name)
        return leading + f"{sb_prefix} {sb_num} — {comp_ru}{dots_rest}" + trailing

    # 5. TOC entries (contain dot leaders, no tabs)
    if ". . ." in core:
        return leading + translate_toc_entry(core) + trailing

    # 6. "Fig.\tPage" header (no-tab variant)
    if core.startswith("Fig."):
        return leading + core.replace("Fig.", "Рис.").replace("Page", "Страница") + trailing

    # 7. Component name with Protective Treatment (fallback for non-dot cases)
    if "Protective Treatment" in core:
        result = translate_component_name(core)
        if result != core:
            return leading + result + trailing

    # 7b. "Component Repairs — Key Diagram Figure NNN" headings
    m_kd = re.match(r'^(.+?)\s+[—–-]\s+Key Diagram\s+Figure\s+(\d+)$', core)
    if m_kd:
        comp = m_kd.group(1).strip()
        fig_num = m_kd.group(2)
        comp_ru = translate_component_name(comp)
        return leading + f"{comp_ru} — Ключевая схема Рисунок {fig_num}" + trailing

    # 7c. "REFER TO FIGURE NNN REPAIR No. ..." pattern (MUST be before generic REPAIR check)
    m_ref_repair = re.match(r'^REFER TO FIGURE\s+(\d+)\s+REPAIR\s+No\.\s+(.+)$', core)
    if m_ref_repair:
        return leading + f"ОБРАТИТЕСЬ К РИСУНКУ {m_ref_repair.group(1)} РЕМОНТ № {m_ref_repair.group(2)}" + trailing

    # 7d. Standalone "REFER TO FIGURE NNN"
    m_ref = re.match(r'^REFER TO FIGURE\s+(\d+)$', core)
    if m_ref:
        return leading + f"ОБРАТИТЕСЬ К РИСУНКУ {m_ref.group(1)}" + trailing

    # 7e. "REPAIR No. X-Y, REPAIR X-Y — X-Y" headings
    if "REPAIR No." in core or "REPAIR " in core:
        result = core.replace("REPAIR No.", "РЕМОНТ №").replace("REPAIR ", "РЕМОНТ ")
        if result != core:
            return leading + result + trailing

    # 8. "Repair No." entries (outside TOC)
    if core.startswith("Repair No."):
        m = re.match(r'^(Repair No\.\s*\d+-\d+)\s+(.*?)$', core)
        if m:
            comp = m.group(2).strip()
            translated_comp = translate_component_name(comp)
            prefix = m.group(1).replace("Repair No.", "Ремонт №")
            return leading + f"{prefix} {translated_comp}" + trailing
        return leading + core.replace("Repair No.", "Ремонт №") + trailing

    # 8d. "Key Diagram Figure NNN" (possibly after partial translation)
    m_kd2 = re.search(r'Key Diagram\s+Figure\s+(\d+)', core)
    if m_kd2:
        result = core.replace(m_kd2.group(0), f"Ключевая схема Рисунок {m_kd2.group(1)}")
        return leading + result + trailing

    # 9. "Page X" reference
    if re.match(r'^Page\s+\d+', core):
        return leading + core.replace("Page", "Стр.") + trailing

    # 10. Exact match in component names, part names, table headers, tool names
    if core in COMPONENT_NAMES:
        return leading + COMPONENT_NAMES[core] + trailing
    if core in PART_NAMES_TABLE:
        return leading + PART_NAMES_TABLE[core] + trailing
    if core in TABLE_HEADERS:
        return leading + TABLE_HEADERS[core] + trailing
    if core in TOOL_NAMES_TABLE:
        return leading + TOOL_NAMES_TABLE[core] + trailing

    # 10a. Protective treatment / material text (bypass quality gate)
    # MUST come BEFORE translate_component_name() which splits on " - " and
    # can intercept PT/material text, replacing only the dash while leaving
    # English words like "Aluminium alloy" untranslated.
    _pt_kw_text = (
        "apply ", "paint", "primer", "refer ",
        "chromic acid", "passivate", "anodise", "anodize",
        "do not", "not paint", "not cadmium",
        "protective treatment", "cadmium plate",
        "chromium plate", "chromium plated", "alocrom",
        "spotface", "areas ", "the thread", "the bush",
        "the hole", "the face", "the lubrication",
        "the identification", "the bearing", "the screw",
        "split pin", "thick", "thickness",
        "externally", "internally", "all over", "but not",
        "including", "optional", "permitted",
        "plate all over", "plate to",
        "aluminium", "stainless", "zinc-nickel",
        "coat ", "color ", "extend onto",
        "chrome", "base material", "where shown", "split line",
        # Drawing annotation keywords
        "bore", "bores", "flange", "knuckle", "lug", "lugs",
        "radius", "plating", "deposit", "overlap", "grinding",
        "termination", "journal", "barrel", "retraction",
        "manifold", "brake", "grease", "chamfer",
        "beyond", "remain", "transition", "permissible",
        "face ", "thread", "limit", "internal", "external",
        "typical", "places", "intersection", "point",
        "drag arm", "for main fitting", "run out", "band",
        # Part 5 repair procedure keywords
        "machine ", "machin", "examine", "shot peen", "grind ",
        "finish grind", "passivat", "identify the part",
        "adhesive", "liner", "emery", "roughen", "bonded",
        "scarf", "contact surface", "clamp ", "repair bush",
        "repair sleeve", "oversize", "sealant", "anodize",
        "de-embrittle", "delamination", "grit blast",
        "sulphamate", "preheat", "oven", "mixture",
        "clean cold water", "brush", "masking tape",
        "minimum of", "correctly", "damage", "wear",
        "corroded", "corrosion", "flaws", "parent metal",
        "bare metal", "surface finish", "micro-inches",
        "micrometers", "repair number", "part number",
        "lug ", " lug", "press pad", "drift ",
        "line ream", "repair liner", "gland housing",
    )
    _core_lower = core.lower()
    if any(kw in _core_lower for kw in _pt_kw_text):
        if core in FIXED:
            return leading + FIXED[core] + trailing
        result = core
        for en, ru in PROCEDURAL_VOCAB:
            result = result.replace(en, ru)
        if result != core:
            return leading + result + trailing

    # 10b. Component name with pattern matching
    result = translate_component_name(core)
    if result != core:
        return leading + result + trailing

    # 10c. Component name + "Only" suffix
    if core.endswith(" Only"):
        base = core[:-5].strip()
        base_ru = translate_component_name(base)
        if base_ru != base:
            return leading + f"{base_ru} Только" + trailing

    # 10d. "- ComponentName" continuation lines
    m_cont = re.match(r'^-\s+(.+)$', core)
    if m_cont:
        inner = m_cont.group(1).strip()
        inner_ru = translate_component_name(inner)
        if inner_ru != inner:
            return leading + f"— {inner_ru}" + trailing

    # 11. "Blank" as page status
    if core == "Blank":
        return leading + "Пусто" + trailing

    # 11a. Standalone "and"/"or" connectors
    if core == "and":
        return leading + "и" + trailing
    if core == "or":
        return leading + "или" + trailing

    # 11a2. "and N" / "or SPEC" continuation fragments
    if re.match(r'^and\s+[\d\(\)\-/A-Z]+$', core):
        return leading + core.replace("and ", "и ") + trailing
    if re.match(r'^or\s+\S+', core):
        return leading + core.replace("or ", "или ") + trailing

    # 11b. Units text like "(110 lbf/in2)."
    if re.match(r'^\(\d+\s+lbf/in2\)\.$', core):
        return text  # keep as-is, it's a measurement

    # 11c. Part references "(XX-NNN) and (XX-NNN)."
    m_refs = re.match(r'^\([\d\-]+\)\s+and\s+\([\d\-]+\)\.$', core)
    if m_refs:
        return leading + core.replace(" and ", " и ") + trailing

    # 11d. "the PART (XX-NNN)." references
    m_part_ref = re.match(r'^the\s+(.+?)\s+\([\d\-A-Za-z]+\)\.$', core)
    if m_part_ref:
        part_name = m_part_ref.group(1)
        # Translate part name via PROCEDURAL_VOCAB lookups
        for en, ru in [("pin", "штифт"), ("nut", "гайку"), ("bolt", "болт"),
                       ("seal", "уплотнение"), ("bush", "втулку"), ("bearing", "подшипник"),
                       ("bracket", "кронштейн"), ("sleeve", "втулку"), ("spacer", "проставку"),
                       ("plate", "пластину"), ("washer", "шайбу"), ("clamp", "хомут"),
                       ("retainer", "фиксатор"), ("cylinder", "цилиндр"),
                       ("split pins", "шплинты"), ("split pin", "шплинт")]:
            if part_name.lower() == en:
                result = core.replace(f"the {part_name}", ru)
                return leading + result + trailing
        # If no match, try procedural
        result = translate_procedural_text(core)
        if result != core:
            return leading + result + trailing

    # 11e. NOTE with tab-separated technical content
    if core.startswith("NOTE:") and "\t" in core:
        result = translate_procedural_text(core)
        if result != core:
            return leading + result + trailing

    # 12. Procedural text fallback (word-level translation)
    if re.search(r'[A-Za-z]{4,}', core):
        result = translate_procedural_text(text)
        if result != text:
            return result

    # 12b. Short English fragments with "and"/"or"
    if " and " in core or " or " in core:
        result = core.replace(" and ", " и ").replace(" or ", " или ")
        if result != core:
            return leading + result + trailing

    return text


def translate_table_cell_text(text: str) -> str:
    """Translate text within a table cell."""
    stripped = text.strip()
    if not stripped:
        return text

    clean = stripped.strip('\n').strip()

    # Exact match in FIXED dict (handles procedural sentences added there)
    if clean in FIXED:
        return text.replace(clean, FIXED[clean])

    # Table headers
    if clean in TABLE_HEADERS:
        return text.replace(clean, TABLE_HEADERS[clean])

    # Section names in revision tables
    if clean in SECTION_NAMES:
        return text.replace(clean, SECTION_NAMES[clean])

    # "Repair No." section refs
    if clean.startswith("Repair No."):
        return text.replace("Repair No.", "Ремонт №")

    # SB titles
    for en_prefix in sorted(SB_TITLE_PARTS.keys(), key=len, reverse=True):
        if clean.startswith(en_prefix):
            return text.replace(clean, translate_sb_title(clean))

    # SB revision numbers
    if clean == "Initial Issue":
        return text.replace(clean, "Первоначальный выпуск")
    if clean == "No effect":
        return text.replace(clean, "Без изменений")

    # "Updated paras X.Y" type entries
    if clean.startswith("Updated paras"):
        return text.replace("Updated paras", "Обновлены пункты")
    if clean.startswith("Updated para"):
        return text.replace("Updated para", "Обновлён пункт")

    # "NNN (Superseded)" pattern
    if "(Superseded)" in clean:
        return text.replace("(Superseded)", "(Заменён)")

    # Standalone continuation fragments
    if clean.startswith("para ") or clean.startswith("paras "):
        result = clean.replace("paras ", "пунктов ").replace("para ", "пункта ")
        result = translate_reason(result)
        return text.replace(clean, result)

    # Reason for change
    if any(phrase in clean for phrase in ["Updated", "Added", "Updated fig", "Updated Messier"]):
        return text.replace(clean, translate_reason(clean))

    # Standalone "figure NNN" or "figures NNN" (continuation lines)
    m_fig = re.match(r'^(figures?)\s+(.*)$', clean)
    if m_fig:
        word = "рисунки" if m_fig.group(1) == "figures" else "рисунок"
        return text.replace(clean, f"{word} {m_fig.group(2)}")

    # Standalone continuation fragments from multi-paragraph reason cells
    if clean.startswith("in para"):
        return text.replace("in para", "в пункте")
    if clean.startswith("in paras"):
        return text.replace("in paras", "в пунктах")

    # "Sliding tube" entries in figure tables
    if "Sliding tube" in clean or "sliding tube" in clean:
        result = clean
        result = result.replace("Sliding tube", "Скользящая труба")
        result = result.replace("sliding tube", "Скользящая труба")
        result = _translate_suffix(result)
        result = result.replace(" and ", " и ").replace(" or ", " или ")
        return text.replace(clean, result)

    # Repair/Oversize descriptions
    if (clean.startswith("Repair to ") or clean.startswith("Repair Bush") or
        clean.startswith("Repair Bearing") or clean.startswith("Repair Sleeve") or
        clean.startswith("Repair sleeve") or clean.startswith("Repair Sleeves") or
        clean.startswith("Repair Threaded") or clean.startswith("Oversize ") or
        clean.startswith("Lower Bearing")):
        fixed = re.sub(r'Installation(Repair No\.)', r'Installation\t\1', clean)
        result = translate_repair_description(fixed)
        return text.replace(clean, result)

    # "Fits and Clearances - Key Diagram" with dot leaders
    if "Key Diagram" in clean and ". . ." in clean:
        result = clean
        result = result.replace("Fits and Clearances - Key Diagram", "Посадки и зазоры — Ключевая схема")
        result = result.replace("Approved Repairs - Key Diagram", "Утверждённые ремонты — Ключевая схема")
        for en, ru in sorted(COMPONENT_NAMES.items(), key=lambda x: -len(x[0])):
            result = result.replace(en, ru)
        return text.replace(clean, result)

    # Table cells with dot leaders (TOC-like entries)
    if ". . ." in clean:
        result = translate_toc_entry(clean)
        if result != clean:
            return text.replace(clean, result)
        # Try word-level translation
        result = translate_procedural_text(clean)
        if result != clean:
            return text.replace(clean, result)

    # "Component Repairs - Key Diagram" pattern
    if "Repairs - Key Diagram" in clean or "Key Diagram" in clean:
        result = clean
        result = result.replace("Approved Repairs - Key Diagram", "Утверждённые ремонты — Ключевая схема")
        result = result.replace("Repairs - Key Diagram", "Ремонты — Ключевая схема")
        for en, ru in sorted(COMPONENT_NAMES.items(), key=lambda x: -len(x[0])):
            result = result.replace(en, ru)
        return text.replace(clean, result)

    # Continuation fragments from multi-paragraph reason cells
    if clean.startswith("Limited to Safran"):
        result = clean
        result = re.sub(r'Safran L[а-яёА-ЯЁ\w]*ing Systems', 'Safran Landing Systems', result)
        result = result.replace("Limited to Safran Landing Systems", "Limited на Safran Landing Systems")
        return text.replace(clean, result)

    # "IPL N" → "ИПД N" (Illustrated Parts List references)
    m_ipl = re.match(r'^IPL\s+(.+)$', clean)
    if m_ipl:
        return text.replace("IPL", "ИПД")

    # "TBA" → "Будет определено"
    if clean == "TBA":
        return text.replace("TBA", "Будет определено")

    # Standalone "install." in table cells
    if clean == "install." or clean == "install":
        return text.replace("install", "установить")

    # "Areas A and B." pattern
    if clean.startswith("Areas "):
        result = clean.replace("Areas", "Участки").replace(" and ", " и ").replace(" or ", " или ")
        return text.replace(clean, result)

    # Numbers/codes with "and"/"or" (continuation fragments)
    # "814, 818 and 823", "10-260 and 10-260A", "1-49, and 1-49A,"
    if re.match(r'^[\d,\s\-/]+ and [\d,\s\-/A-Z]+,?$', clean):
        return text.replace(" and ", " и ")
    # "and (20-220)", "and 460006261", "and PCS-3002", "and 2"
    if re.match(r'^and\s+[\d\(\)\-/A-Z]+$', clean):
        return text.replace("and ", "и ")
    # "460006151/20, 460006151/30 and" (trailing and)
    if re.match(r'^[\d,\s\-/]+ and$', clean):
        return text.replace(" and", " и")

    # "will install" / "will install:" continuation fragments
    if "will" in clean and ("install" in clean or "be installed" in clean):
        result = clean.replace("will install", "будут установлены").replace("will be installed", "будут установлены")
        result = result.replace(" and ", " и ").replace(" or ", " или ")
        return text.replace(clean, result)

    # Continuation fragments starting with "plate" (from "Apply cadmium plate")
    if clean.startswith("plate ") or clean.startswith("plate:"):
        result = clean
        for en, ru in PROCEDURAL_VOCAB:
            result = result.replace(en, ru)
        result = result.replace("plate ", "покрытие ")
        result = result.replace("plate:", "покрытие:")
        if result != clean:
            return text.replace(clean, result)

    # "not paint" continuation fragments
    if clean.startswith("not paint") or clean.startswith("not cadmium"):
        result = translate_procedural_text(clean)
        if result != clean:
            return text.replace(clean, result)

    # Standalone short words in table cells
    if clean == "and":
        return text.replace("and", "и")
    if clean == "REV":
        return text.replace("REV", "РЕД")
    if clean == "Only":
        return text.replace("Only", "Только")
    if clean == "Tables":
        return text.replace("Tables", "Таблицы")

    # Part names in table cells (discard tables, NDT tables)
    if clean in PART_NAMES_TABLE:
        return text.replace(clean, PART_NAMES_TABLE[clean])

    # Tool names in table cells
    if clean in TOOL_NAMES_TABLE:
        return text.replace(clean, TOOL_NAMES_TABLE[clean])

    # Tool function descriptions (cells starting with common patterns)
    if clean.startswith("Remove the") or clean.startswith("Hold the") or \
       clean.startswith("Close the") or clean.startswith("To remove"):
        result = translate_procedural_text(clean)
        if result != clean:
            return text.replace(clean, result)

    # "Use with NNN" patterns in tool tables
    if clean.startswith("Use with"):
        result = clean.replace(" and ", " и ").replace(" and", " и")
        result = result.replace("Use with", "Используется с")
        return text.replace(clean, result)

    # ── Protective treatment table cells ──
    # Broad detection: any cell that looks like protective treatment instructions.
    # Apply PROCEDURAL_VOCAB directly WITHOUT quality gate, because these cells
    # naturally have many spec codes (PCS-, M-DLPS, AMS, etc.) that stay English.
    _pt_keywords = (
        "apply ", "paint", "primer", "refer ",
        "chromic acid", "passivate", "anodise", "anodize",
        "do not", "not paint", "not cadmium",
        "protective treatment", "cadmium plate",
        "chromium plate", "chromium plated", "alocrom",
        "spotface", "areas ", "the thread", "the bush",
        "the hole", "the face", "the lubrication",
        "the identification", "the bearing", "the screw",
        "split pin", "thick", "thickness",
        "externally", "internally", "all over", "but not",
        "including", "optional", "permitted",
        "bore", "bores", "flange", "radius", "chrome",
        "plating", "deposit", "overlap", "grinding",
        "chamfer",
    )
    _clean_lower = clean.lower()
    if any(kw in _clean_lower for kw in _pt_keywords):
        # Try exact match first (already checked above, but just in case)
        if clean in FIXED:
            return text.replace(clean, FIXED[clean])
        # Apply PROCEDURAL_VOCAB directly, no quality gate
        result = clean
        for en, ru in PROCEDURAL_VOCAB:
            result = result.replace(en, ru)
        if result != clean:
            return text.replace(clean, result)

    # Cells starting with spec codes that continue protective treatment text
    # e.g. "M-DLPS100-2. Do not include...", "PCS-2500. Apply primer..."
    if re.match(r'^(M-DL|PCS-|IFC\s|MIL-|AMS\d|NCT\s|DEF\s)', clean) and \
       re.search(r'[A-Za-z]{4,}', clean[8:] if len(clean) > 8 else ''):
        result = clean
        for en, ru in PROCEDURAL_VOCAB:
            result = result.replace(en, ru)
        if result != clean:
            return text.replace(clean, result)

    # Material names (NDT tables and Table 601) — BEFORE page range handlers
    if clean.startswith("Aluminium") or clean.startswith("Stainless steel") or \
       clean.startswith("Stainless Steel") or clean.startswith("Steel") or \
       clean.startswith("Type ") or clean == "Nitrogen" or \
       clean == "Hydraulic fluid" or clean == "White spirit" or \
       clean.startswith("Aluminium alloy"):
        result = clean.replace("Stainless steel", "Нержавеющая сталь")
        result = result.replace("Stainless Steel", "Нержавеющая сталь")
        result = result.replace("Aluminium alloy", "Алюминиевый сплав")
        result = result.replace("Aluminium Alloy", "Алюминиевый сплав")
        result = result.replace("Steel", "Сталь")
        result = result.replace("Type ", "Тип ")
        # Also handle "or" and "to" in material specs
        result = re.sub(r'\bor\b', 'или', result)
        result = re.sub(r'\bto\b', 'до', result)
        if result != clean:
            return text.replace(clean, result)

    # NDT classification phrases
    if "Inclusion class" in clean or "chromium plate" in clean or "Chromium plated" in clean:
        result = translate_procedural_text(clean)
        if result != clean:
            return text.replace(clean, result)

    # Test descriptions in equipment tables
    if "tests" in clean.lower() or "resistance" in clean.lower() or \
       "switch" in clean.lower() or "target" in clean.lower():
        result = translate_procedural_text(clean)
        if result != clean:
            return text.replace(clean, result)

    # "NNN to NNN and" continuation pattern
    if re.match(r'^[\d\s]+to\s+[\d\s]+and$', clean):
        return text.replace(" to ", " до ").replace(" and", " и")

    # Page number ranges with "to": "601 to 603", "504, 507 to 512", "704 to" (continuation)
    if re.search(r'\bto\b', clean) and re.search(r'\d', clean):
        result = re.sub(r'\bto\b', 'до', text)
        result = re.sub(r'\band\b', 'и', result)
        return result

    # Page number ranges with "and" only: "709 and 729"
    if re.match(r'^[\d,\.\s]+(and\s+[\d,\.\s]*)+$', clean):
        return re.sub(r'\band\b', 'и', text)

    # Generic fallback to main translator
    result = translate_text(text)
    return result


# ══════════════════════════════════════════════════════════════════════════════
#  XML-LEVEL TEXT REPLACEMENT WITH FORMATTING PRESERVATION
# ══════════════════════════════════════════════════════════════════════════════

def replace_paragraph_text(para, new_text: str):
    """Replace all text in a paragraph while preserving formatting.

    Works at XML level to handle text inside hyperlinks and other structures.
    Handles tabs properly: if original XML has <w:tab/> elements, splits text
    at tab positions so each segment goes into the correct w:t element.
    Removes empty runs after replacement to keep XML clean.
    """
    element = para._element

    all_t_elements = list(element.iter(qn('w:t')))
    if not all_t_elements:
        return

    # Handle tabs
    tab_runs = []
    for r_elem in element.iter(qn('w:r')):
        if r_elem.find(qn('w:tab')) is not None:
            tab_runs.append(r_elem)

    if tab_runs and '\t' in new_text:
        segments = new_text.split('\t')
        t_groups = [[]]
        for child in element.iter():
            if child.tag == qn('w:r') and child in tab_runs:
                t_groups.append([])
            elif child.tag == qn('w:t'):
                t_groups[-1].append(child)

        for seg_idx, segment in enumerate(segments):
            if seg_idx < len(t_groups) and t_groups[seg_idx]:
                t_groups[seg_idx][0].text = segment
                t_groups[seg_idx][0].set(qn('xml:space'), 'preserve')
                for t_elem in t_groups[seg_idx][1:]:
                    _remove_run_if_empty(t_elem)
            elif seg_idx < len(t_groups):
                if seg_idx > 0 and seg_idx - 1 < len(tab_runs) and segment.strip():
                    tab_run = tab_runs[seg_idx - 1]
                    new_t = etree.SubElement(tab_run, qn('w:t'))
                    new_t.text = segment
                    new_t.set(qn('xml:space'), 'preserve')

        for gi in range(len(segments), len(t_groups)):
            for t_elem in t_groups[gi]:
                _remove_run_if_empty(t_elem)

        needed_tabs = len(segments) - 1
        if len(tab_runs) > needed_tabs:
            for extra_tab_run in tab_runs[needed_tabs:]:
                parent = extra_tab_run.getparent()
                if parent is not None:
                    parent.remove(extra_tab_run)
        return

    # Strategy: put all text in the first non-empty w:t, clear others
    ref_idx = 0
    for i, t_elem in enumerate(all_t_elements):
        if t_elem.text and t_elem.text.strip():
            ref_idx = i
            break

    all_t_elements[ref_idx].text = new_text
    all_t_elements[ref_idx].set(qn('xml:space'), 'preserve')

    for i, t_elem in enumerate(all_t_elements):
        if i != ref_idx:
            _remove_run_if_empty(t_elem)


def _remove_run_if_empty(t_elem):
    """Remove a w:t element's parent run if it's empty, otherwise clear text."""
    run_elem = t_elem.getparent()
    if run_elem is not None and run_elem.tag == qn('w:r'):
        children = [c for c in run_elem if c.tag != qn('w:rPr')]
        if len(children) == 1 and children[0] is t_elem:
            run_parent = run_elem.getparent()
            if run_parent is not None:
                run_parent.remove(run_elem)
                return
    t_elem.text = ""


def process_paragraph(para):
    """Translate a paragraph's text content."""
    original = para.text
    if not original or not original.strip():
        return
    translated = translate_text(original)
    if translated != original:
        replace_paragraph_text(para, translated)


def process_table_cell(cell):
    """Translate all paragraphs within a table cell."""
    for para in cell.paragraphs:
        original = para.text
        if not original or not original.strip():
            continue
        translated = translate_table_cell_text(original)
        if translated != original:
            replace_paragraph_text(para, translated)


# ══════════════════════════════════════════════════════════════════════════════
#  FONT SIZE ADJUSTMENT
# ══════════════════════════════════════════════════════════════════════════════

def _get_font_size_from_element(element):
    """Get font size from a w:r element's w:rPr/w:sz."""
    rPr = element.find(qn('w:rPr'))
    if rPr is not None:
        sz = rPr.find(qn('w:sz'))
        if sz is not None:
            val = sz.get(qn('w:val'))
            if val:
                return int(val)  # half-points
    return None


def _set_font_size_on_element(element, half_points: int):
    """Set font size on a w:r element."""
    rPr = element.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(element, qn('w:rPr'))
        element.insert(0, rPr)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), str(half_points))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = etree.SubElement(rPr, qn('w:szCs'))
    szCs.set(qn('w:val'), str(half_points))


def adjust_font_if_needed(para, original_text: str, new_text: str,
                          default_half_pts: int = 18, in_table: bool = False):
    """Reduce font size if Russian text is significantly longer."""
    if not original_text.strip() or not new_text.strip():
        return
    if new_text == original_text:
        return

    ratio = len(new_text) / max(len(original_text), 1)
    threshold = 1.2 if in_table else 1.3
    if ratio <= threshold:
        return

    element = para._element
    runs = list(element.iter(qn('w:r')))
    if not runs:
        return

    for run_elem in runs:
        current_hp = _get_font_size_from_element(run_elem)
        if current_hp is None:
            current_hp = default_half_pts

        if ratio > 1.7:
            factor = 0.78
        elif ratio > 1.5:
            factor = 0.82
        elif ratio > 1.3:
            factor = 0.88
        else:
            factor = 0.92

        new_hp = max(int(current_hp * factor), 12)  # minimum 6pt
        _set_font_size_on_element(run_elem, new_hp)


# ══════════════════════════════════════════════════════════════════════════════
#  POST-PROCESSING: FORMATTING FIXES
# ══════════════════════════════════════════════════════════════════════════════

def post_process_formatting(doc, table_font_half_pts: int = 18) -> int:
    """Fix formatting issues after translation. Returns count of fixes applied.

    table_font_half_pts: target font size for ALL table cell text in half-points.
                         18 = 9pt. Set to 0 to skip font normalization.
    """
    fix_count = 0

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    text = para.text
                    if not text:
                        continue

                    # Fix 1: Remove trailing tabs in TOC repair entries
                    if text.rstrip() != text.rstrip('\t') and '\t' in text:
                        stripped = text.rstrip('\t')
                        if stripped != text:
                            t_elems = list(para._element.iter(qn('w:t')))
                            for t_elem in t_elems:
                                if t_elem.text and t_elem.text.endswith('\t'):
                                    t_elem.text = t_elem.text.rstrip('\t')
                                    fix_count += 1

                    # Fix 2: Normalize table font size to target
                    if table_font_half_pts > 0:
                        element = para._element
                        for run_elem in element.iter(qn('w:r')):
                            current_hp = _get_font_size_from_element(run_elem)
                            if current_hp != table_font_half_pts:
                                _set_font_size_on_element(run_elem, table_font_half_pts)
                                fix_count += 1

    return fix_count


# ══════════════════════════════════════════════════════════════════════════════
#  VERIFICATION
# ══════════════════════════════════════════════════════════════════════════════

def verify_translation(dst_path: str, verbose: bool = True) -> dict:
    """Verify translated document for remaining English text.

    Returns dict with 'untranslated_paragraphs' and 'untranslated_cells' lists.
    """
    doc = Document(str(dst_path))
    result = {"untranslated_paragraphs": [], "untranslated_cells": []}

    # Check paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if re.search(r'[A-Za-z]{4,}', text) and not is_only_numbers_or_codes(text):
            if any(skip in text for skip in [
                "Safran Landing Systems UK",
                "Cheltenham Road",
                "201587", "EDES2", "MAF1",
                "(c)AC", "A320-", "A321-",
                "32-12-22", "AMS-", "ASTM",
                "MLG", "NLG", "M-DLNDT", "M-DLPS",
                "EASA", "PCS-", "BS3643",
                "MT1025", "MT1026", "Sermetel", "SERMETEL",
                "IFC 40-", "MIL-A-", "PLACES",
            ]):
                continue
            if re.search(r'[А-Яа-яЁё]{2,}', text):
                continue
            result["untranslated_paragraphs"].append((i, text[:100]))

    if verbose:
        if result["untranslated_paragraphs"]:
            print(f"Found {len(result['untranslated_paragraphs'])} potentially untranslated paragraphs:")
            for idx, txt in result["untranslated_paragraphs"][:20]:
                safe = txt.encode(sys.stdout.encoding or 'utf-8', errors='replace').decode(sys.stdout.encoding or 'utf-8', errors='replace')
                print(f"  P{idx}: {safe}")
        else:
            print("All paragraphs appear to be translated!")

    # Check table cells
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    if re.search(r'[A-Za-z]{4,}', text) and not is_only_numbers_or_codes(text):
                        if any(skip in text for skip in [
                            "SAFRAN LANDING SYSTEMS",
                            "Safran Landing Systems",
                            "201587", "EDES2", "MAF1",
                            "(c)AC", "A320-", "A321-",
                            "32-12-22", "AMS-", "ASTM",
                            "MLG", "NLG", "Messier",
                            "SERVICE BULLETIN", "M-DLNDT", "M-DLPS",
                            "EASA", "PCS-", "BS3643",
                            "MT1025", "MT1026", "Cond ",
                            "Sermetel", "IFC 40-", "MIL-A-",
                            "AMS5", "AMS6", "AMS2",
                        ]):
                            continue
                        if re.search(r'[А-Яа-яЁё]{3,}', text):
                            continue
                        result["untranslated_cells"].append((f"T{ti+1}R{ri+1}C{ci+1}", text[:120]))

    if verbose:
        if result["untranslated_cells"]:
            print(f"\nFound {len(result['untranslated_cells'])} potentially untranslated table cells:")
            for loc, txt in result["untranslated_cells"][:20]:
                safe = txt.encode(sys.stdout.encoding or 'utf-8', errors='replace').decode(sys.stdout.encoding or 'utf-8', errors='replace')
                print(f"  {loc}: {safe}")
        else:
            print("All table cells appear to be translated!")

    return result


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN TRANSLATION PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

def translate_document(src: str, dst: str, table_font_half_pts: int = 18):
    """Translate a CMM document from English to Russian.

    Args:
        src: Path to the source .docx file
        dst: Path for the output translated .docx file
        table_font_half_pts: Target font size for table cells (18 = 9pt).
                             Set to 0 to skip font normalization.
    """
    src_path = Path(src)
    dst_path = Path(dst)

    print(f"Loading: {src_path}")
    doc = Document(str(src_path))

    # ── Process paragraphs ──
    translated_count = 0
    for i, para in enumerate(doc.paragraphs):
        original = para.text
        if not original or not original.strip():
            continue
        translated = translate_text(original)
        if translated != original:
            replace_paragraph_text(para, translated)
            adjust_font_if_needed(para, original, translated)
            translated_count += 1

    print(f"Translated {translated_count} paragraphs")

    # ── Process tables ──
    table_translated = 0
    for ti, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text
                    if not original or not original.strip():
                        continue
                    translated = translate_table_cell_text(original)
                    if translated != original:
                        replace_paragraph_text(para, translated)
                        adjust_font_if_needed(para, original, translated, in_table=True)
                        table_translated += 1

    print(f"Translated {table_translated} table cells")

    # ── Post-processing: fix formatting issues ──
    fix_count = post_process_formatting(doc, table_font_half_pts)
    print(f"Applied {fix_count} formatting fixes")

    # ── Process headers and footers ──
    # Text in headers/footers is typically inside textboxes (w:txbxContent),
    # not in simple paragraphs. We process both levels.
    hf_count = 0
    from docx.text.paragraph import Paragraph as _Paragraph
    for section in doc.sections:
        hf_elements = [
            section.header, section.first_page_header, section.even_page_header,
            section.footer, section.first_page_footer, section.even_page_footer,
        ]
        for hf in hf_elements:
            if hf.is_linked_to_previous:
                continue
            # Process textboxes inside header/footer
            for txbx in hf._element.iter(qn('w:txbxContent')):
                for p_elem in txbx.iter(qn('w:p')):
                    para = _Paragraph(p_elem, doc)
                    original = para.text
                    if original and original.strip():
                        translated = translate_hf_text(original)
                        if translated != original:
                            replace_paragraph_text(para, translated)
                            hf_count += 1
            # Also check direct paragraphs (fallback)
            for para in hf.paragraphs:
                original = para.text
                if original and original.strip():
                    translated = translate_hf_text(original)
                    if translated != original:
                        replace_paragraph_text(para, translated)
                        hf_count += 1

    print(f"Translated {hf_count} header/footer elements")

    # ── Process textboxes ──
    tb_count = 0
    body = doc.element.body
    for txbx in body.iter(qn('w:txbxContent')):
        for p in txbx.iter(qn('w:p')):
            from docx.text.paragraph import Paragraph
            para = Paragraph(p, doc)
            original = para.text
            if original and original.strip():
                translated = translate_text(original)
                if translated != original:
                    replace_paragraph_text(para, translated)
                    adjust_font_if_needed(para, original, translated, in_table=True)
                    tb_count += 1

    print(f"Translated {tb_count} textbox elements")

    # ── Save ──
    os.makedirs(str(dst_path.parent), exist_ok=True)
    doc.save(str(dst_path))
    print(f"Saved: {dst_path}")

    # ── Verification ──
    print("\n=== Verification ===")
    verify_translation(str(dst_path))

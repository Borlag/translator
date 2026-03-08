"""
Translate original_new_part1.docx (EN→RU) with formatting preservation.
Aviation CMM technical translation using glossary-based terminology.
"""
import copy
import os
import re
import sys
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn

# ── Paths ──────────────────────────────────────────────────────────────────
SRC = Path(r"C:\Users\Urdul\Desktop\project\translator\for_test\new_formating\section\original_new_part1.docx")
DST = Path(r"C:\Users\Urdul\Desktop\project\translator\for_test\new_formating\section_translate\opus\original_new_part1.docx")

# ── Component name glossary (EN → RU) ─────────────────────────────────────
COMPONENT_NAMES = {
    "Main Landing Gear Leg": "Стойка основного шасси",
    "Main Landing Gear": "Основное шасси",
    "MLG Shock Strut": "Стойка амортизатора основного шасси",
    "Main Fitting Subassembly": "Сборка корпуса стойки",
    "Main Fitting": "Корпус стойки",
    "Sliding Tube": "Скользящая труба",
    "Lower Torque Link": "Нижний шлиц-шарнир",
    "Upper Torque Link": "Верхний шлиц-шарнир",
    "Upper Slave Link": "Верхнее ведомое звено",
    "Lower Slave Link": "Нижнее ведомое звено",
    "Upper Diaphragm Tube Subassembly": "Сборка верхней диафрагменной трубы",
    "Upper Diaphragm Tube": "Верхняя диафрагменная труба",
    "Pivot Bracket": "Поворотный кронштейн",
    "Upper Pivot Bracket": "Верхний поворотный кронштейн",
    "Harness Support Bracket": "Кронштейн крепления жгута",
    "Lock Stay Cardan": "Кардан фиксатора",
    "Transfer Block": "Переходный блок",
    "Transfer block": "Переходный блок",
    "Spherical Bearing": "Сферический подшипник",
    "Pintle pin": "Штифт навеса стойки",
    "Pintle Pin": "Штифт навеса стойки",
    "Forward Pintle Pin": "Передний штифт навеса стойки",
    "Retaining Pin": "Стопорный штифт",
    "Uplock Pin ain Fitting": "Штифт замка убранного положения корпуса стойки",
    "Cylinder": "Цилиндр",
    "Bracket": "Кронштейн",
    "Pin": "Штифт",
    "Spacer": "Проставка",
    "Drag-arm Spacer": "Проставка тяги",
    "Slave Link": "Ведомое звено",
    "Valve Stem": "Шток клапана",
    "Inflation Valve": "Клапан зарядки",
    # Lowercase variants
    "Sliding tube": "Скользящая труба",
    "Transfer block": "Переходный блок",
    "Pintle pin": "Штифт навеса стойки",
    # Repair procedure descriptions
    "Bush": "Втулка",
    "Bushes": "Втулки",
    "Machining": "Механическая обработка",
    "Installation": "Установка",
    "Machining and Installation": "Механическая обработка и установка",
    "Machining and installation": "Механическая обработка и установка",
}

# ── Fixed phrase translations ──────────────────────────────────────────────
FIXED = {
    # Title page
    "MAIN LANDING GEAR LEG": "СТОЙКА ОСНОВНОГО ШАССИ",
    "PART NUMBER": "НОМЕР ДЕТАЛИ",
    "COMPONENT MAINTENANCE MANUAL WITH": "РУКОВОДСТВО ПО ТЕХНИЧЕСКОМУ ОБСЛУЖИВАНИЮ КОМПОНЕНТОВ С",
    "ILLUSTRATED PARTS LIST": "ИЛЛЮСТРИРОВАННЫМ ПЕРЕЧНЕМ ДЕТАЛЕЙ",
    "STATEMENT OF INITIAL CERTIFICATION": "СВИДЕТЕЛЬСТВО О ПЕРВИЧНОЙ СЕРТИФИКАЦИИ",
    "INTENTIONALLY BLANK": "НАМЕРЕННО ОСТАВЛЕНО ПУСТЫМ",

    # Structural headings
    "NEW/REVISED PAGES": "НОВЫЕ/ПЕРЕСМОТРЕННЫЕ СТРАНИЦЫ",
    "REVISION RECORD": "ЗАПИСЬ ИЗМЕНЕНИЙ",
    "RECORD OF REVISIONS": "ЗАПИСЬ ИЗМЕНЕНИЙ",
    "RECORD OF TEMPORARY REVISIONS": "ЗАПИСЬ ВРЕМЕННЫХ ИЗМЕНЕНИЙ",
    "LIST OF SERVICE BULLETINS": "СПИСОК СЕРВИСНЫХ БЮЛЛЕТЕНЕЙ",
    "LIST OF SERVICE BULLETINS (Continued)": "СПИСОК СЕРВИСНЫХ БЮЛЛЕТЕНЕЙ (Продолжение)",
    "LIST OF EFFECTIVE PAGES": "ПЕРЕЧЕНЬ ДЕЙСТВУЮЩИХ СТРАНИЦ",
    "TABLE OF CONTENTS": "СОДЕРЖАНИЕ",
    "ILLUSTRATIONS": "ИЛЛЮСТРАЦИИ",
    "Issued by": "Выпущено",
    "REV NO.": "НОМ. РЕД.",
    "ISSUE DATE": "ДАТА ВЫПУСКА",
    "INSERTION DATE": "ДАТА ВСТАВКИ",
    "BY": "КЕМ",
    "REMOVAL DATE": "ДАТА УДАЛЕНИЯ",

    # TOC section headings
    "Assembly (Including Storage)": "Сборка (включая хранение)",
    "General": "Общие сведения",
    "Procedure": "Процедура",
    "Storage": "Хранение",
    "Fits and Clearances": "Посадки и зазоры",
    "Fits and Clearances Definitions": "Определения посадок и зазоров",
    "Remarks": "Примечания",
    "Torque Data": "Данные моментов затяжки",
    "Special Tools, Fixtures and Equipment": "Специальные инструменты, приспособления и оборудование",
    "Illustrated Parts List": "Иллюстрированный перечень деталей",
    "Introduction": "Введение",
    "How To Use This Illustrated Parts List": "Как пользоваться иллюстрированным перечнем деталей",
    "Vendor Codes, Names and Addresses": "Коды, наименования и адреса поставщиков",
    "Numerical Index": "Числовой указатель",
    "Detailed Parts List": "Подробный список деталей",

    # Figure captions
    "Diagram of Operation": "Схема работы",
    "Electrical Bonding Resistance Test Points": "Точки проверки сопротивления электрического соединения",
    "Protective Treatment": "Защитная обработка",
}

# ── Table header translations ──────────────────────────────────────────────
TABLE_HEADERS = {
    "Subject Reference": "Ссылка на раздел",
    "Remove and Destroy Pages": "Удалить и уничтожить страницы",
    "Insert New/Revised": "Вставить новые/пересмотренные",
    "Pages": "Страницы",
    "Dated": "Дата",
    "Reason for Change": "Причина изменения",
    "REV.\nNo.": "НОМ.\nРЕД.",
    "ISSUE DATE": "ДАТА ВЫПУСКА",
    "DATE INSERTED": "ДАТА ВСТАВКИ",
    "PAGE NUMBER": "НОМЕР СТРАНИЦЫ",
    "BY": "КЕМ",
    "DATE REMOVED": "ДАТА УДАЛЕНИЯ",
    "SB NUMBER": "НОМЕР SB",
    "SB TITLE": "НАЗВАНИЕ SB",
    "SB REVISION NUMBER": "НОМЕР РЕДАКЦИИ SB",
    "DATE INCORPORATED INTO MANUAL": "ДАТА ВКЛЮЧЕНИЯ В РУКОВОДСТВО",
    "COVER SB NO.": "НОМЕР ОХВАТ. SB",
    "PART NUMBER": "НОМЕР ДЕТАЛИ",
    "DASH NO.": "НОМЕР ТИРЕ",
    "MOD. STRIKE NO.": "НОМЕР СНЯТИЯ МОД.",
    "SAFRAN LANDING SYSTEMS MODIFICATION NUMBER": "НОМЕР МОДИФИКАЦИИ SAFRAN LANDING SYSTEMS",
    "SAFRAN LANDING SYSTEMS\nSERVICE BULLETIN NUMBER": "НОМЕР СЕРВИСНОГО БЮЛЛЕТЕНЯ\nSAFRAN LANDING SYSTEMS",
    "SAFRAN LANDING SYSTEMS SERVICE BULLETIN NUMBER": "НОМЕР СЕРВИСНОГО БЮЛЛЕТЕНЯ SAFRAN LANDING SYSTEMS",
    "Page": "Страница",
    "Date": "Дата",
    "Subject": "Раздел",
    "Initial Issue": "Первоначальный выпуск",
    "No effect": "Без изменений",
}

# ── Revision table section name translations ───────────────────────────────
SECTION_NAMES = {
    "Record of Revisions": "Запись изменений",
    "Unit Identification Chart": "Таблица идентификации изделия",
    "List of Effective Pages": "Перечень действующих страниц",
    "Table of Contents": "Содержание",
    "Disassembly": "Разборка",
    "Check": "Проверка",
    "Repair": "Ремонт",
    "Assembly (Including Storage)": "Сборка (включая хранение)",
    "Fits and Clearances": "Посадки и зазоры",
    "Special Tools, Fixtures and Equipment": "Специальные инструменты, приспособления и оборудование",
    "Illustrated Parts List": "Иллюстрированный перечень деталей",
    # List of Effective Pages section names
    "Subject": "Раздел",
    "(Continued)": "(Продолжение)",
    "Testing and Fault": "Проверка и поиск неисправностей",
    "Isolation": "Локализация",
    "Isolation (Continued)": "Локализация (Продолжение)",
    "Cleaning": "Очистка",
    "Repair (Continued)": "Ремонт (Продолжение)",
    "Assembly (Including": "Сборка (включая",
    "Storage)": "хранение)",
    "Storage) (Continued)": "хранение) (Продолжение)",
    "Special Tools, Fixtures": "Специальные инструменты, приспособления",
    "and Equipment": "и оборудование",
}

# ── Reason for change phrase translations ──────────────────────────────────
REASON_PHRASES = {
    "Updated revision status": "Обновлён статус редакции",
    "Updated pages": "Обновлены страницы",
    "Added content": "Добавлено содержание",
    "Updated page numbers": "Обновлены номера страниц",
    "Updated figures": "Обновлены рисунки",
    "Updated figure numbers": "Обновлены номера рисунков",
    "Updated tables": "Обновлены таблицы",
    "Added paras": "Добавлены пункты",
    "Added para": "Добавлен пункт",
    "Updated fig-items": "Обновлены поз. рисунков",
    "Updated fig-item": "Обновлена поз. рисунка",
    "Added fig-item": "Добавлена поз. рисунка",
    "in para": "в пункте",
    "in paras": "в пунктах",
    "and": "и",
    "only": "только",
    "only in para": "только в пункте",
    "Updated Messier-Dowty Limited to Safran Landing Systems": "Обновлено Messier-Dowty Limited на Safran Landing Systems",
    "Updated Messier-Dowty Limited to Safran Landing Sy": "Обновлено Messier-Dowty Limited на Safran Landing Sy",
    "Updated Messier-Dowty": "Обновлено Messier-Dowty",
    "Limited to Safran Landing Systems": "Limited на Safran Landing Systems",
    "Added repair no": "Добавлен ремонт №",
    "Added repair no.": "Добавлен ремонт №",
    "Added caution at para": "Добавлено предупреждение в пункте",
    "Updated material specification": "Обновлена спецификация материала",
    "Updated materia": "Обновлены материа",
    "Updated paras": "Обновлены пункты",
    "Updated IPL figs": "Обновлены рисунки ИПД",
    "to include Ref": "для включения ссылок",
    "Updated figure numbers": "Обновлены номера рисунков",
    "Updated figure": "Обновлён рисунок",
    "Updated figures": "Обновлены рисунки",
    "Added figure": "Добавлен рисунок",
    "Added Ref. Codes": "Добавлены коды ссылок",
    "Updated tables": "Обновлены таблицы",
    "Updated table": "Обновлена таблица",
    # Standalone "Updated" at end of line (continuation to next paragraph)
    ". Updated": ". Обновлены",
    "Updated IPL fig": "Обновлён рисунок ИПД",
    "figure": "рисунок",
    "figures": "рисунки",
    "Codes": "Коды",
    "details": "детали",
}

# ── SB title translations ─────────────────────────────────────────────────
SB_TITLE_PARTS = {
    "MLG - Installation of stub bolt subassembly for th": "ОШ — Установка сборки болта-вставки для",
    "MLG - To allow an increase in aircraft maximum tak": "ОШ — Для обеспечения увеличения максимальной взлётной",
    "MLG -To add tracking numbers to parts listed in Ai": "ОШ — Добавление учётных номеров к деталям, указанным в",
    "MLG - Installation of a 201585 series MLG Leg and ": "ОШ — Установка стойки ОШ серии 201585 и",
    "MLG - Introduction of a new lower bearing subassem": "ОШ — Введение новой сборки нижнего подшипника",
    "MLG - Introduction of new charging labels": "ОШ — Введение новых этикеток зарядки",
    "MLG - Introduction of new 1M and 2M Axle harnesses": "ОШ — Введение новых жгутов осей 1М и 2М",
    "MLG - Introduction of new 1M and 2M Leg Harness an": "ОШ — Введение новых жгутов стойки 1М и 2М и",
    "MLG Leg-Introduction of new retaining pins and a n": "Стойка ОШ — Введение новых стопорных штифтов и н",
    "MLG Leg - Introduction of new retaining pins for t": "Стойка ОШ — Введение новых стопорных штифтов для",
    "MLG Leg - Introduction of a new lower bearing suba": "Стойка ОШ — Введение новой сборки нижнего подшипника",
    "MLG Leg - Barkhausen Noise Inspection of Main Land": "Стойка ОШ — Контроль методом шума Баркгаузена корпуса",
    "MLG Leg - Introduction of a new Main Fitting": "Стойка ОШ — Введение нового корпуса стойки",
    "MLG Leg - Introduction of a new torque link damper": "Стойка ОШ — Введение нового демпфера шлиц-шарнира",
    "MLG Leg - Introduction of a new main fitting subas": "Стойка ОШ — Введение новой сборки корпуса стойки",
    "MLG - Introduction of a new upper pivot bracket": "ОШ — Введение нового верхнего поворотного кронштейна",
    "MLG - Introduction of a new changeover valve stem ": "ОШ — Введение нового штока переключающего клапана",
    "MLG Complete - Modification of the transfer block ": "ОШ в сборе — Модификация переходного блока",
    "MLG - Conversion of low - friction lower - bearing": "ОШ — Замена низкофрикционного нижнего подшипника",
    "MLG complete - Introduction of a new transfer bloc": "ОШ в сборе — Введение нового переходного блока",
}


# ══════════════════════════════════════════════════════════════════════════
#  Helper: translate paragraph/cell text
# ══════════════════════════════════════════════════════════════════════════

def is_only_numbers_or_codes(text: str) -> bool:
    """Return True if text is just numbers, part numbers, dates, page refs."""
    stripped = text.strip()
    if not stripped:
        return True
    # Pure numbers, dates, page refs, codes
    if re.match(r'^[\d\s,\-/\.\t©®]+$', stripped):
        return True
    # Part numbers like 201587001
    if re.match(r'^[\d\s,]+$', stripped):
        return True
    # Dates like "Dec 6/2019", "Mar 18/2025"
    if re.match(r'^(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d+/\d+$', stripped):
        return True
    # Tab-separated page ref + date: "601\tDec 6/2019" or "602\tMar 18/2025"
    if re.match(r'^\d+\t(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d+/\d+$', stripped):
        return True
    # ATA codes like 32-12-22
    if re.match(r'^\d{2}-\d{2}-\d{2}$', stripped):
        return True
    # Single short codes
    if re.match(r'^[A-Z]?\d+[\-\.]\d+', stripped) and len(stripped) < 15:
        return True
    return False


def translate_toc_entry(text: str) -> str:
    """Translate a TOC entry preserving dot leaders and page numbers."""
    # Pattern: "ComponentName . . . . . . . . page_ref\tPageNum"
    # Also: "Repair No. X-Y ComponentName . . . . Repair No. X-Y\tPageNum"

    # Handle "Repair No." pattern in TOC
    m = re.match(r'^(Repair No\.\s*\d+-\d+)\s+(.*?)(\s*\.[\s\.]+.*)', text)
    if m:
        prefix = m.group(1)
        comp_name = m.group(2).strip()
        rest = m.group(3)
        # Translate component name
        translated_comp = translate_component_name(comp_name)
        prefix_ru = prefix.replace("Repair No.", "Ремонт №")
        # Also translate the right-side "Repair No." in the rest
        rest_ru = rest.replace("Repair No.", "Ремонт №")
        return f"{prefix_ru} {translated_comp}{rest_ru}"

    # Handle "Page Repair No." pattern
    m = re.match(r'^Page\s+(Repair No\.\s*\d+-\d+)\s+(.*?)(\s*\.[\s\.]+.*)', text)
    if m:
        prefix = m.group(1)
        comp_name = m.group(2).strip()
        rest = m.group(3)
        translated_comp = translate_component_name(comp_name)
        prefix_ru = prefix.replace("Repair No.", "Ремонт №")
        rest_ru = rest.replace("Repair No.", "Ремонт №")
        return f"Стр. {prefix_ru} {translated_comp}{rest_ru}"

    # Handle figure TOC entries (component name with fig-item refs)
    # e.g., "Main Landing Gear Leg . . . . . . 3"
    m = re.match(r'^(.*?)(\s*\.[\s\.]+\s*\d+.*)$', text)
    if m:
        name_part = m.group(1).strip()
        rest = m.group(2)
        translated_name = translate_component_name(name_part)
        return f"{translated_name}{rest}"

    # Handle tab-separated entries like "Storage\t798.12"
    m = re.match(r'^(.*?)\t(.*)$', text)
    if m:
        name_part = m.group(1).strip()
        page = m.group(2)
        if name_part in FIXED:
            return f"{FIXED[name_part]}\t{page}"
        translated_name = translate_component_name(name_part)
        return f"{translated_name}\t{page}"

    return text


def _translate_suffix(suffix: str) -> str:
    """Translate common suffixes like 'Protective Treatment (Sheet 1)'."""
    result = suffix
    if "Protective Treatment" in result:
        result = result.replace("Protective Treatment", "Защитная обработка")
    if "Sheet" in result:
        result = re.sub(r'Sheet\s+(\d+)', r'Лист \1', result)
    if result.strip() == "Withdrawn" or result.strip() == "(Withdrawn)":
        result = result.replace("Withdrawn", "Отозвано")
    if "Only" in result:
        result = result.replace("Only", "Только")
    return result


def _lookup_base_name(name: str) -> str:
    """Look up base component name in glossaries."""
    if name in COMPONENT_NAMES:
        return COMPONENT_NAMES[name]
    if name in FIXED:
        return FIXED[name]
    return name


def translate_component_name(name: str) -> str:
    """Translate a component name from the glossary."""
    # Check exact match first
    if name in COMPONENT_NAMES:
        return COMPONENT_NAMES[name]
    if name in FIXED:
        return FIXED[name]

    # ── Pattern: "Name (refs) ... or/and (refs) ... - Suffix" ──
    # Most general: split on last " - " to get name_with_refs and suffix
    dash_idx = name.rfind(" - ")
    if dash_idx > 0:
        name_part = name[:dash_idx].strip()
        suffix_part = name[dash_idx + 3:].strip()
        suffix_ru = _translate_suffix(suffix_part)

        # Extract base name (before first parenthesis)
        paren_idx = name_part.find("(")
        if paren_idx > 0:
            base = name_part[:paren_idx].strip()
            refs = name_part[paren_idx:]
            base_ru = _lookup_base_name(base)
            refs_ru = refs.replace("and", "и").replace("or", "или")
            return f"{base_ru} {refs_ru} — {suffix_ru}"
        else:
            base_ru = _lookup_base_name(name_part)
            return f"{base_ru} — {suffix_ru}"

    # ── Pattern: "Name (refs)" without suffix ──
    paren_idx = name.find("(")
    if paren_idx > 0:
        base = name[:paren_idx].strip()
        refs = name[paren_idx:]
        base_ru = _lookup_base_name(base)
        refs_ru = _translate_suffix(refs)
        refs_ru = refs_ru.replace(" and ", " и ").replace(" or ", " или ")
        if base_ru != base:
            return f"{base_ru} {refs_ru}"

    # ── Pattern: "(Withdrawn)" suffix ──
    if name.endswith("(Withdrawn)"):
        base = name.replace("(Withdrawn)", "").strip()
        base_ru = translate_component_name(base)
        return f"{base_ru} (Отозвано)"

    return name


def translate_sb_title(text: str) -> str:
    """Translate SB title text."""
    # Check for exact prefix match (titles are truncated at ~50 chars in tables)
    for en, ru in SB_TITLE_PARTS.items():
        if text.startswith(en):
            return ru + text[len(en):]
    return text


def translate_reason(text: str) -> str:
    """Translate reason-for-change text in revision tables."""
    result = text
    # Normalize multi-line text for matching: collapse \n to space for lookup
    normalized = re.sub(r'\n', ' ', result)
    # Fix garbled text like "Safran Lиing Systems" -> "Safran Landing Systems"
    normalized = re.sub(r'Safran L[а-яёА-ЯЁ\w]*ing Systems', 'Safran Landing Systems', normalized)
    # Apply phrase translations longest-first on normalized text
    sorted_phrases = sorted(REASON_PHRASES.keys(), key=len, reverse=True)
    for en in sorted_phrases:
        ru = REASON_PHRASES[en]
        normalized = normalized.replace(en, ru)
    return normalized


def translate_repair_description(text: str) -> str:
    """Translate 'Repair to Component — Process' descriptions."""
    result = text
    # Split on tab to get description and repair number
    parts = result.split('\t')
    desc = parts[0]
    rest_parts = '\t'.join(parts[1:]) if len(parts) > 1 else ""

    # Replace "Repair No." in rest
    rest_parts = rest_parts.replace("Repair No.", "Ремонт №")

    # Translate the description
    desc = desc.replace("Oversize Bushes", "Ремонтные (увеличенные) втулки")
    desc = desc.replace("Oversize Bush (es)", "Ремонтные (увеличенные) втулки")
    desc = desc.replace("Oversize Bush", "Ремонтная (увеличенная) втулка")
    desc = desc.replace("Oversize Lubrication adapter", "Ремонтный (увеличенный) смазочный адаптер")
    desc = desc.replace("Oversize Transfer Dowel", "Ремонтный (увеличенный) переходной штифт")
    desc = desc.replace("Lower Bearing Subassembly", "Сборка нижнего подшипника")
    desc = desc.replace("Machining and Inner Liner Installation", "Механическая обработка и установка внутреннего вкладыша")
    desc = desc.replace("Machining and Liner Installation", "Механическая обработка и установка вкладыша")
    desc = desc.replace("Repair Bearing", "Ремонт подшипника")
    desc = desc.replace("Repair to ", "Ремонт ")
    desc = desc.replace("Repair Bushes", "Ремонт втулок")
    desc = desc.replace("Repair Bush", "Ремонт втулки")

    # Translate component names in the description
    for en, ru in sorted(COMPONENT_NAMES.items(), key=lambda x: -len(x[0])):
        desc = desc.replace(en, ru)

    # Translate process descriptions
    desc = desc.replace("Machining and Installation", "Механическая обработка и установка")
    desc = desc.replace("Machining and installation", "Механическая обработка и установка")
    desc = desc.replace("Machining", "Механическая обработка")
    desc = desc.replace("Installation", "Установка")
    desc = desc.replace("Cadmium Plating", "Кадмирование")
    desc = desc.replace("Rework", "Доработка")

    # Translate connectors
    desc = desc.replace(" — ", " — ")  # already correct
    desc = desc.replace(" and ", " и ")

    if rest_parts:
        return f"{desc}\t{rest_parts}"
    return desc


def translate_long_text(text: str) -> str:
    """Translate longer text blocks (certification, copyright, etc.)."""
    # Certification statement
    if text.startswith("This manual complies with British Civil Airworthiness"):
        return "Данное руководство соответствует требованиям Британских авиационных правил лётной годности, раздел A, глава A5-3."

    if text.startswith("NOTE: The above certification does not apply"):
        return ("ПРИМЕЧАНИЕ: Вышеуказанная сертификация не распространяется на изменения или "
                "дополнения, внесённые после даты первичной сертификации другими "
                "одобренными организациями. Изменения или дополнения, внесённые другими "
                "одобренными организациями, должны быть отдельно сертифицированы и "
                "зарегистрированы на отдельных учётных листах.")

    # Copyright notice
    if text.startswith("This document and all information contained herein"):
        return ("Настоящий документ и вся содержащаяся в нём информация являются "
                "исключительной собственностью Safran Landing Systems (и/или аффилированных компаний).")

    if text.startswith("No intellectual property rights are granted"):
        return ("Передача данного документа или раскрытие его содержания не предоставляет "
                "никаких прав на интеллектуальную собственность. Воспроизведение данного "
                "документа для третьих лиц запрещено без письменного согласия Safran Landing Systems "
                "(и/или соответствующей аффилированной компании).")

    if text.startswith("Record the issue date and insertion date"):
        return ("Запишите дату выпуска и дату вставки данной редакции в Запись изменений "
                "и сохраните данное Письмо о рассылке.")

    if text.startswith("The technical data in this document"):
        return ("Технические данные в настоящем документе (или файле) могут содержать данные США "
                "и подлежать экспортному контролю в соответствии с")

    return text


def translate_header_line(text: str) -> str:
    """Translate header/footer running text."""
    # Header: "PART No. 201587001 AND 201587002 COMPONENT MAINTENANCE MANUAL MAIN LANDING GEAR LEG"
    m = re.match(r'^PART\s+No\.\s+([\d\s]+)AND\s+([\d]+)\s+COMPONENT MAINTENANCE MANUAL\s+MAIN LANDING GEAR LEG$', text)
    if m:
        return f"ДЕТАЛЬ № {m.group(1)}И {m.group(2)} РУКОВОДСТВО ПО ТЕХНИЧЕСКОМУ ОБСЛУЖИВАНИЮ КОМПОНЕНТОВ СТОЙКА ОСНОВНОГО ШАССИ"

    # Copyright header
    if "SAFRAN LANDING SYSTEMS" in text and "SUBSEQUENT REVISION PAGE DATES" in text:
        m2 = re.match(r'[©®\s]*SAFRAN LANDING SYSTEMS\s+(\d+)\s+\(AND SUBSEQUENT REVISION PAGE DATES\)', text)
        if m2:
            return f"© SAFRAN LANDING SYSTEMS {m2.group(1)} (И ПОСЛЕДУЮЩИЕ ДАТЫ ПЕРЕСМОТРА СТРАНИЦ)"
        # Try with copyright symbol
        return text.replace("AND SUBSEQUENT REVISION PAGE DATES", "И ПОСЛЕДУЮЩИЕ ДАТЫ ПЕРЕСМОТРА СТРАНИЦ")

    return text


def translate_text(text: str) -> str:
    """Main translation function. Returns translated text or original if untranslatable."""
    stripped = text.strip()
    if not stripped:
        return text

    # Skip purely numeric/code content
    if is_only_numbers_or_codes(stripped):
        return text

    # Skip already-garbled (mojibake) content
    if re.search(r'[А-Яа-яЁё]{3,}', stripped):
        # Already contains Russian - skip
        return text

    # Preserve leading/trailing whitespace
    leading = text[:len(text) - len(text.lstrip())]
    trailing = text[len(text.rstrip()):]
    core = text.strip()

    # 1. Exact match in FIXED
    if core in FIXED:
        return leading + FIXED[core] + trailing

    # 2. Long text translations
    for start_phrase in [
        "This manual complies",
        "NOTE: The above certification",
        "This document and all information",
        "No intellectual property rights",
        "Record the issue date",
        "The technical data in this document",
    ]:
        if core.startswith(start_phrase):
            return leading + translate_long_text(core) + trailing

    # 3. Header line
    if "COMPONENT MAINTENANCE MANUAL" in core or "SUBSEQUENT REVISION PAGE DATES" in core:
        result = translate_header_line(core)
        if result != core:
            return leading + result + trailing

    # 4. Protective Treatment patterns (check BEFORE dot leaders)
    if "Protective Treatment" in core:
        # Continuation lines like "- Protective Treatment (Sheet 1) . . . .\t640"
        m_cont = re.match(r'^-\s*(Protective Treatment.*?)(\s*\.[\s\.]+.*)$', core)
        if m_cont:
            suffix_ru = _translate_suffix(m_cont.group(1))
            return leading + f"— {suffix_ru}{m_cont.group(2)}" + trailing
        # Standalone "Protective Treatment - Sheet X . . . .\t663"
        m_pt = re.match(r'^(Protective Treatment\s*-?\s*.*?)(\s*\.[\s\.]+.*)$', core)
        if m_pt:
            suffix_ru = _translate_suffix(m_pt.group(1))
            return leading + f"{suffix_ru}{m_pt.group(2)}" + trailing
        # Component - Protective Treatment without dots
        result = translate_component_name(core)
        if result != core:
            return leading + result + trailing

    # 4b. Tab-separated entries (check BEFORE dot leaders for numbered fig entries)
    if "\t" in core:
        # "Repair No. X-Y\t601\tDate" (List of Effective Pages)
        m_lep = re.match(r'^(Repair No\.\s*\d+-\d+)\t(.*)$', core)
        if m_lep:
            prefix = m_lep.group(1).replace("Repair No.", "Ремонт №")
            return leading + f"{prefix}\t{m_lep.group(2)}" + trailing

        # "604\tBlank"
        m_blank = re.match(r'^(\d+)\tBlank$', core)
        if m_blank:
            return leading + f"{m_blank.group(1)}\tПусто" + trailing

        # "Fig.\tPage"
        if core.startswith("Fig."):
            return leading + core.replace("Fig.", "Рис.").replace("Page", "Страница") + trailing

        # Numbered figure TOC entries: "618\tSliding Tube (18-80)..."
        m_fig = re.match(r'^(\d+)\t(.*)$', core)
        if m_fig:
            fig_num = m_fig.group(1)
            rest = m_fig.group(2)
            # Check if rest contains dot leaders (one-line entry)
            if ". . ." in rest:
                # Try: "CompName . . . .\tPageNum" (with trailing tab+page)
                m_dots = re.match(r'^(.*?)(\s*\.[\s\.]+\s*)(\t?\d+.*)$', rest)
                if m_dots:
                    comp_part = m_dots.group(1).strip()
                    dots = m_dots.group(2)
                    page = m_dots.group(3)
                    translated_comp = translate_component_name(comp_part)
                    return leading + f"{fig_num}\t{translated_comp}{dots}{page}" + trailing
                # Dots at end without page number (trailing dots only)
                m_dots2 = re.match(r'^(.*?)(\s*\.[\s\.]+\s*)$', rest)
                if m_dots2:
                    comp_part = m_dots2.group(1).strip()
                    dots = m_dots2.group(2)
                    # Handle "Component (refs) - " ending (multi-line continuation)
                    m_trailing_dash = re.match(r'^(.*?)\s*-\s*$', comp_part)
                    if m_trailing_dash:
                        # Name ends with dash - continuation on next line
                        inner = m_trailing_dash.group(1).strip()
                        translated_inner = translate_component_name(inner)
                        return leading + f"{fig_num}\t{translated_inner} —{dots}" + trailing
                    translated_comp = translate_component_name(comp_part)
                    return leading + f"{fig_num}\t{translated_comp}{dots}" + trailing
            else:
                # Multi-line entry, just translate the component name part
                translated_rest = translate_component_name(rest.strip())
                return leading + f"{fig_num}\t{translated_rest}" + trailing

        # Non-numeric tab entries
        if not re.match(r'^\d', core):
            return leading + translate_toc_entry(core) + trailing

        return text

    # 5. TOC entries (contain dot leaders, no tabs)
    if ". . ." in core:
        return leading + translate_toc_entry(core) + trailing

    # 6. "Fig.\tPage" header (no-tab variant)
    if core.startswith("Fig."):
        return leading + core.replace("Fig.", "Рис.").replace("Page", "Страница") + trailing

    # 7. Component name with Protective Treatment (fallback for non-dot cases)
    if "Protective Treatment" in core:
        result = translate_component_name(core)
        if result != core:
            return leading + result + trailing

    # 8. "Repair No." entries (outside TOC)
    if core.startswith("Repair No."):
        m = re.match(r'^(Repair No\.\s*\d+-\d+)\s+(.*?)$', core)
        if m:
            comp = m.group(2).strip()
            translated_comp = translate_component_name(comp)
            prefix = m.group(1).replace("Repair No.", "Ремонт №")
            return leading + f"{prefix} {translated_comp}" + trailing
        return leading + core.replace("Repair No.", "Ремонт №") + trailing

    # 9. "Page X" reference
    if re.match(r'^Page\s+\d+', core):
        return leading + core.replace("Page", "Стр.") + trailing

    # 10. Exact match in component names
    if core in COMPONENT_NAMES:
        return leading + COMPONENT_NAMES[core] + trailing

    # 11. "Blank" as page status
    if core == "Blank":
        return leading + "Пусто" + trailing

    return text


def translate_table_cell_text(text: str) -> str:
    """Translate text within a table cell."""
    stripped = text.strip()
    if not stripped:
        return text

    # Strip leading/trailing newlines for matching
    clean = stripped.strip('\n').strip()

    # Table headers
    if clean in TABLE_HEADERS:
        return text.replace(clean, TABLE_HEADERS[clean])

    # Section names in revision tables
    if clean in SECTION_NAMES:
        return text.replace(clean, SECTION_NAMES[clean])

    # "Repair No." section refs
    if clean.startswith("Repair No."):
        return text.replace("Repair No.", "Ремонт №")

    # SB titles
    for en_prefix in SB_TITLE_PARTS:
        if clean.startswith(en_prefix):
            return text.replace(clean, translate_sb_title(clean))

    # SB revision numbers
    if clean == "Initial Issue":
        return text.replace(clean, "Первоначальный выпуск")
    if clean == "No effect":
        return text.replace(clean, "Без изменений")

    # "Updated paras X.Y" type entries (check BEFORE generic reason handler)
    if clean.startswith("Updated paras"):
        return text.replace("Updated paras", "Обновлены пункты")
    if clean.startswith("Updated para"):
        return text.replace("Updated para", "Обновлён пункт")

    # Reason for change
    if any(phrase in clean for phrase in ["Updated", "Added", "Updated fig", "Updated Messier"]):
        return text.replace(clean, translate_reason(clean))

    # Standalone "figure NNN" or "figures NNN, NNN" (continuation lines of reason cells)
    m_fig = re.match(r'^(figures?)\s+(.*)$', clean)
    if m_fig:
        word = "рисунки" if m_fig.group(1) == "figures" else "рисунок"
        return text.replace(clean, f"{word} {m_fig.group(2)}")

    # Standalone continuation fragments from multi-paragraph reason cells
    if clean.startswith("para ") or clean.startswith("paras "):
        result = clean.replace("paras ", "пунктов ").replace("para ", "пункта ")
        return text.replace(clean, result)
    if clean.startswith("in para"):
        return text.replace("in para", "в пункте")
    if clean.startswith("in paras"):
        return text.replace("in paras", "в пунктах")

    # "Sliding tube" entries in figure tables
    if "Sliding tube" in clean or "sliding tube" in clean:
        result = clean
        result = result.replace("Sliding tube", "Скользящая труба")
        result = result.replace("sliding tube", "Скользящая труба")
        result = _translate_suffix(result)
        result = result.replace(" and ", " и ").replace(" or ", " или ")
        return text.replace(clean, result)

    # Repair/Oversize descriptions
    if (clean.startswith("Repair to ") or clean.startswith("Repair Bush") or
        clean.startswith("Repair Bearing") or clean.startswith("Oversize ") or
        clean.startswith("Lower Bearing")):
        # Fix missing separator: "InstallationRepair No." -> "Installation\tRepair No."
        fixed = re.sub(r'Installation(Repair No\.)', r'Installation\t\1', clean)
        result = translate_repair_description(fixed)
        return text.replace(clean, result)

    # "Component Repairs - Key Diagram" pattern
    if "Repairs - Key Diagram" in clean or "Key Diagram" in clean:
        result = clean
        result = result.replace("Approved Repairs - Key Diagram", "Утверждённые ремонты — Ключевая схема")
        result = result.replace("Repairs - Key Diagram", "Ремонты — Ключевая схема")
        # Translate component names
        for en, ru in sorted(COMPONENT_NAMES.items(), key=lambda x: -len(x[0])):
            result = result.replace(en, ru)
        return text.replace(clean, result)

    # Continuation fragments from multi-paragraph reason cells
    if clean.startswith("Limited to Safran"):
        result = clean
        result = re.sub(r'Safran L[а-яёА-ЯЁ\w]*ing Systems', 'Safran Landing Systems', result)
        result = result.replace("Limited to Safran Landing Systems", "Limited на Safran Landing Systems")
        return text.replace(clean, result)

    # Numbers with "and" (continuation lines): "814, 818 and 823"
    if re.match(r'^[\d,\s]+ and [\d,\s]+$', clean):
        return text.replace(" and ", " и ")

    # Generic fallback to main translator
    result = translate_text(text)
    return result


# ══════════════════════════════════════════════════════════════════════════
#  Run-level text replacement with formatting preservation
# ══════════════════════════════════════════════════════════════════════════

def replace_paragraph_text(para, new_text: str):
    """Replace all text in a paragraph while preserving formatting.

    Works at XML level to handle text inside hyperlinks and other structures.
    """
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    element = para._element

    # Collect ALL w:t elements (including those inside hyperlinks)
    all_t_elements = list(element.iter(qn('w:t')))

    if not all_t_elements:
        return

    # Strategy: put all text in the first non-empty w:t, clear others
    ref_idx = 0
    for i, t_elem in enumerate(all_t_elements):
        if t_elem.text and t_elem.text.strip():
            ref_idx = i
            break

    # Set new text on reference element, clear all others
    all_t_elements[ref_idx].text = new_text
    # Preserve spaces via xml:space="preserve"
    all_t_elements[ref_idx].set(qn('xml:space'), 'preserve')

    for i, t_elem in enumerate(all_t_elements):
        if i != ref_idx:
            t_elem.text = ""


def process_paragraph(para):
    """Translate a paragraph's text content."""
    original = para.text
    if not original or not original.strip():
        return

    translated = translate_text(original)
    if translated != original:
        replace_paragraph_text(para, translated)


def process_table_cell(cell):
    """Translate all paragraphs within a table cell."""
    for para in cell.paragraphs:
        original = para.text
        if not original or not original.strip():
            continue
        translated = translate_table_cell_text(original)
        if translated != original:
            replace_paragraph_text(para, translated)


# ══════════════════════════════════════════════════════════════════════════
#  Font size adjustment for longer Russian text
# ══════════════════════════════════════════════════════════════════════════

def adjust_font_if_needed(para, original_text: str, new_text: str):
    """Slightly reduce font size if Russian text is significantly longer."""
    if not original_text.strip() or not new_text.strip():
        return
    if new_text == original_text:
        return

    ratio = len(new_text) / max(len(original_text), 1)

    # Only adjust if text is >30% longer
    if ratio <= 1.3:
        return

    runs = para.runs
    if not runs:
        return

    for run in runs:
        if run.font.size:
            current_size = run.font.size
            # Reduce by 5-15% depending on expansion ratio
            if ratio > 1.5:
                new_size = int(current_size * 0.85)
            elif ratio > 1.3:
                new_size = int(current_size * 0.90)
            else:
                new_size = int(current_size * 0.95)
            run.font.size = new_size


# ══════════════════════════════════════════════════════════════════════════
#  Main translation pipeline
# ══════════════════════════════════════════════════════════════════════════

def main():
    print(f"Loading: {SRC}")
    doc = Document(str(SRC))

    # ── Process paragraphs ──
    translated_count = 0
    for i, para in enumerate(doc.paragraphs):
        original = para.text
        if not original or not original.strip():
            continue

        translated = translate_text(original)
        if translated != original:
            replace_paragraph_text(para, translated)
            adjust_font_if_needed(para, original, translated)
            translated_count += 1

    print(f"Translated {translated_count} paragraphs")

    # ── Process tables ──
    table_translated = 0
    for ti, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text
                    if not original or not original.strip():
                        continue
                    translated = translate_table_cell_text(original)
                    if translated != original:
                        replace_paragraph_text(para, translated)
                        adjust_font_if_needed(para, original, translated)
                        table_translated += 1

    print(f"Translated {table_translated} table cells")

    # ── Process headers and footers ──
    hf_count = 0
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                for para in header.paragraphs:
                    original = para.text
                    if original and original.strip():
                        translated = translate_text(original)
                        if translated != original:
                            replace_paragraph_text(para, translated)
                            hf_count += 1
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for para in footer.paragraphs:
                    original = para.text
                    if original and original.strip():
                        translated = translate_text(original)
                        if translated != original:
                            replace_paragraph_text(para, translated)
                            hf_count += 1

    print(f"Translated {hf_count} header/footer elements")

    # ── Process textboxes (if any exist in the document XML) ──
    tb_count = 0
    body = doc.element.body
    for txbx in body.iter(qn('w:txbxContent')):
        for p in txbx.iter(qn('w:p')):
            from docx.text.paragraph import Paragraph
            para = Paragraph(p, doc)
            original = para.text
            if original and original.strip():
                translated = translate_text(original)
                if translated != original:
                    replace_paragraph_text(para, translated)
                    tb_count += 1

    print(f"Translated {tb_count} textbox elements")

    # ── Save ──
    os.makedirs(str(DST.parent), exist_ok=True)
    doc.save(str(DST))
    print(f"Saved: {DST}")

    # ── Verification summary ──
    print("\n=== Verification ===")
    doc2 = Document(str(DST))
    untranslated = []
    for i, para in enumerate(doc2.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        # Check if English text remains (should be translated)
        if re.search(r'[A-Za-z]{4,}', text) and not is_only_numbers_or_codes(text):
            # Skip known exceptions (company names, codes, addresses)
            if any(skip in text for skip in [
                "Safran Landing Systems UK",
                "Cheltenham Road",
                "201587", "EDES2", "MAF1",
                "(c)AC", "A320-", "A321-",
                "32-12-22", "AMS-", "ASTM",
                "MLG", "NLG",
            ]):
                continue
            # Skip if contains Cyrillic (partially translated or mojibake)
            if re.search(r'[А-Яа-яЁё]{3,}', text):
                continue
            untranslated.append((i, text[:100]))

    if untranslated:
        print(f"Found {len(untranslated)} potentially untranslated paragraphs:")
        for idx, txt in untranslated[:20]:
            print(f"  P{idx}: {txt}")
    else:
        print("All paragraphs appear to be translated!")

    # Check table cells
    untranslated_cells = []
    for ti, table in enumerate(doc2.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    if re.search(r'[A-Za-z]{4,}', text) and not is_only_numbers_or_codes(text):
                        if any(skip in text for skip in [
                            "SAFRAN LANDING SYSTEMS",
                            "Safran Landing Systems",
                            "201587", "EDES2", "MAF1",
                            "(c)AC", "A320-", "A321-",
                            "32-12-22", "AMS-", "ASTM",
                            "MLG", "NLG", "Messier",
                            "SERVICE BULLETIN",
                        ]):
                            continue
                        if re.search(r'[А-Яа-яЁё]{3,}', text):
                            continue
                        untranslated_cells.append((f"T{ti+1}R{ri+1}C{ci+1}", text[:120]))

    if untranslated_cells:
        print(f"\nFound {len(untranslated_cells)} potentially untranslated table cells:")
        for loc, txt in untranslated_cells[:20]:
            print(f"  {loc}: {txt}")
    else:
        print("All table cells appear to be translated!")


if __name__ == "__main__":
    main()

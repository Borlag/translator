"""
CMM Section 2 Translation Engine
Translates aviation CMM documentation from English to Russian.
All translations by Claude claude-sonnet-4-6 — no external models used.

Architecture:
  1. EXACT — full-text exact match dictionary (highest priority)
  2. SENTENCE_TEMPLATES — regex-based full-sentence templates
  3. PHRASE_SUBS — ordered phrase/word substitutions applied to any remaining text
"""
import re
from docx import Document

# ──────────────────────────────────────────────────────────────
# HELPER: Russian grammatical helpers
# ──────────────────────────────────────────────────────────────

# Coating name map (English phrase → Russian noun phrase in appropriate form)
_COAT = {
    'cadmium plate': ('кадмиевое покрытие', 'кадмиевого покрытия', 'кадмиевое покрытие'),
    'cadmium plating': ('кадмирование', 'кадмирования', 'кадмирование'),
    'zinc nickel plate': ('цинко-никелевое покрытие', 'цинко-никелевого покрытия', 'цинко-никелевое покрытие'),
    'chromium plate': ('хромовое покрытие', 'хромового покрытия', 'хромовое покрытие'),
    'chrome plate': ('хромовое покрытие', 'хромового покрытия', 'хромовое покрытие'),
    'primer paint': ('грунтовочная краска', 'грунтовочной краски', 'грунтовочную краску'),
    'paint': ('краска', 'краски', 'краску'),
    'sealant': ('герметик', 'герметика', 'герметик'),
    'loctite': ('Локтайт', 'Локтайта', 'Локтайт'),
    'primer': ('грунтовка', 'грунтовки', 'грунтовку'),
    'sulphamate nickel plate': ('сульфаматное никелевое покрытие', 'сульфаматного никелевого покрытия', 'сульфаматное никелевое покрытие'),
    'sulphamate nickel plating': ('сульфаматное никелевое покрытие', 'сульфаматного никелевого покрытия', 'сульфаматное никелевое покрытие'),
}

_PART = {
    'sleeve': 'вставка', 'sleeves': 'вставки',
    'bush': 'втулка', 'bushes': 'втулки', 'bushing': 'втулка', 'bushings': 'втулки',
    'repair sleeve': 'ремонтная вставка', 'repair sleeves': 'ремонтные вставки',
    'repair bush': 'ремонтная втулка', 'repair bushes': 'ремонтные втулки',
    'pin': 'штифт', 'pins': 'штифты',
    'bolt': 'болт', 'bolts': 'болты',
    'nut': 'гайка', 'nuts': 'гайки',
    'washer': 'шайба', 'washers': 'шайбы',
    'spacer': 'проставка', 'spacers': 'проставки',
    'shim': 'прокладка', 'shims': 'прокладки',
    'bearing': 'подшипник', 'bearings': 'подшипники',
    'seal': 'уплотнение', 'seals': 'уплотнения',
    'ring': 'кольцо', 'rings': 'кольца',
    'label': 'ярлык', 'labels': 'ярлыки',
    'bracket': 'кронштейн', 'brackets': 'кронштейны',
    'clamp': 'хомут', 'clamps': 'хомуты',
    'screw': 'винт', 'screws': 'винты',
    'rivet': 'заклёпка', 'rivets': 'заклёпки',
    'bar': 'штанга', 'bars': 'штанги',
    'spring': 'пружина', 'springs': 'пружины',
    'valve': 'клапан', 'valves': 'клапаны',
    'tube': 'трубка', 'tubes': 'трубки',
    'plug': 'заглушка', 'plugs': 'заглушки',
    'liner': 'вкладыш', 'liners': 'вкладыши',
    'dowel': 'штифт', 'dowels': 'штифты',
    'target': 'мишень', 'targets': 'мишени',
    'harness': 'жгут',
    'connector': 'разъём', 'connectors': 'разъёмы',
}

_AREA = {
    'machined area': 'обработанная область',
    'machined areas': 'обработанные области',
    'reworked area': 'переработанная область',
    'reworked areas': 'переработанные области',
    'repaired area': 'отремонтированная область',
    'repaired areas': 'отремонтированные области',
    'repair area': 'ремонтная область',
    'repair areas': 'ремонтные области',
    'base metal': 'основной металл',
    'mating surface': 'сопрягаемая поверхность',
    'mating surfaces': 'сопрягаемые поверхности',
    'adjacent face': 'смежный торец',
    'adjacent faces': 'смежные торцы',
    'flange face': 'торец фланца',
    'flange faces': 'торцы фланца',
    'bores and faces': 'отверстия и торцы',
    'bore diameter': 'диаметр отверстия',
    'bore diameters': 'диаметры отверстий',
    'inside diameter': 'внутренний диаметр',
    'internal diameter': 'внутренний диаметр',
    'internal diameters': 'внутренние диаметры',
    'outer diameter': 'наружный диаметр',
    'external diameter': 'наружный диаметр',
    'spigot end': 'шейка',
    'outer faces': 'наружные торцы',
    'cross holes': 'поперечные отверстия',
    'cross hole': 'поперечное отверстие',
}

_COMP = {
    'main fitting': 'корпус стойки', 'main fitting subassembly': 'сборка корпуса стойки',
    'sliding tube': 'скользящая труба', 'sliding tube subassembly': 'сборка скользящей трубы',
    'upper torque link': 'верхний шлиц-шарнир', 'lower torque link': 'нижний шлиц-шарнир',
    'upper slave link': 'верхнее ведомое звено', 'lower slave link': 'нижнее ведомое звено',
    'upper stay': 'верхнее звено', 'lower stay': 'нижнее звено',
    'lock stay': 'фиксирующее звено', 'locking stay': 'фиксирующее звено',
    'lock stay cardan': 'кардан фиксирующего звена', 'cardan': 'кардан',
    'cylinder': 'цилиндр',
    'upper diaphragm tube': 'верхняя диафрагменная труба',
    'lower bearing': 'нижний подшипник',
    'lower bearing subassembly': 'сборка нижнего подшипника',
    'shock absorber': 'амортизатор', 'shock absorber subassembly': 'сборка амортизатора',
    'gland housing': 'корпус сальника',
    'retaining pin': 'фиксирующий штифт',
    'upper pivot bracket': 'верхний кронштейн оси',
    'pintle pin': 'штифт навеса', 'cardan pin': 'штифт кардана',
    'slave link': 'ведомое звено',
    'uplock pin': 'штифт замка убранного положения',
    'inflation valve': 'клапан накачки', 'inflation valve subassembly': 'сборка клапана накачки',
    'charging valve': 'зарядный клапан',
    'main landing gear leg': 'стойка основного шасси',
    'torque link': 'шлиц-шарнир',
    'stop ring': 'стопорное кольцо',
    'transfer dowel': 'передаточный штифт', 'transfer dowels': 'передаточные штифты',
    'valve stem': 'шток клапана',
    'proximity switch': 'датчик приближения', 'proximity sensor': 'датчик приближения',
    'switch target': 'мишень датчика',
    'bonding jumper': 'соединительная перемычка',
    'identification washer': 'идентификационная шайба',
    'lubrication fitting': 'смазочный ниппель',
    'cup washer': 'чашеобразная шайба',
    'downlocking spring': 'пружина фиксации',
    'forward pintle bush': 'передняя втулка штифта навеса',
    'housing': 'корпус',
    'backing ring': 'опорное кольцо',
    'o-ring seal': 'уплотнительное кольцо',
    'o-ring': 'уплотнительное кольцо',
}


_COAT_ACC = {
    'cadmium plate': 'кадмиевое покрытие',
    'cadmium plating': 'кадмирование',
    'zinc nickel plate': 'цинко-никелевое покрытие',
    'chromium plate': 'хромовое покрытие',
    'chrome plate': 'хромовое покрытие',
    'primer paint': 'грунтовочную краску',
    'paint': 'краску',
    'sealant': 'герметик',
    'sulphamate nickel plate': 'сульфаматное никелевое покрытие',
    'sulphamate nickel plating': 'сульфаматное никелевое покрытие',
    'primer': 'грунтовку',
    'protective treatment': 'защитное покрытие',
    'protective treatments': 'защитные покрытия',
}


def _coat_nom(s):
    s = s.strip().lower()
    for k, v in _COAT.items():
        if k in s:
            return v[0]
    return s


def _coat_gen(s):
    s = s.strip().lower()
    for k, v in _COAT.items():
        if k in s:
            return v[1]
    return s


def _coat_acc(s):
    """Accusative form for 'Remove the [coating]'."""
    s_low = s.strip().lower()
    for k, v in sorted(_COAT_ACC.items(), key=lambda x: -len(x[0])):
        if k in s_low:
            return v
    return _body(s)


def _comp(s):
    s = s.strip().lower()
    if s in _COMP:
        return _COMP[s]
    for k, v in sorted(_COMP.items(), key=lambda x: -len(x[0])):
        if k in s:
            return _COMP[k]
    return s


def _area(s):
    s = s.strip().lower()
    if s in _AREA:
        return _AREA[s]
    for k, v in sorted(_AREA.items(), key=lambda x: -len(x[0])):
        if k in s:
            return v
    return s


def _fig(s):
    """Translate 'Figure X' or 'Figures X and Y'."""
    s = re.sub(r'\bFigures?\s+(\d+)\s+and\s+(\d+)\b', r'Рисунки \1 и \2', s)
    s = re.sub(r'\bFigures?\s+(\d+)\s+to\s+(\d+)\b', r'Рисунки \1–\2', s)
    s = re.sub(r'\bFigure\s+(\d+)\b', r'Рисунок \1', s)
    return s


def _pcs(s):
    """Pass-through reference codes."""
    return s


def _dim(s):
    """Pass-through dimension strings."""
    return s


def _translate_ref_tail(tail):
    """Translate tail like ': refer to PCS-XXXX and Figure 601.'"""
    tail = re.sub(r'[Rr]efer to ', 'см. ', tail)
    tail = _fig(tail)
    return tail


def _translate_note(note_text):
    """Translate NOTE body text."""
    return translate_text(note_text)


# ──────────────────────────────────────────────────────────────
# 1.  EXACT DICTIONARY  (complete strings, case-sensitive)
# ──────────────────────────────────────────────────────────────
EXACT = {
    # ── Document structure ──
    "Specified Damage and Material Specification.": "Указанное повреждение и спецификация материала.",
    "Specified Damage": "Указанное повреждение",
    "Material Specification": "Спецификация материала",
    "Special Tools": "Специальные инструменты",
    "Special tools are not necessary.": "Специальные инструменты не требуются.",
    "These special tools are necessary:": "Необходимые специальные инструменты:",
    "These materials are necessary:": "Необходимые материалы:",
    "These repair parts are necessary:": "Необходимые запасные части для ремонта:",
    "Repair Parts": "Запасные части для ремонта",
    "NOTE: Alternative equivalents are permitted.": "ПРИМЕЧАНИЕ: Допускается применение эквивалентных аналогов.",
    "NOTE:\tAlternative equivalents are permitted.": "ПРИМЕЧАНИЕ: Допускается применение эквивалентных аналогов.",
    # ── Common figure text annotations ──
    "(WITHOUT SLEEVES)": "(БЕЗ ВСТАВОК)",
    "(WITH SLEEVES)": "(С ВСТАВКАМИ)",
    "(WITH SLEEVES FITTED (4 MAX))": "(С УСТАНОВЛЕННЫМИ ВСТАВКАМИ (МАКС. 4))",
    "SIDE ELEVATION WITHOUT SLEEVES": "ВИД СБОКУ БЕЗ ВСТАВОК",
    "SIDE ELEVATION WITH SLEEVES FITTED (4 MAX)": "ВИД СБОКУ С УСТАНОВЛЕННЫМИ ВСТАВКАМИ (МАКС. 4)",
    "(WITHOUT BUSHES)": "(БЕЗ ВТУЛОК)",
    "(WITH BUSHES)": "(С ВТУЛКАМИ)",
    "WITHOUT BUSHES": "БЕЗ ВТУЛОК",
    "WITH BUSHES INSTALLED": "С УСТАНОВЛЕННЫМИ ВТУЛКАМИ",
    "(BOTH BUSHES)": "(ОБЕ ВТУЛКИ)",
    "(WITHOUT BUSH)": "(БЕЗ ВТУЛКИ)",
    "(WITH BUSH)": "(С ВТУЛКОЙ)",
    "(WITH REPAIR BUSHES)": "(С РЕМОНТНЫМИ ВТУЛКАМИ)",
    "(WITHOUT REPAIR BUSHES)": "(БЕЗ РЕМОНТНЫХ ВТУЛОК)",
    "(WITH REPAIR BUSHES) REFER TO FIGURE 601": "(С РЕМОНТНЫМИ ВТУЛКАМИ) СМ. РИСУНОК 601",
    "(WITH BUSH) REFER TO FIGURE 601": "(С ВТУЛКОЙ) СМ. РИСУНОК 601",
    "(WITOUT REPAIR BUSH)": "(БЕЗ РЕМОНТНОЙ ВТУЛКИ)",
    "(WITHOUT REPAIR BUSH)": "(БЕЗ РЕМОНТНОЙ ВТУЛКИ)",
    "REPAIR BUSH": "РЕМОНТНАЯ ВТУЛКА",
    "REPAIR SLEEVE": "РЕМОНТНАЯ ВСТАВКА",
    "BOTH BUSHES\tSECTION": "ОБЕ ВТУЛКИ   РАЗРЕЗ",
    "BOTH BUSHES": "ОБЕ ВТУЛКИ",
    "COMMON ZONE": "ОБЩАЯ ЗОНА",
    "(FOR LH ONLY)": "(ТОЛЬКО ДЛЯ ЛЕВОГО БОРТА)",
    "(FOR RH ONLY)": "(ТОЛЬКО ДЛЯ ПРАВОГО БОРТА)",
    "MINIMUM WALL THICKNESS TYPICAL": "МИНИМАЛЬНАЯ ТОЛЩИНА СТЕНКИ (ТИПИЧНО)",
    "MINIMUM WALL THICKNESS": "МИНИМАЛЬНАЯ ТОЛЩИНА СТЕНКИ",
    "MINIMUM LUG WIDTH": "МИНИМАЛЬНАЯ ШИРИНА ПРОУШИНЫ",
    "MINIMUM LUG THICKNESS": "МИНИМАЛЬНАЯ ТОЛЩИНА ПРОУШИНЫ",
    "LUG THICKNESS": "ТОЛЩИНА ПРОУШИНЫ",
    "LUG WIDTH": "ШИРИНА ПРОУШИНЫ",
    "MIN WALL THICKNESS": "МИН. ТОЛЩИНА СТЕНКИ",
    "CHECK DIAMETER": "КОНТРОЛЬНЫЙ ДИАМЕТР",
    "CHECK DIA THROUGH": "КОНТРОЛЬНЫЙ ДИАМЕТР НАСКВОЗЬ",
    "NO CADMIUM PLATE": "БЕЗ КАДМИРОВАНИЯ",
    "DO NOT CADMIUM": "НЕ КАДМИРОВАТЬ",
    "PRIMER ONLY": "ТОЛЬКО ГРУНТОВКА",
    "NO PAINT": "БЕЗ КРАСКИ",
    "NO PAINT CHECK DIA THROUGH": "БЕЗ КРАСКИ, КОНТРОЛЬНЫЙ ДИАМЕТР НАСКВОЗЬ",
    "GREASE GROOVE": "СМАЗОЧНАЯ КАНАВКА",
    "GREASE GROOVES": "СМАЗОЧНЫЕ КАНАВКИ",
    "GREASE NIPPLE": "СМАЗОЧНЫЙ НИППЕЛЬ",
    "EXTENT OF SHOT PEEN": "ЗОНА ДРОБЕСТРУЙНОЙ ОБРАБОТКИ",
    "CHROMIUM PLATE TERMINATION": "ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ",
    "AFTER CHROMIUM PLATE": "ПОСЛЕ ХРОМИРОВАНИЯ",
    "THIS BORE TO": "ЭТО ОТВЕРСТИЕ",
    "BE CADMIUM PLATED": "ПОДЛЕЖИТ КАДМИРОВАНИЮ",
    "APPLY FILLET": "НАНЕСТИ ГАЛТЕЛЬ",
    "APPLY FILLET OF SEALANT": "НАНЕСТИ ГАЛТЕЛЬ ИЗ ГЕРМЕТИКА",
    "APPLY LOCTITE GRADE 601: REFER TO PCS-5303.": "НАНЕСТИ ЛОКТАЙТ МАРКИ 601: СМ. PCS-5303.",
    "DEPTH OF SLOT": "ГЛУБИНА ПАЗА",
    "SPOTFACE WITH": "ПОДРЕЗКА ТОРЦА С",
    "BLEND EDGES": "СГЛАДИТЬ КРОМКИ",
    "(0.02 to 0.03in) BLEND EDGES": "(от 0,02 до 0,03 дюйма) СГЛАДИТЬ КРОМКИ",
    "DETAIL A\tDETAIL B": "ДЕТАЛЬ A   ДЕТАЛЬ B",
    "MESSIER-DOWTY gloucester": "MESSIER-DOWTY Gloucester",
    "LABEL TO BE IN LINE WITH LUG": "МЕТКА ДОЛЖНА БЫТЬ В ЛИНИИ С ПРОУШИНОЙ",
    "APPLY CADMIUM OR ZINC-NICKEL PLATE. THE CADMIUM PLATE THICKNESS MUST BE BETWEEN 0,010 and 0,015mm (00004 and 00006in).":
        "НАНЕСТИ КАДМИЕВОЕ ИЛИ ЦИНКО-НИКЕЛЕВОЕ ПОКРЫТИЕ. ТОЛЩИНА КАДМИЕВОГО ПОКРЫТИЯ ДОЛЖНА БЫТЬ ОТ 0,010 ДО 0,015 мм (0,0004–0,0006 ДЮЙМА).",
    "PART SECTION X-X": "РАЗРЕЗ X-X",
    "PART SECTION Z-Z": "РАЗРЕЗ Z-Z",
    "SECTION Z-Z": "РАЗРЕЗ Z-Z",
    "SECTION": "РАЗРЕЗ",
    "SLIDING TUBE SUB ASSEMBLY": "СБОРКА СКОЛЬЗЯЩЕЙ ТРУБЫ",
    "FOR SHOCK ABSORBER": "ДЛЯ АМОРТИЗАТОРА",
    "LOWER BEARING": "НИЖНИЙ ПОДШИПНИК",
    "LOWER BEARING SUBASSEMBLY": "СБОРКА НИЖНЕГО ПОДШИПНИКА",
    "LOWER BEARING SUBASSEMBLY (16-110) OR (16-110A) OR (16-110B) OR (16-110C) OR (16A-110) OR (16A-110A)":
        "СБОРКА НИЖНЕГО ПОДШИПНИКА (16-110) ИЛИ (16-110A) ИЛИ (16-110B) ИЛИ (16-110C) ИЛИ (16A-110) ИЛИ (16A-110A)",
    "NOTE: GLAND HOUSING (16-140) OR (16A-140) OR (16A-140A) NOT SHOWN FOR CLARITY":
        "ПРИМЕЧАНИЕ: КОРПУС САЛЬНИКА (16-140) ИЛИ (16A-140) ИЛИ (16A-140A) НЕ ПОКАЗАН ДЛЯ ЯСНОСТИ",
    "GREASE GROOVE DIMENSIONS AFTER INSTALLATION": "РАЗМЕРЫ СМАЗОЧНОЙ КАНАВКИ ПОСЛЕ УСТАНОВКИ",
    "(0.02 to 0.03in) BLEND EDGES": "(от 0,02 до 0,03 дюйма) СГЛАДИТЬ КРОМКИ",
    "PROXIMITY SWITCH AND TARGET TESTS": "ИСПЫТАНИЯ ДАТЧИКА ПРИБЛИЖЕНИЯ И МИШЕНИ",
    "Proximity switch and target tests": "Испытания датчика приближения и мишени",
    # ── IPC/IPL entries ──
    ".PIN, RETAINING": ".ШТИФТ, ФИКСИРУЮЩИЙ",
    ".WASHER, CUP": ".ШАЙБА, ЧАШЕОБРАЗНАЯ",
    ".WASHER, LOCKING": ".ШАЙБА, СТОПОРНАЯ",
    ".WASHER, IDENTIFICATION": ".ШАЙБА, ИДЕНТИФИКАЦИОННАЯ",
    ".FITTING, LUBRICATION": ".НИППЕЛЬ, СМАЗОЧНЫЙ",
    ".ABSORBER SUBASSEMBLY, SHOCK": ".СБОРКА АМОРТИЗАТОРА, УДАРНОГО ТИПА",
    ".NUT, LOCKING": ".ГАЙКА, САМОКОНТРЯЩАЯСЯ",
    ".BUNG": ".ЗАГЛУШКА",
    ".FITTING SUBASSEMBLY, MAIN, LH": ".СБОРКА КОРПУСА СТОЙКИ, ЛЕВАЯ",
    ".FITTING SUBASSEMBLY, MAIN, RH": ".СБОРКА КОРПУСА СТОЙКИ, ПРАВАЯ",
    "..BUSH, RETRACTION ACTUATOR LUG": "..ВТУЛКА, УШК ПРИВОДА УБОРКИ",
    ".LABEL (CHARGING INSTRUCTION)": ".ЯРЛЫК (ИНСТРУКЦИЯ ПО ЗАРЯДКЕ)",
    ".BEARING, SPHERICAL SPARES FOR": ".ПОДШИПНИК, СФЕРИЧЕСКИЙ (ЗАПЧАСТИ)",
    # ── SB/Revision references ──
    "SUPSD BY ITEM 50C)": "ЗАМЕНЕНА ПОЗИЦИЕЙ 50C)",
    "SUPSD BY ITEM 50D)": "ЗАМЕНЕНА ПОЗИЦИЕЙ 50D)",
    "SUPSD BY ITEM 10C)": "ЗАМЕНЕНА ПОЗИЦИЕЙ 10C)",
    "SUPSD BY ITEM 30C)": "ЗАМЕНЕНА ПОЗИЦИЕЙ 30C)",
    "SUPSD BY ITEM 30D)": "ЗАМЕНЕНА ПОЗИЦИЕЙ 30D)",
    "SUPSD BY ITEM 30E)": "ЗАМЕНЕНА ПОЗИЦИЕЙ 30E)",
    "SUPSDS ITEM 50A)": "ЗАМЕНЯЕТ ПОЗИЦИЮ 50A)",
    "SUPSDS ITEM 30)": "ЗАМЕНЯЕТ ПОЗИЦИЮ 30)",
    "SUPSDS ITEM 30A)": "ЗАМЕНЯЕТ ПОЗИЦИЮ 30A)",
    "SUPSDS ITEM 30B)": "ЗАМЕНЯЕТ ПОЗИЦИЮ 30B)",
    "(ALTERNATIVE TO ITEM 50C)": "(АЛЬТЕРНАТИВА ПОЗИЦИИ 50C)",
    "(ALTERNATIVE TO ITEM 10A": "(АЛЬТЕРНАТИВА ПОЗИЦИИ 10A",
    "(PRE REF. CODE: 2542)": "(ДО КОД ССЫЛКИ: 2542)",
    "(POST REF. CODE: 2542)": "(ПОСЛЕ КОД ССЫЛКИ: 2542)",
    "USE WITH ITEMS 50E AND 50F": "ИСПОЛЬЗОВАТЬ С ПОЗИЦИЯМИ 50E И 50F",
    "(REFER TO FIG 14 TO 17 FOR": "(СМ. РИС. 14–17 ДЛЯ",
    "(REFER TO FIG 13 FOR NHA)": "(СМ. РИС. 13 ДЛЯ СУЗ)",
    "(USE WITH ITEMS 90, 90A, 100": "(ИСПОЛЬЗОВАТЬ С ПОЗИЦИЯМИ 90, 90A, 100",
    "(USE WITH ITEMS 90, 90A, 90B, 90C,": "(ИСПОЛЬЗОВАТЬ С ПОЗИЦИЯМИ 90, 90A, 90B, 90C,",
    "(REFER TO FIGURE 802)": "(СМ. РИСУНОК 802)",
    "(REFER TO FIGURE 806)": "(СМ. РИСУНОК 806)",
    "(REFER TO FIGURE 807)": "(СМ. РИСУНОК 807)",
    "(TURNED THRO 90o COUNTERCLOCKWISE)": "(ПОВЁРНУТО НА 90° ПРОТИВ ЧАСОВОЙ СТРЕЛКИ)",
    # ── Fits & Clearances ──
    "Table 801 (Continued)": "Таблица 801 (Продолжение)",
    "Table 802 (Continued)": "Таблица 802 (Продолжение)",
    "Table 803 (Continued)": "Таблица 803 (Продолжение)",
    "Table 808 (Continued)": "Таблица 808 (Продолжение)",
    "Table 812 (Continued)": "Таблица 812 (Продолжение)",
    "Table 813 (Continued)": "Таблица 813 (Продолжение)",
    "Table 814 (Continued)": "Таблица 814 (Продолжение)",
    "Table 815 (Continued)": "Таблица 815 (Продолжение)",
    "Fits and Clearances Table 803 (Continued)": "Посадки и зазоры, Таблица 803 (Продолжение)",
    "Fits and Clearances Table 808": "Посадки и зазоры, Таблица 808",
    "Fits and Clearances Table 812 (Continued)": "Посадки и зазоры, Таблица 812 (Продолжение)",
    "Fits and Clearances Table 813 (Continued)": "Посадки и зазоры, Таблица 813 (Продолжение)",
    "Fits and Clearances Table 814 (Continued)": "Посадки и зазоры, Таблица 814 (Продолжение)",
    "Fits and Clearances Table 815 (Continued)": "Посадки и зазоры, Таблица 815 (Продолжение)",
    # ── Record repair number ──
    "Record the repair number onto the documentation which is attached to the part. Optionally, identify the part with the Safran Landing Systems repair number 64-4505126-00 adjacent to the existing part number: refer to PCS-6000-07.":
        "Запишите номер ремонта в документацию, прикреплённую к детали. По желанию идентифицируйте деталь номером ремонта Safran Landing Systems 64-4505126-00 рядом с существующим номером детали: см. PCS-6000-07.",
    # ── Surface finish notes with tab ──
    "THE SURFACE FINISH MUST BE\tOR BETTER UNLESS GIVEN DIFFERENTLY. DEBURR THE SHARP EDGES WITH 45 DEGREES CHAMFER OR":
        "ЧИСТОТА ПОВЕРХНОСТИ ДОЛЖНА БЫТЬ НЕ ХУЖЕ УКАЗАННОГО ЗНАЧЕНИЯ, ЕСЛИ НЕ УКАЗАНО ИНАЧЕ. СНИМИТЕ ЗАУСЕНЦЫ НА ОСТРЫХ КРОМКАХ С ФАСКОЙ 45 ГРАДУСОВ ИЛИ",
    "THE SURFACE FINISH MUST BE\tOR BETTER UNLESS GIVEN DIFFERENTLY. APPLY LOCTITE GRADE 601: REFER TO PCS-5303.":
        "ЧИСТОТА ПОВЕРХНОСТИ ДОЛЖНА БЫТЬ НЕ ХУЖЕ УКАЗАННОГО ЗНАЧЕНИЯ, ЕСЛИ НЕ УКАЗАНО ИНАЧЕ. НАНЕСИТЕ ЛОКТАЙТ МАРКИ 601: СМ. PCS-5303.",
    # ── Cautions ──
    "CAUTION:\tFOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO MESSIER-DOWTY LIMITED: REFER TO GUIDE-CS-001.":
        "ВНИМАНИЕ: ПРИ ПОВРЕЖДЕНИЯХ, ПРЕВЫШАЮЩИХ ПРЕДЕЛЫ ДАННОЙ СХЕМЫ РЕМОНТА, ОБРАТИТЕСЬ В MESSIER-DOWTY LIMITED: СМ. GUIDE-CS-001.",
    "CAUTION:\tFOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO MESSIER-DOWTY LTD: REFER TO GUIDE-CS-001.":
        "ВНИМАНИЕ: ПРИ ПОВРЕЖДЕНИЯХ, ПРЕВЫШАЮЩИХ ПРЕДЕЛЫ ДАННОЙ СХЕМЫ РЕМОНТА, ОБРАТИТЕСЬ В MESSIER-DOWTY LTD: СМ. GUIDE-CS-001.",
    "CAUTION:\tFOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO SAFRAN LANDING SYSTEMS: REFER TO GUIDE-CS-001.":
        "ВНИМАНИЕ: ПРИ ПОВРЕЖДЕНИЯХ, ПРЕВЫШАЮЩИХ ПРЕДЕЛЫ ДАННОЙ СХЕМЫ РЕМОНТА, ОБРАТИТЕСЬ В SAFRAN LANDING SYSTEMS: СМ. GUIDE-CS-001.",
    # ── Common notes/cautions ──
    "CAUTION:\tDO NOT BLOCK THE HOLE IN THE BEARING (20-300): IT IS A DRAIN.":
        "ВНИМАНИЕ: НЕ ЗАКРЫВАЙТЕ ОТВЕРСТИЕ В ПОДШИПНИКЕ (20-300): ЭТО ДРЕНАЖНОЕ ОТВЕРСТИЕ.",
    # ── Sulphamate nickel thickness ──
    "THE THICKNESS OF THE SULPHAMATE NICKEL PLATING AFTER MACHINING MUST BE BETWEEN 0,050 and 0,760mm (0.0020 and 0.0299in).":
        "ТОЛЩИНА СУЛЬФАМАТНОГО НИКЕЛЕВОГО ПОКРЫТИЯ ПОСЛЕ МЕХАНИЧЕСКОЙ ОБРАБОТКИ ДОЛЖНА БЫТЬ ОТ 0,050 ДО 0,760 мм (0,0020–0,0299 ДЮЙМА).",
    # ── Tool descriptions ──
    "Extraction Tube": "Трубка для извлечения",
    "Lifting Bar Assembly": "Сборка подъёмной штанги",
    "Transport and Build Trolley": "Транспортировочная и монтажная тележка",
    "Support Arms": "Опорные рычаги",
    "Towing Frame": "Буксировочная рама",
    "Remove the forward pintle bush (20-250A)": "Снимите переднюю втулку штифта навеса (20-250A)",
    # ── Instruction sentence fragments ──
    "Procedure (Refer to Figures 601 and 602).": "Процедура (см. Рисунки 601 и 602).",
    "Procedure (Refer to Figures 601 to 603)": "Процедура (см. Рисунки 601–603)",
    "Machine diameter(s) A to remove the damage or wear or corrosion to": "Расточите диаметр(ы) A для устранения повреждения, износа или коррозии до",
    "Measure and record the diameter(s) A.": "Измерьте и запишите диаметр(ы) A.",
    "Measure and record the new diameter A.": "Измерьте и запишите новый диаметр A.",
    "Machine diameter Z, use the formula:": "Расточите диаметр Z по формуле:",
    "Machine diameter A, use the formula:": "Расточите диаметр A по формуле:",
    "Examine the machined areas for flaws: refer to PCS-3100, inclusion Class 4 and PCS-3600.":
        "Осмотрите обработанные области на наличие дефектов: см. PCS-3100, класс включений 4 и PCS-3600.",
    "Shot peen the reworked areas only: refer to PCS-2300.": "Выполните дробеструйную обработку только переработанных областей: см. PCS-2300.",
    "Apply applicable paint to the repaired areas: refer to PCS-2500.": "Нанесите соответствующую краску на отремонтированные области: см. PCS-2500.",
    "Examine the part to make sure that you have obeyed all the repair instructions correctly.":
        "Осмотрите деталь, чтобы убедиться в точном соблюдении всех инструкций по ремонту.",
    "Machine the chamfers as shown: refer to Figure 602.": "Снимите фаски согласно чертежу: см. Рисунок 602.",
    "Machine the chamfers as shown: refer to Figure 601.": "Снимите фаски согласно чертежу: см. Рисунок 601.",
    "Machine the radii and chamfers as shown: refer to Figure 601.": "Обработайте радиусы и фаски согласно чертежу: см. Рисунок 601.",
    "Machine the bore of sleeve(s) to the dimensions shown: refer to Figure 601.":
        "Расточите отверстие вставки(ок) до размеров, указанных на рисунке: см. Рисунок 601.",
    "Machine the chamfer as shown: refer to M-DLPS900 and Figure 601.":
        "Снимите фаску согласно чертежу: см. M-DLPS900 и Рисунок 601.",
    "NOTE: Make sure that the sleeve chamfer is protruded out of the lug face during the installation and do not remain in the lug width.":
        "ПРИМЕЧАНИЕ: Убедитесь, что фаска вставки выступает над торцом проушины при установке и не остаётся в ширине проушины.",
    "NOTE: Make sure that the sleeve chamfer is machined and do not remain in the lug width.":
        "ПРИМЕЧАНИЕ: Убедитесь, что фаска вставки обработана и не остаётся в ширине проушины.",
    "Apply paint locally to the torque link, but not to the repair bushes: refer to PCS-2500.":
        "Нанесите краску локально на шлиц-шарнир, но не на ремонтные втулки: см. PCS-2500.",
    "Apply paint locally to upper diaphragm tube: refer to PCS-2500.":
        "Нанесите краску локально на верхнюю диафрагменную трубу: см. PCS-2500.",
    "Do this procedure, if there is wear or damage to one or more bore diameter A:":
        "Выполните данную процедуру при наличии износа или повреждения одного или нескольких диаметров отверстия A:",
    "NOTE: The cup washers (13-20) are safetied after testing: refer to para AG.":
        "ПРИМЕЧАНИЕ: Чашеобразные шайбы (13-20) законтриваются после испытания: см. пункт AG.",
    "NOTE: Make sure the internal diameter of the bushes (16-130 or 16A-130A) are between the dimensions as shown: refer to Figure 712.":
        "ПРИМЕЧАНИЕ: Убедитесь, что внутренний диаметр втулок (16-130 или 16A-130A) соответствует размерам, указанным на рисунке: см. Рисунок 712.",
    "Install the split pin (6-10) and safety it: refer to M-DLPS1011-1.":
        "Установите шплинт (6-10) и законтрите его: см. M-DLPS1011-1.",
    "Install the split pin (6-60) and safety it: refer to M-DLPS1011-1.":
        "Установите шплинт (6-60) и законтрите его: см. M-DLPS1011-1.",
    "CAUTION:\tFOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO MESSIER-DOWTY LIMITED: REFER TO GUIDANCE DOCUMENTATION.":
        "ВНИМАНИЕ: ПРИ ПОВРЕЖДЕНИЯХ, ПРЕВЫШАЮЩИХ ДОПУСТИМЫЕ ПРЕДЕЛЫ ДАННОЙ СХЕМЫ РЕМОНТА, НАПРАВЬТЕ ЗАПРОС В MESSIER-DOWTY LIMITED: СМ. РУКОВОДЯЩУЮ ДОКУМЕНТАЦИЮ.",
    # ── IPC/IPL ──
    "-ITEM NOT ILLUSTRATED": "-ДЕТАЛЬ НЕ ПОКАЗАНА",
    # ── Common procedure text ──
    "Materials are not necessary.": "Материалы не требуются.",
    "Repair parts are not necessary.": "Запасные части для ремонта не требуются.",
    "INTENTIONALLY BLANK": "НАМЕРЕННО ОСТАВЛЕНО ПУСТЫМ",
    "PART No. 201587001 AND 201587002 COMPONENT MAINTENANCE MANUAL MAIN LANDING GEAR LEG":
        "РУКОВОДСТВО ПО ТО КОМПОНЕНТОВ СТОЙКИ ОСНОВНОГО ШАССИ № 201587001 И 201587002",
    # ── CAUTION variants ──
    "CAUTION: FOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO MESSIER-DOWTY LIMITED: REFER TO GUIDE-CS-001.":
        "ВНИМАНИЕ: ПРИ ПОВРЕЖДЕНИЯХ, ПРЕВЫШАЮЩИХ ДОПУСТИМЫЕ ПРЕДЕЛЫ ДАННОЙ СХЕМЫ РЕМОНТА, ОБРАТИТЕСЬ В MESSIER-DOWTY LIMITED: СМ. GUIDE-CS-001.",
    "CAUTION: FOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO MESSIER-DOWTY LTD: REFER TO GUIDE-CS-001.":
        "ВНИМАНИЕ: ПРИ ПОВРЕЖДЕНИЯХ, ПРЕВЫШАЮЩИХ ДОПУСТИМЫЕ ПРЕДЕЛЫ ДАННОЙ СХЕМЫ РЕМОНТА, ОБРАТИТЕСЬ В MESSIER-DOWTY LTD: СМ. GUIDE-CS-001.",
    "CAUTION: FOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO SAFRAN LANDING SYSTEMS: REFER TO GUIDE-CS-001.":
        "ВНИМАНИЕ: ПРИ ПОВРЕЖДЕНИЯХ, ПРЕВЫШАЮЩИХ ДОПУСТИМЫЕ ПРЕДЕЛЫ ДАННОЙ СХЕМЫ РЕМОНТА, ОБРАТИТЕСЬ В SAFRAN LANDING SYSTEMS: СМ. GUIDE-CS-001.",
    # ── TYPICAL N PLACES ──
    "(TYPICAL 2 PLACES)": "(ТИПИЧНО В 2 МЕСТАХ)",
    "(TYPICAL 4 PLACES)": "(ТИПИЧНО В 4 МЕСТАХ)",
    "(TYPICAL 6 PLACES)": "(ТИПИЧНО В 6 МЕСТАХ)",
    "TYPICAL 2 PLACES": "ТИПИЧНО В 2 МЕСТАХ",
    "TYPICAL 4 PLACES": "ТИПИЧНО В 4 МЕСТАХ",
    "MAX 4 PLACES": "МАКС. В 4 МЕСТАХ",
    "2 PLACES WITH SLEEVES FITTED": "2 МЕСТА С УСТАНОВЛЕННЫМИ ВСТАВКАМИ",
    "X 45 DEGREES CHAMFER TYPICAL 2 PLACES": "× 45 ГРАДУСОВ ФАСКА ТИПИЧНО В 2 МЕСТАХ",
    "THICKNESS TYPICAL 2 PLACES": "ТОЛЩИНА ТИПИЧНО В 2 МЕСТАХ",
    "OVERSIZE BUSH (REFER TO TABLE 601) 4 OFF": "РЕМОНТНАЯ ВТУЛКА (СМ. ТАБЛИЦУ 601) 4 ШТ.",
    "DIA. H  15 DEGREES": "ДИАМ. H  15 ГРАДУСОВ",
    # ── REPAIR BUSH labels ──
    "REPAIR BUSH 450258145": "РЕМОНТНАЯ ВТУЛКА 450258145",
    "REPAIR BUSH 450258806": "РЕМОНТНАЯ ВТУЛКА 450258806",
    "REPAIR BUSH 450217818": "РЕМОНТНАЯ ВТУЛКА 450217818",
    "REPAIR BUSH 450217819": "РЕМОНТНАЯ ВТУЛКА 450217819",
    "REPAIR BUSH 450217864": "РЕМОНТНАЯ ВТУЛКА 450217864",
    "REPAIR BUSH 450217865": "РЕМОНТНАЯ ВТУЛКА 450217865",
    # ── Notes ──
    "NOTE: Install the bush by heating the housing and cooling the bush only.":
        "ПРИМЕЧАНИЕ: Установите втулку, нагрев корпус и охладив только втулку.",
    "NOTE: Install the bushes by heating the housing and cooling the bushes only.":
        "ПРИМЕЧАНИЕ: Установите втулки, нагрев корпус и охладив только втулки.",
    "APPLY JOINTING COMPOUND TO THE BORE ONLY TO THIS AREA":
        "НАНЕСТИ МОНТАЖНЫЙ СОСТАВ НА ОТВЕРСТИЕ ТОЛЬКО В ДАННОЙ ЗОНЕ",
    "APPLY FILLET OF SEALANT (TYPICAL 2 PLACES)": "НАНЕСТИ ГАЛТЕЛЬ ИЗ ГЕРМЕТИКА (ТИПИЧНО В 2 МЕСТАХ)",
    "AFTER GRINDING THE CHROMIUM PLATE": "ПОСЛЕ ШЛИФОВКИ ХРОМОВОГО ПОКРЫТИЯ",
    "(63 micro-inches).": "(63 микродюйма).",
    "(63 micro-inches) or better: refer to PCS-4100 and Figure 601.":
        "(63 микродюйма) или лучше: см. PCS-4100 и Рисунок 601.",
    "(0.4587in) MIN WALL": "(0,4587 дюйма) МИН. ТОЛЩИНА СТЕНКИ",
    "Do the above step three more times.": "Повторите вышеуказанный шаг ещё три раза.",
    "Do this step three more times.": "Повторите данный шаг ещё три раза.",
    "Do the above step once more.": "Повторите вышеуказанный шаг ещё один раз.",
    "Select the applicable repair bushes from Table 601 for diameters A (qty 4).":
        "Выберите применимые ремонтные втулки из Таблицы 601 для диаметров A (кол-во 4).",
    "Apply protective treatment to the retaining pin: refer to REPAIR.":
        "Нанесите защитное покрытие на фиксирующий штифт: см. REPAIR.",
    "Restore the protective treatments to the cylinder: refer to REPAIR.":
        "Восстановите защитные покрытия на цилиндре: см. REPAIR.",
    "Apply paint locally to the torque link, but not to the repair bushes: refer to REPAIR.":
        "Нанесите краску локально на шлиц-шарнир, но не на ремонтные втулки: см. REPAIR.",
    "If there is evidence of delamination, remove the sulphamate nickel plate and do the repair again.":
        "При наличии признаков расслоения снимите сульфаматное никелевое покрытие и повторите ремонт.",
    "If there is sign of delamination, remove the sulphamate nickel plate and do the repair procedure again.":
        "При наличии признаков расслоения снимите сульфаматное никелевое покрытие и повторите процедуру ремонта.",
    "Use of 5x or 10x magnification. Examine the edges of sulphamate nickel plate and make sure that they are correctly bonded.":
        "При увеличении 5× или 10×. Осмотрите кромки сульфаматного никелевого покрытия и убедитесь в их надлежащем сцеплении.",
    "Grit blast the area to be sulphamate nickel plated: refer to PCS-2610. Make sure that the upper diaphragm tube is correctly masked.":
        "Выполните пескоструйную обработку зоны для нанесения сульфаматного никелевого покрытия: см. PCS-2610. Убедитесь, что верхняя диафрагменная труба правильно замаскирована.",
    "Calculate the diameter G and dimension H for the repair bushes, use the formulas:":
        "Рассчитайте диаметр G и размер H для ремонтных втулок по формулам:",
    "Calculate the dimensions of the repair bush 450217818: refer to Figure 601 and Figure 603. Use the formula:":
        "Рассчитайте размеры ремонтной втулки 450217818: см. Рисунок 601 и Рисунок 603. Используйте формулу:",
    "Calculate the dimensions of the repair bush 450217819: refer to Figure 601 and Figure 603. Use the formula:":
        "Рассчитайте размеры ремонтной втулки 450217819: см. Рисунок 601 и Рисунок 603. Используйте формулу:",
    "Procedure (Refer to Figures 601, 602 and 603)": "Процедура (см. Рисунки 601, 602 и 603)",
    "Machine the repair bushes to the diameter calculated and as shown: refer to Figure 602. Make the surface finish 1,6 micrometers (63 micro-inches).":
        "Расточите ремонтные втулки до рассчитанного диаметра согласно чертежу: см. Рисунок 602. Чистота поверхности должна составлять 1,6 мкм (63 микродюйма).",
    "Machine the repair bushes to the dimensions shown and calculated. Machine the inside face of the bush flange to make dimension H: refer to Figure 602. Make the surface finish 1,6 micrometers (63 micro-inches).":
        "Расточите ремонтные втулки до указанных и рассчитанных размеров. Расточьте внутреннюю поверхность фланца втулки для обеспечения размера H: см. Рисунок 602. Чистота поверхности должна составлять 1,6 мкм (63 микродюйма).",
    "Machine the repair bush to the dimensions shown and calculated: refer to Figure 602. Make the surface finish 1,6 micrometers (63 micro-inches).":
        "Расточите ремонтную втулку до указанных и рассчитанных размеров: см. Рисунок 602. Чистота поверхности должна составлять 1,6 мкм (63 микродюйма).",
    "Machine the flange faces of the repair bushes to get the correct dimensions after installation:  refer  to  Figure 602.  Make  the  surface  finish  1,6 micrometers (63 micro-inches).":
        "Расточьте торцы фланцев ремонтных втулок для получения правильных размеров после установки: см. Рисунок 602. Чистота поверхности должна составлять 1,6 мкм (63 микродюйма).",
    "Remove the paint locally from the torque link: refer to PCS-2700.":
        "Снимите краску локально со шлиц-шарнира: см. PCS-2700.",
    "Remove the cadmium plate from the torque link. Refer to PCS-2101.":
        "Снимите кадмиевое покрытие со шлиц-шарнира. См. PCS-2101.",
    "Remove the paint from the upper torque link: refer to PCS-2700.":
        "Снимите краску с верхнего шлиц-шарнира: см. PCS-2700.",
    "Remove the cadmium plate from the upper torque link. Refer to PCS-2101.":
        "Снимите кадмиевое покрытие с верхнего шлиц-шарнира. См. PCS-2101.",
    "Remove the paint locally from the lower torque link: refer to PCS-2700.":
        "Снимите краску локально с нижнего шлиц-шарнира: см. PCS-2700.",
    "Remove the cadmium plate locally from the lower torque link: refer to PCS-2101.":
        "Снимите кадмиевое покрытие локально с нижнего шлиц-шарнира: см. PCS-2101.",
    "Remove the paint locally from the upper slave link: refer to PCS-2700.":
        "Снимите краску локально с верхнего ведомого звена: см. PCS-2700.",
    "Remove the paint locally from the upper diaphragm tube: refer to PCS-2700.":
        "Снимите краску локально с верхней диафрагменной трубы: см. PCS-2700.",
    "Remove the cadmium plate locally from the upper diaphragm tube: refer to PCS-2101.":
        "Снимите кадмиевое покрытие локально с верхней диафрагменной трубы: см. PCS-2101.",
    "SECTION Z-Z (WITH BUSHES) REFER TO FIGURE 601": "РАЗРЕЗ Z-Z (С ВТУЛКАМИ) СМ. РИСУНОК 601",
    "Use Press Pad 460004330/146 and Drift 460004331/8 to install the repair bush 450217818. Use press pad 460004330/97 to install the repair bush 450217819: refer to M-DLPS1011-14 and Figure 603.":
        "Используйте прессовую подушку 460004330/146 и оправку 460004331/8 для установки ремонтной втулки 450217818. Используйте прессовую подушку 460004330/97 для установки ремонтной втулки 450217819: см. M-DLPS1011-14 и Рисунок 603.",
    "Use the Press Pad 460004330/255 and install the repair bushes: refer to M-DLPS1011-20. Use zinc loaded mastinox (made from Mastinox D40, Material Ref. Item TBA and Zinc Powder, Material Ref. Item TBA)":
        "Используйте прессовую подушку 460004330/255 и установите ремонтные втулки: см. M-DLPS1011-20. Используйте Mastinox с добавкой цинка (из Mastinox D40, Материал Поз. TBA и цинкового порошка, Материал Поз. TBA)",
    "Use the Press Pad 460004330/255 and install the repair bushes: refer to M-DLPS1011-20. Use electrically conducting Mastinox (made from Mastinox D40, Material Ref. Item 05-533 and Zinc Powder, Material":
        "Используйте прессовую подушку 460004330/255 и установите ремонтные втулки: см. M-DLPS1011-20. Используйте электропроводящий Mastinox (из Mastinox D40, Материал Поз. 05-533 и цинкового порошка, Материал",
    "Use the press pad 460004330/127 and the drift 460004331/21 and install the repair bushes: refer to M-DLPS1011-20. Use electrically conducting Mastinox (made from Mastinox D40, Material Ref. Item 05-53":
        "Используйте прессовую подушку 460004330/127 и оправку 460004331/21 и установите ремонтные втулки: см. M-DLPS1011-20. Используйте электропроводящий Mastinox (из Mastinox D40, Материал Поз. 05-53",
    "Use the press pad 460004330/127 and drift 460004331/21 and install the repair bushes: refer to M-DLPS1011-20. Use electrically conducting mastinox (made from Mastinox D40, Material Ref. Item 05-533 an":
        "Используйте прессовую подушку 460004330/127 и оправку 460004331/21 и установите ремонтные втулки: см. M-DLPS1011-20. Используйте электропроводящий Mastinox (из Mastinox D40, Материал Поз. 05-533 и",
    # ── Full text variants (continue from truncated entries above) ──
    "Use the Press Pad 460004330/255 and install the repair bushes: refer to M-DLPS1011-20. Use electrically conducting Mastinox (made from Mastinox D40, Material Ref. Item 05-533 and Zinc Powder, Material Ref. Item TBA).":
        "Используйте прессовую подушку 460004330/255 и установите ремонтные втулки: см. M-DLPS1011-20. Используйте электропроводящий Mastinox (из Mastinox D40, Материал Поз. 05-533 и цинкового порошка, Материал Поз. TBA).",
    "Use the press pad 460004330/127 and the drift 460004331/21 and install the repair bushes: refer to M-DLPS1011-20. Use electrically conducting Mastinox (made from Mastinox D40, Material Ref. Item 05-533 and Zinc Powder, Material Ref. Item TBA).":
        "Используйте прессовую подушку 460004330/127 и оправку 460004331/21 и установите ремонтные втулки: см. M-DLPS1011-20. Используйте электропроводящий Mastinox (из Mastinox D40, Материал Поз. 05-533 и цинкового порошка, Материал Поз. TBA).",
    "Use the press pad 460004330/127 and drift 460004331/21 and install the repair bushes: refer to M-DLPS1011-20. Use electrically conducting mastinox (made from Mastinox D40, Material Ref. Item 05-533 and Zinc powder, Material Ref. Item TBA).":
        "Используйте прессовую подушку 460004330/127 и оправку 460004331/21 и установите ремонтные втулки: см. M-DLPS1011-20. Используйте электропроводящий Mastinox (из Mastinox D40, Материал Поз. 05-533 и цинкового порошка, Материал Поз. TBA).",
    # ── More unique entries ──
    "Grit blast only the area that is to be sulphamate nickel plated: refer to PCS-2610. Make sure that the upper diaphragm tube is correctly masked.":
        "Выполните пескоструйную обработку только зоны, на которую будет нанесено сульфаматное никелевое покрытие: см. PCS-2610. Убедитесь, что верхняя диафрагменная труба правильно замаскирована.",
    "CAUTION: APPLY TEMPORARY CORROSION AND DAMAGE PROTECTION: REFER TO PCS-2800.":
        "ВНИМАНИЕ: НАНЕСИТЕ ВРЕМЕННУЮ ЗАЩИТУ ОТ КОРРОЗИИ И ПОВРЕЖДЕНИЙ: СМ. PCS-2800.",
    "Glass bead peen the machined areas: refer to M-DLPS134.":
        "Выполните дробеструйную обработку стеклянными шариками обработанных областей: см. M-DLPS134.",
    "Calculate the dimensions for the repair bush, use the formula:":
        "Рассчитайте размеры ремонтной втулки по формуле:",
    "Install the repair threaded insert to the transfer block: refer to Figure 602.":
        "Установите ремонтную резьбовую вставку в блок передачи: см. Рисунок 602.",
    "(WITHOUT THREADED INSERT)": "(БЕЗ РЕЗЬБОВОЙ ВСТАВКИ)",
    "(WITH THREADED INSERT)": "(С РЕЗЬБОВОЙ ВСТАВКОЙ)",
    "REPAIR THREADED INSERT": "РЕМОНТНАЯ РЕЗЬБОВАЯ ВСТАВКА",
    "Refer to M-DLPS615. Apply protective varnish to the wiring diagram plate (1-110).":
        "См. M-DLPS615. Нанесите защитный лак на пластину со схемой электрических соединений (1-110).",
    "Do the electrical bonding tests.": "Выполните испытания электрического соединения.",
    "If the bonding resistance of any of the bushes is outside the limits, then remove the defective bush and install the bush again.":
        "Если сопротивление соединения любой из втулок выходит за пределы допустимых значений, снимите дефектную втулку и установите её снова.",
    "SLIDING TUBE": "СКОЛЬЗЯЩАЯ ТРУБА",
    "A\tSLIDING TUBE": "A   СКОЛЬЗЯЩАЯ ТРУБА",
    "BUSH OUTER DIA. AND FLANGE UNDERSIDE ONLY": "ТОЛЬКО НАРУЖНЫЙ ДИАМЕТР ВТУЛКИ И НИЖНЯЯ СТОРОНА ФЛАНЦА",
    "TO THIS AREA": "НА ЭТУ ЗОНУ",
    "Hold the cylinder (17-230A) with the Bench Clamp MT1025 and the Holding Blocks 460006406.":
        "Зафиксируйте цилиндр (17-230A) с помощью настольных тисков MT1025 и удерживающих блоков 460006406.",
    "Push the piston (17-200) into the cylinder (17-230A) to retain the bearing\t(17-180) inside the cylinder (17-230A).":
        "Вставьте поршень (17-200) в цилиндр (17-230A), чтобы зафиксировать подшипник (17-180) внутри цилиндра (17-230A).",
    "CAUTION: MAKE SURE THAT YOU DO THE SPECIAL DIMENSION CHECK OF THE ROD (17-160): REFER TO CHECK. IF NOT, IT CAN CAUSE DAMAGE TO THE COMPONENT.":
        "ВНИМАНИЕ: УБЕДИТЕСЬ В ВЫПОЛНЕНИИ СПЕЦИАЛЬНОЙ ПРОВЕРКИ РАЗМЕРОВ ШТОКА (17-160): СМ. CHECK. В ПРОТИВНОМ СЛУЧАЕ КОМПОНЕНТ МОЖЕТ БЫТЬ ПОВРЕЖДЁН.",
    "Assemble the washer (17-170) and the rod (17-160) into the cylinder (17-230A) and carefully install the nut assembly (17-150) over the rod (17-160) and into the cylinder (17-230A).":
        "Установите шайбу (17-170) и шток (17-160) в цилиндр (17-230A), после чего аккуратно установите гаечную сборку (17-150) на шток (17-160) и в цилиндр (17-230A).",
    "Keep the tab washer (17-90A) in position and insert the jacking dome (17-80A).":
        "Удерживайте стопорную шайбу (17-90A) в положении и вставьте домкратный купол (17-80A).",
    "DRILLING AREA 1 SCREW FOR EACH AREA": "ЗОНА СВЕРЛЕНИЯ 1 ВИНТ НА КАЖДУЮ ЗОНУ",
    "RIGHT CONFIGURATION": "ПРАВАЯ КОНФИГУРАЦИЯ",
    "ILLUSTRATED PARTS LIST": "ИЛЛЮСТРИРОВАННЫЙ ПЕРЕЧЕНЬ ДЕТАЛЕЙ",
    "Remove the cadmimum plate from the upper diaphragm tube: refer to PCS-2101.":
        "Снимите кадмиевое покрытие с верхней диафрагменной трубы: см. PCS-2101.",
    "Remove the cadmimum plate from the upper doaphgram tube: refer to PCS-2101.":
        "Снимите кадмиевое покрытие с верхней диафрагменной трубы: см. PCS-2101.",
    "(WITHOUT REPAIR BEARING)": "(БЕЗ РЕМОНТНОГО ПОДШИПНИКА)",
    "(WITH REPAIR BEARING)": "(С РЕМОНТНЫМ ПОДШИПНИКОМ)",
    "APPLY LIGHT COAT OF PRIMER": "НАНЕСТИ ТОНКИЙ СЛОЙ ГРУНТОВКИ",
    "If necessary, remove the paint from the bracket: refer to PCS-2700.":
        "При необходимости снимите краску с кронштейна: см. PCS-2700.",
    "Calculate the diameters for the repair bushes, use the formulae:":
        "Рассчитайте диаметры для ремонтных втулок по формулам:",
    "Use press pad 460004330/146 and drift 460004331/8 to install the repair bush 450266362. Use press pad 460004330/148 to install the repair bush 450266361: refer to M-DLPS1011-24.":
        "Используйте прессовую подушку 460004330/146 и оправку 460004331/8 для установки ремонтной втулки 450266362. Используйте прессовую подушку 460004330/148 для установки ремонтной втулки 450266361: см. M-DLPS1011-24.",
    "Apply paint all over the bracket, but not to the bushes, alocromed spotfaces and holes: refer to PCS-2500.":
        "Нанесите краску по всей поверхности кронштейна, кроме втулок, подрезанных торцов с покрытием Alodine и отверстий: см. PCS-2500.",
    "REMOVE SHARP CORNERS": "СНЯТЬ ОСТРЫЕ УГЛЫ",
    "BREAK EDGE": "СНЯТЬ КРОМКУ",
    "BLENDED EDGES": "СГЛАЖЕННЫЕ КРОМКИ",
    "BOTH SIDES": "ОБЕ СТОРОНЫ",
    "INSTALLATION AND MACHINING": "УСТАНОВКА И МЕХАНИЧЕСКАЯ ОБРАБОТКА",
    "PRODUCE CROSS": "ВЫПОЛНИТЬ ПЕРЕКРЁСТНУЮ РИСКУ",
    "ASSEMBLY (INCLUDING STORAGE)": "СБОРКА (ВКЛЮЧАЯ ХРАНЕНИЕ)",
    "REPAIR POSITIONS": "ПОЗИЦИИ РЕМОНТА",
    "MAIN FITTING": "КОРПУС СТОЙКИ",
    "NO ARDROX BEYOND THIS POINT": "ARDROX НЕ НАНОСИТЬ ПОСЛЕ ЭТОЙ ТОЧКИ",
    "THIS LUG (2 PLACES) MUST NOT": "ЭТА ПРОУШИНА (2 МЕСТА) НЕ ДОЛЖНА",
    "If the bare metal is not damaged or corroded:": "Если основной металл не повреждён и не покрыт коррозией:",
    "If the bare metal is damaged or corroded:": "Если основной металл повреждён или покрыт коррозией:",
    "Assemble the two halves of the split housing.": "Соберите две половины разъёмного корпуса.",
    "If the pins are damaged, remove and discard the locating pins: refer to Figure 601.":
        "Если штифты повреждены, снимите и утилизируйте фиксирующие штифты: см. Рисунок 601.",
    "Make sure that the work area, the tools and the equipment are clean.":
        "Убедитесь в чистоте рабочей зоны, инструментов и оборудования.",
    "Make sure that all of the parts are correct to the data given in FITS AND CLEARANCES.":
        "Убедитесь в соответствии всех деталей данным, приведённым в разделе ПОСАДКИ И ЗАЗОРЫ.",
    "Machine parts to the tolerances specified in M-DLPS900 unless different instructions are in the procedures.":
        "Обрабатывайте детали с допусками, указанными в M-DLPS900, если в процедурах не указано иное.",
    "Torque all parts that have threads: refer to M-DLPS1002-1, unless different torque values are in FITS AND CLEARANCES.":
        "Затяните все резьбовые детали: см. M-DLPS1002-1, если иные значения моментов затяжки не указаны в ПОСАДКАХ И ЗАЗОРАХ.",
    "The electrical bonding resistance must not be more than 1 milliohm.":
        "Сопротивление электрического соединения не должно превышать 1 миллиом.",
    "Apply Ardrox AV100D.": "Нанесите Ardrox AV100D.",
    "Install the lubrication fittings: refer to PCS-7310": "Установите смазочные ниппели: см. PCS-7310",
    "Install the drag arm lower bush (20-360A).": "Установите нижнюю втулку рычага лобового сопротивления (20-360A).",
    "Forward pintle bush (20-250A)": "Передняя втулка штифта навеса (20-250A)",
    "Drag arm lower bush (20-360A)": "Нижняя втулка рычага лобового сопротивления (20-360A)",
    "Passivate the repair bearing: refer to AMS2700.": "Пассивируйте ремонтный подшипник: см. AMS2700.",
    "These repair parts are necessary": "Необходимые запасные части для ремонта",
    "INSERT\t(WITH THREADED INSERT)": "ВСТАВКА   (С РЕЗЬБОВОЙ ВСТАВКОЙ)",
    "(WITH THREADED INSERT)\tINSERT": "(С РЕЗЬБОВОЙ ВСТАВКОЙ)   ВСТАВКА",
    "Repair to Threaded Insert - Machining and Installation Figure 602":
        "Ремонт с резьбовой вставкой – Механическая обработка и установка Рисунок 602",
    "Repair Threaded Insert - Machining and Installation Figure 602":
        "Ремонтная резьбовая вставка – Механическая обработка и установка Рисунок 602",
    "This Repair, Messier-Dowty Limited Repair No. 450237480, has been withdrawn from use.":
        "Данный ремонт, Ремонт Messier-Dowty Limited № 450237480, изъят из обращения.",
    "CAUTION : FOR DEVIATIONS OUTSIDE THE LIMITS OF THIS REPAIR SCHEME CONTACT M-DL GLOUCESTER.":
        "ВНИМАНИЕ: ПРИ ОТКЛОНЕНИЯХ, ВЫХОДЯЩИХ ЗА ПРЕДЕЛЫ ДАННОЙ СХЕМЫ РЕМОНТА, ОБРАТИТЕСЬ В M-DL GLOUCESTER.",
    "Use the milliohm meter to measure the electrical bonding resistance of the bushes that follow:":
        "Используйте миллиомметр для измерения сопротивления электрического соединения следующих втулок:",
    "Use the Press Pad 460006268 and install the lubrication adapter (20-220), the identification washer (20-210) and the lubrication fitting (20-200A).":
        "Используйте прессовую подушку 460006268 и установите переходник смазки (20-220), идентификационную шайбу (20-210) и смазочный ниппель (20-200A).",
    "Use the Press Pad 460004330/147 to install the repair sleeve to the upper pivot bracket: refer to M-DLPS1011-5, PCS-5303 and Figure 602. Use Loctite grade 601, Material Ref. Item TBA.":
        "Используйте прессовую подушку 460004330/147 для установки ремонтной вставки в верхний кронштейн оси: см. M-DLPS1011-5, PCS-5303 и Рисунок 602. Используйте Loctite марки 601, Материал Поз. TBA.",
    "Use the Press Pad 460004330/91 to install the bearing (20-280): refer to M-DLPS1011-20. The ends of the bearing (20-280) must align with or go below the end of the bearing hole.":
        "Используйте прессовую подушку 460004330/91 для установки подшипника (20-280): см. M-DLPS1011-20. Торцы подшипника (20-280) должны быть заподлицо или ниже края отверстия под подшипник.",
    "CAUTION: MAKE SURE THAT SEALANT DOES NOT CAUSE A BLOCKAGE IN THE DRAIN HOLE NEAR THE BEARING (20-300).":
        "ВНИМАНИЕ: УБЕДИТЕСЬ, ЧТО ГЕРМЕТИК НЕ ПЕРЕКРЫВАЕТ ДРЕНАЖНОЕ ОТВЕРСТИЕ РЯДОМ С ПОДШИПНИКОМ (20-300).",
    "WARNING: DO NOT GET HYDRAULIC FLUID ON YOUR SKIN OR IN YOUR EYES. DO NOT BREATHE THE FUMES. ONLY USE IN A LOCATION THAT HAS A CONTINUOUS FLOW OF CLEAN AIR. HYDRAULIC FLUID IS POISONOUS AND DANGEROUS.":
        "ПРЕДУПРЕЖДЕНИЕ: НЕ ДОПУСКАЙТЕ ПОПАДАНИЯ ГИДРАВЛИЧЕСКОЙ ЖИДКОСТИ НА КОЖУ ИЛИ В ГЛАЗА. НЕ ВДЫХАЙТЕ ПАРЫ. ИСПОЛЬЗУЙТЕ ТОЛЬКО В ХОРОШО ПРОВЕТРИВАЕМОМ МЕСТЕ. ГИДРАВЛИЧЕСКАЯ ЖИДКОСТЬ ЯДОВИТА И ОПАСНА.",
    "Figure 703. Make sure that the primer paint is not visible at the joints after you apply the sealant.":
        "Рисунок 703. Убедитесь, что грунтовочная краска не видна на соединениях после нанесения герметика.",
    "Apply a fillet of sealant, Material Ref. Item TBA, to the joints between the repair bush and the transfer block: refer to PCS-7200 and Figure 602.":
        "Нанесите галтель из герметика, Материал Поз. TBA, на соединение ремонтной втулки и блока передачи: см. PCS-7200 и Рисунок 602.",
    "Machine the repair bush to the dimensions shown and calculated: refer to Figure 602. Make the surface finish 1,6 micrometers (63 micro-inches). Mchine face W to get the necessary flange thickness.":
        "Расточите ремонтную втулку до указанных и рассчитанных размеров: см. Рисунок 602. Чистота поверхности должна составлять 1,6 мкм (63 микродюйма). Расточьте поверхность W для получения необходимой толщины фланца.",
    "Put the rectangular seal 450237796 and the backing ring 450237797 in a bag and attach the bag to the cylinder. During assembly, use the rectangular seal 450237796 and the backing ring 450237797 in position of (17-210) and (17-220).":
        "Поместите прямоугольное уплотнение 450237796 и опорное кольцо 450237797 в пакет и прикрепите его к цилиндру. При сборке используйте прямоугольное уплотнение 450237796 и опорное кольцо 450237797 вместо (17-210) и (17-220).",
    "Apply sulphamate nickel plate to the reworked areas: refer to MIL-STD-868, solution 2 and Figure 601. The sulphamate nickel plate thickness must be sufficient to get the correct dimensions after machining the sulphamate nickel plate.":
        "Нанесите сульфаматное никелевое покрытие на переработанные области: см. MIL-STD-868, раствор 2 и Рисунок 601. Толщина сульфаматного никелевого покрытия должна быть достаточной для получения правильных размеров после его механической обработки.",
    "Apply sulphamate nickel plate to the reworked areas: refer to MIL-STD-868, solution 2 and Figure 601.The sulphamate nickel plate thickness must be sufficient to get the correct dimensions after machining the sulphamate nickel plate.":
        "Нанесите сульфаматное никелевое покрытие на переработанные области: см. MIL-STD-868, раствор 2 и Рисунок 601. Толщина сульфаматного никелевого покрытия должна быть достаточной для получения правильных размеров после его механической обработки.",
    "NOTE: If the repair is applied to the diameters A, B and C, apply sulphamate nickel plate: refer to Figure 601 and identify the part with the Safran Landing Systems repair number 64-4505141-00-ABC adjacent to the existing part number: refer to PCS-6000-07.":
        "ПРИМЕЧАНИЕ: Если ремонт применяется к диаметрам A, B и C, нанесите сульфаматное никелевое покрытие: см. Рисунок 601 и идентифицируйте деталь с номером ремонта Safran Landing Systems 64-4505141-00-ABC рядом с существующим номером детали: см. PCS-6000-07.",
    "Apply paint to the lock stay cardan but not to the repair bush(es): refer to PCS-2500.":
        "Нанесите краску на кардан фиксирующего звена, но не на ремонтную(ые) втулку(и): см. PCS-2500.",
    "Apply paint locally to the torque link, but not to the repair bushes: refer to PCS-2500.":
        "Нанесите краску локально на шлиц-шарнир, но не на ремонтные втулки: см. PCS-2500.",
    # ── Short labels and headings ──
    "Torque Data": "Данные момента затяжки",
    "JACKING DOME": "ДОМКРАТНЫЙ КУПОЛ",
    "Hold the Unit": "Удержание узла",
    "NUMERICAL INDEX": "ЧИСЛОВОЙ УКАЗАТЕЛЬ",
    "Numerical Index (where applicable)": "Числовой указатель (при наличии)",
    "LEFT CONFIGURATION": "ЛЕВАЯ КОНФИГУРАЦИЯ",
    "Assembly Clearance": "Монтажный зазор",
    "DETAILED PARTS LIST": "ПОДРОБНЫЙ ПЕРЕЧЕНЬ ДЕТАЛЕЙ",
    "Detailed Parts List": "Подробный перечень деталей",
    "a Detailed Parts List": "подробный перечень деталей",
    "FITS AND CLEARANCES": "ПОСАДКИ И ЗАЗОРЫ",
    "Allowable Clearance": "Допустимый зазор",
    "Initial Manufacturing Limits": "Исходные производственные пределы",
    "Torque Data Table 823": "Данные момента затяжки, Таблица 823",
    "Torque Data Table 823 (Continued)": "Данные момента затяжки, Таблица 823 (Продолжение)",
    "The Adapter 460006237": "Адаптер 460006237",
    "a Numerical Index (where applicable)": "числовой указатель (при наличии)",
    "Repeat the above step.": "Повторите вышеуказанный шаг.",
    "Vendor Codes, Names and Addresses": "Коды, наименования и адреса поставщиков",
    "The Lifting Bar 460006208": "Подъёмная штанга 460006208",
    "The Jacking Dome Adapter 460006223": "Адаптер домкратного купола 460006223",
    "The Support Arms 460006215": "Опорные рычаги 460006215",
    "The Towing Frame 460006216": "Буксировочная рама 460006216",
    "The Spherical Bearing Locator 460007282": "Фиксатор сферического подшипника 460007282",
    "The Pintle Location Assembly 460007281": "Сборка позиционирования штифта навеса 460007281",
    "The Location Frame 460007234 (for left configuration units) OR": "Рама позиционирования 460007234 (для узлов левой конфигурации) ИЛИ",
    "The Location Frame 460007235 (for right configuration units).": "Рама позиционирования 460007235 (для узлов правой конфигурации).",
    "SPECIAL TOOLS, FIXTURES AND EQUIPMENT": "СПЕЦИАЛЬНЫЕ ИНСТРУМЕНТЫ, ПРИСПОСОБЛЕНИЯ И ОБОРУДОВАНИЕ",
    "APPLY JOINTING COMPOUND TO": "НАНЕСТИ МОНТАЖНЫЙ СОСТАВ НА",
    "APPLY JOINTING COMPOUND TO CYLINDER": "НАНЕСТИ МОНТАЖНЫЙ СОСТАВ НА ЦИЛИНДР",
    "APPLY JOINTING COMPOUND TO SLIDING TUBE (18-80B) OR": "НАНЕСТИ МОНТАЖНЫЙ СОСТАВ НА СКОЛЬЗЯЩУЮ ТРУБУ (18-80B) ИЛИ",
    "APPLY ARDROX AV100D ONLY TO THIS AREA": "НАНЕСТИ ARDROX AV100D ТОЛЬКО НА ЭТУ ЗОНУ",
    "ARDROX AV100D TO OVERLAP WITH ARDROX AV100D APPLIED AT": "ARDROX AV100D ДОЛЖЕН ПЕРЕКРЫВАТЬСЯ С ARDROX AV100D, НАНЕСЁННЫМ НА",
    "TOUCH UP PAINT EXTERNALLY ON LOCK WASHER AND JACKING DOME": "ПОДКРАСИТЬ СНАРУЖИ СТОПОРНУЮ ШАЙБУ И ДОМКРАТНЫЙ КУПОЛ",
    "BUILD-UP TO MAKE SURE A CLEAR DRAINAGE PATH EXISTS": "НАПЛАВКА ДЛЯ ОБЕСПЕЧЕНИЯ СВОБОДНОГО ДРЕНАЖНОГО ПУТИ",
    "MAIN FITTING SUBASSEMBLY STAGE REFER TO FIGURE 701 (SHEET 2)": "ЭТАП СБОРКИ КОРПУСА СТОЙКИ СМ. РИСУНОК 701 (ЛИСТ 2)",
    "EACH LOCKING PLATE (15-80) TO BE DRILLED WITHIN THE DRILLING AREAS": "КАЖДАЯ СТОПОРНАЯ ПЛАСТИНА (15-80) ДОЛЖНА БЫТЬ ПРОСВЕРЛЕНА В ЗОНАХ СВЕРЛЕНИЯ",
    "UPPER DIAPHRAGM TUBE (15-390)": "ВЕРХНЯЯ ДИАФРАГМЕННАЯ ТРУБА (15-390)",
    "UPPER DIAPHRAGM TUBE (15-390A)": "ВЕРХНЯЯ ДИАФРАГМЕННАЯ ТРУБА (15-390A)",
    "TWO PIECE STOP ASSEMBLY (15-110)": "СБОРКА ДВУХКОМПОНЕНТНОГО УПОРА (15-110)",
    "(TYPICAL 2 PLACES) APPLY PRIMER PAINT TO BUSH DIA": "(ТИПИЧНО В 2 МЕСТАХ) НАНЕСТИ ГРУНТОВОЧНУЮ КРАСКУ НА ДИАМЕТР ВТУЛКИ",
    "VIEWS ARE FROM THE JACKING DOME (17-80) END OF THE SLIDING TUBE ASSEMBLY (17-240)": "ВИДЫ СО СТОРОНЫ ДОМКРАТНОГО КУПОЛА (17-80) СБОРКИ СКОЛЬЗЯЩЕЙ ТРУБЫ (17-240)",
    "(FEMALE) THREAD DO NOT APPLY JOINTING COMPOUND TO JACKING DOME (MALE) THREAD": "(ГАЙКА) РЕЗЬБА НЕ НАНОСИТЬ МОНТАЖНЫЙ СОСТАВ НА (БОЛТ) РЕЗЬБУ ДОМКРАТНОГО КУПОЛА",
    # ── Part references ──
    "The cap screws (7-30)": "Болты (7-30)",
    "The cap screws (4-90)": "Болты (4-90)",
    "The cap screws (6-120)": "Болты (6-120)",
    "The lock plate (15-180)": "Стопорная пластина (15-180)",
    "The lock plate (17-120)": "Стопорная пластина (17-120)",
    "The level tube (15-300)": "Трубка уровня (15-300)",
    "The locking plate (15-80)": "Стопорная пластина (15-80)",
    "The clapper seat (15-230)": "Седло откидного клапана (15-230)",
    "The transfer dowels (12-120)": "Фиксирующие штифты (12-120)",
    "The recoil orifice plate (15-70)": "Пластина дросселя отдачи (15-70)",
    "The upper pivot bracket (10-160)": "Верхний кронштейн оси (10-160)",
    "The bottom of the bracket (6-170)": "Нижняя часть кронштейна (6-170)",
    "The diaphragm subassembly (15-190)": "Сборка диафрагмы (15-190)",
    "The hole through the damper (9-160)": "Отверстие в амортизаторе (9-160)",
    "The two piece stop with inserts (15-130)": "Двухкомпонентный упор со вставками (15-130)",
    "The lubrication shaft subassembly (10-90)": "Сборка смазочного вала (10-90)",
    "The compression orifice plate (15-220)": "Пластина дросселя сжатия (15-220)",
    "Slave Link Subassembly (6-190)": "Сборка ведомого звена (6-190)",
    "Lower Slave Link Subassembly (6-290)": "Сборка нижнего ведомого звена (6-290)",
    "Lower Torque Link Subassembly (11-150)": "Сборка нижнего шлиц-шарнира (11-150)",
    "Upper Torque Link Subassembly (10-170)": "Сборка верхнего шлиц-шарнира (10-170)",
    "The shank and below the head of the cap screws (11-10)": "Стержень и место под головкой болтов (11-10)",
    "The Illustrated Parts List contains:": "Иллюстрированный перечень деталей содержит:",
    # ── Figure references ──
    "Sliding Tube Subassembly Figure 18": "Сборка скользящей трубы Рисунок 18",
    "Main Landing Gear Leg Figure 1": "Стойка основного шасси Рисунок 1",
    "Main Landing Gear Leg Figure 2 - Sheet 1": "Стойка основного шасси Рисунок 2 – Лист 1",
    "Main Landing Gear Leg Figure 2 - Sheet 2": "Стойка основного шасси Рисунок 2 – Лист 2",
    "Main Landing Gear Leg Figure 3": "Стойка основного шасси Рисунок 3",
    "Main Landing Gear Leg Figure 4": "Стойка основного шасси Рисунок 4",
    "Main Landing Gear Leg Figure 5": "Стойка основного шасси Рисунок 5",
    "Main Landing Gear Leg Figure 6": "Стойка основного шасси Рисунок 6",
    "Main Landing Gear Leg Figure 7": "Стойка основного шасси Рисунок 7",
    "Main Landing Gear Leg Figure 8": "Стойка основного шасси Рисунок 8",
    "Main Landing Gear Leg Figure 9": "Стойка основного шасси Рисунок 9",
    "Main Landing Gear Leg Figure 10": "Стойка основного шасси Рисунок 10",
    "Main Landing Gear Leg Figure 11": "Стойка основного шасси Рисунок 11",
    "Main Landing Gear Leg Figure 12": "Стойка основного шасси Рисунок 12",
    "Main Landing Gear Leg Figure 13": "Стойка основного шасси Рисунок 13",
    "Main Landing Gear Leg Figure 14": "Стойка основного шасси Рисунок 14",
    "Main Landing Gear Leg Figure 15": "Стойка основного шасси Рисунок 15",
    "Main Landing Gear Leg Figure 17": "Стойка основного шасси Рисунок 17",
    "Main Landing Gear Leg Figure 19": "Стойка основного шасси Рисунок 19",
    "Main Landing Gear Leg Figure 20": "Стойка основного шасси Рисунок 20",
    "Fits and Clearances - Key Diagram Figure 801": "Посадки и зазоры – Ключевая диаграмма Рисунок 801",
    "Fits and Clearances - Key Diagram Figure 802": "Посадки и зазоры – Ключевая диаграмма Рисунок 802",
    "Fits and Clearances Definitions (Refer to Tables 801 to 822)": "Определения посадок и зазоров (см. Таблицы 801–822)",
    "Pivot Bracket Subassembly (7-120) and Harness Support Bracket (7-100)": "Сборка кронштейна оси (7-120) и опорный кронштейн жгута (7-100)",
    # ── Procedures ──
    "Do para (1) and (2) two more times.": "Выполните пункты (1) и (2) ещё два раза.",
    "Do para (4) and (5) three more times.": "Выполните пункты (4) и (5) ещё три раза.",
    "Do para (11) and (12) three more times.": "Выполните пункты (11) и (12) ещё три раза.",
    "Do para (33) and (34) two more times.": "Выполните пункты (33) и (34) ещё два раза.",
    "Do the below procedure:": "Выполните следующую процедуру:",
    "Lubricate the areas that follow:": "Смажьте следующие зоны:",
    "Install these Special Tools in the main fitting subassembly (20-90):": "Установите следующие специальные инструменты в сборку корпуса стойки (20-90):",
    "Use these Special Tools to hold the unit as necessary during the procedure:": "Используйте следующие специальные инструменты для удержания узла по мере необходимости в ходе процедуры:",
    "Install the main fitting subassembly (20-90) in the Transport and Build Trolley 460006213.": "Установите сборку корпуса стойки (20-90) на транспортировочную и монтажную тележку 460006213.",
    "Torque the parts in Table 823 to their applicable values.": "Затяните детали, перечисленные в Таблице 823, до применимых значений момента затяжки.",
    "Reduce the torque on the locking nut (19-52) to zero.": "Уменьшите момент затяжки на самоконтрящейся гайке (19-52) до нуля.",
    "Reduce the torque on the retaining pins (13-10) to zero.": "Уменьшите момент затяжки на фиксирующих штифтах (13-10) до нуля.",
    "Reduce the torque on the nut subassembly (17-130) to zero.": "Уменьшите момент затяжки на гаечной сборке (17-130) до нуля.",
    "Remove the sharp edges: refer to M-DLPS900.": "Снимите острые кромки: см. M-DLPS900.",
    "Measure the angular difference (D) in alignment.": "Измерьте угловое отклонение (D) при совмещении.",
    "all necessary data for the procurement of parts.": "все необходимые данные для приобретения деталей.",
    "Install the dust cap (9-170) and the clamp (9-165).": "Установите пылезащитный колпачок (9-170) и хомут (9-165).",
    "Assemble the two piece stop subassembly (15-110).": "Соберите сборку двухкомпонентного упора (15-110).",
    "Install the clapper seat (15-230) over the baffle (15-240).": "Установите седло откидного клапана (15-230) на перегородку (15-240).",
    "Install the clapper seat (15-230) and the compression orifice plate (15-220) on the baffle (15-240).": "Установите седло откидного клапана (15-230) и пластину дросселя сжатия (15-220) на перегородку (15-240).",
    "Assemble the baffle (15-240) over the compression orifice plate (15-220) and the diaphragm (15-210A).": "Соберите перегородку (15-240) поверх пластины дросселя сжатия (15-220) и диафрагмы (15-210A).",
    "Assemble the two piece stop subassembly (15-110) over the upper diaphragm tube subassembly (15-360A).": "Соберите сборку двухкомпонентного упора (15-110) поверх сборки верхней диафрагменной трубы (15-360A).",
    "Install the bearing (15-270) over the upper diaphragm tube subassembly (15-360A).": "Установите подшипник (15-270) на сборку верхней диафрагменной трубы (15-360A).",
    "Apply final torque to the diaphragm subassembly (15-190A) to the value 200 to": "Затяните сборку диафрагмы (15-190A) окончательным моментом от 200 до",
    "AL. Set the Proximity Switch (7-230) and the Target (7-180)": "AL. Установите датчик приближения (7-230) и мишень (7-180)",
    "The ends of the cap screws (7-220), the nuts (7-200) and the washers (7-210). AM. Set the Proximity Switch (7-40) and the Target (6-130)": "Торцы болтов (7-220), гаек (7-200) и шайб (7-210). AM. Установите датчик приближения (7-40) и мишень (6-130)",
    "AJ. Do the main landing gear leg tests: refer to TESTING AND FAULT ISOLATION. AK. Retaining Pins (13-10)": "AJ. Выполните испытания стойки основного шасси: см. ИСПЫТАНИЯ И ПОИСК НЕИСПРАВНОСТЕЙ. AK. Фиксирующие штифты (13-10)",
    "AO. Do the electrical bonding resistance tests: refer to TESTING AND FAULT ISOLATION. AP. Examine the paint finish for damage. If there is damage, repair it: refer to REPAIR.": "AO. Выполните испытания сопротивления электрического соединения: см. ИСПЫТАНИЯ И ПОИСК НЕИСПРАВНОСТЕЙ. AP. Осмотрите лакокрасочное покрытие на наличие повреждений. При наличии повреждений выполните ремонт: см. REPAIR.",
    "AN. Do the tests and adjustments of the proximity switches (7-40 and 7-230): refer to TESTING AND FAULT ISOLATION.": "AN. Выполните испытания и регулировку датчиков приближения (7-40 и 7-230): см. ИСПЫТАНИЯ И ПОИСК НЕИСПРАВНОСТЕЙ.",
    "Do the piston leakage test. Refer to TESTING AND FAULT ISOLATION.": "Выполните испытание поршня на герметичность. См. ИСПЫТАНИЯ И ПОИСК НЕИСПРАВНОСТЕЙ.",
    "Do the piston leakage test: refer to TESTING AND FAULT ISOLATION.": "Выполните испытание поршня на герметичность: см. ИСПЫТАНИЯ И ПОИСК НЕИСПРАВНОСТЕЙ.",
    "Install the split pins (5-320A) and safety them: refer to PCS-7610.": "Установите шплинты (5-320A) и законтрите их: см. PCS-7610.",
    "Install the split pins (8-40) and safety them: refer to M-DLPS1011-1.": "Установите шплинты (8-40) и законтрите их: см. M-DLPS1011-1.",
    "Install the split pins (7-60) and safety them: refer to M-DLPS1011-1.": "Установите шплинты (7-60) и законтрите их: см. M-DLPS1011-1.",
    "Install the split pins (4-110) and safety them: refer to M-DLPS1011-1.": "Установите шплинты (4-110) и законтрите их: см. M-DLPS1011-1.",
    "Install the split pins (12-10) and safety them: refer to M-DLPS1011-1.": "Установите шплинты (12-10) и законтрите их: см. M-DLPS1011-1.",
    "Install the split pins (8-100) and safety them: refer to M-DLPS1011-1.": "Установите шплинты (8-100) и законтрите их: см. M-DLPS1011-1.",
    "Install the split pins (2-230) and safety them: refer to M-DLPS1011-1.": "Установите шплинты (2-230) и законтрите их: см. M-DLPS1011-1.",
    "Install the split pins (4-270) and safety them: refer to M-DLPS1011-1.": "Установите шплинты (4-270) и законтрите их: см. M-DLPS1011-1.",
    "Measure the gap between the target (6-130) and the proximity switch (7-40).": "Измерьте зазор между мишенью (6-130) и датчиком приближения (7-40).",
    "Measure the gap between the target (7-180) and the proximity switch (7-230).": "Измерьте зазор между мишенью (7-180) и датчиком приближения (7-230).",
    "Turn the valve stem (12-90) counterclockwise until it stops and install the cap screw (1).": "Поверните шток клапана (12-90) против часовой стрелки до упора и установите болт (1).",
    "Use the Lock Punch 460006589 to safety the locking washer (19-54) in four places at the same distance apart: refer to M-DLPS1011-1.": "Используйте ударный инструмент 460006589 для контровки стопорной шайбы (19-54) в четырёх равноудалённых точках: см. M-DLPS1011-1.",
    "Hold the cylinder (17-230) with the Bench Clamp MT1025 and Holding Blocks MT1026/63.": "Зафиксируйте цилиндр (17-230) с помощью настольных тисков MT1025 и удерживающих блоков MT1026/63.",
    "CAUTION: MAKE SURE THAT YOU DO THE SPECIAL DIMENSION CHECK OF THE ROD (17-160): REFER TO CHECK, PARA 3. A. IF NOT, IT CAN CAUSE DAMAGE TO THE COMPONENT.": "ВНИМАНИЕ: УБЕДИТЕСЬ В ВЫПОЛНЕНИИ СПЕЦИАЛЬНОЙ ПРОВЕРКИ РАЗМЕРОВ ШТОКА (17-160): СМ. CHECK, ПУН. 3.A. В ПРОТИВНОМ СЛУЧАЕ КОМПОНЕНТ МОЖЕТ БЫТЬ ПОВРЕЖДЁН.",
    "CAUTION: DO NOT USE WATER TO HEAT THE BACKING RINGS. WATER ABSORBED BY THE BACKING RINGS WILL CAUSE CORROSION.": "ВНИМАНИЕ: НЕ ИСПОЛЬЗУЙТЕ ВОДУ ДЛЯ НАГРЕВА ОПОРНЫХ КОЛЕЦ. ВОДА, ПОГЛОЩЁННАЯ ОПОРНЫМИ КОЛЬЦАМИ, ВЫЗОВЕТ КОРРОЗИЮ.",
    "CAUTION: DO NOT CAUSE A BLOCKAGE WITH THE SEALANT IN THE CHANNELS FORMED BETWEEN THE HOUSING (12-170) AND THE MAIN FITTING SUBASSEMBLY (20-90).": "ВНИМАНИЕ: НЕ ДОПУСКАЙТЕ ЗАСОРЕНИЯ ГЕРМЕТИКОМ КАНАЛОВ, ОБРАЗОВАННЫХ МЕЖДУ КОРПУСОМ (12-170) И СБОРКОЙ КОРПУСА СТОЙКИ (20-90).",
    "CAUTION: THE TWO PIECE STOP WITH INSERTS (15-130) IS MADE FROM TWO PARTS THAT MAKE A SET. DO NOT USE WITH A PART FROM ANOTHER SET.": "ВНИМАНИЕ: ДВУХКОМПОНЕНТНЫЙ УПОР СО ВСТАВКАМИ (15-130) СОСТОИТ ИЗ ДВУХ ЧАСТЕЙ, ОБРАЗУЮЩИХ КОМПЛЕКТ. НЕ ИСПОЛЬЗОВАТЬ С ДЕТАЛЬЮ ИЗ ДРУГОГО КОМПЛЕКТА.",
    "CAUTION: MAKE SURE THAT YOU ASSEMBLE THE LONGER SLEEVE (9-200) IN THE UPPER TORQUE LINK SUBASSEMBLY (10-170) AND THE SHORTER SLEEVE (9-90) IN THE LOWER TORQUE LINK SUBASSEMBLY (11-150).": "ВНИМАНИЕ: УБЕДИТЕСЬ, ЧТО БОЛЕЕ ДЛИННАЯ ВСТАВКА (9-200) УСТАНОВЛЕНА В СБОРКЕ ВЕРХНЕГО ШЛИЦ-ШАРНИРА (10-170), А БОЛЕЕ КОРОТКАЯ ВСТАВКА (9-90) – В СБОРКЕ НИЖНЕГО ШЛИЦ-ШАРНИРА (11-150).",
    "CAUTION: ARDROX AV100D MUST NOT ENTER THE BORE OF THE INFLATION VALVE (13-110) IN THE UPPER DIAPHRAGM TUBE (15-390).": "ВНИМАНИЕ: ARDROX AV100D НЕ ДОЛЖЕН ПОПАДАТЬ В ОТВЕРСТИЕ КЛАПАНА НАКАЧКИ (13-110) В ВЕРХНЕЙ ДИАФРАГМЕННОЙ ТРУБЕ (15-390).",
    "Install the two backing rings into the groove in the piston (17-200). Use string or tape to help the backing rings into the groove.": "Установите два опорных кольца в канавку поршня (17-200). Используйте нить или ленту для укладки опорных колец в канавку.",
    "Install the piston (17-200) into the cylinder (17-230A) until the seal (17-190) is held inside the cylinder (17-230A) and stop.": "Вставьте поршень (17-200) в цилиндр (17-230A) до тех пор, пока уплотнение (17-190) не зафиксируется внутри цилиндра (17-230A), и остановитесь.",
    "Install the piston (17-200) into the cylinder (17-230A) until the seal\t(17-190) is held inside the cylinder (17-230A) and stop.": "Вставьте поршень (17-200) в цилиндр (17-230A) до тех пор, пока уплотнение (17-190) не зафиксируется внутри цилиндра (17-230A), и остановитесь.",
    "Carefully insert the cylinder (17-230A) into the sliding tube subassembly (17-240B)": "Аккуратно вставьте цилиндр (17-230A) в сборку скользящей трубы (17-240B)",
    "Carefully insert the cylinder (17-230A) into the sliding tube subassembly (17-240B) or (17-240C). Align the holes for the valve support (17-50).": "Аккуратно вставьте цилиндр (17-230A) в сборку скользящей трубы (17-240B) или (17-240C). Совместите отверстия для опоры клапана (17-50).",
    "To the sliding tube (18-80B) where the cylinder (17-230A) touches,": "На скользящую трубу (18-80B) в месте контакта с цилиндром (17-230A),",
    "Install the valve support (17-50) to the sliding tube subassembly (17-240) with the flat washers (17-40A) and the screw cap (17-30).": "Установите опору клапана (17-50) в сборку скользящей трубы (17-240) с помощью плоских шайб (17-40A) и крышки (17-30).",
    "Install the valve support (17-50) to the sliding tube subassembly (17-240) with the flat washers (17-40A) and the cap screws (17-30).": "Установите опору клапана (17-50) в сборку скользящей трубы (17-240) с помощью плоских шайб (17-40A) и болтов (17-30).",
    "Assemble the cylinder (17-230) and its related parts to the sliding tube subassembly (17-240 or 17-240A). Align the holes for the valve support (17-50).": "Присоедините цилиндр (17-230) и связанные с ним детали к сборке скользящей трубы (17-240 или 17-240A). Совместите отверстия для опоры клапана (17-50).",
    "Align the upper torque link subassembly (10-170) and the lower torque link subassembly (11-150): put the spacer (9-80) between them and install the pin (9-70).": "Совместите сборку верхнего шлиц-шарнира (10-170) и сборку нижнего шлиц-шарнира (11-150): поместите проставку (9-80) между ними и установите штифт (9-70).",
    "Install the screws (15-90) and the tab washers (15-100) through the locking plates (15-80) into the two piece stop with inserts (15-130).": "Установите винты (15-90) и стопорные шайбы (15-100) через стопорные пластины (15-80) в двухкомпонентный упор со вставками (15-130).",
    "Install the nuts (2-210), the retainers (2-220) and the transfer block subassembly (2-290B) and align the threaded insert (2-50).": "Установите гайки (2-210), фиксаторы (2-220) и сборку блока передачи (2-290B) и совместите резьбовую вставку (2-50).",
    "Install the nuts (2-210), the retainers (2-220) and the transfer block subassembly (2-290 and 2-290A) and align the threaded insert (2-50).": "Установите гайки (2-210), фиксаторы (2-220) и сборку блока передачи (2-290 и 2-290A) и совместите резьбовую вставку (2-50).",
    "Carefully position the baffle (15-240) over the rod (17-160) and insert the upper diaphragm tube (15-390A) into the sliding piston (18-80B) or(18-80C) or (18-80F)": "Аккуратно установите перегородку (15-240) на шток (17-160) и вставьте верхнюю диафрагменную трубу (15-390A) в скользящий поршень (18-80B) или (18-80C) или (18-80F)",
    "Put the piston (17-200) and Keep Ring DRT68792 over the open end of the cylinder (17-230). Push the piston (17-200) until the sealing ring and the backing rings are held in the cylinder (17-230): remove the Keep Ring DRT68792.": "Наденьте поршень (17-200) и удерживающее кольцо DRT68792 на открытый конец цилиндра (17-230). Вдавите поршень (17-200) до тех пор, пока уплотнительное и опорные кольца не зафиксируются в цилиндре (17-230): снимите удерживающее кольцо DRT68792.",
    "Lubricate the Keep Ring DRT68792 with hydraulic fluid, Material Ref. Item 02-501. Install the Keep Ring DRT68792 over the piston (17-200), the sealing ring and the backing rings.": "Смажьте удерживающее кольцо DRT68792 гидравлической жидкостью, Материал Поз. 02-501. Установите удерживающее кольцо DRT68792 на поршень (17-200), уплотнительное и опорные кольца.",
    "Lubricate the lubrication fittings with grease: refer to M-DLPS1005-1 and PCS-7300. Make sure that the grease paths are not blocked and the grease flows smoothly.": "Смажьте смазочные ниппели консистентной смазкой: см. M-DLPS1005-1 и PCS-7300. Убедитесь, что каналы смазки не заблокированы и смазка поступает свободно.",
    "Lubricate the lubrication fittings (17-270A) or (18-52) with grease: refer to M-DLPS1005-1 and PCS-7300. Make sure the grease paths are not blocked and the grease flows smoothly.": "Смажьте смазочные ниппели (17-270A) или (18-52) консистентной смазкой: см. M-DLPS1005-1 и PCS-7300. Убедитесь, что каналы смазки не заблокированы и смазка поступает свободно.",
    "Apply Zinc loaded Jointing compound, Molykote 111 to the sliding tube (18-80B)": "Нанесите монтажный состав с добавкой цинка, Molykote 111, на скользящую трубу (18-80B)",
    "Use the milliohmeter to measure the electrical bonding resistance of the bushes (18-40). The electrical bonding resistance must not be more than 1 milliohm.": "Используйте миллиомметр для измерения сопротивления электрического соединения втулок (18-40). Сопротивление электрического соединения не должно превышать 1 миллиом.",
    "Remove the sliding tube subassembly (17-240) and its related parts from the Build Trolley 460007240: use the Lifting Tackle 460006211.": "Снимите сборку скользящей трубы (17-240) и связанные с ней детали с монтажной тележки 460007240: используйте подъёмные приспособления 460006211.",
    "Install the Sliding Tube Subassembly (17-240) and its Related Parts": "Установка сборки скользящей трубы (17-240) и связанных деталей",
    "Install the sliding tube subassembly (17-240) and its related parts in the main fitting subassembly (20-90). Align the holes in the upper diaphragm tube subassembly (15-360) and the main fitting subassembly (20-90) for the pin (13-190).": "Установите сборку скользящей трубы (17-240) и связанные с ней детали в сборку корпуса стойки (20-90). Совместите отверстия в сборке верхней диафрагменной трубы (15-360) и сборке корпуса стойки (20-90) для штифта (13-190).",
    "Move the lower bearing subassembly (16-110) into the main fitting subassembly (20-90). Align the holes for the retaining pins (13-10).": "Переместите сборку нижнего подшипника (16-110) в сборку корпуса стойки (20-90). Совместите отверстия для фиксирующих штифтов (13-10).",
    "Carefully install the inflation valve subassembly (13-90) through the pin (13-190) into the upper diaphragm tube subassembly (15-360).": "Аккуратно установите сборку клапана накачки (13-90) через штифт (13-190) в сборку верхней диафрагменной трубы (15-360).",
    "Refer to Figure 702 and M-DLPS405-10. Bond the wiring diagram plate (1-110) to the main fitting subassembly (20-90).": "См. Рисунок 702 и M-DLPS405-10. Приклейте пластину со схемой электрических соединений (1-110) к сборке корпуса стойки (20-90).",
    "Refer to the DETAILED PARTS LIST and identify the label (17-290 only) with the shock absorber subassembly part number of the modification standard to which the unit is being assembled. Also mark the serial number on the label (18-70): refer to PCS-6000-05.": "Обратитесь к ПОДРОБНОМУ ПЕРЕЧНЮ ДЕТАЛЕЙ и нанесите на ярлык (только 17-290) номер детали сборки амортизатора в соответствии со стандартом модификации, применяемым при сборке узла. Также нанесите серийный номер на ярлык (18-70): см. PCS-6000-05.",
    "Refer to M-DLPS1005-1. Lubricate these parts with Hydraulic Fluid MIL-H-5606:": "См. M-DLPS1005-1. Смажьте следующие детали гидравлической жидкостью MIL-H-5606:",
    "Refer to M-DLPS1011-1. Lubricate these parts with Hydraulic Fluid MIL-H-5606:": "См. M-DLPS1011-1. Смажьте следующие детали гидравлической жидкостью MIL-H-5606:",
    "Lubricate these parts with Hydraulic fluid MIL-PRF-5606, Material Ref. Item TBA:": "Смажьте следующие детали гидравлической жидкостью MIL-PRF-5606, Материал Поз. TBA:",
    "Refer to M-DLPS709-14. Apply jointing compound, Material Ref. Item TBA, to these parts:": "См. M-DLPS709-14. Нанесите монтажный состав, Материал Поз. TBA, на следующие детали:",
    "Refer to M-DLPS615. Apply protective varnish to the wiring diagram plate (1-110).": "См. M-DLPS615. Нанесите защитный лак на пластину со схемой электрических соединений (1-110).",
    "Identify the part with the Messier-Dowty Limited repair number 450267365 and the repair number shown in REPAIR No. 18-4 adjacent to the part number: refer to PCS-6000-04 or PCS-6000-06.": "Идентифицируйте деталь с номером ремонта Messier-Dowty Limited 450267365 и номером ремонта, указанным в РЕМОНТ № 18-4, рядом с номером детали: см. PCS-6000-04 или PCS-6000-06.",
    "Identify the part with the Messier-Dowty Limited repair number 450267365 and the repair number shown in REPAIR No. 18-4 adjacent to the part number: refer to PCS-6000-07.": "Идентифицируйте деталь с номером ремонта Messier-Dowty Limited 450267365 и номером ремонта, указанным в РЕМОНТ № 18-4, рядом с номером детали: см. PCS-6000-07.",
    "Remove the paint from the second stage cylinder: refer to PCS-2700.": "Снимите краску с цилиндра второй ступени: см. PCS-2700.",
    "Identify the transfer block with the Messier-Dowty Limited repair number 450266420 adjacent to the part number: refer to PCS-6000-05.": "Идентифицируйте блок передачи с номером ремонта Messier-Dowty Limited 450266420 рядом с номером детали: см. PCS-6000-05.",
    "Identify the transfer block with the Messier-Dowty Limited repair number 450266420 adjacent to the part number: refer to PCS-6000-07.": "Идентифицируйте блок передачи с номером ремонта Messier-Dowty Limited 450266420 рядом с номером детали: см. PCS-6000-07.",
    "Identify the transfer block with the Messier-Dowty Limited repair number 450266421 adjacent to the part number: refer to PCS-6000-05.": "Идентифицируйте блок передачи с номером ремонта Messier-Dowty Limited 450266421 рядом с номером детали: см. PCS-6000-05.",
    "Identify the transfer block with the Messier-Dowty Limited repair number 450266421 adjacent to the part number: refer to PCS-6000-07.": "Идентифицируйте блок передачи с номером ремонта Messier-Dowty Limited 450266421 рядом с номером детали: см. PCS-6000-07.",
    "Identify the repair threaded insert with the Messier-Dowty Limited repair number 450266420 adjacent to the part number: refer to PCS-6000-05.": "Идентифицируйте ремонтную резьбовую вставку с номером ремонта Messier-Dowty Limited 450266420 рядом с номером детали: см. PCS-6000-05.",
    "Identify the repair threaded insert with the Messier-Dowty Limited repair number 450266420 adjacent to the part number: refer to PCS-6000-07.": "Идентифицируйте ремонтную резьбовую вставку с номером ремонта Messier-Dowty Limited 450266420 рядом с номером детали: см. PCS-6000-07.",
    "Identify the repair threaded insert with the Messier-Dowty Limited repair number 450266421 adjacent to the part number: refer to PCS-6000-05.": "Идентифицируйте ремонтную резьбовую вставку с номером ремонта Messier-Dowty Limited 450266421 рядом с номером детали: см. PCS-6000-05.",
    "Identify the repair threaded insert with the Messier-Dowty Limited repair number 450266421 adjacent to the part number: refer to PCS-6000-07.": "Идентифицируйте ремонтную резьбовую вставку с номером ремонта Messier-Dowty Limited 450266421 рядом с номером детали: см. PCS-6000-07.",
    "Identify the split housing halves with the Messier-Dowty Limited repair number 450266520 adjacent to the part number: refer to PCS-6000-04.": "Идентифицируйте половины разъёмного корпуса с номером ремонта Messier-Dowty Limited 450266520 рядом с номером детали: см. PCS-6000-04.",
    "Calculate the diameter C for the repair threaded insert, use the formula:": "Рассчитайте диаметр C для ремонтной резьбовой вставки по формуле:",
    "Calculate the diameter C for the repair threaded insert, use formula:": "Рассчитайте диаметр C для ремонтной резьбовой вставки по формуле:",
    "Use the Press Pad 460006620 and the Drift 460004331/2 and install the drag arm upper bushes (20-390A): refer to PCS-5105-2. Make sure that the drag arm upper bushes are aligned as shown in Figure 703.": "Используйте прессовую подушку 460006620 и оправку 460004331/2 и установите верхние втулки рычага лобового сопротивления (20-390A): см. PCS-5105-2. Убедитесь, что верхние втулки рычага лобового сопротивления совмещены, как показано на Рисунке 703.",
    "Use the Press Pad 460004330/169 and the Drift 460004331/2 and install the drag arm lower bush (20-360A): refer to PCS-5105-2.": "Используйте прессовую подушку 460004330/169 и оправку 460004331/2 и установите нижнюю втулку рычага лобового сопротивления (20-360A): см. PCS-5105-2.",
    "Use the Pull Bar DRT66012 and the Bush Assembly Tool DRT68300 and install the bushes (20-320) while the primer paint is still wet: refer to PCS-5120.": "Используйте монтажную тягу DRT66012 и инструмент для установки втулок DRT68300 и установите втулки (20-320), пока грунтовочная краска ещё влажная: см. PCS-5120.",
    "Use the Press Pad Assembly 460006603, the Guide Bush 460006604 and the Alignment Bar 460006601 and install the retraction actuator lug bush (20-230A) while the primer paint is still wet: refer to PCS-5105-4.": "Используйте сборку прессовой подушки 460006603, направляющую втулку 460006604 и центрирующую штангу 460006601 и установите втулку ушка привода уборки (20-230A), пока грунтовочная краска ещё влажная: см. PCS-5105-4.",
    "Use the Press Pad 460006268 and install the lubrication adapters (20-190), the identification washers (20-180) and the lubrication fittings (20-170A).": "Используйте прессовую подушку 460006268 и установите переходники смазки (20-190), идентификационные шайбы (20-180) и смазочные ниппели (20-170A).",
    "Use the Press Pad 460006268 and install the lubrication adapter (20-160), the identification washer (20-150) and the lubrication fitting (20-140A).": "Используйте прессовую подушку 460006268 и установите переходник смазки (20-160), идентификационную шайбу (20-150) и смазочный ниппель (20-140A).",
    "Use the Press Pad 460006268 and install the lubrication adapter (20-130), the identification washer (20-120) and the lubrication fitting (20-110A).": "Используйте прессовую подушку 460006268 и установите переходник смазки (20-130), идентификационную шайбу (20-120) и смазочный ниппель (20-110A).",
    "Machine (do not grind) the sulphamate nickel plate to the dimensions shown: refer to M-DLPS900, M-DLPS1000 and Figure 602. Make the surface finish 3,2 micrometers (125 micro-inches) for diameter A and faces E and F. Make the surface finish 1,6 micrometers (63 micro-inches) for the adjacent diameter C.": "Расточите (не шлифовать) сульфаматное никелевое покрытие до указанных размеров: см. M-DLPS900, M-DLPS1000 и Рисунок 602. Чистота поверхности должна составлять 3,2 мкм (125 микродюймов) для диаметра A и поверхностей E и F. Чистота поверхности должна составлять 1,6 мкм (63 микродюйма) для смежного диаметра C.",
    "Lubricate the bushes (10-240 and 10-250) through the lubrication fittings (10-180 and": "Смажьте втулки (10-240 и 10-250) через смазочные ниппели (10-180 и",
    "Lubricate the bushes (11-220 and 11-230) through the lubrication fittings (11-160 and": "Смажьте втулки (11-220 и 11-230) через смазочные ниппели (11-160 и",
    "The mating faces of the main fitting subassembly (20-90) and the bracket assembly (5-10).": "Сопрягаемые поверхности сборки корпуса стойки (20-90) и сборки кронштейна (5-10).",
    "The mating faces of the bracket subassembly (4-330) and the main fitting subassembly (20-90)": "Сопрягаемые поверхности сборки кронштейна (4-330) и сборки корпуса стойки (20-90)",
    "The mating faces of the main fitting subassembly (20-90) and the bracket subassembly (5-90).": "Сопрягаемые поверхности сборки корпуса стойки (20-90) и сборки кронштейна (5-90).",
    "The mating faces of the main fitting subassembly (20-90) and the bracket subassembly (5-270).": "Сопрягаемые поверхности сборки корпуса стойки (20-90) и сборки кронштейна (5-270).",
    "The mating faces of the transfer block subassembly (2-290B) and the main fitting sub-assembly (20-90)": "Сопрягаемые поверхности сборки блока передачи (2-290B) и сборки корпуса стойки (20-90)",
    "The mating faces of the transfer block subassembly (2-290 and 2-290A) and the main fitting sub-assembly (20-90)": "Сопрягаемые поверхности сборки блока передачи (2-290 и 2-290A) и сборки корпуса стойки (20-90)",
    "The faces of the bracket subassembly (8-90) and the sliding tube subassembly (17-240) that will touch": "Поверхности сборки кронштейна (8-90) и сборки скользящей трубы (17-240), которые будут соприкасаться",
    "The harness support bracket (7-100) where it will touch the main fitting subassembly (20-90)": "Опорный кронштейн жгута (7-100) в месте касания сборки корпуса стойки (20-90)",
    "The pivot bracket subassembly (7-120) where it will touch the main fitting subassembly (20-90)": "Сборка кронштейна оси (7-120) в месте касания сборки корпуса стойки (20-90)",
    "These materials are necessary to prevent damage to the unit during usual storage conditions.": "Данные материалы необходимы для защиты узла от повреждений в обычных условиях хранения.",
    "Remove the Bench Clamp MT1025, the Holding Blocks 460006406 and the Torque Reactor 460007278.": "Снимите настольные тиски MT1025, удерживающие блоки 460006406 и реактор момента затяжки 460007278.",
    "Remove the Bench Clamp MT1025, the Holding Blocks 460006406 and the Torque Reactor 460006407 from the diaphragm tube subassembly (15-360A).": "Снимите настольные тиски MT1025, удерживающие блоки 460006406 и реактор момента затяжки 460006407 со сборки диафрагменной трубы (15-360A).",
    "Put a heat shrink sleeve (1-41) and a ferrule (1-43) over each end of the Bowden cable (1-45).": "Наденьте термоусадочную трубку (1-41) и обжимную гильзу (1-43) на каждый конец троса Боудена (1-45).",
    "Put the heat shrink sleeves (1-41) over the ferrules (1-43) and shrink into place: refer to M-DLPS821.": "Наденьте термоусадочные трубки (1-41) на обжимные гильзы (1-43) и усадите на место: см. M-DLPS821.",
    "Make a loop in each end of the Bowden cable (1-45) to get the dimensions shown, and insert the ends into the applicable ferrule (1-43). Crimp the ferrules (1-43) onto the ends of the Bowden cable (1-45).": "Сформируйте петлю на каждом конце троса Боудена (1-45) для получения указанных размеров и вставьте концы в соответствующую обжимную гильзу (1-43). Обожмите гильзы (1-43) на концах троса Боудена (1-45).",
    "Assemble one end of the Bowden cable (1-45) to the cross bolt (1-49) and the other end of the Bowden cable (1-45) to the cross bolt (1-47).": "Присоедините один конец троса Боудена (1-45) к поперечному болту (1-49), а другой конец троса Боудена (1-45) – к поперечному болту (1-47).",
    "NOTE: Make sure the holes in the new inner liner align with the holes in the lower bearing subassembly.": "ПРИМЕЧАНИЕ: Убедитесь, что отверстия в новом внутреннем вкладыше совпадают с отверстиями в сборке нижнего подшипника.",
    "NOTE: Make sure that the holes in the liner (16A-117) are aligned with the holes in the lower bearing subassembly (16A-113).": "ПРИМЕЧАНИЕ: Убедитесь, что отверстия во вкладыше (16A-117) совмещены с отверстиями в сборке нижнего подшипника (16A-113).",
    "NOTE: All dimensions mentioned in the Fits and Clearances table do not include cadmium plate, unless otherwise stated.": "ПРИМЕЧАНИЕ: Все размеры в таблице посадок и зазоров не включают кадмиевое покрытие, если не указано иное.",
    "NOTE: Align the slot in the lock plate (15-180) with the tongue on the upper diaphragm tube subassembly (15-360). If necessary, turn the lock plate (15-180) to align them correctly.": "ПРИМЕЧАНИЕ: Совместите паз в стопорной пластине (15-180) с выступом на сборке верхней диафрагменной трубы (15-360). При необходимости поверните стопорную пластину (15-180) для правильного совмещения.",
    "NOTE: If necessary, turn the lock plate (15-180) upside down to align the lock plate (15-180) and the diaphragm tube subassembly (15-360A) correctly.": "ПРИМЕЧАНИЕ: При необходимости переверните стопорную пластину (15-180) вверх дном для правильного совмещения стопорной пластины (15-180) и сборки диафрагменной трубы (15-360A).",
    "While the primer is wet, install the bushes (3-150): use the Press Pad 460004330/130 and a drift.": "Пока грунтовка ещё влажная, установите втулки (3-150): используйте прессовую подушку 460004330/130 и оправку.",
    "While the primer is wet, install the bush (15-380) in the upper diaphragm tube (15-390): use the Press Pad 460004330/135.": "Пока грунтовка ещё влажная, установите втулку (15-380) в верхнюю диафрагменную трубу (15-390): используйте прессовую подушку 460004330/135.",
    "While the primer is wet, install the bush (15-370) in the upper diaphragm tube (15-390): use the Press Pad 460004330/134.": "Пока грунтовка ещё влажная, установите втулку (15-370) в верхнюю диафрагменную трубу (15-390): используйте прессовую подушку 460004330/134.",
    "Install the two piece stop subassembly (15-110) over the upper diaphragm tube subassembly (15-360A).": "Установите сборку двухкомпонентного упора (15-110) на сборку верхней диафрагменной трубы (15-360A).",
    "Use the Assembly and Extraction Tool 460006410 to install the level tube (15-300) to the upper diaphragm tube (15-390A).": "Используйте монтажно-демонтажный инструмент 460006410 для установки трубки уровня (15-300) в верхнюю диафрагменную трубу (15-390A).",
    "Install the O-ring seal (15-310) on the level tube (15-300): use the Assembly and Extraction Tool 460006410 to install the level tube (15-300) in the upper diaphragm tube subassembly (15-360).": "Установите уплотнительное кольцо (15-310) на трубку уровня (15-300): используйте монтажно-демонтажный инструмент 460006410 для установки трубки уровня (15-300) в сборку верхней диафрагменной трубы (15-360).",
    # ── Recurring sentence fragments ──
    "Make the surface finish 1,6 micrometers (63 micro-inches).": "Чистота поверхности должна составлять 1,6 мкм (63 микродюйма).",
    "Make the surface finish 3,2 micrometers (125 micro-inches).": "Чистота поверхности должна составлять 3,2 мкм (125 микродюймов).",
    "The cadmium plate thickness must be between 0,010 and 0,015 mm (0.0004 and 0.0006 in).": "Толщина кадмиевого покрытия должна быть от 0,010 до 0,015 мм (0,0004–0,0006 дюйма).",
    "The cadmium plate thickness must be between 0,010 and 0,020 mm (0.0004 and 0.0008 in).": "Толщина кадмиевого покрытия должна быть от 0,010 до 0,020 мм (0,0004–0,0008 дюйма).",
    "The cadmium plate thickness must be between 0,010 and 0,020 mm (0.0004 and 0.0008 in): refer to PCS-2100 or PCS-2141.": "Толщина кадмиевого покрытия должна быть от 0,010 до 0,020 мм (0,0004–0,0008 дюйма): см. PCS-2100 или PCS-2141.",
    "No bare metal is permitted.": "Обнажение основного металла не допускается.",
    "The surface finish must be 1,6 micrometers (63 micro-inches).": "Чистота поверхности должна составлять 1,6 мкм (63 микродюйма).",
    "Do not reduce the lug width E below": "Не уменьшайте ширину проушины E ниже",
    "Stress relieve the reworked areas for 4 hours at 185 to 195 oC (366 to 384 oF): refer to PCS-2101.":
        "Снимите напряжения в переработанных областях в течение 4 ч при 185–195 °C (366–384 °F): см. PCS-2101.",
    # ── Recurring trailing sentences (high-frequency) ──
    "The surface finish must be 2,5 micrometers (100 micro-inches).": "Чистота поверхности должна составлять 2,5 мкм (100 микродюймов).",
    "The surface finish must be 1,6 micrometers (63 micro-inches) or better: refer to M-DLPS900, M-DLPS1000 and Figure 601.":
        "Чистота поверхности должна составлять 1,6 мкм (63 микродюйма) или лучше: см. M-DLPS900, M-DLPS1000 и Рисунок 601.",
    "The surface finish must be 1,6 micrometers (63 micro-inches) or better.": "Чистота поверхности должна составлять 1,6 мкм (63 микродюйма) или лучше.",
    "The surface finish must be 3,2 micrometers (125 micro-inches).": "Чистота поверхности должна составлять 3,2 мкм (125 микродюймов).",
    "THE SURFACE FINISH MUST BE\tOR BETTER UNLESS GIVEN DIFFERENTLY.": "ЧИСТОТА ПОВЕРХНОСТИ ДОЛЖНА БЫТЬ НЕ ХУЖЕ УКАЗАННОГО ЗНАЧЕНИЯ, ЕСЛИ НЕ УКАЗАНО ИНАЧЕ.",
    "DEBURR THE SHARP EDGES WITH 45 DEGREES CHAMFER OR": "СНИМИТЕ ЗАУСЕНЦЫ НА ОСТРЫХ КРОМКАХ С ФАСКОЙ 45 ГРАДУСОВ ИЛИ",
    "UNLESS GIVEN DIFFERENTLY.": "ЕСЛИ НЕ УКАЗАНО ИНАЧЕ.",
    "Record the repair number onto the documentation which is attached to the part.":
        "Запишите номер ремонта в документацию, прикреплённую к детали.",
    "The plating thickness must not be more": "Толщина покрытия не должна превышать",
    "Examine the edges of sulphamate nickel plate and make sure that they are correctly bonded.":
        "Осмотрите кромки сульфаматного никелевого покрытия и убедитесь в их надлежащем сцеплении.",
    "Examine the edges of sulphamate nickel plate to make sure they are properly bonded: use 5 or 10X magnification.":
        "Осмотрите кромки сульфаматного никелевого покрытия для проверки надлежащего сцепления: используйте увеличение 5× или 10×.",
    "The bush flange thickness must be equal within 0,1 mm (0.0039 in) after machining.":
        "Толщина фланца втулки после механической обработки должна быть одинаковой с допуском 0,1 мм (0,0039 дюйма).",
    "IF NOT, IT CAN CAUSE DAMAGE TO THE COMPONENT.": "В ПРОТИВНОМ СЛУЧАЕ КОМПОНЕНТ МОЖЕТ БЫТЬ ПОВРЕЖДЁН.",
    "The breakout load must be between 0,339 and 0,904 N m (3 and 8 lbf in).":
        "Момент страгивания должен быть от 0,339 до 0,904 Н·м (3–8 фунт-дюйм).",
    "The shot peen can extend to area outside limits to between 0,25 and 0,50 mm (0.010 and 0.020 in).":
        "Зона дробеструйной обработки может выходить за пределы на 0,25–0,50 мм (0,010–0,020 дюйма).",
    "Apply protective varnish to the wiring diagram plate (1-110).":
        "Нанесите защитный лак на пластину со схемой электрических соединений (1-110).",
    "-3 DEGREES TO +3 DEGREES SLOT ORIENTATION IMPORTANT": "ОТ -3 ДО +3 ГРАДУСОВ ОРИЕНТАЦИЯ ПАЗА ВАЖНА",
    "APPLY PRIMER PAINT TO BUSH DIA AND FLANGE UNDERSIDE ONLY": "НАНЕСТИ ГРУНТОВОЧНУЮ КРАСКУ ТОЛЬКО НА ДИАМЕТР ВТУЛКИ И НИЖНЮЮ СТОРОНУ ФЛАНЦА",
    "CAUTION: MAKE SURE THAT YOU DO THE SPECIAL DIMENSION CHECK OF THE ROD (17-160): REFER TO CHECK.":
        "ВНИМАНИЕ: УБЕДИТЕСЬ В ВЫПОЛНЕНИИ СПЕЦИАЛЬНОЙ ПРОВЕРКИ РАЗМЕРОВ ШТОКА (17-160): СМ. CHECK.",
    "Make sure that you do not exceed a torque value of 160 N m (118 lbf ft).":
        "Убедитесь, что момент затяжки не превышает 160 Н·м (118 фунт-фут).",
    "Use electrically conducting Mastinox (made from Mastinox D40, Material Ref.":
        "Используйте электропроводящий Mastinox (из Mastinox D40, Материал Поз.",
    "Examine the upper diaphram tube for flaws: refer to PCS-3100, inclusion class 3.":
        "Осмотрите верхнюю диафрагменную трубу на наличие дефектов: см. PCS-3100, класс включений 3.",
    "Refer to PCS-2500: apply primer to the mating surfaces of the bush (15-370) and the upper diaphragm tube (15-390).": "См. PCS-2500: нанесите грунтовку на сопрягаемые поверхности втулки (15-370) и верхней диафрагменной трубы (15-390).",
    "Refer to PCS-2500: apply primer to the mating surfaces of the bush (15-380) and the upper diaphragm tube (15-390).": "См. PCS-2500: нанесите грунтовку на сопрягаемые поверхности втулки (15-380) и верхней диафрагменной трубы (15-390).",
    "Apply a fillet of Sealant around the joints between the bush (15-370) and the upper diaphragm tube (15-390A): refer to PCS-7200 and Figure 717. Make sure that the primer paint is not visible at the joints after you apply the sealant.": "Нанесите галтель из герметика вокруг соединений между втулкой (15-370) и верхней диафрагменной трубой (15-390A): см. PCS-7200 и Рисунок 717. Убедитесь, что грунтовочная краска не видна на соединениях после нанесения герметика.",
    "Apply a fillet of Sealant around the joints between the bush (15-380) and the upper diaphragm tube (15-390A): refer to PCS-7200 and Figure 7006. Make sure that the primer paint is not visible at the joints after you apply the sealant.": "Нанесите галтель из герметика вокруг соединений между втулкой (15-380) и верхней диафрагменной трубой (15-390A): см. PCS-7200 и Рисунок 7006. Убедитесь, что грунтовочная краска не видна на соединениях после нанесения герметика.",
    "Install the locking plates (15-80) in the grooves of upper bearing housing (15-40A) and mark the hole locations. Remove the locking plates (15-80) and drill the holes. The holes must align with the holes in the two piece stop with inserts (15-130): refer to Figure 718.": "Установите стопорные пластины (15-80) в канавки верхнего корпуса подшипника (15-40A) и отметьте расположение отверстий. Снимите стопорные пластины (15-80) и просверлите отверстия. Отверстия должны совпадать с отверстиями в двухкомпонентном упоре со вставками (15-130): см. Рисунок 718.",
    "Insert the upper bearing housing (15-40A) over the upper diaphragm tube (15-390A) and slide it towards the upper end of upper diaphragm tube (15-390A) to facilitate the installation of two piece stop subassembly (15-110).": "Наденьте верхний корпус подшипника (15-40A) на верхнюю диафрагменную трубу (15-390A) и сдвиньте его к верхнему концу верхней диафрагменной трубы (15-390A) для облегчения установки сборки двухкомпонентного упора (15-110).",
    "Use the Bench Clamp MT1025, the Holding Blocks 460006406 and the Torque Reactor 460006407 to install the diaphragm subassembly (15-190A) and the related parts to the diaphragm tube subassembly (15-360A).": "Используйте настольные тиски MT1025, удерживающие блоки 460006406 и реактор момента затяжки 460006407 для установки сборки диафрагмы (15-190A) и связанных деталей в сборку диафрагменной трубы (15-360A).",
    "Use the Bench Clamp MT1025, the Holding Blocks 460006406 and the Torque Reactor 460007278: hold the upper diaphragm tube subassembly (15-360 only). Install the diaphragm subassembly (15-190 only): use the Torque Adapter 460007283 to torque it to between 200 and 300 N m (148 and 221 lbf ft).": "Используйте настольные тиски MT1025, удерживающие блоки 460006406 и реактор момента затяжки 460007278: удерживайте сборку верхней диафрагменной трубы (только 15-360). Установите сборку диафрагмы (только 15-190): используйте адаптер момента затяжки 460007283 и затяните моментом от 200 до 300 Н·м (148–221 фунт-фут).",
    "Install the 2M electrical axle harness (11-50) and the 1M electrical axle harness (11-40): refer to Figure 721.": "Установите электрический жгут оси 2M (11-50) и электрический жгут оси 1M (11-40): см. Рисунок 721.",
    "Use Lockwire to safety the ground stud subassembly (5-390A) to the bolt (5-395): refer to PCS-7610. Put heat shrinkable tubing over the lockwire.": "Законтрите сборку заземляющей шпильки (5-390A) к болту (5-395) контровочной проволокой: см. PCS-7610. Наденьте термоусадочную трубку на контровочную проволоку.",
    "Apply Ardrox AV100D, Material Ref. Item TBA, to the joints between the pin (13-190), main fitting subassembly (20-90) and the upper diaphragm tube (15-390). Make sure that Ardrox is sprayed to overlap the Ardrox sprayed during main fitting assembly at step paragraph A.(53): refer to Figure 701 and 717.": "Нанесите Ardrox AV100D, Материал Поз. TBA, на соединения между штифтом (13-190), сборкой корпуса стойки (20-90) и верхней диафрагменной трубой (15-390). Убедитесь, что Ardrox нанесён с перекрытием Ardrox, нанесённого при сборке корпуса стойки на шаге пункта A.(53): см. Рисунки 701 и 717.",
    "Apply sufficient amount of Ardrox AV100D, Material Ref. Item TBA, through the hole of the pin (13-190) to cover the remaining upper surfaces of the upper diaphragm tube (15-390). Visually examine to make sure that you have applied Ardrox AV100D, Material Ref. Item TBA, to the complete upper surfaces of the upper diaphragm tube (15-390).": "Нанесите достаточное количество Ardrox AV100D, Материал Поз. TBA, через отверстие штифта (13-190) для покрытия оставшихся верхних поверхностей верхней диафрагменной трубы (15-390). Визуально проверьте, что Ardrox AV100D, Материал Поз. TBA, нанесён на все верхние поверхности верхней диафрагменной трубы (15-390).",
    "Install the bushes (16-130) in the gland housing (16-140): refer to M-DLPS1011-14. Ream the internal diameters of the bushes (16-130) to the dimension given in FITS AND CLEARANCES, Figure 815, reference letter A. The centers of the bushes (16-130) must be the same as the center of the bush hole: the tolerance is 0,0200 mm (0.00078 in).": "Установите втулки (16-130) в корпус сальника (16-140): см. M-DLPS1011-14. Развёртайте внутренние диаметры втулок (16-130) до размера, указанного в ПОСАДКАХ И ЗАЗОРАХ, Рисунок 815, буква A. Центры втулок (16-130) должны совпадать с центром отверстия под втулку: допуск составляет 0,0200 мм (0,00078 дюйма).",
    "or 16A-110E) onto the sliding tube subassembly (17-240). and use the Assembly Sleeve 460006405 and install the lower bearing subassembly (16-110D or 16A-110E) to the sliding tube subassembly (17-240).": "или 16A-110E) на сборку скользящей трубы (17-240). и используйте монтажную втулку 460006405 и установите сборку нижнего подшипника (16-110D или 16A-110E) в сборку скользящей трубы (17-240).",
    "The Numerical Index is to help you find part numbers in the Detailed Parts List.": "Числовой указатель служит для поиска номеров деталей в Подробном перечне деталей.",
    "The Shelf Life Limitation (SLL) for the main landing gear leg (1-1) is 120 months.": "Срок хранения (SLL) стойки основного шасси (1-1) составляет 120 месяцев.",
    "Items A and B - The maximum end float at assembly between the upper torque link subassembly (10-170) and the main fitting subassembly (20-90) and (20-100) is 1,156 mm": "Позиции A и B – Максимальный осевой зазор при сборке между сборкой верхнего шлиц-шарнира (10-170) и сборкой корпуса стойки (20-90) и (20-100) составляет 1,156 мм",
    "Items C1 and C2 - There are two standards of main fitting (20-410) and (20-420) in production: any one of the two fits is permitted.": "Позиции C1 и C2 – В производстве применяются два стандарта корпуса стойки (20-410) и (20-420): допускается любая из двух посадок.",
    "Items J1 and J2 - There are two standards of bracket (4-140) and (4-150) in production: any one of the two fits is permitted.": "Позиции J1 и J2 – В производстве применяются два стандарта кронштейна (4-140) и (4-150): допускается любая из двух посадок.",
    "If necessary, strike through the old part number and identify the new lower bearing subassembly with the new inner liner: refer to PCS-6000-07.": "При необходимости зачеркните старый номер детали и идентифицируйте новую сборку нижнего подшипника с новым внутренним вкладышем: см. PCS-6000-07.",
    "In general part numbers are in THE DISASSEMBLY sequence. The parts are indented to show that they are related to the next higher assembly (NHA).": "Как правило, номера деталей приведены в последовательности РАЗБОРКИ. Детали записаны с отступом для обозначения их принадлежности к следующей сборке более высокого уровня (СВУ).",
    "The Part Numbers that are shown (NP) in the Detailed Parts List are non-procurable items. Unless the part has been superseded the next higher assembly must be installed.": "Номера деталей, обозначенные (NP) в Подробном перечне деталей, являются непоставляемыми позициями. Если деталь не заменена, необходимо установить сборку следующего более высокого уровня.",
    "This parts list illustrates and identifies each part of the component in this CMM. Use it to identify parts and to help with provisioning.": "Данный перечень деталей иллюстрирует и идентифицирует каждую деталь компонента в настоящем руководстве CMM. Используйте его для идентификации деталей и при снабжении.",
    "The Total Required column (TTL REQ.) shows the total necessary each time the part number is shown in the Detailed Parts List.": "Колонка «Общее количество» (TTL REQ.) показывает общее необходимое количество каждый раз, когда номер детали указан в Подробном перечне деталей.",
    "Deleted or superseded part numbers have these identifications:": "Удалённые или замененные номера деталей имеют следующие обозначения:",
    "This figure is the maximum allowable clearance between two parts which are assembled together. A minus sign (-) shows an interference fit.": "Данное значение является максимально допустимым зазором между двумя собранными деталями. Знак минус (-) указывает на посадку с натягом.",
    "These figures are the maximum and minimum clearances when two parts are assembled together. A minus sign (-) shows an interference fit.": "Данные значения являются максимальным и минимальным зазорами при сборке двух деталей. Знак минус (-) указывает на посадку с натягом.",
    "These figures are the maximum and minimum dimensions of new parts.\tThe difference between the two dimensions is the tolerance.": "Данные значения являются максимальными и минимальными размерами новых деталей. Разница между двумя размерами является допуском.",
    "These figures are the dimensions to which parts can wear and be used: the difference between the two dimensions must not be more than the allowable clearance.": "Данные значения представляют собой размеры, до которых детали могут изношены и продолжать использоваться: разница между двумя размерами не должна превышать допустимый зазор.",
    "The quantity in the Units per Assembly column is the quantity necessary for the next higher assembly. AR in the Units per Assembly column shows that the quantity of parts to be used is as required. RF in the Units per Assembly column shows that the part is for reference only.": "Количество в колонке «Количество на сборку» — это количество, необходимое для сборки следующего более высокого уровня. AR в колонке «Количество на сборку» указывает, что количество используемых деталей — по необходимости. RF в колонке «Количество на сборку» означает, что деталь приведена только для справки.",
    "The Effectivity Code (EFF. CODE) agrees with that of the next higher assembly. The effectivity code also shows if subassemblies and details are applicable to their next higher assembly or subassembly. When an item is applicable to all units the Effectivity Code column will be empty. The effectivity code usage is specific to the IPL figure to which it applies.":
        "Код применяемости (EFF. CODE) соответствует коду сборки следующего более высокого уровня. Код применяемости также указывает, применимы ли сборочные единицы и детали к их сборке или сборочной единице следующего более высокого уровня. Если позиция применима ко всем изделиям, колонка кода применяемости будет пустой. Использование кода применяемости характерно для соответствующего рисунка ИПЧ.",
    "NOTE: Below you will find a list of vendor codes associated to this component maintenance manual. For the latest vendor name and address details associated to these codes, please refer to the Safran Landing Systems Technical Publications on-line service - document titled - \ufffdList of Contacts\ufffd.":
        "ПРИМЕЧАНИЕ: Ниже приведён список кодов поставщиков, связанных с данным руководством CMM. Актуальные наименования и адреса поставщиков по этим кодам см. в онлайн-сервисе технических публикаций Safran Landing Systems – документ «Список контактов».",
    # ── Common procedural paragraphs that break in _body() ──
    "Specified Damage and Material Specification": "Указанное повреждение и спецификация материала",
    "(REFER TO SHEET 2)": "(СМ. ЛИСТ 2)",
    "(REFER TO SHEET 3)": "(СМ. ЛИСТ 3)",
    "(REFER TO FIGURE 801)": "(СМ. РИСУНОК 801)",
    "REFER TO TABLE 601": "СМ. ТАБЛИЦУ 601",
    "REFER TO FIGURE 805": "СМ. РИСУНОК 805",
    "Materials": "Материалы",
    "INCLUSIVE CHAMFER": "ВКЛЮЧИТЕЛЬНАЯ ФАСКА",
    "CHROMIUM PLATE DEPOSIT": "СЛОЙ ХРОМОВОГО ПОКРЫТИЯ",
    "x 60 DEGREES INCLUSIVE CHAMFER": "x 60 ГРАДУСОВ ВКЛЮЧИТЕЛЬНАЯ ФАСКА",
    "x 60 DEGREE CHAMFER BOTH BUSHES": "x 60 ГРАДУСОВ ФАСКА ОБЕ ВТУЛКИ",
    "or better.": "или лучше.",
    "or better": "или лучше",
    "Do the above step two more times.": "Повторите вышеуказанный шаг ещё два раза.",
    "Do the above step two more times": "Повторите вышеуказанный шаг ещё два раза",
}

# ──────────────────────────────────────────────────────────────
# 2.  SENTENCE TEMPLATES  (regex + lambda)
# Ordered from most-specific to least-specific.
# ──────────────────────────────────────────────────────────────

def _ref_tail(s):
    """'refer to PCS-2101 and Figure 602.' → 'см. PCS-2101 и Рисунок 602.'"""
    s = s.strip()
    s = re.sub(r'[Rr]efer to ', 'см. ', s)
    s = re.sub(r'[Ff]igures?\s+(\d+)\s+to\s+(\d+)', r'Рисунки \1–\2', s)
    s = re.sub(r'[Ff]igures?\s+(\d+)\s+and\s+(\d+)', r'Рисунки \1 и \2', s)
    s = re.sub(r'[Ff]igure\s+(\d+)', r'Рисунок \1', s)
    s = re.sub(r'\band\b', 'и', s)
    s = re.sub(r'\bor\b', 'или', s)
    return s


def _inc(s):
    """'inclusion Class 3' → 'класс включений 3'"""
    return re.sub(r'inclusion [Cc]lass\s+(\w+)', r'класс включений \1', s)


def _body(s):
    """Generic body translate — replaces known EN words with Russian."""
    s = s.strip()
    # ── Component name phrases (MUST be before individual words) ──
    s = re.sub(r'\bmain fitting subassembly\b', 'сборка корпуса стойки', s, flags=re.I)
    s = re.sub(r'\bmain fitting\b', 'корпус стойки', s, flags=re.I)
    s = re.sub(r'\bsliding tube subassembly\b', 'сборка скользящей трубы', s, flags=re.I)
    s = re.sub(r'\bsliding tube\b', 'скользящая труба', s, flags=re.I)
    s = re.sub(r'\bupper torque link\b', 'верхний шлиц-шарнир', s, flags=re.I)
    s = re.sub(r'\blower torque link\b', 'нижний шлиц-шарнир', s, flags=re.I)
    s = re.sub(r'\btorque link\b', 'шлиц-шарнир', s, flags=re.I)
    s = re.sub(r'\bupper slave link\b', 'верхнее ведомое звено', s, flags=re.I)
    s = re.sub(r'\blower slave link\b', 'нижнее ведомое звено', s, flags=re.I)
    s = re.sub(r'\bslave link\b', 'ведомое звено', s, flags=re.I)
    s = re.sub(r'\bupper stay\b', 'верхнее звено', s, flags=re.I)
    s = re.sub(r'\blower stay\b', 'нижнее звено', s, flags=re.I)
    s = re.sub(r'\block stay cardan\b', 'кардан фиксирующего звена', s, flags=re.I)
    s = re.sub(r'\blocking stay\b', 'фиксирующее звено', s, flags=re.I)
    s = re.sub(r'\block stay\b', 'фиксирующее звено', s, flags=re.I)
    s = re.sub(r'\bupper diaphragm tube\b', 'верхняя диафрагменная труба', s, flags=re.I)
    s = re.sub(r'\blower diaphragm tube\b', 'нижняя диафрагменная труба', s, flags=re.I)
    s = re.sub(r'\bdiaphragm tube\b', 'диафрагменная труба', s, flags=re.I)
    s = re.sub(r'\bupper bearing housing\b', 'верхний корпус подшипника', s, flags=re.I)
    s = re.sub(r'\blower bearing housing\b', 'нижний корпус подшипника', s, flags=re.I)
    s = re.sub(r'\bbearing housing\b', 'корпус подшипника', s, flags=re.I)
    s = re.sub(r'\bupper pivot bracket\b', 'верхний кронштейн оси', s, flags=re.I)
    s = re.sub(r'\blower pivot bracket\b', 'нижний кронштейн оси', s, flags=re.I)
    s = re.sub(r'\bpivot bracket\b', 'кронштейн оси', s, flags=re.I)
    s = re.sub(r'\bgland housing\b', 'корпус сальника', s, flags=re.I)
    s = re.sub(r'\bgland nut\b', 'сальниковая гайка', s, flags=re.I)
    s = re.sub(r'\bshock absorber subassembly\b', 'сборка амортизатора', s, flags=re.I)
    s = re.sub(r'\bshock absorber\b', 'амортизатор', s, flags=re.I)
    s = re.sub(r'\bretaining pins?\b', 'фиксирующий штифт', s, flags=re.I)
    s = re.sub(r'\bcharging valves?\b', 'зарядный клапан', s, flags=re.I)
    s = re.sub(r'\bwiper rings?\b', 'грязесъёмное кольцо', s, flags=re.I)
    s = re.sub(r'\bwiper seals?\b', 'грязесъёмное уплотнение', s, flags=re.I)
    s = re.sub(r'\bground studs?\b', 'шпилька заземления', s, flags=re.I)
    s = re.sub(r'\bstatic discharge connectors?\b', 'разъём снятия статического заряда', s, flags=re.I)
    s = re.sub(r'\buplock hooks?\b', 'крюк верхнего замка', s, flags=re.I)
    s = re.sub(r'\buplock\b', 'верхний замок', s, flags=re.I)
    s = re.sub(r'\banti-corrosion compounds?\b', 'антикоррозионный состав', s, flags=re.I)
    s = re.sub(r'\banti-corrosion\b', 'антикоррозионный', s, flags=re.I)
    s = re.sub(r'\bgrooved spherical bearings?\b', 'подшипник сферический с канавкой', s, flags=re.I)
    s = re.sub(r'\bself lubricating bearings?\b', 'самосмазывающийся подшипник', s, flags=re.I)
    s = re.sub(r'\bspherical bearings?\b', 'сферический подшипник', s, flags=re.I)
    s = re.sub(r'\bstaked bearing assemblyi?e?s?\b', 'узел подшипника с развальцовкой', s, flags=re.I)
    s = re.sub(r'\bbearing assembly\b', 'узел подшипника', s, flags=re.I)
    s = re.sub(r'\bnut assembly\b', 'узел гайки', s, flags=re.I)
    s = re.sub(r'\blaminated shims?\b', 'пакетная прокладка', s, flags=re.I)
    s = re.sub(r'\bbacking plates?\b', 'подкладная пластина', s, flags=re.I)
    s = re.sub(r'\btab washer(?:s)?\b', 'стопорная шайба', s, flags=re.I)
    s = re.sub(r'\bpin spanner\b', 'штифтовой ключ', s, flags=re.I)
    s = re.sub(r'\bwire locking\b', 'контровка проволокой', s, flags=re.I)
    s = re.sub(r'\bwire rope\b', 'трос', s, flags=re.I)
    # Coatings first (longer phrases)
    s = re.sub(r'\bzinc nickel plat(?:e|ing|ed)\b', 'цинко-никелевое покрытие', s, flags=re.I)
    s = re.sub(r'\bcadmium plat(?:e|ing|ed)\b', 'кадмиевое покрытие', s, flags=re.I)
    s = re.sub(r'\bchromium plat(?:e|ing|ed)\b', 'хромовое покрытие', s, flags=re.I)
    s = re.sub(r'\bchrome plat(?:e|ing|ed)\b', 'хромовое покрытие', s, flags=re.I)
    s = re.sub(r'\bprimer paint\b', 'грунтовочная краска', s, flags=re.I)
    # Areas
    s = re.sub(r'\bmachined (?:area|areas)\b', 'обработанные области', s, flags=re.I)
    s = re.sub(r'\breworked (?:area|areas)\b', 'переработанные области', s, flags=re.I)
    s = re.sub(r'\brepaired (?:area|areas)\b', 'отремонтированные области', s, flags=re.I)
    # Parts
    s = re.sub(r'\brepair sleeves?\b', 'ремонтная вставка', s, flags=re.I)
    s = re.sub(r'\brepair bushes?\b', 'ремонтная втулка', s, flags=re.I)
    s = re.sub(r'\bsleeve(?:s)?\b', 'вставка', s, flags=re.I)
    s = re.sub(r'\bbushes?\b', 'втулка', s, flags=re.I)
    s = re.sub(r'\bbushings?\b', 'втулка', s, flags=re.I)
    s = re.sub(r'\bpin(?:s)?\b', 'штифт', s, flags=re.I)
    s = re.sub(r'\bbolt(?:s)?\b', 'болт', s, flags=re.I)
    s = re.sub(r'\bnut(?:s)?\b', 'гайка', s, flags=re.I)
    s = re.sub(r'\bwasher(?:s)?\b', 'шайба', s, flags=re.I)
    s = re.sub(r'\bspacer(?:s)?\b', 'проставка', s, flags=re.I)
    s = re.sub(r'\bbearing(?:s)?\b', 'подшипник', s, flags=re.I)
    s = re.sub(r'\bseal(?:s)?\b', 'уплотнение', s, flags=re.I)
    s = re.sub(r'\bring(?:s)?\b', 'кольцо', s, flags=re.I)
    s = re.sub(r'\bspring(?:s)?\b', 'пружина', s, flags=re.I)
    s = re.sub(r'\bvalve(?:s)?\b', 'клапан', s, flags=re.I)
    s = re.sub(r'\bhousing\b', 'корпус', s, flags=re.I)
    # Structural features
    s = re.sub(r'\bbore(?:s)?\b', 'отверстие', s, flags=re.I)
    s = re.sub(r'\bdiameter(?:s)?\b', 'диаметр', s, flags=re.I)
    s = re.sub(r'\bface(?:s)?\b', 'торец', s, flags=re.I)
    s = re.sub(r'\bchamfer(?:s)?\b', 'фаска', s, flags=re.I)
    s = re.sub(r'\bradii\b', 'радиусы', s, flags=re.I)
    s = re.sub(r'\bradius\b', 'радиус', s, flags=re.I)
    s = re.sub(r'\blug(?:s)?\b', 'проушина', s, flags=re.I)
    s = re.sub(r'\bflange(?:s)?\b', 'фланец', s, flags=re.I)
    s = re.sub(r'\bthread(?:s)?\b', 'резьба', s, flags=re.I)
    s = re.sub(r'\bhole(?:s)?\b', 'отверстие', s, flags=re.I)
    s = re.sub(r'\bshank(?:s)?\b', 'хвостовик', s, flags=re.I)
    s = re.sub(r'\bhead(?:s)?\b', 'головка', s, flags=re.I)
    s = re.sub(r'\bjoint(?:s)?\b', 'соединение', s, flags=re.I)
    s = re.sub(r'\bslot(?:s)?\b', 'паз', s, flags=re.I)
    s = re.sub(r'\bgroove(?:s)?\b', 'канавка', s, flags=re.I)
    s = re.sub(r'\bspigot\b', 'шейка', s, flags=re.I)
    s = re.sub(r'\bcross hole(?:s)?\b', 'поперечное отверстие', s, flags=re.I)
    # Materials/substances
    s = re.sub(r'\bsealant\b', 'герметик', s, flags=re.I)
    s = re.sub(r'\blockwire\b', 'контровочная проволока', s, flags=re.I)
    s = re.sub(r'\bgrease\b', 'смазка', s, flags=re.I)
    s = re.sub(r'\bpaint\b', 'краска', s, flags=re.I)
    s = re.sub(r'\bprimer\b', 'грунтовка', s, flags=re.I)
    s = re.sub(r'\bloctite\b', 'Локтайт', s, flags=re.I)
    # Damage types
    s = re.sub(r'\bcorrosion\b', 'коррозия', s, flags=re.I)
    s = re.sub(r'\bdamage\b', 'повреждение', s, flags=re.I)
    s = re.sub(r'\bwear\b', 'износ', s, flags=re.I)
    s = re.sub(r'\bflaw(?:s)?\b', 'дефект', s, flags=re.I)
    # Dimensional/quality terms
    s = re.sub(r'\bthickness\b', 'толщина', s, flags=re.I)
    s = re.sub(r'\bdimension(?:s)?\b', 'размер', s, flags=re.I)
    s = re.sub(r'\btolerance(?:s)?\b', 'допуск', s, flags=re.I)
    s = re.sub(r'\bsurface finish\b', 'чистота поверхности', s, flags=re.I)
    s = re.sub(r'\bmaterial\b', 'материал', s, flags=re.I)
    # Actions/adjectives
    s = re.sub(r'\binstalled\b', 'установленный', s, flags=re.I)
    s = re.sub(r'\bmachined\b', 'обработанный', s, flags=re.I)
    s = re.sub(r'\breworked\b', 'переработанный', s, flags=re.I)
    s = re.sub(r'\badjacent\b', 'смежный', s, flags=re.I)
    s = re.sub(r'\bnecessary\b', 'необходимый', s, flags=re.I)
    s = re.sub(r'\bapplicable\b', 'соответствующий', s, flags=re.I)
    s = re.sub(r'\bcorrect\b', 'правильный', s, flags=re.I)
    s = re.sub(r'\bcorrectly\b', 'правильно', s, flags=re.I)
    s = re.sub(r'\bsufficiently\b', 'достаточно', s, flags=re.I)
    s = re.sub(r'\bminimum\b', 'минимальный', s, flags=re.I)
    s = re.sub(r'\bmaximum\b', 'максимальный', s, flags=re.I)
    s = re.sub(r'\binternal\b', 'внутренний', s, flags=re.I)
    s = re.sub(r'\bexternal\b', 'внешний', s, flags=re.I)
    s = re.sub(r'\bouter\b', 'наружный', s, flags=re.I)
    s = re.sub(r'\binner\b', 'внутренний', s, flags=re.I)
    s = re.sub(r'\blocally\b', 'локально', s, flags=re.I)
    # New coatings
    s = re.sub(r'\bsulphamate nickel plat(?:e|ing|ed)\b', 'сульфаматное никелевое покрытие', s, flags=re.I)
    s = re.sub(r'\bsulphamate nickel\b', 'сульфаматный никель', s, flags=re.I)
    s = re.sub(r'\bnickel plate(?:d)?\b', 'никелевое покрытие', s, flags=re.I)
    s = re.sub(r'\bnickel plating\b', 'никелевое покрытие', s, flags=re.I)
    # More verbs
    s = re.sub(r'\bGrit blast\b', 'Пескоструйная обработка', s, flags=re.I)
    s = re.sub(r'\bgrit blast(?:ed|ing)?\b', 'пескоструйная обработка', s, flags=re.I)
    s = re.sub(r'\bgrind\b', 'шлифовать', s, flags=re.I)
    s = re.sub(r'\bhone\b', 'хонинговать', s, flags=re.I)
    s = re.sub(r'\bream\b', 'развертывать', s, flags=re.I)
    s = re.sub(r'\bdeburr\b', 'снять заусенцы', s, flags=re.I)
    s = re.sub(r'\bspot face\b', 'подрезка торца', s, flags=re.I)
    s = re.sub(r'\bspotface\b', 'подрезка торца', s, flags=re.I)
    s = re.sub(r'\bmask\b', 'маскировать', s, flags=re.I)
    s = re.sub(r'\bmasked\b', 'замаскированный', s, flags=re.I)
    # More terms
    s = re.sub(r'\bwithin the dimensions shown\b', 'в пределах указанных размеров', s, flags=re.I)
    s = re.sub(r'\bwithin the dimensions\b', 'в пределах размеров', s, flags=re.I)
    s = re.sub(r'\bwithin the limits\b', 'в пределах допустимых значений', s, flags=re.I)
    s = re.sub(r'\bto remove the minimum amount of material\b', 'для снятия минимального количества материала', s, flags=re.I)
    s = re.sub(r'\bminimum amount of material\b', 'минимальное количество материала', s, flags=re.I)
    s = re.sub(r'\bbase metal\b', 'основной металл', s, flags=re.I)
    s = re.sub(r'\bcorrect centres\b', 'правильными центрами', s, flags=re.I)
    s = re.sub(r'\bcorrect center\b', 'правильный центр', s, flags=re.I)
    s = re.sub(r'\bmagnification\b', 'увеличение', s, flags=re.I)
    s = re.sub(r'\bcut-out\b', 'вырез', s, flags=re.I)
    s = re.sub(r'\balignment\b', 'совмещение', s, flags=re.I)
    s = re.sub(r'\bintensity\b', 'интенсивность', s, flags=re.I)
    s = re.sub(r'\bde-embrittlement\b', 'устранение хрупкости', s, flags=re.I)
    s = re.sub(r'\boversize\b', 'увеличенный размер', s, flags=re.I)
    s = re.sub(r'\bconducting\b', 'проводящий', s, flags=re.I)
    s = re.sub(r'\belectrically\b', 'электрически', s, flags=re.I)
    s = re.sub(r'\bsolution\b', 'раствор', s, flags=re.I)
    s = re.sub(r'\boperation\b', 'операция', s, flags=re.I)
    s = re.sub(r'\bprofile\b', 'профиль', s, flags=re.I)
    s = re.sub(r'\bblend\b', 'сгладить', s, flags=re.I)
    s = re.sub(r'\bprimary\b', 'первичный', s, flags=re.I)
    s = re.sub(r'\bsufficient\b', 'достаточный', s, flags=re.I)
    s = re.sub(r'\bincluding\b', 'включая', s, flags=re.I)
    s = re.sub(r'\bexcept\b', 'кроме', s, flags=re.I)
    s = re.sub(r'\badjacent to\b', 'рядом с', s, flags=re.I)
    s = re.sub(r'\bflush with\b', 'заподлицо с', s, flags=re.I)
    s = re.sub(r'\bwill touch\b', 'будет касаться', s, flags=re.I)
    s = re.sub(r'\bhas been\b', 'был', s, flags=re.I)
    s = re.sub(r'\bless than\b', 'менее', s, flags=re.I)
    s = re.sub(r'\bmore than\b', 'более', s, flags=re.I)
    s = re.sub(r'\bnot more than\b', 'не более', s, flags=re.I)
    s = re.sub(r'\bnot less than\b', 'не менее', s, flags=re.I)
    s = re.sub(r'\bafter machining\b', 'после механической обработки', s, flags=re.I)
    s = re.sub(r'\bafter installation\b', 'после установки', s, flags=re.I)
    s = re.sub(r'\bbefore installation\b', 'до установки', s, flags=re.I)
    s = re.sub(r'\bto get\b', 'для получения', s, flags=re.I)
    s = re.sub(r'\bto make\b', 'для обеспечения', s, flags=re.I)
    s = re.sub(r'\bin place of\b', 'вместо', s, flags=re.I)
    s = re.sub(r'\bone or more\b', 'одного или нескольких', s, flags=re.I)
    s = re.sub(r'\bdo not reduce\b', 'не уменьшайте', s, flags=re.I)
    s = re.sub(r'\bdo not machine\b', 'не обрабатывайте', s, flags=re.I)
    # Matl Ref
    s = re.sub(r'\bMaterial Ref\.\s*Item\b', 'Материал Поз.', s, flags=re.I)
    s = re.sub(r'\bMaterial Ref\b', 'Материал Поз.', s, flags=re.I)
    s = re.sub(r'\bRef\. Item\b', 'Поз.', s, flags=re.I)
    s = re.sub(r'\bonly\b', 'только', s, flags=re.I)
    s = re.sub(r'\ball over\b', 'по всей поверхности', s, flags=re.I)
    s = re.sub(r'\bflush\b', 'заподлицо', s, flags=re.I)
    s = re.sub(r'\bwet\b', 'влажный', s, flags=re.I)
    s = re.sub(r'\boptionally\b', 'по желанию', s, flags=re.I)
    s = re.sub(r'\bequivalents?\b', 'эквивалент', s, flags=re.I)
    s = re.sub(r'\bpermitted\b', 'допускается', s, flags=re.I)
    s = re.sub(r'\brequired\b', 'требуется', s, flags=re.I)
    s = re.sub(r'\bshown\b', 'показанный', s, flags=re.I)
    s = re.sub(r'\bindicated\b', 'указанный', s, flags=re.I)
    s = re.sub(r'\bcalculated\b', 'рассчитанный', s, flags=re.I)
    s = re.sub(r'\bformula\b', 'формула', s, flags=re.I)
    s = re.sub(r'\bspecification(?:s)?\b', 'спецификация', s, flags=re.I)
    # Prepositions / connectors (translate common ones in context)
    s = re.sub(r'\bbetween\b', 'от', s, flags=re.I)
    s = re.sub(r'\bfollowing\b', 'следующий', s, flags=re.I)
    s = re.sub(r'\bthat follow\b', 'перечисленные ниже', s, flags=re.I)
    s = re.sub(r'\bmust be\b', 'должна быть', s, flags=re.I)
    s = re.sub(r'\bdo not\b', 'не', s, flags=re.I)
    s = re.sub(r'\bDo not\b', 'Не', s)
    # Figures
    s = re.sub(r'\bFigures?\s+(\d+)\s+to\s+(\d+)\b', r'Рисунки \1–\2', s)
    s = re.sub(r'\bFigures?\s+(\d+)\s+and\s+(\d+)\b', r'Рисунки \1 и \2', s)
    s = re.sub(r'\bFigure\s+(\d+)\b', r'Рисунок \1', s)
    # Refer to
    s = re.sub(r'[Rr]efer to ', 'см. ', s)
    # SB references
    s = re.sub(r'\bPre SB\b', 'до SB', s, flags=re.I)
    s = re.sub(r'\bPost SB\b', 'после SB', s, flags=re.I)
    s = re.sub(r'\bPRE SB\b', 'ДО SB', s)
    s = re.sub(r'\bPOST SB\b', 'ПОСЛЕ SB', s)
    s = re.sub(r'\bInclusion [Cc]lass\b', 'класс включений', s, flags=re.I)
    # ── Common phrases that appear as trailing sentences ──
    s = re.sub(r'Make the surface finish\s+([\d,\.]+)\s+micrometers?\s*\((\d+)\s*micro[- ]inches?\)', r'Чистота поверхности должна составлять \1 мкм (\2 микродюймов)', s)
    s = re.sub(r'The cadmium plate thickness must be between\s+', 'Толщина кадмиевого покрытия должна быть от ', s)
    s = re.sub(r'The (?:zinc nickel|zinc-nickel) plate thickness must be between\s+', 'Толщина цинко-никелевого покрытия должна быть от ', s)
    s = re.sub(r'The plating thickness must be between\s+', 'Толщина покрытия должна быть от ', s)
    s = re.sub(r'The sulphamate nickel plate thickness must be sufficient to get the correct dimensions after machining',
               'Толщина сульфаматного никелевого покрытия должна быть достаточной для получения правильных размеров после механической обработки', s)
    s = re.sub(r'\bNo bare metal is permitted\b', 'Обнажение основного металла не допускается', s, flags=re.I)
    s = re.sub(r'\bStress relieve\b', 'Снимите напряжения', s)
    s = re.sub(r'\bbare metal\b', 'основной металл', s, flags=re.I)
    s = re.sub(r'\bmicrometers?\b', 'мкм', s, flags=re.I)
    s = re.sub(r'\bmicro[- ]inches?\b', 'микродюймов', s, flags=re.I)
    # ── Assembly procedure words ──
    s = re.sub(r'\bspotfaces?\b', 'подрезки торцов', s, flags=re.I)
    s = re.sub(r'\bpiston\b', 'поршень', s, flags=re.I)
    s = re.sub(r'\bjacking dome\b', 'домкратный купол', s, flags=re.I)
    s = re.sub(r'\bbaffle\b', 'перегородка', s, flags=re.I)
    s = re.sub(r'\bdiaphragm\b', 'диафрагма', s, flags=re.I)
    s = re.sub(r'\borifice\b', 'дроссель', s, flags=re.I)
    s = re.sub(r'\bcompression\b', 'сжатия', s, flags=re.I)
    s = re.sub(r'\brecoil\b', 'отдачи', s, flags=re.I)
    s = re.sub(r'\bcap screws?\b', 'болт', s, flags=re.I)
    s = re.sub(r'\block plates?\b', 'стопорная пластина', s, flags=re.I)
    s = re.sub(r'\bsplit pins?\b', 'шплинт', s, flags=re.I)
    s = re.sub(r'\btab washers?\b', 'стопорная шайба', s, flags=re.I)
    s = re.sub(r'\blocking washers?\b', 'стопорная шайба', s, flags=re.I)
    s = re.sub(r'\blocking nuts?\b', 'самоконтрящаяся гайка', s, flags=re.I)
    s = re.sub(r'\bidentification washers?\b', 'идентификационная шайба', s, flags=re.I)
    s = re.sub(r'\blubrication fittings?\b', 'смазочный ниппель', s, flags=re.I)
    s = re.sub(r'\blubrication adapters?\b', 'переходник смазки', s, flags=re.I)
    s = re.sub(r'\bpress pad\b', 'прессовая подушка', s, flags=re.I)
    s = re.sub(r'\bthe drift\b', 'оправка', s, flags=re.I)
    s = re.sub(r'\bcentres?\b', 'центр', s, flags=re.I)
    s = re.sub(r'\bcenter\b', 'центр', s, flags=re.I)
    s = re.sub(r'\brod\b', 'шток', s, flags=re.I)
    s = re.sub(r'\blevel tube\b', 'трубка уровня', s, flags=re.I)
    s = re.sub(r'\bbacking rings?\b', 'опорное кольцо', s, flags=re.I)
    s = re.sub(r'\bsealing rings?\b', 'уплотнительное кольцо', s, flags=re.I)
    s = re.sub(r'\btransfer blocks?\b', 'блок передачи', s, flags=re.I)
    s = re.sub(r'\bthreaded inserts?\b', 'резьбовая вставка', s, flags=re.I)
    s = re.sub(r'\blocating pins?\b', 'фиксирующий штифт', s, flags=re.I)
    s = re.sub(r'\bjig\b', 'приспособление', s, flags=re.I)
    s = re.sub(r'\bfixture\b', 'приспособление', s, flags=re.I)
    # ── Verbs commonly found in procedures ──
    s = re.sub(r'\bcarefully\b', 'аккуратно', s, flags=re.I)
    s = re.sub(r'\binsert\b', 'вставьте', s, flags=re.I)
    s = re.sub(r'\bpush\b', 'вдавите', s, flags=re.I)
    s = re.sub(r'\bpull\b', 'извлеките', s, flags=re.I)
    s = re.sub(r'\bturn\b', 'поверните', s, flags=re.I)
    s = re.sub(r'\brotate\b', 'поверните', s, flags=re.I)
    s = re.sub(r'\bhold\b', 'удерживайте', s, flags=re.I)
    s = re.sub(r'\bkeep\b', 'удерживайте', s, flags=re.I)
    s = re.sub(r'\bposition\b', 'положение', s, flags=re.I)
    s = re.sub(r'\bidentify\b', 'идентифицируйте', s, flags=re.I)
    s = re.sub(r'\bstriking?\b', 'зачеркните', s, flags=re.I)
    s = re.sub(r'\bStrike through\b', 'Зачеркните', s)
    s = re.sub(r'\bslide\b', 'сдвиньте', s, flags=re.I)
    s = re.sub(r'\bprotect\b', 'защитите', s, flags=re.I)
    s = re.sub(r'\bprotective\b', 'защитный', s, flags=re.I)
    s = re.sub(r'\bapply\b', 'нанесите', s, flags=re.I)
    s = re.sub(r'\brepeat\b', 'повторите', s, flags=re.I)
    s = re.sub(r'\bstop\b', 'остановитесь', s, flags=re.I)
    s = re.sub(r'\bremove\b', 'снимите', s, flags=re.I)
    s = re.sub(r'\bexamine\b', 'осмотрите', s, flags=re.I)
    s = re.sub(r'\binstall\b', 'установите', s, flags=re.I)
    s = re.sub(r'\bsafety\b', 'законтрите', s, flags=re.I)
    s = re.sub(r'\bmeasure\b', 'измерьте', s, flags=re.I)
    s = re.sub(r'\brecord\b', 'запишите', s, flags=re.I)
    s = re.sub(r'\bcalculate\b', 'рассчитайте', s, flags=re.I)
    s = re.sub(r'\bselect\b', 'выберите', s, flags=re.I)
    s = re.sub(r'\bverify\b', 'проверьте', s, flags=re.I)
    s = re.sub(r'\bcheck\b', 'проверьте', s, flags=re.I)
    s = re.sub(r'\bclean\b', 'очистите', s, flags=re.I)
    s = re.sub(r'\bdrain\b', 'дренажный', s, flags=re.I)
    s = re.sub(r'\balign\b', 'совместите', s, flags=re.I)
    s = re.sub(r'\baligned\b', 'совмещённый', s, flags=re.I)
    s = re.sub(r'\bset\b', 'установите', s, flags=re.I)
    s = re.sub(r'\btest(?:s)?\b', 'испытание', s, flags=re.I)
    s = re.sub(r'\btorque\b', 'момент затяжки', s, flags=re.I)
    s = re.sub(r'\bhand lap\b', 'притрите вручную', s, flags=re.I)
    # ── More component words ──
    s = re.sub(r'\bupper\b', 'верхний', s, flags=re.I)
    s = re.sub(r'\blower\b', 'нижний', s, flags=re.I)
    s = re.sub(r'\bmain\b', 'главный', s, flags=re.I)
    s = re.sub(r'\bfittings?\b', 'фитинг', s, flags=re.I)
    s = re.sub(r'\bbrackets?\b', 'кронштейн', s, flags=re.I)
    s = re.sub(r'\btube\b', 'труба', s, flags=re.I)
    s = re.sub(r'\btubes\b', 'трубы', s, flags=re.I)
    s = re.sub(r'\bcylinder\b', 'цилиндр', s, flags=re.I)
    s = re.sub(r'\blinks?\b', 'звено', s, flags=re.I)
    s = re.sub(r'\bpivot\b', 'ось', s, flags=re.I)
    s = re.sub(r'\bpintle\b', 'шкворень', s, flags=re.I)
    s = re.sub(r'\barms?\b', 'рычаг', s, flags=re.I)
    s = re.sub(r'\bsupport\b', 'опора', s, flags=re.I)
    s = re.sub(r'\bsliding\b', 'скользящий', s, flags=re.I)
    s = re.sub(r'\bslave\b', 'ведомый', s, flags=re.I)
    s = re.sub(r'\bcharging\b', 'зарядный', s, flags=re.I)
    s = re.sub(r'\bstay\b', 'звено', s, flags=re.I)
    s = re.sub(r'\bgland\b', 'сальник', s, flags=re.I)
    s = re.sub(r'\bconnectors?\b', 'разъём', s, flags=re.I)
    s = re.sub(r'\bwire\b', 'проволока', s, flags=re.I)
    s = re.sub(r'\bstuds?\b', 'шпилька', s, flags=re.I)
    s = re.sub(r'\bcaps?\b', 'крышка', s, flags=re.I)
    s = re.sub(r'\bcovers?\b', 'крышка', s, flags=re.I)
    s = re.sub(r'\bplates?\b', 'пластина', s, flags=re.I)
    s = re.sub(r'\btabs?\b', 'лепесток', s, flags=re.I)
    s = re.sub(r'\blaminated\b', 'пакетный', s, flags=re.I)
    s = re.sub(r'\bbacking\b', 'подкладной', s, flags=re.I)
    s = re.sub(r'\bsplit\b', 'разрезной', s, flags=re.I)
    s = re.sub(r'\bretaining\b', 'фиксирующий', s, flags=re.I)
    s = re.sub(r'\blocking\b', 'стопорный', s, flags=re.I)
    s = re.sub(r'\bmating\b', 'сопрягаемый', s, flags=re.I)
    s = re.sub(r'\bends?\b', 'торец', s, flags=re.I)
    s = re.sub(r'\bunderside\b', 'нижняя сторона', s, flags=re.I)
    s = re.sub(r'\boutside\b', 'наружная сторона', s, flags=re.I)
    s = re.sub(r'\binside\b', 'внутри', s, flags=re.I)
    # ── More adjectives/adverbs ──
    s = re.sub(r'\bdamaged\b', 'повреждённый', s, flags=re.I)
    s = re.sub(r'\bworn\b', 'изношенный', s, flags=re.I)
    s = re.sub(r'\bsmallest\b', 'наименьший', s, flags=re.I)
    s = re.sub(r'\blargest\b', 'наибольший', s, flags=re.I)
    s = re.sub(r'\bclosest\b', 'ближайший', s, flags=re.I)
    s = re.sub(r'\boriginal\b', 'оригинальный', s, flags=re.I)
    s = re.sub(r'\bexternally\b', 'снаружи', s, flags=re.I)
    s = re.sub(r'\binternally\b', 'изнутри', s, flags=re.I)
    s = re.sub(r'\bcooling\b', 'охлаждение', s, flags=re.I)
    s = re.sub(r'\bheating\b', 'нагрев', s, flags=re.I)
    s = re.sub(r'\bapproximately\b', 'приблизительно', s, flags=re.I)
    s = re.sub(r'\bmeasured\b', 'измеренный', s, flags=re.I)
    # ── More verbs ──
    s = re.sub(r'\bmake\b', 'обеспечьте', s, flags=re.I)
    s = re.sub(r'\bmust\b', 'должен', s, flags=re.I)
    s = re.sub(r'\bcontinue\b', 'продолжите', s, flags=re.I)
    s = re.sub(r'\bleave\b', 'оставьте', s, flags=re.I)
    s = re.sub(r'\bensure\b', 'обеспечьте', s, flags=re.I)
    s = re.sub(r'\bplace\b', 'поместите', s, flags=re.I)
    s = re.sub(r'\bopen\b', 'откройте', s, flags=re.I)
    s = re.sub(r'\bclose\b', 'закройте', s, flags=re.I)
    s = re.sub(r'\bfill\b', 'заполните', s, flags=re.I)
    s = re.sub(r'\bcut\b', 'отрежьте', s, flags=re.I)
    s = re.sub(r'\bfit\b', 'подогнать', s, flags=re.I)
    s = re.sub(r'\btighten\b', 'затяните', s, flags=re.I)
    s = re.sub(r'\bloosen\b', 'ослабьте', s, flags=re.I)
    s = re.sub(r'\bpress\b', 'запрессуйте', s, flags=re.I)
    s = re.sub(r'\bmarke?d?\b', 'отметьте', s, flags=re.I)
    s = re.sub(r'\bcoat(?:ed)?\b', 'покройте', s, flags=re.I)
    # ── More nouns ──
    s = re.sub(r'\brunout\b', 'биение', s, flags=re.I)
    s = re.sub(r'\bclearances?\b', 'зазор', s, flags=re.I)
    s = re.sub(r'\blabels?\b', 'бирка', s, flags=re.I)
    s = re.sub(r'\baxis\b', 'ось', s, flags=re.I)
    s = re.sub(r'\bweight\b', 'масса', s, flags=re.I)
    s = re.sub(r'\bload\b', 'нагрузка', s, flags=re.I)
    s = re.sub(r'\bforce\b', 'сила', s, flags=re.I)
    s = re.sub(r'\bpressure\b', 'давление', s, flags=re.I)
    s = re.sub(r'\btemperature\b', 'температура', s, flags=re.I)
    s = re.sub(r'\bvendor\b', 'поставщик', s, flags=re.I)
    s = re.sub(r'\bcodes?\b', 'код', s, flags=re.I)
    s = re.sub(r'\bmanual\b', 'руководство', s, flags=re.I)
    s = re.sub(r'\bcomponents?\b', 'компонент', s, flags=re.I)
    s = re.sub(r'\bmaintenance\b', 'техническое обслуживание', s, flags=re.I)
    s = re.sub(r'\blist\b', 'перечень', s, flags=re.I)
    s = re.sub(r'\blatest\b', 'последний', s, flags=re.I)
    s = re.sub(r'\bassociated\b', 'связанный', s, flags=re.I)
    # ── More prepositions/connectors ──
    s = re.sub(r'\bbut\b', 'но', s, flags=re.I)
    s = re.sub(r'\bacross\b', 'через', s, flags=re.I)
    s = re.sub(r'\balong\b', 'вдоль', s, flags=re.I)
    s = re.sub(r'\baround\b', 'вокруг', s, flags=re.I)
    s = re.sub(r'\babove\b', 'выше', s, flags=re.I)
    s = re.sub(r'\bbelow\b', 'ниже', s, flags=re.I)
    s = re.sub(r'\bover\b', 'по', s, flags=re.I)
    s = re.sub(r'\bunder\b', 'под', s, flags=re.I)
    s = re.sub(r'\bbetter\b', 'лучше', s, flags=re.I)
    s = re.sub(r'\bgiven\b', 'указанным', s, flags=re.I)
    s = re.sub(r'\breference\b', 'ссылка', s, flags=re.I)
    s = re.sub(r'\bletter\b', 'буква', s, flags=re.I)
    s = re.sub(r'\bthen\b', 'затем', s, flags=re.I)
    s = re.sub(r'\balso\b', 'также', s, flags=re.I)
    s = re.sub(r'\bonly\b', 'только', s, flags=re.I)
    s = re.sub(r'\byou\b', '', s, flags=re.I)
    s = re.sub(r'\bcan\b', 'можно', s, flags=re.I)
    s = re.sub(r'\bwill\b', '', s, flags=re.I)
    s = re.sub(r'\bshould\b', 'следует', s, flags=re.I)
    s = re.sub(r'\bshall\b', 'следует', s, flags=re.I)
    s = re.sub(r'\binches\b', 'дюймов', s, flags=re.I)
    s = re.sub(r'\btwo\b', 'два', s, flags=re.I)
    s = re.sub(r'\bthree\b', 'три', s, flags=re.I)
    s = re.sub(r'\bfour\b', 'четыре', s, flags=re.I)
    s = re.sub(r'\bfive\b', 'пять', s, flags=re.I)
    s = re.sub(r'\bone\b', 'один', s, flags=re.I)
    s = re.sub(r'\bsee\b', 'см.', s, flags=re.I)
    s = re.sub(r'\buse\b', 'используйте', s, flags=re.I)
    s = re.sub(r'\brelated\b', 'связанный', s, flags=re.I)
    s = re.sub(r'\bapprox\b', 'приблиз.', s, flags=re.I)
    s = re.sub(r'\bFITS AND CLEARANCES\b', 'ПОСАДКИ И ЗАЗОРЫ', s)
    s = re.sub(r'\bFits and Clearances\b', 'Посадки и зазоры', s)
    # ── Common connectors & articles (ALL with re.I) ──
    s = re.sub(r'\bthe\s+', '', s, flags=re.I)
    s = re.sub(r'\band/or\b', 'и/или', s, flags=re.I)
    s = re.sub(r'\band\b', 'и', s, flags=re.I)
    s = re.sub(r'\bor\b', 'или', s, flags=re.I)
    s = re.sub(r'\bof\b', '', s, flags=re.I)
    s = re.sub(r'\bto\b', 'до', s, flags=re.I)
    s = re.sub(r'\bfor\b', 'для', s, flags=re.I)
    s = re.sub(r'\bfrom\b', 'с', s, flags=re.I)
    s = re.sub(r'\binto\b', 'в', s, flags=re.I)
    s = re.sub(r'\bin\b', 'в', s, flags=re.I)
    s = re.sub(r'\bon\b', 'на', s, flags=re.I)
    s = re.sub(r'\bat\b', 'при', s, flags=re.I)
    s = re.sub(r'\bwith\b', 'с', s, flags=re.I)
    s = re.sub(r'\bwithout\b', 'без', s, flags=re.I)
    s = re.sub(r'\buntil\b', 'до тех пор пока', s, flags=re.I)
    s = re.sub(r'\bwhile\b', 'пока', s, flags=re.I)
    s = re.sub(r'\bwhen\b', 'когда', s, flags=re.I)
    s = re.sub(r'\bwhere\b', 'где', s, flags=re.I)
    s = re.sub(r'\bhave\b', '', s, flags=re.I)
    s = re.sub(r'\bhas\b', '', s, flags=re.I)
    s = re.sub(r'\bare\b', '', s, flags=re.I)
    s = re.sub(r'\bis\b', '', s, flags=re.I)
    s = re.sub(r'\bnot\b', 'не', s, flags=re.I)
    s = re.sub(r'\bnew\b', 'новый', s, flags=re.I)
    s = re.sub(r'\bold\b', 'старый', s, flags=re.I)
    s = re.sub(r'\bduring\b', 'при', s, flags=re.I)
    s = re.sub(r'\bafter\b', 'после', s, flags=re.I)
    s = re.sub(r'\bbefore\b', 'до', s, flags=re.I)
    s = re.sub(r'\bthis\b', 'данный', s, flags=re.I)
    s = re.sub(r'\bthat\b', '', s, flags=re.I)
    s = re.sub(r'\bthese\b', 'данные', s, flags=re.I)
    s = re.sub(r'\bthose\b', '', s, flags=re.I)
    s = re.sub(r'\beach\b', 'каждый', s, flags=re.I)
    s = re.sub(r'\bevery\b', 'каждый', s, flags=re.I)
    s = re.sub(r'\ball\b', 'все', s, flags=re.I)
    s = re.sub(r'\bany\b', 'любой', s, flags=re.I)
    s = re.sub(r'\bother\b', 'другой', s, flags=re.I)
    s = re.sub(r'\bboth\b', 'оба', s, flags=re.I)
    s = re.sub(r'\bits\b', 'его', s, flags=re.I)
    s = re.sub(r'\bit\b', 'его', s, flags=re.I)
    s = re.sub(r'\bif\b', 'если', s, flags=re.I)
    s = re.sub(r'\bdo\b', '', s, flags=re.I)
    s = re.sub(r'\bstill\b', 'ещё', s, flags=re.I)
    s = re.sub(r'\bthere\b', '', s, flags=re.I)
    s = re.sub(r'\btheir\b', 'их', s, flags=re.I)
    s = re.sub(r'\bwidth\b', 'ширина', s, flags=re.I)
    s = re.sub(r'\blength\b', 'длина', s, flags=re.I)
    s = re.sub(r'\bdepth\b', 'глубина', s, flags=re.I)
    s = re.sub(r'\bheight\b', 'высота', s, flags=re.I)
    s = re.sub(r'\bgap\b', 'зазор', s, flags=re.I)
    s = re.sub(r'\bhours?\b', 'ч', s, flags=re.I)
    s = re.sub(r'\bpart\b', 'деталь', s, flags=re.I)
    s = re.sub(r'\bparts\b', 'детали', s, flags=re.I)
    s = re.sub(r'\bunit\b', 'узел', s, flags=re.I)
    s = re.sub(r'\bunits\b', 'узлы', s, flags=re.I)
    s = re.sub(r'\bside\b', 'сторона', s, flags=re.I)
    s = re.sub(r'\barea\b', 'зона', s, flags=re.I)
    s = re.sub(r'\bareas\b', 'зоны', s, flags=re.I)
    s = re.sub(r'\bedge(?:s)?\b', 'кромка', s, flags=re.I)
    s = re.sub(r'\bsurface\b', 'поверхность', s, flags=re.I)
    s = re.sub(r'\bfinish\b', 'чистота', s, flags=re.I)
    s = re.sub(r'\bvalue(?:s)?\b', 'значение', s, flags=re.I)
    s = re.sub(r'\blimit(?:s)?\b', 'предел', s, flags=re.I)
    s = re.sub(r'\bsame\b', 'тот же', s, flags=re.I)
    s = re.sub(r'\bconfiguration\b', 'конфигурация', s, flags=re.I)
    s = re.sub(r'\bprocedure\b', 'процедура', s, flags=re.I)
    s = re.sub(r'\bstep\b', 'шаг', s, flags=re.I)
    s = re.sub(r'\bparagraph\b', 'пункт', s, flags=re.I)
    s = re.sub(r'\bpara\b', 'пункт', s, flags=re.I)
    s = re.sub(r'\brepair\b', 'ремонт', s, flags=re.I)
    s = re.sub(r'\bnumber\b', 'номер', s, flags=re.I)
    s = re.sub(r'\bexisting\b', 'существующий', s, flags=re.I)
    s = re.sub(r'\bsubassembly\b', 'сборка', s, flags=re.I)
    s = re.sub(r'\bassembly\b', 'сборка', s, flags=re.I)
    s = re.sub(r'\binstallation\b', 'установка', s, flags=re.I)
    s = re.sub(r'\bmachining\b', 'механическая обработка', s, flags=re.I)
    s = re.sub(r'\bjointing compound\b', 'монтажный состав', s, flags=re.I)
    s = re.sub(r'\bcompound\b', 'состав', s, flags=re.I)
    s = re.sub(r'\bhydraulic fluid\b', 'гидравлическая жидкость', s, flags=re.I)
    s = re.sub(r'\bhydraulic\b', 'гидравлический', s, flags=re.I)
    s = re.sub(r'\belectrical bonding\b', 'электрическое соединение', s, flags=re.I)
    s = re.sub(r'\bresistance\b', 'сопротивление', s, flags=re.I)
    s = re.sub(r'\bmilliohm\b', 'миллиом', s, flags=re.I)
    s = re.sub(r'\bpaint\b', 'краска', s, flags=re.I)
    s = re.sub(r'\badapter\b', 'переходник', s, flags=re.I)
    s = re.sub(r'\bzero\b', 'ноль', s, flags=re.I)
    s = re.sub(r'\bas\b', 'как', s, flags=re.I)
    s = re.sub(r'\bso\b', 'так', s, flags=re.I)
    s = re.sub(r'\bbe\b', '', s, flags=re.I)
    s = re.sub(r'\bby\b', '', s, flags=re.I)
    s = re.sub(r'\bno\b', 'нет', s, flags=re.I)
    s = re.sub(r'\bsure\b', '', s, flags=re.I)
    # Clean up multiple spaces
    s = re.sub(r'  +', ' ', s).strip()
    return s


def _verb_phrase(verb_en, verb_ru, rest):
    """Translate 'Verb the rest...' construction."""
    rest = _body(rest)
    return f"{verb_ru} {rest}"


# Sentence-level template patterns
# Each entry: (compiled_regex, translator_function)
_TEMPLATES = []


def _T(pattern, fn, flags=re.IGNORECASE | re.DOTALL):
    _TEMPLATES.append((re.compile(pattern, flags), fn))


# ─── Machine X to remove damage/wear within dimensions: refer to M-DLPS... ───
_T(
    r'^Machine\s+(the\s+)?(.+?)\s+(?:sufficiently\s+)?to remove the (?:minimum amount of material (?:necessary\s+)?to remove the\s+)?(?:damage|wear|corrosion)(?:\s+or\s+(?:corrosion|wear|damage))?\s+within the dimensions shown(?:\s+and\s+as\s+shown)?\s*(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Расточите {_body((m.group(1) or '') + m.group(2))} для устранения дефектов в пределах указанных размеров{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)
_T(
    r'^Machine\s+(the\s+)?(.+?)\s+(?:sufficiently\s+)?to remove (?:the\s+)?(?:minimum amount of material to remove the\s+)?(?:damage|wear|corrosion)(?:\s+or\s+(?:corrosion|wear|damage))?\s*(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Расточите {_body((m.group(1) or '') + m.group(2))} для устранения дефектов{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)
# ─── Machine X to remove damage ... Do not machine/reduce X more than Y ───
_T(
    r'^Machine\s+(the\s+)?(.+?)\s+to remove the minimum amount of material to remove the (?:wear or damage|damage or wear)\.?\s*Do not\s+(?:machine|reduce)\s+(.+?)\s+(?:more than|to less than)\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Расточите {_body((m.group(1) or '') + m.group(2))} для снятия минимального количества материала с целью устранения износа или повреждения. Не обрабатывайте {_body(m.group(3))} до менее {m.group(4)}{(': ' + _ref_tail('refer to '+m.group(5))) if m.group(5) else ''}."
)

# ─── Apply sealant, Material Ref. Item XXXX, to the joints between X and Y ───
_T(
    r'^Apply sealant,?\s+(?:Material Ref\.?\s*)?(?:Item\s+)?([\w\-]+),?\s+to the joints? between (the\s+)?(.+?) and (?:the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Нанесите герметик, Материал Поз. {m.group(1)}, на соединение между {_body((m.group(2) or '') + m.group(3))} и {_body(m.group(4))}{(': ' + _ref_tail('refer to '+m.group(5))) if m.group(5) else ''}."
)
_T(
    r'^Apply sealant,?\s+(?:Material Ref\.?\s*)?(?:Item\s+)?([\w\-]+),?\s+to (?:the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Нанесите герметик, Материал Поз. {m.group(1)}, на {_body(m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)
_T(
    r'^Apply\s+(a fillet of sealant|a fillet of Sealant)\s+(?:around\s+)?(?:the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Нанесите галтель из герметика на {_body(m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)
_T(
    r'^Apply a (?:line of red silicone anti-tamper sealant|line of [^:,]+sealant) (?:across|around|to) (.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Нанесите полосу силиконового противовскрытного герметика по {_body(m.group(1))}{(': ' + _ref_tail('refer to '+m.group(2))) if m.group(2) else ''}."
)

# ─── Around the joint between X and Y ───
_T(
    r'^Around the joints? between (the\s+)?(.+?) and (?:the\s+)?(.+?)$',
    lambda m: f"Вокруг соединения между {_body((m.group(1) or '') + m.group(2))} и {_body(m.group(3))}"
)
_T(
    r'^Around the joints? between (the\s+)?(.+?),? (?:the\s+)?(.+?) and (?:the\s+)?(.+?)$',
    lambda m: f"Вокруг соединений между {_body((m.group(1) or '') + m.group(2))}, {_body(m.group(3))} и {_body(m.group(4))}"
)

# ─── De-embrittle the reworked areas for N hours at X ───
_T(
    r'^De-embrittle\s+(the\s+)?(.+?)\s+for\s+(\d+)\s+hours?\s+at\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Выполните устранение водородной хрупкости {_body((m.group(1) or '') + m.group(2))} в течение {m.group(3)} ч при {m.group(4)}{(': см. ' + m.group(5)) if m.group(5) else ''}."
)

# ─── Examine the ground chromium plate for flaws ───
_T(
    r'^Examine\s+the\s+ground\s+chromium\s+plate\s+for\s+flaws?\s*(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Осмотрите шлифованное хромовое покрытие на наличие дефектов{(': ' + _ref_tail('refer to '+m.group(1))) if m.group(1) else ''}."
)

# ─── If necessary, hone or hand ream the bore diameter ───
_T(
    r'^If necessary,?\s+hone or hand ream\s+(the\s+)?(.+?)\s+to\s+(?:the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"При необходимости хонингуйте или вручную разверните {_body((m.group(1) or '') + m.group(2))} до {_body(m.group(3))}{(': ' + _ref_tail('refer to '+m.group(4))) if m.group(4) else ''}."
)

# ─── Grit blast ───
_T(
    r'^Grit blast\s+(the\s+)?(.+?)\s+(?:to be\s+)?(?:sulphamate nickel plated)?(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Выполните пескоструйную обработку {_body((m.group(1) or '') + m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── Apply sulphamate nickel plate ───
_T(
    r'^Apply sulphamate nickel plate (?:to\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Нанесите сульфаматное никелевое покрытие на {_body(m.group(1))}{(': ' + _ref_tail('refer to '+m.group(2))) if m.group(2) else ''}."
)

# ─── Machine (do not grind) the sulphamate nickel plate ───
_T(
    r'^Machine \(do not grind\) (?:the\s+)?(.+?) to the dimensions shown(?:\s*:\s*refer to\s+(.+?))?\.?(.*)$',
    lambda m: f"Расточите (не шлифовать) {_body(m.group(1))} до указанных размеров{(': ' + _ref_tail('refer to '+m.group(2))) if m.group(2) else ''}.{m.group(3)}"
)

# ─── After machining X, Y must not be less than Z ───
_T(
    r'^After machining (.+?),\s*(the\s+)?(.+?)\s+must not be less than\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"После обработки {_body(m.group(1))}, {_body((m.group(2) or '') + m.group(3))} не должен быть менее {m.group(4)}{(': ' + _ref_tail('refer to '+m.group(5))) if m.group(5) else ''}."
)

# ─── Use of Xx magnification. Examine ... ───
_T(
    r'^Use of\s+(\d+x.*?)\.\s+(.+?)$',
    lambda m: f"Используйте увеличение {m.group(1)}. {translate_text(m.group(2))}"
)

# ─── Lubricate [part] with [lubricant] ───
_T(
    r'^Lubricate\s+(the\s+)?(.+?)\s+with\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Смажьте {_body((m.group(1) or '') + m.group(2))} {_body(m.group(3))}{(': ' + _ref_tail('refer to '+m.group(4))) if m.group(4) else ''}."
)

# ─── Put [part] in/on/to [place] ───
_T(
    r'^Put\s+(the\s+)?(.+?)\s+(?:in|on|into|to)\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Поместите {_body((m.group(1) or '') + m.group(2))} в {_body((m.group(3) or '') + m.group(4))}{(': ' + _ref_tail('refer to '+m.group(5))) if m.group(5) else ''}."
)

# ─── For [application], do [procedure] ───
_T(
    r'^For (?:the\s+)?(.+?),\s*do\s+(.+?)$',
    lambda m: f"Для {_body(m.group(1))} выполните {_body(m.group(2))}"
)

# ─── [specific] subassembly for Shock Absorber ───
_T(
    r'^(.+?)\s+(?:for\s+)?Shock Absorber Subassembly\s+(.+?)$',
    lambda m: f"{translate_text(m.group(1))} для сборки амортизатора {m.group(2)}"
)

# ─── Sliding tube related entries ───
_T(
    r'^(?:The\s+)?(?:internal|inner)\s+(?:diameter|liner)\s+of\s+(.+?)\s+\((.+?)\)(.*?)$',
    lambda m: f"Внутренний диаметр {_body(m.group(1))} ({m.group(2)}){m.group(3)}"
)
_T(
    r'^(?:The\s+)?(?:external|outer)\s+diameter\s+of\s+(.+?)\s+\((.+?)\)(.*?)$',
    lambda m: f"Наружный диаметр {_body(m.group(1))} ({m.group(2)}){m.group(3)}"
)

# ─── Install the identification washers and lubrication fittings in the retaining pins ───
_T(
    r'^Install the identification washers? \((.+?)\) and the lubrication fittings? \((.+?)\) in the retaining pins? \((.+?)\)\.',
    lambda m: f"Установите идентификационные шайбы ({m.group(1)}) и смазочные ниппели ({m.group(2)}) в фиксирующие штифты ({m.group(3)})."
)

# ─── Apply Molykote to the locations that follow ───
_T(
    r'^Apply Molykote\s+(\d+)\s+to (?:the\s+)?(.+?)(?:\s*\(refer to (.+?)\))?:?$',
    lambda m: f"Нанесите Molykote {m.group(1)} на {_body(m.group(2))}{(' (см. '+m.group(3)+')') if m.group(3) else ''}:"
)
_T(
    r'^Apply Molykote\s+(\d+)\s+to the locations that follow\s*\(refer to (.+?)\):$',
    lambda m: f"Нанесите Molykote {m.group(1)} на следующие зоны (см. {m.group(2)}):"
)

# ─── Machine the ends of the installed repair sleeve(s). Flush with... ───
_T(
    r'^Machine the ends of (?:the\s+)?(.+?)\.\s*Flush with (?:the\s+)?(.+?) and prepare (?:the\s+)?(.+?)\s+as shown(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Обработайте торцы {_body(m.group(1))}. Выровняйте заподлицо с {_body(m.group(2))} и обработайте {_body(m.group(3))} согласно чертежу{(': ' + _ref_tail('refer to '+m.group(4))) if m.group(4) else ''}."
)

# ─── Record the repair number onto documentation ───
_T(
    r'^Record the repair number onto (?:the\s+)?(.+?) which is attached to (?:the\s+)?(.+?)\.\s*(?:Optionally,\s+)?(.+?)$',
    lambda m: f"Запишите номер ремонта в {_body(m.group(1))}, прикреплённую к {_body(m.group(2))}. По желанию, {translate_text(m.group(3))}"
)

# ─── Install repair sleeve/bush flush with X face C flush with D ───
_T(
    r'^Install (?:the\s+)?(.+?)\s+(?:to\s+)?(.+?) with (.+?) flush with (.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Установите {_body(m.group(1))} в {_body(m.group(2))} так, чтобы {_body(m.group(3))} было заподлицо с {_body(m.group(4))}{(': ' + _ref_tail('refer to '+m.group(5))) if m.group(5) else ''}."
)

# ─── Machine the bore and the ends of the repair sleeve(s) to get the dimensions ───
_T(
    r'^Machine (?:the\s+)?(.+?) and (?:the\s+)?(.+?) (?:of\s+)?(?:the\s+)?(.+?) to get (?:the\s+)?(.+?) as shown(?::\s*refer to\s+(.+?))?\.?(.*)$',
    lambda m: f"Расточите {_body(m.group(1))} и {_body(m.group(2))} {_body(m.group(3))} до получения {_body(m.group(4))} согласно чертежу{(': ' + _ref_tail('refer to '+m.group(5))) if m.group(5) else ''}.{translate_text(m.group(6)) if m.group(6) else ''}"
)

# ─── Apply [coating] [to/over] [all over] [location] [but not to ...] ───
_T(
    r'^Apply\s+(cadmium plate|zinc nickel plate|chromium plate|primer paint|paint|sealant|sulphamate nickel plate)\s+((?:all over|externally|locally),?\s*)?(?:but not (?:to\s+)?(?:the\s+)?(.+?),?\s+)?(?:to\s+)?(?:the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Нанесите {_coat_nom(m.group(1))} {({'externally': 'снаружи ', 'externally,': 'снаружи, ', 'locally': 'локально ', 'locally,': 'локально, ', 'all over': 'по всей поверхности ', 'all over,': 'по всей поверхности, '}.get((m.group(2) or '').strip(), m.group(2) or ''))}{('кроме ' + _body(m.group(3)) + ', ') if m.group(3) else ''}на {_body(m.group(4))}{(': ' + _ref_tail('refer to '+m.group(5))) if m.group(5) else ''}."
)

# ─── Apply [coating] [to/over] [all over], except where indicated ───
_T(
    r'^Apply\s+(cadmium plate|zinc nickel plate|chromium plate|primer paint|paint|sealant)\s+(?:all over)?,?\s*except where indicated(?::\s*refer to\s+(.+?))?\.?(.*)$',
    lambda m: f"Нанесите {_coat_nom(m.group(1))} по всей поверхности, кроме указанных мест{(': ' + _ref_tail('refer to '+m.group(2))) if m.group(2) else ''}.{m.group(3)}"
)

# ─── NOTE: This operation includes N hours de-embrittlement ───
_T(
    r'^NOTE: This operation includes\s+(\d+)\s+hours? de-embrittlement at\s+(.+?)\.',
    lambda m: f"ПРИМЕЧАНИЕ: Данная операция включает {m.group(1)} ч устранения водородной хрупкости при {m.group(2)}."
)

# ─── If this repair is applied to diameters ... ───
_T(
    r'^NOTE: If this repair is applied to (.+?), (?:identify|apply|examine)\s+(.+?)$',
    lambda m: f"ПРИМЕЧАНИЕ: Если данный ремонт применяется к {_body(m.group(1))}, {translate_text(m.group(2))}"
)

# ─── Apply [coating] [to/over] [target]. [Additional sentence] ───
# ─── Apply [coating] [to/over/all over/externally] [target] [.qualifier] [refer to ...] ───
_T(
    r'^Apply\s+(cadmium plate|zinc nickel plate|chromium plate|chrome plate|primer paint|paint|sealant|grease|Molykote\s+\d+|loctite grade\s+\d+|a fillet of [Ss]ealant|a fillet of red silicone[^,]*sealant|a fillet of [Ss]ealant|a line of [^:,]+sealant)\s+(.*?)$',
    lambda m: f"Нанесите {_coat_nom(m.group(1))} {_body(m.group(2))}"
)

# ─── The [coating] thickness must be between X: refer to ... ───
_T(
    r'^The\s+(cadmium plate|zinc nickel plate|chromium plate|chrome plate|plating)\s+thickness\s+must be\s+between\s+(.+?):\s*refer to\s+(.+?)\.',
    lambda m: f"Толщина {_coat_gen(m.group(1))} должна быть от {m.group(2)}: {_ref_tail('refer to '+m.group(3))}."
)
_T(
    r'^The\s+(cadmium plate|zinc nickel plate|chromium plate|plating)\s+thickness\s+must be\s+between\s+(.+?)\.',
    lambda m: f"Толщина {_coat_gen(m.group(1))} должна быть от {m.group(2)}."
)

# ─── Machine [part/dim] [to/by] [action] [refer to ...] ───
_T(
    r'^Machine\s+(.+?)\s+(?:sufficiently )?to remove the (?:damage|wear|corrosion|minimum amount of material[^:.]*)(.*?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Расточите {_body(m.group(1))} для устранения дефекта{_body(m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)
_T(
    r'^Machine\s+(the (?:bore|diameter|internal diameter|inside diameter|external diameter|outer diameter|flange|end|chamfer|ends|radii|cross holes|spotface|face)\s+(?:of\s+)?.*?)\s+to\s+(?:the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Расточите {_body(m.group(1))} до {_body(m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)
_T(
    r'^Machine\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Расточите {_body(m.group(1))}{(': ' + _ref_tail('refer to '+m.group(2))) if m.group(2) else ''}."
)

# ─── Examine [part] for flaws/damage ───
_T(
    r'^Examine\s+(the\s+)?(.+?)\s+for\s+(flaws?|damage.*?)[,:]\s*refer to\s+(.+?)(?:,\s*(.+?))?\.',
    lambda m: f"Осмотрите {_body((m.group(1) or '') + m.group(2))} на наличие дефектов: {_ref_tail('refer to ' + m.group(4))}{(', ' + _inc(m.group(5))) if m.group(5) else ''}."
)
_T(
    r'^Examine\s+(the\s+)?(.+?)\s+for\s+(flaws?|damage)(.*)$',
    lambda m: f"Осмотрите {_body((m.group(1) or '') + m.group(2))} на наличие дефектов{_body(m.group(4))}."
)
_T(
    r'^Examine\s+(the\s+)?(.+?)\s+for\s+(cracks?|corrosion)(.*)$',
    lambda m: f"Осмотрите {_body((m.group(1) or '') + m.group(2))} на наличие {'трещин' if 'crack' in m.group(3).lower() else 'коррозии'}{_body(m.group(4))}."
)

# ─── Shot peen ───
_T(
    r'^Shot peen\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Выполните дробеструйную обработку {_body(m.group(1))}{(': ' + _ref_tail('refer to '+m.group(2))) if m.group(2) else ''}."
)

# ─── Install [part] [using/with/to/...] [refer to ...] ───
_T(
    r'^Install\s+(the\s+)?(.+?)\s+(?:using\s+|with\s+|to\s+|in\s+|while\s+)(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Установите {_body((m.group(1) or '') + m.group(2))} {_body(m.group(3))}{(': ' + _ref_tail('refer to '+m.group(4))) if m.group(4) else ''}."
)
_T(
    r'^Install\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Установите {_body((m.group(1) or '') + m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── Remove [part] [from ...] [refer to ...] ───
_T(
    r'^Remove\s+(the\s+)?(.+?)\s+from\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Снимите {_body((m.group(1) or '') + m.group(2))} с {_body(m.group(3))}{(': ' + _ref_tail('refer to '+m.group(4))) if m.group(4) else ''}."
)
_T(
    r'^Remove\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Снимите {_body((m.group(1) or '') + m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── Identify the part with the [company] repair number [num] ───
_T(
    r'^Identify the part with the (Messier-Dowty Limited|Safran Landing Systems) repair number\s+(\S+)\s+adjacent to the part number(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Идентифицируйте деталь по номеру ремонта {m.group(1)} {m.group(2)} рядом с номером детали{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)
_T(
    r'^Identify the part with the (Messier-Dowty Limited|Safran Landing Systems) repair number\s+(.+?)\.',
    lambda m: f"Идентифицируйте деталь по номеру ремонта {m.group(1)} {m.group(2)}."
)

# ─── Locally [verb] [part] ───
_T(
    r'^Locally apply\s+(.+?)\s+to\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Локально нанесите {_coat_nom(m.group(1))} на {_body(m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)
_T(
    r'^Locally\s+(cadmium plate|anodise|cadmium|nodise)\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Локально нанесите {'кадмиевое покрытие' if 'cadmium' in m.group(1).lower() else 'анодное покрытие'} на {_body((m.group(2) or '') + m.group(3))}{(': ' + _ref_tail('refer to '+m.group(4))) if m.group(4) else ''}."
)

# ─── Torque [part] to between X and Y ───
_T(
    r'^Torque\s+(the\s+)?(.+?)\s+to\s+between\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Затяните {_body((m.group(1) or '') + m.group(2))} моментом от {m.group(3)}{(': ' + _ref_tail('refer to '+m.group(4))) if m.group(4) else ''}."
)
_T(
    r'^Torque\s+(the\s+)?(.+?)\s+to\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Затяните {_body((m.group(1) or '') + m.group(2))} моментом {m.group(3)}{(': ' + _ref_tail('refer to '+m.group(4))) if m.group(4) else ''}."
)

# ─── Safety the [part] / Safety [part] with lockwire ───
_T(
    r'^Safety\s+(the\s+)?(.+?)\s+(?:to\s+(.+?)\s+)?with\s+lockwire(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Законтрите {_body((m.group(1) or '') + m.group(2))}{(' к ' + _body(m.group(3))) if m.group(3) else ''} контровочной проволокой{(': ' + _ref_tail('refer to '+m.group(4))) if m.group(4) else ''}."
)
_T(
    r'^Safety\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Законтрите {_body((m.group(1) or '') + m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── Use [tool/pad] and install ───
_T(
    r'^Use\s+(the\s+)?(.+?)\s+(?:and\s+)?(?:to\s+)?install\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Используя {_body((m.group(1) or '') + m.group(2))}, установите {_body((m.group(3) or '') + m.group(4))}{(': ' + _ref_tail('refer to '+m.group(5))) if m.group(5) else ''}."
)
_T(
    r'^Use\s+(the\s+)?(.+?)\s+(?:to\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Используйте {_body((m.group(1) or '') + m.group(2))} для {_body(m.group(3))}{(': ' + _ref_tail('refer to '+m.group(4))) if m.group(4) else ''}."
)

# ─── Make sure [clause] ───
_T(
    r'^Make sure\s+(?:that\s+)?(.+?)\.',
    lambda m: f"Убедитесь в том, что {_body(m.group(1))}."
)

# ─── Measure and record ───
_T(
    r'^Measure and record\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Измерьте и запишите {_body(m.group(1))}{(': ' + _ref_tail('refer to '+m.group(2))) if m.group(2) else ''}."
)

# ─── Prepare [part] ───
_T(
    r'^Prepare\s+(the\s+)?(.+?)\s+(?:with\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Подготовьте {_body((m.group(1) or '') + m.group(2))} {_body(m.group(3))}{(': ' + _ref_tail('refer to '+m.group(4))) if m.group(4) else ''}."
)

# ─── Calculate [dim] ───
_T(
    r'^Calculate\s+(the\s+)?(.+?),\s*use (?:the\s+)?formula:(.*)$',
    lambda m: f"Рассчитайте {_body((m.group(1) or '') + m.group(2))} по формуле:{m.group(3)}"
)
_T(
    r'^Calculate\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Рассчитайте {_body((m.group(1) or '') + m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── Assemble [part] ───
_T(
    r'^Assemble\s+(the\s+)?(.+?)\s+(?:to\s+|for\s+|and\s+)(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Соберите {_body((m.group(1) or '') + m.group(2))} {_body(m.group(3))}{(': ' + _ref_tail('refer to '+m.group(4))) if m.group(4) else ''}."
)
_T(
    r'^Assemble\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Соберите {_body((m.group(1) or '') + m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── Align [part] ───
_T(
    r'^Align\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Совместите {_body((m.group(1) or '') + m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── Check / Verify ───
_T(
    r'^(?:Check|Verify)\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Проверьте {_body((m.group(1) or '') + m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── If necessary [verb] ───
_T(
    r'^If necessary\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"При необходимости {_body(m.group(1))}{(': ' + _ref_tail('refer to '+m.group(2))) if m.group(2) else ''}."
)

# ─── Clean [part] ───
_T(
    r'^Clean\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Очистите {_body((m.group(1) or '') + m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── Repair to [Component] Figure N - Sheet M ───
_T(
    r'^Repair to (.+?)\s+Figure\s+(\d+)\s*[-–]\s*Sheet\s+(\d+)(.*?)$',
    lambda m: f"Ремонт {_comp_gen(m.group(1))} Рисунок {m.group(2)} – Лист {m.group(3)}{m.group(4)}"
)
_T(
    r'^Repair to (.+?)\s+-\s+Machining Figure\s+(\d+)(.*?)$',
    lambda m: f"Ремонт {_comp_gen(m.group(1))} – Механическая обработка Рисунок {m.group(2)}{m.group(3)}"
)
_T(
    r'^Repair to (.+?)\s+Figure\s+(\d+)(.*?)$',
    lambda m: f"Ремонт {_comp_gen(m.group(1))} Рисунок {m.group(2)}{m.group(3)}"
)

# ─── Repair Bush/Sleeve - Machining and Installation Figure N ───
_T(
    r'^Repair (?:Bush|Sleeve)\s*-\s*Machining(?: and Installation)?\s+Figure\s+(\d+)(.*?)$',
    lambda m: f"Ремонтная втулка – Механическая обработка{'и установка ' if 'Installation' in m.group(0) else ' '}Рисунок {m.group(1)}{m.group(2)}"
)

# ─── [Component] - Installation Figure N (Sheet M of K) ───
_T(
    r'^(.+?)\s*(?:Assembly and\s+)?Installation Figure\s+(\d+)(.*?)$',
    lambda m: f"{_comp(m.group(1))} – Установка Рисунок {m.group(2)}{m.group(3)}"
)

# ─── Repair No. XX-X [Component] ([PN]) ───
_T(
    r'^Repair No\.\s*(\d+[-–]\d+)\s+(.+?)\s+\((.+?)\)\s*$',
    lambda m: f"Ремонт № {m.group(1)} {_comp(m.group(2))} ({m.group(3)})"
)

# ─── Fits and Clearances (Table N) Figure N ───
_T(
    r'^Fits and Clearances\s*\(Table\s+(\d+)\)\s+Figure\s+(\d+)(.*?)$',
    lambda m: f"Посадки и зазоры (Таблица {m.group(1)}) Рисунок {m.group(2)}{m.group(3)}"
)
_T(
    r'^Fits and Clearances Table\s+(\d+)(.*?)$',
    lambda m: f"Посадки и зазоры, Таблица {m.group(1)}{m.group(2).replace('(Continued)', '(Продолжение)')}"
)
_T(
    r'^Table\s+(\d+)\s*\(Continued\)(.*?)$',
    lambda m: f"Таблица {m.group(1)} (Продолжение){m.group(2)}"
)

# ─── Grease Groove Dimensions After Installation in the [housing] ───
_T(
    r'^Grease Groove Dimensions After Installation in the (.+?)\s*\((.+?)\)\s*$',
    lambda m: f"Размеры смазочной канавки после установки в {_comp(m.group(1))} ({m.group(2)})"
)

# ─── Assembly and Installation of [X] (PN) for [Y] (PN) ───
_T(
    r'^(?:Assembly and\s+)?Installation of\s+(.+?)\s+\((.+?)\)\s+for\s+(.+?)\s+\((.+?)\)(.*?)$',
    lambda m: f"{'Сборка и установка' if 'Assembly' in m.group(0) else 'Установка'} {_comp_gen(m.group(1))} ({m.group(2)}) для {_comp_gen(m.group(3))} ({m.group(4)}){m.group(5)}"
)

# ─── Move the [part] into [place] ───
_T(
    r'^Move\s+(the\s+)?(.+?)\s+into\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Переместите {_body((m.group(1) or '') + m.group(2))} в {_body((m.group(3) or '') + m.group(4))}{(': ' + _ref_tail('refer to '+m.group(5))) if m.group(5) else ''}."
)

# ─── Damage or wear to [part] ───
_T(
    r'^Damage or (?:wear|corrosion) to\s+(.+?)\.$',
    lambda m: f"Повреждение или износ {_body(m.group(1))}."
)
_T(
    r'^Damage, wear or corrosion to\s+(.+?)\.$',
    lambda m: f"Повреждение, износ или коррозия {_body(m.group(1))}."
)
_T(
    r'^Damage or (?:wear|corrosion) to (?:the\s+)?(.+?) and/or (.+?)\.$',
    lambda m: f"Повреждение или износ {_body(m.group(1))} и/или {_body(m.group(2))}."
)

# ─── To repair [damage] to [location] ───
_T(
    r'^To repair\s+(.+?)\s+(?:to|at|of)\s+(.+?)\.$',
    lambda m: f"Для ремонта {_body(m.group(1))} на {_body(m.group(2))}."
)

# ─── Do this procedure if [condition] ───
_T(
    r'^Do this procedure if\s+(.+?)[:\.]$',
    lambda m: f"Выполните данную процедуру, если {_body(m.group(1))}."
)
_T(
    r'^Do this procedure if\s+(.+?)$',
    lambda m: f"Выполните данную процедуру, если {_body(m.group(1))}"
)

# ─── [Component] Subassembly (...) for [application] ───
_T(
    r'^(.+?)\s+Subassembly\s+\((.+?)\)\s+for\s+(.+?)$',
    lambda m: f"Сборка {_comp(m.group(1))} ({m.group(2)}) для {_comp(m.group(3))}"
)

# ─── Pre/Post SB ... ───
_T(
    r'^(Pre|Post)\s+SB\s+([\d\-]+(?:\s+or\s+[\d\-]+)*):?\s+(.+)$',
    lambda m: f"{'До' if m.group(1).lower()=='pre' else 'После'} SB {m.group(2)}: {_body(m.group(3))}"
)
_T(
    r'^(PRE|POST)\s+SB\s+([\d\-]+(?:\s+OR\s+[\d\-]+)*):?\s+(.+)$',
    lambda m: f"{'ДО' if m.group(1).upper()=='PRE' else 'ПОСЛЕ'} SB {m.group(2)}: {_body(m.group(3))}"
)
_T(
    r'^(PRE|POST)\s+REF\.\s+CODE:\s*(\d+):?\s+(.+)$',
    lambda m: f"{'ДО' if m.group(1).upper()=='PRE' else 'ПОСЛЕ'} КОД ССЫЛКИ: {m.group(2)}: {_body(m.group(3))}"
)

# ─── NOTE: ... / CAUTION: ... ───
_T(
    r'^NOTE:\s*(.+)$',
    lambda m: f"ПРИМЕЧАНИЕ: {_body(m.group(1))}"
)
_T(
    r'^CAUTION:\s*(.+)$',
    lambda m: f"ВНИМАНИЕ: {_body(m.group(1))}"
)
_T(
    r'^WARNING:\s*(.+)$',
    lambda m: f"ПРЕДУПРЕЖДЕНИЕ: {_body(m.group(1))}"
)

# ─── Refer to Figure N / para AG ───
_T(
    r'^Refer to Figure\s+(\d+)\.\s+(.+)$',
    lambda m: f"см. Рисунок {m.group(1)}. {_body(m.group(2))}"
)
_T(
    r'^Refer to Figure\s+(\d+)\.\s+(.+)$',
    lambda m: f"см. Рисунок {m.group(1)}. {_body(m.group(2))}"
)

# ─── [Part name] ... - Installation Figure N (Sheet M of K) ───
_T(
    r'^(.+?)\s+(?:for\s+)?Shock Absorber Subassembly\s*(.*)$',
    lambda m: f"{_comp(m.group(1))} для сборки амортизатора {m.group(2)}"
)

# ─── Apply primer paint to [area] ───
_T(
    r'^Apply primer paint to\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Нанесите грунтовочную краску на {_body(m.group(1))}{(': ' + _ref_tail('refer to '+m.group(2))) if m.group(2) else ''}."
)

# ─── Anodise [part] ───
_T(
    r'^Anodis[e]?\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Анодируйте {_body((m.group(1) or '') + m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── Bond the [part] to [surface] ───
_T(
    r'^Bond\s+(the\s+)?(.+?)\s+to\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Приклейте {_body((m.group(1) or '') + m.group(2))} к {_body((m.group(3) or '') + m.group(4))}{(': ' + _ref_tail('refer to '+m.group(5))) if m.group(5) else ''}."
)

# ─── Record [action] ───
_T(
    r'^Record\s+(the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Запишите {_body((m.group(1) or '') + m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── THE SURFACE FINISH MUST BE ... ───
_T(
    r'^THE SURFACE FINISH MUST BE(.+?)(?:UNLESS GIVEN DIFFERENTLY)?\.?\s*(.*)$',
    lambda m: f"ЧИСТОТА ПОВЕРХНОСТИ ДОЛЖНА БЫТЬ{m.group(1).rstrip()}{'ЕСЛИ НЕ УКАЗАНО ИНАЧЕ' if 'UNLESS GIVEN' in m.group(0).upper() else ''}. {m.group(2)}"
)

# ─── X DEGREES CHAMFER / INCLUSIVE CHAMFER ───
_T(
    r'^(.+?)\s*x\s*(.+?)\s*DEGREES?\s*(INCLUSIVE\s+)?CHAMFER(?:\s+(\d+\s+PLACES?))?(.*)$',
    lambda m: f"{m.group(1)} × {m.group(2)} {'ВКЛЮЧИТЕЛЬНО ' if m.group(3) else ''}ГРАДУСОВ ФАСКА{(' ' + m.group(4).replace('PLACES', 'МЕСТ').replace('PLACE', 'МЕСТО')) if m.group(4) else ''}{m.group(5)}"
)

# ─── [PN] mm MINIMUM WALL THICKNESS ───
_T(
    r'^(.+?mm\s+\(.+?\))\s+MINIMUM WALL THICKNESS(.*)$',
    lambda m: f"{m.group(1)} МИН. ТОЛЩИНА СТЕНКИ{m.group(2)}"
)

# ─── CHECK DIAMETER (TYPICAL N PLACES) ───
_T(
    r'^CHECK DIAMETER\s*\(TYPICAL\s+(\d+)\s+PLACES?\)(.*)$',
    lambda m: f"КОНТРОЛЬНЫЙ ДИАМЕТР (ТИПИЧНО {m.group(1)} МЕСТ){m.group(2)}"
)

# ─── [value] LENGTH OF CHROMIUM PLATE ───
_T(
    r'^(.+?)\s+LENGTH OF CHROMIUM PLATE(.*)$',
    lambda m: f"{m.group(1)} ДЛИНА ХРОМОВОГО ПОКРЫТИЯ{m.group(2)}"
)

# ─── REPAIR BUSH XXXXXXXX / REPAIR SLEEVE XXXXXXXX ───
_T(
    r'^REPAIR BUSH\s+(\S+)\s*$',
    lambda m: f"РЕМОНТНАЯ ВТУЛКА {m.group(1)}"
)
_T(
    r'^REPAIR SLEEVE\s+(\S+)\s*$',
    lambda m: f"РЕМОНТНАЯ ВСТАВКА {m.group(1)}"
)

# ─── TYPICAL N PLACES ───
_T(
    r'^\(?TYPICAL\s+(\d+)\s+PLACES?\)?\s*$',
    lambda m: f"(ТИПИЧНО В {m.group(1)} МЕСТАХ)"
)
_T(
    r'^MAX\s+(\d+)\s+PLACES?\s*$',
    lambda m: f"МАКС. В {m.group(1)} МЕСТАХ"
)

# ─── Remove [coating] locally from [part] ───
_T(
    r'^Remove (?:the\s+)?(.+?)\s+locally\s+from (?:the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Снимите {_coat_acc(m.group(1))} локально с {_body(m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)
_T(
    r'^Remove (?:the\s+)?(.+?)\s+from (?:the\s+)?(.+?)(?:[.:]\s*[Rr]efer to\s+(.+?))?\.?$',
    lambda m: f"Снимите {_coat_acc(m.group(1))} с {_body(m.group(2))}{('. ' + _ref_tail('Refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── Restore [treatments] to [part] ───
_T(
    r'^Restore (?:the\s+)?(.+?)\s+to (?:the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Восстановите {_body(m.group(1))} на {_body(m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── If there is evidence/sign of delamination ───
_T(
    r'^If there is (?:evidence|sign) of delamination,\s*(.+?)\.$',
    lambda m: f"При наличии признаков расслоения {_body(m.group(1))}."
)

# ─── Do the above step N more times ───
_T(
    r'^Do the (?:above|this)\s+step\s+(.+?)\s+more times?\.$',
    lambda m: f"Повторите {'данный' if 'this' in m.group(0).lower() else 'вышеуказанный'} шаг ещё {m.group(1)} раз{'а' if m.group(1) in ('2','3','4') else ''}."
)

# ─── Select applicable repair bushes from Table N ───
_T(
    r'^Select (?:the\s+)?applicable repair bushes? from Table\s+(\d+)\s+for (.+?)\.',
    lambda m: f"Выберите применимые ремонтные втулки из Таблицы {m.group(1)} для {_body(m.group(2))}."
)

# ─── Apply protective treatment ───
_T(
    r'^Apply protective treatment(?:s)? to (?:the\s+)?(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Нанесите защитное покрытие на {_body(m.group(1))}{(': ' + _ref_tail('refer to '+m.group(2))) if m.group(2) else ''}."
)

# ─── Grit blast area to be [coating]: Make sure X is masked ───
_T(
    r'^Grit blast the area to be (.+?) plated:\s*refer to\s+(.+?)\.\s*Make sure that (?:the\s+)?(.+?) is correctly masked\.$',
    lambda m: f"Выполните пескоструйную обработку зоны для нанесения {_body(m.group(1))} покрытия: см. {m.group(2)}. Убедитесь, что {_body(m.group(3))} правильно замаскирован(а)."
)

# ─── SECTION Z-Z (WITH/WITHOUT BUSHES) REFER TO FIGURE N ───
_T(
    r'^SECTION\s+([A-Z]\-[A-Z])\s+\((.+?)\)\s+REFER TO FIGURE\s+(\d+)\s*$',
    lambda m: f"РАЗРЕЗ {m.group(1)} ({m.group(2).replace('WITH BUSHES', 'С ВТУЛКАМИ').replace('WITHOUT BUSHES', 'БЕЗ ВТУЛОК').replace('WITH SLEEVES', 'С ВСТАВКАМИ').replace('WITHOUT SLEEVES', 'БЕЗ ВСТАВОК')}) СМ. РИСУНОК {m.group(3)}"
)

# ─── Make/The surface finish must be N micrometers ───
_T(
    r'^Make the surface finish\s+([\d,\.]+)\s+micrometers?\s*\((.+?)\)\.?$',
    lambda m: f"Чистота поверхности должна составлять {m.group(1)} мкм ({m.group(2)})."
)
_T(
    r'^The surface finish must be\s+([\d,\.]+)\s+micrometers?\s*\((.+?)\)(?:\s+or better)?(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Чистота поверхности должна составлять {m.group(1)} мкм ({m.group(2)}){' или лучше' if 'better' in m.group(0) else ''}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── The [coating] thickness must be between X and Y ───
_T(
    r'^The (?:cadmium plate|zinc nickel plate|chromium plate|plating|sulphamate nickel plate) thickness must (?:be between|not be (?:more|less) than)\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Толщина покрытия {'должна быть от' if 'between' in m.group(0) else ('не должна превышать' if 'more' in m.group(0) else 'не должна быть менее')} {m.group(1)}{(': ' + _ref_tail('refer to '+m.group(2))) if m.group(2) else ''}."
)

# ─── Apply zinc loaded jointing compound ───
_T(
    r'^Apply [Zz]inc loaded [Jj]ointing compound,?\s*Molykote\s+(\d+)\s+to\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Нанесите монтажный состав с добавкой цинка, Molykote {m.group(1)}, на {_body(m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── NOTE: Install the bush by heating the housing to N °C ───
_T(
    r'^NOTE:\s*Install the bush(?:es)? by heating the housing(?: to (.+?))? and cooling the bush(?:es)? only\.$',
    lambda m: f"ПРИМЕЧАНИЕ: Установите втулк{'и' if 'bushes' in m.group(0) else 'у'}, нагрев корпус{(' до '+m.group(1)) if m.group(1) else ''} и охладив только втулк{'и' if 'bushes' in m.group(0) else 'у'}."
)

# ─── NOTE: You can lubricate the seals with grease ───
_T(
    r'^NOTE:\s*You can lubricate the seals? with grease,?\s*(.+?)$',
    lambda m: f"ПРИМЕЧАНИЕ: Допускается смазка уплотнений консистентной смазкой, {_body(m.group(1))}"
)

# ─── Torque the bolts/nuts to the value N ───
_T(
    r'^Torque\s+(?:the\s+)?(.+?)\s+to the value\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Затяните {_body(m.group(1))} моментом {m.group(2)}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── Use the Pin Spanner NN and torque the nut assembly to N ───
_T(
    r'^Use the (.+?) and torque\s+(.+?)\s+to\s+(.+?)(?:\s+then\s+(.+?))?\.?$',
    lambda m: f"Используйте {_body(m.group(1))} и затяните {_body(m.group(2))} моментом {m.group(3)}{(', затем ' + _body(m.group(4))) if m.group(4) else ''}."
)

# ─── Apply [compound] to [part] where shown ───
_T(
    r'^Apply\s+(.+?)\s+to\s+(.+?)\s+(?:bores?\s+)?(?:where shown\s*)?(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Нанесите {_body(m.group(1))} на {_body(m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── Use Press Pad XXXXX and install ───
_T(
    r'^Use (?:the\s+)?[Pp]ress [Pp]ad\s+([\w/]+)\s+and\s+(?:the\s+)?(?:[Dd]rift|[Oo]utside\s+[Pp]unch)\s+([\w/]+)\s+(?:to\s+|and\s+)?install\s+(.+?)(?::\s*refer to\s+(.+?))?\.?(.*)$',
    lambda m: f"Используйте прессовую подушку {m.group(1)} и оправку {m.group(2)} для установки {_body(m.group(3))}{(': ' + _ref_tail('refer to '+m.group(4))) if m.group(4) else ''}.{m.group(5)}"
)
_T(
    r'^Use (?:the\s+)?[Pp]ress [Pp]ad\s+([\w/]+)\s+(?:to\s+|and\s+)?install\s+(.+?)(?::\s*refer to\s+(.+?))?\.?(.*)$',
    lambda m: f"Используйте прессовую подушку {m.group(1)} для установки {_body(m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}.{m.group(4)}"
)

# ─── Calculate the diameter/dimension for repair busheS ───
_T(
    r'^Calculate (?:the\s+)?(.+?) for (?:the\s+)?(.+?),\s*use (?:the\s+)?formulas?:?\s*$',
    lambda m: f"Рассчитайте {_body(m.group(1))} для {_body(m.group(2))} по формулам:"
)
_T(
    r'^Calculate (?:the\s+)?(?:dimensions? of\s+)?(?:the\s+)?(.+?):\s*refer to\s+(.+?)\.\s*[Uu]se (?:the\s+)?formula:?\s*$',
    lambda m: f"Рассчитайте размеры {_body(m.group(1))}: {_ref_tail('refer to '+m.group(2))}. Используйте формулу:"
)


# ─── Standalone Figure/Table references ───
_T(
    r'^Figure\s+(\d+)\s*[-–]?\s*(Sheet\s+(\d+).*)?\s*$',
    lambda m: f"Рисунок {m.group(1)}{(' – Лист ' + m.group(3) + m.group(2)[len('Sheet '+m.group(3)):]) if m.group(2) else ''}"
)
_T(
    r'^Figure\s+(\d+)\s*$',
    lambda m: f"Рисунок {m.group(1)}"
)
_T(
    r'^Table\s+(\d+)\s*$',
    lambda m: f"Таблица {m.group(1)}"
)

# ─── Do not remove/make/machine more than ... ───
_T(
    r'^Do not (?:remove|machine)\s+(?:the\s+)?(?:faces?\s+)?more than\s+(.+?)(?:\s+of\s+material(?:\s+from\s+(?:each\s+of\s+)?(?:the\s+)?(.+?))?)?\.?$',
    lambda m: f"Не снимайте более {m.group(1)} материала{(' с каждого ' + _body(m.group(2))) if m.group(2) else ''}."
)
_T(
    r'^Do not make\s+(?:the\s+)?(.+?)\s+more than\s+(.+?)\.?$',
    lambda m: f"Не увеличивайте {_body(m.group(1))} более {m.group(2)}."
)
_T(
    r'^Do not machine the faces? more than\s+(.+?)\.?$',
    lambda m: f"Не обрабатывайте торец более {m.group(1)}."
)

# ─── Diameter X must be between Y ───
_T(
    r'^(?:The\s+)?[Dd]iameter\s+([A-Z])\s+must be between\s+(.+?)\.?$',
    lambda m: f"Диаметр {m.group(1)} должен быть от {m.group(2)}."
)
_T(
    r'^(?:The\s+)?[Dd]iameter\s+must be between\s+(.+?)\.?$',
    lambda m: f"Диаметр должен быть от {m.group(1)}."
)

# ─── The minimum wall thickness must be X ───
_T(
    r'^(?:The\s+)?(?:and\s+)?(?:the\s+)?minimum wall thickness must be\s+(.+?)\.?$',
    lambda m: f"Минимальная толщина стенки должна быть {m.group(1)}."
)

# ─── Lubricate the threads/part with grease/oil ───
_T(
    r'^Lubricate the (?:threads?|bearings?|seals?)\s+of\s+(?:the\s+)?(.+?)\s+with\s+(.+?)(?::\s*refer to\s+(.+?))?\.?$',
    lambda m: f"Смажьте резьбу {_body(m.group(1))} {_body(m.group(2))}{(': ' + _ref_tail('refer to '+m.group(3))) if m.group(3) else ''}."
)

# ─── Apply one/two coat(s) of X to Y ───
_T(
    r'^Apply\s+(?:one|two|three)\s+coats?\s+(?:of\s+)?(.+?),?\s+(?:Material\s+)?(?:Ref\.?\s*)?(?:Item\s+)?([\w\-]+),?\s+(?:to|on)\s+(?:the\s+)?(.+?)\.?$',
    lambda m: f"Нанесите слой {_body(m.group(1))}, Поз. {m.group(2)}, на {_body(m.group(3))}."
)

# ─── Leave approximately X between/around Y ───
_T(
    r'^Leave approximately\s+(.+?)\s+(?:between|around|of)\s+(.+?)\.?$',
    lambda m: f"Оставьте приблизительно {m.group(1)} между {_body(m.group(2))}."
)

# ─── For the shock absorber configuration (X or Y or ...): verb ... ───
_T(
    r'^For (?:the\s+)?shock absorber configuration\s+\((.+?)\):\s*(.+?)$',
    lambda m: f"Для конфигурации амортизатора ({m.group(1)}): {_body(m.group(2))}"
)

# ─── [Verb] [part] related components ───
_T(
    r'^(.+?) related components\.?\s*(.*)$',
    lambda m: f"{_body(m.group(1))} связанные компоненты. {_body(m.group(2))}" if m.group(2) else f"{_body(m.group(1))} связанные компоненты."
)


def _comp(s):
    """Translate component name (nominative)."""
    s_low = s.strip().lower()
    for k, v in sorted(_COMP.items(), key=lambda x: -len(x[0])):
        if s_low == k or s_low.startswith(k + ' '):
            return v + s[len(k):] if s_low.startswith(k) else v
    # Fallback: apply _body word substitution
    return _body(s)


def _comp_gen(s):
    """Translate component name (genitive)."""
    _GEN = {
        'main fitting': 'корпуса стойки',
        'main fitting subassembly': 'сборки корпуса стойки',
        'sliding tube': 'скользящей трубы',
        'sliding tube subassembly': 'сборки скользящей трубы',
        'upper torque link': 'верхнего шлиц-шарнира',
        'lower torque link': 'нижнего шлиц-шарнира',
        'upper slave link': 'верхнего ведомого звена',
        'lower slave link': 'нижнего ведомого звена',
        'upper stay': 'верхнего звена',
        'lower stay': 'нижнего звена',
        'lock stay': 'фиксирующего звена',
        'locking stay': 'фиксирующего звена',
        'lock stay cardan': 'кардана фиксирующего звена',
        'cylinder': 'цилиндра',
        'upper diaphragm tube': 'верхней диафрагменной трубы',
        'lower bearing': 'нижнего подшипника',
        'lower bearing subassembly': 'сборки нижнего подшипника',
        'shock absorber': 'амортизатора',
        'shock absorber subassembly': 'сборки амортизатора',
        'gland housing': 'корпуса сальника',
        'retaining pin': 'фиксирующего штифта',
        'upper pivot bracket': 'верхнего кронштейна оси',
        'main landing gear leg': 'стойки основного шасси',
        'torque link': 'шлиц-шарнира',
        'repair bush': 'ремонтной втулки',
        'repair sleeve': 'ремонтной вставки',
    }
    s_low = s.strip().lower()
    if s_low in _GEN:
        return _GEN[s_low]
    return _comp(s)


# ──────────────────────────────────────────────────────────────
# 3.  MAIN translate_text()
# ──────────────────────────────────────────────────────────────

def _translate_single(t):
    """Translate a single sentence/fragment."""
    if not t or not t.strip():
        return t
    t = t.strip()

    # 1. Exact match
    if t in EXACT:
        return _postprocess(EXACT[t])
    tnp = t.rstrip('.')
    if tnp in EXACT:
        return _postprocess(EXACT[tnp] + ('.' if t.endswith('.') else ''))

    # 2. Sentence templates
    for pat, fn in _TEMPLATES:
        m = pat.match(t)
        if m:
            try:
                result = fn(m)
                if result and result != t:
                    return _postprocess(result)
            except Exception:
                pass

    # 3. Generic body substitution
    return _postprocess(_body(t))


def _postprocess(s):
    """Final pass: clean up any remaining common English words in translated text."""
    # ── Fix Figure/Table glued to text ──
    s = re.sub(r'([A-Za-zа-яёА-ЯЁ])(?=Figure\s*\d)', r'\1. ', s)
    s = re.sub(r'([A-Za-zа-яёА-ЯЁ])(?=Table\s*\d)', r'\1. ', s)
    # ── Fix (s) → (ы), (es) → (ы) in Russian context ──
    s = re.sub(r'(?<=\w)\(s\)', '(ы)', s)
    s = re.sub(r'(?<=\w)\(es\)', '(ы)', s)
    # ── Fix duplicate Figure/Table refs: "Рисунок 601.Figure 601" → "Рисунок 601." ──
    s = re.sub(r'(Рисунок\s*\d+)\s*\.?\s*Figure\s*\d+', r'\1', s)
    s = re.sub(r'(Рисунки\s*\d+[–\-]\d+)\s*\.?\s*Figure\s*\d+', r'\1', s)
    s = re.sub(r'(Таблица\s*\d+)\s*\.?\s*Table\s*\d+', r'\1', s)
    s = re.sub(r'\.?(Figure\s+\d+)\s*$', '', s)  # trailing "Figure 601" after period
    s = re.sub(r'\.?(Table\s+\d+)\s*$', '', s)    # trailing "Table 601"
    # ── Domain-specific missing words ──
    s = re.sub(r'\bDe-embrittle\b', 'Выполните устранение водородной хрупкости', s)
    s = re.sub(r'\bde-embrittle\b', 'выполните устранение водородной хрупкости', s)
    s = re.sub(r'\bexternally\b', 'снаружи', s, flags=re.I)
    s = re.sub(r'\bplating thickness\b', 'толщина покрытия', s, flags=re.I)
    s = re.sub(r'\bplating\b', 'покрытия', s, flags=re.I)
    s = re.sub(r'\bprepare\b', 'подготовьте', s, flags=re.I)
    s = re.sub(r'\bfollow\b', 'совпадать с', s, flags=re.I)
    s = re.sub(r'\bmeasured\b', 'измеренный', s, flags=re.I)
    s = re.sub(r'\bspecified\b', 'указанный', s, flags=re.I)
    s = re.sub(r'\bsufficiently\b', 'достаточно', s, flags=re.I)
    s = re.sub(r'\bsuffciently\b', 'достаточно', s, flags=re.I)
    s = re.sub(r'\bwithin\b', 'в пределах', s, flags=re.I)
    s = re.sub(r'\bintensity\b', 'интенсивность', s, flags=re.I)
    s = re.sub(r'\bhand ream\b', 'вручную развернуть', s, flags=re.I)
    s = re.sub(r'\bhand\b', 'вручную', s, flags=re.I)
    s = re.sub(r'\bpainted\b', 'окрашенные', s, flags=re.I)
    s = re.sub(r'\breduce\b', 'уменьшите', s, flags=re.I)
    s = re.sub(r'\btogether\b', 'вместе', s, flags=re.I)
    s = re.sub(r'\bforward\b', 'передний', s, flags=re.I)
    s = re.sub(r'\bcommon\b', 'общий', s, flags=re.I)
    s = re.sub(r'\bblock\b', 'блок', s, flags=re.I)
    s = re.sub(r'\bsealing\b', 'уплотнительный', s, flags=re.I)
    s = re.sub(r'\binflation\b', 'наполнение', s, flags=re.I)
    s = re.sub(r'\bretainers?\b', 'фиксатор', s, flags=re.I)
    s = re.sub(r'\blabels?\b', 'ярлык', s, flags=re.I)
    s = re.sub(r'\bdrag\b', 'тяга', s, flags=re.I)
    s = re.sub(r'\bidentification\b', 'идентификация', s, flags=re.I)
    s = re.sub(r'\blubrication\b', 'смазка', s, flags=re.I)
    s = re.sub(r'\bservice\b', 'обслуживание', s, flags=re.I)
    s = re.sub(r'\bairline\b', 'авиакомпания', s, flags=re.I)
    s = re.sub(r'\bdetails?\b', 'подробности', s, flags=re.I)
    s = re.sub(r'\btype\b', 'тип', s, flags=re.I)
    s = re.sub(r'\binsert(?:s)?\b', 'вставка', s, flags=re.I)
    s = re.sub(r'\bremove\b', 'удалите', s, flags=re.I)
    s = re.sub(r'\bdepth\b', 'глубина', s, flags=re.I)
    s = re.sub(r'\bground\b', 'шлифованный', s, flags=re.I)
    s = re.sub(r'\bmachine\b', 'обработайте', s, flags=re.I)
    s = re.sub(r'\bflush\b', 'заподлицо', s, flags=re.I)
    s = re.sub(r'\binstalled\b', 'установленный', s, flags=re.I)
    s = re.sub(r'\bnecessary\b', 'необходимо', s, flags=re.I)
    s = re.sub(r'\bshown\b', 'показанный', s, flags=re.I)
    s = re.sub(r'\bexisting\b', 'существующий', s, flags=re.I)
    s = re.sub(r'\baxis\b', 'ось', s, flags=re.I)
    s = re.sub(r'\brepair number\b', 'номер ремонта', s, flags=re.I)
    s = re.sub(r'\bdeposit\b', 'слой покрытия', s, flags=re.I)
    s = re.sub(r'\btermination\b', 'граница', s, flags=re.I)
    s = re.sub(r'\binclusive\b', 'включительно', s, flags=re.I)
    s = re.sub(r'\bdegree\b', 'градус', s, flags=re.I)
    s = re.sub(r'\bflaws?\b', 'дефект', s, flags=re.I)
    s = re.sub(r'\btighten\b', 'затяните', s, flags=re.I)
    s = re.sub(r'\btorque\b', 'момент затяжки', s, flags=re.I)
    s = re.sub(r'\balign\b', 'совместите', s, flags=re.I)
    s = re.sub(r'\bnotch\b', 'паз', s, flags=re.I)
    s = re.sub(r'\bbase material\b', 'основной материал', s, flags=re.I)
    s = re.sub(r'\bbase metal\b', 'основной металл', s, flags=re.I)
    s = re.sub(r'\bbase\b', 'основной', s, flags=re.I)
    s = re.sub(r'\bcorroded\b', 'корродированный', s, flags=re.I)
    s = re.sub(r'\bcorrosion\b', 'коррозия', s, flags=re.I)
    s = re.sub(r'\bPassivate\b', 'Пассивируйте', s)
    s = re.sub(r'\bpassivate\b', 'пассивируйте', s)
    s = re.sub(r'\bANODISE\b', 'АНОДИРОВАТЬ', s)
    s = re.sub(r'\banodise\b', 'анодировать', s)
    s = re.sub(r'\bAPPLIED\b', 'НАНЕСЁННОЕ', s)
    s = re.sub(r'\bDRILLED\b', 'СВЕРЛЁНОЕ', s)
    s = re.sub(r'\bMAJOR\b', 'НАРУЖНЫЙ', s)
    s = re.sub(r'\bFULL\b', 'ПОЛНЫЙ', s)
    s = re.sub(r'\bFINE\b', 'ЧИСТОВОЙ', s)
    s = re.sub(r'\bGRINDING\b', 'ШЛИФОВАНИЯ', s)
    s = re.sub(r'\bgrinding\b', 'шлифования', s)
    s = re.sub(r'\bTERMINATE\b', 'ЗАКАНЧИВАТЬСЯ', s)
    s = re.sub(r'\bterminate\b', 'заканчиваться', s)
    s = re.sub(r'\bsolution\b', 'раствор', s, flags=re.I)
    s = re.sub(r'\bremaining\b', 'оставшийся', s, flags=re.I)
    s = re.sub(r'\boverlaps?\b', 'перекрывает', s, flags=re.I)
    s = re.sub(r'\bgive\b', 'обеспечить', s, flags=re.I)
    s = re.sub(r'\bsecond stage\b', 'второй ступени', s, flags=re.I)
    s = re.sub(r'\brestore\b', 'восстановить', s, flags=re.I)
    s = re.sub(r'\bscheme\b', 'схема', s, flags=re.I)
    s = re.sub(r'\blight coat\b', 'тонкий слой', s, flags=re.I)
    s = re.sub(r'\blight\b', 'лёгкий', s, flags=re.I)
    s = re.sub(r'\bcoat\b', 'слой', s, flags=re.I)
    s = re.sub(r'\balternative\b', 'альтернативный', s, flags=re.I)
    s = re.sub(r'\bequivalents?\b', 'эквивалент', s, flags=re.I)
    s = re.sub(r'\bpermitted\b', 'допускается', s, flags=re.I)
    s = re.sub(r'\btransfer\b', 'перенесите', s, flags=re.I)
    s = re.sub(r'\bvalue\b', 'значение', s, flags=re.I)
    s = re.sub(r'\bfinal\b', 'окончательный', s, flags=re.I)
    s = re.sub(r'\bapplication\b', 'нанесение', s, flags=re.I)
    s = re.sub(r'\bwrench\b', 'ключ', s, flags=re.I)
    s = re.sub(r'\bspanner\b', 'ключ', s, flags=re.I)
    s = re.sub(r'\bcrowfoot\b', 'накидной', s, flags=re.I)
    s = re.sub(r'\bstem\b', 'шток', s, flags=re.I)
    s = re.sub(r'\bshape\b', 'форма', s, flags=re.I)
    s = re.sub(r'\bline\b', 'линия', s, flags=re.I)
    s = re.sub(r'\bdrift\b', 'оправка', s, flags=re.I)
    s = re.sub(r'\bbar\b', 'штанга', s, flags=re.I)
    s = re.sub(r'\bshaft\b', 'вал', s, flags=re.I)
    s = re.sub(r'\blift\b', 'подъём', s, flags=re.I)
    s = re.sub(r'\bcross\b', 'поперечный', s, flags=re.I)
    s = re.sub(r'\bbung\b', 'пробка', s, flags=re.I)
    s = re.sub(r'\bdowels?\b', 'штифт', s, flags=re.I)
    s = re.sub(r'\battaching\b', 'крепёжный', s, flags=re.I)
    s = re.sub(r'\bdatum\b', 'база', s, flags=re.I)
    s = re.sub(r'\bwedge\b', 'клин', s, flags=re.I)
    s = re.sub(r'\breaction\b', 'реактивный', s, flags=re.I)
    s = re.sub(r'\badjust\b', 'отрегулируйте', s, flags=re.I)
    s = re.sub(r'\bsubassembl(?:y|ies)\b', 'подсборка', s, flags=re.I)
    s = re.sub(r'\bclass\b', 'класс', s, flags=re.I)
    s = re.sub(r'\battach\b', 'прикрепите', s, flags=re.I)
    s = re.sub(r'\bsheet\b', 'лист', s, flags=re.I)
    s = re.sub(r'\babsorber\b', 'амортизатор', s, flags=re.I)
    s = re.sub(r'\bgrind\b', 'шлифуйте', s, flags=re.I)
    s = re.sub(r'\brequired\b', 'требуется', s, flags=re.I)
    s = re.sub(r'\bwhere\b', 'где', s, flags=re.I)
    s = re.sub(r'\binto\b', 'в', s, flags=re.I)
    s = re.sub(r'\busing\b', 'используя', s, flags=re.I)
    s = re.sub(r'\bnew\b', 'новый', s, flags=re.I)
    s = re.sub(r'\bcomplete\b', 'полный', s, flags=re.I)
    s = re.sub(r'\bapplicable\b', 'соответствующий', s, flags=re.I)
    s = re.sub(r'\bsufficient\b', 'достаточный', s, flags=re.I)
    s = re.sub(r'\bminimum\b', 'минимальный', s, flags=re.I)
    s = re.sub(r'\bmaximum\b', 'максимальный', s, flags=re.I)
    s = re.sub(r'\bwidth\b', 'ширина', s, flags=re.I)
    s = re.sub(r'\blength\b', 'длина', s, flags=re.I)
    s = re.sub(r'\barea\b', 'область', s, flags=re.I)
    s = re.sub(r'\bwear\b', 'износ', s, flags=re.I)
    s = re.sub(r'\bdamage\b', 'повреждение', s, flags=re.I)
    s = re.sub(r'\bexcess\b', 'избыток', s, flags=re.I)
    s = re.sub(r'\bconcentric\b', 'концентрический', s, flags=re.I)
    s = re.sub(r'\bcoaxial\b', 'соосный', s, flags=re.I)
    s = re.sub(r'\bclean\b', 'очистите', s, flags=re.I)
    s = re.sub(r'\bthat\b', 'что', s, flags=re.I)
    s = re.sub(r'\bthis\b', 'данный', s, flags=re.I)
    s = re.sub(r'\bit\b', '', s, flags=re.I)
    s = re.sub(r'\bis\b', '', s, flags=re.I)
    s = re.sub(r'\bare\b', '', s, flags=re.I)
    s = re.sub(r'\bwas\b', 'был', s, flags=re.I)
    s = re.sub(r'\bbeen\b', '', s, flags=re.I)
    s = re.sub(r'\bhave\b', '', s, flags=re.I)
    s = re.sub(r'\bhas\b', '', s, flags=re.I)
    s = re.sub(r'\bcan\b', 'может', s, flags=re.I)
    s = re.sub(r'\bnot\b', 'не', s, flags=re.I)
    s = re.sub(r'\bwill\b', 'будет', s, flags=re.I)
    # Common words that slip through templates
    s = re.sub(r'\bor\b', 'или', s)
    s = re.sub(r'\bOR\b', 'ИЛИ', s)
    s = re.sub(r'\band\b', 'и', s)
    s = re.sub(r'\bAND\b', 'И', s)
    s = re.sub(r'\bthe\b', '', s, flags=re.I)
    s = re.sub(r'\bof\b', '', s, flags=re.I)
    s = re.sub(r'\bto\b', 'до', s)
    s = re.sub(r'\bTO\b', 'ДО', s)
    s = re.sub(r'\bOF\b', '', s)
    s = re.sub(r'\bTHE\b', '', s)
    s = re.sub(r'\bfor\b', 'для', s, flags=re.I)
    s = re.sub(r'\bin\b', 'в', s, flags=re.I)
    s = re.sub(r'\bon\b', 'на', s, flags=re.I)
    s = re.sub(r'\bwith\b', 'с', s, flags=re.I)
    s = re.sub(r'\bas\b', 'как', s, flags=re.I)
    s = re.sub(r'\brefer\b', 'см.', s, flags=re.I)
    s = re.sub(r'\bonly\b', 'только', s, flags=re.I)
    s = re.sub(r'\bbut\b', 'но', s, flags=re.I)
    s = re.sub(r'\bnot\b', 'не', s, flags=re.I)
    s = re.sub(r'\bmust\b', 'должен', s, flags=re.I)
    s = re.sub(r'\bbe\b', '', s, flags=re.I)
    s = re.sub(r'\bbetween\b', 'от', s, flags=re.I)
    s = re.sub(r'\bafter\b', 'после', s, flags=re.I)
    s = re.sub(r'\bbefore\b', 'до', s, flags=re.I)
    s = re.sub(r'\bbelow\b', 'ниже', s, flags=re.I)
    s = re.sub(r'\babove\b', 'выше', s, flags=re.I)
    s = re.sub(r'\bDo\b', '', s)
    s = re.sub(r'\bdo\b', '', s)
    # Component names
    s = re.sub(r'\bmain fitting\b', 'корпус стойки', s, flags=re.I)
    s = re.sub(r'\bsliding tube\b', 'скользящая труба', s, flags=re.I)
    s = re.sub(r'\bshock absorber\b', 'амортизатор', s, flags=re.I)
    s = re.sub(r'\bupper\b', 'верхний', s, flags=re.I)
    s = re.sub(r'\blower\b', 'нижний', s, flags=re.I)
    s = re.sub(r'\bbush(?:es)?\b', 'втулка', s, flags=re.I)
    s = re.sub(r'\bbearing\b', 'подшипник', s, flags=re.I)
    s = re.sub(r'\bbracket\b', 'кронштейн', s, flags=re.I)
    s = re.sub(r'\btube\b', 'труба', s, flags=re.I)
    s = re.sub(r'\bcylinder\b', 'цилиндр', s, flags=re.I)
    s = re.sub(r'\blink\b', 'звено', s, flags=re.I)
    s = re.sub(r'\bslave\b', 'ведомый', s, flags=re.I)
    s = re.sub(r'\bpivot\b', 'ось', s, flags=re.I)
    s = re.sub(r'\bfitting\b', 'фитинг', s, flags=re.I)
    s = re.sub(r'\bsurface\b', 'поверхность', s, flags=re.I)
    s = re.sub(r'\bfinish\b', 'чистота', s, flags=re.I)
    s = re.sub(r'\bmicrometers?\b', 'мкм', s, flags=re.I)
    s = re.sub(r'\bmicro[- ]inches?\b', 'микродюймов', s, flags=re.I)
    s = re.sub(r'\binches\b', 'дюймов', s, flags=re.I)
    s = re.sub(r'\bthickness\b', 'толщина', s, flags=re.I)
    s = re.sub(r'\bcadmium plate\b', 'кадмиевое покрытие', s, flags=re.I)
    s = re.sub(r'\bplate\b', 'пластина', s, flags=re.I)
    s = re.sub(r'\bdiameter\b', 'диаметр', s, flags=re.I)
    s = re.sub(r'\bdamaged\b', 'повреждённый', s, flags=re.I)
    s = re.sub(r'\bworn\b', 'изношенный', s, flags=re.I)
    s = re.sub(r'\breference\b', 'ссылка', s, flags=re.I)
    s = re.sub(r'\bletter\b', 'буква', s, flags=re.I)
    s = re.sub(r'\bgiven\b', 'указанным', s, flags=re.I)
    s = re.sub(r'\bbetter\b', 'лучше', s, flags=re.I)
    s = re.sub(r'\bover\b', 'по', s, flags=re.I)
    s = re.sub(r'\ball\b', 'все', s, flags=re.I)
    s = re.sub(r'\bthen\b', 'затем', s, flags=re.I)
    s = re.sub(r'\bcontinue\b', 'продолжите', s, flags=re.I)
    s = re.sub(r'\bMake\b', 'Обеспечьте', s)
    s = re.sub(r'\bmake\b', 'обеспечьте', s)
    s = re.sub(r'\bsure\b', '', s, flags=re.I)
    s = re.sub(r'\bmore\b', 'более', s, flags=re.I)
    s = re.sub(r'\bthan\b', 'чем', s, flags=re.I)
    s = re.sub(r'\bfrom\b', 'с', s, flags=re.I)
    s = re.sub(r'\beach\b', 'каждый', s, flags=re.I)
    s = re.sub(r'\bby\b', '', s, flags=re.I)
    s = re.sub(r'\bFigure\b', 'Рисунок', s)
    s = re.sub(r'\bTable\b', 'Таблица', s)
    s = re.sub(r'\bSheet\b', 'Лист', s)
    # "FITS AND CLEARANCES" uppercase
    s = re.sub(r'\bFITS AND CLEARANCES\b', 'ПОСАДКИ И ЗАЗОРЫ', s)
    # Hydraulic
    s = re.sub(r'\bhydraulic\b', 'гидравлический', s, flags=re.I)
    s = re.sub(r'\badapter\b', 'переходник', s, flags=re.I)
    s = re.sub(r'\bzero\b', 'ноль', s, flags=re.I)
    s = re.sub(r'\bprocedure\b', 'процедура', s, flags=re.I)
    s = re.sub(r'\btheir\b', 'их', s, flags=re.I)
    s = re.sub(r'\bLubricate\b', 'Смажьте', s)
    s = re.sub(r'\blubricate\b', 'смажьте', s)
    s = re.sub(r'\bRefer\b', 'См.', s)
    # Fix words glued to numbers (no word boundary between digit and letter)
    s = re.sub(r'(?<=\d)only(?=[^a-zA-Z]|$)', ' только', s, flags=re.I)
    s = re.sub(r'only(?=[A-Z])', 'только ', s)  # "onlyFITS" → "только FITS"
    # More missing words
    s = re.sub(r'\bfillet\b', 'галтель', s, flags=re.I)
    s = re.sub(r'\bvisible\b', 'видимый', s, flags=re.I)
    s = re.sub(r'\bpitch\b', 'шаг', s, flags=re.I)
    s = re.sub(r'\bdrill\b', 'сверло', s, flags=re.I)
    s = re.sub(r'\bpoint\b', 'точка', s, flags=re.I)
    s = re.sub(r'\bsharp\b', 'острый', s, flags=re.I)
    s = re.sub(r'\bcorners?\b', 'угол', s, flags=re.I)
    s = re.sub(r'\bhousings?\b', 'корпус', s, flags=re.I)
    s = re.sub(r'\blocally\b', 'локально', s, flags=re.I)
    s = re.sub(r'\bblended\b', 'сглаженный', s, flags=re.I)
    s = re.sub(r'\binclusion class\b', 'класс включений', s, flags=re.I)
    s = re.sub(r'\bretraction actuator\b', 'привод уборки', s, flags=re.I)
    s = re.sub(r'\bPneumatic Pump\b', 'пневматический насос', s)
    s = re.sub(r'\bReactor Pad\b', 'реакторная подушка', s)
    s = re.sub(r'\bwiring diagram\b', 'схема электропроводки', s, flags=re.I)
    s = re.sub(r'\bstaked\b', 'развальцованный', s, flags=re.I)
    s = re.sub(r'\bstaking\b', 'развальцовка', s, flags=re.I)
    s = re.sub(r'\bbreak out\b', 'момент страгивания', s, flags=re.I)
    s = re.sub(r'\bSpray\b', 'Распылите', s)
    s = re.sub(r'\bspray\b', 'распылите', s)
    s = re.sub(r'\b[Aa]dhesive\b', 'клей', s)
    s = re.sub(r'\bgrade\b', 'марка', s, flags=re.I)
    s = re.sub(r'\binstead\b', 'вместо', s, flags=re.I)
    s = re.sub(r'\blockstay\b', 'фиксирующее звено', s, flags=re.I)
    s = re.sub(r'\bcardan\b', 'кардан', s, flags=re.I)
    s = re.sub(r'\bDrift\b', 'оправка', s)
    s = re.sub(r'\bamount\b', 'количество', s, flags=re.I)
    s = re.sub(r'\bincrease\b', 'увеличивайте', s, flags=re.I)
    s = re.sub(r'\bcenters?\b', 'центр', s, flags=re.I)
    s = re.sub(r'\bcentres?\b', 'центр', s, flags=re.I)
    s = re.sub(r'\bstraight\b', 'прямой', s, flags=re.I)
    s = re.sub(r'\bequal\b', 'равный', s, flags=re.I)
    s = re.sub(r'\bletters?\b', 'буква', s, flags=re.I)
    s = re.sub(r'\bblocked\b', 'заблокированный', s, flags=re.I)
    s = re.sub(r'\bLine up Tool\b', 'инструмент выравнивания', s, flags=re.I)
    s = re.sub(r'\bAssemble\b', 'Соберите', s)
    s = re.sub(r'\bassemble\b', 'соберите', s)
    s = re.sub(r'\bPOST\b', 'ПОСЛЕ', s)
    s = re.sub(r'\bEFFECTIVE\b', 'ДЕЙСТВИТЕЛЬНАЯ', s)
    s = re.sub(r'\bTAPPED\b', 'НАРЕЗАННАЯ', s)
    s = re.sub(r'\bPITCH\b', 'ШАГ', s)
    s = re.sub(r'\bproperly\b', 'правильно', s, flags=re.I)
    s = re.sub(r'\bbonded\b', 'склеенный', s, flags=re.I)
    s = re.sub(r'\bthey\b', 'они', s, flags=re.I)
    s = re.sub(r'\bunless\b', 'если не', s, flags=re.I)
    s = re.sub(r'\bdifferent\b', 'другой', s, flags=re.I)
    s = re.sub(r'\binstructions?\b', 'указание', s, flags=re.I)
    s = re.sub(r'\bprocedures?\b', 'процедура', s, flags=re.I)
    s = re.sub(r'\bup\b', 'вверх', s, flags=re.I)
    s = re.sub(r'\bdown\b', 'вниз', s, flags=re.I)
    s = re.sub(r'\bback\b', 'назад', s, flags=re.I)
    s = re.sub(r'\bgo\b', '', s, flags=re.I)
    s = re.sub(r'\binitial\b', 'начальный', s, flags=re.I)
    s = re.sub(r'\bcondition\b', 'состояние', s, flags=re.I)
    s = re.sub(r'\bthrough\b', 'через', s, flags=re.I)
    s = re.sub(r'\bproximity\b', 'датчик приближения', s, flags=re.I)
    s = re.sub(r'\bswitch\b', 'выключатель', s, flags=re.I)
    s = re.sub(r'\bpiece\b', 'деталь', s, flags=re.I)
    s = re.sub(r'\bthreads?\b', 'резьба', s, flags=re.I)
    s = re.sub(r'\bsurfaces?\b', 'поверхность', s, flags=re.I)
    s = re.sub(r'\blocating\b', 'фиксирующий', s, flags=re.I)
    s = re.sub(r'\beight\b', 'восемь', s, flags=re.I)
    s = re.sub(r'\b[Ff]igures?\b', 'Рисунок', s)
    # FITS AND CLEARANCES in all forms
    s = re.sub(r'FITSAND CLEARANCES', 'ПОСАДКИ И ЗАЗОРЫ', s)
    s = re.sub(r'FITS AND CLEARANCES', 'ПОСАДКИ И ЗАЗОРЫ', s)
    s = re.sub(r'FITS И ЗАЗОРЫ', 'ПОСАДКИ И ЗАЗОРЫ', s)
    s = re.sub(r'ПОСАДКИ И ЗАЗОРЫ И ЗАЗОРЫ', 'ПОСАДКИ И ЗАЗОРЫ', s)
    s = re.sub(r'Fits and Clearances', 'Посадки и зазоры', s)
    # Titanine JC is a brand name — keep as-is
    # More domain words
    s = re.sub(r'\block\b', 'стопорный', s, flags=re.I)
    s = re.sub(r'\bliner\b', 'вкладыш', s, flags=re.I)
    s = re.sub(r'\bsection\b', 'секция', s, flags=re.I)
    s = re.sub(r'\boutward\b', 'наружу', s, flags=re.I)
    s = re.sub(r'\ballow\b', 'дайте', s, flags=re.I)
    s = re.sub(r'\bscrew\b', 'винт', s, flags=re.I)
    s = re.sub(r'\bfluid\b', 'жидкость', s, flags=re.I)
    s = re.sub(r'\bMaterial\b', 'Материал', s)
    s = re.sub(r'\bElectrical\b', 'Электрический', s)
    s = re.sub(r'\bAxle\b', 'осевой', s)
    s = re.sub(r'\bHarness\b', 'жгут', s)
    s = re.sub(r'\bcable\b', 'кабель', s, flags=re.I)
    s = re.sub(r'\bsymmetrical\b', 'симметричный', s, flags=re.I)
    s = re.sub(r'\bsolder\b', 'припой', s, flags=re.I)
    s = re.sub(r'\btin\b', 'олово', s, flags=re.I)
    s = re.sub(r'\bBowden\b', 'Боуден', s)
    s = re.sub(r'\bAlloy\b', 'сплав', s, flags=re.I)
    s = re.sub(r'\bshims?\b', 'прокладка', s, flags=re.I)
    s = re.sub(r'\bTemporarily\b', 'Временно', s)
    s = re.sub(r'\btemporarily\b', 'временно', s)
    s = re.sub(r'\bsides?\b', 'сторона', s, flags=re.I)
    s = re.sub(r'\bcup\b', 'чашка', s, flags=re.I)
    s = re.sub(r'\bput\b', 'наносите', s, flags=re.I)
    s = re.sub(r'\bobtained\b', 'полученный', s, flags=re.I)
    s = re.sub(r'\bminus\b', 'минус', s, flags=re.I)
    s = re.sub(r'\bmedian\b', 'медиана', s, flags=re.I)
    s = re.sub(r'\btarget\b', 'мишень', s, flags=re.I)
    s = re.sub(r'\bwashers?\b', 'шайба', s, flags=re.I)
    s = re.sub(r'\bpoints?\b', 'точка', s, flags=re.I)
    s = re.sub(r'\bflats?\b', 'лыска', s, flags=re.I)
    s = re.sub(r'\bbend\b', 'загните', s, flags=re.I)
    s = re.sub(r'\bagainst\b', 'на', s, flags=re.I)
    s = re.sub(r'\bsmall\b', 'небольшое', s, flags=re.I)
    s = re.sub(r'\bquantity\b', 'количество', s, flags=re.I)
    s = re.sub(r'\blayer\b', 'слой', s, flags=re.I)
    s = re.sub(r'\bsuitable\b', 'подходящий', s, flags=re.I)
    s = re.sub(r'\bsmooth\b', 'гладкий', s, flags=re.I)
    s = re.sub(r'\bedged\b', 'кромкой', s, flags=re.I)
    s = re.sub(r'\btenons?\b', 'шип', s, flags=re.I)
    s = re.sub(r'\bengaged\b', 'зацепление', s, flags=re.I)
    s = re.sub(r'\bdry\b', 'высохнуть', s, flags=re.I)
    s = re.sub(r'\bmade\b', 'сделанный', s, flags=re.I)
    s = re.sub(r'\bpowder\b', 'порошок', s, flags=re.I)
    s = re.sub(r'\bfind\b', 'найти', s, flags=re.I)
    s = re.sub(r'\bsubassemblies\b', 'подсборки', s, flags=re.I)
    s = re.sub(r'\bLifting\b', 'подъёмный', s)
    s = re.sub(r'\bTackle\b', 'приспособление', s)
    s = re.sub(r'\bbonding\b', 'соединительный', s, flags=re.I)
    s = re.sub(r'\bcotter\b', 'шплинт', s, flags=re.I)
    s = re.sub(r'\bFREE BAND\b', 'СВОБОДНАЯ ПОЛОСА', s)
    s = re.sub(r'\bdimensions?\b', 'размер', s, flags=re.I)
    # Table/IPL specific words
    s = re.sub(r'\bMAIN LANDING GEAR LEG\b', 'СТОЙКА ОСНОВНОГО ШАССИ', s)
    s = re.sub(r'\bMain Landing Gear Leg\b', 'Стойка основного шасси', s)
    s = re.sub(r'\bmain landing gear leg\b', 'стойка основного шасси', s)
    s = re.sub(r'\bLANDING GEAR LEG\b', 'ШАССИ СТОЙКА', s)
    s = re.sub(r'\bPER ASSY\b', 'НА СБОРКУ', s)
    s = re.sub(r'\bPer Assy\b', 'На сборку', s)
    s = re.sub(r'\bGEAR\b', 'ШАССИ', s)
    s = re.sub(r'\bLEG\b', 'СТОЙКА', s)
    s = re.sub(r'\bPER\b', 'НА', s)
    s = re.sub(r'\bASSY\b', 'СБОРКА', s)
    s = re.sub(r'\bALTERNATIVE\b', 'АЛЬТЕРНАТИВА', s)
    s = re.sub(r'\bAllowable\b', 'Допустимый', s)
    s = re.sub(r'\ballowable\b', 'допустимый', s)
    s = re.sub(r'\bITEMS?\b', 'ПОЗИЦИЯ', s)
    s = re.sub(r'\bFigures?\b', 'Рисунок', s)
    s = re.sub(r'\babsorber\b', 'амортизатор', s, flags=re.I)
    s = re.sub(r'\bshock\b', 'ударный', s, flags=re.I)
    s = re.sub(r'\bgear\b', 'шасси', s, flags=re.I)
    s = re.sub(r'\bleg\b', 'стойка', s, flags=re.I)
    s = re.sub(r'\bSteel\b', 'Сталь', s)
    s = re.sub(r'\bAluminium\b', 'Алюминий', s)
    s = re.sub(r'\bBronze\b', 'Бронза', s)
    s = re.sub(r'\bStrength\b', 'Прочность', s)
    s = re.sub(r'\bSPHERICAL\b', 'СФЕРИЧЕСКИЙ', s)
    s = re.sub(r'\bUSED\b', 'ИСПОЛЬЗУЕТСЯ', s)
    # Final catch-all for common words
    s = re.sub(r'\band\b', 'и', s, flags=re.I)
    s = re.sub(r'\bor\b', 'или', s, flags=re.I)
    s = re.sub(r'\bthe\b', '', s, flags=re.I)
    s = re.sub(r'\bof\b', '', s, flags=re.I)
    s = re.sub(r'\bto\b', 'до', s, flags=re.I)
    s = re.sub(r'\bfor\b', 'для', s, flags=re.I)
    s = re.sub(r'\bwith\b', 'с', s, flags=re.I)
    s = re.sub(r'\bin\b', 'в', s, flags=re.I)
    s = re.sub(r'\bon\b', 'на', s, flags=re.I)
    # Clean up multiple spaces
    s = re.sub(r'  +', ' ', s).strip()
    return s


def translate_text(text):
    """Translate an English CMM text string to Russian.
    Handles multi-sentence paragraphs by splitting and translating each part.
    """
    if not text or not text.strip():
        return text
    t = text.strip()

    # 1. Try full text as-is first (exact match is best)
    if t in EXACT:
        return _postprocess(EXACT[t])
    tnp = t.rstrip('.')
    if tnp in EXACT:
        return _postprocess(EXACT[tnp] + ('.' if t.endswith('.') else ''))

    # 2. Pre-process: insert space before Figure/Table glued to text
    t = re.sub(r'F\.igure', 'Figure', t)  # Fix broken "F.igure" cross-refs
    t = re.sub(r'(\))\.?(Figure\s+\d+)', r'\1. \2', t)
    t = re.sub(r'(\.)(?=Figure\s+\d+)', r'\1 ', t)
    t = re.sub(r'(\))\.?(Table\s+\d+)', r'\1. \2', t)
    t = re.sub(r'([a-z)])(?=Figure\s+\d)', r'\1. ', t)
    t = re.sub(r'([A-Za-z])(?=Figure\s+\d)', r'\1. ', t)  # "CLEARANCESFigure" / "AFigure" → "... . Figure"
    t = re.sub(r'([A-Za-z])(?=Table\s+\d)', r'\1. ', t)
    # Temperature units
    t = re.sub(r'(\d)\s*oC\b', r'\1 °C', t)
    t = re.sub(r'(\d)\s*oF\b', r'\1 °F', t)

    # 3. Split into sentences FIRST, then translate each
    sentences = re.split(r'(?<=\.)\s+(?=[A-ZА-ЯЁ])', t)
    if len(sentences) > 1:
        translated_parts = []
        for s in sentences:
            # Skip standalone Figure/Table refs — they're already translated inline
            if re.match(r'^(?:Figure|Table)\s+\d+\.?$', s.strip()):
                continue
            translated_parts.append(_translate_single(s))
        result = ' '.join(translated_parts)
        return _postprocess(result)

    # 4. Single sentence: try templates, then _body
    for pat, fn in _TEMPLATES:
        m = pat.match(t)
        if m:
            try:
                result = fn(m)
                if result and result != t:
                    return _postprocess(result)
            except Exception:
                pass

    # 5. Single sentence fallback
    return _postprocess(_body(t))


# ──────────────────────────────────────────────────────────────
# DOCUMENT TRANSLATION
# ──────────────────────────────────────────────────────────────

_SKIP_EN = {
    'PCS', 'REF', 'QTY', 'TBA', 'OFF', 'REV', 'DIN', 'AMS', 'ASTM', 'MIL',
    'SAE', 'ISO', 'CMM', 'SBN', 'SB', 'AMM', 'IPC', 'SRM', 'BDM', 'NDT',
    'DLPS', 'SAFRAN', 'MESSIER', 'DOWTY', 'LANDING', 'SYSTEMS', 'LIMITED',
    'MOLYKOTE', 'MOBIL', 'TYPE', 'CLASS', 'LOCTITE', 'SERMETEL', 'MEK',
}


def _needs_translation(txt):
    words = re.findall(r'[A-Za-z]{3,}', txt)
    real = [w for w in words if w.upper() not in _SKIP_EN and
            not re.match(r'^[A-Z]{1,4}[\d\-]+$', w)]
    return len(real) >= 2


def _set_para_text(para, new_text):
    """Replace paragraph text while keeping formatting of first run.
    Also removes cross-reference field elements that cause duplicate Figure/Table refs.
    """
    from lxml import etree
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    # Remove hyperlink and fldSimple elements (direct children only)
    for child in list(para._element):
        if child.tag in (W + 'hyperlink', W + 'fldSimple'):
            para._element.remove(child)

    # Remove runs that are part of field character sequences
    # Field sequences: fldChar(begin) ... instrText ... fldChar(separate) ... result ... fldChar(end)
    in_field = False
    runs_to_remove = []
    for child in list(para._element):
        if child.tag == W + 'r':
            has_fld_char = child.find(W + 'fldChar')
            has_instr = child.find(W + 'instrText')
            if has_fld_char is not None:
                fld_type = has_fld_char.get(W + 'fldCharType', '')
                if fld_type == 'begin':
                    in_field = True
                    runs_to_remove.append(child)
                elif fld_type == 'end':
                    runs_to_remove.append(child)
                    in_field = False
                elif fld_type == 'separate':
                    runs_to_remove.append(child)
            elif has_instr is not None:
                runs_to_remove.append(child)
            elif in_field:
                runs_to_remove.append(child)
    for r in runs_to_remove:
        try:
            para._element.remove(r)
        except ValueError:
            pass

    runs = para.runs
    if not runs:
        # Fall back to direct XML — set all <w:t> elements
        first = True
        for r in para._element.iter(W + 't'):
            if first:
                r.text = new_text
                first = False
            else:
                r.text = ''
        return
    # Put everything into first run, clear the rest
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''


def translate_document(input_path, output_path):
    print(f"Loading: {input_path}")
    doc = Document(input_path)
    n_para = n_cell = 0

    # Paragraphs
    total = len(doc.paragraphs)
    for i, para in enumerate(doc.paragraphs):
        if i % 2000 == 0:
            print(f"  Para {i}/{total} (translated={n_para})...")
        txt = para.text.strip()
        if not txt or not _needs_translation(txt):
            continue
        ru = translate_text(txt)
        if ru != txt:
            _set_para_text(para, ru)
            n_para += 1

    print(f"  Translated {n_para} paragraphs")

    # Tables
    total_t = len(doc.tables)
    for t_idx, table in enumerate(doc.tables):
        if t_idx % 60 == 0:
            print(f"  Table {t_idx}/{total_t}...")
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if not txt or not _needs_translation(txt):
                    continue
                ru = translate_text(txt)
                if ru != txt:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            _set_para_text(para, ru)
                            ru = ''  # only first para
                            n_cell += 1
                            break

    print(f"  Translated {n_cell} table cells")
    print(f"Saving: {output_path}")
    doc.save(output_path)
    print("Done!")
    return n_para, n_cell


if __name__ == '__main__':
    src = r'C:\Users\Urdul\Desktop\project\translator\for_test\new_formating\section_translate\section_2_opus\part_2_orig_claude.docx'
    dst = r'C:\Users\Urdul\Desktop\project\translator\for_test\new_formating\section_translate\section_2_opus\part_2_ru_draft.docx'
    translate_document(src, dst)

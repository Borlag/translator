from __future__ import annotations

import argparse
import json
import logging
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.run import Run

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from src.docxru.config import load_config
from src.docxru.llm import build_llm_client
from src.docxru.pipeline import _read_optional_text

_TOOL_HEADERS = {"special tool", "equipment"}
_FUNCTION_HEADERS = {"function"}


def _normalize_header(text: str) -> str:
    return " ".join((text or "").replace("\n", " ").split()).strip().lower()


def _capture_run_style(run: Run | None) -> dict[str, Any]:
    if run is None:
        return {}
    style: dict[str, Any] = {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
    }
    try:
        style["font_name"] = run.font.name
        style["font_size"] = run.font.size
        style["font_color"] = run.font.color.rgb
    except Exception:
        pass
    return style


def _apply_run_style(run: Run, style: dict[str, Any]) -> None:
    if not style:
        return
    run.bold = style.get("bold")
    run.italic = style.get("italic")
    run.underline = style.get("underline")
    try:
        run.font.name = style.get("font_name")
        run.font.size = style.get("font_size")
        if style.get("font_color") is not None:
            run.font.color.rgb = style.get("font_color")
    except Exception:
        return


def _set_paragraph_text(paragraph, text: str) -> None:
    first_run = paragraph.runs[0] if paragraph.runs else None
    style = _capture_run_style(first_run)
    for child in list(paragraph._p.iterchildren()):
        local = child.tag.split("}")[-1].lower()
        if local in {"r", "hyperlink"}:
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    _apply_run_style(run, style)


def _coerce_response_map(raw: str) -> dict[str, dict[str, str]]:
    payload = json.loads(raw)
    if isinstance(payload, dict) and isinstance(payload.get("translations"), list):
        entries = payload["translations"]
    elif isinstance(payload, list):
        entries = payload
    else:
        entries = [payload]

    out: dict[str, dict[str, str]] = {}
    for item in entries:
        if not isinstance(item, dict):
            continue
        item_id = str(item.get("id") or "").strip()
        tool_ru = str(item.get("tool_ru") or "").strip()
        function_ru = str(item.get("function_ru") or "").strip()
        if not item_id:
            continue
        out[item_id] = {"tool_ru": tool_ru, "function_ru": function_ru}
    return out


def _build_client(config_path: Path):
    cfg = load_config(config_path)
    logger = logging.getLogger("fix_tool_tables")
    system_prompt = _read_optional_text(cfg.llm.system_prompt_path, logger, "system prompt")
    glossary_text = _read_optional_text(cfg.llm.glossary_path, logger, "glossary")
    client = build_llm_client(
        provider=cfg.llm.provider,
        model=cfg.llm.model,
        temperature=cfg.llm.temperature,
        timeout_s=cfg.llm.timeout_s,
        max_output_tokens=cfg.llm.max_output_tokens,
        source_lang=cfg.llm.source_lang,
        target_lang=cfg.llm.target_lang,
        base_url=cfg.llm.base_url,
        custom_system_prompt=system_prompt,
        glossary_text=glossary_text,
        glossary_prompt_text=glossary_text if cfg.llm.glossary_in_prompt else None,
        prompt_examples_mode=cfg.llm.prompt_examples_mode,
        reasoning_effort=cfg.llm.reasoning_effort,
        prompt_cache_key=cfg.llm.prompt_cache_key,
        prompt_cache_retention=cfg.llm.prompt_cache_retention,
        structured_output_mode="strict",
    )
    return cfg, client


def _iter_table_rows(source_doc: Document, target_doc: Document) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for table_idx, (source_table, target_table) in enumerate(zip(source_doc.tables, target_doc.tables)):
        if not source_table.rows or not target_table.rows:
            continue
        source_headers = [_normalize_header(cell.text) for cell in source_table.rows[0].cells]
        if len(source_headers) < 3:
            continue
        if source_headers[1] not in _TOOL_HEADERS or source_headers[2] not in _FUNCTION_HEADERS:
            continue

        row_count = min(len(source_table.rows), len(target_table.rows))
        for row_idx in range(1, row_count):
            source_row = source_table.rows[row_idx]
            target_row = target_table.rows[row_idx]
            if len(source_row.cells) < 3 or len(target_row.cells) < 3:
                continue

            source_tool = source_row.cells[1].text.strip()
            source_function = source_row.cells[2].text.strip()
            current_tool = target_row.cells[1].text.strip()
            current_function = target_row.cells[2].text.strip()
            if not source_tool and not source_function:
                continue

            items.append(
                {
                    "id": f"t{table_idx}r{row_idx}",
                    "table_idx": table_idx,
                    "row_idx": row_idx,
                    "part_no": source_row.cells[0].text.strip(),
                    "source_tool": source_tool,
                    "source_function": source_function,
                    "current_tool_ru": current_tool,
                    "current_function_ru": current_function,
                }
            )
    return items


def _build_prompt(batch: list[dict[str, Any]]) -> str:
    return (
        "TASK: FIX_TOOL_TABLE_ROWS\n"
        "You correct Russian translations in aviation maintenance tables with columns "
        "Part No., Special Tool/Equipment, Function.\n"
        'Return ONLY valid JSON object {"translations":[{"id":"...","tool_ru":"...","function_ru":"..."}]}.\n'
        "Rules:\n"
        "- Use source_tool and source_function as the meaning ground truth.\n"
        "- Translate Special Tool/Equipment as a proper Russian technical tool name.\n"
        "- If source_tool contains line breaks, treat them as layout only and translate the whole term semantically.\n"
        "- Translate Function as a concise table label for the tool purpose, not as a raw literal imperative step.\n"
        "- Preferred Function patterns:\n"
        "  Hold the pin -> Для удержания штифта\n"
        "  Remove the bush -> Для снятия втулки\n"
        "  Remove the bushes -> Для снятия втулок\n"
        "  Use with 460006406 -> Использовать с 460006406\n"
        "  Main landing gear leg (1-1) tests -> Испытания стойки основного шасси (1-1)\n"
        "- Preserve part numbers, references like (9-70), and identifiers exactly.\n"
        "- Keep wording short enough for a table cell.\n"
        "- Return plain Russian text only, without explanations.\n\n"
        "INPUT_JSON:\n"
        + json.dumps(batch, ensure_ascii=False)
    )


def _apply_batch(target_doc: Document, batch: list[dict[str, Any]], translations: dict[str, dict[str, str]]) -> int:
    applied = 0
    for item in batch:
        translated = translations.get(item["id"])
        if translated is None:
            continue
        row = target_doc.tables[item["table_idx"]].rows[item["row_idx"]]
        tool_para = row.cells[1].paragraphs[0] if row.cells[1].paragraphs else row.cells[1].add_paragraph()
        function_para = row.cells[2].paragraphs[0] if row.cells[2].paragraphs else row.cells[2].add_paragraph()
        _set_paragraph_text(tool_para, translated.get("tool_ru", "") or item["current_tool_ru"])
        _set_paragraph_text(function_para, translated.get("function_ru", "") or item["current_function_ru"])
        applied += 1
    return applied


def main() -> None:
    parser = argparse.ArgumentParser(description="Fix tool/function table translations in a translated DOCX.")
    parser.add_argument("--source", required=True, help="Path to the original English DOCX.")
    parser.add_argument("--translated", required=True, help="Path to the translated Russian DOCX.")
    parser.add_argument("--output", required=True, help="Path to save the corrected DOCX.")
    parser.add_argument("--config", default="config/config.agent_openai.yaml", help="Path to docxru YAML config.")
    parser.add_argument("--batch-size", type=int, default=8, help="Rows per LLM request.")
    args = parser.parse_args()

    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")
    _, client = _build_client(Path(args.config))

    source_doc = Document(args.source)
    target_doc = Document(args.translated)
    items = _iter_table_rows(source_doc, target_doc)
    logging.info("Found %d candidate tool/function rows.", len(items))

    total_applied = 0
    batch_size = max(1, int(args.batch_size))
    for start in range(0, len(items), batch_size):
        batch = items[start : start + batch_size]
        raw = client.translate(_build_prompt(batch), {"task": "batch_translate"})
        translations = _coerce_response_map(raw)
        applied = _apply_batch(target_doc, batch, translations)
        total_applied += applied
        logging.info(
            "Applied batch %d-%d (%d rows).",
            start + 1,
            start + len(batch),
            applied,
        )

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    target_doc.save(output_path)
    logging.info("Saved corrected DOCX: %s (rows updated: %d)", output_path, total_applied)


if __name__ == "__main__":
    main()

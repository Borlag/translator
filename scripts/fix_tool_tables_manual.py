from __future__ import annotations

import argparse
from pathlib import Path
import zipfile

from docx import Document
from docx.text.run import Run
from lxml import etree

_TOOL_HEADERS = {"special tool", "equipment"}
_FUNCTION_HEADERS = {"function"}
_WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

_TOOL_TRANSLATIONS = {
    "28 vdc power supply": "Источник питания 28 В пост. тока",
    "adapter": "Адаптер",
    "alignment bar": "Центровочная штанга",
    "assembly sleeve": "Монтажная втулка",
    "assembly tool": "Монтажный инструмент",
    "assembly/extraction tool": "Монтажно-демонтажный инструмент",
    "bench clamp": "Верстачный зажим",
    "bolt": "Болт",
    "bottom press adapter": "Нижний адаптер для пресса",
    "build trolley": "Монтажная тележка",
    "charging adapter": "Заправочный адаптер",
    "crowfoot wrench": "Рожковая насадка",
    "cutter": "Резец",
    "drift": "Выколотка",
    "extraction bar": "Выпрессовочная штанга",
    "extraction pad": "Выпрессовочная проставка",
    "extraction tube": "Выпрессовочная трубка",
    "extractor": "Съемник",
    "extractor pad and drawbolt": "Выпрессовочная проставка и стяжной болт",
    "extractor plate": "Пластина съемника",
    "guide bush": "Направляющая втулка",
    "holding blocks": "Удерживающие блоки",
    "holding fixture": "Удерживающее приспособление",
    "hydraulic test rig": "Стенд гидравлических испытаний",
    "hydraulic-pneumatic pump set": "Гидропневматический насосный комплект",
    "jacking dome adapter": "Адаптер опоры домкрата",
    "lampbox": "Ламповый блок",
    "lifting bar assembly": "Подъемная траверса",
    "lifting tackle": "Подъемная оснастка",
    "load cell and adapter": "Датчик нагрузки и адаптер",
    "loading press": "Пресс для нагружения",
    "location frame": "Установочная рама",
    "milliohmmeter megger,": "Миллиомметр Megger,",
    "nitrogen supply": "Система подачи азота",
    "offset adapter": "Адаптер со смещением",
    "pin spanner": "Штифтовый ключ",
    "pintle location assembly": "Установочная сборка штифта навеса",
    "press adapter": "Адаптер для пресса",
    "press pad": "Нажимная проставка",
    "press pad and drawbolt": "Нажимная проставка и стяжной болт",
    "press pad assembly": "Нажимная проставка в сборе",
    "spacer": "Проставка",
    "spherical bearing locator": "Локатор сферического подшипника",
    "support arms": "Опорные рычаги",
    "torque adapter": "Адаптер для динамометрического ключа",
    "torque reaction adapter": "Адаптер реактивного момента",
    "torque reactor": "Реактивная опора",
    "towing frame": "Буксировочная рама",
    "transport and build trolley": "Транспортно-монтажная тележка",
    "turner inflation equipment": "Оборудование Turner для заправки азотом",
    "type bt51": "Тип BT51",
}

_EXACT_FUNCTION_TRANSLATIONS = {
    "": "",
    "(2-250a)": "(2-250A)",
    "electrical bonding resistance": "Измерение сопротивления электрического соединения",
    "main landing gear leg (1-1) tests": "Испытания стойки основного шасси (1-1)",
    "proximity switch and target tests": "Испытания датчика приближения и мишени",
    "tests": "",
    "hold the main landing gear leg (1-1),": "Удерживать стойку основного шасси (1-1)",
    "hold the main landing gear leg (1-1), use with 460007281 and 460007282": (
        "Удерживать стойку основного шасси (1-1); использовать с 460007281 и 460007282"
    ),
    "to get the correct dimension across the repair bushes": "Обеспечить требуемый размер по ремонтным втулкам",
    "finish machine the repair bush 450258806": (
        "Выполнить окончательную механическую обработку ремонтной втулки 450258806"
    ),
    "use with press pad 460006600": "Использовать с нажимной проставкой 460006600",
    "use with press pad 460006603": "Использовать с нажимной проставкой 460006603",
}

_OBJECT_TRANSLATIONS = {
    "bearing": "подшипник",
    "bearing (20-260)": "подшипник (20-260)",
    "bearing (20-270)": "подшипник (20-270)",
    "bearing (20-280)": "подшипник (20-280)",
    "bearing (20-290)": "подшипник (20-290)",
    "bearing (20-300)": "подшипник (20-300)",
    "bearing (20-310)": "подшипник (20-310)",
    "bearing (4-340) and the bush (20-360)": "подшипник (4-340) и втулку (20-360)",
    "bearing (4-350)": "подшипник (4-350)",
    "bearings (20-230, 20-240 and 20-290)": "подшипники (20-230, 20-240 и 20-290)",
    "bearings (5-280 and 5-290) and the bushes (20-380)": "подшипники (5-280 и 5-290) и втулки (20-380)",
    "blank bush 450237817": "втулку-заготовку 450237817",
    "bush (15-370)": "втулку (15-370)",
    "bush (15-380)": "втулку (15-380)",
    "bush (18-50)": "втулку (18-50)",
    "bush (2-310)": "втулку (2-310)",
    "bush (2-320)": "втулку (2-320)",
    "bush (6-210 and 8-160)": "втулку (6-210 и 8-160)",
    "bush (6-220 and 8-150)": "втулку (6-220 и 8-150)",
    "bush 450237800": "втулку 450237800",
    "bush made from 450237811": "втулку, изготовленную из 450237811",
    "bushes (10-250, 11-230, 18-40 and 20-330)": "втулки (10-250, 11-230, 18-40 и 20-330)",
    "bushes (18-20)": "втулки (18-20)",
    "bushes (18-30)": "втулки (18-30)",
    "bushes (20-320)": "втулки (20-320)",
    "bushes (20-340 and 20-350)": "втулки (20-340 и 20-350)",
    "bushes (20-340) and (20-350)": "втулки (20-340) и (20-350)",
    "bushes (20-390)": "втулки (20-390)",
    "bushes (7-130)": "втулки (7-130)",
    "charging valves (13-60 and 17-20)": "заправочные клапаны (13-60 и 17-20)",
    "cylinder (17-230)": "цилиндр (17-230)",
    "diaphragm subassembly (15-190)": "узел диафрагмы (15-190)",
    "drag arm sleeve (20-370a)": "втулку тягового рычага (20-370A)",
    "forward pintle bush": "переднюю втулку штифта навеса",
    "forward pintle bush (20-250a)": "переднюю втулку штифта навеса (20-250A)",
    "jacking dome (17-80)": "опору домкрата (17-80)",
    "level tube (15-300)": "трубку уровня (15-300)",
    "locking nut (19-52)": "стопорную гайку (19-52)",
    "lower bearing subassembly": "нижний узел подшипника",
    "lubrication adapters (18-60), (20-130), (20-160), (20-190) and (20-220)": (
        "смазочные адаптеры (18-60), (20-130), (20-160), (20-190) и (20-220)"
    ),
    "main landing gear leg (1-1)": "стойку основного шасси (1-1)",
    "main landing gear leg (1-1) (left configuration)": (
        "стойку основного шасси (1-1) (левая конфигурация)"
    ),
    "main landing gear leg (1-2) (right configuration)": (
        "стойку основного шасси (1-2) (правая конфигурация)"
    ),
    "nut (9-50)": "гайку (9-50)",
    "nut subassembly (17-130)": "узел гайки (17-130)",
    "nuts (14-60)": "гайки (14-60)",
    "oversize bearing": "подшипник увеличенного размера",
    "oversize bearing 450258809": "подшипник увеличенного размера 450258809",
    "oversize bearing 450258810": "подшипник увеличенного размера 450258810",
    "oversize bush(es)": "втулку(и) увеличенного размера",
    "oversize bushes": "втулки увеличенного размера",
    "oversize lubrication adaptor": "смазочный адаптер увеличенного размера",
    "oversize lubrication adapters": "смазочные адаптеры увеличенного размера",
    "pin (9-70)": "штифт (9-70)",
    "repair bearing": "ремонтный подшипник",
    "repair bearing 450266081": "ремонтный подшипник 450266081",
    "repair bush": "ремонтную втулку",
    "repair bush 450217851": "ремонтную втулку 450217851",
    "repair bush 450217852": "ремонтную втулку 450217852",
    "repair bush 450237810": "ремонтную втулку 450237810",
    "repair bush 450237819": "ремонтную втулку 450237819",
    "repair bush 450258800": "ремонтную втулку 450258800",
    "repair bush 450258806": "ремонтную втулку 450258806",
    "repair bush 450258811": "ремонтную втулку 450258811",
    "repair bush 450258812": "ремонтную втулку 450258812",
    "repair bush 450266800": "ремонтную втулку 450266800",
    "repair bush(es)": "ремонтную втулку(и)",
    "repair bushes": "ремонтные втулки",
    "sliding tube subassembly (17-240) and related parts": (
        "узел скользящей трубы (17-240) и связанные детали"
    ),
    "upper bearing housing (15-40)": "корпус верхнего подшипника (15-40)",
    "upper diaphragm tube (15-390)": "верхнюю диафрагменную трубку (15-390)",
}

_BODY_TEXT_REPLACEMENTS: tuple[tuple[str, str], ...] = (
    ("адаптер противодействия крутящему моменту", "адаптер реактивного момента"),
    ("адаптер крутящего момента", "адаптер для динамометрического ключа"),
    ("адаптер для момента затяжки", "адаптер для динамометрического ключа"),
    ("адаптер момента затяжки", "адаптер для динамометрического ключа"),
    ("реактор момента затяжки", "реактивная опора"),
    ("реактор крутящего момента", "реактивная опора"),
    ("инструмент сборки/извлечения", "монтажно-демонтажный инструмент"),
    ("гидропневматическую насосную установку", "гидропневматический насосный комплект"),
    ("сборку прижимной подушки", "нажимную проставку в сборе"),
    ("сборку прижимной пластины", "нажимную проставку в сборе"),
    ("пресс накладку", "нажимную проставку"),
    ("пресс-накладку", "нажимную проставку"),
    ("экстракционную трубку", "выпрессовочную трубку"),
    ("переднюю штифтовую втулку", "переднюю втулку штифта навеса"),
    ("извлекательную подушку", "выпрессовочную проставку"),
    ("извлекательную балку", "выпрессовочную штангу"),
    ("извлекательный стержень", "выпрессовочную штангу"),
    ("извлекатель ", "съемник "),
    ("экстрактор ", "съемник "),
    ("настольный зажим", "верстачный зажим"),
    ("стендовом хомуте", "верстачном зажиме"),
    ("зажима стенда", "верстачного зажима"),
    (
        "адаптер для динамометрического ключа 460007283, реактивная опора 460007278",
        "адаптер для динамометрического ключа 460007283, реактивную опору 460007278",
    ),
    ("домкратный купол", "опору домкрата"),
    ("подузла диафрагмы", "узла диафрагмы"),
    ("пластина компрессионного жиклёра", "пластину компрессионного жиклера"),
    ("перегородка (", "перегородку ("),
    ("(20-370A только)", "(только 20-370A)"),
    ("0,250 mm", "0,250 мм"),
    ("0.010 in", "0,010 дюйма"),
)


def _normalize_space(text: str) -> str:
    return " ".join((text or "").replace("\n", " ").split()).strip()


def _normalize_key(text: str) -> str:
    return _normalize_space(text).lower()


def _capture_run_style(run: Run | None) -> dict[str, object]:
    if run is None:
        return {}
    style: dict[str, object] = {"bold": run.bold, "italic": run.italic, "underline": run.underline}
    try:
        style["font_name"] = run.font.name
        style["font_size"] = run.font.size
    except Exception:
        pass
    return style


def _apply_run_style(run: Run, style: dict[str, object]) -> None:
    if not style:
        return
    run.bold = style.get("bold")
    run.italic = style.get("italic")
    run.underline = style.get("underline")
    try:
        run.font.name = style.get("font_name")
        run.font.size = style.get("font_size")
    except Exception:
        return


def _set_cell_text(cell, text: str) -> None:
    first_paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    first_style = first_paragraph.style
    first_alignment = first_paragraph.alignment
    first_run = first_paragraph.runs[0] if first_paragraph.runs else None
    run_style = _capture_run_style(first_run)

    for paragraph in list(cell.paragraphs)[1:]:
        p = paragraph._element
        p.getparent().remove(p)

    for child in list(first_paragraph._p.iterchildren()):
        local = child.tag.split("}")[-1].lower()
        if local in {"r", "hyperlink"}:
            first_paragraph._p.remove(child)

    new_run = first_paragraph.add_run(text)
    _apply_run_style(new_run, run_style)
    first_paragraph.style = first_style
    first_paragraph.alignment = first_alignment


def _paragraph_text_nodes(paragraph: etree._Element) -> list[etree._Element]:
    return paragraph.xpath(".//w:t", namespaces=_WORD_NS)


def _set_xml_text(nodes: list[etree._Element], value: str) -> None:
    if not nodes:
        return
    first = nodes[0]
    first.text = value
    if value[:1].isspace() or value[-1:].isspace():
        first.set(_XML_SPACE, "preserve")
    elif _XML_SPACE in first.attrib:
        del first.attrib[_XML_SPACE]
    for node in nodes[1:]:
        node.text = ""
        if _XML_SPACE in node.attrib:
            del node.attrib[_XML_SPACE]


def _rewrite_body_text(docx_path: Path) -> int:
    changed = 0
    tmp_path = docx_path.with_suffix(".tmp.docx")

    with zipfile.ZipFile(docx_path) as src, zipfile.ZipFile(tmp_path, "w") as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename != "word/document.xml":
                dst.writestr(item, data)
                continue

            root = etree.fromstring(data)
            for paragraph in root.xpath(".//w:p", namespaces=_WORD_NS):
                nodes = _paragraph_text_nodes(paragraph)
                if not nodes:
                    continue
                current = "".join(node.text or "" for node in nodes)
                updated = current
                for source_text, target_text in _BODY_TEXT_REPLACEMENTS:
                    updated = updated.replace(source_text, target_text)
                if updated != current:
                    _set_xml_text(nodes, updated)
                    changed += 1
            dst.writestr(item, etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes"))

    tmp_path.replace(docx_path)
    return changed


def _translate_tool(text: str) -> str:
    key = _normalize_key(text)
    if key not in _TOOL_TRANSLATIONS:
        raise KeyError(f"Missing tool translation: {text!r}")
    return _TOOL_TRANSLATIONS[key]


def _translate_use_rest(text: str) -> str:
    key = _normalize_key(text)
    if key == "press pad 460006600":
        return "нажимной проставкой 460006600"
    if key == "press pad 460006603":
        return "нажимной проставкой 460006603"
    return text.replace("\n", " ").replace(" and ", " и ")


def _translate_object(text: str) -> str:
    key = _normalize_key(text)
    if key not in _OBJECT_TRANSLATIONS:
        raise KeyError(f"Missing function object translation: {text!r}")
    return _OBJECT_TRANSLATIONS[key]


def _translate_function(text: str) -> str:
    flat = _normalize_space(text)
    key = flat.lower()
    if key in _EXACT_FUNCTION_TRANSLATIONS:
        return _EXACT_FUNCTION_TRANSLATIONS[key]

    patterns = (
        ("use with ", lambda rest: f"Использовать с {_translate_use_rest(rest)}"),
        ("hold the ", lambda rest: f"Удерживать {_translate_object(rest)}"),
        ("remove the ", lambda rest: f"Снять {_translate_object(rest)}"),
        ("remove ", lambda rest: f"Снять {_translate_object(rest)}"),
        ("install the ", lambda rest: f"Установить {_translate_object(rest)}"),
        ("install ", lambda rest: f"Установить {_translate_object(rest)}"),
        ("close the ", lambda rest: f"Закрыть {_translate_object(rest)}"),
        ("lift the ", lambda rest: f"Поднять {_translate_object(rest)}"),
        ("to remove the ", lambda rest: f"Снять {_translate_object(rest)}"),
        ("to remove ", lambda rest: f"Снять {_translate_object(rest)}"),
        ("to install the ", lambda rest: f"Установить {_translate_object(rest)}"),
        ("to install ", lambda rest: f"Установить {_translate_object(rest)}"),
    )
    for prefix, builder in patterns:
        if key.startswith(prefix):
            return builder(flat[len(prefix) :])

    raise KeyError(f"Missing function translation: {text!r}")


def _iter_target_rows(source_doc: Document, target_doc: Document):
    for table_idx, (source_table, target_table) in enumerate(zip(source_doc.tables, target_doc.tables)):
        if not source_table.rows or not target_table.rows or len(source_table.rows[0].cells) < 3:
            continue
        headers = [_normalize_key(cell.text) for cell in source_table.rows[0].cells[:3]]
        if headers[1] not in _TOOL_HEADERS or headers[2] not in _FUNCTION_HEADERS:
            continue
        row_count = min(len(source_table.rows), len(target_table.rows))
        for row_idx in range(1, row_count):
            source_row = source_table.rows[row_idx]
            target_row = target_table.rows[row_idx]
            if len(source_row.cells) < 3 or len(target_row.cells) < 3:
                continue
            yield table_idx, row_idx, source_row, target_row


def main() -> None:
    parser = argparse.ArgumentParser(description="Deterministic manual fix for tool/function tables.")
    parser.add_argument("--source", required=True)
    parser.add_argument("--translated", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    source_doc = Document(args.source)
    target_doc = Document(args.translated)

    updated_rows = 0
    for table_idx, row_idx, source_row, target_row in _iter_target_rows(source_doc, target_doc):
        source_tool = source_row.cells[1].text
        source_function = source_row.cells[2].text

        tool_ru = _translate_tool(source_tool)
        function_ru = _translate_function(source_function)

        if table_idx == 34 and row_idx == 4:
            function_ru = "Испытания сопротивления электрического соединения"
        if table_idx == 34 and row_idx == 5:
            function_ru = ""

        _set_cell_text(target_row.cells[1], tool_ru)
        _set_cell_text(target_row.cells[2], function_ru)
        updated_rows += 1

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    target_doc.save(output_path)
    changed_paragraphs = _rewrite_body_text(output_path)
    print(f"Saved {output_path} (updated rows: {updated_rows}, updated paragraphs: {changed_paragraphs})")


if __name__ == "__main__":
    main()

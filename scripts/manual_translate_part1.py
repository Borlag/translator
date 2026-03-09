from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import tempfile
from collections import Counter
from dataclasses import replace
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCXRU_SRC = ROOT / ".claude" / "worktrees" / "competent-jones" / "src"
if str(DOCXRU_SRC) not in sys.path:
    sys.path.insert(0, str(DOCXRU_SRC))

from docxru.com_word import update_fields_via_com
from docxru.config import PipelineConfig
from docxru.docx_reader import collect_segments
from docxru.layout_check import validate_layout
from docxru.layout_fix import fix_expansion_issues
from docxru.tagging import _apply_style, _clear_paragraph_runs, paragraph_to_tagged


LATIN_RE = re.compile(r"[A-Za-z]")
CODE_LIKE_RE = re.compile(
    r"^(?:"
    r"[0-9.,/ -]+|"
    r"\d+(?:-\d+)+(?:[A-Z])?|"
    r"[A-Z]{1,6}[0-9A-Z./-]*|"
    r"\([a-z]\)[A-Z0-9-]+|"
    r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{1,2}/\d{4}|"
    r"Page\s+\d+\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{1,2}/\d{4}|"
    r"[0-9]+\s+to\s+[0-9.]+(?:\s+and\s+[0-9.]+)?|"
    r"[0-9]+\s+and\s+[0-9]+|"
    r"[0-9]+\s+\t\s+[0-9A-Za-z./ -]+"
    r")$"
)
RESERVED_UNCHANGED = {
    "©",
    "SAFRAN",
    "SAFRAN LANDING SYSTEMS",
    "CAGE: K0654",
    "M-D",
    "Safran Landing Systems UK Ltd",
    "SAFRAN LANDING SYSTEMS UK Ltd",
    "Cheltenham Road, Gloucester, GL2 9QH, England",
    "EDES2-0005-2253",
    "EDES2-0005-2255",
    "EDES2-0005-2542",
    "MAF1-0005-2155",
    "MAF1-0005-2156",
    "MAF1-0005-2157",
    "MAF1-0005-2160",
    "MAF1-0005-2161",
    "MAF1-0005-2162",
    "MAF1-0005-2164",
    "MAF1-0005-2180",
    "MAF1-0005-2207",
    "Mastinox D40",
    "Molykote 111",
}

MONTH_MAP: dict[str, str] = {
    "Jan": "янв.",
    "Feb": "февр.",
    "Mar": "мар.",
    "Apr": "апр.",
    "May": "мая",
    "Jun": "июн.",
    "Jul": "июл.",
    "Aug": "авг.",
    "Sep": "сент.",
    "Oct": "окт.",
    "Nov": "нояб.",
    "Dec": "дек.",
}

MONTH_NUMBER: dict[str, str] = {
    "Jan": "01",
    "Feb": "02",
    "Mar": "03",
    "Apr": "04",
    "May": "05",
    "Jun": "06",
    "Jul": "07",
    "Aug": "08",
    "Sep": "09",
    "Oct": "10",
    "Nov": "11",
    "Dec": "12",
}


EXACT_MAP: dict[str, str] = {
    "MAIN LANDING GEAR LEG": "СТОЙКА ОСНОВНОГО ШАССИ",
    "PART NUMBER": "НОМЕР ДЕТАЛИ",
    "COMPONENT MAINTENANCE MANUAL WITH": "РУКОВОДСТВО ПО ТЕХНИЧЕСКОМУ ОБСЛУЖИВАНИЮ КОМПОНЕНТА С",
    "ILLUSTRATED PARTS LIST": "ИЛЛЮСТРИРОВАННЫЙ ПЕРЕЧЕНЬ ДЕТАЛЕЙ",
    "STATEMENT OF INITIAL CERTIFICATION": "ЗАЯВЛЕНИЕ О ПЕРВОНАЧАЛЬНОЙ СЕРТИФИКАЦИИ",
    "This manual complies with British Civil Airworthiness Requirements, Section A, Chapter A5-3.": (
        "Настоящее руководство соответствует Британским требованиям летной годности гражданской авиации, раздел A, глава A5-3."
    ),
    "NOTE: The above certification does not apply to revisions or amendments made after the date of initial certification by other Approved Organisations. Revisions or Amendments made by other Approved Organisations must each be separately certified and recorded on separate record sheets.": (
        "ПРИМЕЧАНИЕ. Указанная выше сертификация не распространяется на ревизии или изменения, внесенные после даты первоначальной сертификации другими одобренными организациями. Каждая ревизия или изменение, выполненные другими одобренными организациями, должны быть сертифицированы отдельно и зарегистрированы на отдельных листах учета."
    ),
    "SAFRAN LANDING SYSTEMS 2016 (AND SUBSEQUENT REVISION PAGE DATES)": "SAFRAN LANDING SYSTEMS 2016 (И ПОСЛЕДУЮЩИЕ ДАТЫ РЕВИЗИОННЫХ СТРАНИЦ)",
    "SAFRAN LANDING SYSTEMS 2025 (AND SUBSEQUENT REVISION PAGE DATES)": "SAFRAN LANDING SYSTEMS 2025 (И ПОСЛЕДУЮЩИЕ ДАТЫ РЕВИЗИОННЫХ СТРАНИЦ)",
    "© SAFRAN LANDING SYSTEMS 2016 (AND SUBSEQUENT REVISION PAGE DATES)": "© SAFRAN LANDING SYSTEMS 2016 (И ПОСЛЕДУЮЩИЕ ДАТЫ РЕВИЗИОННЫХ СТРАНИЦ)",
    "© SAFRAN LANDING SYSTEMS 2025 (AND SUBSEQUENT REVISION PAGE DATES)": "© SAFRAN LANDING SYSTEMS 2025 (И ПОСЛЕДУЮЩИЕ ДАТЫ РЕВИЗИОННЫХ СТРАНИЦ)",
    "This document and all information contained herein is the sole property of Safran Landing Systems (and/or its affiliated companies).": (
        "Настоящий документ и вся содержащаяся в нем информация являются исключительной собственностью Safran Landing Systems (и/или ее аффилированных компаний)."
    ),
    "No intellectual property rights are granted by the delivery of this document or the disclosure of its content. This document shall not be reproduced to a third party without the express written consent of Safran Landing Systems (and/or the appropriate affiliated company).": (
        "Передача настоящего документа или раскрытие его содержания не предоставляет никаких прав интеллектуальной собственности. Воспроизведение документа для третьих лиц допускается только при наличии прямого письменного согласия Safran Landing Systems (и/или соответствующей аффилированной компании)."
    ),
    "PART No. 201587001 AND 201587002 COMPONENT MAINTENANCE MANUAL MAIN LANDING GEAR LEG": (
        "ДЕТАЛЬ № 201587001 И 201587002 РУКОВОДСТВО ПО ТЕХНИЧЕСКОМУ ОБСЛУЖИВАНИЮ КОМПОНЕНТА СТОЙКА ОСНОВНОГО ШАССИ"
    ),
    "INTENTIONALLY BLANK": "ПРЕДНАМЕРЕННО ОСТАВЛЕНО ПУСТЫМ",
    "NEW/REVISED PAGES": "НОВЫЕ/ПЕРЕСМОТРЕННЫЕ СТРАНИЦЫ",
    "REVISION RECORD": "ЖУРНАЛ РЕВИЗИЙ",
    "RECORD OF REVISIONS": "ЖУРНАЛ РЕВИЗИЙ",
    "RECORD OF TEMPORARY REVISIONS": "ЖУРНАЛ ВРЕМЕННЫХ РЕВИЗИЙ",
    "LIST OF SERVICE BULLETINS": "СПИСОК СЕРВИСНЫХ БЮЛЛЕТЕНЕЙ",
    "LIST OF EFFECTIVE PAGES": "ПЕРЕЧЕНЬ ДЕЙСТВУЮЩИХ СТРАНИЦ",
    "TABLE OF CONTENTS": "СОДЕРЖАНИЕ",
    "ILLUSTRATIONS": "ИЛЛЮСТРАЦИИ",
    "Record the issue date and insertion date of this revision in the Record of Revisions and retain this Letter of Transmittal.": (
        "Зарегистрируйте дату выпуска и дату внесения настоящей ревизии в журнале ревизий и сохраните данное сопроводительное письмо."
    ),
    "Issued by": "ВЫПУЩЕНО",
    "Cheltenham Road, Gloucester, GL2 9QH, England Telephone: +44 (0) 1452 712424 Fax: +44 (0) 1452 713821": (
        "Cheltenham Road, Gloucester, GL2 9QH, England. Тел.: +44 (0) 1452 712424. Факс: +44 (0) 1452 713821"
    ),
    "Telephone: +44 (0) 1452 712424 www.safran-landing-systems.com": "Тел.: +44 (0) 1452 712424  www.safran-landing-systems.com",
    "UNIT IDENTIFICATION CHART": "ТАБЛИЦА ИДЕНТИФИКАЦИИ АГРЕГАТА",
    "COMPONENT MAINTENANCE MANUAL 32-12-22 MAIN LANDING GEAR LEG": (
        "РУКОВОДСТВО ПО ТЕХНИЧЕСКОМУ ОБСЛУЖИВАНИЮ КОМПОНЕНТА 32-12-22 СТОЙКА ОСНОВНОГО ШАССИ"
    ),
    "Fig.\tPage": "Рис.\tСтр.",
    "Fig.	Page": "Рис.\tСтр.",
    "Fig. Page": "Рис.\tСтр.",
    "Рис.\tPage": "Рис.\tСтр.",
    "Рис. Page": "Рис.\tСтр.",
    "Page": "Стр.",
    "Date": "Дата",
    "Blank": "Пусто",
    "604 Blank": "604\tПусто",
    "Pages (Continued)": "Страницы (Продолжение)",
    "Pages (Продолжение)": "Страницы (Продолжение)",
    "Initial Issue": "Первоначальный выпуск",
    "No effect": "Без влияния",
    "SURFACE": "ПОВЕРХНОСТИ",
    "DEPOSIT": "НАПЛЫВ",
    "FULL CHROME": "СПЛОШНОЕ ХРОМОВОЕ ПОКРЫТИЕ",
    "FULL CHROME PLATING THICKNESS": "ТОЛЩИНА СПЛОШНОГО ХРОМОВОГО ПОКРЫТИЯ",
    "PLATING THICKNESS": "ТОЛЩИНА ПОКРЫТИЯ",
    "PLATING LIMIT": "ГРАНИЦА ПОКРЫТИЯ",
    "EXTERNAL THICK ZINC-NICKEL": "НАРУЖНОЕ ТОЛСТОЕ ЦИНК-НИКЕЛЕВОЕ ПОКРЫТИЕ",
    "INTERNAL THICK ZINC-NICKEL": "ВНУТРЕННЕЕ ТОЛСТОЕ ЦИНК-НИКЕЛЕВОЕ ПОКРЫТИЕ",
    "EXTERNAL SERMETEL LIMIT": "НАРУЖНАЯ ГРАНИЦА ПОКРЫТИЯ SERMETEL",
    "INTERNAL SERMETEL LIMIT": "ВНУТРЕННЯЯ ГРАНИЦА ПОКРЫТИЯ SERMETEL",
    "AXLE BORE": "ОТВЕРСТИЕ ОСИ",
    "AXLE NUT CROSS BOLT HOLES": "ОТВЕРСТИЯ ПОПЕРЕЧНОГО БОЛТА ГАЙКИ ОСИ",
    "LUG BORES": "ОТВЕРСТИЯ ПРОУШИН",
    "DRAG ARM HOLES": "ОТВЕРСТИЯ ТЯГИ СКЛАДЫВАНИЯ",
    "TYPICAL 2 TRANSFER BLOCK LUGS": "ТИПОВО 2 ПРОУШИНЫ ПЕРЕХОДНОГО БЛОКА",
    "TYPICAL 2 BRAKE MANIFOLD LUGS": "ТИПОВО 2 ПРОУШИНЫ ТОРМОЗНОГО КОЛЛЕКТОРА",
    "2 LUGS": "2 ПРОУШИНЫ",
    "PLACES": "МЕСТА",
    "EXTENT": "ПРОТЯЖЕННОСТЬ",
    "( 2 PLACES )": "(2 МЕСТА)",
    "EXTENT OF FINE LIMIT DIA.": "ПРОТЯЖЕННОСТЬ УЧАСТКА МЕНЬШЕГО ПРЕД. ДИАМ.",
    "OF FINE LIMIT DIA.": "УЧАСТКА МЕНЬШЕГО ПРЕД. ДИАМ.",
    "UPPER DIAPHRAGM TUBE CROSS BORE": "ПОПЕРЕЧНОЕ ОТВЕРСТИЕ ВЕРХНЕЙ ДИАФРАГМЕННОЙ ТРУБЫ",
    "DIAMETER THRU BORE INCLUDING CHAMFERS": "ДИАМЕТР СКВОЗНОГО ОТВЕРСТИЯ ВКЛЮЧАЯ ФАСКИ",
    "DIAMETERS THRU BORES INCLUDING CHAMFERS": "ДИАМЕТРЫ СКВОЗНЫХ ОТВЕРСТИЙ ВКЛЮЧАЯ ФАСКИ",
    "(BORE AND CHAMFER INCLUDED)": "(ВКЛЮЧАЯ ОТВЕРСТИЕ И ФАСКУ)",
    "KNUCKLE BORES": "ОТВЕРСТИЯ ШАРНИРА",
    "PINTLE CROSS BORES": "ПОПЕРЕЧНЫЕ ОТВЕРСТИЯ ШТИФТА НАВЕСА СТОЙКИ",
    "PINTLE BORES": "ОТВЕРСТИЯ ШТИФТА НАВЕСА СТОЙКИ",
    "RETRACTION BORES": "ОТВЕРСТИЯ УЗЛА УБОРКИ",
    "TORQUE LINK AND RETAINING PIN BORES": "ОТВЕРСТИЯ ШЛИЦ-ШАРНИРА И СТОПОРНОГО ШТИФТА",
    "NO CADMIUM PLATE OR PAINT BEYOND THIS LINE": "НИ КАДМИЕВОЕ ПОКРЫТИЕ, НИ КРАСКА НЕ ДОЛЖНЫ ВЫХОДИТЬ ЗА ЭТУ ЛИНИЮ",
    "TWO PLACES": "2 МЕСТА",
    "A (2 PLACES)": "A (2 МЕСТА)",
    "C (2 PLACES)": "C (2 МЕСТА)",
    "C 2 LUGS C": "C 2 ПРОУШИНЫ C",
    "B HOLE": "ОТВЕРСТИЕ B",
    "HOLE": "ОТВЕРСТИЕ",
    "SPOTFACE": "ПОДРЕЗКА ПЛОЩАДКИ",
    "SPOTF": "ПОДРЕЗКА ПЛОЩАДКИ",
    "CHAMFER": "ФАСКА",
    "CHAMFE": "ФАСКА",
    "BARREL": "ЦИЛИНДРИЧЕСКАЯ ЧАСТЬ",
    "FACE": "ПОВЕРХНОСТЬ",
    "ON FACE": "НА ТОРЦЕ",
    "ZINC NICKEL PLATE": "ЦИНК-НИКЕЛЕВОЕ ПОКРЫТИЕ",
    "CADMIUM PLATE C": "КАДМИЕВОЕ ПОКРЫТИЕ C",
    "CADMIUM PLATE AND PAINT TO OVERLAP ON CHROMIUM RADIUS": "КАДМИЕВОЕ ПОКРЫТИЕ И КРАСКА ДОЛЖНЫ ПЕРЕКРЫВАТЬСЯ НА ХРОМОВОМ РАДИУСЕ",
    "CADMIUM PLATE AND PAINT TO OVERLAP": "КАДМИЕВОЕ ПОКРЫТИЕ И КРАСКА ДОЛЖНЫ ПЕРЕКРЫВАТЬСЯ",
    "PRIMER PAINT": "ГРУНТОВОЧНАЯ КРАСКА",
    "LIMIT OF A": "ГРАНИЦА A",
    "LIMIT OF D": "ГРАНИЦА D",
    "CHROME": "ХРОМ",
    "RAD IUS": "РАДИУС",
    "TYPICA": "ТИПОВО",
    "MIN.": "МИН.",
    "MAX.": "МАКС.",
    "TYPICAL": "ТИПОВО",
    "WORKING DIA.": "РАБОЧИЙ ДИАМ.",
    "This document and all information contained herein is the sole property of Safran Landing Systems (and/or its affiliated companies).": (
        "Настоящий документ и вся содержащаяся в нем информация являются исключительной собственностью Safran Landing Systems (и/или ее аффилированных компаний)."
    ),
    "The technical data in this document (or file) may contain US data and be controlled for export under the Export Administration Regulations (EAR), 15 CFR Parts 730-774, ECCN: 9E991. Violations of these laws may be subject to fines and penalties under the Export Administration Act.": (
        "Технические данные, содержащиеся в настоящем документе (или файле), могут включать данные США и подпадать под экспортный контроль в соответствии с Правилами экспортного администрирования (EAR), 15 CFR части 730-774, ECCN: 9E991. Нарушение этих требований может повлечь штрафы и иные санкции в соответствии с Законом об экспортном администрировании."
    ),
    "MLG - Installation of stub bolt subassembly for the forward pintle pin in place of the cross bolt.": (
        "MLG - Установка подсборки короткого болта для переднего штифта навеса стойки вместо поперечного болта."
    ),
    "MLG - To allow an increase in aircraft maximum take-off weight to 93 tonne.": (
        "MLG - Обеспечение увеличения максимальной взлетной массы самолета до 93 т."
    ),
    "MLG -To add tracking numbers to parts listed in Airbus Airworthiness Limitations Section (ALS).": (
        "MLG - Добавление номеров отслеживания к деталям, перечисленным в разделе ограничений летной годности Airbus (ALS)."
    ),
    "MLG - Installation of a 201585 series MLG Leg and Dressings where a 201387 MLG Leg and Dressings has been installed.": (
        "MLG - Установка стойки MLG серии 201585 и ее навесных элементов вместо ранее установленной стойки MLG серии 201387 с навесными элементами."
    ),
    "MLG -To add tracking numbers to parts listed in Airbus Maintenance Planning Document, Section 9-1. (Torque link apex pin nut)": (
        "MLG - Добавление номеров отслеживания к деталям, перечисленным в разделе 9-1 документа планирования технического обслуживания Airbus (гайка штифта вершины шлиц-шарнира)."
    ),
    "MLG - Introduction of a new lower bearing subassembly.": "MLG - Введение новой сборки нижнего подшипника.",
    "MLG - Introduction of new charging labels": "MLG - Введение новых табличек заправки.",
    "MLG - Introduction of new 1M and 2M Axle harnesses": "MLG - Введение новых жгутов оси 1M и 2M.",
    "MLG - Introduction of new 1M and 2M Leg Harness and of new 1M and 2M Axle Harnesses": (
        "MLG - Введение новых жгутов стойки 1M и 2M и новых жгутов оси 1M и 2M."
    ),
    "MLG Leg-Introduction of new retaining pins and a new lower bearing subassembly with a new self lubricating liner": (
        "Стойка MLG - Введение новых стопорных штифтов и новой сборки нижнего подшипника с новым самосмазывающимся вкладышем."
    ),
    "MLG Leg - Introduction of new retaining pins for the lower bearing subassembly": (
        "Стойка MLG - Введение новых стопорных штифтов для сборки нижнего подшипника."
    ),
    "MLG Leg - Introduction of a new lower bearing subassembly with a new low friction inner liner": (
        "Стойка MLG - Введение новой сборки нижнего подшипника с новым внутренним вкладышем с низким коэффициентом трения."
    ),
    "MLG Leg - Barkhausen Noise Inspection of Main Landing Gear Sliding Tube Axles.": (
        "Стойка MLG - Контроль методом шума Баркгаузена осей скользящей трубы основной стойки шасси."
    ),
    "MLG Leg - Introduction of a new Main Fitting": "Стойка MLG - Введение нового корпуса стойки.",
    "MLG Leg - Introduction of a new torque link damper unit": (
        "Стойка MLG - Введение нового демпферного узла шлиц-шарнира."
    ),
    "MLG Leg - Introduction of a new main fitting subassembly and related parts": (
        "Стойка MLG - Введение новой сборки корпуса стойки и сопутствующих деталей."
    ),
    "MLG - Introduction of a new upper pivot bracket": "MLG - Введение нового верхнего кронштейна шарнира.",
    "MLG - Introduction of a new changeover valve stem and housing": (
        "MLG - Введение нового штока и корпуса переключающего клапана."
    ),
    "MLG Complete - Modification of the transfer block subassembly": (
        "MLG в сборе - Модификация подсборки переходного блока."
    ),
    "MLG - Conversion of low - friction lower - bearing MLG to standard lower - bearing MLG": (
        "MLG - Переоборудование MLG с нижним подшипником пониженного трения в стандартную конфигурацию нижнего подшипника."
    ),
    "MLG complete - Introduction of a new transfer block subassembly": (
        "MLG в сборе - Введение новой подсборки переходного блока."
    ),
    "Updated revision status": "Обновлен статус ревизии",
    "Added Ref. Codes 2253 and 2255 details": "Добавлены сведения по кодам ссылок 2253 и 2255",
    "Updated pages": "Обновлены страницы",
    "Added content. Updated page numbers. Updated figure numbers": (
        "Добавлено содержимое. Обновлены номера страниц. Обновлены номера рисунков"
    ),
    "Added para 2.P.(31)": "Добавлен п. 2.P.(31)",
    "Updated tables 501 and 502": "Обновлены таблицы 501 и 502",
    "Updated table 601. Updated caution at para": "Обновлена таблица 601. Обновлено предупреждение в п.",
    "3.C. Updated figure titles. Deleted figures 626, 627, 649, 650, 653": (
        "3.C. Обновлены наименования рисунков. Удалены рисунки 626, 627, 649, 650, 653"
    ),
    "and 654. Updated figure 626. Added figures 642": "и 654. Обновлен рисунок 626. Добавлены рисунки 642",
    "648. Updated table 602. Updated figure numbers": "648. Обновлена таблица 602. Обновлены номера рисунков",
    "Added fig-item (18-80A) in para 1. Updated Messier-Dowty Limited to Safran Landing Systems": (
        "Добавлена позиция рисунка (18-80A) в п. 1. Наименование Messier-Dowty Limited заменено на Safran Landing Systems"
    ),
    "Added fig-item (18-80A) in paras 1. Updated Messier-Dowty Limited to Safran Landing Systems": (
        "Добавлена позиция рисунка (18-80A) в пп. 1. Наименование Messier-Dowty Limited заменено на Safran Landing Systems"
    ),
    "Added fig-item (18-80A) in paras 1. and 1.A.(2)": "Добавлена позиция рисунка (18-80A) в пп. 1 и 1.A.(2)",
    "Added fig-item (18-80A) in para 1. Updated Messier-Dowty Limited to Safran Landing Systems. Updated conversion value in figure 602": (
        "Добавлена позиция рисунка (18-80A) в п. 1. Наименование Messier-Dowty Limited заменено на Safran Landing Systems. Обновлено значение пересчета на рисунке 602"
    ),
    "Added fig-item (18-80A) in para 1. Updated material specification in para 1.D.(1). Updated Messier-Dowty Limited to Safran Landing System. Updated": (
        "Добавлена позиция рисунка (18-80A) в п. 1. Обновлена спецификация материала в п. 1.D.(1). Наименование Messier-Dowty Limited заменено на Safran Landing Systems. Обновлен"
    ),
    "figure 603": "рисунок 603",
    "Added fig-item (18-80A)": "Добавлена позиция рисунка (18-80A)",
    "in para 1. Updated Messier-Dowty Limited to Safran Landing Systems": (
        "в п. 1. Наименование Messier-Dowty Limited заменено на Safran Landing Systems"
    ),
    "para 1. Updated Messier-Dowty Limited to Safran Landing Systems": (
        "п. 1. Наименование Messier-Dowty Limited заменено на Safran Landing Systems"
    ),
    "Updated Messier-Dowty Limited to Safran Landing Systems": (
        "Наименование Messier-Dowty Limited заменено на Safran Landing Systems"
    ),
    "Updated fig-items in para 1. Updated Messier-Dowty Limited to Safran Landing Systems": (
        "Обновлены позиции рисунков в п. 1. Наименование Messier-Dowty Limited заменено на Safran Landing Systems"
    ),
    "Updated fig-items in para 1": "Обновлены позиции рисунков в п. 1",
    "Updated fig-items in para 1. Added caution at para 1.E. Updated Messier-Dowty Limited to Safran Landing Systems": (
        "Обновлены позиции рисунков в п. 1. Добавлено предупреждение в п. 1.E. Наименование Messier-Dowty Limited заменено на Safran Landing Systems"
    ),
    "Updated Messier-Dowty": "Наименование Messier-Dowty",
    "Limited to Safran Landing Systems. Updated table 601": (
        "изменено на Safran Landing Systems. Обновлена таблица 601"
    ),
    "Updated fig-item (2-340) only in para 1": "Обновлена только позиция рисунка (2-340) в п. 1",
    "Updated fig-item (2-350) only in para 1": "Обновлена только позиция рисунка (2-350) в п. 1",
    "Updated paras 1.H.(1),": "Обновлены пп. 1.H.(1),",
    "1.I.(1), 2.E, 2.F, 2.G,": "1.I.(1), 2.E, 2.F, 2.G,",
    "2.H, 2.J, 2.K, 2.M, 2.N": "2.H, 2.J, 2.K, 2.M, 2.N",
    "and 2.O. Added paras": "и 2.O. Добавлены пп.",
    "2.I and 2.L. Added figure 713. Updated": "2.I и 2.L. Добавлен рисунок 713. Обновлены",
    "figures 705, 706, 707,": "рисунки 705, 706, 707,",
    "708 and 710. Updated figure numbers": "708 и 710. Обновлены номера рисунков",
    "Updated figure 815.": "Обновлен рисунок 815.",
    "Updated tables 813,": "Обновлены таблицы 813,",
    "814, 818 and 823": "814, 818 и 823",
    "Updated para 1.A": "Обновлен п. 1.A",
    "Updated IPL figs 13 to 18 to include Ref.": "Обновлены рисунки IPL 13-18 с включением кодов ссылок",
    "Codes: 2253 and 2255. Updated IPL fig 15.": "2253 и 2255. Обновлен рисунок IPL 15.",
    "Subject Reference": "Тема/ссылка",
    "Remove and Destroy Pages": "Изъять и уничтожить страницы",
    "Insert New/Revised": "Вставить новые/пересмотренные",
    "Reason for Change": "Причина изменения",
    "Pages": "Страницы",
    "Dated": "Дата",
    "REV. No.": "№ РЕВ.",
    "REV.": "РЕВ.",
    "No.": "№",
    "ISSUE DATE": "ДАТА ВЫПУСКА",
    "DATE INSERTED": "ДАТА ВНЕСЕНИЯ",
    "PAGE NUMBER": "НОМЕР СТРАНИЦЫ",
    "BY": "КЕМ",
    "DATE REMOVED": "ДАТА УДАЛЕНИЯ",
    "SB NUMBER": "№ SB",
    "SB TITLE": "НАИМЕНОВАНИЕ SB",
    "SB REVISION NUMBER": "№ РЕВИЗИИ SB",
    "DATE INCORPORATED INTO MANUAL": "ДАТА ВКЛЮЧЕНИЯ В РУКОВОДСТВО",
    "COVER SB NO.": "№ ОХВАТЫВАЮЩЕГО SB",
    "SAFRAN LANDING SYSTEMS SERVICE BULLETIN NUMBER": "НОМЕР СЕРВИСНОГО БЮЛЛЕТЕНЯ SAFRAN LANDING SYSTEMS",
    "SERVICE BULLETIN NUMBER": "НОМЕР СЕРВИСНОГО БЮЛЛЕТЕНЯ",
    "INSERTION DATE": "ДАТА ВНЕСЕНИЯ",
    "REV NO.": "№ РЕВ.",
    "DASH NO.": "№ ПОЗ.",
    "SAFRAN LANDING SYSTEMS MODIFICATION NUMBER": "НОМЕР МОДИФИКАЦИИ SAFRAN LANDING SYSTEMS",
    "MOD. STRIKE NO.": "№ ВНЕДРЕНИЯ МОДИФ.",
    "List of Effective": "Перечень действующих",
    "Unit Identification": "Идентификация агрегата",
    "Record of Temporary": "Журнал временных",
    "Revisions": "ревизий",
    "Assembly (Including": "Сборка (включая",
    "Storage)": "хранение)",
    "Storage) (Continued)": "хранение) (Продолжение)",
    "Assembly (Including Storage)": "Сборка (включая хранение)",
    "Updated fig-items in": "Обновлены позиции рисунков в",
    "Description and": "Описание и",
    "Operation": "работа",
    "Testing and Fault": "Испытания и поиск",
    "Isolation": "неисправностей",
    "Isolation (Continued)": "неисправностей (Продолжение)",
    "Isolation (Продолжение)": "неисправностей (Продолжение)",
    "List of Service Bulletins": "Список сервисных бюллетеней",
    "Upper Pivot Bracket (10-160) Only - Protective Treatment": (
        "Только верхний кронштейн шарнира (10-160) - Защитная обработка"
    ),
    "Spacer (4-180) Only - Protective Treatment": "Только проставка (4-180) - Защитная обработка",
    "Uplock Pin ain Fitting (5-400A) - Protective Treatment": (
        "Штифт замка убранного положения корпуса стойки (5-400A) - Защитная обработка"
    ),
    "Штифт замка убранного положения ain Fitting (5-400A) - Защитная обработка": (
        "Штифт замка убранного положения корпуса стойки (5-400A) - Защитная обработка"
    ),
    "Torque Link Repairs - Key Diagram": "Ремонты шлиц-шарнира - Схема расположения",
    "Torque Link Ремонтs - Схема расположения": "Ремонты шлиц-шарнира - Схема расположения",
    "Main Fitting Repairs - Key Diagram": "Ремонты корпуса стойки - Схема расположения",
    "Sliding Tube Repairs - Key Diagram": "Ремонты скользящей трубы - Схема расположения",
    "Upper Diaphragm Tube Repairs - Key Diagram": "Ремонты верхней диафрагменной трубы - Схема расположения",
    "Cylinder Repairs - Key Diagram": "Ремонты цилиндра - Схема расположения",
    "Transfer Block Repairs - Key Diagram.": "Ремонты переходного блока - Схема расположения.",
    "Harness Support Bracket Repairs - Key Diagram": "Ремонты кронштейна крепления жгута - Схема расположения",
    "Upper Pivot Bracket Repairs - Key Diagram": "Ремонты верхнего кронштейна шарнира - Схема расположения",
    "Корпус стойки Ремонтs - Схема расположения": "Ремонты корпуса стойки - Схема расположения",
    "Скользящая труба Ремонтs - Схема расположения": "Ремонты скользящей трубы - Схема расположения",
    "Верхняя диафрагменная труба Ремонтs - Схема расположения": (
        "Ремонты верхней диафрагменной трубы - Схема расположения"
    ),
    "Цилиндр Ремонтs - Схема расположения": "Ремонты цилиндра - Схема расположения",
    "Переходной блок Ремонтs - Схема расположения.": "Ремонты переходного блока - Схема расположения.",
    "Кронштейн крепления жгута Ремонтs - Схема расположения": (
        "Ремонты кронштейна крепления жгута - Схема расположения"
    ),
    "Верхний кронштейн шарнира Ремонтs - Схема расположения": (
        "Ремонты верхнего кронштейна шарнира - Схема расположения"
    ),
    "Верхний кронштейн шарнира (10-160) Только - Защитная обработка": (
        "Только верхний кронштейн шарнира (10-160) - Защитная обработка"
    ),
    "Проставка (4-180) Только - Защитная обработка": "Только проставка (4-180) - Защитная обработка",
    "Oversize Bush(es) - Machining and Installation Repair No. 11-2": (
        "Ремонтная (увеличенная) втулка - Механическая обработка и установка\tРемонт № 11-2"
    ),
    "Oversize Bush(es) - Механическая обработка и установка Ремонт № 11-2": (
        "Ремонтная (увеличенная) втулка - Механическая обработка и установка\tРемонт № 11-2"
    ),
    "Repair to Bushes - Machining and Installation Repair No. 11-13": (
        "Ремонт втулок - Механическая обработка и установка\tРемонт № 11-13"
    ),
    "Ремонт Bushes - Механическая обработка и установка Ремонт № 11-13": (
        "Ремонт втулок - Механическая обработка и установка\tРемонт № 11-13"
    ),
    "Page Repair No. 11-37 Main Fitting . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . Repair No. 11-37 601": (
        "Стр. Ремонт № 11-37 Корпус стойки . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . Ремонт № 11-37\t601"
    ),
    "Page Ремонт № 11-37 Корпус стойки . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . Ремонт № 11-37 601": (
        "Стр. Ремонт № 11-37 Корпус стойки . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . Ремонт № 11-37\t601"
    ),
    "CAUTION: DO NOT USE CHLORINATED SOLVENTS. CHLORINATED SOLVENTS CAN MIX WITH VERY SMALL QUANTITIES OF WATER IN HYDRAULIC SYSTEMS TO MAKE HYDROCHLORIC ACID. HYDROCHLORIC ACID WILL CAUSE CORROSION ON METAL SURFACES.": (
        "ОСТОРОЖНО: НЕ ИСПОЛЬЗУЙТЕ ХЛОРИРОВАННЫЕ РАСТВОРИТЕЛИ. ХЛОРИРОВАННЫЕ РАСТВОРИТЕЛИ МОГУТ СМЕШИВАТЬСЯ С ОЧЕНЬ МАЛЫМИ КОЛИЧЕСТВАМИ ВОДЫ В ГИДРАВЛИЧЕСКИХ СИСТЕМАХ С ОБРАЗОВАНИЕМ СОЛЯНОЙ КИСЛОТЫ. СОЛЯНАЯ КИСЛОТА ВЫЗЫВАЕТ КОРРОЗИЮ МЕТАЛЛИЧЕСКИХ ПОВЕРХНОСТЕЙ."
    ),
    "All the materials in this manual have a Ref. Item identification. This is the reference item number of the material in the Aircraft Manufacturer’s Consumable Materials List.": (
        "Все материалы, приведенные в настоящем руководстве, имеют код ссылки. Это номер ссылки материала в перечне расходных материалов изготовителя воздушного судна."
    ),
    "All references in this manual are to the left configuration of the unit unless the instructions tell you differently.": (
        "Если в указаниях не оговорено иное, все ссылки в настоящем руководстве относятся к левому исполнению агрегата."
    ),
    "NOTE: Nitrogen will be released through the charging valve (13-60) as the piston (17-200) moves.": (
        "ПРИМЕЧАНИЕ. При перемещении поршня (17-200) азот будет выходить через заправочный клапан (13-60)."
    ),
    "NOTE: The charging valve (17-20) must be open to let the unit extend fully.": (
        "ПРИМЕЧАНИЕ. Заправочный клапан (17-20) должен быть открыт, чтобы агрегат полностью выдвинулся."
    ),
    "NOTE: The thread size is M142 x 1.5 pitch - 5h6h to BS3643.": (
        "ПРИМЕЧАНИЕ.\tРазмер резьбы: M142 x 1.5, поле допуска 5h6h по BS3643."
    ),
    "NOTE: Alternative equivalents are permitted.": "ПРИМЕЧАНИЕ. Допускаются альтернативные эквиваленты.",
    "CAUTION: DO NOT USE A PRESSURE OF MORE THAN 7,58 BAR (110 LBF/IN2).": (
        "ОСТОРОЖНО: НЕ ИСПОЛЬЗУЙТЕ ДАВЛЕНИЕ БОЛЕЕ 7,58 БАР (110 ФУНТ/ДЮЙМ2)."
    ),
    "CAUTION: DO NOT CAUSE DAMAGE TO THE PAINT FINISH.": (
        "ОСТОРОЖНО: НЕ ДОПУСТИТЕ ПОВРЕЖДЕНИЯ ЛАКОКРАСОЧНОГО ПОКРЫТИЯ."
    ),
    "AECMA Simplified English to PSC-85-16598 is used in this manual.": (
        "В настоящем руководстве используется упрощенный английский язык AECMA в соответствии с PSC-85-16598."
    ),
    "Parts of permanent assemblies that are not correctly attached.": (
        "Детали постоянных сборок, закрепленные ненадлежащим образом."
    ),
    "There is an error because of the pressure gauge capacity.": (
        "Из-за предела измерения манометра возникает погрешность."
    ),
    "If there is an error because of the gauge capacity:": (
        "Если возникает погрешность из-за диапазона прибора:"
    ),
    "Discard parts that you must not use again. These include:": (
        "Забракуйте детали, повторное использование которых не допускается. К ним относятся:"
    ),
    "Examine the unit for damage before you start the tests.": (
        "Перед началом испытаний осмотрите агрегат на наличие повреждений."
    ),
    "The test fluid must be clean: refer to M-DLPS910-1.": (
        "Испытательная жидкость должна быть чистой: см. M-DLPS910-1."
    ),
    "Examine each part for these types of damage:": (
        "Осмотрите каждую деталь на наличие следующих видов повреждений:"
    ),
    "Where K = 273 for temperatures in C": "Где K = 273 для температур в °C",
    "(459 for temperatures in F)": "(459 для температур в °F)",
    "They are a set: keep them together.": "Они образуют комплект: храните их вместе.",
    "This equipment is necessary:": "Необходимо следующее оборудование:",
    "These special tools are necessary:": "Необходимы следующие специальные инструменты:",
    "These materials are necessary:": "Необходимы следующие материалы:",
    "Clean the part: refer to para 2.A.": "Очистите деталь: см. п. 2.A.",
    "Unless instructions are different:": "Если в указаниях не сказано иное:",
    "Keep the unit in this condition for a minimum of six hours.": (
        "Выдержите агрегат в этом состоянии не менее шести часов."
    ),
    "The Location Frame 460007235 (for right configuration units).": (
        "Установочная рама 460007235 (для агрегатов правого исполнения)."
    ),
    "The Location Frame 460007234 (for left configuration units)": (
        "Установочная рама 460007234 (для агрегатов левого исполнения)."
    ),
    "Make a record of the ambient temperature, (T1).": (
        "Запишите температуру окружающей среды (T1)."
    ),
    "Measure the ambient temperature, (T2).": (
        "Измерьте температуру окружающей среды (T2)."
    ),
    "The inflation equipment must be to MIL-G-8348.": (
        "Заправочное оборудование должно соответствовать MIL-G-8348."
    ),
    "Remove the wire thread insert. Make sure that broken pieces do not stay in the hole.": (
        "Удалите резьбовую вставку. Убедитесь, что в отверстии не осталось обломков."
    ),
    "Bend the outer coil of the wire thread insert to the centre of the hole.": (
        "Отогните наружный виток резьбовой вставки к центру отверстия."
    ),
    "If necessary, remove the wire thread inserts:": (
        "При необходимости удалите резьбовые вставки:"
    ),
    "Use the Milliohmmeter Megger, Type BT51, to measure the electrical bonding resistance.": (
        "Используйте миллиомметр Megger, тип BT51, для измерения сопротивления электрического соединения."
    ),
    "Use the Crowfoot Wrench T14500 to remove the charging valve (13-60). Remove the O-ring seal (13-67) from the charging valve (13-60).": (
        "Используйте рожковый ключ типа Crowfoot T14500, чтобы снять заправочный клапан (13-60). Снимите с заправочного клапана (13-60) уплотнительное кольцо круглого сечения (13-67)."
    ),
    "Use the Crowfoot Wrench T14500 to remove the charging valve (17-20). Remove the O-ring seal (17-27) from the charging valve (17-20).": (
        "Используйте рожковый ключ типа Crowfoot T14500, чтобы снять заправочный клапан (17-20). Снимите с заправочного клапана (17-20) уплотнительное кольцо круглого сечения (17-27)."
    ),
    "Use the Charging Adapter 460002502 to connect the hydraulic test rig to the charging valve (13-60).": (
        "Используйте заправочный адаптер 460002502, чтобы подключить гидравлический испытательный стенд к заправочному клапану (13-60)."
    ),
    "Make sure that all of the nitrogen pressure has been released: remove the charging valve (17-20).": (
        "Убедитесь, что давление азота полностью сброшено: снимите заправочный клапан (17-20)."
    ),
    "Release the lock indentations of the locking washer (19-54).": (
        "Отогните участки стопорения стопорной шайбы (19-54)."
    ),
    "Release the pressure in the gauge.": "Сбросьте давление в манометре.",
    "Release the nitrogen pressure.": "Сбросьте давление азота.",
    "The Bottom Press Adapter 460007260.": "Нижний нажимной адаптер 460007260.",
    "The Jacking Dome Adapter 460006223": "Адаптер поддомкратного купола 460006223",
    "The Holding Fixture 460006231.": "Удерживающее приспособление 460006231.",
    "Trolley Support Arms Towing Frame": "Тележка, опорные рычаги, буксировочная рама",
    "Hydraulic fluid . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . Material Ref. Item 02-501": (
        "Гидравлическая жидкость . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . Код ссылки материала 02-501"
    ),
    "Static dis- charge con- nector": "Разъем отвода статического электричества",
    "Unserviceable screw threads.": "Резьба, непригодная к эксплуатации.",
    "P5A = P1A + (P2A - P4A) OR": "P5A = P1A + (P2A - P4A) ИЛИ",
    "10-123-11MD or": "10-123-11MD или",
    "BEARING (20-250)": "ПОДШИПНИК (20-250)",
    "Damper (9-160)": "ДЕМПФЕР (9-160)",
    "Parts 1": "Части 1",
    "- Chromium": "- Хромовое покрытие",
    "Cond H1025": "сост. H1025",
    "or": "или",
    "This manual contains Description, Operation, Maintenance procedures and an Illustrated Parts List (IPL). IPL Figure and Item numbers in parentheses follow the part name to identify them.": (
        "Настоящее руководство содержит разделы «Описание», «Работа», процедуры технического обслуживания и иллюстрированный перечень деталей (IPL). Номера рисунков и позиций IPL в скобках следуют за наименованием детали для ее идентификации."
    ),
    "A Unit Identification Chart is included to show the modification status of the unit. The modification status is related to the unit part number by the dash number: the dash number is marked on the unit name plate adjacent to the part number.": (
        "В руководство включена таблица идентификации агрегата, показывающая статус модификации агрегата. Статус модификации связан с номером детали агрегата через номер исполнения: номер исполнения нанесен на табличку агрегата рядом с номером детали."
    ),
    "All dimensions and quantities in this manual are in SI units with Imperial units in parentheses. A comma shows a decimal part of an SI unit. A full point shows a decimal part of an Imperial unit.": (
        "Все размеры и количества, приведенные в настоящем руководстве, указаны в единицах СИ, а британские единицы приведены в скобках. Запятая обозначает десятичную часть единицы СИ. Точка обозначает десятичную часть британской единицы."
    ),
    "This manual refers to Process Specifications (M-DLPS and PCS) and Non-destructive Tests (M-DLNDT). These are available within the Safran Landing Systems Technical Publications on-line service.": (
        "В настоящем руководстве приведены ссылки на технологические спецификации (M-DLPS и PCS) и документы по неразрушающему контролю (M-DLNDT). Они доступны в онлайн-службе технических публикаций Safran Landing Systems."
    ),
    "Use approved persons and good aircraft engineering practice for all procedures in this manual.": (
        "Для всех процедур, приведенных в настоящем руководстве, используйте уполномоченный персонал и соблюдайте надлежащую авиационно-техническую практику."
    ),
    "The repairs in this CMM have been approved under Airbus’ EASA Design Organisation Approval No. EASA.21J.031.": (
        "Ремонты, приведенные в данном CMM, одобрены в рамках одобрения конструкторской организации Airbus EASA No. EASA.21J.031."
    ),
    "On occasion a REF. CODE can be identified in the NOMENCLATURE column in the DETAILED PARTS LIST. This is a Safran Landing Systems reference code and is used for cross-reference purposes only.": (
        "Иногда в графе NOMENCLATURE подробного списка деталей может указываться REF. CODE. Это код ссылки Safran Landing Systems, используемый только для перекрестных ссылок."
    ),
    "The accuracy and the adequacy of the instructions in this CMM have been technically verified by shop verification (performed or simulated) or by similarity with manufacturing instructions or with component maintenance manuals instructions from other programs that have been verified in shop.": (
        "Точность и достаточность указаний, приведенных в данном CMM, технически подтверждены производственной проверкой (выполненной или смоделированной) либо по аналогии с производственными инструкциями или с инструкциями руководств по техническому обслуживанию компонентов из других программ, прошедших производственную проверку."
    ),
    "Safran Landing Systems UK Ltd Component Maintenance Manual, Main Landing Gear Leg and Dressings, 32-12-21.": (
        "Safran Landing Systems UK Ltd. Руководство по техническому обслуживанию компонента. Стойка основного шасси и навесные элементы. 32-12-21."
    ),
    "Safran Landing Systems UK Ltd Component Maintenance Manual, Axle Harness 1M and 2M, 32-12-29.": (
        "Safran Landing Systems UK Ltd. Руководство по техническому обслуживанию компонента. Жгут оси 1M и 2M. 32-12-29."
    ),
    "The main landing gear leg is a two stage, telescopic shock absorber.": (
        "Стойка основного шасси представляет собой двухступенчатый телескопический амортизатор."
    ),
    "Description (Refer to Figures 1 and 2)": "Описание (См. рисунки 1 и 2)",
    "The main landing gear leg has a sliding tube subassembly that operates in a main fitting subassembly. The sliding tube subassembly operates through a lower bearing subassembly. The lower bearing subassembly also seals the sliding tube subassembly in the main fitting subassembly.": (
        "Стойка основного шасси имеет подсборку скользящей трубы, работающую в подсборке корпуса стойки. Подсборка скользящей трубы проходит через подсборку нижнего подшипника. Подсборка нижнего подшипника также герметизирует подсборку скользящей трубы в подсборке корпуса стойки."
    ),
    "An upper torque link subassembly attaches to the main fitting subassembly. A lower torque link subassembly attaches to the sliding tube subassembly. A damper attaches to the upper torque link subassembly. A pin installs through the damper and connects the upper and lower torque link subassemblies.": (
        "К подсборке корпуса стойки крепится подсборка верхнего шлиц-шарнира. К подсборке скользящей трубы крепится подсборка нижнего шлиц-шарнира. К подсборке верхнего шлиц-шарнира крепится демпфер. Через демпфер устанавливается штифт, соединяющий подсборки верхнего и нижнего шлиц-шарнира."
    ),
    "A slave link subassembly and a lower slave link subassembly attach opposite the upper and lower torque link subassemblies.": (
        "Подсборка ведомого звена и подсборка нижнего ведомого звена крепятся напротив подсборок верхнего и нижнего шлиц-шарнира."
    ),
    "A rod and a cylinder install in the sliding tube subassembly. A piston installs in the cylinder. An upper diaphragm tube subassembly installs in the main fitting subassembly. A baffle, a compression orifice plate and a diaphragm subassembly install in the upper diaphragm tube subassembly. The rod goes through the baffle.": (
        "В подсборку скользящей трубы устанавливаются шток и цилиндр. В цилиндр устанавливается поршень. В подсборку корпуса стойки устанавливается подсборка верхней диафрагменной трубы. В подсборку верхней диафрагменной трубы устанавливаются перегородка, пластина дроссельного отверстия сжатия и подсборка диафрагмы. Шток проходит через перегородку."
    ),
    "An upper bearing housing installs between the top of the sliding tube subassembly and the main fitting subassembly. A recoil orifice plate operates in the upper bearing housing.": (
        "Корпус верхнего подшипника устанавливается между верхней частью подсборки скользящей трубы и подсборкой корпуса стойки. В корпусе верхнего подшипника работает пластина дроссельного отверстия отбоя."
    ),
    "Operation (Refer to Figure 2)": "Работа (См. рисунок 2)",
    "Compression": "Сжатие",
    "The sliding tube subassembly moves into the main fitting subassembly. The subsequent decrease in volume causes hydraulic fluid to flow through the upper bearing housing: the recoil orifice plate moves and slows the flow of hydraulic fluid. The decrease in volume also causes hydraulic fluid to move through the diaphragm and lift the compression orifice plate: the hydraulic fluid flows through the baffle and into the upper diaphragm tube subassembly. This slows the speed of the compression.": (
        "Подсборка скользящей трубы входит в подсборку корпуса стойки. Последующее уменьшение объема вызывает перетекание гидравлической жидкости через корпус верхнего подшипника: пластина дроссельного отверстия отбоя перемещается и замедляет поток гидравлической жидкости. Уменьшение объема также вызывает прохождение гидравлической жидкости через диафрагму и подъем пластины дроссельного отверстия сжатия: гидравлическая жидкость проходит через перегородку и поступает в подсборку верхней диафрагменной трубы. Это замедляет скорость сжатия."
    ),
    "Hydraulic fluid that moves into the upper diaphragm tube compresses the nitrogen in the main fitting subassembly and the upper diaphragm tube subassembly. As the pressure of the nitrogen increases, the hydraulic fluid in the rod moves against the piston. The piston is pushed into the cylinder and compresses the nitrogen in it. This slows the speed of the compression more.": (
        "Гидравлическая жидкость, поступающая в верхнюю диафрагменную трубу, сжимает азот в подсборке корпуса стойки и в подсборке верхней диафрагменной трубы. По мере увеличения давления азота гидравлическая жидкость в штоке перемещается к поршню. Поршень вдавливается в цилиндр и сжимает находящийся в нем азот. Это дополнительно замедляет скорость сжатия."
    ),
    "WARNING: DO NOT GET HYDRAULIC FLUID ON YOUR SKIN OR IN YOUR EYES. DO NOT BREATHE THE FUMES. ONLY USE IN A LOCATION THAT HAS A CONTINUOUS FLOW OF CLEAN AIR. HYDRAULIC FLUID IS POISONOUS AND DANGEROUS.": (
        "ПРЕДУПРЕЖДЕНИЕ: НЕ ДОПУСКАЙТЕ ПОПАДАНИЯ ГИДРАВЛИЧЕСКОЙ ЖИДКОСТИ НА КОЖУ ИЛИ В ГЛАЗА. НЕ ВДЫХАЙТЕ ПАРЫ. ИСПОЛЬЗУЙТЕ ТОЛЬКО В МЕСТЕ С ПОСТОЯННЫМ ПРИТОКОМ ЧИСТОГО ВОЗДУХА. ГИДРАВЛИЧЕСКАЯ ЖИДКОСТЬ ЯДОВИТА И ОПАСНА."
    ),
    "WARNING: DO NOT GET CLEANING AGENTS ON YOUR SKIN, IN YOUR EYES OR NEAR A FLAME. DO NOT BREATHE THE FUMES. ONLY USE IN A LOCATION THAT HAS A CONTINUOUS FLOW OF CLEAN AIR. CLEANING AGENTS ARE POISONOUS AND FLAMMABLE.": (
        "ПРЕДУПРЕЖДЕНИЕ: НЕ ДОПУСКАЙТЕ ПОПАДАНИЯ ОЧИСТИТЕЛЕЙ НА КОЖУ, В ГЛАЗА ИЛИ В ЗОНУ ОТКРЫТОГО ПЛАМЕНИ. НЕ ВДЫХАЙТЕ ПАРЫ. ИСПОЛЬЗУЙТЕ ТОЛЬКО В МЕСТЕ С ПОСТОЯННЫМ ПРИТОКОМ ЧИСТОГО ВОЗДУХА. ОЧИСТИТЕЛИ ЯДОВИТЫ И ОГНЕОПАСНЫ."
    ),
    "WARNING: DO NOT GET PAINT STRIPPER ON YOUR SKIN, IN YOUR EYES OR NEAR A FLAME. DO NOT BREATHE THE FUMES. ONLY USE IN A LOCATION THAT HAS A CONTINUOUS FLOW OF CLEAN AIR. PAINT STRIPPER IS POISONOUS AND FLAMMABLE.": (
        "ПРЕДУПРЕЖДЕНИЕ: НЕ ДОПУСКАЙТЕ ПОПАДАНИЯ СМЫВКИ КРАСКИ НА КОЖУ, В ГЛАЗА ИЛИ В ЗОНУ ОТКРЫТОГО ПЛАМЕНИ. НЕ ВДЫХАЙТЕ ПАРЫ. ИСПОЛЬЗУЙТЕ ТОЛЬКО В МЕСТЕ С ПОСТОЯННЫМ ПРИТОКОМ ЧИСТОГО ВОЗДУХА. СМЫВКА КРАСКИ ЯДОВИТА И ОГНЕОПАСНА."
    ),
    "WARNING: RELEASE ALL NITROGEN PRESSURE BEFORE YOU REMOVE THE CHARGING VALVES (13-60, 17-20).": (
        "ПРЕДУПРЕЖДЕНИЕ: ПЕРЕД СНЯТИЕМ ЗАПРАВОЧНЫХ КЛАПАНОВ (13-60, 17-20) ПОЛНОСТЬЮ СБРОСЬТЕ ДАВЛЕНИЕ АЗОТА."
    ),
    "CAUTION: DO NOT PUT AN END LOAD OF MORE THAN 5,08 TONNES (5 TONS) ON THE MAIN LANDING GEAR LEG (1-1).": (
        "ОСТОРОЖНО: НЕ ПРИКЛАДЫВАЙТЕ К СТОЙКЕ ОСНОВНОГО ШАССИ (1-1) ОСЕВУЮ НАГРУЗКУ БОЛЕЕ 5,08 ТОННЫ (5 ТОНН)."
    ),
    "CAUTION: DISCARD THE SCREWS (15-90) AND THE LOCKING PLATES (15-80) WHEN REMOVED.": (
        "ОСТОРОЖНО: ПОСЛЕ СНЯТИЯ ВЫБРАСЫВАЙТЕ ВИНТЫ (15-90) И СТОПОРНЫЕ ПЛАСТИНЫ (15-80)."
    ),
    "Recoil": "Отбой",
    "After compression, the nitrogen pressure in the cylinder pushes the piston to the end of the cylinder: hydraulic fluid moves out of the cylinder and into the rod. The nitrogen pressure in the main fitting subassembly and the upper diaphragm subassembly pushes the hydraulic fluid through the baffle: the compression orifice plate is pushed against the diaphragm subassembly and limits the flow of hydraulic fluid through it. This slows the speed of the recoil. The sliding tube subassembly moves out of the main fitting subassembly.": (
        "После сжатия давление азота в цилиндре перемещает поршень к концу цилиндра: гидравлическая жидкость выходит из цилиндра и поступает в шток. Давление азота в подсборке корпуса стойки и в подсборке верхней диафрагмы проталкивает гидравлическую жидкость через перегородку: пластина дроссельного отверстия сжатия прижимается к подсборке диафрагмы и ограничивает поток гидравлической жидкости через нее. Это замедляет скорость отбоя. Подсборка скользящей трубы выходит из подсборки корпуса стойки."
    ),
    "The Upper and Lower Torque Link Subassemblies": "Подсборки верхнего и нижнего шлиц-шарнира",
    "The upper and lower torque link subassemblies prevent the sliding tube subassembly from turning in the main fitting subassembly.": (
        "Подсборки верхнего и нижнего шлиц-шарнира предотвращают поворот подсборки скользящей трубы в подсборке корпуса стойки."
    ),
    "The damper controls the movement of the upper and lower torque link subassemblies.": (
        "Демпфер управляет перемещением подсборок верхнего и нижнего шлиц-шарнира."
    ),
    "The hydraulic test rig must have a hand pump and a power pump. The power pump must have a controlled flow of not less than 4,5 l/min (4.62 in3/sec).": (
        "Гидравлический испытательный стенд должен иметь ручной насос и силовой насос. Силовой насос должен обеспечивать регулируемый расход не менее 4,5 л/мин (4.62 in3/sec)."
    ),
    "The temperature of the test fluid must be between 20 and 40 C (68 and 104 F).": (
        "Температура испытательной жидкости должна быть в пределах от 20 до 40 °C (68-104 °F)."
    ),
    "During all hydraulic tests, the unit and the test circuit must be hydraulically full.": (
        "Во время всех гидравлических испытаний агрегат и испытательный контур должны быть полностью заполнены гидравлической жидкостью."
    ),
    "During the proximity switch tests the ambient temperature must be between 15 and 25 C (59 and 77 F).": (
        "Во время испытаний датчиков приближения температура окружающей среды должна быть в пределах от 15 до 25 °C (59-77 °F)."
    ),
    "Piston (17-200) Leakage Tests": "Испытания поршня (17-200) на герметичность",
    "Use the Charging Adapter 460002502 and the Turner Inflation Equipment T14218: connect the charging valve (17-20) to the nitrogen supply. Open the charging valve (17-20).": (
        "Используйте заправочный адаптер 460002502 и заправочное оборудование Turner T14218: подключите заправочный клапан (17-20) к источнику азота. Откройте заправочный клапан (17-20)."
    ),
    "Slowly increase the nitrogen pressure to between 9,32 and 10,68 bar (135 and 155 lbf/in2). Make a record of the pressure. Close the charging valve (17-20) and hold the nitrogen pressure for 15 minutes.": (
        "Медленно увеличьте давление азота до 9,32-10,68 бар (135-155 lbf/in2). Запишите давление. Закройте заправочный клапан (17-20) и выдержите давление азота в течение 15 минут."
    ),
    "Open the charging valve (17-20) and measure the nitrogen pressure: it must be the same as the record in para (2). Leakage must not occur.": (
        "Откройте заправочный клапан (17-20) и измерьте давление азота: оно должно совпадать со значением, записанным в п. (2). Утечка не допускается."
    ),
    "Disconnect the nitrogen supply and remove the Turner Inflation Equipment T14218 and the Charging Adapter 460002502.": (
        "Отсоедините источник азота и снимите заправочное оборудование Turner T14218 и заправочный адаптер 460002502."
    ),
    "Refer to ASSEMBLY: install the charging valve (17-20) and complete the assembly procedure.": (
        "См. раздел «Сборка»: установите заправочный клапан (17-20) и завершите процедуру сборки."
    ),
    "Main Landing Gear Leg (1-1) Tests": "Испытания стойки основного шасси (1-1)",
    "Initial Operations": "Подготовительные операции",
    "Use these special tools to install the main landing gear leg (1-1) vertically in the loading press:": (
        "Используйте следующие специальные инструменты, чтобы установить стойку основного шасси (1-1) вертикально в нагрузочный пресс:"
    ),
    "Assemble the Load Cell and Adapter 460006232 and the Offset Adapter 460006234 to the main landing gear leg (1-1).": (
        "Установите на стойку основного шасси (1-1) тензодатчик Load Cell с адаптером 460006232 и смещенный адаптер 460006234."
    ),
    "Procedure to Fill and Pressurize the Main Landing Gear Leg (1-1)": (
        "Процедура заполнения и наддува стойки основного шасси (1-1)"
    ),
    "Make sure that there is no pressure in the main landing gear leg (1-1): open the charging valves (13-60 and 17-20).": (
        "Убедитесь в отсутствии давления в стойке основного шасси (1-1): откройте заправочные клапаны (13-60 и 17-20)."
    ),
    "Slowly increase the hydraulic pressure to between 13,11 and 14,48 bar (190 and 210 lbf/in2) and let the unit extend fully.": (
        "Медленно увеличьте гидравлическое давление до 13,11-14,48 бар (190-210 lbf/in2) и дайте агрегату полностью выдвинуться."
    ),
    "Release the hydraulic pressure and fully close the unit.": (
        "Сбросьте гидравлическое давление и полностью закройте агрегат."
    ),
    "Do para (c) and (d) until the hydraulic fluid that comes out of the unit does not have air in it.": (
        "Выполняйте пп. (c) и (d), пока выходящая из агрегата гидравлическая жидкость не перестанет содержать воздух."
    ),
    "Fully close the unit and disconnect the hydraulic test rig.": (
        "Полностью закройте агрегат и отсоедините гидравлический испытательный стенд."
    ),
    "Use the Charging Adapter 460002502 and the Turner Inflation Equipment T14218 to connect the nitrogen supply to the charging valve (13-60).": (
        "Используйте заправочный адаптер 460002502 и заправочное оборудование Turner T14218, чтобы подключить источник азота к заправочному клапану (13-60)."
    ),
    "Slowly increase the nitrogen pressure until the unit starts to extend. Hold the pressure and fully extend the unit. The pressure must not be more than 7,58 bar (110 lbf/in2).": (
        "Медленно увеличивайте давление азота, пока агрегат не начнет выдвигаться. Удерживайте давление и полностью выдвиньте агрегат. Давление не должно превышать 7,58 бар (110 lbf/in2)."
    ),
    "Refer to Figure 101 and measure the dimension X: it must be between 483,05 and 487,85 mm (19.017 and 19.207 in).": (
        "См. рисунок 101 и измерьте размер X: он должен быть в пределах от 483,05 до 487,85 мм (19.017-19.207 in)."
    ),
    "Use the Charging Adapter 460002502 and the Turner Inflation Equipment T14218 to connect the charging valve (17-20) to the nitrogen supply.": (
        "Используйте заправочный адаптер 460002502 и заправочное оборудование Turner T14218, чтобы подключить заправочный клапан (17-20) к источнику азота."
    ),
    "Slowly increase the nitrogen pressure to between 13,11 and 14,48 bar (190 and 210 lbf/in2).": (
        "Медленно увеличьте давление азота до 13,11-14,48 бар (190-210 lbf/in2)."
    ),
    "Slowly increase the nitrogen pressure to between 67,59 and 70,34 bar (980 and1020 lbf/in2).": (
        "Медленно увеличьте давление азота до 67,59-70,34 бар (980-1020 lbf/in2)."
    ),
    "Close the charging valve (17-20); use the Crowfoot Wrench T14500 to torque it to between 5,7 and 7,9 N m (50 and 70 lbf in).": (
        "Закройте заправочный клапан (17-20); используйте рожковый ключ типа Crowfoot T14500, чтобы затянуть его моментом 5,7-7,9 Н·м (50-70 lbf in)."
    ),
    "Use the Turner Inflation Equipment T14218 and the Charging Adapter 460002502 to connect the nitrogen supply to the charging valve (13-60).": (
        "Используйте заправочное оборудование Turner T14218 и заправочный адаптер 460002502, чтобы подключить источник азота к заправочному клапану (13-60)."
    ),
    "Slowly increase the nitrogen pressure to between 6,90 and 8,27 bar (100 and 120 lbf/in2).": (
        "Медленно увеличьте давление азота до 6,90-8,27 бар (100-120 lbf/in2)."
    ),
    "Close the charging valve (13-60); use the Crowfoot Wrench T14500 to torque it to between 5,7 and 7,9 N m (50 and 70 lbf in).": (
        "Закройте заправочный клапан (13-60); используйте рожковый ключ типа Crowfoot T14500, чтобы затянуть его моментом 5,7-7,9 Н·м (50-70 lbf in)."
    ),
    "Slowly increase the nitrogen pressure to between 9,32 and 10,68 bar (135 and": (
        "Медленно увеличьте давление азота до 9,32-10,68 бар (135 и"
    ),
    "155 lbf/in2). Make a record of the pressure. Close the charging valve (17-20) and hold the nitrogen pressure for 15 minutes.": (
        "155 lbf/in2). Запишите давление. Закройте заправочный клапан (17-20) и выдержите давление азота в течение 15 минут."
    ),
    "Slowly increase the nitrogen pressure until the unit starts to extend. Hold the pressure and fully extend the unit. The pressure must not be more than 7,58 bar": (
        "Медленно увеличивайте давление азота, пока агрегат не начнет выдвигаться. Удерживайте давление и полностью выдвиньте агрегат. Давление не должно превышать 7,58 бар"
    ),
    "Main Landing Gear Leg (1-1) Figure 101": "Стойка основного шасси (1-1), рисунок 101",
    "Make a record of the nitrogen pressure at the charging valve (13-60), (P1A) and the charging valve (17-20), (P1B).": (
        "Запишите давление азота на заправочном клапане (13-60), (P1A), и на заправочном клапане (17-20), (P1B)."
    ),
    "Compare the pressures P1A and P2A and compare the pressures P1B and P2B. The pressures P1A and P2A must be the same and the pressures P1B and P2B must be the same, unless:": (
        "Сравните давления P1A и P2A, а также давления P1B и P2B. Давления P1A и P2A должны совпадать, и давления P1B и P2B также должны совпадать, если только:"
    ),
    "There is a difference between the temperatures T1 and T2": (
        "Имеется разница между температурами T1 и T2"
    ),
    "If there is a difference between the temperatures T1 and T2, calculate the correct value for the nitrogen pressures (these will be P3A and P3B) and adjust the pressures to the corrected values. Use the formula:": (
        "Если имеется разница между температурами T1 и T2, рассчитайте скорректированные значения давления азота (это будут P3A и P3B) и доведите давление до скорректированных значений. Используйте формулу:"
    ),
    "OR": "ИЛИ",
    "Z =  1 for pressures in bar": "Z = 1 для давлений в барах",
    "Main Landing Gear Leg (1-1) Figure 101": "Стойка основного шасси (1-1), рисунок 101",
    "Open the charging valve (13-60) and measure the nitrogen pressure, (P4A).": (
        "Откройте заправочный клапан (13-60) и измерьте давление азота, (P4A)."
    ),
    "Open the charging valve (17-20) and measure the nitrogen pressure, (P4B).": (
        "Откройте заправочный клапан (17-20) и измерьте давление азота, (P4B)."
    ),
    "Calculate the correct values for the nitrogen pressures (these will be P5A and P5B) and adjust the pressures to the corrected values. Use the formula:": (
        "Рассчитайте скорректированные значения давления азота (это будут P5A и P5B) и доведите давление до скорректированных значений. Используйте формулу:"
    ),
    "P5A = P1A + (P2A - P4A) OR": "P5A = P1A + (P2A - P4A) ИЛИ",
    "Prepare for Transport and Storage": "Подготовка к транспортированию и хранению",
    "Open the charging valve (13-60) and reduce the nitrogen pressure to between 3,45 and 4,82 bar (50 and 70 lbf/in2).": (
        "Откройте заправочный клапан (13-60) и уменьшите давление азота до 3,45-4,82 бар (50-70 lbf/in2)."
    ),
    "Open the charging valve (17-20) and reduce the nitrogen pressure to between 3,45 and 4,82 bar (50 and 70 lbf/in2).": (
        "Откройте заправочный клапан (17-20) и уменьшите давление азота до 3,45-4,82 бар (50-70 lbf/in2)."
    ),
    "Write this data on a label and attach it to the unit: THE GEAR MUST BE INFLATED TO THE APPROPRIATE PRESSURES BEFORE BEING PLACED IN SERVICE.": (
        "Нанесите эти данные на табличку и прикрепите ее к агрегату: СТОЙКА ДОЛЖНА БЫТЬ ЗАПРАВЛЕНА ДО СООТВЕТСТВУЮЩИХ ДАВЛЕНИЙ ПЕРЕД ВВОДОМ В ЭКСПЛУАТАЦИЮ."
    ),
    "Complete the torque procedure for the retaining pins (13-10): refer to ASSEMBLY.": (
        "Выполните процедуру затяжки стопорных штифтов (13-10): см. раздел «Сборка»."
    ),
    "Proximity Switches (7-40 and 7-230) Adjustments and Tests": (
        "Регулировка и испытания датчиков приближения (7-40 и 7-230)"
    ),
    "Use the loading press: set the dimension between the pins (10-80 and 11-130) to between 632,80 and 636,95 mm (24.9134 and 25.0767 in).": (
        "Используйте нагрузочный пресс: установите размер между штифтами (10-80 и 11-130) в пределах 632,80-636,95 мм (24.9134-25.0767 in)."
    ),
    "Adjust the spacers (6-140, 7-50, 7-190 and 7-240) or laminated shims (6-140A, 7-50A, 7-90A and 7-240A): refer to ASSEMBLY.": (
        "Отрегулируйте проставки (6-140, 7-50, 7-190 и 7-240) или наборные регулировочные прокладки (6-140A, 7-50A, 7-90A и 7-240A): см. раздел «Сборка»."
    ),
    "NOTE: If the calculated gap is in the tolerance, the spacers (6-140, 7-50, 7-190 and 7-240) or laminated shims (6-140A, 7-50A, 7-90A and 7-240A) are not necessary.": (
        "ПРИМЕЧАНИЕ. Если рассчитанный зазор находится в пределах допуска, проставки (6-140, 7-50, 7-190 и 7-240) или наборные регулировочные прокладки (6-140A, 7-50A, 7-90A и 7-240A) не требуются."
    ),
    "Connect the 28 VDC power supply, the Lampbox 460005842 and the main landing gear leg (1-1).": (
        "Подключите источник питания 28 В пост. тока, блок Lampbox 460005842 и стойку основного шасси (1-1)."
    ),
    "Use the loading press to fully extend the main landing gear leg (1-1).": (
        "Используйте нагрузочный пресс, чтобы полностью выдвинуть стойку основного шасси (1-1)."
    ),
    "Use the loading press to slowly close the main landing gear leg (1-1):": (
        "Используйте нагрузочный пресс, чтобы медленно закрыть стойку основного шасси (1-1):"
    ),
    "The proximity switch (7-230) must operate before the main landing gear leg (1-1) has closed by 26,00 mm (1.0236 in).": (
        "Датчик приближения (7-230) должен сработать до того, как стойка основного шасси (1-1) закроется на 26,00 мм (1.0236 in)."
    ),
    "The proximity switch (7-40) must operate before the main landing gear leg (1-1) has closed by 29,30 mm (1.1535 in).": (
        "Датчик приближения (7-40) должен сработать до того, как стойка основного шасси (1-1) закроется на 29,30 мм (1.1535 in)."
    ),
    "Do para (4) and (5) again.": "Повторите пп. (4) и (5).",
    "Disconnect the 28 VDC supply and the Lampbox 460005842.": (
        "Отсоедините питание 28 В пост. тока и блок Lampbox 460005842."
    ),
    "Remove the main landing gear leg (1-1) from the loading press.": (
        "Снимите стойку основного шасси (1-1) с нагрузочного пресса."
    ),
    "NOTE: Make sure that the main landing gear leg (1-1) is electrically isolated from the equipment that is used to hold it.": (
        "ПРИМЕЧАНИЕ. Убедитесь, что стойка основного шасси (1-1) электрически изолирована от оборудования, используемого для ее удержания."
    ),
    "Measure between the bearing (20-250) and the test points given in Table 101. The electrical bonding resistance must not be more than the limit given in Table 101.": (
        "Измерьте между подшипником (20-250) и точками проверки, указанными в таблице 101. Сопротивление электрического соединения не должно превышать предельное значение, указанное в таблице 101."
    ),
    "Measure between the axle of the sliding tube subassembly (17-240) and the test points given in Table 102. The electrical bonding resistance must not be more than the limit given in Table 102.": (
        "Измерьте между осью подсборки скользящей трубы (17-240) и точками проверки, указанными в таблице 102. Сопротивление электрического соединения не должно превышать предельное значение, указанное в таблице 102."
    ),
    "NOTE: Refer to TESTING AND FAULT ISOLATION to find the necessary level of disassembly. This will give the condition of the component or the possible cause of its malfunction.": (
        "ПРИМЕЧАНИЕ. См. раздел «Испытания и поиск неисправностей», чтобы определить необходимый уровень разборки. Это позволит установить состояние компонента или возможную причину его неисправности."
    ),
    "Make sure that the work area, the tools and the equipment are clean.": (
        "Убедитесь, что рабочая зона, инструменты и оборудование чистые."
    ),
    "Procedure (Refer to IPL Figures 1 to 20)": "Процедура (См. рисунки IPL 1-20)",
    "Use these special tools as necessary during the procedure to lift and to hold the unit:": (
        "Используйте следующие специальные инструменты по мере необходимости в ходе процедуры для подъема и удержания агрегата:"
    ),
    "The Lifting Bar Assembly 460006208": "Подъемная штанга в сборе 460006208",
    "The Spherical Bearing Locator 460007282": "Установочный шаблон сферического подшипника 460007282",
    "The Pintle Location Assembly 460007281": "Установочное приспособление штифта навеса стойки 460007281",
    "Assemble the Load Cell and Adapter 460006232 and the Offset Adapter 460006234 to the main landing gear leg (1-1).": (
        "Установите на стойку основного шасси (1-1) тензодатчик с адаптером 460006232 и смещенный адаптер 460006234."
    ),
    "Open the charging valve (13-60) and release the nitrogen pressure. Do not close the charging valve (13-60).": (
        "Откройте заправочный клапан (13-60) и стравите давление азота. Не закрывайте заправочный клапан (13-60)."
    ),
    "and": "и",
    "and 460006261": "и 460006261",
    "Type 1 or 17-4PH to": "тип 1 или 17-4PH по",
    "Measure the nitrogen pressure at the charging valve (13-60), (P2A) and the charging valve (17-20), (P2B).": (
        "Измерьте давление азота у заправочного клапана (13-60), (P2A), и у заправочного клапана (17-20), (P2B)."
    ),
    "Z = 1 for pressures in bar": "Z = 1 для давлений в бар",
    "(15 for pressures in lbf/in2)": "(15 для давлений в lbf/in2)",
    "Repair Sleeves - Machining and Installation\tRepair No. 12-2": (
        "Ремонтные втулки - механическая обработка и установка\tРемонт № 12-2"
    ),
    "Repair Sleeves - Machining and Installation\tRepair No. 12-5": (
        "Ремонтные втулки - механическая обработка и установка\tРемонт № 12-5"
    ),
    "Oversize Thread Insert - Installation\tRepair No. 18-2": (
        "Резьбовая вставка увеличенного размера - установка\tРемонт № 18-2"
    ),
    "Labels (20-10, 20-30, 20-40, 20-60 and 20-80) and wiring diagram plate (1-110) . . . .\n- Installation . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .": (
        "Таблички (20-10, 20-30, 20-40, 20-60 и 20-80) и табличка электрической схемы (1-110) . . . .\n- Установка . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . ."
    ),
    "Proximity Switch (7-230) and Target (7-180) - Adjustment. . . . . . . . . . . . . . . . . . . . . .": (
        "Датчик приближения (7-230) и мишень (7-180) - Регулировка. . . . . . . . . . . . . . . . . . . . . ."
    ),
    "Proximity Switch (7-40) and Target (6-130) - Adjustment. . . . . . . . . . . . . . . . . . . . . .": (
        "Датчик приближения (7-40) и мишень (6-130) - Регулировка. . . . . . . . . . . . . . . . . . . . . ."
    ),
    "Load Cell and Adapter": "Тензодатчик и адаптер",
    "Remove the charging valves (13-60 and 17-20)": "Снимите заправочные клапаны (13-60 и 17-20)",
    "Remove the forward pintle bush (20-250A)": "Снимите втулку переднего штифта навеса (20-250A)",
    "Remove the bearings (5-280 and 5-290) and the bushes (20-380)": (
        "Снимите подшипники (5-280 и 5-290) и втулки (20-380)"
    ),
    "Remove the bearings (20-230, 20-240 and 20-290)": (
        "Снимите подшипники (20-230, 20-240 и 20-290)"
    ),
    "Lifting Bar Assembly": "Подъемная штанга в сборе",
    "Lift the sliding tube subassembly (17-240) and related parts": (
        "Поднимать подсборку скользящей трубы (17-240) и связанные с ней детали"
    ),
    "Transport and Build\nTrolley Support Arms Towing Frame\nJacking Dome Adapter\n\nAdapter": (
        "Транспортировочная и сборочная\nтележка, опорные рычаги, буксировочная рама\nАдаптер поддомкратного купола\n\nАдаптер"
    ),
    "Use with 460006232, 460006263\nand 460006261": (
        "Использовать с 460006232, 460006263\nи 460006261"
    ),
    "Use with 460006232, 460006263 and 460006261": "Использовать с 460006232, 460006263 и 460006261",
    "Assembly/Extraction Tool": "Сборочно-демонтажный инструмент",
    "Extractor Pad and Drawbolt": "Съемная опора и вытяжной болт",
    "\nTo remove the forward pintle bush (2-250A)": "\nДля снятия втулки переднего штифта навеса (2-250A)",
    "Hold the main landing gear leg (1-1) (left configuration)\n\nHold the main landing gear leg (1-2) (right configuration)": (
        "Удерживать стойку основного шасси (1-1) (левая конфигурация)\n\nУдерживать стойку основного шасси (1-2) (правая конфигурация)"
    ),
    "Hold the sliding tube subassembly (17-240) and related parts": (
        "Удерживать подсборку скользящей трубы (17-240) и связанные с ней детали"
    ),
    "M-D\nSpec": "Спецификация\nM-D",
    "M-D Spec": "Спецификация\nM-D",
    "PCS-3100\nand\n\nM-DLNDT3\nParts 1\nand 2": "PCS-3100\nи\n\nM-DLNDT3\nЧасти 1\nи 2",
    "PCS-3100 and M-DLNDT3 Parts 1 and 2": "PCS-3100\nи\n\nM-DLNDT3\nЧасти 1\nи 2",
    "and\n\nM-DLNDT3\nParts 1": "и\n\nM-DLNDT3\nЧасти 1",
    "and M-DLNDT3 Parts 1": "и\n\nM-DLNDT3\nЧасти 1",
    "and 2": "и 2",
    "and\n\nPCS-3002": "и\n\nPCS-3002",
    "and PCS-3002": "и\n\nPCS-3002",
    "Spec": "Спецификация",
    "Steel, 35CD4 or\n4340 to AMS6414 or 35NCD16 to NCT 10-123-11 MD": (
        "Сталь, 35CD4 или\n4340 по AMS6414 или 35NCD16 по NCT 10-123-11 MD"
    ),
    "Stainless Steel, Z15CN17-03\nType 1 or 17-4PH to\nAMS 5604/5643": (
        "Коррозионностойкая сталь, Z15CN17-03\nтип 1 или 17-4PH по\nAMS 5604/5643"
    ),
    "Stainless Steel, Z15CN17-03 Type 1 or 17-4PH to AMS 5604/5643": (
        "Коррозионностойкая сталь, Z15CN17-03\nтип 1 или 17-4PH по\nAMS 5604/5643"
    ),
    "Aluminium Alloy, L113 or\nL168-T6511": "Алюминиевый сплав, L113 или\nL168-T6511",
    "Aluminium Alloy, BS L168 or\nBS 2L93 or 7075-T73, T7351, T73510\nor T73511 to MTL-2701": (
        "Алюминиевый сплав, BS L168 или\nBS 2L93 или 7075-T73, T7351, T73510\nили T73511 по MTL-2701"
    ),
    "Post SB 201-32-22: cut the Bowden cable (1-45) and remove the cross bolts (1-47 and 1-49).": (
        "ПОСЛЕ SB 201-32-22: разрежьте трос Боудена (1-45) и снимите поперечные болты (1-47 и 1-49)."
    ),
    "Remove the bolt (2-10), the washer (2-20), the pin (2-30), the spacer (2-40) and the": (
        "Снимите болт (2-10), шайбу (2-20), штифт (2-30), проставку (2-40) и"
    ),
    "Release the tab washer (2-150). Remove the bolt (2-140), the tab washer (2-150) and the bonding cable (2-160).": (
        "Отогните усик отгибной шайбы (2-150). Снимите болт (2-140), отгибную шайбу (2-150) и кабель заземления (2-160)."
    ),
    "Remove the bolts (5-20 and 5-40), the washers (5-30 and 5-50), the nut (5-60) and the": (
        "Снимите болты (5-20 и 5-40), шайбы (5-30 и 5-50), гайку (5-60) и"
    ),
    "Use the Drift 460004331/7 and the Extractor 460006151/24 to remove the bearings (5-280 and 5-290) from the bracket (5-300).": (
        "Используйте выколотку 460004331/7 и съемник 460006151/24 для снятия подшипников (5-280 и 5-290) с кронштейна (5-300)."
    ),
    "Remove the split pin (6-60), the slotted nut (6-70), the washer (6-80) and the pivot pin (6-90).": (
        "Снимите шплинт (6-60), прорезную гайку (6-70), шайбу (6-80) и шарнирный штифт (6-90)."
    ),
    "Remove the slave link subassembly (6-190) and its attached parts.": (
        "Снимите подсборку ведомого звена (6-190) и присоединенные к ней детали."
    ),
    "NOTE: If the calculated gap is in the tolerance, the spacers (6-140) or the laminated shim (6-140A) is not installed.": (
        "ПРИМЕЧАНИЕ. Если рассчитанный зазор находится в пределах допуска, проставки (6-140) или наборная регулировочная прокладка (6-140A) не устанавливаются."
    ),
    "Remove the grooved spherical bearing (6-300) or the self lubricating bearing (6-300A) from the lower slave link (6-310).": (
        "Снимите сферический подшипник с канавкой (6-300) или самосмазывающийся подшипник (6-300A) с нижнего ведомого звена (6-310)."
    ),
    "NOTE: If the calculated gap is in the tolerance, the spacer (7-50) or the laminated shim (7-50A) is not installed.": (
        "ПРИМЕЧАНИЕ. Если рассчитанный зазор находится в пределах допуска, проставка (7-50) или наборная регулировочная прокладка (7-50A) не устанавливается."
    ),
    "NOTE: If the calculated gap is in the tolerance, the spacer (7-190) or the laminated shim (7-190A) is not installed.": (
        "ПРИМЕЧАНИЕ. Если рассчитанный зазор находится в пределах допуска, проставка (7-190) или наборная регулировочная прокладка (7-190A) не устанавливается."
    ),
    "Remove the housing (12-170) and its related parts.": "Снимите корпус (12-170) и связанные с ним детали.",
    "Remove the plate (13-100) from the inflation valve (13-110).": (
        "Снимите пластину (13-100) с заправочного клапана (13-110)."
    ),
    "Remove the split pin (13-140), the nut (13-150), the washers (13-160), the bolt (13-170) and the stop ring (13-180).": (
        "Снимите шплинт (13-140), гайку (13-150), шайбы (13-160), болт (13-170) и стопорное кольцо (13-180)."
    ),
    "Remove the shock absorber subassembly (13-50) and its related parts from the main fitting subassembly (20-90).": (
        "Снимите подсборку амортизатора (13-50) и связанные с ней детали со сборки корпуса стойки (20-90)."
    ),
    "Remove the upper diaphragm tube subassembly (15-360) and its related parts.": (
        "Снимите подсборку верхней диафрагменной трубы (15-360) и связанные с ней детали."
    ),
    "Remove the O-ring seal (16-20A or 16A-20A), the backing rings (16-30 or 16A-30), the O-ring seal (16-40A or 16A-40A) and the backing": (
        "Снимите уплотнительное кольцо круглого сечения (16-20A или 16A-20A), опорные кольца (16-30 или 16A-30), уплотнительное кольцо круглого сечения (16-40A или 16A-40A) и опорные"
    ),
    "rings (16-50 or 16A-50).": "кольца (16-50 или 16A-50).",
    "Remove the common lower bearing bushes (16-130A or 16A-130B) from the lower bearing housing (16-140B or 16A-140C).": (
        "Снимите общие втулки нижнего подшипника (16-130A или 16A-130B) с корпуса нижнего подшипника (16-140B или 16A-140C)."
    ),
    "Release the tab washer (17-110) and remove the bolts (17-100). Remove the tab washers (17-110).": (
        "Отогните усик отгибной шайбы (17-110) и снимите болты (17-100). Снимите отгибные шайбы (17-110)."
    ),
    "Remove the lock plate (17-120) and use the Pin Spanner 460007284 to remove the nut subassembly (17-130). Remove the rod (17-160) and the washer (17-170).": (
        "Снимите стопорную пластину (17-120) и используйте штифтовой ключ 460007284, чтобы снять подсборку гайки (17-130). Снимите шток (17-160) и шайбу (17-170)."
    ),
    "Use the Extractor 460006151/47 and the Drift 460004331/1 to remove the bearings (20-230 and 20-240).": (
        "Используйте съемник 460006151/47 и выколотку 460004331/1 для снятия подшипников (20-230 и 20-240)."
    ),
    "Paint the part: refer to REPAIR.": "Окрасьте деталь: см. раздел РЕМОНТ.",
    "Spring Data": "Данные пружины",
    "Slowly open the charging valve (17-20) and release all of the second stage nitrogen pressure.": (
        "Медленно откройте заправочный клапан (17-20) и стравите все давление азота второй ступени."
    ),
    "Slowly open the charging valve (13-60) and release all of the first stage nitrogen pressure.": (
        "Медленно откройте заправочный клапан (13-60) и стравите все давление азота первой ступени."
    ),
    "Use the Lifting Tackle 460006211 and install the sliding tube subassembly (17-240) in the Build Trolley 460007240.": (
        "Используйте подъемную оснастку 460006211 и установите подсборку скользящей трубы (17-240) в сборочную тележку 460007240."
    ),
    "Remove the upper bearing housing (15-40) and related parts as follows:": (
        "Снимите корпус верхнего подшипника (15-40) и связанные с ним детали следующим образом:"
    ),
    "Use Pin Spanner 460007279 to remove the upper bearing housing (15-40).": (
        "Используйте штифтовой ключ 460007279, чтобы снять корпус верхнего подшипника (15-40)."
    ),
    "Remove and discard the pins (15-120).": "Снимите и выбросьте штифты (15-120).",
    "Use the Torque Adapter 460007283, the Torque Reactor 460007278, the Holding Blocks 460006406 and the Bench Clamp MT1025 to remove the diaphragm subassembly (15-190), the compression orifice plate (15-220), the clapper seat (15-230) and the baffle (15-240).": (
        "Используйте моментный адаптер 460007283, упор реактивного момента 460007278, удерживающие блоки 460006406 и слесарные тиски MT1025, чтобы снять подсборку диафрагмы (15-190), пластину дроссельного отверстия сжатия (15-220), седло клапана (15-230) и перегородку (15-240)."
    ),
    "Use the Assembly/Extraction Tool 460006410 to remove the level tube (15-300) and remove the O-ring seal (15-310).": (
        "Используйте сборочно-демонтажный инструмент 460006410, чтобы снять уровневую трубку (15-300) и уплотнительное кольцо круглого сечения (15-310)."
    ),
    "Remove the lower bearing subassembly (16-110) and its related parts from the sliding tube subassembly (17-240).": (
        "Снимите сборку нижнего подшипника (16-110) и связанные с ней детали с подсборки скользящей трубы (17-240)."
    ),
    "Remove the lower bearing subassembly (16A-110D) and its related parts from the sliding tube subassembly (17-240).": (
        "Снимите сборку нижнего подшипника (16A-110D) и связанные с ней детали с подсборки скользящей трубы (17-240)."
    ),
    "Remove the inner liner (16A-117) from the lower bearing subassembly (16A-110D) and discard it.": (
        "Снимите внутренний вкладыш (16A-117) со сборки нижнего подшипника (16A-110D) и выбросьте его."
    ),
    "Remove the lower bearing (16A-150A) from the lower bearing housing subassembly (16A-120B). Discard the machined lower bearing (16A-150A).": (
        "Снимите нижний подшипник (16A-150A) с подсборки корпуса нижнего подшипника (16A-120B). Обработанный нижний подшипник (16A-150A) выбросьте."
    ),
    "Lower Bearing Subassembly (16-110D or 16A-110E) Post Ref. Code: 2253": (
        "Сборка нижнего подшипника (16-110D или 16A-110E), код ссылки ПОСЛЕ: 2253"
    ),
    "Remove the lower bearing subassembly (16-110D or 16A-110E) and its related parts from the sliding tube subassembly (17-240).": (
        "Снимите сборку нижнего подшипника (16-110D или 16A-110E) и связанные с ней детали с подсборки скользящей трубы (17-240)."
    ),
    "Remove the cap screws (17-30), the washers (17-40) and the valve support (17-50).": (
        "Снимите винты с цилиндрической головкой (17-30), шайбы (17-40) и опору клапана (17-50)."
    ),
    "Release the lock washer (17-90) and use the Torque Adapter 460006404 to remove the jacking dome (17-80). Remove the lock washer (17-90).": (
        "Отогните стопорную шайбу (17-90) и используйте моментный адаптер 460006404, чтобы снять поддомкратный купол (17-80). Снимите стопорную шайбу (17-90)."
    ),
    "Remove the cylinder (17-230) and its related parts from the sliding tube subassembly (17-240).": (
        "Снимите цилиндр (17-230) и связанные с ним детали с подсборки скользящей трубы (17-240)."
    ),
    "Hold the cylinder (17-230) in the Bench Clamp MT1025 and Holding Blocks MT1026/63.": (
        "Закрепите цилиндр (17-230) в слесарных тисках MT1025 и удерживающих блоках MT1026/63."
    ),
    "Remove the lubrication fittings (17-270) and the identification washers (17-280) from the sliding tube subassembly (17-240).": (
        "Снимите смазочные штуцеры (17-270) и идентификационные шайбы (17-280) с подсборки скользящей трубы (17-240)."
    ),
    "Remove the lubrication fittings (18-52) and identification washers (18-54).": (
        "Снимите смазочные штуцеры (18-52) и идентификационные шайбы (18-54)."
    ),
    "Remove the label (18-70) from the sliding tube (18-80).": (
        "Снимите табличку (18-70) со скользящей трубы (18-80)."
    ),
    "Use the Torque Adapter 460007232 to remove the locking nut (19-52). Remove the locking washer (19-54) and the outer race and the ball of the spherical bearing (19-50).": (
        "Используйте моментный адаптер 460007232, чтобы снять стопорную гайку (19-52). Снимите стопорную шайбу (19-54), наружную обойму и шарик сферического подшипника (19-50)."
    ),
    "NOTE: The outer race and the ball are parts of the spherical bearing (19-550).": (
        "ПРИМЕЧАНИЕ. Наружная обойма и шарик являются частями сферического подшипника (19-50)."
    ),
    "Remove the lubrication fitting (20-110) and the identification washer (20-120). Remove the lubrication adapter (20-130).": (
        "Снимите смазочный штуцер (20-110) и идентификационную шайбу (20-120). Снимите смазочный адаптер (20-130)."
    ),
    "Remove the lubrication fitting (20-140) and the identification washer (20-150). Remove the lubrication adapter (20-160).": (
        "Снимите смазочный штуцер (20-140) и идентификационную шайбу (20-150). Снимите смазочный адаптер (20-160)."
    ),
    "Remove the lubrication fitting (20-170) and the identification washer (20-180). Remove the lubrication adapter (20-190).": (
        "Снимите смазочный штуцер (20-170) и идентификационную шайбу (20-180). Снимите смазочный адаптер (20-190)."
    ),
    "Remove the lubrication fittings (20-200) and identification washers (20-210). Remove the lubrication adapters (20-220).": (
        "Снимите смазочные штуцеры (20-200) и идентификационные шайбы (20-210). Снимите смазочные адаптеры (20-220)."
    ),
    "Use the Hydraulic-Pneumatic Pump Set 460006497, the Bolt 460006498/7, the Press Pad 460006499/25 and the Extraction Tube 460004680 and remove the forward pintle bush (20-250A).": (
        "Используйте гидропневматическую насосную установку 460006497, болт 460006498/7, нажимную опору 460006499/25 и выпрессовочную трубку 460004680, чтобы снять втулку переднего штифта навеса (20-250A)."
    ),
    "Use the Extractor Plate 460007259/460006151/9 and the Drift 460004331/21 to remove the bushes (20-340 and 20-350).": (
        "Используйте плиту съемника 460007259/460006151/9 и выколотку 460004331/21, чтобы снять втулки (20-340 и 20-350)."
    ),
    "Use the Press Pad Assembly 460006267 and remove the drag arm sleeve (20-370A only).": (
        "Используйте сборку нажимной опоры 460006267 и снимите втулку тяги складывания (только 20-370A)."
    ),
    "Clean all the metal parts with white spirit, Material Ref. Item 11-524. Make sure that you fully remove all sealants, adhesives and jointing compounds.": (
        "Очистите все металлические детали уайт-спиритом, код ссылки материала 11-524. Убедитесь, что полностью удалены все герметики, клеевые составы и герметизирующие составы."
    ),
    "Use clean PVC or polythene gloves to prevent corrosion of metal parts.": (
        "Используйте чистые перчатки из ПВХ или полиэтилена, чтобы предотвратить коррозию металлических деталей."
    ),
    "Prevent corrosion of the metal parts that you do not immediately use for assembly procedures: refer to PCS-2800.": (
        "Предотвратите коррозию металлических деталей, которые не будут немедленно использованы при сборке: см. PCS-2800."
    ),
    "The procedure to examine the parts is in two levels:": (
        "Процедура проверки деталей состоит из двух уровней:"
    ),
    "Visually examine each part. Carefully examine changes of section and areas which contact sealing rings.": (
        "Визуально осмотрите каждую деталь. Особенно внимательно осмотрите переходы сечений и участки, контактирующие с уплотнительными кольцами."
    ),
    "Deterioration of protective treatment.": "Повреждение защитной обработки.",
    "Distortion and/or cracks.": "Деформация и/или трещины.",
    "Wear or fretting.": "Износ или фреттинг-коррозия.",
    "Scores, dents or burrs.": "Риски, вмятины или заусенцы.",
    "Measure all the parts that are in FITS AND CLEARANCES and compare with the dimensions in the table.": (
        "Измерьте все детали, указанные в разделе ПОСАДКИ И ЗАЗОРЫ, и сравните результаты с размерами, приведенными в таблице."
    ),
    "Special Dimension Check:": "Специальная проверка размеров:",
    "Examine the rod (17-160) for the diameter of radial damping holes. The diameter of each hole must be between 5,40 and 5,60 mm (0.213 and 0.220 in).": (
        "Проверьте шток (17-160) по диаметру радиальных демпфирующих отверстий. Диаметр каждого отверстия должен быть в пределах 5,40-5,60 мм (0.213-0.220 in)."
    ),
    "Examine the thread form of the diaphragm subassembly (15-190) and diaphragm (15-210A) with shadow graph projection.": (
        "Проверьте профиль резьбы подсборки диафрагмы (15-190) и диафрагмы (15-210A) с помощью теневой проекции."
    ),
    "Examine the 4 holes in the sliding tube (18-80) where the bracket (8-170) installs, for burrs. If you find burrs contact Safran Landing Systems who will supply an applicable repair.": (
        "Проверьте 4 отверстия в скользящей трубе (18-80), в которые устанавливается кронштейн (8-170), на наличие заусенцев. При обнаружении заусенцев свяжитесь с Safran Landing Systems, которая предоставит соответствующий ремонт."
    ),
    "NOTE: Use a good light source and 10x magnification to view the area, to look for burrs.": (
        "ПРИМЕЧАНИЕ. Используйте хороший источник света и 10-кратное увеличение для осмотра участка на наличие заусенцев."
    ),
    "Examine all parts shown in Tables 501 and 502 to the applicable NDT and information given.": (
        "Проверьте все детали, указанные в таблицах 501 и 502, соответствующим методом НК и с учетом приведенных указаний."
    ),
    "CAUTION: YOU MUST DISASSEMBLE ALL PARTS, THIS WILL INCLUDE THE BUSHES, THEY MUST BE REMOVED AND DISCARDED. YOU MUST APPLY THE NDT INSPECTION TO THE DETAIL LEVEL PART ONLY AS IDENTIFIED IN TABLES 501 AND 502. IF THE BUSHES ARE NOT REMOVED THE INSPECTION IS NOT COMPLETE FOR THE DETAIL PART AND DAMAGE CAN OCCUR.": (
        "ОСТОРОЖНО. НЕОБХОДИМО ПОЛНОСТЬЮ РАЗОБРАТЬ ВСЕ ДЕТАЛИ, ВКЛЮЧАЯ ВТУЛКИ; ИХ НЕОБХОДИМО СНЯТЬ И ВЫБРОСИТЬ. КОНТРОЛЬ НК СЛЕДУЕТ ВЫПОЛНЯТЬ ТОЛЬКО ДЛЯ ДЕТАЛИ НА УРОВНЕ ДЕТАЛИЗАЦИИ, УКАЗАННОЙ В ТАБЛИЦАХ 501 И 502. ЕСЛИ ВТУЛКИ НЕ СНЯТЫ, КОНТРОЛЬ ДЛЯ ДАННОЙ ДЕТАЛИ НЕ ЯВЛЯЕТСЯ ПОЛНЫМ, И ВОЗМОЖНО ПОВРЕЖДЕНИЕ."
    ),
    "Parts that are included in Tables 501 and 502 must be fully disassembled to the lowest detail level for NDT inspection. This includes the removal of all of the bushes.": (
        "Детали, указанные в таблицах 501 и 502, должны быть полностью разобраны до наименьшего уровня детализации для контроля НК. Это включает снятие всех втулок."
    ),
    "Examination of Magnetic Steel Parts by Non-destructive Testing Table 501": (
        "Контроль магнитных стальных деталей методами неразрушающего контроля. Таблица 501"
    ),
    "Examination of Non-Magnetic Parts by Non-destructive Testing Table 502": (
        "Контроль немагнитных деталей методами неразрушающего контроля. Таблица 502"
    ),
    "Repair Levels": "Уровни ремонта",
    "There are two levels of repair procedure for parts that are found to be unserviceable after inspection: refer to CHECK.": (
        "Существуют два уровня процедуры ремонта деталей, признанных негодными после контроля: см. ПРОВЕРКА."
    ),
    "Repair of surface damage.": "Ремонт поверхностных повреждений.",
    "Repair of wear or damage with an approved Messier-Dowty Limited or Safran Landing Systems repair.": (
        "Ремонт износа или повреждений с применением утвержденного ремонта Messier-Dowty Limited или Safran Landing Systems."
    ),
    "Repair isolated external scores, smooth dents and abrasions, that have no cracks and no effect on internal dimensions: refer to para (2). Such damage must not be:": (
        "Устраните отдельные наружные риски, плавные вмятины и потертости, не имеющие трещин и не влияющие на внутренние размеры: см. п. (2). Такие повреждения не должны быть:"
    ),
    "Less than one diameter from a hole and less than 6,35 mm (0.250 in) from a bearing surface": (
        "На расстоянии менее одного диаметра от отверстия и менее 6,35 мм (0.250 in) от поверхности под подшипник"
    ),
    "Remove burrs, corrosion and sharp edges: the area of damage must not be more than 645 mm2 (1.0 in2) for each 6450 mm2 (10.0 in2). Subsequently, remove 0,127 mm (0.0050 in) more of the material and repair the protective treatment.": (
        "Удалите заусенцы, коррозию и острые кромки: площадь повреждения не должна превышать 645 мм2 (1.0 in2) на каждые 6450 мм2 (10.0 in2). Затем снимите дополнительно 0,127 мм (0.0050 in) материала и восстановите защитную обработку."
    ),
    "In a bore that will not seal, ignore abrasions and small scores that have no burrs. If there are burrs, remove them plus 0,127 mm (0.0050 in) of material from the area. Repair the protective treatment.": (
        "В отверстии, не требующем герметизации, допускается не учитывать потертости и мелкие риски без заусенцев. Если имеются заусенцы, удалите их вместе с 0,127 мм (0.0050 in) материала с данного участка. Восстановите защитную обработку."
    ),
    "In a bore that will seal, polish scores to remove them. Make sure that the surface finish, concentricity and fits and clearances do not change.": (
        "В отверстии, подлежащем герметизации, отполируйте риски до их удаления. Убедитесь, что шероховатость поверхности, соосность, посадки и зазоры не изменились."
    ),
    "Remove burrs from external screw threads.": "Удалите заусенцы с наружной резьбы.",
    "Approved repairs are in para 4. The repairs in this CMM have been approved under Airbus’ EASA Design Organisation Approval No. EASA.21J.031.": (
        "Утвержденные ремонты приведены в п. 4. Ремонты, указанные в данном CMM, одобрены в рамках разрешения EASA на деятельность конструкторской организации Airbus № EASA.21J.031."
    ),
    "Unless instructions are different in the approved repair, the applicable tolerances are:": (
        "Если в утвержденном ремонте не указано иное, применяются следующие допуски:"
    ),
    "General tolerance: + or - 0,25 mm (0.010 in)": "Общий допуск: + или - 0,25 мм (0.010 in)",
    "Holes that are drilled or machined: + 0,25 to - 0,05 mm (+ 0.010 to - 0.002 in)": (
        "Для просверленных или механически обработанных отверстий: +0,25 до -0,05 мм (+0.010 до -0.002 in)"
    ),
    "Angular tolerance: + or - 0,5 degree.": "Угловой допуск: + или - 0,5°.",
    "Before you repair a part that is identified with a concession, salvage or repair number, write to Safran Landing Systems for approval. Such numbers are adjacent to the part number, for example:": (
        "Перед ремонтом детали, обозначенной номером отступления, восстановления или ремонта, направьте запрос в Safran Landing Systems для получения одобрения. Такие номера указываются рядом с номером детали, например:"
    ),
    "If the repairs in this manual cannot correct the wear or damage to the part, write to Safran Landing Systems: refer to M-DLPS3002.": (
        "Если ремонты, приведенные в данном руководстве, не позволяют устранить износ или повреждение детали, направьте запрос в Safran Landing Systems: см. M-DLPS3002."
    ),
    "Clean the parts after repair: refer to CLEANING.": "Очистите детали после ремонта: см. ОЧИСТКА.",
    "Identify the parts after repair with the Messier-Dowty Limited or Safran Landing Systems Repair Number: refer to the applicable repair for instructions.": (
        "Нанесите на детали после ремонта ремонтный номер Messier-Dowty Limited или Safran Landing Systems: см. соответствующий ремонт для указаний."
    ),
    "Repair damage to small areas of cadmium plated surfaces: refer to PCS-2141.": (
        "Устраните повреждения небольших участков кадмированного покрытия: см. PCS-2141."
    ),
    "Repair damage to small areas of anodized surfaces: refer to PCS-2220.": (
        "Устраните повреждения небольших участков анодированной поверхности: см. PCS-2220."
    ),
    "Repair damage to small areas of paint finish: refer to M-DLPS1003-1, use paint to PCS-2500.": (
        "Устраните повреждения небольших участков лакокрасочного покрытия: см. M-DLPS1003-1, использовать краску по PCS-2500."
    ),
    "Chip damage of less than 10,0 mm2 (0.015 in2) can be restored with Sermetel 249 with Sermetel 273 catalyst: refer to M-DLPS637 (cold rework only).": (
        "Сколы площадью менее 10,0 мм2 (0.015 in2) допускается восстанавливать материалом Sermetel 249 с катализатором Sermetel 273: см. M-DLPS637 (только холодная доработка)."
    ),
    "Protective Treatment Replacement": "Замена защитной обработки",
    "Protective treatment replacement procedures and the applicable parts are given in Table 601.": (
        "Процедуры замены защитной обработки и соответствующие детали приведены в таблице 601."
    ),
    "Protective Treatment - Sequence of Application": "Защитная обработка - Последовательность нанесения",
    "If you apply protective treatment processes that include Sermetel W to any ultra high tensile (UHT) steel part, the sequence of the processes is important. The sequence of the protective treatment processes must be as follows:": (
        "Если к любой детали из сверхвысокопрочной стали (UHT) применяются процессы защитной обработки, включающие Sermetel W, последовательность процессов имеет принципиальное значение. Процессы защитной обработки должны выполняться в следующей последовательности:"
    ),
    "Surface Damage": "Повреждение поверхности",
    "Identification": "Идентификация",
    "Cadmium Plated Surfaces": "Кадмированные поверхности",
    "Anodized Surfaces": "Анодированные поверхности",
    "Paint Finish": "Лакокрасочное покрытие",
    "Item No.": "Поз. №",
    "Concession - CON 14235": "Отступление - CON 14235",
    "Salvage - 440015644": "Восстановление - 440015644",
    "NOTE:-": "ПРИМЕЧАНИЕ:",
    "Chromium plate processes.": "Процессы хромирования.",
    "Cadmium plate processes.": "Процессы кадмирования.",
    "Sermetel W processes.": "Процессы нанесения покрытия Sermetel W.",
    "Paint processes.": "Процессы окраски.",
    "More than 19,00 mm (0.750 in) in length": "Длиной более 19,00 мм (0.750 in)",
    "More than 0,76 mm (0.030 in) in depth": "Глубиной более 0,76 мм (0.030 in)",
    "On a radius.": "На радиусном переходе.",
    "15,00 mm (0.591 in) diameter areas around the holes on the inside face of one flange.": (
        "Участки диаметром 15,00 мм (0.591 in) вокруг отверстий на внутренней поверхности одного фланца."
    ),
    "The cadmium plate must overlap the chromium plate run out. Bare metal not permitted.": (
        "Кадмиевое покрытие должно перекрывать зону схода хромового покрытия. Оголенный металл не допускается."
    ),
    "Apply over layer with red color paint AVIOX 77702": "Нанести верхний слой красной краски AVIOX 77702",
    "Make sure that the outer diameter is not more than 12,540 mm (0.493 in) after the paint.": (
        "Убедитесь, что наружный диаметр после окраски не превышает 12,540 мм (0.493 in)."
    ),
    "E. Do not apply zinc-nickel plate in areas F.": "E. Не наносить цинк-никелевое покрытие на участки F.",
    "PRIMER PAINT TO PCS-2500 OVER SERMETEL W TO LENGTH 305,00mm (12.000in) for (18-80)": (
        "ГРУНТОВОЧНАЯ КРАСКА ПО PCS-2500 ПОВЕРХ SERMETEL W НА ДЛИНЕ 305,00mm (12.000in) ДЛЯ (18-80)"
    ),
    "220,00mm (8.661in) for (18-80A)": "220,00mm (8.661in) ДЛЯ (18-80A)",
    "SERMETEL W ON INTERNAL": "ПОКРЫТИЕ SERMETEL W НА ВНУТРЕННЕМ",
    "DIAMETER OVER LENGTH A A": "ДИАМЕТРЕ ПО ДЛИНЕ A-A",
    "TO IFC 40-860-03MD": "ПО IFC 40-860-03MD",
    "SERMETEL W TO IFC 40-860-03MD": "ПОКРЫТИЕ SERMETEL W ПО IFC 40-860-03MD",
    "SERMETEL W TO": "ПОКРЫТИЕ SERMETEL W ПО",
    "SERMETEL W": "ПОКРЫТИЕ SERMETEL W",
    "Sermetel W": "Покрытие Sermetel W",
    "CHROMIUM PLATE": "ХРОМОВОЕ ПОКРЫТИЕ",
    "PRIMER PAINT TO": "ГРУНТОВОЧНАЯ КРАСКА ПО",
    "PRIMER PAINT TO PCS-2500": "ГРУНТОВОЧНАЯ КРАСКА ПО PCS-2500",
    "PCS-2500. Do not paint:": "PCS-2500. Не окрашивать:",
    "LENGTH OF CADMIUM PLATE": "ДЛИНА КАДМИЕВОГО ПОКРЫТИЯ",
    "Zinc-nickel plate is optional on areas": "Цинк-никелевое покрытие допускается на участках",
    "(HOLE TO DEPTH OF 16,0mm (0.63in)": "(ОТВЕРСТИЕ НА ГЛУБИНУ 16,0mm (0.63in)",
    "TO A DEPTH OF 16,000mm (0.6299in)": "НА ГЛУБИНУ 16,000mm (0.6299in)",
    "M-DLPS102-1. Do not include areas": "M-DLPS102-1. Не включать участки",
    "Apply primer 463-12-8.": "Нанести грунт 463-12-8.",
    "Apply cadmium plate: refer to": "Нанесите кадмиевое покрытие: см.",
    "Apply cadmium plate all over: refer to M-DLPS100-1.": (
        "Нанесите кадмиевое покрытие по всей поверхности: см. M-DLPS100-1."
    ),
    "Apply cadmium plate all over: refer to M-DLPS100-2. Make the cadmium plate thickness between 0,010": (
        "Нанесите кадмиевое покрытие по всей поверхности: см. M-DLPS100-2. Толщина кадмиевого покрытия должна составлять 0,010"
    ),
    "Apply cadmium plate: refer to M-DLPS100-2. Paint: refer to": (
        "Нанесите кадмиевое покрытие: см. M-DLPS100-2. Окраска: см."
    ),
    "Apply paint: refer to M-DLPS1003-1 and PCS-2500. Do not paint:": (
        "Нанесите лакокрасочное покрытие: см. M-DLPS1003-1 и PCS-2500. Не окрашивать:"
    ),
    "the 1,5 mm (0.060 in) hole for the Bowden cable(1-45).": (
        "Отверстие 1,5 мм (0.060 in) под трос Боудена (1-45)."
    ),
    "Refer to Figure 639. Apply cadmium plate: refer to PCS-2100. Do not include areas that have chromium plate. Make the cadmium plate thickness between 0,010 and 0,020": (
        "См. рисунок 639. Нанесите кадмиевое покрытие: см. PCS-2100. Не включать участки с хромовым покрытием. Толщина кадмиевого покрытия должна быть в пределах 0,010-0,020"
    ),
    "Apply primer paint to the areas A: refer to PCS-2500. Apply paint all over: refer to PCS-2500. Do not apply paint to:": (
        "Нанесите грунтовочную краску на участки A: см. PCS-2500. Нанесите лакокрасочное покрытие по всей поверхности: см. PCS-2500. Не наносить лакокрасочное покрытие на:"
    ),
    "M-DLPS100-1. The cadmium plate must be 0,010 to 0,015 mm (0.0004 to 0.0006 in) thick. Apply paint all over: refer to PCS-2500. Do not apply paint to the thread or to the surfaces that enter the transfer block (2-340 and 2-350)": (
        "M-DLPS100-1. Кадмиевое покрытие должно иметь толщину 0,010-0,015 мм (0.0004-0.0006 in). Нанесите лакокрасочное покрытие по всей поверхности: см. PCS-2500. Не наносить лакокрасочное покрытие на резьбу и поверхности, входящие в переходной блок (2-340 и 2-350)."
    ),
    "M-DLPS100-2. The cadmium plate must be 0,010 to 0,015 mm (0.0004 to 0.0006 in) thick.": (
        "M-DLPS100-2. Кадмиевое покрытие должно иметь толщину 0,010-0,015 мм (0.0004-0.0006 in)."
    ),
    "PCS-2500. Apply primer paint only to the contact faces of the flanges. Do not paint:": (
        "PCS-2500. Нанесите грунтовочную краску только на контактные поверхности фланцев. Не окрашивать:"
    ),
    "Apply one layer of primer paint only to the areas A: refer to PCS-2500. Apply paint all over as per PCS-2500 but not to:": (
        "Нанесите один слой грунтовочной краски только на участки A: см. PCS-2500. Нанесите лакокрасочное покрытие по всей поверхности по PCS-2500, но не на:"
    ),
    "M-DLPS1003-1 and PCS-2500. Do": "M-DLPS1003-1 и PCS-2500. Не",
    "M-DLPS102-1. Apply Alocrom 1200 to the spotfaces A: refer to": (
        "M-DLPS102-1. Нанесите Alocrom 1200 на подрезки площадок A: см."
    ),
    "Paint external areas only: refer to M-DLPS1003-1 and PCS-2500.": (
        "Окрашивать только наружные участки: см. M-DLPS1003-1 и PCS-2500."
    ),
    "Apply primer paint only to the contact face. Paint must not go in the bores.": (
        "Нанесите грунтовочную краску только на контактную поверхность. Краска не должна попадать в отверстия."
    ),
    "Refer to Figure 636. Apply cadmium": "См. рисунок 636. Нанесите кадмиевое",
    "plate all over including the holes less than diameter 10 mm (0.393 in) but not to the chromium plated areas and areas A: refer to PCS-2100. Make the cadmium plate thickness between 0,010 and 0,020 mm": (
        "покрытие по всей поверхности, включая отверстия диаметром менее 10 мм (0.393 in), но не на участки с хромовым покрытием и участки A: см. PCS-2100. Толщина кадмиевого покрытия должна быть в пределах 0,010-0,020 мм"
    ),
    "(0.0004 and 0.0008 in). The cadmium plate must overlap the chromium plate run out. The cadmium plate is optional on the lubrication fitting bores where the lubrication adaptors (20-130),": (
        "(0.0004 and 0.0008 in). Кадмиевое покрытие должно перекрывать зону схода хромового покрытия. Кадмиевое покрытие допускается в отверстиях под смазочные штуцеры, где устанавливаются смазочные адаптеры (20-130),"
    ),
    "(20-160), (20-190) and (20-220) will": "(20-160), (20-190) и (20-220) будут",
    "Apply paint all over but not on the chromium plated areas, the areas A, B, C and on the lubrication fitting bores where the lubrication adaptors (20-130), (20-160),": (
        "Нанесите лакокрасочное покрытие по всей поверхности, но не на участки с хромовым покрытием, участки A, B, C и отверстия под смазочные штуцеры, где устанавливаются смазочные адаптеры (20-130), (20-160),"
    ),
    "Apply paint all over but not on the chromium plated areas, the areas A, B, C and on the lubrication fitting bores where the lubrication adaptors (20-130), (20-160), (20-": (
        "Нанесите лакокрасочное покрытие по всей поверхности, но не на участки с хромовым покрытием, участки A, B, C и отверстия под смазочные штуцеры, где устанавливаются смазочные адаптеры (20-130), (20-160), (20-"
    ),
    "Refer to Figure 631. Chromic acid": "См. рисунок 631. Анодируйте в хромовой",
    "anodise all over, but not the spotface A and area B: refer to M-DLPS102-1. Apply Alocrom 1200 to the spotface A and area B: refer to M-DLPS114. Apply paint all over but not to the spotface A, areas B and holes C. Apply only primer paint to the holes C: refer to PCS-2500.": (
        "анодируйте по всей поверхности, кроме подрезки площадки A и участка B: см. M-DLPS102-1. Нанесите Alocrom 1200 на подрезку площадки A и участок B: см. M-DLPS114. Нанесите лакокрасочное покрытие по всей поверхности, но не на подрезку площадки A, участки B и отверстия C. Нанесите только грунтовочную краску в отверстия C: см. PCS-2500."
    ),
    "Refer to Figure 601. Anodise all over, but not areas A. Apply Alocrom 1200 to areas A: refer to PCS-2220. Paint all over but do not include on areas A and bores B. Apply primer to bores B: refer to PCS-2500.": (
        "См. рисунок 601. Выполните анодирование по всей поверхности, кроме участков A. Нанесите Alocrom 1200 на участки A: см. PCS-2220. Нанесите лакокрасочное покрытие по всей поверхности, но не на участки A и отверстия B. Нанесите грунт в отверстия B: см. PCS-2500."
    ),
    "Before installation of bushes: Apply primer paint to areas A but not to the areas B: refer to PCS-2500.": (
        "Перед установкой втулок: нанесите грунтовочную краску на участки A, но не на участки B: см. PCS-2500."
    ),
    "After installation of bushes: Apply paint to the areas C but not to the areas D: refer to PCS-2500.": (
        "После установки втулок: нанесите лакокрасочное покрытие на участки C, но не на участки D: см. PCS-2500."
    ),
    "Refer to Figure 616. Apply cadmium plate internally and externally on area A: refer to M-DLPS100-2. Do not apply cadmium plate to bores B": (
        "См. рисунок 616. Нанесите кадмиевое покрытие изнутри и снаружи на участок A: см. M-DLPS100-2. Не наносить кадмиевое покрытие на отверстия B"
    ),
    "PCS-2100. Make the cadmium plate thickness 0,010 to 0,015 mm": (
        "PCS-2100. Толщина кадмиевого покрытия должна составлять 0,010-0,015 мм"
    ),
    "Refer to para 3.C. and Figure 618.": "См. п. 3.C. и рисунок 618.",
    "Apply cadmium plate to PCS-2100. The cadmium plate thickness should be between 0,010 and 0,015 mm (0.0004 and 0.0006 in). Do not apply cadmium plate:": (
        "Нанесите кадмиевое покрытие по PCS-2100. Толщина кадмиевого покрытия должна быть в пределах 0,010-0,015 мм (0.0004 и 0.0006 in). Не наносить кадмиевое покрытие:"
    ),
    "where identified on Figure 618. Apply Sermetel W where shown: refer to IFC 40-860-03MD. Apply primer as shown and finish paint to PCS-2500. Do not apply paint:": (
        "Где указано на рисунке 618. Нанесите покрытие Sermetel W в указанных местах: см. IFC 40-860-03MD. Нанесите грунт, как показано, и финишное лакокрасочное покрытие по PCS-2500. Не наносить лакокрасочное покрытие:"
    ),
    "Apply sermetel W only to the areas C: refer to IFC 40-860-03MD.": (
        "Нанесите покрытие Sermetel W только на участки C: см. IFC 40-860-03MD."
    ),
    "Sermetel is optional in areas D. If sermetel is not applied in areas D, apply cadmium plate to areas D. Make the sermetel W coating thickness between 0,025 and": (
        "Покрытие Sermetel допускается на участках D. Если покрытие Sermetel не наносится на участках D, нанесите кадмиевое покрытие на участки D. Толщина покрытия Sermetel W должна быть в пределах 0,025"
    ),
    "0,050 mm (0.001 and 0.002 in). The Sermetel W coating must overlap the chromium plated areas and cadmium plated areas.": (
        "0,050 мм (0.001 и 0.002 in). Покрытие Sermetel W должно перекрывать участки с хромовым и кадмиевым покрытием."
    ),
    "not paint the holes and areas A, B and C. Apply primer to areas C: refer to PCS-2500. Refer to Figure 602.": (
        "Не окрашивать отверстия и участки A, B и C. Нанесите грунт на участки C: см. PCS-2500. См. рисунок 602."
    ),
    "Refer to Figure 602. Chromic acid anodize: refer to MIL-A-8625 Type 1B, Class 1, do not include areas A. Apply Alocrom to areas A: refer to PCS-2220, Type 2.": (
        "См. рисунок 602. Выполните анодирование в хромовой кислоте по MIL-A-8625, тип 1B, класс 1, не включая участки A. Нанесите Alocrom на участки A: см. PCS-2220, тип 2."
    ),
    "Paint: refer to M-DLPS1003-1 and PCS-2500. Apply primer only to the flanges of the bushes. Do not paint:": (
        "Окраска: см. M-DLPS1003-1 и PCS-2500. Наносите грунт только на фланцы втулок. Не окрашивать:"
    ),
    "the holes for the lubrication fittings.": "отверстия под смазочные штуцеры.",
    "M-DLPS102-1. Apply primer: refer to PCS-2500. Do not include the holes for lubrication": (
        "M-DLPS102-1. Нанесите грунт: см. PCS-2500. Не включать отверстия под смазку"
    ),
    "M-DLPS100-1. The cadmium plate must be 0,004 to 0,005 mm (0.00015 to 0.0002 in) thick. Paint to": (
        "M-DLPS100-1. Толщина кадмиевого покрытия должна составлять 0,004-0,005 мм (0.00015-0.0002 in). Окраска по"
    ),
    "PCS-2500 but not on the threaded diameter, shank and the adjacent flange face. Apply a light coat of primer paint only to the bore for the spherical bearing (4-50).": (
        "PCS-2500, но не на резьбовой диаметр, хвостовик и прилегающую поверхность фланца. Нанесите тонкий слой грунтовочной краски только в отверстие под сферический подшипник (4-50)."
    ),
    "Passivate: refer to AMS2700": "Пассивируйте: см. AMS2700",
    "Passivate: refer to AMS2700.": "Пассивируйте: см. AMS2700.",
    "Protective treatment is not necessary": "Защитная обработка не требуется",
    "Upper diaphragm tube sub-assembly": "Подсборка верхней диафрагменной трубы",
    "Locking nut": "Стопорная гайка",
    "Apply cadmium plate, but not to the": "Нанести кадмиевое покрытие, но не на",
    "Do not paint areas F.": "Не окрашивать участки F.",
    "NO CADMIUM PLATE NO PAINT": "БЕЗ КАДМИЕВОГО ПОКРЫТИЯ НЕ ОКРАШИВАТЬ",
    "NO CADMIUM PLATE NO PAINT (2 PLACES)": "БЕЗ КАДМИЕВОГО ПОКРЫТИЯ НЕ ОКРАШИВАТЬ (2 МЕСТА)",
    "NO PAINT": "НЕ ОКРАШИВАТЬ",
    "NO PAINT (Qty 24)": "НЕ ОКРАШИВАТЬ (КОЛ-ВО 24)",
    "NO PAINT 13,00mm (0.512in) DIA.": "НЕ ОКРАШИВАТЬ 13,00mm (0.512in) ДИАМ.",
    "the area that has chromium plate": "участок с хромовым покрытием",
    "areas that have chromium plate": "участки с хромовым покрытием",
    "chromium plated area C": "участок C с хромовым покрытием",
    "the 3 holes in face D": "3 отверстия на поверхности D",
    "the 3 holes in face D.": "3 отверстия на поверхности D.",
    "the two radial holes.": "два радиальных отверстия.",
    "the hole through the threads.": "отверстие через резьбу.",
    "to the areas where Sermetel W is applied": "на участки, где нанесено покрытие Sermetel W",
    "to the threaded surfaces.": "на резьбовые поверхности.",
    "the thread undercut": "подрез резьбы",
    "not paint the undercut below the head.": "не окрашивать подрез под головкой.",
    "The threads": "Резьбу",
    "the threads": "резьбу",
    "the thread": "резьбу",
    "the holes": "отверстия",
    "the split pin hole.": "отверстие под шплинт.",
    "the thread and undercut": "резьбу и подрез",
    "the two holes through the end. Apply a thin coat of primer paint to the holes through the end: refer to PCS-2500.": (
        "Два отверстия в торце. Нанесите тонкий слой грунтовочной краски в отверстия торца: см. PCS-2500."
    ),
    "areas A": "участки A",
    "areas B": "участки B",
    "areas C.": "участки C.",
    "the areas A.": "участки A.",
    "the area B": "участок B",
    "B INTERNALLY": "B ВНУТРИ",
    "C INTERNALLY": "C ВНУТРИ",
    "FACE D": "ПОВЕРХНОСТЬ D",
    "0.0006 in) thick.": "0.0006 in) толщиной.",
    "Figure Deleted Figure 609": "Рисунок удален, рисунок 609",
    "Figure Deleted Figure 611": "Рисунок удален, рисунок 611",
    "paint:": "окрасить:",
    "not cadmium plate:": "не кадмировать:",
    "AMS5659 condition H1025": "AMS5659, состояние H1025",
    "MIL-A-8625 Type IB, Class 1.": "MIL-A-8625, тип IB, класс 1.",
    "C B SECTION Y-Y": "C B СЕЧЕНИЕ Y-Y",
    "X V SECTION U-U": "X V СЕЧЕНИЕ U-U",
    "B PAINT TO": "B ОКРАСКА ПО",
    "CAUTION: DO NOT REPAIR A PART WITH A PROCEDURE THAT IS NOT APPROVED.": (
        "ПРЕДОСТЕРЕЖЕНИЕ: НЕ РЕМОНТИРУЙТЕ ДЕТАЛЬ ПО НЕУТВЕРЖДЕННОЙ ПРОЦЕДУРЕ."
    ),
    "CAUTION: YOU MUST COMPLETE THE PROCESSES THAT FOLLOW IN THE SEQUENCE SHOWN. FAILURE TO DO THE PROCESSES IN THE CORRECT SEQUENCE CAN DAMAGE THE SLIDING TUBE (18-80) OR (18-80A) OR": (
        "ПРЕДОСТЕРЕЖЕНИЕ: НЕОБХОДИМО ВЫПОЛНЯТЬ ПРИВЕДЕННЫЕ НИЖЕ ПРОЦЕССЫ В УКАЗАННОЙ ПОСЛЕДОВАТЕЛЬНОСТИ. НАРУШЕНИЕ ЭТОЙ ПОСЛЕДОВАТЕЛЬНОСТИ МОЖЕТ ПОВРЕДИТЬ СКОЛЬЗЯЩУЮ ТРУБУ (18-80) ИЛИ (18-80A) ИЛИ"
    ),
    "(18-80B) OR REDUCE THE EFFECT OF THE PROTECTIVE TREATMENTS.": (
        "(18-80B) ИЛИ СНИЗИТЬ ЭФФЕКТИВНОСТЬ ЗАЩИТНЫХ ОБРАБОТОК."
    ),
    "(0.002 to 0.005in) DEEP": "(0.002-0.005in) ГЛУБ.",
    "(0 to 0.078in) RUNOUT": "(0-0.078in) БИЕНИЕ",
    "(1.161 to 1.200in) DIA. SPOTFACE (Qty. 24)": (
        "(1.161-1.200in) ДИАМ. ПОДРЕЗКА ПЛОЩАДКИ (КОЛ-ВО 24)"
    ),
    "(0 to 0.078in) RUNOUT (BOTH SIDES)": "(0-0.078in) БИЕНИЕ (С ОБЕИХ СТОРОН)",
    "LIMIT OF SERMETEL W TERMINATION FROM CENTER 84,00 to 96,00mm": (
        "ПРЕДЕЛ ПОКРЫТИЯ SERMETEL W ОТ ЦЕНТРА 84,00-96,00mm"
    ),
    "LIMIT OF SERMETEL W TERMINATION FROM CENTER 95,50 TO 107,50MM": (
        "ПРЕДЕЛ ПОКРЫТИЯ SERMETEL W ОТ ЦЕНТРА 95,50-107,50MM"
    ),
    "INTERNALLY": "ВНУТРИ",
    "16,0mm (0,63in) DIA. THIS FACE ONLY": "16,0mm (0,63in) ДИАМ. ТОЛЬКО НА ЭТОЙ ПОВЕРХНОСТИ",
    "22,00mm (0.866in) DIA. THIS FACE ONLY": "22,00mm (0.866in) ДИАМ. ТОЛЬКО НА ЭТОЙ ПОВЕРХНОСТИ",
    "FROM THIS SURFACE)": "ОТ ЭТОЙ ПОВЕРХНОСТИ)",
    "IN THIS FACE ONLY": "ТОЛЬКО НА ЭТОЙ ПОВЕРХНОСТИ",
    "FROM OUTSIDE FACE": "ОТ НАРУЖНОЙ ПОВЕРХНОСТИ",
    "Apply cadmium plate all over: refer to M-DLPS100-1. The cadmium plate must be 0,010 to 0,015 mm (0.0004 to 0.0006 in) thick. Apply paint all over: refer to PCS-2500. Do not apply paint to the thread or to the surfaces that enter the transfer block (2-340 and 2-350)": (
        "Нанесите кадмиевое покрытие по всей поверхности: см. M-DLPS100-1. Кадмиевое покрытие должно иметь толщину 0,010-0,015 мм (0.0004-0.0006 in). Нанесите лакокрасочное покрытие по всей поверхности: см. PCS-2500. Не наносить лакокрасочное покрытие на резьбу и поверхности, входящие в переходной блок (2-340 и 2-350)."
    ),
    "Apply cadmium plate: refer to M-DLPS100-2. The cadmium plate must be 0,010 to 0,015 mm (0.0004 to 0.0006 in) thick. Paint: refer to M-DLPS1003-1 and PCS-2500. Do not paint areas A and B. Apply primer to area A: refer to Figure 604.": (
        "Нанесите кадмиевое покрытие: см. M-DLPS100-2. Кадмиевое покрытие должно иметь толщину 0,010-0,015 мм (0.0004-0.0006 in). Окраска: см. M-DLPS1003-1 и PCS-2500. Не окрашивать участки A и B. Нанесите грунт на участок A: см. рисунок 604."
    ),
    "Apply cadmium plate all over: refer to M-DLPS100-2S. The cadmium plate must be 0,010 to 0,015 mm (0.0004 to 0.0006 in) thick.": (
        "Нанесите кадмиевое покрытие по всей поверхности: см. M-DLPS100-2S. Кадмиевое покрытие должно иметь толщину 0,010-0,015 мм (0.0004-0.0006 in)."
    ),
    "Refer to Figure 626. Chromic acid anodise all over but not the spotfaces A: refer to M-DLPS102-1. Apply Alocrom 1200 to the areas A: refer to M-DLPS114. Apply paint all over but not to the areas A, B, C and D: refer to PCS-2500. Primer paint only on faces D. Apply light coat of primer to area B.": (
        "См. рисунок 626. Анодируйте в хромовой кислоте по всей поверхности, кроме подрезок площадок A: см. M-DLPS102-1. Нанесите Alocrom 1200 на участки A: см. M-DLPS114. Нанесите лакокрасочное покрытие по всей поверхности, но не на участки A, B, C и D: см. PCS-2500. Нанесите грунтовочную краску только на поверхности D. Нанесите тонкий слой грунта на участок B."
    ),
    "Refer to Figure 627. Chromic acid anodise all over but not the spotfaces A: refer to MIL-A-8625 Type 1B, Class 1. Apply Alocrom 1200 to the areas A: refer to PCS-2220 Type 2. Apply one coat of primer to the areas B: refer to PCS-2500. Apply primer only to the areas D: refer to PCS-2500. Apply paint all over but not to the areas A, C and D: refer to PCS-2500.": (
        "См. рисунок 627. Анодируйте в хромовой кислоте по всей поверхности, кроме подрезок площадок A: см. MIL-A-8625, тип 1B, класс 1. Нанесите Alocrom 1200 на участки A: см. PCS-2220, тип 2. Нанесите один слой грунта на участки B: см. PCS-2500. Нанесите грунт только на участки D: см. PCS-2500. Нанесите лакокрасочное покрытие по всей поверхности, но не на участки A, C и D: см. PCS-2500."
    ),
    "Apply cadmium plate all over: refer to M-DLPS100-2. Make the cadmium plate thickness between 0,010 and 0,015 mm (0.0004 and 0.0005 in).": (
        "Нанесите кадмиевое покрытие по всей поверхности: см. M-DLPS100-2. Толщина кадмиевого покрытия должна составлять 0,010-0,015 мм (0.0004-0.0005 in)."
    ),
    "Refer to Figure 634. Chromic acid anodise all over but not the spotface A: refer to M-DLPS102-1. Apply Alocrom 1200 to the areas A: refer to M-DLPS114. Apply paint all over but not to the areas A, B, C and D: refer to PCS-2500. Apply a light coat of primer to the hole B: refer to PCS-2500. Apply only primer to face D: refer to PCS-2500.": (
        "См. рисунок 634. Анодируйте в хромовой кислоте по всей поверхности, кроме подрезки площадки A: см. M-DLPS102-1. Нанесите Alocrom 1200 на участки A: см. M-DLPS114. Нанесите лакокрасочное покрытие по всей поверхности, но не на участки A, B, C и D: см. PCS-2500. Нанесите тонкий слой грунта в отверстие B: см. PCS-2500. Нанесите грунт только на поверхность D: см. PCS-2500."
    ),
    "Refer to Figure 613. Apply cadmium plate: refer to M-DLPS100-2. Do not include areas that have chromium plate. Paint area A: refer to M-DLPS1003-1 and PCS-2500.": (
        "См. рисунок 613. Нанесите кадмиевое покрытие: см. M-DLPS100-2. Не включать участки с хромовым покрытием. Нанесите лакокрасочное покрытие на участок A: см. M-DLPS1003-1 и PCS-2500."
    ),
    "Refer to Figure 613. Apply cadmium plate: refer to PCS-2101. Make the cadmium plate thickness between 0,010 and 0,015 mm (0.0004 and 0.0005 in). Do not include areas that have chromium plate. Paint area A: refer to PCS-2500.": (
        "См. рисунок 613. Нанесите кадмиевое покрытие: см. PCS-2101. Толщина кадмиевого покрытия должна быть в пределах 0,010-0,015 мм (0.0004-0.0005 in). Не включать участки с хромовым покрытием. Нанесите лакокрасочное покрытие на участок A: см. PCS-2500."
    ),
    "Apply cadmium plate: refer to M-DLPS100-2. Paint: refer to M-DLPS1003-1 and PCS-2500. Do not paint the screw threads and the face that touches the wheel bearings.": (
        "Нанесите кадмиевое покрытие: см. M-DLPS100-2. Окраска: см. M-DLPS1003-1 и PCS-2500. Не окрашивать резьбу и поверхность, контактирующую с колесными подшипниками."
    ),
    "Refer to Figure 630. Before installation of bushes: Apply primer paint to areas A but not to the areas B: refer to PCS-2500. After installation of bushes: Apply paint to the areas C but not to the areas D: refer to PCS-2500.": (
        "См. рисунок 630. Перед установкой втулок: нанесите грунтовочную краску на участки A, но не на участки B: см. PCS-2500. После установки втулок: нанесите лакокрасочное покрытие на участки C, но не на участки D: см. PCS-2500."
    ),
    "Refer to Figure 623. Apply cadmium plate internally and externally over area A: refer to PCS-2101. Make the cadmium plate thickness between 0,010 and 0,015 mm (0.0004 and 0.0006 in).": (
        "См. рисунок 623. Нанесите кадмиевое покрытие изнутри и снаружи на участок A: см. PCS-2101. Толщина кадмиевого покрытия должна быть в пределах 0,010-0,015 мм (0.0004-0.0006 in)."
    ),
    "Refer to Figure 617. Cadmium plate all over to M-DLPS100-2. Make cadmium plate thickness 0,010 to 0,015 mm (0.0004 to 0.0006 in). Do not cadmium plate: the area B chromium plated area C the 3 holes in face D Paint areas A: refer to PCS-2500. Primer paint only on face D and areas E and F including the chamfer. Do not paint: chromium plated area C the 3 holes in face D.": (
        "См. рисунок 617. Нанесите кадмиевое покрытие по всей поверхности по M-DLPS100-2. Толщина кадмиевого покрытия должна составлять 0,010-0,015 мм (0.0004-0.0006 in). Не кадмировать: участок B, участок C с хромовым покрытием и 3 отверстия на поверхности D. Нанесите лакокрасочное покрытие на участки A: см. PCS-2500. Нанесите грунтовочную краску только на поверхность D и участки E и F, включая фаску. Не окрашивать: участок C с хромовым покрытием и 3 отверстия на поверхности D."
    ),
    "Refer to Figure 624. Apply cadmium plate all over but not to the areas A: refer to PCS-2101. Cadmium plate is optional on radii and chamfer B. Make the cadmium plate thickness between 0,010 and 0,015 mm (0.0004 and 0.0006 in). Apply only primer paint to areas C including chamfer: refer to PCS-2500. Apply paint to areas D: refer to PCS-2500. No bare cadmium permitted.": (
        "См. рисунок 624. Нанесите кадмиевое покрытие по всей поверхности, но не на участки A: см. PCS-2101. Кадмиевое покрытие допускается на радиусные участки и фаску B. Толщина кадмиевого покрытия должна быть в пределах 0,010-0,015 мм (0.0004-0.0006 in). Нанесите только грунтовочную краску на участки C, включая фаску: см. PCS-2500. Нанесите лакокрасочное покрытие на участки D: см. PCS-2500. Оголенный кадмий не допускается."
    ),
    "Refer to para 3.C. and Figure 618. Apply cadmium plate to PCS-2100. The cadmium plate thickness should be between 0,010 and 0,015 mm (0.0004 and 0.0006 in). Do not apply cadmium plate: to the areas where Sermetel W is applied to the chromium plated areas where identified on Figure 618. Apply Sermetel W where shown: refer to IFC 40-860-03MD. Apply primer as shown and finish paint to PCS-2500. Do not apply paint: where identified on Figure 618 to the chromium plated areas. to the threaded surfaces.": (
        "См. п. 3.C. и рисунок 618. Нанесите кадмиевое покрытие по PCS-2100. Толщина кадмиевого покрытия должна быть в пределах 0,010-0,015 мм (0.0004 и 0.0006 in). Не наносить кадмиевое покрытие: на участки, где нанесено покрытие Sermetel W, на участки с хромовым покрытием и в местах, указанных на рисунке 618. Нанесите покрытие Sermetel W в указанных местах: см. IFC 40-860-03MD. Нанесите грунт, как показано, и финишное лакокрасочное покрытие по PCS-2500. Не наносить лакокрасочное покрытие: в местах, указанных на рисунке 618, на участки с хромовым покрытием и на резьбовые поверхности."
    ),
    "Refer to Figure 625. Apply cadmium plate all over but not to the chromium plated areas A and areas B and C: refer to PCS-2100. Make the cadmium plate thickness between 0,010 and 0,020 mm (0.0004 and 0.0008 in). Apply sermetel W only to the areas C: refer to IFC 40-860-03MD. Sermetel is optional in areas D. If sermetel is not applied in areas D, apply cadmium plate to areas D. Make the sermetel W coating thickness between 0,025 and 0,050 mm (0.001 and 0.002 in). The Sermetel W coating must overlap the chromium plated areas and cadmium plated areas. Apply primer paint only to the areas E. Apply paint all over but not to the chromium plated areas A, areas E and F: refer to PCS-2500. Do not paint areas F.": (
        "См. рисунок 625. Нанесите кадмиевое покрытие по всей поверхности, но не на участки с хромовым покрытием A и участки B и C: см. PCS-2100. Толщина кадмиевого покрытия должна быть в пределах 0,010-0,020 мм (0.0004 и 0.0008 in). Нанесите покрытие Sermetel W только на участки C: см. IFC 40-860-03MD. Покрытие Sermetel допускается на участках D. Если покрытие Sermetel не наносится на участках D, нанесите кадмиевое покрытие на участки D. Толщина покрытия Sermetel W должна быть в пределах 0,025-0,050 мм (0.001 и 0.002 in). Покрытие Sermetel W должно перекрывать участки с хромовым и кадмиевым покрытием. Нанесите грунтовочную краску только на участки E. Нанесите лакокрасочное покрытие по всей поверхности, но не на участки с хромовым покрытием A и участки E и F: см. PCS-2500. Не окрашивать участки F."
    ),
    "Refer to Figure 619. Apply primer paint only to the areas D: refer to Figure 619 and PCS-2500. Refer to PCS-2500 and apply paint all over externally, but not to: the bearings, bush bores and flanges the lubrication fittings and their identification washers the holes (with or without threads) the areas A and D. Refer to PCS-2500 and apply paint internally along surface B, but not along surface C.": (
        "См. рисунок 619. Нанесите грунтовочную краску только на участки D: см. рисунок 619 и PCS-2500. По PCS-2500 нанесите лакокрасочное покрытие по всей наружной поверхности, но не на: подшипники, отверстия под втулки и фланцы, смазочные штуцеры и их идентификационные шайбы, отверстия (с резьбой или без резьбы), участки A и D. По PCS-2500 нанесите лакокрасочное покрытие изнутри вдоль поверхности B, но не вдоль поверхности C."
    ),
    "Refer to Figures 620 and 621. Apply cadmium plate to M-DLPS131, do not apply cadmium plate to areas A. Primer paint all over but not in holes and on areas identified B and the 22,0 mm (0.87 in) diameters C: refer to PCS-2500.": (
        "См. рисунки 620 и 621. Нанесите кадмиевое покрытие по M-DLPS131, не наносите кадмиевое покрытие на участки A. Нанесите грунтовочную краску по всей поверхности, но не в отверстия, не на обозначенные участки B и не на диаметры C 22,0 мм (0.87 in): см. PCS-2500."
    ),
    "Refer to Figure 622. Apply cadmium plate all over but not to the chromium plated areas and areas A: refer to M-DLPS131. Apply brush cadmium plate to the areas D: refer to M-DLPS137. Paint all over externally and internally to areas B but not to the chromium plated areas, the bush and bearing bores, the chamfers, the lubrication fitting bores and areas C and D: refer to PCS-2500.": (
        "См. рисунок 622. Нанесите кадмиевое покрытие по всей поверхности, но не на участки с хромовым покрытием и участки A: см. M-DLPS131. Нанесите кистью кадмиевое покрытие на участки D: см. M-DLPS137. Нанесите лакокрасочное покрытие по всей поверхности снаружи и изнутри на участки B, но не на участки с хромовым покрытием, отверстия под втулки и подшипники, фаски, отверстия под смазочные штуцеры и участки C и D: см. PCS-2500."
    ),
    "Refer to Figure 636. Apply cadmium plate all over including the holes less than diameter 10 mm (0.393 in) but not to the chromium plated areas and areas A: refer to PCS-2100. Make the cadmium plate thickness between 0,010 and 0,020 mm (0.0004 and 0.0008 in). The cadmium plate must overlap the chromium plate run out. The cadmium plate is optional on the lubrication fitting bores where the lubrication adaptors (20-130), (20-160), (20-190) and (20-220) will install. Apply primer paint only to areas B: refer to PCS-2500. Apply wet primer to PCS-2804 or apply resin to PCS-2802 to the area D. Apply paint all over but not on the chromium plated areas, the areas A, B, C and on the lubrication fitting bores where the lubrication adaptors (20-130), (20-160), (20-190) and (20-220) will install: refer to PCS-2500. Paint finish is optional in areas E.": (
        "См. рисунок 636. Нанесите кадмиевое покрытие по всей поверхности, включая отверстия диаметром менее 10 мм (0.393 in), но не на участки с хромовым покрытием и участки A: см. PCS-2100. Толщина кадмиевого покрытия должна быть в пределах 0,010-0,020 мм (0.0004 и 0.0008 in). Кадмиевое покрытие должно перекрывать зону схода хромового покрытия. Кадмиевое покрытие допускается в отверстиях под смазочные штуцеры, где будут устанавливаться смазочные адаптеры (20-130), (20-160), (20-190) и (20-220). Нанесите грунтовочную краску только на участки B: см. PCS-2500. Нанесите жидкий грунт по PCS-2804 или смолу по PCS-2802 на участок D. Нанесите лакокрасочное покрытие по всей поверхности, но не на участки с хромовым покрытием, участки A, B, C и отверстия под смазочные штуцеры, где будут устанавливаться смазочные адаптеры (20-130), (20-160), (20-190) и (20-220): см. PCS-2500. Лакокрасочное покрытие допускается на участках E."
    ),
    "Refer to Figure 636. Apply cadmium plate all over including the holes less than diameter 10 mm (0.393 in) but not to the chromium plated areas and areas A: refer to PCS-2100. Make the cadmium plate thickness between 0,010 and 0,020 mm (0.0004 and 0.0008 in). The cadmium plate must overlap the chromium plate run out. The cadmium plate is optional on the lubrication fitting bores where the lubrication adaptors (20-130), (20-160), (20-190) and (20-220) will install. Apply primer paint only to areas B: refer to PCS-2500. Apply wet primer to PCS-2804 or apply resin to PCS-2802 to the area D. Apply paint all over but not on the chromium plated areas, the areas A, B, C and on the lubrication fitting bores where the lubrication adaptors (20-130), (20-160), (20- 190) and (20-220) will install: refer to PCS-2500.": (
        "См. рисунок 636. Нанесите кадмиевое покрытие по всей поверхности, включая отверстия диаметром менее 10 мм (0.393 in), но не на участки с хромовым покрытием и участки A: см. PCS-2100. Толщина кадмиевого покрытия должна быть в пределах 0,010-0,020 мм (0.0004 и 0.0008 in). Кадмиевое покрытие должно перекрывать зону схода хромового покрытия. Кадмиевое покрытие допускается в отверстиях под смазочные штуцеры, где будут устанавливаться смазочные адаптеры (20-130), (20-160), (20-190) и (20-220). Нанесите грунтовочную краску только на участки B: см. PCS-2500. Нанесите жидкий грунт по PCS-2804 или смолу по PCS-2802 на участок D. Нанесите лакокрасочное покрытие по всей поверхности, но не на участки с хромовым покрытием, участки A, B, C и отверстия под смазочные штуцеры, где будут устанавливаться смазочные адаптеры (20-130), (20-160), (20-190) и (20-220): см. PCS-2500."
    ),
    "Refer to Figure 636. Apply cadmium plate all over including the holes less than diameter 10 mm (0.393 in) but not to the chromium plated areas and areas A: refer to PCS-2100. Make the cadmium plate thickness between 0,010 and 0,020 mm (0.0004 and 0.0008 in). The cadmium plate must overlap the chromium plate run out. The cadmium plate is optional on the lubrication fitting bores where the lubrication adaptors (20-130), (20-160), (20-190) and (20-220) will install. Apply primer paint only to areas B: refer to PCS-2500. Apply wet primer to PCS-2804 or apply resin to PCS-2802 to the area D. Apply paint all over but not on the chromium plated areas, the areas A, B, C and on the lubrication fitting bores where the lubrication adaptors (20-130), (20-160), (20- 190) and (20-220) will install: refer to PCS-2500. Paint finish is optional in areas E.": (
        "См. рисунок 636. Нанесите кадмиевое покрытие по всей поверхности, включая отверстия диаметром менее 10 мм (0.393 in), но не на участки с хромовым покрытием и участки A: см. PCS-2100. Толщина кадмиевого покрытия должна быть в пределах 0,010-0,020 мм (0.0004 и 0.0008 in). Кадмиевое покрытие должно перекрывать зону схода хромового покрытия. Кадмиевое покрытие допускается в отверстиях под смазочные штуцеры, где будут устанавливаться смазочные адаптеры (20-130), (20-160), (20-190) и (20-220). Нанесите грунтовочную краску только на участки B: см. PCS-2500. Нанесите жидкий грунт по PCS-2804 или смолу по PCS-2802 на участок D. Нанесите лакокрасочное покрытие по всей поверхности, но не на участки с хромовым покрытием, участки A, B, C и отверстия под смазочные штуцеры, где будут устанавливаться смазочные адаптеры (20-130), (20-160), (20-190) и (20-220): см. PCS-2500. Лакокрасочное покрытие допускается на участках E."
    ),
    "M-DLPS100-2. The cadmium plate must be 0,010 to 0,015 mm (0.0004 to 0.0006 in) thick. Paint: refer to": (
        "M-DLPS100-2. Кадмиевое покрытие должно иметь толщину 0,010-0,015 мм (0.0004-0.0006 in). Окраска: см."
    ),
    "not paint areas A and B. Apply primer to area A: refer to Figure 604.": (
        "Не окрашивать участки A и B. Нанесите грунт на участок A: см. рисунок 604."
    ),
    "Refer to Figure 634. Chromic acid anodise all over but not the spotface A: refer to M-DLPS102-1. Apply Alocrom 1200 to the areas A: refer to M-DLPS114. Apply paint all over but not to the areas A, B, C and D: refer to PCS-2500. Apply a light coat of primer to the hole B: refer to": (
        "См. рисунок 634. Анодируйте в хромовой кислоте по всей поверхности, кроме подрезки площадки A: см. M-DLPS102-1. Нанесите Alocrom 1200 на участки A: см. M-DLPS114. Нанесите лакокрасочное покрытие по всей поверхности, но не на участки A, B, C и D: см. PCS-2500. Нанесите тонкий слой грунта в отверстие B: см."
    ),
    "PCS-2500. Apply only primer to face D: refer to PCS-2500.": (
        "PCS-2500. Нанесите грунт только на поверхность D: см. PCS-2500."
    ),
    "Refer to Figure 613. Apply cadmium plate: refer to M-DLPS100-2. Do not include areas that have chromium plate. Paint area A: refer to": (
        "См. рисунок 613. Нанесите кадмиевое покрытие: см. M-DLPS100-2. Не включать участки с хромовым покрытием. Нанесите лакокрасочное покрытие на участок A: см."
    ),
    "not paint the screw threads and the face that touches the wheel bearings.": (
        "Не окрашивать резьбу и поверхность, контактирующую с колесными подшипниками."
    ),
    "Refer to Figure 630. Before installation of bushes: Apply primer paint to areas A but not to the areas B: refer to PCS-2500.": (
        "См. рисунок 630. Перед установкой втулок: нанесите грунтовочную краску на участки A, но не на участки B: см. PCS-2500."
    ),
    "Refer to Figure 623. Apply cadmium": "См. рисунок 623. Нанесите кадмиевое",
    "plate internally and externally over area A: refer to PCS-2101. Make the cadmium plate thickness between 0,010 and 0,015 mm (0.0004 and": (
        "покрытие изнутри и снаружи на участок A: см. PCS-2101. Толщина кадмиевого покрытия должна быть в пределах 0,010-0,015 мм (0.0004 и"
    ),
    "Refer to Figure 617. Cadmium plate all over to M-DLPS100-2. Make cadmium plate thickness 0,010 to 0,015 mm (0.0004 to 0.0006 in). Do": (
        "См. рисунок 617. Нанесите кадмиевое покрытие по всей поверхности по M-DLPS100-2. Толщина кадмиевого покрытия должна составлять 0,010-0,015 мм (0.0004-0.0006 in). Не"
    ),
    "Paint areas A: refer to PCS-2500. Primer paint only on face D and areas E and F including the chamfer. Do not paint:": (
        "Нанесите лакокрасочное покрытие на участки A: см. PCS-2500. Нанесите грунтовочную краску только на поверхность D и участки E и F, включая фаску. Не окрашивать:"
    ),
    "Refer to Figure 624. Apply cadmium plate all over but not to the areas A: refer to PCS-2101. Cadmium plate is optional on radii and chamfer B. Make the cadmium plate thickness between 0,010 and 0,015 mm": (
        "См. рисунок 624. Нанесите кадмиевое покрытие по всей поверхности, но не на участки A: см. PCS-2101. Кадмиевое покрытие допускается на радиусные участки и фаску B. Толщина кадмиевого покрытия должна быть в пределах 0,010-0,015 мм"
    ),
    "Apply only primer paint to areas C including chamfer: refer to": (
        "Нанесите только грунтовочную краску на участки C, включая фаску: см."
    ),
    "to the chromium plated areas": "на участки с хромовым покрытием",
    "to the chromium plated areas.": "на участки с хромовым покрытием.",
    "the areas A and D.": "участки A и D.",
    "Refer to PCS-2500 and apply paint internally along surface B, but not along surface C.": (
        "По PCS-2500 нанесите лакокрасочное покрытие изнутри вдоль поверхности B, но не вдоль поверхности C."
    ),
    "M-DLPS137. Paint all over externally and internally to areas B but not to the chromium plated areas, the bush and bearing bores, the chamfers, the lubrication fitting bores and areas C and D: refer to PCS-2500.": (
        "M-DLPS137. Нанесите лакокрасочное покрытие по всей поверхности снаружи и изнутри на участки B, но не на участки с хромовым покрытием, отверстия под втулки и подшипники, фаски, отверстия под смазочные штуцеры и участки C и D: см. PCS-2500."
    ),
    "and 0,015 mm (0.0004": "и 0,015 мм (0.0004",
    "and 0.0005 in).": "и 0.0005 in).",
    "Areas A and B.": "Участки A и B.",
}


PART5_EXACT_MAP: dict[str, str] = {
    "Specified Damage and Material Specification.": "Указанное повреждение и спецификация материала.",
    "Specified Damage and Material Specification": "Указанное повреждение и спецификация материала",
    "Specified Damage": "Указанное повреждение",
    "Damaged or loose liner.": "Поврежденный или ослабленный вкладыш.",
    "Loose or damaged liner.": "Ослабленный или поврежденный вкладыш.",
    "Special tools are not necessary.": "Специальные инструменты не требуются.",
    "Materials are not necessary.": "Материалы не требуются.",
    "Repair parts are not necessary.": "Ремонтные детали не требуются.",
    "These repair parts are necessary:": "Необходимы следующие ремонтные детали:",
    "Repair loose but undamaged liner:": "Ремонт ослабленного, но неповрежденного вкладыша:",
    "Repair damaged liner:": "Ремонт поврежденного вкладыша:",
    "Remove all of the adhesive from the external diameter of the gland housing and the internal diameter of the liner: refer to CLEANING.": (
        "Удалите весь клей с наружного диаметра корпуса сальника и внутреннего диаметра вкладыша: см. ОЧИСТКА."
    ),
    "Measure the diameter of the gland housing across the contact area: this must be between 211,328 and 211,368 mm (8.3200 and 8.3215 in).": (
        "Измерьте диаметр корпуса сальника по контактному участку: он должен быть в пределах 211,328-211,368 mm (8.3200-8.3215 in)."
    ),
    "Temporarily put the liner in position on the gland housing. If necessary, cut the liner at the scarf joints to adjust its length.": (
        "Временно установите вкладыш на корпус сальника. При необходимости подрежьте вкладыш по косым стыкам, чтобы подогнать его по длине."
    ),
    "Clean the contact surfaces of the liner and the gland housing: use the cleaning tissues, Material Ref. Item TBA, and the cleaning agent, Material Ref. Item 11-583.": (
        "Очистите контактные поверхности вкладыша и корпуса сальника: используйте салфетки для очистки, код ссылки материала TBA, и очиститель, код ссылки материала 11-583."
    ),
    "Apply adhesive, Material Ref. Item 08-722, to the contact surface of the liner near to the scarf joints: refer to M-DLPS724.": (
        "Нанесите клей, код ссылки материала 08-722, на контактную поверхность вкладыша возле косых стыков: см. M-DLPS724."
    ),
    "Assemble the liner to the gland housing.": "Установите вкладыш на корпус сальника.",
    "Examine the part to make sure that you have obeyed all the repair instructions correctly.": (
        "Проверьте деталь и убедитесь, что все указания по ремонту выполнены правильно."
    ),
    "Examine the part to make sure that you have obeyed the repair instructions correctly.": (
        "Проверьте деталь и убедитесь, что указания по ремонту выполнены правильно."
    ),
    "Remove the damaged liner and remove all of the adhesive from the external diameter of the gland housing: refer to CLEANING.": (
        "Снимите поврежденный вкладыш и удалите весь клей с наружного диаметра корпуса сальника: см. ОЧИСТКА."
    ),
    "Machine the gland housing to the dimensions given in Figure 601 with a surface finish of 3,2 micrometers (125 micro-inches).": (
        "Обработайте корпус сальника до размеров, указанных на рисунке 601, с шероховатостью поверхности 3,2 micrometers (125 micro-inches)."
    ),
    "Examine the gland housing for flaws: refer to M-DLNDT8.": "Осмотрите корпус сальника на наличие дефектов: см. M-DLNDT8.",
    "Apply Alocrom, Material Ref. Item 13-501, to the machined areas: refer to PCS-2220.": (
        "Нанесите Alocrom, код ссылки материала 13-501, на обработанные участки: см. PCS-2220."
    ),
    "Temporarily put the repair liner in position on the gland housing. If necessary, cut the repair liner at the scarf joints to adjust its length.": (
        "Временно установите ремонтный вкладыш на корпус сальника. При необходимости подрежьте ремонтный вкладыш по косым стыкам, чтобы подогнать его по длине."
    ),
    "Apply adhesive PVC tape, Material Ref. Item TBA, around the gland housing to the sides of and touching the repair liner. Make sure that the edges of the adhesive PVC tape, Material Ref. Item TBA, bond tightly to the gland housing.": (
        "Наклейте клейкую ленту ПВХ, код ссылки материала TBA, вокруг корпуса сальника по обе стороны от ремонтного вкладыша вплотную к нему. Убедитесь, что края клейкой ленты ПВХ, код ссылки материала TBA, плотно приклеены к корпусу сальника."
    ),
    "Use the Emery cloth, 60-100 grit, Material Ref. Item TBA, to roughen the surfaces to be bonded. Do not damage the edges of the adhesive PVC tape, Material Ref. Item TBA.": (
        "Используйте шлифовальную шкурку зернистостью 60-100, код ссылки материала TBA, чтобы зашероховать склеиваемые поверхности. Не повредите края клейкой ленты ПВХ, код ссылки материала TBA."
    ),
    "Use the cleaning agent, Material Ref. Item 11-583, and cleaning tissues, Material Ref. Item TBA, to clean the roughened surfaces.": (
        "Используйте очиститель, код ссылки материала 11-583, и салфетки для очистки, код ссылки материала TBA, чтобы очистить зашерохованные поверхности."
    ),
    "Preheat an oven to between 35 and 45 oC (95 and 113 oF).": "Предварительно нагрейте печь до 35-45 oC (95-113 oF).",
    "Prepare a surface treatment mixture of 1 part by volume of Accomet C, Material Ref. Item TBA, and 4 parts by volume of clean cold water.": (
        "Приготовьте смесь для обработки поверхности: 1 часть по объему Accomet C, код ссылки материала TBA, и 4 части по объему чистой холодной воды."
    ),
    "Use a brush to apply a smooth layer of the prepared surface treatment mixture to the contact surface of the gland housing.": (
        "Кистью нанесите ровный слой приготовленной смеси для обработки поверхности на контактную поверхность корпуса сальника."
    ),
    "Put the gland housing in the preheated oven for a minimum of 4 minutes and until the applied surface treatment mixture is dry.": (
        "Поместите корпус сальника в предварительно нагретую печь минимум на 4 минуты, пока нанесенная смесь для обработки поверхности не высохнет."
    ),
    "Alternative procedure for paragraphs (j) to (m): apply Alocrom to the contact surface of the gland housing: refer to PCS-2220, type 2.": (
        "Альтернативная процедура для пунктов (j)-(m): нанесите Alocrom на контактную поверхность корпуса сальника: см. PCS-2220, тип 2."
    ),
    "Use a brush to apply Araldite, 2015, Material Ref. Item TBA, to the gland housing.": (
        "Кистью нанесите Araldite 2015, код ссылки материала TBA, на корпус сальника."
    ),
    "Assemble the repair liner to the gland housing and use masking tape, Material Ref. Item 08-715, to hold it in that position. Use one layer of masking tape, Material Ref. Item 08-715, at each side of the repair liner. The masking tape must be sufficiently wide to bond to the repair liner and the adhesive PVC tape, Material Ref. Item TBA: make sure the ends touch but do not overlap.": (
        "Установите ремонтный вкладыш на корпус сальника и зафиксируйте его в этом положении малярной лентой, код ссылки материала 08-715. Используйте по одному слою малярной ленты, код ссылки материала 08-715, с каждой стороны ремонтного вкладыша. Малярная лента должна быть достаточно широкой, чтобы приклеиться к ремонтному вкладышу и клейкой ленте ПВХ, код ссылки материала TBA; убедитесь, что концы ленты соприкасаются, но не перекрываются."
    ),
    "Clamp the repair liner to the gland housing using an applicable tool.": (
        "Прижмите ремонтный вкладыш к корпусу сальника подходящим приспособлением."
    ),
    "Put the gland housing in the preheated oven, kept at between 35 and 45 °C (95 and 113 °F), for 345 to 375 minutes.": (
        "Поместите корпус сальника в предварительно нагретую печь, поддерживаемую при 35-45 °C (95-113 °F), на 345-375 минут."
    ),
    "Remove the gland housing from the oven and allow to cool for a minimum of 30 minutes.": (
        "Выньте корпус сальника из печи и дайте ему остыть не менее 30 минут."
    ),
    "Machine the diameter and width of the repair liner to the dimensions given in Figure 601 then remove the adhesive PVC tape, Material Ref. Item TBA.": (
        "Обработайте диаметр и ширину ремонтного вкладыша до размеров, указанных на рисунке 601, затем удалите клейкую ленту ПВХ, код ссылки материала TBA."
    ),
    "Remove all of the tape and clean the parts as necessary.": "Удалите всю ленту и при необходимости очистите детали.",
    "Remove the used adhesive from the external diameter of the gland housing and the internal diameter of the liner: refer to CLEANING.": (
        "Удалите использованный клей с наружного диаметра корпуса сальника и внутреннего диаметра вкладыша: см. ОЧИСТКА."
    ),
    "Put the liner in position on the gland housing. If necessary, cut the liner at the scarf joints to adjust its length.": (
        "Установите вкладыш на корпус сальника. При необходимости подрежьте вкладыш по косым стыкам, чтобы подогнать его по длине."
    ),
    "Clean the surfaces of the liner and the gland housing that will touch: use the cleaning tissues, Material Ref. Item TBA and the cleaning agent, Material Ref. Item 11-583.": (
        "Очистите соприкасающиеся поверхности вкладыша и корпуса сальника: используйте салфетки для очистки, код ссылки материала TBA, и очиститель, код ссылки материала 11-583."
    ),
    "Apply adhesive, Material Ref. Item 08-722, to the cleaned surface of the liner near to the scarf joints: refer to M-DLPS724.": (
        "Нанесите клей, код ссылки материала 08-722, на очищенную поверхность вкладыша возле косых стыков: см. M-DLPS724."
    ),
    "Remove the damaged liner and remove the used adhesive from the external diameter of the gland housing: refer to CLEANING.": (
        "Снимите поврежденный вкладыш и удалите использованный клей с наружного диаметра корпуса сальника: см. ОЧИСТКА."
    ),
    "Put the repair liner in position on the gland housing. If necessary, cut the repair liner at the scarf joints to adjust its length.": (
        "Установите ремонтный вкладыш на корпус сальника. При необходимости подрежьте ремонтный вкладыш по косым стыкам, чтобы подогнать его по длине."
    ),
    "Use the Emery cloth, 60-100 grit, Material Ref. Item TBA, to make rough, the surfaces that will bond. Do not damage the edges of the adhesive PVC tape, Material Ref. Item TBA.": (
        "Используйте шлифовальную шкурку зернистостью 60-100, код ссылки материала TBA, чтобы зашероховать склеиваемые поверхности. Не повредите края клейкой ленты ПВХ, код ссылки материала TBA."
    ),
    "Use the cleaning agent, Material Ref. Item 11-583, and cleaning tissues, Material Ref. Item TBA, to clean the surfaces made rough.": (
        "Используйте очиститель, код ссылки материала 11-583, и салфетки для очистки, код ссылки материала TBA, чтобы очистить зашерохованные поверхности."
    ),
    "Set the temperature of an oven to between 35 and 45 oC (95 and 113 oF).": (
        "Установите температуру печи в пределах 35-45 oC (95-113 oF)."
    ),
    "Prepare a mixture of 1 part by volume of Accomet C, Material Ref. Item TBA, and 4 parts by volume of clean cold water.": (
        "Приготовьте смесь из 1 части по объему Accomet C, код ссылки материала TBA, и 4 частей по объему чистой холодной воды."
    ),
    "Use a brush to apply a flat layer of the prepared mixture to the surfaces of the gland housing made rough.": (
        "Кистью нанесите ровный слой приготовленной смеси на зашерохованные поверхности корпуса сальника."
    ),
    "Put the gland housing in the oven for a minimum of 4 minutes and until the applied mixture is dry.": (
        "Поместите корпус сальника в печь минимум на 4 минуты, пока нанесенная смесь не высохнет."
    ),
    "Alternative procedure for paragraphs (g) to (j): apply Alocrom to the contact surface of the gland housing: refer to PCS-2220, type 2.": (
        "Альтернативная процедура для пунктов (g)-(j): нанесите Alocrom на контактную поверхность корпуса сальника: см. PCS-2220, тип 2."
    ),
    "Assemble the repair liner to the gland housing and hold it tightly with the masking tape, Material Ref. Item 08-715. Bond the masking tape, Material Ref. Item 08-715, to the adjacent adhesive PVC tape, Material Ref. Item TBA. Use one layer of the masking tape, Material Ref. Item 08-715, only at each side of the repair liner: make sure that the ends touch but do not overlap.": (
        "Установите ремонтный вкладыш на корпус сальника и плотно зафиксируйте его малярной лентой, код ссылки материала 08-715. Приклейте малярную ленту, код ссылки материала 08-715, к соседней клейкой ленте ПВХ, код ссылки материала TBA. Используйте только по одному слою малярной ленты, код ссылки материала 08-715, с каждой стороны ремонтного вкладыша; убедитесь, что концы ленты соприкасаются, но не перекрываются."
    ),
    "Use a clamp to attach the repair liner to the gland housing.": (
        "Используйте зажим, чтобы прижать ремонтный вкладыш к корпусу сальника."
    ),
    "Put the gland housing in the oven, kept at between 35 and 45 oC (95 and 113 oF), for 345 to 375 minutes.": (
        "Поместите корпус сальника в печь, поддерживаемую при 35-45 oC (95-113 oF), на 345-375 минут."
    ),
    "Remove the gland housing from the oven and let its temperature decrease for a minimum of 30 minutes.": (
        "Выньте корпус сальника из печи и дайте ему охлаждаться не менее 30 минут."
    ),
    "Machine the diameter and width of the repair liner to the dimensions shown before you remove the adhesive PVC tape, Material Ref. Item TBA.": (
        "Обработайте диаметр и ширину ремонтного вкладыша до указанных размеров, затем удалите клейкую ленту ПВХ, код ссылки материала TBA."
    ),
    "Damage or wear to the diameters A and B.": "Повреждение или износ диаметров A и B.",
    "Remove the chromium plate from diameters A and B.": "Удалите хромовое покрытие с диаметров A и B.",
    "Machine the diameters A and B to remove damage or wear after removal of the chromium plate. Remove the minimum amount of material necessary, to the dimensions shown in Figure 601, to remove the damage or wear.": (
        "Обработайте диаметры A и B после удаления хромового покрытия, чтобы устранить повреждение или износ. Удаляйте минимально необходимое количество материала, доводя до размеров, указанных на рисунке 601."
    ),
    "Shot peen the machined areas: refer to M-DLPS123.": "Выполните дробеструйное упрочнение обработанных участков: см. M-DLPS123.",
    "Apply chromium plate to the diameters A and B: refer to M-DLPS101-2. Make sure the chromium plate terminations are smooth: refer to M-DLPS1031.": (
        "Нанесите хромовое покрытие на диаметры A и B: см. M-DLPS101-2. Убедитесь, что границы хромового покрытия выполнены плавно: см. M-DLPS1031."
    ),
    "Finish grind the pin to the dimensions shown in Figure 601.": (
        "Окончательно отшлифуйте штифт до размеров, указанных на рисунке 601."
    ),
    "Damage or wear to the diameter A.": "Повреждение или износ диаметра A.",
    "Remove the chromium plate only from the diameter A to show the parent metal.": (
        "Удалите хромовое покрытие только с диаметра A, чтобы вскрыть основной металл."
    ),
    "If the parent metal is not damaged or worn:": "Если основной металл не поврежден и не изношен:",
    "Apply chromium plate to the diameter A: refer to PCS-2110 Type B. Refer to Figure 601 for chromium plate termination information.": (
        "Нанесите хромовое покрытие на диаметр A: см. PCS-2110, тип B. Информацию о границе хромового покрытия см. на рисунке 601."
    ),
    "Grind the diameter A to between 15,960 and 15,987 mm (0.6284 and 0.6294 in). The surface finish must be 0,8 micrometers (32 micro-inches).": (
        "Отшлифуйте диаметр A до 15,960-15,987 mm (0.6284-0.6294 in). Шероховатость поверхности должна быть 0,8 micrometers (32 micro-inches)."
    ),
    "If the parent metal is damaged or worn:": "Если основной металл поврежден или изношен:",
    "Machine the diameter A sufficiently to remove damage and wear, down to a minimum of 15,377 mm (0.6054 in).": (
        "Обработайте диаметр A настолько, чтобы удалить повреждение и износ, но не менее чем до 15,377 mm (0.6054 in)."
    ),
    "Remove the chromium plate from diameter A.": "Удалите хромовое покрытие с диаметра A.",
    "Machine diameter A to remove the damage or wear after removal of the chromium plate. The minimum diameter is 24,36 mm (0.960 in) and the surface finish must be 0,4 micrometers (16 micro-inches).": (
        "Обработайте диаметр A после удаления хромового покрытия, чтобы устранить повреждение или износ. Минимальный диаметр должен составлять 24,36 mm (0.960 in), а шероховатость поверхности - 0,4 micrometers (16 micro-inches)."
    ),
    "If diameter A has been machined (see para (2)), shot peen the machined area: refer to M-DLPS123.": (
        "Если диаметр A обрабатывался (см. пункт (2)), выполните дробеструйное упрочнение обработанного участка: см. M-DLPS123."
    ),
    "Apply chromium plate to diameter A as shown with a minimum chromium plate thickness of 0,075 mm (0.0030 in): refer to M-DLPS101-2.": (
        "Нанесите хромовое покрытие на диаметр A, как показано, с минимальной толщиной 0,075 mm (0.0030 in): см. M-DLPS101-2."
    ),
    "Machine diameter A to between 24,947 and 24,980 mm (0.9822 and 0.9834 in) with a surface finish of 0,4 micrometers (16 micro-inches).": (
        "Обработайте диаметр A до 24,947-24,980 mm (0.9822-0.9834 in) с шероховатостью поверхности 0,4 micrometers (16 micro-inches)."
    ),
    "Examine the chromium plated surface for flaws: refer to M-DLNDT3.": (
        "Осмотрите поверхность с хромовым покрытием на наличие дефектов: см. M-DLNDT3."
    ),
    "Apply cadmium plate to the machined areas but not where chromium plated: refer to PCS-2141.": (
        "Нанесите кадмиевое покрытие на обработанные участки, но не на участки с хромовым покрытием: см. PCS-2141."
    ),
    "Spray primer paint lightly on the cadmium plated surface: refer to PCS-2500.": (
        "Слегка распылите грунтовочную краску на кадмированную поверхность: см. PCS-2500."
    ),
    "Refer to PCS-6000-07 and identify the part with the applicable Messier-Dowty Limited repair number, adjacent to the part number, after painting:": (
        "См. PCS-6000-07 и после окраски нанесите рядом с номером детали применимый номер ремонта Messier-Dowty Limited:"
    ),
    "450237490A if only the chromium plate has been repaired, or": (
        "450237490A, если был отремонтирован только слой хромового покрытия, или"
    ),
    "450237490B if chromium plate and parent metal have been repaired.": (
        "450237490B, если были отремонтированы хромовое покрытие и основной металл."
    ),
    "Machine diameter A sufficiently to remove the damage or wear and to the dimensions shown in Figure 601.": (
        "Обработайте диаметр A настолько, чтобы устранить повреждение или износ и получить размеры, указанные на рисунке 601."
    ),
    "Shot peen the pin: refer to M-DLPS123.": "Выполните дробеструйное упрочнение штифта: см. M-DLPS123.",
    "Apply chromium plate to diameter A to give a minimum diameter of 20,120 mm (0.7921 in): refer to M-DLPS101-2P.": (
        "Нанесите хромовое покрытие на диаметр A до достижения минимального диаметра 20,120 mm (0.7921 in): см. M-DLPS101-2P."
    ),
    "Grind diameter A to the dimensions shown with a surface finish of 0,4 micrometers (16 micro-inches). Refer to M-DLPS1031-3 and M-DLPS1031-6 for the chromium plate terminations where shown.": (
        "Отшлифуйте диаметр A до указанных размеров с шероховатостью поверхности 0,4 micrometers (16 micro-inches). В местах, где это показано, границы хромового покрытия выполняйте по M-DLPS1031-3 и M-DLPS1031-6."
    ),
    "Paint the pin all over, but not on the threads, the thread undercut and the chromium plated areas: refer to PCS-2500.": (
        "Окрасьте штифт по всей поверхности, но не по резьбе, проточке под резьбой и участкам с хромовым покрытием: см. PCS-2500."
    ),
    "Damage or corrosion to diameter A.": "Повреждение или коррозия диаметра A.",
    "Remove the chromium plate from diameter A: refer to PCS-2110.": "Удалите хромовое покрытие с диаметра A: см. PCS-2110.",
    "If the bare metal is not damaged or corroded:": "Если оголенный металл не поврежден и не корродирован:",
    "Shot peen diameter A: refer to M-DLPS123.": "Выполните дробеструйное упрочнение диаметра A: см. M-DLPS123.",
    "Apply chromium plate to diameter A: refer to PCS-2110, Type C. Refer to Figure 601 for chromium plate termination information.": (
        "Нанесите хромовое покрытие на диаметр A: см. PCS-2110, тип C. Информацию о границе хромового покрытия см. на рисунке 601."
    ),
    "Grind diameter A to between 59,951 and 59,970 mm (2.3603 and 2.3610 in). The surface finish must be 0,8 micrometers (32 micro-inches).": (
        "Отшлифуйте диаметр A до 59,951-59,970 mm (2.3603-2.3610 in). Шероховатость поверхности должна быть 0,8 micrometers (32 micro-inches)."
    ),
    "If the bare metal is damaged or corroded:": "Если оголенный металл поврежден или корродирован:",
    "Machine diameter A just sufficiently to remove the damage or corrosion: refer to M-DLPS1004-4-1. The diameter must not be less than 59,370 mm": (
        "Обработайте диаметр A ровно настолько, чтобы удалить повреждение или коррозию: см. M-DLPS1004-4-1. Диаметр не должен быть менее 59,370 mm"
    ),
    "(2.3374 in). The surface finish must be 1,6 micrometers (63 micro-inches).": (
        "(2.3374 in). Шероховатость поверхности должна быть 1,6 micrometers (63 micro-inches)."
    ),
    "Shot peen and apply chromium plate to diameter A: refer to M-DLPS123 and PCS-2110, Type C. Refer to Figure 601 for chromium plate termination information.": (
        "Выполните дробеструйное упрочнение и нанесите хромовое покрытие на диаметр A: см. M-DLPS123 и PCS-2110, тип C. Информацию о границе хромового покрытия см. на рисунке 601."
    ),
    "Machine diameter A sufficiently to remove the damage or wear and to between 15,90 and 16,50 mm (0.626 and 0.649 in): refer to M-DLPS1004-4-1. Machine a surface finish of 1,6 micrometers (63 micro-inches).": (
        "Обработайте диаметр A настолько, чтобы удалить повреждение или износ, и доведите его до 15,90-16,50 mm (0.626-0.649 in): см. M-DLPS1004-4-1. Шероховатость поверхности должна быть 1,6 micrometers (63 micro-inches)."
    ),
    "Measure the new diameter A.": "Измерьте новый диаметр A.",
    "Calculate diameter C of the repair bushes:": "Рассчитайте диаметр C ремонтных втулок:",
    "C = A (as measured) + 0,010 to 0,039 mm (0.0004 to 0.0015 in).": (
        "C = A (по результатам измерения) + 0,010-0,039 mm (0.0004-0.0015 in)."
    ),
    "Machine the repair bushes to the dimensions shown and calculated. Machine a surface finish of 1,6 micrometers (63 micro-inches).": (
        "Обработайте ремонтные втулки до указанных и рассчитанных размеров. Шероховатость поверхности должна быть 1,6 micrometers (63 micro-inches)."
    ),
    "Apply adhesive, Material Ref. Item 08-665, to the outside diameter of the repair bushes: refer to M-DLPS709-6.": (
        "Нанесите клей, код ссылки материала 08-665, на наружный диаметр ремонтных втулок: см. M-DLPS709-6."
    ),
    "Install the repair bushes to diameter A until aligned with the outside diameter of the pin: refer to M-DLPS1011-5.": (
        "Установите ремонтные втулки в диаметр A до совмещения с наружным диаметром штифта: см. M-DLPS1011-5."
    ),
    "Machine the repair bushes to the contour of the pin and the dimensions shown.": (
        "Обработайте ремонтные втулки по контуру штифта и до указанных размеров."
    ),
    "Damage or wear to the diameter A and/or the adjacent inside face.": (
        "Повреждение или износ диаметра A и/или прилегающего внутреннего торца."
    ),
    "Machine the diameter A to remove the minimum amount of material necessary to remove the damage or wear: refer to Figure 601 for the dimensions. The surface finish must be 1,6 micrometers (63 micro-inches).": (
        "Обработайте диаметр A, снимая минимально необходимое количество материала для устранения повреждения или износа: размеры см. на рисунке 601. Шероховатость поверхности должна быть 1,6 micrometers (63 micro-inches)."
    ),
    "Machine the face Q as necessary to remove damage or wear: refer to Figure 601 for the dimensions. The surface finish must be 1,6 micrometers (63 micro-inches).": (
        "При необходимости обработайте торец Q, чтобы устранить повреждение или износ: размеры см. на рисунке 601. Шероховатость поверхности должна быть 1,6 micrometers (63 micro-inches)."
    ),
    "Machine the face Q as necessary to remove the damage or wear: refer to Figure 601 for the dimensions. The surface finish must be 1,6 micrometers (63 micro-inches).": (
        "При необходимости обработайте торец Q, чтобы устранить повреждение или износ: размеры см. на рисунке 601. Шероховатость поверхности должна быть 1,6 micrometers (63 micro-inches)."
    ),
    "Measure and make a record of the new diameter A and the thickness of the lug D.": (
        "Измерьте и зафиксируйте новый диаметр A и толщину проушины D."
    ),
    "Prepare the machined surfaces of the bracket:": "Подготовьте обработанные поверхности кронштейна:",
    "Calculate the diameter B and the dimension C of the oversize bearing:": (
        "Рассчитайте диаметр B и размер C ремонтного подшипника увеличенного размера:"
    ),
    "B = A (as measured) - 0,006 to + 0,023 mm (- 0.0002 to + 0.0009 in)": (
        "B = A (по результатам измерения) - 0,006 до +0,023 mm (-0.0002 до +0.0009 in)"
    ),
    "C = D (as measured) - 0,15 to - 0,25 mm (- 0.006 to - 0.010 in).": (
        "C = D (по результатам измерения) - 0,15 до -0,25 mm (-0.006 до -0.010 in)."
    ),
    "C = D (as measured) - 0,30 to - 0,40 mm (- 0.011 to - 0.015 in).": (
        "C = D (по результатам измерения) - 0,30 до -0,40 mm (-0.011 до -0.015 in)."
    ),
    "Machine the oversize bearing to the dimensions shown and calculated: the surface finish must be 1,6 micrometers (63 micro-inches).": (
        "Обработайте ремонтный подшипник увеличенного размера до указанных и рассчитанных размеров; шероховатость поверхности должна быть 1,6 micrometers (63 micro-inches)."
    ),
    "Passivate the oversize bearing: refer to M-DLPS124.": (
        "Пассивируйте ремонтный подшипник увеличенного размера: см. M-DLPS124."
    ),
    "Apply cadmium plate to the oversize bearing, but not on the internal diameter. The thickness of the cadmium plate must be between 0,010 and 0,015 mm (0.0004 and 0.0006 in): refer to M-DLPS100-2S.": (
        "Нанесите кадмиевое покрытие на ремонтный подшипник увеличенного размера, но не на внутренний диаметр. Толщина кадмиевого покрытия должна быть 0,010-0,015 mm (0.0004-0.0006 in): см. M-DLPS100-2S."
    ),
    "Use the Press Pad 460004330/85 and Drift 460004331/7 to install the oversize bearing to the bracket: refer to M-DLPS1011-14. Check line ream the oversize bearing to the dimension shown in Figure 601.": (
        "Используйте нажимную опору 460004330/85 и выколотку 460004331/7, чтобы установить ремонтный подшипник увеличенного размера в кронштейн: см. M-DLPS1011-14. Затем выполните контрольную линейную развертку ремонтного подшипника до размера, указанного на рисунке 601."
    ),
    "Use the Press Pad 460004330/136 and Drift 460004331/7 to install the oversize bearing to the bracket: refer to M-DLPS1011-14. Check line ream the oversize bearing to the dimension shown in Figure 601.": (
        "Используйте нажимную опору 460004330/136 и выколотку 460004331/7, чтобы установить ремонтный подшипник увеличенного размера в кронштейн: см. M-DLPS1011-14. Затем выполните контрольную линейную развертку ремонтного подшипника до размера, указанного на рисунке 601."
    ),
    "Apply sealant, Material Ref. Item 09-510A, to seal between the ends of the oversize bearing and the bracket: refer to M-DLPS709-19.": (
        "Нанесите герметик, код ссылки материала 09-510A, чтобы герметизировать зазор между торцами ремонтного подшипника увеличенного размера и кронштейном: см. M-DLPS709-19."
    ),
    "Machine diameter A sufficiently to remove the damage or wear: refer to Figure 601 for the dimensions. The surface finish must be 1,6 micrometers (63 micro-inches).": (
        "Обработайте диаметр A настолько, чтобы устранить повреждение или износ: размеры см. на рисунке 601. Шероховатость поверхности должна быть 1,6 micrometers (63 micro-inches)."
    ),
    "MACHINE TO 19,380mm (0.763in) MIN. CHROMIUM PLATE TO 20,120mm (0.7921in)": (
        "ОБРАБОТАТЬ ДО 19,380mm (0.763in) МИН. ХРОМИРОВАТЬ ДО 20,120mm (0.7921in)"
    ),
    "THE CHROMIUM PLATE MUST STOP IN THIS LENGTH. AN IRREGULAR LINE IS PERMITTED": (
        "ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ ДОЛЖНА НАХОДИТЬСЯ В ПРЕДЕЛАХ ЭТОЙ ДЛИНЫ. ДОПУСКАЕТСЯ НЕРОВНАЯ ЛИНИЯ"
    ),
    "IN THIS LENGTH. AN IRREGULAR LINE IS PERMITTED.": (
        "В ПРЕДЕЛАХ ЭТОЙ ДЛИНЫ. ДОПУСКАЕТСЯ НЕРОВНАЯ ЛИНИЯ."
    ),
    "2,00mm (0.079in) MAX. CHROMIUM PLATE MUST STOP": (
        "2,00mm (0.079in) МАКС. ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ"
    ),
    "2,00mm (0.079in) MIN. LENGTH": "2,00mm (0.079in) МИН. ДЛИНА",
    "CHROMIUM PLATE TEMINATION TO M-DLPS1031-6": "ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ ПО M-DLPS1031-6",
    "CHROMIUM PLATE TEMINATION TO M-DLPS1031-3": "ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ ПО M-DLPS1031-3",
    "CHROMIUM PLATE TEMINATION TO M-DLPS1031-1": "ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ ПО M-DLPS1031-1",
    "BEFORE CHROMIUM PLATE IS APPLIED": "ДО НАНЕСЕНИЯ ХРОМОВОГО ПОКРЫТИЯ",
    "SULPHAMATE NICKEL PLATE DEPOSIT": "СЛОЙ СУЛЬФАМАТНОГО НИКЕЛЕВОГО ПОКРЫТИЯ",
    "IT CAN BE FINISHED BY GRINDING": "МОЖНО ДОВЕСТИ ШЛИФОВАНИЕМ",
    "Measure the new diameter A.": "Измерьте новый диаметр A.",
    "SMOOTH BLEND M-DLPS1031-6": "СГЛАДИТЬ ПО M-DLPS1031-6",
    "LENGTH OF CHROMIUM PLATE": "ДЛИНА ХРОМОВОГО ПОКРЫТИЯ",
    "CHROMIUM PLATED LENGTH": "ДЛИНА ХРОМОВОГО ПОКРЫТИЯ",
    "BEFORE CHROMIUM PLATE": "ДО ХРОМИРОВАНИЯ",
    "AFTER CHROMIUM PLATE": "ПОСЛЕ ХРОМИРОВАНИЯ",
    "NOT CHROMIUM PLATED": "БЕЗ ХРОМОВОГО ПОКРЫТИЯ",
    "MIN. WALL THICKNESS": "МИН. ТОЛЩИНА СТЕНКИ",
    "CHROMIUM DEPOSIT": "СЛОЙ ХРОМОВОГО ПОКРЫТИЯ",
    "MAKE EDGES SMOOTH": "СГЛАДЬТЕ КРОМКИ",
    "AFTER GRINDING": "ПОСЛЕ ШЛИФОВАНИЯ",
    "NOT TO SCALE": "НЕ В МАСШТАБЕ",
    "15 DEGREES": "15°",
    "RAD. MIN.": "МИН. РАД.",
    "LENGTH OF": "ДЛИНА",
    "SHOT PEEN": "ДРОБЕСТРУЙНОЕ УПРОЧНЕНИЕ",
    "DIM. C": "РАЗМ. C",
    "DIM. D": "РАЗМ. D",
    "Damage or wear to diameter A.": "Повреждение или износ диаметра A.",
    "Apply cadmium plate to the machined areas: refer to PCS-2141.": (
        "Нанесите кадмиевое покрытие на обработанные участки: см. PCS-2141."
    ),
    "Apply cadmium plate to all surfaces of the repair bushes: refer to M-DLPS100-1.": (
        "Нанесите кадмиевое покрытие на все поверхности ремонтных втулок: см. M-DLPS100-1."
    ),
    "Remove the cadmium plate from the pin: refer to PCS-2100.": (
        "Удалите кадмиевое покрытие со штифта: см. PCS-2100."
    ),
    "Machine the diameter A to remove the minimum amount of material necessary to remove the damage or wear, restore the 20,00 mm (0.787 in) radius in two places as shown: refer to M-DLPS1004-4-1 and Figure 601. Do not machine diameter A more than 33,99 mm (1.3383 in). Make the surface finish 1,6 micrometers (63 micro-inches).": (
        "Обработайте диаметр A, снимая минимально необходимое количество материала для устранения повреждения или износа, и, как показано, восстановите радиус 20,00 mm (0.787 in) в двух местах: см. M-DLPS1004-4-1 и рисунок 601. Не обрабатывайте диаметр A более чем до 33,99 mm (1.3383 in). Шероховатость поверхности должна быть 1,6 micrometers (63 micro-inches)."
    ),
    "Machine the radii to the dimensions as shown: refer to Figure 601.": (
        "Обработайте радиусные участки до размеров, как показано на рисунке 601."
    ),
    "Examine the part for flaws: refer to PCS-3100 inclusion class 4 and PCS-3600.": (
        "Осмотрите деталь на наличие дефектов: см. PCS-3100, класс включений 4, и PCS-3600."
    ),
    "Shot peen the reworked areas over lapping adjacent angled surface by at least 6,35 mm (0.250 in). Almen A intensity to be 0,20 to 0,030 mm (0.008 to 0.012 in): refer to M-DLPS123.": (
        "Выполните дробеструйное упрочнение доработанных участков с перекрытием прилегающей наклонной поверхности не менее чем на 6,35 mm (0.250 in). Интенсивность по Almen A должна составлять 0,20-0,030 mm (0.008-0.012 in): см. M-DLPS123."
    ),
    "Grit blast the reworked areas: refer to PCS-2610.": (
        "Выполните пескоструйную обработку доработанных участков: см. PCS-2610."
    ),
    "Apply sulphamate nickel plate to the reworked areas: refer to MIL STD 868A solution 2, PCS-2120 and Figure 601. The sulphamate nickel plate thickness must be sufficient to get the correct diameter after machining. Make sure that the cross hole and the lubrication holes are sufficiently masked: refer to Figure 601.": (
        "Нанесите сульфаматное никелевое покрытие на доработанные участки: см. MIL STD 868A, раствор 2, PCS-2120 и рисунок 601. Толщина сульфаматного никелевого покрытия должна быть достаточной для получения требуемого диаметра после механической обработки. Убедитесь, что поперечное отверстие и смазочные отверстия достаточно замаскированы: см. рисунок 601."
    ),
    "NOTE: The above procedure includes de-embrittle for 23 hours at 185oC to 195oC (366oF to 383oF).": (
        "ПРИМЕЧАНИЕ. Приведенная выше процедура включает снятие водородной хрупкости в течение 23 часов при 185oC-195oC (366oF-383oF)."
    ),
    "Machine (do not grind) the sulphamate nickel plate to get the bore diameter of 32,468 to 32,500 mm (1.2783 to 1.2795 in): refer to Figure 601. Make the surface finish 1,6": (
        "Обработайте (не шлифуйте) сульфаматное никелевое покрытие до получения диаметра отверстия 32,468-32,500 mm (1.2783-1.2795 in): см. рисунок 601. Шероховатость поверхности должна быть 1,6"
    ),
    "Machine the 20,000 mm (0.7874 in) radii as shown: refer to Figure 601.": (
        "Обработайте радиусы 20,000 mm (0.7874 in), как показано на рисунке 601."
    ),
    "If the pin base metal has been machined, examine the machined areas for flaws: refer to PCS-3600.": (
        "Если основной металл штифта подвергался обработке, осмотрите обработанные участки на наличие дефектов: см. PCS-3600."
    ),
    "NOTE: The above procedure includes de-embrittle for 4 hours at 185oC to 195oC (366oF to 383oF).": (
        "ПРИМЕЧАНИЕ. Приведенная выше процедура включает снятие водородной хрупкости в течение 4 часов при 185oC-195oC (366oF-383oF)."
    ),
    "Grit blast the sulphamate nickel area: refer to PCS-2610. Make sure that the remainder of the pin is sufficiently masked.": (
        "Выполните пескоструйную обработку участка с сульфаматным никелевым покрытием: см. PCS-2610. Убедитесь, что остальная часть штифта достаточно замаскирована."
    ),
    "Examine the edges of sulphamate nickel plate to make sure they are properly bonded: use 5 or 10X magnification.": (
        "Осмотрите кромки сульфаматного никелевого покрытия и убедитесь в надежности их сцепления, используя увеличение 5x или 10x."
    ),
    "If there is evidence of delamination, remove the sulphamate nickel plate and do the repair again.": (
        "Если имеются признаки расслоения, удалите сульфаматное никелевое покрытие и выполните ремонт повторно."
    ),
    "Apply cadmium plate all over the pin except where chromium plated: refer to PCS-2100 and Figure 601. The cadmium plate thickness must be between 0,010 and 0,020 mm (0.0004 and 0.0008 in). Make sure the sulphamate nickel plate is fully encapsulated by cadmium plate.": (
        "Нанесите кадмиевое покрытие по всей поверхности штифта, кроме участков с хромовым покрытием: см. PCS-2100 и рисунок 601. Толщина кадмиевого покрытия должна быть 0,010-0,020 mm (0.0004-0.0008 in). Убедитесь, что сульфаматное никелевое покрытие полностью перекрыто кадмиевым покрытием."
    ),
    "Apply paint to the reworked areas: refer to REPAIR and PCS-2500.": (
        "Нанесите лакокрасочное покрытие на доработанные участки: см. РЕМОНТ и PCS-2500."
    ),
    "Machine diameter A just sufficiently to remove the damage or corrosion and to between 14,25 and 14,75 mm (0.561 and 0.581 in): refer to M-DLPS1004-4-1. The surface finish must be 1,6 micrometers (63 micro-inches).": (
        "Обработайте диаметр A ровно настолько, чтобы удалить повреждение или коррозию, и доведите его до 14,25-14,75 mm (0.561-0.581 in): см. M-DLPS1004-4-1. Шероховатость поверхности должна быть 1,6 micrometers (63 micro-inches)."
    ),
    "Measure and record the new diameters A.": "Измерьте и зафиксируйте новые диаметры A.",
    "Shot peen the reworked areas: refer to M-DLPS123 and Figure 601.": (
        "Выполните дробеструйное упрочнение доработанных участков: см. M-DLPS123 и рисунок 601."
    ),
    "Locally apply cadmium plate to the reworked areas: refer to PCS-2141.": (
        "Локально нанесите кадмиевое покрытие на доработанные участки: см. PCS-2141."
    ),
    "Calculate diameter B of the repair sleeves:": "Рассчитайте диаметр B ремонтных втулок:",
    "Dia. B = Dia. A (as measured) + 0,010 to 0,039 mm (0.0004 to 0.0015 in).": (
        "ДИАМ. B = ДИАМ. A (по результатам измерения) + 0,010-0,039 mm (0.0004-0.0015 in)."
    ),
    "Refer to Figure 601 and machine the repair sleeves to the dimensions shown and calculated.": (
        "См. рисунок 601 и обработайте ремонтные втулки до указанных и рассчитанных размеров."
    ),
    "Apply cadmium plate to the outside diameter and the chamfer of the repair sleeves: refer to PCS-2101.": (
        "Нанесите кадмиевое покрытие на наружный диаметр и фаску ремонтных втулок: см. PCS-2101."
    ),
    "Apply adhesive, Material Ref. Item 08-665, to the outside diameter of the repair sleeves: refer to PCS-5303.": (
        "Нанесите клей, код ссылки материала 08-665, на наружный диаметр ремонтных втулок: см. PCS-5303."
    ),
    "Install the repair sleeves to diameters A until aligned with the outside diameter of the pin: refer to M-DLPS1011-5.": (
        "Установите ремонтные втулки в диаметры A до совмещения с наружным диаметром штифта: см. M-DLPS1011-5."
    ),
    "Machine the bores of the repair sleeves to a diameter between 12,700 and 12,733 mm (0.5000 and 0.5013 in). Machine the inner and outer ends of both sleeves to inner and outer profile of the pin. Make the radius between 0,5 and 0,75 mm (0.020 and 0.029 in) at inside and outside of both the sleeves: refer to Figure 601.": (
        "Обработайте отверстия ремонтных втулок до диаметра 12,700-12,733 mm (0.5000-0.5013 in). Обработайте внутренние и наружные концы обеих втулок по внутреннему и наружному профилю штифта. Выполните радиус 0,5-0,75 mm (0.020-0.029 in) с внутренней и наружной стороны обеих втулок: см. рисунок 601."
    ),
    "Locally apply cadmium plate to the reworked areas of the sleeves: refer to PCS-2141.": (
        "Локально нанесите кадмиевое покрытие на доработанные участки втулок: см. PCS-2141."
    ),
    "Machine the chamfer to the dimensions shown in Figure 601.": (
        "Обработайте фаску до размеров, указанных на рисунке 601."
    ),
    "anodize the surfaces: refer to M-DLPS102-1, or": "анодируйте поверхности: см. M-DLPS102-1, или",
    "apply Alocrom to the surfaces: refer to PCS-2220.": "нанесите Alocrom на поверхности: см. PCS-2220.",
    "Machine the diameter A sufficiently to remove the damage or wear: refer to Figure 601 for the dimensions. The surface finish must be 1,6 micrometers (63 micro-inches).": (
        "Обработайте диаметр A настолько, чтобы устранить повреждение или износ: размеры см. на рисунке 601. Шероховатость поверхности должна быть 1,6 micrometers (63 micro-inches)."
    ),
    "CAUTION: FOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO MESSIER-DOWTY LIMITED: REFER TO GUIDE-CS-001.": (
        "ПРЕДУПРЕЖДЕНИЕ. При повреждении, превышающем пределы данной схемы ремонта, обратитесь в Messier-Dowty Limited: см. GUIDE-CS-001."
    ),
    "CAUTION: FOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO MESSIER-DOWTY LTD: REFER TO GUIDE-CS-001.": (
        "ПРЕДУПРЕЖДЕНИЕ. При повреждении, превышающем пределы данной схемы ремонта, обратитесь в Messier-Dowty Ltd: см. GUIDE-CS-001."
    ),
}

EXACT_MAP.update(PART5_EXACT_MAP)


PHRASE_RULES: list[tuple[str, str]] = [
    ("(Continued)", "(Продолжение)"),
    ("(Withdrawn)", "(Аннулирован)"),
    ("(Superseded)", "(Заменен)"),
    ("How To Use This Illustrated Parts List", "Как пользоваться данным иллюстрированным перечнем деталей"),
    ("Vendor Codes, Names and Addresses", "Коды поставщиков, наименования и адреса"),
    ("Detailed Parts List", "Подробный список деталей"),
    ("Numerical Index", "Числовой указатель"),
    ("Equipment and Materials", "Оборудование и материалы"),
    ("Description and Operation", "Описание и работа"),
    ("Testing and Fault Isolation", "Испытания и поиск неисправностей"),
    ("Special Detailed Inspection", "Специальный детальный осмотр"),
    ("Detailed Inspection", "Детальный осмотр"),
    ("Test Conditions", "Условия испытаний"),
    ("Fault Isolation", "Поиск неисправностей"),
    ("Special Tools, Fixtures and Equipment", "Специальные инструменты, приспособления и оборудование"),
    ("Fits and Clearances Definitions", "Определения посадок и зазоров"),
    ("Fits and Clearances", "Посадки и зазоры"),
    ("Electrical Bonding Resistance Test Points", "Точки проверки сопротивления электрического соединения"),
    ("Diagram of Operation", "Схема работы"),
    ("Illustrated Parts List", "Иллюстрированный перечень деталей"),
    ("Component Maintenance Manual", "Руководство по техническому обслуживанию компонента"),
    ("Main Landing Gear Leg", "Стойка основного шасси"),
    ("List of Effective Pages", "Перечень действующих страниц"),
    ("List of Service Bulletins", "Список сервисных бюллетеней"),
    ("Record of Temporary Revisions", "Журнал временных ревизий"),
    ("Record of Revisions", "Журнал ревизий"),
    ("Revision Record", "Журнал ревизий"),
    ("Unit Identification Chart", "Таблица идентификации агрегата"),
    ("Letter of Transmittal", "Сопроводительное письмо"),
    ("Title Page", "Титульный лист"),
    ("Table of Contents", "Содержание"),
    ("Illustrations", "Иллюстрации"),
    ("Repair Procedure Conditions", "Условия выполнения процедуры ремонта"),
    ("Approved Repairs", "Утвержденные ремонты"),
    ("Protective Treatment", "Защитная обработка"),
    ("Surface Damage", "Повреждение поверхности"),
    ("Identification", "Идентификация"),
    ("Cadmium Plated Surfaces", "Кадмированные поверхности"),
    ("Anodized Surfaces", "Анодированные поверхности"),
    ("Paint Finish", "Лакокрасочное покрытие"),
    ("Key Diagram", "Схема расположения"),
    ("Airbus Maintenance Planning Document", "документ Airbus по планированию технического обслуживания"),
    ("Maintenance Planning Document", "документ по планированию технического обслуживания"),
    ("Torque Link Repairs", "Ремонты шлиц-шарнира"),
    ("Torque Link Ремонтs", "Ремонты шлиц-шарнира"),
    ("Корпус стойки Ремонтs", "Ремонты корпуса стойки"),
    ("Скользящая труба Ремонтs", "Ремонты скользящей трубы"),
    ("Верхняя диафрагменная труба Ремонтs", "Ремонты верхней диафрагменной трубы"),
    ("Цилиндр Ремонтs", "Ремонты цилиндра"),
    ("Переходной блок Ремонтs", "Ремонты переходного блока"),
    ("Кронштейн крепления жгута Ремонтs", "Ремонты кронштейна крепления жгута"),
    ("Верхний кронштейн шарнира Ремонтs", "Ремонты верхнего кронштейна шарнира"),
    ("Torque Data", "Данные по моментам затяжки"),
    ("Remarks", "Примечания"),
    ("Introduction", "Введение"),
    ("General", "Общие сведения"),
    ("Procedure", "Процедура"),
    ("Storage", "Хранение"),
    ("Including", "включая"),
    ("Description", "Описание"),
    ("Operation", "Работа"),
    ("Data", "Данные"),
    ("Cleaning", "Очистка"),
    ("Check", "Проверка"),
    ("Assembly", "Сборка"),
    ("Disassembly", "Разборка"),
    ("Repair", "Ремонт"),
    ("Tables", "таблицы"),
    ("Table", "таблица"),
    ("and/or", "и/или"),
    (" and ", " и "),
    (" or ", " или "),
    ("EXTERNAL SERMETEL LIMIT", "НАРУЖНАЯ ГРАНИЦА ПОКРЫТИЯ SERMETEL"),
    ("INTERNAL SERMETEL LIMIT", "ВНУТРЕННЯЯ ГРАНИЦА ПОКРЫТИЯ SERMETEL"),
    ("PRIMER PAINT", "ГРУНТОВОЧНАЯ КРАСКА"),
    ("EXTENT", "ПРОТЯЖЕННОСТЬ"),
    ("DIA ", "ДИАМ. "),
    ("RAD IUS", "РАДИУС"),
    ("TYPICA", "ТИПОВО"),
    ("FOR MAIN FITTING", "ДЛЯ КОРПУСА СТОЙКИ"),
    ("DETAIL ", "ДЕТАЛЬ "),
    ("VIEW ", "ВИД "),
    ("SECTION ", "СЕЧЕНИЕ "),
    ("(BOTH SIDES)", "(С ОБЕИХ СТОРОН)"),
    ("KNUCKLE HOLES", "ОТВЕРСТИЯ ШАРНИРА"),
    ("Messier-Dowty Limited or Safran Landing Systems Repair No.", "Номер ремонта Messier-Dowty Limited или Safran Landing Systems"),
    ("Applicable Part", "Применяемая деталь"),
    ("NO ZINC-NICKEL OR PAINT DEPOSIT TO REMAIN ON OR PROUD OF WORKING DIA. AFTER GRINDING CHROME", "НИ ЦИНК-НИКЕЛЕВОЕ, НИ ЛАКОКРАСОЧНОЕ ПОКРЫТИЕ НЕ ДОЛЖНО ОСТАВАТЬСЯ НА РАБОЧЕМ ДИАМ. ИЛИ ВЫСТУПАТЬ НАД НИМ ПОСЛЕ ШЛИФОВАНИЯ ХРОМА"),
    ("NO ZINC-NICKEL OR PAINT DEPOSIT TO REMAIN", "НИ ЦИНК-НИКЕЛЕВОЕ, НИ ЛАКОКРАСОЧНОЕ ПОКРЫТИЕ НЕ ДОЛЖНО ОСТАВАТЬСЯ"),
    ("ON OR PROUD OF WORKING DIA. AFTER GRINDING CHROME", "НА РАБОЧЕМ ДИАМ. ИЛИ ВЫСТУПАТЬ НАД НИМ ПОСЛЕ ШЛИФОВАНИЯ ХРОМА"),
    ("NO ZINC-NICKEL OR PAINT DEPOSIT", "НИ ЦИНК-НИКЕЛЕВОЕ, НИ ЛАКОКРАСОЧНОЕ ПОКРЫТИЕ"),
    ("NO ZINC-NICKEL OR PAINT ON THIS", "НЕ ДОПУСКАЕТСЯ ЦИНК-НИКЕЛЕВОЕ ИЛИ ЛАКОКРАСОЧНОЕ ПОКРЫТИЕ НА ЭТОЙ"),
    ("TO REMAIN ON WORKING DIA. AFTER CHROME PLATING", "НЕ ДОЛЖНО ОСТАВАТЬСЯ НА РАБОЧЕМ ДИАМ. ПОСЛЕ ХРОМИРОВАНИЯ"),
    ("CHROME PLATING WILL TERMINATE ANYWHERE ON THE CHAMFER", "ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ МОЖЕТ РАСПОЛАГАТЬСЯ В ЛЮБОМ МЕСТЕ НА ФАСКЕ"),
    ("BREAK FLANGE FACE CHROME TERMINATION (TYPICAL)", "ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ НА ПОВЕРХНОСТИ ТОРМОЗНОГО ФЛАНЦА (ТИПОВО)"),
    ("JOURNAL A,B,C CHROME TERMINATION (TYPICAL)", "ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ ШЕЕК A, B, C (ТИПОВО)"),
    ("JOURNAL A OUTER CHROME TERMINATION (TYPICAL)", "НАРУЖНАЯ ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ ШЕЙКИ A (ТИПОВО)"),
    ("JOURNAL C INNER CHROME TERMINATION (TYPICAL)", "ВНУТРЕННЯЯ ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ ШЕЙКИ C (ТИПОВО)"),
    ("HPC SEAL ABUTMENT LOWER CHROME TERMINATION", "НИЖНЯЯ ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ УПОРА УПЛОТНЕНИЯ HPC"),
    ("HPC SEAL ABUTMENT UPPER CHROME TERMINATION", "ВЕРХНЯЯ ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ УПОРА УПЛОТНЕНИЯ HPC"),
    ("BARREL OUTER DIA. LOWER CHROME TERMINATION", "НИЖНЯЯ ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ НАРУЖ. ДИАМ. ЦИЛИНДРИЧЕСКОЙ ЧАСТИ"),
    ("BARREL OUTER DIA. UPPER CHROME TERMINATION", "ВЕРХНЯЯ ГРАНИЦА ХРОМОВОГО ПОКРЫТИЯ НАРУЖ. ДИАМ. ЦИЛИНДРИЧЕСКОЙ ЧАСТИ"),
    ("FULL CHROME PLATING THICKNESS", "ТОЛЩИНА СПЛОШНОГО ХРОМОВОГО ПОКРЫТИЯ"),
    ("CHROME RUN OUT BAND", "ПОЛОСА СХОДА ХРОМОВОГО ПОКРЫТИЯ"),
    ("RUN OUT BAND", "ПОЛОСА СХОДА ПОКРЫТИЯ"),
    ("PAINT DEPOSIT OVERLAP", "ПЕРЕКРЫТИЕ ЛАКОКРАСОЧНОГО ПОКРЫТИЯ"),
    ("ZINC-NICKEL DEPOSIT OVERLAP AND", "ПЕРЕКРЫТИЕ ЦИНК-НИКЕЛЕВОГО ПОКРЫТИЯ И"),
    ("ZINC-NICKEL DEPOSIT OVERLAP", "ПЕРЕКРЫТИЕ ЦИНК-НИКЕЛЕВОГО ПОКРЫТИЯ"),
    ("DEPOSIT OVERLAP", "ПЕРЕКРЫТИЕ ПОКРЫТИЯ"),
    ("CHROME PLATING DEPOSIT", "НАПЛЫВ ХРОМОВОГО ПОКРЫТИЯ"),
    ("CHROME PLATING", "ХРОМОВОЕ ПОКРЫТИЕ"),
    ("CHROMIUM PLATED SURFACE", "ПОВЕРХНОСТЬ С ХРОМОВЫМ ПОКРЫТИЕМ"),
    ("DIA. AFTER CHROME PLATING", "ДИАМ. ПОСЛЕ ХРОМИРОВАНИЯ"),
    ("DIA. AFTER GRINDING CHROME", "ДИАМ. ПОСЛЕ ШЛИФОВАНИЯ ХРОМА"),
    ("DIA. AFTER GRINDING", "ДИАМ. ПОСЛЕ ШЛИФОВАНИЯ"),
    ("EDGE BLENDED WITH A SMOOTH TRANSITION", "КРОМКА СГЛАЖЕНА С ПЛАВНЫМ ПЕРЕХОДОМ"),
    ("EDGE BLENDED WITH", "КРОМКА СГЛАЖЕНА С"),
    ("A SMOOTH TRANSITION WAVY OR IRREGULAR LINE PERMISSIBLE", "ПЛАВНЫЙ ПЕРЕХОД. ДОПУСКАЕТСЯ ВОЛНИСТАЯ ИЛИ НЕРОВНАЯ ЛИНИЯ"),
    ("WAVY OR IRREGULAR LINE PERMISSIBLE", "ДОПУСКАЕТСЯ ВОЛНИСТАЯ ИЛИ НЕРОВНАЯ ЛИНИЯ"),
    ("A SMOOTH TRANSITION", "ПЛАВНЫЙ ПЕРЕХОД"),
    ("(TO HPC ABUTMENT FACE)", "(ДО УПОРНОГО ТОРЦА HPC)"),
    ("TO RADIUS INTERSECTION POINT", "ДО ТОЧКИ ПЕРЕСЕЧЕНИЯ С РАДИУСОМ"),
    ("CROSS BORE FOR", "ПОПЕРЕЧНОЕ ОТВЕРСТИЕ ДЛЯ"),
    ("CHANGE OVER VALVE HOLES AND LUGS", "ОТВЕРСТИЯ И ПРОУШИНЫ ПЕРЕКЛЮЧАЮЩЕГО КЛАПАНА"),
    ("CHANGE OVER VALVE HOLES", "ОТВЕРСТИЯ ПЕРЕКЛЮЧАЮЩЕГО КЛАПАНА"),
    ("LOWER CARDAN BORE", "ОТВЕРСТИЕ НИЖНЕГО КАРДАНА"),
    ("LOWER DOOR LUGS", "ПРОУШИНЫ НИЖНЕЙ ДВЕРЦЫ"),
    ("UPPER DOOR LUGS", "ПРОУШИНЫ ВЕРХНЕЙ ДВЕРЦЫ"),
    ("UPLOCK LUGS", "ПРОУШИНЫ ЗАМКА УБРАННОГО ПОЛОЖЕНИЯ"),
    ("KNUCKLE TOOLING LUG", "ТЕХНОЛОГИЧЕСКАЯ ПРОУШИНА ШАРНИРА"),
    ("TOOLING LUG", "ТЕХНОЛОГИЧЕСКАЯ ПРОУШИНА"),
    ("GREASE HOLES", "СМАЗОЧНЫЕ ОТВЕРСТИЯ"),
    ("LOCK LINK BORE", "ОТВЕРСТИЕ ЗВЕНА ФИКСАТОРА"),
    ("BRAKE FLANGE", "ТОРМОЗНОЙ ФЛАНЕЦ"),
    ("EXTENT OF FINE LIMIT DIA.", "ПРОТЯЖЕННОСТЬ УЧАСТКА МЕНЬШЕГО ПРЕД. ДИАМ."),
    ("OF FINE LIMIT DIA.", "УЧАСТКА МЕНЬШЕГО ПРЕД. ДИАМ."),
    ("PRIMER PAINT FACE C", "ГРУНТОВОЧНАЯ КРАСКА, ПОВЕРХНОСТЬ C"),
    ("C AFTER THREAD", "C ПОСЛЕ НАРЕЗАНИЯ РЕЗЬБЫ"),
    ("DIA.", "ДИАМ."),
    ("CHAMFERS", "ФАСКИ"),
    ("CHAMFER", "ФАСКА"),
    ("RADIUS", "РАДИУС"),
    ("TYPICAL", "ТИПОВО"),
    ("Repair No.", "Ремонт №"),
    ("Repair to", "Ремонт"),
    ("Repair Bushes", "Ремонтные втулки"),
    ("Repair Bush(es)", "Ремонтная втулка(и)"),
    ("Repair Bush", "Ремонтная втулка"),
    ("Repair Bearing", "Подшипник"),
    ("Oversize Bushes", "Ремонтные втулки увеличенного размера"),
    ("Oversize Lubrication adapter", "Увеличенный смазочный адаптер"),
    ("Oversize Transfer Dowel", "Увеличенный переходной штифт"),
    ("Machining and Installation", "механическая обработка и установка"),
    ("Machining and installation", "механическая обработка и установка"),
    ("Machining", "механическая обработка"),
    ("Installation", "установка"),
    ("Inner Liner", "внутренний вкладыш"),
    ("Liner", "вкладыш"),
    ("Chromium Plate Termination", "завершение хромового покрытия"),
    ("Sheet", "лист"),
    ("Only", "только"),
    ("Vendor", "Поставщик"),
    ("Subject", "Тема"),
    ("Chart", "Таблица"),
    ("Fig.", "Рис."),
    ("ain Fitting", "корпуса стойки"),
    ("Main Fitting Subassembly", "Сборка корпуса стойки"),
    ("Main Fitting", "Корпус стойки"),
    ("Sliding Tube", "Скользящая труба"),
    ("Lower Bearing Subassembly", "Сборка нижнего подшипника"),
    ("Lower Torque Link", "Нижний шлиц-шарнир"),
    ("Upper Torque Link", "Верхний шлиц-шарнир"),
    ("Upper Slave Link", "Верхнее ведомое звено"),
    ("Lower Slave Link", "Нижнее ведомое звено"),
    ("Upper Diaphragm Tube Subassembly", "Сборка верхней диафрагменной трубы"),
    ("Upper Diaphragm Tube", "Верхняя диафрагменная труба"),
    ("Upper Pivot Bracket", "Верхний кронштейн шарнира"),
    ("Pivot Bracket", "Кронштейн шарнира"),
    ("Harness Support Bracket", "Кронштейн крепления жгута"),
    ("Transfer Block", "Переходной блок"),
    ("transfer block", "переходной блок"),
    ("Lock Stay Cardan", "Кардан фиксатора"),
    ("Spherical Bearing", "Сферический подшипник"),
    ("Pintle Pin", "Штифт навеса стойки"),
    ("Forward Pintle Pin", "Передний штифт навеса стойки"),
    ("Retaining Pin", "Стопорный штифт"),
    ("Valve Stem", "Шток клапана"),
    ("Inflation Valve", "Заправочный клапан"),
    ("Uplock Pin", "Штифт замка убранного положения"),
    ("Pivot Pin", "Шарнирный штифт"),
    ("Pin", "Штифт"),
    ("Bracket", "Кронштейн"),
    ("Slave Link", "Ведомое звено"),
    ("Cylinder", "Цилиндр"),
    ("Spacer", "Проставка"),
    ("Drag-arm Spacer", "Проставка тяги складывания"),
    ("Special Tools, Fixtures", "Специальные инструменты, приспособления"),
    ("and Equipment", "и оборудование"),
    ("Reference Publications", "Справочные публикации"),
    ("Application of Jointing Compound", "Нанесение герметизирующего состава"),
    ("Electrical Bonding Resistance Tests", "Проверка сопротивления электрического соединения"),
    ("Leakage Tests", "испытания на герметичность"),
    ("Maintenance procedures", "процедуры технического обслуживания"),
    ("Process Specifications", "технологические спецификации"),
    ("Non-destructive Tests", "неразрушающий контроль"),
    ("Technical Publications", "технические публикации"),
    ("on-line service", "онлайн-служба"),
    ("Imperial units", "британские единицы"),
    ("SI units", "единицы СИ"),
    ("Figure and Item numbers", "номера рисунков и позиций"),
    ("Item No.", "Поз. №"),
    ("Refer to Figures", "См. рисунки"),
    ("Refer to Figure", "См. рисунок"),
    ("Subassemblies", "Подсборки"),
    ("Subassembly", "Подсборка"),
    ("subassemblies", "подсборки"),
    ("subassembly", "подсборка"),
    ("Sliding Tube Subassembly", "Подсборка скользящей трубы"),
    ("Bracket Subassembly", "Подсборка кронштейна"),
    ("Bolt Subassembly", "Подсборка болта"),
    ("Installation of Bushes", "Установка втулок"),
    ("Installation of Labels", "Установка табличек"),
    ("Assembly of Lower Bearing Subassembly", "Сборка нижнего подшипника"),
    ("Seal Configuration", "Конфигурация уплотнения"),
    ("Configuration", "Конфигурация"),
    ("Crimping of the Pin", "Расклепка штифта"),
    ("Application of Ardrox AV100D to the", "Нанесение Ardrox AV100D на"),
    ("to the Upper Diaphragm Tube", "на верхнюю диафрагменную трубу"),
    ("to the Pin", "на штифт"),
    ("Ardrox Application", "Нанесение Ardrox"),
    ("Dimensions After Installation in the Gland Housing", "Размеры после установки в корпус сальника"),
    ("Grease Groove", "Смазочная канавка"),
    ("Gland Housing", "Корпус сальника"),
    ("Lower Bearing", "Нижний подшипник"),
    ("Turner Inflation Equipment", "Заправочное оборудование Turner"),
    ("Axle Harness", "Жгут оси"),
    ("Dressings", "Навесные элементы"),
    ("Hydraulic-Pneumatic Pump Set", "Гидропневматическая насосная установка"),
    ("Hydraulic Test Rig", "Гидравлический испытательный стенд"),
    ("Nitrogen Supply", "Источник азота"),
    ("28 VDC Power Supply", "Источник питания 28 В пост. тока"),
    ("Crowfoot Wrench", "Рожковый ключ типа Crowfoot"),
    ("Charging Adapter", "Заправочный адаптер"),
    ("Holding Fixture", "Удерживающее приспособление"),
    ("Holding Blocks", "Удерживающие блоки"),
    ("Location Frame", "Установочная рама"),
    ("Torque Adapter", "Моментный адаптер"),
    ("Bottom Press Adapter", "Нижний нажимной адаптер"),
    ("Jacking Dome Adapter", "Адаптер поддомкратного купола"),
    ("Press Adapter", "Нажимной адаптер"),
    ("Extractor Plate", "Плита съемника"),
    ("Offset Adapter", "Смещенный адаптер"),
    ("Extraction Tube", "Выпрессовочная трубка"),
    ("Extraction Pad", "Выпрессовочная опора"),
    ("Extraction Bar", "Выпрессовочная штанга"),
    ("Lifting Tackle", "Подъемная оснастка"),
    ("Transport and Build Trolley", "Тележка для транспортировки и сборки"),
    ("Build Trolley", "Сборочная тележка"),
    ("Support Arms", "Опорные рычаги"),
    ("Towing Frame", "Буксировочная рама"),
    ("Bench Clamp", "Слесарные тиски"),
    ("Milliohmmeter Megger, Type BT51", "Миллиомметр Megger, тип BT51"),
    ("Proximity switch connector shell", "Корпус разъема датчика приближения"),
    ("Static discharge connector", "Разъем отвода статического электричества"),
    ("Heat shrink sleeve", "Термоусадочная трубка"),
    ("Ferrule", "Обжимная втулка"),
    ("Lampbox", "Ламповый блок"),
    ("Rod end", "Шарнирный наконечник"),
    ("Bowden cable", "Трос Боудена"),
    ("Labels", "Таблички"),
    ("Label", "Табличка"),
    ("Bushes", "Втулки"),
    ("Bush", "Втулка"),
    ("cap screws", "винты с цилиндрической головкой"),
    ("cap screw", "винт с цилиндрической головкой"),
    ("shims", "регулировочные прокладки"),
    ("shim", "регулировочная прокладка"),
    ("laminated shims", "наборные регулировочные прокладки"),
    ("laminated shim", "наборная регулировочная прокладка"),
    ("Bolts", "Болты"),
    ("Nuts", "Гайки"),
    ("Washers", "Шайбы"),
    ("Spacers", "Проставки"),
    ("lubrication fittings", "смазочные штуцеры"),
    ("lubrication fitting", "смазочный штуцер"),
    ("ground stud", "заземляющий штырь"),
    ("outer race", "наружная обойма"),
    ("retainers", "фиксаторы"),
    ("wedge", "клин"),
    ("target", "мишень"),
    ("diaphragm", "диафрагма"),
    ("baffle", "перегородка"),
    ("rod", "шток"),
    ("piston", "поршень"),
    ("shock absorber", "амортизатор"),
    ("unit", "агрегат"),
    ("Sealing ring", "Уплотнительное кольцо"),
    ("Seal", "Уплотнение"),
    ("Sleeve", "Втулка"),
    ("Washer", "Шайба"),
    ("Screws", "Винты"),
    ("Bolt", "Болт"),
    ("Nut", "Гайка"),
    ("Bung", "Заглушка"),
    ("Adapter", "Адаптер"),
    ("Extractor", "Съемник"),
    ("Drift", "Выколотка"),
    ("Press Pad", "Нажимная опора"),
    ("Joint seal", "Герметизирующее уплотнение"),
    ("Wiper ring", "Грязесъемное кольцо"),
    ("washer(s)", "шайба(ы)"),
    ("Backing rings", "Опорные кольца"),
    ("Backing ring", "Опорное кольцо"),
    ("Split pins", "Шплинты"),
    ("Split pin", "Шплинт"),
    ("Retaining pins", "Стопорные штифты"),
    ("Retaining pin", "Стопорный штифт"),
    ("Locking pins", "Стопорные штифты"),
    ("Locking pin", "Стопорный штифт"),
    ("Locking plate", "Стопорная пластина"),
    ("Lock plate", "Стопорная пластина"),
    ("Locking washer", "Стопорная шайба"),
    ("Lock washer", "Стопорная шайба"),
    ("Tab washers", "Отгибные шайбы"),
    ("Tab washer", "Отгибная шайба"),
    ("Cup washers", "Чашечные шайбы"),
    ("Cup washer", "Чашечная шайба"),
    ("O-ring seals", "Уплотнительные кольца круглого сечения"),
    ("O-ring seal", "Уплотнительное кольцо круглого сечения"),
    ("Wire Thread Inserts", "Резьбовые вставки"),
    ("Wire Thread Insert", "Резьбовая вставка"),
    ("wire thread inserts", "резьбовые вставки"),
    ("wire thread insert", "резьбовая вставка"),
    ("Oversize Thread Insert", "Резьбовая вставка увеличенного размера"),
    ("Threaded insert", "Резьбовая вставка"),
    ("Retaining ring", "Стопорное кольцо"),
    ("Level tube", "Уровневая трубка"),
    ("Clapper seat", "Седло клапана"),
    ("Compression orifice plate", "Пластина дроссельного отверстия сжатия"),
    ("Recoil orifice plate", "Пластина дроссельного отверстия отбоя"),
    ("Special Bolt", "Специальный болт"),
    ("Cross Bolt", "Поперечный болт"),
    ("Hydraulic fluid", "Гидравлическая жидкость"),
    ("Weight without hydraulic fluid", "Масса без гидравлической жидкости"),
    ("Weight with hydraulic fluid", "Масса с гидравлической жидкостью"),
    ("approximately", "приблизительно"),
    ("Nitrogen", "Азот"),
    ("Paint Removal", "Удаление краски"),
    ("Examine Parts Visually", "Визуально осмотрите детали"),
    ("Examine Dimensions", "Проверьте размеры"),
    ("Application of", "Нанесение"),
    ("Jointing Compound", "герметизирующего состава"),
    ("Hole Locations", "Расположение отверстий"),
    ("Materials", "Материалы"),
    ("Material Type", "Тип материала"),
    ("Material", "Материал"),
    ("Part Name", "Наименование детали"),
    ("Function", "Назначение"),
    ("Equipment", "Оборудование"),
    ("Special Tool", "Специальный инструмент"),
    ("Special Tools", "Специальные инструменты"),
    ("Part No.", "Номер детали"),
    ("EASA No.", "EASA №"),
    ("Name", "Наименование"),
    ("Repair Sleeves", "Ремонтные втулки"),
    ("Bearings", "Подшипники"),
    ("Bearing", "Подшипник"),
    ("Damper", "Демпфер"),
    ("Parts", "Части"),
    ("Ref. Item identification", "код ссылки"),
    ("Material Ref. Item", "Код ссылки материала"),
    ("Ref. Item", "Код ссылки"),
    ("LIMIT VALUE MILLIOHMS", "ПРЕДЕЛЬНОЕ ЗНАЧЕНИЕ, мОм"),
    ("TEST POINT", "ТОЧКА ПРОВЕРКИ"),
    ("IPL FIGURE AND ITEM No.", "РИС./ПОЗ. IPL"),
    ("IPL Fig/Item No.", "Рис./поз. IPL"),
    ("IPL Fig/Item", "Рис./поз. IPL"),
    ("Fig Item No.", "Рис./поз. №"),
    ("Fig Item", "Рис./поз."),
    ("Use with", "Использовать с"),
    ("Use the", "Используйте"),
    ("Hold the main landing gear leg", "Удерживать стойку основного шасси"),
    ("Hold the upper diaphragm tube", "Удерживать верхнюю диафрагменную трубу"),
    ("Hold the", "Удерживать"),
    ("Lift the", "Поднимать"),
    ("Main landing gear leg (1-1) tests", "Испытания стойки основного шасси (1-1)"),
    ("Proximity switch and target tests", "Испытания датчика приближения и мишени"),
    ("Use the Extractor", "Используйте съемник"),
    ("Use the Torque Adapter", "Используйте моментный адаптер"),
    ("Use the Charging Adapter", "Используйте заправочный адаптер"),
    ("Use the Crowfoot Wrench", "Используйте рожковый ключ типа Crowfoot"),
    ("to remove the bushes", "для снятия втулок"),
    ("to remove the bush", "для снятия втулки"),
    ("to remove the bearings", "для снятия подшипников"),
    ("to remove the bearing", "для снятия подшипника"),
    ("to remove the nuts", "для снятия гаек"),
    ("to remove the nut", "для снятия гайки"),
    ("to remove the forward pintle bush", "для снятия втулки переднего штифта навеса"),
    ("to hold the pin", "для удержания штифта"),
    ("to remove the charging valve", "для снятия заправочного клапана"),
    ("to remove the lubrication adapters", "для снятия смазочных адаптеров"),
    ("Remove the bushes", "Снимите втулки"),
    ("Remove the bush", "Снимите втулку"),
    ("Remove the", "Снимите"),
    ("Remove the bearings", "Снимите подшипники"),
    ("Remove the bearing", "Снимите подшипник"),
    ("Remove the nuts", "Снимите гайки"),
    ("Remove the nut", "Снимите гайку"),
    ("Remove and discard", "Снимите и выбросьте"),
    ("Open the charging valve", "Откройте заправочный клапан"),
    ("Close the charging valve", "Закройте заправочный клапан"),
    ("Close the charging valves", "Закрыть заправочные клапаны"),
    ("Remove the locking plate", "Снимите стопорную пластину"),
    ("Remove the locking nut", "Снимите стопорную гайку"),
    ("Remove the stop rings", "Снимите стопорные кольца"),
    ("Remove the jacking dome", "Снимите поддомкратный купол"),
    ("Remove the level tube", "Снимите уровневую трубку"),
    ("Remove the upper bearing housing", "Снимите корпус верхнего подшипника"),
    ("Remove the transfer dowels", "Снимите переходные штифты"),
    ("Remove the wiring diagram plate", "Снимите табличку электрической схемы"),
    ("Remove the lubrication adapters", "Снимите смазочные адаптеры"),
    ("Remove drag arm sleeve", "Снимите втулку тяги складывания"),
    ("Remove the 2M electrical axle harness", "Снимите электрический жгут оси 2M"),
    ("Remove the 1M electrical axle harness", "Снимите электрический жгут оси 1M"),
    ("Remove the charging valves", "Снимите заправочные клапаны"),
    ("Remove the lubrication fittings", "Снимите смазочные штуцеры"),
    ("Remove the lubrication fitting", "Снимите смазочный штуцер"),
    ("Remove the identification washers", "Снимите идентификационные шайбы"),
    ("Remove the identification washer", "Снимите идентификационную шайбу"),
    ("Remove the bonding cable", "Снимите кабель заземления"),
    ("Remove the related parts", "Снимите связанные детали"),
    ("wiring diagram plate", "табличка электрической схемы"),
    ("identification washers", "идентификационные шайбы"),
    ("identification washer", "идентификационная шайба"),
    ("bonding cable", "кабель заземления"),
    ("related parts", "связанные детали"),
    ("attached parts", "присоединенные детали"),
    ("its related parts", "связанные с ней детали"),
    ("its attached parts", "присоединенные к ней детали"),
    ("charging valves", "заправочные клапаны"),
    ("forward pintle bush", "втулка переднего штифта навеса"),
    ("cardan assembly", "кардан в сборе"),
    ("Cardan Assembly", "Кардан в сборе"),
    ("Lock Stay Cardan Subassembly", "Подсборка кардана фиксатора"),
    ("lock stay cardan subassembly", "подсборка кардана фиксатора"),
    ("Proximity Switches", "Датчики приближения"),
    ("Proximity Switch", "Датчик приближения"),
    ("proximity switches", "датчики приближения"),
    ("proximity switch", "датчик приближения"),
    ("Adjustment.", "Регулировка."),
    ("harness support bracket", "кронштейн крепления жгута"),
    ("harness support", "опора жгута"),
    ("pivot bracket subassembly", "подсборка кронштейна шарнира"),
    ("lubrication shaft subassembly", "подсборка смазочного вала"),
    ("dust cap", "пылезащитный колпачок"),
    ("clamp", "хомут"),
    ("slotted nut", "прорезная гайка"),
    ("grooved spherical bearing", "сферический подшипник с канавкой"),
    ("self lubricating bearing", "самосмазывающийся подшипник"),
    ("Torque Reaction Adapter", "Адаптер реактивного момента"),
    ("Press Pad Assembly", "Сборка нажимной опоры"),
    ("Extraction Tube", "Выпрессовочная трубка"),
    ("Extractor Plate", "Плита съемника"),
    ("Assembly/Extraction Tool", "Сборочно-демонтажный инструмент"),
    ("Extractor Pad and Drawbolt", "Съемная опора и вытяжной болт"),
    ("Lifting Bar Assembly", "Подъемная штанга в сборе"),
    ("Pintle Location Assembly", "Установочное приспособление штифта навеса"),
    ("Spherical Bearing Locator", "Установочный шаблон сферического подшипника"),
    ("Pin Spanner", "Штифтовой ключ"),
    ("left configuration", "левая конфигурация"),
    ("right configuration", "правая конфигурация"),
    ("Build Trolley", "Сборочная тележка"),
    ("Transport and Build", "Транспортировочная и сборочная"),
    ("Release the tab washer", "Отогните усик отгибной шайбы"),
    ("Release the tab washers", "Отогните усики отгибных шайб"),
    ("Release the cup washers", "Освободите чашечные шайбы"),
    ("Remove the two piece stop with inserts", "Снимите двухсоставной упор с вставками"),
    ("Remove the damaged paint", "Удалите поврежденное лакокрасочное покрытие"),
    ("refer to", "см."),
    ("Remove the bung", "Снимите заглушку"),
    ("Dry all the metal parts.", "Высушите все металлические детали."),
    ("Upper dia- phragm tube", "Верхняя диафрагменная труба"),
    ("The Adapter", "Адаптер"),
    ("POST SB", "ПОСЛЕ SB"),
    ("PRE SB", "ДО SB"),
    ("Housing", "Корпус"),
    ("Torque Reaction", "Реактивный момент"),
    ("Torque Reactor", "Упор реактивного момента"),
    ("Loading Press", "Нагрузочный пресс"),
    ("White spirit", "Уайт-спирит"),
    ("Jacking dome", "Поддомкратный купол"),
    ("Compression", "Сжатие"),
    ("Recoil", "Отбой"),
    ("Apply cadmium plate all over", "Нанесите кадмиевое покрытие по всей поверхности"),
    ("Apply cadmium plate", "Нанесите кадмиевое покрытие"),
    ("Paint: refer to", "Окраска: см."),
    ("and PCS-2500. Do not paint:", "и PCS-2500. Не окрашивать:"),
    ("Apply paint all over", "Нанесите лакокрасочное покрытие по всей поверхности"),
    ("Apply paint", "Нанесите лакокрасочное покрытие"),
    ("Apply paint all over but not to", "Нанесите лакокрасочное покрытие по всей поверхности, но не на"),
    ("Apply paint all over but not on", "Нанесите лакокрасочное покрытие по всей поверхности, но не на"),
    ("Apply paint all over externally, but not to", "Нанесите лакокрасочное покрытие по всей наружной поверхности, но не на"),
    ("Apply paint internally along", "Нанесите лакокрасочное покрытие изнутри вдоль"),
    ("Apply primer paint", "Нанесите грунтовочную краску"),
    ("Apply primer paint only to", "Нанесите грунтовочную краску только на"),
    ("Apply only primer paint to", "Нанесите только грунтовочную краску на"),
    ("Apply primer", "Нанесите грунт"),
    ("Apply primer to", "Нанесите грунт на"),
    ("Apply one layer of primer paint", "Нанесите один слой грунтовочной краски"),
    ("Apply one layer of primer paint only to", "Нанесите один слой грунтовочной краски только на"),
    ("Apply one coat of primer", "Нанесите один слой грунта"),
    ("Apply one coat of primer to", "Нанесите один слой грунта на"),
    ("Apply a thin coat of primer paint", "Нанесите тонкий слой грунтовочной краски"),
    ("Apply a thin coat of primer paint to", "Нанесите тонкий слой грунтовочной краски на"),
    ("Apply a light coat of primer", "Нанесите тонкий слой грунта"),
    ("Apply a light coat of primer to", "Нанесите тонкий слой грунта на"),
    ("Apply wet primer", "Нанесите жидкий грунт"),
    ("Apply wet primer to", "Нанесите жидкий грунт по"),
    ("apply resin to", "нанесите смолу по"),
    ("apply resin", "нанесите смолу"),
    ("Apply Alocrom 1200", "Нанесите Alocrom 1200"),
    ("Apply Alocrom 1200 to", "Нанесите Alocrom 1200 на"),
    ("Apply brush cadmium plate", "Нанесите кистью кадмиевое покрытие"),
    ("Do not apply paint", "Не наносить лакокрасочное покрытие"),
    ("Do not paint", "Не окрашивать"),
    ("Do not include", "Не включать"),
    ("Passivate all over", "Пассивируйте по всей поверхности"),
    ("Passivate", "Пассивируйте"),
    ("Chromic acid anodise all over but not", "Анодируйте в хромовой кислоте по всей поверхности, кроме"),
    ("Chromic acid anodize all over but not", "Анодируйте в хромовой кислоте по всей поверхности, кроме"),
    ("Chromic acid anodise all over", "Анодируйте в хромовой кислоте по всей поверхности"),
    ("Chromic acid anodize all over", "Анодируйте в хромовой кислоте по всей поверхности"),
    ("Chromic acid anodise", "Анодируйте в хромовой кислоте"),
    ("Chromic acid anodize", "Анодируйте в хромовой кислоте"),
    ("Anodise all over, but not", "Анодируйте по всей поверхности, кроме"),
    ("Anodise all over", "Анодируйте по всей поверхности"),
    ("all over", "по всей поверхности"),
    ("Make the cadmium plate thickness between", "Толщина кадмиевого покрытия должна быть в пределах"),
    ("The cadmium plate must be", "Кадмиевое покрытие должно иметь толщину"),
    ("The cadmium plate must overlap the chromium plate run out.", "Кадмиевое покрытие должно перекрывать зону схода хромового покрытия."),
    ("Bare metal not permitted.", "Оголенный металл не допускается."),
    ("The cadmium plate is optional on", "Кадмиевое покрытие допускается на"),
    ("Cadmium plate is optional on", "Кадмиевое покрытие допускается на"),
    ("No bare cadmium permitted.", "Оголенный кадмий не допускается."),
    ("Before installation of bushes:", "Перед установкой втулок:"),
    ("After installation of bushes:", "После установки втулок:"),
    ("sub-assembly", "подсборка"),
    ("spotfaces", "подрезки площадок"),
    ("spotface", "подрезка площадки"),
    ("hole for the Bowden cable", "отверстие под трос Боудена"),
    ("the split pin hole", "отверстие под шплинт"),
    ("the thread and undercut", "резьбу и подрез"),
    ("the two holes through the end", "два отверстия в торце"),
    ("the axial hole and chamfers", "осевое отверстие и фаски"),
    ("the screw threads and the face that touches the wheel bearings", "резьбу и поверхность, контактирующую с колесными подшипниками"),
    ("the lubrication fittings and their identification washers", "смазочные штуцеры и их идентификационные шайбы"),
    ("the bearings, bush bores and flanges", "подшипники, отверстия под втулки и фланцы"),
    ("the holes (with or without threads)", "отверстия (с резьбой или без резьбы)"),
    ("the lubrication fitting bores", "отверстия под смазочные штуцеры"),
    ("lubrication fitting bores", "отверстия под смазочные штуцеры"),
    ("lubrication adaptors", "смазочные адаптеры"),
    ("lubrication adapters", "смазочные адаптеры"),
    ("that have chromium plate", "с хромовым покрытием"),
    ("that has chromium plate", "с хромовым покрытием"),
    ("the chromium plated areas", "участки с хромовым покрытием"),
    ("chromium plated areas", "участки с хромовым покрытием"),
    ("to the spotfaces", "на подрезки площадок"),
    ("to spotfaces", "на подрезки площадок"),
    ("but not to the spotfaces", "но не на подрезки площадок"),
    ("but not the spotfaces", "кроме подрезок площадок"),
    ("but not to", "но не на"),
    ("but not on", "но не на"),
    ("but not in", "но не в"),
    ("only to", "только на"),
    ("only on", "только на"),
    ("to the areas", "на участки"),
    ("to areas", "на участки"),
    ("to the area", "на участок"),
    ("to area", "на участок"),
    ("to the holes", "в отверстия"),
    ("where identified on Figure", "где указано на рисунке"),
    ("externally", "снаружи"),
    ("internally", "изнутри"),
    ("bores", "отверстия"),
    ("areas", "участки"),
    ("area", "участок"),
    ("flanges", "фланцы"),
    ("radii", "радиусные участки"),
    ("chamfer", "фаска"),
    ("where shown", "где показано"),
    ("where the lubrication adaptors", "где смазочные адаптеры"),
    ("where the lubrication adapters", "где смазочные адаптеры"),
    ("along surface", "вдоль поверхности"),
    ("Paint finish is optional in areas", "Лакокрасочное покрытие допускается на участках"),
    ("will install", "будут устанавливаться"),
    ("install.", "устанавливаться."),
    ("Specification", "Спецификация"),
    ("Stop ring", "Стопорное кольцо"),
    ("Transfer dowel", "Переходной штифт"),
    ("Two piece stop", "Двухсоставной упор"),
    ("Valve support", "Опора клапана"),
    (" or ", " или "),
    ("Inclusion class", "Класс включений"),
    ("on areas without chromium plate", "на участках без хромового покрытия"),
    ("Chromium plated areas", "Участки с хромовым покрытием"),
    ("plated areas", "участки с покрытием"),
    ("Stainless Steel", "Коррозионностойкая сталь"),
    ("Aluminium Alloy", "Алюминиевый сплав"),
    ("Steel", "Сталь"),
    ("Not applicable.", "Не применяется."),
    ("To be given subsequently.", "Будет указано позднее."),
    ("Corrosion.", "Коррозия."),
    ("Type 1B", "тип 1B"),
    ("Type 2", "тип 2"),
    ("Class 1", "класс 1"),
    ("Type1", "тип 1"),
    ("Qty.", "Кол-во"),
    ("RUNOUT", "БИЕНИЕ"),
    ("DEEP", "ГЛУБ."),
    ("INTERNALLY", "ВНУТРИ"),
    ("THIS FACE ONLY", "ТОЛЬКО НА ЭТОЙ ПОВЕРХНОСТИ"),
    ("IN THIS FACE ONLY", "ТОЛЬКО НА ЭТОЙ ПОВЕРХНОСТИ"),
    ("FROM THIS SURFACE", "ОТ ЭТОЙ ПОВЕРХНОСТИ"),
    ("FROM OUTSIDE FACE", "ОТ НАРУЖНОЙ ПОВЕРХНОСТИ"),
    ("LIMIT OF SERMETEL W TERMINATION FROM CENTER", "ПРЕДЕЛ ПОКРЫТИЯ SERMETEL W ОТ ЦЕНТРА"),
]


PART5_PHRASE_RULES: list[tuple[str, str]] = [
    ("Repair Parts", "Ремонтные детали"),
    ("Repair Part", "Ремонтная деталь"),
    ("Repair sleeve (Qty 2)", "Ремонтная втулка (кол-во 2)"),
    ("Repair bush (Qty 2)", "Ремонтная втулка (кол-во 2)"),
    ("Oversize bearing", "ремонтный подшипник увеличенного размера"),
    ("Install the oversize bearing", "Установить ремонтный подшипник увеличенного размера"),
    ("Install oversize bearing", "Установить ремонтный подшипник увеличенного размера"),
    ("Cleaning tissues", "салфетки для очистки"),
    ("Cleaning agent", "очиститель"),
    ("Adhesive PVC tape", "клейкая лента ПВХ"),
    ("Masking tape", "малярная лента"),
    ("Emery cloth, 60-100 grit", "шлифовальная шкурка, зернистость 60-100"),
    ("Adhesive (Loctite Grade 601)", "Клей (Loctite, марка 601)"),
    ("Adhesive", "клей"),
    ("Sealant", "герметик"),
    ("DIAMETER MINIMUM", "МИН. ДИАМЕТР"),
    ("CENTERS TYPICAL", "МЕЖДУ ОСЯМИ, ТИПОВО"),
    ("TO M-DLPS1031-6", "ПО M-DLPS1031-6"),
    ("ground chromium plate", "отшлифованное хромовое покрытие"),
    ("micrometers", "мкм"),
    ("micro-inches", "мкдюймов"),
]

PHRASE_RULES.extend(PART5_PHRASE_RULES)


PART6_PHRASE_RULES: list[tuple[str, str]] = [
    (
        "CAUTION: DO NOT USE A MECHANICAL MOP POLISHER TO GET THE SURFACE FINISH.",
        "ПРЕДОСТЕРЕЖЕНИЕ. НЕ ИСПОЛЬЗУЙТЕ МЕХАНИЧЕСКУЮ ПОЛИРОВАЛЬНУЮ МАШИНУ ДЛЯ ПОЛУЧЕНИЯ ТРЕБУЕМОЙ ЧИСТОТЫ ПОВЕРХНОСТИ.",
    ),
    (
        "CAUTION: DO NOT MACHINE ALL OF THE FLANGE FACE.",
        "ПРЕДОСТЕРЕЖЕНИЕ. НЕ ОБРАБАТЫВАЙТЕ МЕХАНИЧЕСКИ ВСЮ ПОВЕРХНОСТЬ ФЛАНЦА.",
    ),
    (
        "THE CHROMIUM PLATE MUST NOT EXTEND BEYOND THE DIMENSIONS SHOWN",
        "ХРОМОВОЕ ПОКРЫТИЕ НЕ ДОЛЖНО ВЫХОДИТЬ ЗА ПОКАЗАННЫЕ РАЗМЕРЫ",
    ),
    (
        "MAX. CHROMIUM PLATE MUST TERMINATE IN THIS LENGTH.",
        "МАКС. ХРОМОВОЕ ПОКРЫТИЕ ДОЛЖНО ЗАКАНЧИВАТЬСЯ В ПРЕДЕЛАХ ЭТОЙ ДЛИНЫ.",
    ),
    (
        "CHROMIUM PLATE MUST TERMINATE WITHIN THIS LENGTH.",
        "ХРОМОВОЕ ПОКРЫТИЕ ДОЛЖНО ЗАКАНЧИВАТЬСЯ В ПРЕДЕЛАХ ЭТОЙ ДЛИНЫ.",
    ),
    ("CHROMIUM PLATE DEPOSIT", "Наплыв хромового покрытия"),
    ("CHROMIUM PLATING", "Хромовое покрытие"),
    ("SERMETAL COATING AND", "Покрытие SERMETEL и"),
    ("SERMETEL COATING", "Покрытие SERMETEL"),
    ("UNPLATED LENGTH", "Длина без покрытия"),
    ("LENGTH OF NICKEL PLATE", "Длина никелевого покрытия"),
    ("LENGTH OF CHROMIUM", "Длина хромового покрытия"),
    ("DO NOT CADMIUM PLATE", "Не наносить кадмиевое покрытие"),
    (
        "ALUMINIUM COAT OPTIONAL ON THESE SURFACES",
        "Алюминиевое покрытие допускается на этих поверхностях",
    ),
    ("APPLY COAT OF ALUMINIUM (IVD)", "Нанести покрытие из алюминия (IVD)"),
    ("IVD COATING OPTIONAL IN END FACE", "Покрытие IVD допускается на торце"),
    ("IVD COATING OPTIONAL", "Покрытие IVD допускается"),
    ("SMOOTH EDGE", "Сглаженная кромка"),
    ("EDGE SMOOTHED OUT", "Кромка сглажена"),
    ("EDGE SMOOTHED", "Кромка сглажена"),
    ("ENLARGED DETAIL", "Увеличенный фрагмент"),
    ("MINIMUM DIAMETER BEFORE CHROMIUM PLATE", "Мин. диаметр до хромирования"),
    ("MINIMUM WALL THICKNESS", "Мин. толщина стенки"),
    ("MINIMUM LUG WIDTH", "Мин. ширина проушины"),
    ("PAINT PERMITTED", "Краски допускаются"),
    ("TRACES OF", "Следы"),
    ("IF THE BASE METAL IS NOT DAMAGED", "Если основной металл не поврежден"),
    ("IF THE BASE METAL IS DAMAGED", "Если основной металл поврежден"),
    ("WAVY IRREGULAR LINE IS PERMISSIBLE.", "Допускается волнистая или неровная линия."),
    ("TERMINATION TO M-DLPS1031-1", "Граница по M-DLPS1031-1"),
    ("TERMINATION TO M-DLPS1031-5", "Граница по M-DLPS1031-5"),
    ("TERMINATION TO M-DLPS1031-7", "Граница по M-DLPS1031-7"),
    ("BEFORE CHROMIUM", "До хромирования"),
    ("AFTER CHROMIUM", "После хромирования"),
    ("DEGREES REF.", "ГРАДУСОВ СПРАВ."),
    ("DEGREES", "ГРАДУСОВ"),
    ("MINUTES", "МИНУТ"),
    ("RAD. REF.", "РАД. СПРАВ."),
    ("(REF)", "(СПРАВ.)"),
    ("REF.", "СПРАВ."),
    ("NOTE:", "ПРИМЕЧАНИЕ:"),
    ("DIM. H", "РАЗМ. H"),
    ("FACE B", "Поверхность B"),
    ("FACE C", "Поверхность C"),
    ("FACE D", "Поверхность D"),
    ("FACE E", "Поверхность E"),
    ("FACE R", "Поверхность R"),
    ("(as measured)", "(по результатам измерения)"),
    ("B = A (as measured) - 0,006 mm (0.0002 in) to + 0,028 mm (0.0011 in).", "B = A (по результатам измерения) - 0,006 мм (0.0002 дюйм) до + 0,028 мм (0.0011 дюйм)."),
    ("- 0,10 to + 0,10 mm (- 0.004 to + 0.004 in).", "- 0,10 до + 0,10 мм (- 0.004 до + 0.004 дюйм)."),
    ("77,80mm (3.063in) MINIMUM 20,25mm (0.797in)", "77,80мм (3.063дюйм) МИН. 20,25мм (0.797дюйм)"),
    ("(63) PLATE", "(63) покрытие"),
    ("(10) PLATE", "(10) покрытие"),
    ("Z PLATE", "Z покрытие"),
    ("FINISH PAINT", "Лакокрасочное покрытие"),
    ("Alignment bar", "Выверочная штанга"),
    ("Aluminium Bronze", "Алюминиевая бронза"),
    ("Bronze, UZ 19A6", "Бронза, UZ 19A6"),
    ("Bronze, UZ19A6", "Бронза, UZ19A6"),
    ("Zinc Powder", "Цинковый порошок"),
    ("Guide Bush", "Направляющая втулка"),
    ("Qty", "Кол-во"),
    ("Page ", "Стр. "),
    ("Apply ", "Нанесите "),
    ("Check the ", "Проверьте "),
    ("Calculate ", "Рассчитайте "),
    ("Do this procedure", "Выполните данную процедуру"),
    ("Examine ", "Осмотрите "),
    ("Finish grind ", "Выполните окончательное шлифование "),
    ("Grind ", "Отшлифуйте "),
    ("Grit blast ", "Выполните абразивоструйную обработку "),
    ("If necessary, ", "При необходимости "),
    ("Install ", "Установите "),
    ("Locally ", "Локально "),
    ("Machine ", "Обработайте "),
    ("Make sure that ", "Убедитесь, что "),
    ("Measure and record ", "Измерьте и запишите "),
    ("Refer to ", "См. "),
    ("Shot peen ", "Дробеструйно упрочните "),
    ("Stress relieve ", "Выполните снятие напряжений "),
    ("Use ", "Используйте "),
    ("Make the surface finish ", "Обеспечьте шероховатость поверхности "),
    ("The surface finish must be ", "Шероховатость поверхности должна быть "),
    ("The chromium plate thickness must be between ", "Толщина хромового покрытия должна быть "),
    ("The cadmium plate thickness must be between ", "Толщина кадмиевого покрытия должна быть "),
    ("The thickness of the coating must be between ", "Толщина покрытия должна быть "),
    ("The thickness of the IVD layer must be between ", "Толщина слоя IVD должна быть "),
    ("must be sufficient to give a minimum thickness of ", "должна обеспечивать минимальную толщину "),
    ("or better", "или лучше"),
    ("correctly masked", "правильно замаскирована"),
    ("correct angular position", "правильное угловое положение"),
    ("the bond is satisfactory", "сцепление удовлетворительное"),
    ("do the repair again", "повторите ремонт"),
    ("damage or corrosion", "повреждение или коррозию"),
    ("corrosion or damage", "коррозию или повреждение"),
    ("damage or wear", "повреждение или износ"),
    ("wear or damage", "износ или повреждение"),
    ("the damage or corrosion", "повреждение или коррозию"),
    ("the corrosion or damage", "коррозию или повреждение"),
    ("the damage or wear", "повреждение или износ"),
    ("the wear or damage", "износ или повреждение"),
    ("diameter(s)", "диаметр(ы)"),
    ("diameters", "диаметры"),
    ("diameter", "диаметр"),
    ("flange face(s)", "поверхность(и) фланца(ев)"),
    ("flange faces", "поверхности фланцев"),
    ("flange face", "поверхность фланца"),
    ("brake flange surface", "поверхность тормозного фланца"),
    ("brake flange(s)", "тормозной(ые) фланец(ы)"),
    ("brake flanges", "тормозные фланцы"),
    ("areas adjacent to the flanges", "участки, прилегающие к фланцам"),
    ("the areas adjacent to the flanges", "участки, прилегающие к фланцам"),
    ("the areas that do not have chromium plate", "участки без хромового покрытия"),
    ("machined areas", "обработанные участки"),
    ("machined area", "обработанный участок"),
    ("machined surfaces", "обработанные поверхности"),
    ("reworked areas", "восстановленные участки"),
    ("reworked area", "восстановленный участок"),
    ("repaired area", "отремонтированный участок"),
    ("affected area", "затронутый участок"),
    ("smoothed and polished areas", "сглаженные и отполированные участки"),
    ("base metal", "основной металл"),
    ("parent metal", "основной металл"),
    ("bare metal", "основной металл"),
    ("part number", "номер детали"),
    ("repair number", "номер ремонта"),
    ("band(s)", "полосу(ы)"),
    ("the band", "полосу"),
    ("the bands", "полосы"),
    ("bore diameters", "диаметры отверстий"),
    ("bore diameter", "диаметр отверстия"),
    ("the bore", "отверстие"),
    ("bores", "отверстия"),
    ("bore", "отверстие"),
    ("repair bush(es)", "ремонтную(ые) втулку(и)"),
    ("repair bushes", "ремонтные втулки"),
    ("repair bush", "ремонтную втулку"),
    ("guide bushes", "направляющие втулки"),
    ("bush flanges", "фланцы втулок"),
    ("repair bush flanges", "фланцы ремонтных втулок"),
    ("areas shown", "указанные участки"),
    ("the dimensions shown", "указанные размеры"),
    ("dimensions shown", "указанные размеры"),
    ("the dimensions", "размеры"),
    ("formula", "формулу"),
    ("surface treatment", "покрытие"),
    ("sulphamate nickel plate", "сульфаматно-никелевое покрытие"),
    ("chromium plate", "хромовое покрытие"),
    ("cadmium plate", "кадмиевое покрытие"),
    ("electrically conducting", "электропроводящий"),
    ("rubberised sealant", "прорезиненный герметик"),
    ("around the joints between", "по периметру стыков между"),
    ("to the joints between", "в стыки между"),
]

PHRASE_RULES.extend(PART6_PHRASE_RULES)


PART6_EXACT_MAP: dict[str, str] = {
    "Measure and record the new diameter(s) A.": "Измерьте и запишите новый диаметр(ы) A.",
    "Measure and record the new diameter A.": "Измерьте и запишите новый диаметр A.",
    "Apply Sermetel W protective treatment: refer to M-DLPS637 and REPAIR, Protective Treatment.": "Нанесите защитное покрытие Sermetel W: см. M-DLPS637 и раздел РЕМОНТ, Защитная обработка.",
    "Apply Sermetel to the sliding tube but not on the chromium plated or the cadmium plated areas: refer to REPAIR and M-DLPS637.": "Нанесите покрытие Sermetel на скользящую трубу, но не на участки с хромовым или кадмиевым покрытием: см. РЕМОНТ и M-DLPS637.",
    "Apply paint to the sliding tube but not on the chromium plated areas: refer to REPAIR and PCS-2500.": "Нанесите лакокрасочное покрытие на скользящую трубу, но не на участки с хромовым покрытием: см. РЕМОНТ и PCS-2500.",
    "Apply primer paint to the sliding tube where the bush flanges will touch: refer to PCS-2500.": "Нанесите грунтовочную краску на скользящую трубу в местах прилегания фланцев втулок: см. PCS-2500.",
    "Apply primer paint to the sliding tube where the repair bush flanges will touch: refer to PCS-2500.": "Нанесите грунтовочную краску на скользящую трубу в местах прилегания фланцев ремонтных втулок: см. PCS-2500.",
    "Examine the machined area for flaws: refer to PCS-3600 and PCS-3100, inclusion class 4.": "Осмотрите обработанный участок на наличие дефектов: см. PCS-3600 и PCS-3100, класс включений 4.",
    "Examine the machined areas for flaws: refer to PCS-3600 and PCS-3100, inclusion class 4.": "Осмотрите обработанные участки на наличие дефектов: см. PCS-3600 и PCS-3100, класс включений 4.",
    "Examine the reworked area for flaws: refer to PCS-3600 and PCS-3100, inclusion class 4.": "Осмотрите восстановленный участок на наличие дефектов: см. PCS-3600 и PCS-3100, класс включений 4.",
    "If the base metal is damaged or worn:": "Если основной металл поврежден или изношен:",
    "If the base metal is not damaged or worn:": "Если основной металл не поврежден и не изношен:",
    "If the base metal is not damaged or corroded:": "Если основной металл не поврежден и не поражен коррозией:",
    "If the bare metal is damaged or corroded:": "Если основной металл поврежден или поражен коррозией:",
    "Machine the chamfers and radii to the dimensions shown: refer to Figure 601.": "Обработайте фаски и радиусы по указанным размерам: см. рисунок 601.",
    "Machine the chamfer and radii to the dimensions shown: refer to Figure 601.": "Обработайте фаску и радиусы по указанным размерам: см. рисунок 601.",
    "Machine the chamfers and radii as shown: refer to Figure 602.": "Обработайте фаски и радиусы, как показано: см. рисунок 602.",
    "Machine the radii to the dimensions as shown: refer to Figure 601.": "Обработайте радиусы по указанным размерам: см. рисунок 601.",
    "Remove the Sermetel from the sliding tube, if necessary: refer to M-DLPS637.": "При необходимости удалите покрытие Sermetel со скользящей трубы: см. M-DLPS637.",
    "Remove the cadmium plate from the sliding tube, if necessary: refer to PCS-2100.": "При необходимости удалите кадмиевое покрытие со скользящей трубы: см. PCS-2100.",
    "Remove the paint from the sliding tube, if necessary: refer to PCS-2700 and Figure 601.": "При необходимости удалите лакокрасочное покрытие со скользящей трубы: см. PCS-2700 и рисунок 601.",
    "Remove the paint locally from the sliding tube: refer to PCS-2700.": "Локально удалите лакокрасочное покрытие со скользящей трубы: см. PCS-2700.",
    "Remove the sermetel from the sliding tube: refer to M-DLPS637.": "Удалите покрытие Sermetel со скользящей трубы: см. M-DLPS637.",
    "Remove the Sermetel layer from the sliding tube: refer to REPAIR and M-DLPS637.": "Локально удалите слой Sermetel со скользящей трубы: см. РЕМОНТ и M-DLPS637.",
    "Remove the cadmium plate locally from the sliding tube: refer to PCS-2100.": "Локально удалите кадмиевое покрытие со скользящей трубы: см. PCS-2100.",
    "Remove the chromium plate from sliding tube: refer to PCS-2110, type C.": "Удалите хромовое покрытие со скользящей трубы: см. PCS-2110, тип C.",
    "Remove the chromium plate locally from the sliding tube: refer to PCS-2110, type C and Figure 601.": "Локально удалите хромовое покрытие со скользящей трубы: см. PCS-2110, тип C, и рисунок 601.",
    "Remove the chromium plate from diameter(s) A: refer to PCS-2110 Type C.": "Удалите хромовое покрытие с диаметра(ов) A: см. PCS-2110, тип C.",
    "Remove the chromium plate from diameter(s) B: refer to PCS-2110 Type C and Figure 601.": "Удалите хромовое покрытие с диаметра(ов) B: см. PCS-2110, тип C, и рисунок 601.",
    "Remove the chromium plate from diameter(s) C: refer to PCS-2110 Type C and Figure 602.": "Удалите хромовое покрытие с диаметра(ов) C: см. PCS-2110, тип C, и рисунок 602.",
    "Remove the chromium plate from diameter A: refer to PCS-2110.": "Удалите хромовое покрытие с диаметра A: см. PCS-2110.",
    "Remove the chromium plate from diameter A: refer to PCS-2110, type C.": "Удалите хромовое покрытие с диаметра A: см. PCS-2110, тип C.",
    "Remove the chromium plate band(s) from the brake flange(s): refer to M-DLPS101.": "Удалите полосы хромового покрытия с тормозного(ых) фланца(ев): см. M-DLPS101.",
    "Remove the Sermetel W surface treatment from the sliding tube: refer to M-DLPS637.": "Удалите покрытие Sermetel W со скользящей трубы: см. M-DLPS637.",
    "Shot peen the machined area: refer to M-DLPS123.": "Дробеструйно упрочните обработанный участок: см. M-DLPS123.",
    "Shot peen the machined areas: refer to M-DLPS123.": "Дробеструйно упрочните обработанные участки: см. M-DLPS123.",
    "Shot peen the machined areas: refer to PCS-2300.": "Дробеструйно упрочните обработанные участки: см. PCS-2300.",
    "Shot peen the reworked areas: refer to M-DLPS123.": "Дробеструйно упрочните восстановленные участки: см. M-DLPS123.",
    "Stress relieve the sliding tube for 4 hours at 185 to 195 oC (366 to 384 oF).": "Выполните снятие напряжений скользящей трубы в течение 4 ч при 185-195 oC (366-384 oF).",
    "De-embrittle the sliding tube for 23 hours at 185 to 195oC (366 to 384oF).": "Проведите обезводораживание скользящей трубы в течение 23 ч при 185-195oC (366-384oF).",
    "De-embrittle the sliding tube for 4 hours at 185 to 195oC (366 to 384oF).": "Проведите обезводораживание скользящей трубы в течение 4 ч при 185-195oC (366-384oF).",
    "Apply cadmium plate to the areas that do not have chromium plate: refer to PCS-2141.": "Нанесите кадмиевое покрытие на участки без хромового покрытия: см. PCS-2141.",
    "Apply chromium plate to diameter A: refer to PCS-2110 Type C. Refer to Figure 601 for chromium plate termination information.": "Нанесите хромовое покрытие на диаметр A: см. PCS-2110, тип C. Сведения о границе хромового покрытия см. на рисунке 601.",
    "Apply chromium plate to diameter A: refer to PCS-2110 Type C. Refer to M-DLPS1031 and Figure 602 for chromium plate termination information.": "Нанесите хромовое покрытие на диаметр A: см. PCS-2110, тип C. Сведения о границе хромового покрытия см. в M-DLPS1031 и на рисунке 602.",
    "Apply chromium plate to diameter A: refer to PCS-2110, type C: refer to Figure 601.": "Нанесите хромовое покрытие на диаметр A: см. PCS-2110, тип C, и рисунок 601.",
    "Apply chromium plate to the diameter A: refer to PCS-2110, type C and Figure 601. Make the surface finish 1,6 micrometers (64 micro-inches). The chromium plate thickness must be between 0,020 and 0,025 mm (0.0008 and 0.0010 in).": "Нанесите хромовое покрытие на диаметр A: см. PCS-2110, тип C, и рисунок 601. Обеспечьте шероховатость поверхности 1,6 мкм (64 мкдюймов). Толщина хромового покрытия должна быть 0,020-0,025 мм (0.0008-0.0010 in).",
    "Apply chromium plate to the diameter A: refer to PCS-2110 type C and Figure 601. Make the chromium plate thickness between 0,020 and 0,025 mm (0.0008 and 0.0010 in). Make the surface finish 1,6 micrometers (63 micro-inches).": "Нанесите хромовое покрытие на диаметр A: см. PCS-2110, тип C, и рисунок 601. Толщина хромового покрытия должна быть 0,020-0,025 мм (0.0008-0.0010 in). Обеспечьте шероховатость поверхности 1,6 мкм (63 мкдюймов).",
    "Apply chromium plate to diameter(s) A: refer to PCS-2110 Type C, M-DLPS1031-1, M-DLPS1031-6 and Figure 601. The chromium plate thickness must be sufficient to give a minimum thickness of 0,10 mm (0.004 in) after grinding.": "Нанесите хромовое покрытие на диаметр(ы) A: см. PCS-2110, тип C, M-DLPS1031-1, M-DLPS1031-6 и рисунок 601. Толщина хромового покрытия должна обеспечивать минимальную толщину 0,10 мм (0.004 in) после шлифования.",
    "Apply chromium plate to diameter(s) B: refer to PCS-2110 Type C, M-DLPS1031-1, M-DLPS1031-6 and Figure 601. The chromium plate thickness must be sufficient to give a minimum thickness of 0,10 mm (0.004 in) after grinding.": "Нанесите хромовое покрытие на диаметр(ы) B: см. PCS-2110, тип C, M-DLPS1031-1, M-DLPS1031-6 и рисунок 601. Толщина хромового покрытия должна обеспечивать минимальную толщину 0,10 мм (0.004 in) после шлифования.",
    "Apply chromium plate to diameter(s) C: refer to PCS-2110 Type C, M-DLPS1031-6, Figure 601 and Figure 602. The chromium plate thickness must be sufficient to give a minimum thickness of 0,10 mm (0.004 in) after grinding.": "Нанесите хромовое покрытие на диаметр(ы) C: см. PCS-2110, тип C, M-DLPS1031-6, рисунок 601 и рисунок 602. Толщина хромового покрытия должна обеспечивать минимальную толщину 0,10 мм (0.004 in) после шлифования.",
    "Apply chromium plate to diameter(s) C: refer to PCS-2110 Type C, M-DLPS1031-2, M-DLPS1031-6 and Figure 602. The chromium plate thickness must be sufficient to give a minimum thickness of 0,10 mm (0.004 in) after grinding.": "Нанесите хромовое покрытие на диаметр(ы) C: см. PCS-2110, тип C, M-DLPS1031-2, M-DLPS1031-6 и рисунок 602. Толщина хромового покрытия должна обеспечивать минимальную толщину 0,10 мм (0.004 in) после шлифования.",
    "Apply chromium plate to diameter(s) A: refer to PCS-2110 Type C, M-DLPS1031-5 and Figure 602. The chromium plate thickness must be between 0,020 and 0,025 mm (0.0008 and 0.0010 in).": "Нанесите хромовое покрытие на диаметр(ы) A: см. PCS-2110, тип C, M-DLPS1031-5 и рисунок 602. Толщина хромового покрытия должна быть 0,020-0,025 мм (0.0008-0.0010 in).",
    "Apply chromium plate to make new band(s) on the flange face(s) using one of these procedures: refer to Figure 601.": "Нанесите хромовое покрытие для образования новых полос на поверхности(ях) фланца(ев), используя одну из следующих процедур: см. рисунок 601.",
    "Apply chromium plate bands to the flange(s): refer to M-DLPS101-7. Grind the chromium plate to get the dimensions shown in Figure 601: the surface finish must be 2,5 micrometers (100 micro-inches). Examine the ground chromium plate for flaws: refer to M-DLNDT3.": "Нанесите полосы хромового покрытия на фланец(ы): см. M-DLPS101-7. Отшлифуйте хромовое покрытие до размеров, указанных на рисунке 601; шероховатость поверхности должна быть 2,5 мкм (100 мкдюймов). Осмотрите отшлифованное хромовое покрытие на наличие дефектов: см. M-DLNDT3.",
    "Apply chromium plate bands to the flange(s) to get a deposit that agrees with the dimensions shown in Figure 601: refer to M-DLPS101-7. The surface finish must be 2,5 micrometers (100 micro-inches).": "Нанесите полосы хромового покрытия на фланец(ы) так, чтобы наплыв соответствовал размерам, указанным на рисунке 601: см. M-DLPS101-7. Шероховатость поверхности должна быть 2,5 мкм (100 мкдюймов).",
    "Apply chromium plate to the bands: refer to M-DLPS101-7. Grind the chromium plate to get the dimensions shown: the surface finish must be 2,5 micrometers (100 micro-inches). Examine the ground chromium plate for flaws: refer to M-DLNDT3.": "Нанесите хромовое покрытие на полосы: см. M-DLPS101-7. Отшлифуйте хромовое покрытие до указанных размеров; шероховатость поверхности должна быть 2,5 мкм (100 мкдюймов). Осмотрите отшлифованное хромовое покрытие на наличие дефектов: см. M-DLNDT3.",
    "Apply chromium plate to the band to the dimensions shown: refer to M-DLPS101-7. The surface finish must be 2,5 micrometers (100 micro-inches).": "Нанесите хромовое покрытие на полосу по указанным размерам: см. M-DLPS101-7. Шероховатость поверхности должна быть 2,5 мкм (100 мкдюймов).",
    "Refer to Figure 601 and M-DLPS101-7. Apply chromium plate to make new bands on the brake flange surface: make the chromium plate thicker in the areas where the corrosion or damage has been removed.": "См. рисунок 601 и M-DLPS101-7. Нанесите хромовое покрытие для образования новых полос на поверхности тормозного фланца; в местах удаления коррозии или повреждений хромовое покрытие должно быть толще.",
    "Refer to M-DLPS200-4 and locally grit blast the sliding tube: give protection from the grit blast to the areas adjacent to the flanges.": "См. M-DLPS200-4 и выполните локальную абразивоструйную обработку скользящей трубы; защитите прилегающие к фланцам участки от абразивной струи.",
    "Apply IVD alloy coating with chromation yellow all over the bushes, but not to the bores: refer to ICT 40-893-01MD and Figure 602. The thickness of the coating must be between 0,007 and 0,012 mm (0.0003 and 0.0005 in).": "Нанесите покрытие из сплава IVD с желтым хроматированием на втулки по всей поверхности, кроме отверстий: см. ICT 40-893-01MD и рисунок 602. Толщина покрытия должна быть 0,007-0,012 мм (0.0003-0.0005 in).",
    "Apply a layer of IVD alloy with elecromation yellow all over the repair bush(es) but not in the bores: refer to ICT 40-893-01MD and Figure 603. The thickness of the IVD layer must be between 0,0075 and 0,0125 mm (0.0003 and 0.0005 in).": "Нанесите слой сплава IVD с желтым хроматированием на ремонтную(ые) втулку(и) по всей поверхности, кроме отверстий: см. ICT 40-893-01MD и рисунок 603. Толщина слоя IVD должна быть 0,0075-0,0125 мм (0.0003-0.0005 in).",
    "Use the Press Pad 460004330/105 and install the repair bushes: refer to M-DLPS1011-20.": "Используйте нажимную опору 460004330/105 и установите ремонтные втулки: см. M-DLPS1011-20.",
    "Use Press Pad 460004330/117 and install the repair bushes: refer to M-DLPS1011-20. Use electrically conducting Mastinox (made from Mastinox D40, Material Ref. Item 05-533 and Zinc powder, Material Ref. Item TBA): refer to M-DLPS709-14.": "Используйте нажимную опору 460004330/117 и установите ремонтные втулки: см. M-DLPS1011-20. Используйте электропроводящий состав Mastinox (из Mastinox D40, код ссылки материала 05-533, и цинкового порошка, код ссылки материала TBA): см. M-DLPS709-14.",
    "Use the Press Pad 460004330/131 and install the repair bush: refer to M-DLPS1011-14. Use zinc loaded Molykote (made from Molykote 111, Material Ref. Item 04-512 and Zinc powder, Material Ref. Item TBA): refer to PCS-7304. The slot in the repair bush must be in the correct angular position: refer to Figure 602.": "Используйте нажимную опору 460004330/131 и установите ремонтную втулку: см. M-DLPS1011-14. Используйте Molykote, насыщенный цинком (из Molykote 111, код ссылки материала 04-512, и цинкового порошка, код ссылки материала TBA): см. PCS-7304. Паз в ремонтной втулке должен находиться в правильном угловом положении: см. рисунок 602.",
    "Use the alignment bar 460006246, the press pad assembly 460006250 and the guide bushes 460006251 and 460006252 and install the oversize bushes: refer to PCS-5105-2.": "Используйте выверочную штангу 460006246, сборку нажимной опоры 460006250 и направляющие втулки 460006251 и 460006252 и установите ремонтные втулки увеличенного размера: см. PCS-5105-2.",
    "Machine the face B to produce a flange thickness of 2,00 to 2,05 mm (0.079 to 0.080 in): refer to Figure 602. Make the surface finish 1,6 micrometers (63 micro-inches).": "Обработайте поверхность B так, чтобы получить толщину фланца 2,00-2,05 мм (0.079-0.080 in): см. рисунок 602. Обеспечьте шероховатость поверхности 1,6 мкм (63 мкдюймов).",
    "Machine the repair bushes to the dimensions shown and calculated: refer to Figure 602. Make the surface finish 1,6 micrometers (63 micro-inches).": "Обработайте ремонтные втулки до указанных и расчетных размеров: см. рисунок 602. Обеспечьте шероховатость поверхности 1,6 мкм (63 мкдюймов).",
    "Machine the repair bushes to the dimensions shown and calculated. Machine face C to get the correct dimensions after installation: refer to Figure 602. Make the surface finish 1,6 micrometers (63 micro-inches). The bushes must not protrude through the lug after installation.": "Обработайте ремонтные втулки до указанных и расчетных размеров. Обработайте поверхность C для получения правильных размеров после установки: см. рисунок 602. Обеспечьте шероховатость поверхности 1,6 мкм (63 мкдюймов). После установки втулки не должны выступать через проушину.",
    "Machine the face R to get the flange thickness of between 2,00 and 2,05 mm (0.079 and 0.081 in).": "Обработайте поверхность R так, чтобы получить толщину фланца 2,00-2,05 мм (0.079-0.081 in).",
    "Apply sealant, Material Ref. Item 09-510A, around the joints between the repair bushes and the sliding tube: refer to PCS-7200 and Figure 602.": "Нанесите герметик, код ссылки материала 09-510A, по периметру стыков между ремонтными втулками и скользящей трубой: см. PCS-7200 и рисунок 602.",
    "Apply a bead of sealant, Material Ref. Item 09-510A, to the joints between the repair bush and the sliding tube: refer to PCS-7200 and Figure 602.": "Нанесите валик герметика, код ссылки материала 09-510A, в стыки между ремонтной втулкой и скользящей трубой: см. PCS-7200 и рисунок 602.",
    "Apply Sulphamate Nickel plate to diameter A: refer to PCS-2120 for sequence of operations and to MIL-STD-868A solution 2 and Figure 601.": "Нанесите сульфаматно-никелевое покрытие на диаметр A: см. PCS-2120 по последовательности операций, а также MIL-STD-868A, раствор 2, и рисунок 601.",
    "Machine the sulphamate nickel plate (do not grind), the machined surface must be 0,00 to 0,05 mm (0.000 to 0.002 in) above the brake flange surface. The surface finish must be 1,6 micrometers (63 micro-inches).": "Обработайте сульфаматно-никелевое покрытие резанием (не шлифовать); обработанная поверхность должна выступать над поверхностью тормозного фланца на 0,00-0,05 мм (0.000-0.002 in). Шероховатость поверхности должна быть 1,6 мкм (63 мкдюймов).",
    "Use 5 or 10 times magnification to examine the sulphamate nickel plate to make sure the bond is satisfactory: do the repair again if the bond is not satisfactory.": "Используя увеличение 5x или 10x, осмотрите сульфаматно-никелевое покрытие, чтобы убедиться в удовлетворительном сцеплении; если сцепление неудовлетворительное, повторите ремонт.",
    "This Repair, Messier-Dowty Limited Repair No. 450258400, has been superseded by Repair No. 9-10, Messier-Dowty Limited Repair No. 450258401.": "Настоящий ремонт, ремонт Messier-Dowty Limited № 450258400, заменен ремонтом № 9-10, ремонт Messier-Dowty Limited № 450258401.",
    "Only required, if the parent metal was machined at step (g), examine the machined area for flaws: refer to PCS-3600.": "Только если на этапе (g) выполнялась механическая обработка основного металла, осмотрите обработанный участок на наличие дефектов: см. PCS-3600.",
    "Only required, if the parent metal was machined at step (g), examine the machined area for flaws: refer to PCS-3100, inclusion class 4.": "Только если на этапе (g) выполнялась механическая обработка основного металла, осмотрите обработанный участок на наличие дефектов: см. PCS-3100, класс включений 4.",
    "Use the repair number 450258401A if there was no damage to the base metal or": "Используйте номер ремонта 450258401A, если основной металл не был поврежден, или",
    "Use the repair number 450258401B if there was damage to the base metal.": "Используйте номер ремонта 450258401B, если основной металл был поврежден.",
    "Use the repair number 450258401C if there was no damage to the base metal or": "Используйте номер ремонта 450258401C, если основной металл не был поврежден, или",
    "Use the repair number 450258401D if there was damage to the base metal.": "Используйте номер ремонта 450258401D, если основной металл был поврежден.",
    "Use the repair number 450258401E if there was no damage to the base metal or": "Используйте номер ремонта 450258401E, если основной металл не был поврежден, или",
    "Use the repair number 450258401F if there was damage to the base metal.": "Используйте номер ремонта 450258401F, если основной металл был поврежден.",
    "Material Spcification": "Спецификация материала",
    "Install oversize bushes": "Установите ремонтные втулки увеличенного размера",
    "Install the repair bush": "Установите ремонтную втулку",
    "Install the repair bush(es)": "Установите ремонтную(ые) втулку(и)",
    "Install the repair bushes": "Установите ремонтные втулки",
    "Tool Part No.": "Номер детали инструмента",
    "Repair Bush (Qty 18)": "Ремонтная втулка (Кол-во 18)",
    "To install the repair bush 450266800": "Для установки ремонтной втулки 450266800",
}

PART6_EXACT_MAP_2: dict[str, str] = {
    "CAUTION: REPAIR WILL NOT BE PERMITTED BEYOND THE LIMITS OF THIS REPAIR SCHEME.": "ПРЕДУПРЕЖДЕНИЕ. Ремонт за пределами данной схемы ремонта не допускается.",
    "CAUTION: FOR DAMAGE MORE THAN THE LIMITS OF THIS REPAIR SCHEME, WRITE TO SAFRAN LANDING SYSTEMS: REFER TO GUIDE-CS-001.": "ПРЕДУПРЕЖДЕНИЕ. При повреждении, превышающем пределы данной схемы ремонта, обратитесь в Safran Landing Systems: см. GUIDE-CS-001.",
    "CAUTION: FOR DEVIATIONS OUTSIDE THE LIMITS OF THIS REPAIR SCHEME CONTACT SAFRAN LANDING SYSTEMS.": "ПРЕДУПРЕЖДЕНИЕ. При отклонениях за пределами данной схемы ремонта обратитесь в Safran Landing Systems.",
    "TERMINATE IN THIS LENGTH. WAVY ИЛИ IRREGULAR LINE IS PERMITTED": "ЗАКАНЧИВАЕТСЯ В ПРЕДЕЛАХ ЭТОЙ ДЛИНЫ. ДОПУСКАЕТСЯ ВОЛНИСТАЯ ИЛИ НЕРОВНАЯ ЛИНИЯ",
    "ДИАМЕТР ДО ХРОМИРОВАНИЯ PLATE": "ДИАМЕТР ДО ХРОМИРОВАНИЯ",
    "ДИАМЕТР AFTER GRINDING OF ХРОМОВОЕ ПОКРЫТИЕ": "ДИАМЕТР ПОСЛЕ ШЛИФОВАНИЯ ХРОМОВОГО ПОКРЫТИЯ",
    "MAXIMUM ХРОМОВОЕ ПОКРЫТИЕ MUST TERMINATE IN THIS LENGTH. WAVY ИЛИ IRREGULAR LINE IS": "МАКСИМАЛЬНОЕ ХРОМОВОЕ ПОКРЫТИЕ ДОЛЖНО ЗАКАНЧИВАТЬСЯ В ПРЕДЕЛАХ ЭТОЙ ДЛИНЫ. ВОЛНИСТАЯ ИЛИ НЕРОВНАЯ ЛИНИЯ",
    "PERMITTED": "ДОПУСКАЕТСЯ",
    "ПРОТЯЖЕННОСТЬ OF SHOT PEENING": "ПРОТЯЖЕННОСТЬ ДРОБЕСТРУЙНОГО УПРОЧНЕНИЯ",
    "2,00mm (0.079in) МАКС. ХРОМОВОЕ ПОКРЫТИЕ ДОЛЖНО ЗАКАНЧИВАТЬСЯ В ПРЕДЕЛАХ ЭТОЙ ДЛИНЫ. WAVY ИЛИ IRREGULAR": "2,00mm (0.079in) МАКС. ХРОМОВОЕ ПОКРЫТИЕ ДОЛЖНО ЗАКАНЧИВАТЬСЯ В ПРЕДЕЛАХ ЭТОЙ ДЛИНЫ. ВОЛНИСТАЯ ИЛИ НЕРОВНАЯ",
    "LINE IS PERMITTED": "ЛИНИЯ ДОПУСКАЕТСЯ",
    "WAVY ИЛИ IRREGULAR LINE IS PERMITTED": "ВОЛНИСТАЯ ИЛИ НЕРОВНАЯ ЛИНИЯ ДОПУСКАЕТСЯ",
    "RADIUS": "РАДИУС",
    "(0.060-0.079in) РАДИУС ДО ХРОМИРОВАНИЯ PLATE": "(0.060-0.079in) РАДИУС ДО ХРОМИРОВАНИЯ",
    "177,68-177,72mm (6.995-6.997in) ДИАМЕТР AFTER GRINDING ХРОМОВОЕ ПОКРЫТИЕ": "177,68-177,72mm (6.995-6.997in) ДИАМЕТР ПОСЛЕ ШЛИФОВАНИЯ ХРОМОВОГО ПОКРЫТИЯ",
    "(0.0008-0.0018in) THICKNESS OF ХРОМОВОЕ ПОКРЫТИЕ ABOVE OUTER SURFACE OF FLANGE": "(0.0008-0.0018in) ТОЛЩИНА ХРОМОВОГО ПОКРЫТИЯ НАД НАРУЖНОЙ ПОВЕРХНОСТЬЮ ФЛАНЦА",
    "(0-0.078in) CHROMIUM БИЕНИЕ": "(0-0.078in) БИЕНИЕ ХРОМОВОГО ПОКРЫТИЯ",
    "DEGREES": "ГРАДУСОВ",
    "ГРУНТОВОЧНАЯ КРАСКА WHERE ВТУЛКА FLANGE WILL TOUCH": "ГРУНТОВОЧНАЯ КРАСКА В МЕСТЕ КАСАНИЯ ФЛАНЦА ВТУЛКИ",
    "( WITHOUT РЕМОНТНЫЕ ВТУЛКИ)": "(БЕЗ РЕМОНТНЫХ ВТУЛОК)",
    "ДИАМЕТР ПОСЛЕ ХРОМИРОВАНИЯ PLATE 135,860-135,923mm": "ДИАМЕТР ПОСЛЕ ХРОМИРОВАНИЯ 135,860-135,923mm",
    "ДИАМЕТР AFTER МЕХАНИЧЕСКАЯ ОБРАБОТКА NICKEL PLATE": "ДИАМЕТР ПОСЛЕ МЕХАНИЧЕСКОЙ ОБРАБОТКИ НИКЕЛЕВОГО ПОКРЫТИЯ",
    "ДИАМЕТР BEFORE NICKEL И ХРОМОВОЕ ПОКРЫТИЕ 136,410mm (5.3700in) МАКСИМУМ": "ДИАМЕТР ДО НИКЕЛЕВОГО И ХРОМОВОГО ПОКРЫТИЯ 136,410mm (5.3700in) МАКСИМУМ",
    "(WITHOUT ВТУЛКИ)": "(БЕЗ ВТУЛКИ)",
    "CAUTION : FOR DEVIATIONS OUTSIDE THE LIMITS OF THIS РЕМОНТ SCHEME CONTACT SAFRAN LANDING SYSTEMS.": "ПРЕДУПРЕЖДЕНИЕ. При отклонениях за пределами данной схемы ремонта обратитесь в Safran Landing Systems.",
    "1,00mm (0.039in) ФАСКА (2 PLACES)": "1,00mm (0.039in) ФАСКА (2 МЕСТА)",
    "30 ГРАДУСОВ ФАСКА (2 PLACES)": "30 ГРАДУСОВ ФАСКА (2 МЕСТА)",
    "(0.551-0.630in) НАНЕСИТЕ ELECTRICALLY": "(0.551-0.630in) НАНЕСИТЕ ЭЛЕКТРОПРОВОДЯЩИЙ",
    "CONDUCTING MOLYKOTE 111 ИЛИ ПРОРЕЗИНЕННЫЙ ГЕРМЕТИК": "ЭЛЕКТРОПРОВОДЯЩИЙ MOLYKOTE 111 ИЛИ ПРОРЕЗИНЕННЫЙ ГЕРМЕТИК",
    "(0.472-0.551in) НАНЕСИТЕ ELECTRICALLY": "(0.472-0.551in) НАНЕСИТЕ ЭЛЕКТРОПРОВОДЯЩИЙ",
    "CONDUCTING MOLYKOTE 111 ИЛИ ПРОРЕЗИНЕННЫЙ ГЕРМЕТИК ON BOTH ВТУЛКИ TO": "ЭЛЕКТРОПРОВОДЯЩИЙ MOLYKOTE 111 ИЛИ ПРОРЕЗИНЕННЫЙ ГЕРМЕТИК НА ОБЕ ВТУЛКИ ДО",
    "WITHOUT OVERFLOW ON FACES OF ВТУЛКИ": "БЕЗ ИЗБЫТКА НА ПОВЕРХНОСТЯХ ВТУЛОК",
    "(0.098-0.138in) x 60 ГРАДУСОВ INCLUSIVE": "(0.098-0.138in) x 60 ГРАДУСОВ ВКЛЮЧИТЕЛЬНО",
    "(WITHOUT ВТУЛКА)": "(БЕЗ ВТУЛКИ)",
    "-3-3 ГРАДУСОВ ORIENTATION IS IMPORTANT": "-3-3 ГРАДУСОВ ОРИЕНТАЦИЯ ВАЖНА",
    "x 60 ГРАДУСОВ INCLUSIVE ФАСКА": "x 60 ГРАДУСОВ ВКЛЮЧИТЕЛЬНО ФАСКА",
    "СКОЛЬЗЯЩАЯ ТРУБА OUTSIDE ДИАМЕТР": "СКОЛЬЗЯЩАЯ ТРУБА НАРУЖНЫЙ ДИАМЕТР",
    "TERMINATION ПО M-DLPS1031-6": "ГРАНИЦА ПО M-DLPS1031-6",
    "ЗАВЕРШЕНИЕ ХРОМОВОГО ПОКРЫТИЯ ТИПОВО AROUND 4 SLOTS. IRREGULAR LINE IS PERMITTED": "ЗАВЕРШЕНИЕ ХРОМОВОГО ПОКРЫТИЯ ТИПОВО ВОКРУГ 4 ПАЗОВ. НЕРОВНАЯ ЛИНИЯ ДОПУСКАЕТСЯ",
    "ГРАНИЦА ПО M-DLPS1031-6. IRREGULAR LINE IS PERMITTED": "ГРАНИЦА ПО M-DLPS1031-6. НЕРОВНАЯ ЛИНИЯ ДОПУСКАЕТСЯ",
    "0,0-2,0mm (0.00-0.08in) ХРОМОВОЕ ПОКРЫТИЕ MUST TERMINATE WITHIN THIS LENGTH.": "0,0-2,0mm (0.00-0.08in) ХРОМОВОЕ ПОКРЫТИЕ ДОЛЖНО ЗАКАНЧИВАТЬСЯ В ПРЕДЕЛАХ ЭТОЙ ДЛИНЫ.",
    "WAVY ИЛИ IRREGULAR LINE IS PERMITTED.": "ВОЛНИСТАЯ ИЛИ НЕРОВНАЯ ЛИНИЯ ДОПУСКАЕТСЯ.",
    "1,0-3,00mm (0.04-0.118in) ХРОМОВОЕ ПОКРЫТИЕ ДОЛЖНО ЗАКАНЧИВАТЬСЯ В ПРЕДЕЛАХ ЭТОЙ ДЛИНЫ. WAVY ИЛИ IRREGULAR LINE IS PERMITTED.": "1,0-3,00mm (0.04-0.118in) ХРОМОВОЕ ПОКРЫТИЕ ДОЛЖНО ЗАКАНЧИВАТЬСЯ В ПРЕДЕЛАХ ЭТОЙ ДЛИНЫ. ВОЛНИСТАЯ ИЛИ НЕРОВНАЯ ЛИНИЯ ДОПУСКАЕТСЯ.",
    "ДИАМЕТР ДО ХРОМИРОВАНИЯ PLATE ДИАМ. C": "ДИАМЕТР ДО ХРОМИРОВАНИЯ ДИАМ. C",
    "131,294-131,357mm (5.1690-5.1715in) AFTER GRINDING": "131,294-131,357mm (5.1690-5.1715in) ПОСЛЕ ШЛИФОВАНИЯ",
    "WAVY ИЛИ IRREGULAR LINE IS PERMISSIBLE.": "ВОЛНИСТАЯ ИЛИ НЕРОВНАЯ ЛИНИЯ ДОПУСКАЕТСЯ.",
    "130,683mm (5.145in) MINIMUM ДИАМЕТР": "130,683mm (5.145in) МИНИМАЛЬНЫЙ ДИАМЕТР",
    "131,294-131,357mm (5.1690-5.1715in) AFTER GRINDING НАПЛЫВ ХРОМОВОГО ПОКРЫТИЯ": "131,294-131,357mm (5.1690-5.1715in) ПОСЛЕ ШЛИФОВАНИЯ НАПЛЫВА ХРОМОВОГО ПОКРЫТИЯ",
    "ДИАМ. A 6 HOLES МАКСИМУМ": "ДИАМ. A 6 ОТВЕРСТИЙ МАКСИМУМ",
    "0,0-0,8mm (0.00-0.03in) ХРОМОВОЕ ПОКРЫТИЕ TO TERMINATE OVER THIS LENGTH": "0,0-0,8mm (0.00-0.03in) ХРОМОВОЕ ПОКРЫТИЕ ДОЛЖНО ЗАКАНЧИВАТЬСЯ В ПРЕДЕЛАХ ЭТОЙ ДЛИНЫ",
    "0,50-2,00mm (0.020-0.080in) ХРОМОВОЕ ПОКРЫТИЕ TO TERMINATE OVER THIS LENGTH": "0,50-2,00mm (0.020-0.080in) ХРОМОВОЕ ПОКРЫТИЕ ДОЛЖНО ЗАКАНЧИВАТЬСЯ В ПРЕДЕЛАХ ЭТОЙ ДЛИНЫ",
    "0,5mm (0.020in) x 60 ГРАДУСОВ INCLUSIVE ФАСКА": "0,5mm (0.020in) x 60 ГРАДУСОВ ВКЛЮЧИТЕЛЬНО ФАСКА",
    "Apply chromium plate to diameter A: refer to PCS-2110, type C and Figure 601.": "Нанесите хромовое покрытие на диаметр A: см. PCS-2110, тип C, и рисунок 601.",
    "Apply cadmium plate to the sliding tube but not on the chromium plated areas: refer to REPAIR and PCS-2100. Make the cadmium plate thickness between 0,010 and 0,020 mm (0.0004 and 0.0008 in).": "Нанесите кадмиевое покрытие на скользящую трубу, но не на участки с хромовым покрытием: см. РЕМОНТ и PCS-2100. Толщина кадмиевого покрытия должна быть 0,010-0,020 мм (0.0004-0.0008 in).",
    "Apply paint locally to the sliding tube: refer to REPAIR and PCS-2500.": "Локально нанесите лакокрасочное покрытие на скользящую трубу: см. РЕМОНТ и PCS-2500.",
    "Grind diameter A to between 177,68 and 177,72 mm (6.995 and 6.997 in). The surface finish must be 0,25 micrometers (10 micro-inches).": "Отшлифуйте диаметр A до 177,68-177,72 мм (6.995-6.997 in). Шероховатость поверхности должна быть 0,25 мкм (10 мкдюймов).",
    "Grind diameter A to between 30,950 and 30,975 mm (1.2185 and 1.2194 in). The surface finish must be 0,8 micrometers (31 micro-inches).": "Отшлифуйте диаметр A до 30,950-30,975 мм (1.2185-1.2194 in). Шероховатость поверхности должна быть 0,8 мкм (31 мкдюймов).",
    "Apply Sermetel to the repaired area: refer to M-DLPS637 and Figure 602.": "Нанесите покрытие Sermetel на отремонтированный участок: см. M-DLPS637 и рисунок 602.",
    "Apply cadmium plate all over to the repair bush: refer to PCS-2101. The cadmium plate thickness must be between 0,005 and 0,009 mm (0.0002 and 0.0004 in).": "Нанесите кадмиевое покрытие на ремонтную втулку по всей поверхности: см. PCS-2101. Толщина кадмиевого покрытия должна быть 0,005-0,009 мм (0.0002-0.0004 in).",
    "Apply cadmium plate to the repair bushes, but not to the bores and the flange faces: refer to PCS-2101 or PCS-2141.": "Нанесите кадмиевое покрытие на ремонтные втулки, но не на отверстия и поверхности фланцев: см. PCS-2101 или PCS-2141.",
    "Apply cadmium plate to the reworked areas: refer to PCS-2100 or PCS-2141.": "Нанесите кадмиевое покрытие на восстановленные участки: см. PCS-2100 или PCS-2141.",
    "Apply cadmium plate to the reworked areas: refer to PCS-2100. The cadmium plate thickness must be between 0,010 and 0,015 mm (0.0004 and 0.0006 in).": "Нанесите кадмиевое покрытие на восстановленные участки: см. PCS-2100. Толщина кадмиевого покрытия должна быть 0,010-0,015 мм (0.0004-0.0006 in).",
    "Apply cadmium plate to the reworked areas: refer to PCS-2141.": "Нанесите кадмиевое покрытие на восстановленные участки: см. PCS-2141.",
    "Apply cadmium to the repair bushes, but not to the bores and the flange faces: refer to PCS-2101 and Figure 602. The cadmium plate thickness must be between 0,010 and 0,015 mm (0.0004 and 0.0006 in).": "Нанесите кадмиевое покрытие на ремонтные втулки, но не на отверстия и поверхности фланцев: см. PCS-2101 и рисунок 602. Толщина кадмиевого покрытия должна быть 0,010-0,015 мм (0.0004-0.0006 in).",
    "Apply chromium plate to the reworked diameter(s) A: refer to PCS-2110, type C, M-DLPS1031-5, M-DLPS1031-7 and Figure 601. Make the chromium plate thickness between 0,020 and 0,025 mm (0.0008 and 0.0010 in).": "Нанесите хромовое покрытие на восстановленный(ые) диаметр(ы) A: см. PCS-2110, тип C, M-DLPS1031-5, M-DLPS1031-7 и рисунок 601. Толщина хромового покрытия должна быть 0,020-0,025 мм (0.0008-0.0010 in).",
    "Apply paint to the repaired area, but not to the bushes: refer to PCS-2500.": "Нанесите лакокрасочное покрытие на отремонтированный участок, но не на втулки: см. PCS-2500.",
    "Apply paint to the repaired area: refer to PCS-2500 and REPAIR.": "Нанесите лакокрасочное покрытие на отремонтированный участок: см. PCS-2500 и РЕМОНТ.",
    "Apply paint to the repaired area: refer to REPAIR and PCS-2500.": "Нанесите лакокрасочное покрытие на отремонтированный участок: см. РЕМОНТ и PCS-2500.",
    "Apply primer paint to the repair bush but not to the areas shown: refer to PCS-2500 and Figure 602.": "Нанесите грунтовочную краску на ремонтную втулку, кроме указанных участков: см. PCS-2500 и рисунок 602.",
    "Apply sermetel to the sliding tube: refer to REPAIR and M-DLPS637.": "Нанесите покрытие Sermetel на скользящую трубу: см. РЕМОНТ и M-DLPS637.",
    "Apply sulphamate nickel plate to the smoothed and polished areas. The sulphamate nickel plate must fill the smoothed and polished areas and be 0,25 mm (0.010 in) above the flange face. Refer to MIL STD 868A, solution 2.": "Нанесите сульфаматно-никелевое покрытие на сглаженные и отполированные участки. Сульфаматно-никелевое покрытие должно заполнить эти участки и выступать над поверхностью фланца на 0,25 мм (0.010 in). См. MIL STD 868A, раствор 2.",
    "Apply the electrically conducting Molykote or rubberised sealant to the oversize bushes and the sliding tube over the dimensions shown: refer to Figure 602. Use electrically conducting Molykote (made from Molykote 111, Material Ref. Item 04-512 and Zinc powder, Material Ref. Item TBA): refer to PCS-7304. Or use electrically conducting rubberised sealant, Material Ref. Item 09-581: refer to IFC30-145-03MD.": "Нанесите электропроводящий Molykote или прорезиненный герметик на ремонтные втулки увеличенного размера и скользящую трубу в пределах указанных размеров: см. рисунок 602. Используйте электропроводящий Molykote (из Molykote 111, код ссылки материала 04-512, и цинкового порошка, код ссылки материала TBA): см. PCS-7304. Либо используйте электропроводящий прорезиненный герметик, код ссылки материала 09-581: см. IFC30-145-03MD.",
    "Calculate diameter B of each repair bush (qty 2), use the formula:": "Рассчитайте диаметр B каждой ремонтной втулки (кол-во 2) по формуле:",
    "Calculate diameter C and dimension H of each repair bush (qty 2):": "Рассчитайте диаметр C и размер H каждой ремонтной втулки (кол-во 2):",
    "Calculate diameter C for the repair bushes, use the formula:": "Рассчитайте диаметр C для ремонтных втулок по формуле:",
    "Calculate the diameter C for the repair bush, use the formula:": "Рассчитайте диаметр C для ремонтной втулки по формуле:",
    "Check the bore diameter of the repair bushes: refer to Figure 602.": "Проверьте диаметр отверстий ремонтных втулок: см. рисунок 602.",
    "Check the bore diameters of the repair bushes: refer to Figure 602.": "Проверьте диаметры отверстий ремонтных втулок: см. рисунок 602.",
    "Check the bore of the repair bush: refer to Figure 602.": "Проверьте отверстие ремонтной втулки: см. рисунок 602.",
    "Check the bores of the repair bushes: refer to Figure 602.": "Проверьте отверстия ремонтных втулок: см. рисунок 602.",
    "Damage or corrosion to the chromium plate band(s) on one or both brake flanges.": "Повреждение или коррозия полос(ы) хромового покрытия на одном или обоих тормозных фланцах.",
    "Damage or wear to diameter A and/or adjacent face.": "Повреждение или износ диаметра A и/или прилегающей поверхности.",
    "Damage or wear to diameter(s) A and faces D and E.": "Повреждение или износ диаметра(ов) A и поверхностей D и E.",
    "Damage or wear to diameter(s) A and/or the adjacent face(s) B.": "Повреждение или износ диаметра(ов) A и/или прилегающей(их) поверхности(ей) B.",
    "Damage or wear to diameters A and/or B and/or C.": "Повреждение или износ диаметров A и/или B и/или C.",
    "Damage or wear to the diameter(s) A.": "Повреждение или износ диаметра(ов) A.",
    "Dia. B = Dia. A (as measured) - 0,005 to + 0,041 mm (- 0.0002 to + 0.0016 in).": "ДИАМ. B = ДИАМ. A (по результатам измерения) - 0,005 до +0,041 мм (-0.0002 до +0.0016 in).",
    "Dia. C = Dia. A (as measured) + 0,034 to -0,004 mm (+ 0.0013 to -0.0002 in).": "ДИАМ. C = ДИАМ. A (по результатам измерения) +0,034 до -0,004 мм (+0.0013 до -0.0002 in).",
    "Dia. C = Dia. A (as measured) + 0,090 to + 0,139 mm (+ 0.0035 to + 0.0055 in). Dim. H = Dim. G (as measured) (spotface to grease hole center line)": "ДИАМ. C = ДИАМ. A (по результатам измерения) +0,090 до +0,139 мм (+0.0035 до +0.0055 in). РАЗМ. H = РАЗМ. G (по результатам измерения) (от подрезки площадки до осевой линии смазочного отверстия).",
    "Do this procedure if both the chromium plate and the base metal are damaged.": "Выполните данную процедуру, если повреждены как хромовое покрытие, так и основной металл.",
    "Do this procedure if there is a wear or damage to diameter(s) A:": "Выполните данную процедуру, если имеется износ или повреждение диаметра(ов) A:",
    "Do this procedure if there is corrosion or damage to the parent metal more than 0,25 mm (0.010 in) deep.": "Выполните данную процедуру, если коррозия или повреждение основного металла имеют глубину более 0,25 мм (0.010 in).",
    "Do this procedure if there is corrosion or damage to the parent metal not more than 0,25 mm (0.010 in) deep.": "Выполните данную процедуру, если глубина коррозии или повреждения основного металла не превышает 0,25 мм (0.010 in).",
    "Do this procedure if there is damage or wear to diameter(s) A and or adjacent face(s) B:": "Выполните данную процедуру, если имеется повреждение или износ диаметра(ов) A и/или прилегающей(их) поверхности(ей) B:",
    "Do this procedure if there is no damage to the parent metal:": "Выполните данную процедуру, если основной металл не поврежден:",
    "Do this procedure if there is wear or damage to diameter A:": "Выполните данную процедуру, если имеется износ или повреждение диаметра A:",
    "Do this procedure, if diameter(s) A is damaged or worn:": "Выполните данную процедуру, если диаметр(ы) A поврежден(ы) или изношен(ы):",
    "Do this procedure, if diameter(s) B is damaged or worn:": "Выполните данную процедуру, если диаметр(ы) B поврежден(ы) или изношен(ы):",
    "Do this procedure, if diameter(s) C is damaged or worn:": "Выполните данную процедуру, если диаметр(ы) C поврежден(ы) или изношен(ы):",
    "Examine the chromium plate surface: refer to M-DLNDT3.": "Осмотрите поверхность хромового покрытия: см. M-DLNDT3.",
    "Grind the chromium plate to the dimension shown: refer to Figure 601. The surface finish must be 2,5 micrometers (100 micro-inches).": "Отшлифуйте хромовое покрытие до указанного размера: см. рисунок 601. Шероховатость поверхности должна быть 2,5 мкм (100 мкдюймов).",
    "Grit blast the shot peened area: refer to PCS-2610. Make sure that the sliding tube is correctly masked.": "Выполните абразивоструйную обработку дробеструйно упрочненного участка: см. PCS-2610. Убедитесь, что скользящая труба правильно замаскирована.",
    "If necessary, hone the bore diameter of repair bushes with a surface finish of 2,5 micrometers (100 micro-inches): refer to Figure 602.": "При необходимости хонінгуйте отверстия ремонтных втулок до шероховатости поверхности 2,5 мкм (100 мкдюймов): см. рисунок 602.",
    "If necessary, hone the bore diameters of the repair bushes to the dimensions shown: refer to Figure 602. Make the surface finish 2,6 micrometers (100 micro-inches).": "При необходимости хонінгуйте диаметры отверстий ремонтных втулок до указанных размеров: см. рисунок 602. Обеспечьте шероховатость поверхности 2,6 мкм (100 мкдюймов).",
    "If necessary, hone the bores of the repair bushes to the dimensions shown: refer to Figure 602. Make the surface finish 2,5 micrometers (100 micro-inches).": "При необходимости хонінгуйте отверстия ремонтных втулок до указанных размеров: см. рисунок 602. Обеспечьте шероховатость поверхности 2,5 мкм (100 мкдюймов).",
    "If necessary, hone the repair bush bore to the dimensions shown: refer to Figure 602.": "При необходимости хонінгуйте отверстие ремонтной втулки до указанных размеров: см. рисунок 602.",
    "If necessary, machine the diameter(s) A to remove the minimum amount of material necessary to remove the damage or wear: refer to M-DLPS1004-4-1 and Figure 602. Do not make the diameter A more than 21,073 mm (0.8296 in). Make the surface finish 1,6 micrometers (63 micro-inches) or better.": "При необходимости обработайте диаметр(ы) A, снимая минимально необходимое количество материала для удаления повреждения или износа: см. M-DLPS1004-4-1 и рисунок 602. Не увеличивайте диаметр A более 21,073 мм (0.8296 in). Обеспечьте шероховатость поверхности 1,6 мкм (63 мкдюймов) или лучше.",
    "If the base metal is damaged or corroded:": "Если основной металл поврежден или поражен коррозией:",
    "If the parent metal was machined at para (h), examine the affected area for flaws: refer to M-DLNDT2, inclusion class 4.": "Если в п. (h) выполнялась механическая обработка основного металла, осмотрите затронутый участок на наличие дефектов: см. M-DLNDT2, класс включений 4.",
    "Locally make the area smooth and polish just sufficiently to remove the corrosion or damage from the face, the minimum flange width is 11,6 mm (0.456 in).": "Локально сгладьте участок и отполируйте его ровно настолько, чтобы удалить коррозию или повреждение с поверхности; минимальная ширина фланца должна быть 11,6 мм (0.456 in).",
    "Locally make the area smooth and polish just sufficiently to remove the corrosion or damage from the surface: the minimum flange width is 11,6 mm (0.456 in).": "Локально сгладьте участок и отполируйте его ровно настолько, чтобы удалить коррозию или повреждение с поверхности; минимальная ширина фланца должна быть 11,6 мм (0.456 in).",
    "Locally remove the Sermetel layer from the sliding tube: refer to REPAIR and M-DLPS637.": "Локально удалите слой Sermetel со скользящей трубы: см. РЕМОНТ и M-DLPS637.",
    "Locally remove the chromium plate from the sliding tube: refer to REPAIR PCS-2110 and Figures 601 and 602.": "Локально удалите хромовое покрытие со скользящей трубы: см. РЕМОНТ, PCS-2110, рисунки 601 и 602.",
    "Locally shot peen the sliding tube: refer to M-DLPS123, Almen A intensity of 0,010 to 0,014 mm (0.0004 to 0.0006 in).": "Локально выполните дробеструйное упрочнение скользящей трубы: см. M-DLPS123, интенсивность Almen A 0,010-0,014 мм (0.0004-0.0006 in).",
    "Locally shot peen the sliding tube: refer to M-DLPS123.": "Локально выполните дробеструйное упрочнение скользящей трубы: см. M-DLPS123.",
    "M-DLPS900, M-DL PS1000 and Figure 601. Do not make diameter(s) A more than 18,568 mm (0.7310 in). Make the surface finish 1,6 micrometers (63 micro-inches).": "M-DLPS900, M-DL PS1000 и рисунок 601. Не увеличивайте диаметр(ы) A более 18,568 мм (0.7310 in). Обеспечьте шероховатость поверхности 1,6 мкм (63 мкдюймов).",
    "Machine (do not grind) the sulphamate nickel plate to the dimensions shown: refer to Figure 601.": "Обработайте сульфаматно-никелевое покрытие резанием (не шлифовать) до указанных размеров: см. рисунок 601.",
    "Machine diameter A just sufficiently to remove the damage or corrosion: refer to M-DLPS1004-4-1. The diameter must not be less than 177,07 mm (6.971 in). The surface finish must be 1,6 micrometers (63 micro-inches).": "Обработайте диаметр A ровно настолько, чтобы удалить повреждение или коррозию: см. M-DLPS1004-4-1. Диаметр не должен быть менее 177,07 мм (6.971 in). Шероховатость поверхности должна быть 1,6 мкм (63 мкдюймов).",
    "Machine diameter A to remove the damage or corrosion within the dimensions shown: refer to M-DLPS1004-4-1 and Figure 601.": "Обработайте диаметр A для удаления повреждения или коррозии в пределах указанных размеров: см. M-DLPS1004-4-1 и рисунок 601.",
    "Machine diameter B of the repair bush(es), use formula:": "Обработайте диаметр B ремонтной(ых) втулки(ок) по формуле:",
    "Machine diameter(s) A sufficiently to remove the damage or wear within the dimensions shown: refer to M-DLPS1004-4-1 and Figure 601. Make the surface finish 1,6 micrometers (63 micro-inches).": "Обработайте диаметр(ы) A настолько, чтобы удалить повреждение или износ в пределах указанных размеров: см. M-DLPS1004-4-1 и рисунок 601. Обеспечьте шероховатость поверхности 1,6 мкм (63 мкдюймов).",
    "Machine diameter(s) A sufficiently to remove the damage or wear: refer to": "Обработайте диаметр(ы) A настолько, чтобы удалить повреждение или износ: см.",
    "Machine diameter(s) A sufficiently to remove the minimum amount of material necessary to remove the wear or damage within the dimensions shown: refer to M-DLPS1004-4-1 and Figure 601. Make the surface finish 1,6 micrometers (63 micro-inches).": "Обработайте диаметр(ы) A, снимая минимально необходимое количество материала для устранения износа или повреждения в пределах указанных размеров: см. M-DLPS1004-4-1 и рисунок 601. Обеспечьте шероховатость поверхности 1,6 мкм (63 мкдюймов).",
    "Machine face D of the repair bush to get the correct dimensions: refer to Figure 602. Make the surface finish 1,6 micrometers (63 micro-inches).": "Обработайте поверхность D ремонтной втулки для получения правильных размеров: см. рисунок 602. Обеспечьте шероховатость поверхности 1,6 мкм (63 мкдюймов).",
    "Machine the adjacent face(s) B using spotface cutter sufficiently to remove the damage or wear within the dimensions shown: refer to M-DLPS1004-4-1 and Figure 601. Make the surface finish 1,6 micrometers (63 micro-inches).": "Обработайте прилегающую(ие) поверхность(и) B с помощью подрезной фрезы настолько, чтобы удалить повреждение или износ в пределах указанных размеров: см. M-DLPS1004-4-1 и рисунок 601. Обеспечьте шероховатость поверхности 1,6 мкм (63 мкдюймов).",
    "Machine the adjacent spotfaces D and E to remove any damage or wear, within the dimensions shown: refer to Figure 601. Make the surface finish 1,6 micrometers (63 micro-inches).": "Обработайте прилегающие подрезки площадок D и E для удаления любых повреждений или износа в пределах указанных размеров: см. рисунок 601. Обеспечьте шероховатость поверхности 1,6 мкм (63 мкдюймов).",
    "Machine the chamfers to the dimensions shown: refer to Figure 602.": "Обработайте фаски по указанным размерам: см. рисунок 602.",
    "Machine the diameter A and/or the adjacent face to remove the damage or wear within the dimensions shown: refer to M-DLPS1004-4-1 and Figure 601. Do not increase diameter A more than 48,125 mm (1.8947 in). Make the surface finish 1,6 micrometers (63 micro-inches).": "Обработайте диаметр A и/или прилегающую поверхность для удаления повреждения или износа в пределах указанных размеров: см. M-DLPS1004-4-1 и рисунок 601. Не увеличивайте диаметр A более 48,125 мм (1.8947 in). Обеспечьте шероховатость поверхности 1,6 мкм (63 мкдюймов).",
    "Machine the diameter(s) A to remove the damage or wear within the dimensions shown: refer to M-DLPS1004-4-1 and Figure 601. Make the surface finish 0,8 micrometers (32 micro-inches).": "Обработайте диаметр(ы) A для удаления повреждения или износа в пределах указанных размеров: см. M-DLPS1004-4-1 и рисунок 601. Обеспечьте шероховатость поверхности 0,8 мкм (32 мкдюймов).",
    "Machine the diameter(s) B to remove the damage or wear within the dimensions shown: refer to M-DLPS1004-4-1 and Figure 601. Make the surface finish 0,8 micrometers (32 micro-inches).": "Обработайте диаметр(ы) B для удаления повреждения или износа в пределах указанных размеров: см. M-DLPS1004-4-1 и рисунок 601. Обеспечьте шероховатость поверхности 0,8 мкм (32 мкдюймов).",
    "Machine the diameter(s) C to remove the damage or wear within the dimensions shown: refer to M-DLPS1004-4-1 and Figure 602. Make the surface finish 0,8 micrometers (32 micro-inches).": "Обработайте диаметр(ы) C для удаления повреждения или износа в пределах указанных размеров: см. M-DLPS1004-4-1 и рисунок 602. Обеспечьте шероховатость поверхности 0,8 мкм (32 мкдюймов).",
    "Machine the dimension across the bush flanges between 150,917 and 150,957 mm (5.9416 and 5.9432 in). The symmetry must be as shown: refer to Figure 602.": "Обработайте размер по фланцам втулок до 150,917-150,957 мм (5.9416-5.9432 in). Симметрия должна соответствовать рисунку 602.",
    "Machine the repair bush to the dimensions shown and calculated. Machine face D to get the correct dimension after installation: refer to M-DLPS900 and Figure 602. Make the surface finish 1,6 micrometers (63 micro-inches).": "Обработайте ремонтную втулку до указанных и расчетных размеров. Обработайте поверхность D для получения правильного размера после установки: см. M-DLPS900 и рисунок 602. Обеспечьте шероховатость поверхности 1,6 мкм (63 мкдюймов).",
    "Machine the repair bush to the dimensions shown and calculated: refer to M-DLPS900 and Figure 602.": "Обработайте ремонтную втулку до указанных и расчетных размеров: см. M-DLPS900 и рисунок 602.",
    "NOTE: Refer to Repair No. 9-5, if only the chromium plate is damaged.": "ПРИМЕЧАНИЕ. См. ремонт № 9-5, если повреждено только хромовое покрытие.",
    "NOTE: This operation includes 18 hours de-embrittle at 177 to 205 oC (350 to 400 oF).": "ПРИМЕЧАНИЕ. Данная операция включает 18 ч обезводораживания при 177-205 oC (350-400 oF).",
    "NOTE: This operation includes 23 hours de-embrittle at 177 to 205 oC (350 to 400 oF).": "ПРИМЕЧАНИЕ. Данная операция включает 23 ч обезводораживания при 177-205 oC (350-400 oF).",
    "NOTE: This operation includes 4 hours de-embrittle at 177 to 205 oC (350 to 400 oF).": "ПРИМЕЧАНИЕ. Данная операция включает 4 ч обезводораживания при 177-205 oC (350-400 oF).",
    "Refer to M-DLPS1004-4-1 and machine diameter A just sufficiently to remove the damage or corrosion: the diameter must not be less than 30,365 mm (1.1956 in). The surface finish must be 1,6 micrometers (63 micro-inches).": "См. M-DLPS1004-4-1 и обработайте диаметр A ровно настолько, чтобы удалить повреждение или коррозию; диаметр не должен быть менее 30,365 мм (1.1956 in). Шероховатость поверхности должна быть 1,6 мкм (63 мкдюймов).",
    "Shot peen the flange face(s) including the sulphamate nickel plated areas: refer to M-DLPS123, Almen A intensity of 0,010 to 0,014 mm (0.0004 to 0.0006 in).": "Дробеструйно упрочните поверхность(и) фланца(ев), включая участки с сульфаматно-никелевым покрытием: см. M-DLPS123, интенсивность Almen A 0,010-0,014 мм (0.0004-0.0006 in).",
}

PART7_SHARED_EXACT_MAP: dict[str, str] = {
    "POINT C": "ТОЧКА C",
    "POINT D": "ТОЧКА D",
    "Y POINT C": "Y ТОЧКА C",
    "DIMENSION B": "РАЗМЕР B",
    "DIMENSION E": "РАЗМЕР E",
    "MAXIMUM (REFERENCE)": "МАКСИМУМ (СПРАВ.)",
    "WITHOUT BUSHES": "БЕЗ ВТУЛОК",
    "WITH BUSHES": "С ВТУЛКАМИ",
    "WITH REPAIR BUSHES": "С РЕМОНТНЫМИ ВТУЛКАМИ",
    "WITH REPAIR BUSHES INSTALLED": "С УСТАНОВЛЕННЫМИ РЕМОНТНЫМИ ВТУЛКАМИ",
    "WITHOUT REPAIR BUSHES": "БЕЗ РЕМОНТНЫХ ВТУЛОК",
    "(WITHOUT BEARING)": "(БЕЗ ПОДШИПНИКА)",
    "SMOOTH TO RADIUS. 2 PLACES": "СГЛАДИТЬ ДО РАДИУСА. 2 МЕСТА",
    "SEALANT TO PCS-7200": "ГЕРМЕТИК ПО PCS-7200",
    "APPLY SEALANT TO PCS-7200": "НАНЕСТИ ГЕРМЕТИК ПО PCS-7200",
    "APPLY SEALANT TO PCS-7200 TYPICAL 6 PLACES": "НАНЕСТИ ГЕРМЕТИК ПО PCS-7200, ТИПОВО 6 МЕСТА",
    "PART SECTION Z-Z": "ЧАСТИЧНОЕ СЕЧЕНИЕ Z-Z",
    "DRAIN HOLE J": "ДРЕНАЖНОЕ ОТВЕРСТИЕ J",
    "REFERENCE BOTTOM OF MAIN FITTING": "СПРАВ. НИЖНЯЯ ЧАСТЬ КОРПУСА СТОЙКИ",
    "MACHINING": "МЕХАНИЧЕСКАЯ ОБРАБОТКА",
    "THICKNESS": "ТОЛЩИНА",
    "CHAMFER (2 PLACES)": "ФАСКА (2 МЕСТА)",
    "Press Pad and Drawbolt": "Нажимная опора и стяжной болт",
    "Cutter": "Резец",
    "Finish machine the repair bush 450258806": "Окончательно обработать ремонтную втулку 450258806",
    "Blank bush": "Заготовка втулки",
    "Blank bushes": "Заготовки втулок",
    "Blank bearing": "Заготовка подшипника",
    "Electrically conducting Mastinox to M-DLPS709-14": "Электропроводящий Mastinox по M-DLPS709-14",
    "Electrically conducting Mastinox D40": "Электропроводящий Mastinox D40",
    "Electrically conducting Mastinox D40: refer to M-DLPS709-14": "Электропроводящий Mastinox D40: см. M-DLPS709-14",
    "Electrically conducting zinc loaded Mastinox D40: refer to M-DLPS709-14": (
        "Электропроводящий Mastinox D40 с цинковым наполнителем: см. M-DLPS709-14"
    ),
    "Electrically conducting zinc loaded Molykote 111: refer to PCS-7304": (
        "Электропроводящий Molykote 111 с цинковым наполнителем: см. PCS-7304"
    ),
    "Jointing Compound, zinc loaded Mastinox D40: refer to M-DLPS709-14": (
        "Монтажный состав Mastinox D40 с цинковым наполнителем: см. M-DLPS709-14"
    ),
    "Aluminium bronze AMS4590 or AMS4881 centrifugally cast": (
        "Алюминиевая бронза AMS4590 или AMS4881, центробежного литья"
    ),
    "Oversize Bush(es) - Machining and Installation Figure 602": (
        "Ремонтные втулки увеличенного размера - Механическая обработка и установка рисунок 602"
    ),
}

EXACT_MAP.update(PART6_EXACT_MAP)
EXACT_MAP.update(PART6_EXACT_MAP_2)
EXACT_MAP.update(PART7_SHARED_EXACT_MAP)


def _translate_fragment(text: str) -> str:
    return translate_text(text)


def _regex_translate_examine_for_flaws(match: re.Match[str]) -> str:
    subject = _translate_fragment(match.group(1))
    refs = _translate_fragment(match.group(2))
    return f"Осмотрите {subject} на наличие дефектов: см. {refs}."


def _regex_translate_identify_repair_number(match: re.Match[str]) -> str:
    number = match.group(1)
    refs = _translate_fragment(match.group(2))
    return f"Нанесите рядом с номером детали номер ремонта Messier-Dowty Limited {number}: см. {refs}."


def _regex_translate_identify_safran_repair_number(match: re.Match[str]) -> str:
    number = match.group(1)
    refs = _translate_fragment(match.group(2))
    return f"Нанесите рядом с номером детали номер ремонта Safran Landing Systems {number}: см. {refs}."


def _regex_translate_identify_safran_repair_number_no_value(match: re.Match[str]) -> str:
    refs = _translate_fragment(match.group(1))
    return f"Нанесите рядом с номером детали номер ремонта Safran Landing Systems: см. {refs}."


REGEX_RULES: list[tuple[re.Pattern[str], str | callable]] = [
    (
        re.compile(r"^Examine the (.+?) for flaws: refer to (.+)\.$", re.IGNORECASE),
        _regex_translate_examine_for_flaws,
    ),
    (
        re.compile(
            r"^Identify the part with the Messier-Dowty Limited repair number ([0-9A-Z]+) adjacent to the part number: refer to (.+)\.$",
            re.IGNORECASE,
        ),
        _regex_translate_identify_repair_number,
    ),
    (
        re.compile(
            r"^Identify the part with the Safran Landing Systems repair number ([0-9A-Z]+) adjacent to the part number: refer to (.+)\.$",
            re.IGNORECASE,
        ),
        _regex_translate_identify_safran_repair_number,
    ),
    (
        re.compile(
            r"^Identify the part with the Safran Landing Systems repair number adjacent to the part number: refer to (.+)\.$",
            re.IGNORECASE,
        ),
        _regex_translate_identify_safran_repair_number_no_value,
    ),
    (re.compile(r"^VIEW ON ARROW\s+(.+)$", re.IGNORECASE), r"ВИД ПО СТРЕЛКЕ \1"),
    (re.compile(r"^VIEW\s+(.+)$", re.IGNORECASE), r"ВИД \1"),
    (re.compile(r"^SECTION\s+(.+)$", re.IGNORECASE), r"СЕЧЕНИЕ \1"),
    (re.compile(r"^DETAIL\s+(.+)$", re.IGNORECASE), r"ДЕТАЛЬ \1"),
    (re.compile(r"^and\s+(?=\d)", re.IGNORECASE), "и "),
    (re.compile(r"^(.+?)\s+DIA\s+\((\d+)\s+PLACES\)$", re.IGNORECASE), r"\1 ДИАМ. (\2 МЕСТА)"),
    (re.compile(r"^(.+?)\s+DIA\.$", re.IGNORECASE), r"\1 ДИАМ."),
    (re.compile(r"^(.+?)\s+DIA$", re.IGNORECASE), r"\1 ДИАМ."),
    (re.compile(r"^(.+?)\s+\(DIAMETER\)$", re.IGNORECASE), r"\1 (ДИАМЕТР)"),
    (re.compile(r"^\(DIAMETER\)$", re.IGNORECASE), "(ДИАМЕТР)"),
    (re.compile(r"^DIAMETER\s+(.+)$", re.IGNORECASE), r"ДИАМЕТР \1"),
    (re.compile(r"^(.+?)\s+DIAMETER$", re.IGNORECASE), r"\1 ДИАМЕТР"),
    (re.compile(r"^(.+?)\s+\((\d+)\s+HOLES\)$", re.IGNORECASE), r"\1 (\2 ОТВЕРСТИЯ)"),
    (re.compile(r"^(.+?)\s+(\d+)\s+HOLES$", re.IGNORECASE), r"\1 \2 ОТВЕРСТИЯ"),
    (re.compile(r"^(\d+)\s+HOLES$", re.IGNORECASE), r"\1 ОТВЕРСТИЯ"),
    (re.compile(r"^(.+?)\s+\((\d+)\s+DIAMETERS\)$", re.IGNORECASE), r"\1 (\2 ДИАМЕТРА)"),
    (re.compile(r"^\((\d+)\s+PLACES\)$", re.IGNORECASE), r"(\1 МЕСТА)"),
    (re.compile(r"^(\d+)\s+PLACES$", re.IGNORECASE), r"\1 МЕСТА"),
    (re.compile(r"^(.+?)\s+(\d+)\s+PLACES$", re.IGNORECASE), r"\1 \2 МЕСТА"),
    (re.compile(r"^(\d+)\s+PLACES\s+(.+)$", re.IGNORECASE), r"\1 МЕСТА \2"),
    (re.compile(r"TYPICAL\s+(\d+)\s+PLACES", re.IGNORECASE), r"ТИПОВО \1 МЕСТА"),
    (re.compile(r"TYPICAL\s+(\d+)\s+HOLES", re.IGNORECASE), r"ТИПОВО \1 ОТВЕРСТИЯ"),
    (re.compile(r"TYPICAL\s+(\d+)\s+LUGS", re.IGNORECASE), r"ТИПОВО \1 ПРОУШИНЫ"),
    (re.compile(r"^\((\d+)\s+POSITIONS\)$", re.IGNORECASE), r"(\1 ПОЗИЦИИ)"),
    (re.compile(r"^(.+?)\s+(\d+)\s+POSITIONS$", re.IGNORECASE), r"\1 \2 ПОЗИЦИИ"),
    (re.compile(r"^\(BOTH HOLES\)$", re.IGNORECASE), "(ОБА ОТВЕРСТИЯ)"),
    (re.compile(r"^\(BOTH SIDES\)$", re.IGNORECASE), "(С ОБЕИХ СТОРОН)"),
    (re.compile(r"^\(INNER DIAMETER\)$", re.IGNORECASE), "(ВНУТРЕННИЙ ДИАМЕТР)"),
    (re.compile(r"^\(CHAMFER\)$", re.IGNORECASE), "(ФАСКА)"),
    (re.compile(r"^IN HOLES \(Qty (\d+)\)$", re.IGNORECASE), r"В ОТВЕРСТИЯХ (КОЛ-ВО \1)"),
    (re.compile(r"^TYPICAL\s+(\d+)\s+PLACES$", re.IGNORECASE), r"ТИПОВОЕ ИСПОЛНЕНИЕ, \1 МЕСТА"),
    (re.compile(r"^(.+?)\s+RAD\.$", re.IGNORECASE), r"\1 РАД."),
    (re.compile(r"^(.+?)\s+MAXIMUM$", re.IGNORECASE), r"\1 МАКСИМУМ"),
    (re.compile(r"^(.+?)\s+MINIMUM$", re.IGNORECASE), r"\1 МИНИМУМ"),
    (re.compile(r"^(.+?)\s+REFERENCE$", re.IGNORECASE), r"\1 СПРАВ."),
    (re.compile(r"^(.+?)\s+\(REFERENCE\)$", re.IGNORECASE), r"\1 (СПРАВ.)"),
    (re.compile(r"^POINT\s+(.+)$", re.IGNORECASE), r"ТОЧКА \1"),
    (re.compile(r"^(.+?)\s+POINT\s+(.+)$", re.IGNORECASE), r"\1 ТОЧКА \2"),
    (re.compile(r"^DIMENSION\s+(.+)$", re.IGNORECASE), r"РАЗМЕР \1"),
    (re.compile(r"^FACE\s+(.+)$", re.IGNORECASE), r"ПОВЕРХНОСТЬ \1"),
    (re.compile(r"^PLANE PASSES THROUGH$", re.IGNORECASE), "ПЛОСКОСТЬ ПРОХОДИТ ЧЕРЕЗ"),
    (re.compile(r"^PLANE PASSES$", re.IGNORECASE), "ПЛОСКОСТЬ ПРОХОДИТ"),
    (re.compile(r"^THROUGH POINT\s+(.+)$", re.IGNORECASE), r"ЧЕРЕЗ ТОЧКУ \1"),
    (re.compile(r"^(.+?)\s+RADIUS TYPICAL$", re.IGNORECASE), r"\1 РАДИУС ТИПОВО"),
    (re.compile(r"^(.+?)\s+RADIUS TYPICAL AROUND LUG$", re.IGNORECASE), r"\1 РАДИУС, ТИПОВО ВОКРУГ ПРОУШИНЫ"),
    (re.compile(r"^(.+?)\s+TYPICAL BOTH BORES$", re.IGNORECASE), r"\1 ТИПОВО ДЛЯ ОБОИХ ОТВЕРСТИЙ"),
    (re.compile(r"^(.+?)\s+CORNER RADIUS \((\d+)\s+PLACES\)$", re.IGNORECASE), r"\1 РАДИУС СКРУГЛЕНИЯ (\2 МЕСТА)"),
    (re.compile(r"^(.+?)\s+CORNER RADIUS$", re.IGNORECASE), r"\1 РАДИУС СКРУГЛЕНИЯ"),
    (re.compile(r"^(.+?)\s+BELOW THIS SURFACE$", re.IGNORECASE), r"\1 НИЖЕ ЭТОЙ ПОВЕРХНОСТИ"),
    (re.compile(r"^(.+?)\s+OVER RADIUS\.$", re.IGNORECASE), r"\1 ПО РАДИУСУ."),
    (re.compile(r"^SMOOTH BLEND TO LARGER RADIUS$", re.IGNORECASE), "ПЛАВНО СОПРЯЧЬ С БОЛЬШИМ РАДИУСОМ"),
    (re.compile(r"^D1 BEFORE MACHINING$", re.IGNORECASE), "D1 ДО МЕХАНИЧЕСКОЙ ОБРАБОТКИ"),
    (re.compile(r"^D2 AFTER MACHINING$", re.IGNORECASE), "D2 ПОСЛЕ МЕХАНИЧЕСКОЙ ОБРАБОТКИ"),
    (re.compile(r"^x\s+(.+?)\s+DEGREE(?:S)?\s+CHAMFER$", re.IGNORECASE), r"x \1 ГРАДУСОВ ФАСКА"),
    (re.compile(r"^X\s+(.+?)\s+DEGREE(?:S)?\s+CHAMFER$", re.IGNORECASE), r"X \1 ГРАДУСОВ ФАСКА"),
    (re.compile(r"^(.+?)\s+X\s+(\d+(?:\.\d+)?)\s+DEGREE(?:S)?\s+CHAMFER$", re.IGNORECASE), r"\1 X \2 ГРАДУСОВ ФАСКА"),
    (re.compile(r"^Page\s+(\d+)\s+of\s+(\d+)$", re.IGNORECASE), r"Страница \1 из \2"),
    (re.compile(r"^Page\s+(\d+)\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+(\d{1,2}/\d{4})$", re.IGNORECASE), r"Страница \1 \2 \3"),
    (re.compile(r"^Page\s+(\d+)$", re.IGNORECASE), r"Страница \1"),
    (re.compile(r"^Added repair no\.\s+([0-9-]+)$", re.IGNORECASE), r"Добавлен ремонт № \1"),
    (re.compile(r"\(Sheet\s+(\d+)\s+of\s+(\d+)\)", re.IGNORECASE), r"(Лист \1 из \2)"),
    (re.compile(r"\(Лист\s+(\d+)\s+of\s+(\d+)\)", re.IGNORECASE), r"(Лист \1 из \2)"),
]


def _letters_only(text: str) -> str:
    return "".join(ch for ch in text if ch.isalpha())


def _apply_case_style(source: str, target: str) -> str:
    letters = _letters_only(source)
    if not letters:
        return target
    if letters.isupper():
        return target.upper()
    if source[:1].isupper() and source[1:2].islower():
        return target[:1].upper() + target[1:]
    return target


def _replace_phrase(text: str, source: str, target: str) -> str:
    pattern = re.compile(re.escape(source), re.IGNORECASE)

    def repl(match: re.Match[str]) -> str:
        return _apply_case_style(match.group(0), target)

    return pattern.sub(repl, text)


def _cleanup_mixed_translation(text: str) -> str:
    cleaned = re.sub(r"\b(?:the|a|an)\s+(?=[А-Яа-яЁё])", "", text, flags=re.IGNORECASE)
    cleaned = re.sub(r"(?<![A-Za-z])or(?=(?:\s|$))", "или", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"(?<![A-Za-z])and(?=(?:\s|$))", "и", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"(?<![A-Za-z])from(?=(?:\s|$))", "из", cleaned, flags=re.IGNORECASE)
    cleaned = cleaned.replace("Установка of ", "Установка ")
    cleaned = cleaned.replace("Сборка of ", "Сборка ")
    cleaned = cleaned.replace("Уплотнение Configuration", "Конфигурация уплотнения")
    cleaned = cleaned.replace("Уплотнение Конфигурация", "Конфигурация уплотнения")
    cleaned = cleaned.replace("M-D\nSpec", "Спецификация\nM-D")
    cleaned = cleaned.replace("Electrical Жгут оси", "Электрический жгут оси")
    cleaned = cleaned.replace(" Figure ", " рисунок ")
    cleaned = cleaned.replace(" to AMS", " по AMS")
    cleaned = cleaned.replace(" to MTL", " по MTL")
    cleaned = cleaned.replace(" to NCT", " по NCT")
    cleaned = cleaned.replace(" to BS", " по BS")
    cleaned = cleaned.replace(" to MAT", " по MAT")
    cleaned = cleaned.replace("Шплинтs", "Шплинты")
    cleaned = cleaned.replace("подшипникs", "подшипники")
    cleaned = cleaned.replace("Подшипникs", "Подшипники")
    cleaned = cleaned.replace("Втулкаs", "Втулки")
    cleaned = cleaned.replace("Уплотнениеs", "Уплотнения")
    cleaned = cleaned.replace("Процедураs", "процедуры")
    cleaned = cleaned.replace("Ремонтная детальs", "Ремонтные детали")
    cleaned = cleaned.replace("Ремонтная втулкаs", "Ремонтные втулки")
    cleaned = cleaned.replace(")-", ") -")
    cleaned = cleaned.replace("ФАСКАS", "ФАСКИ")
    cleaned = cleaned.replace("ФАСКАs", "ФАСКИ")
    return cleaned


def _word_find_literal(text: str) -> str:
    out = text.replace("^", "^^")
    out = out.replace("\t", "^t")
    out = out.replace("\n", "^l")
    return out


def _translate_month_date_fragments(text: str) -> str:
    pattern = re.compile(r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+(\d{1,2})/(\d{4})\b")

    def repl(match: re.Match[str]) -> str:
        month = MONTH_NUMBER[match.group(1)]
        day = int(match.group(2))
        year = match.group(3)
        return f"{day:02d}.{month}.{year}"

    return pattern.sub(repl, text)


def _normalize_numeric_ranges(text: str) -> str:
    translated = re.sub(r"(?<=\d)\s+to\s+(?=\d)", "-", text, flags=re.IGNORECASE)
    translated = re.sub(r"(?<=\d)\s+and\s+(?=\d)", ", ", translated, flags=re.IGNORECASE)
    translated = re.sub(r"(?<=\d)\s+to\s*$", "-", translated, flags=re.IGNORECASE)
    translated = re.sub(r"(?<=\d)\s+and\s*$", ",", translated, flags=re.IGNORECASE)
    return translated


def _looks_code_like(text: str) -> bool:
    stripped = " ".join(text.split())
    if not stripped:
        return True
    if stripped in RESERVED_UNCHANGED:
        return True
    tokens = stripped.split()
    month_tokens = {"Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"}
    if tokens and all(
        token in {"Blank"} or token in month_tokens or bool(re.fullmatch(r"[0-9./,-]+", token))
        for token in tokens
    ):
        return True
    return bool(CODE_LIKE_RE.fullmatch(stripped))


def translate_text(text: str) -> str:
    original = text
    if not LATIN_RE.search(text):
        return text
    stripped = " ".join(text.split())
    if stripped in EXACT_MAP:
        return EXACT_MAP[stripped]
    month_translated = _translate_month_date_fragments(original)
    range_normalized = _normalize_numeric_ranges(month_translated)
    if range_normalized != original and not LATIN_RE.search(range_normalized):
        return range_normalized
    if _looks_code_like(stripped):
        if range_normalized != original:
            return range_normalized
        if month_translated != original:
            return month_translated
        return original

    translated = range_normalized
    for pattern, replacement in REGEX_RULES:
        translated = pattern.sub(replacement, translated)
    stripped_after_regex = " ".join(translated.split())
    if stripped_after_regex in EXACT_MAP:
        return EXACT_MAP[stripped_after_regex]
    for source, target in sorted(PHRASE_RULES, key=lambda item: len(item[0]), reverse=True):
        translated = _replace_phrase(translated, source, target)
    translated = _cleanup_mixed_translation(translated)
    translated = translated.replace("  ", " ")
    return translated


def collect_unique_english_strings(docx_path: Path) -> list[str]:
    doc = Document(str(docx_path))
    segments = collect_segments(doc, include_headers=True, include_footers=True)
    unique: list[str] = []
    seen: set[str] = set()
    for seg in segments:
        text = seg.source_plain
        if not LATIN_RE.search(text):
            continue
        normalized = " ".join(text.split())
        if normalized in seen:
            continue
        seen.add(normalized)
        unique.append(text)
    unique.sort(key=len, reverse=True)
    return unique


def build_translation_map(unique_strings: list[str]) -> dict[str, str]:
    mapping: dict[str, str] = {}
    for source in unique_strings:
        target = translate_text(source)
        if target != source:
            mapping[source] = target
    return dict(sorted(mapping.items(), key=lambda item: len(item[0]), reverse=True))


def find_unresolved(unique_strings: list[str], mapping: dict[str, str]) -> list[str]:
    unresolved: list[str] = []
    for source in unique_strings:
        if source in mapping:
            continue
        normalized = " ".join(source.split())
        if normalized in RESERVED_UNCHANGED:
            continue
        if _looks_code_like(normalized):
            continue
        unresolved.append(source)
    return unresolved


def _set_paragraph_text_minimal(paragraph, text: str) -> None:
    _, spans, _inline = paragraph_to_tagged(paragraph)
    span = spans[0] if spans else None
    _clear_paragraph_runs(paragraph)
    run = paragraph.add_run(text)
    if span is not None:
        _apply_style(run, span)


def replace_segments_direct(docx_path: Path, mapping: dict[str, str], *, min_len: int = 0) -> Counter:
    doc = Document(str(docx_path))
    segments = collect_segments(doc, include_headers=True, include_footers=True)
    counts: Counter[str] = Counter()
    changed = False
    for seg in segments:
        source = seg.source_plain
        target = mapping.get(source)
        if not target or len(source) < min_len:
            continue
        _set_paragraph_text_minimal(seg.paragraph_ref, target)
        counts[source] += 1
        changed = True
    if changed:
        doc.save(str(docx_path))
    return counts


def _iter_table_paragraphs(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                yield from _iter_table_paragraphs(cell.tables)


def _iter_nested_tables(tables):
    for table in tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_nested_tables(cell.tables)


def _iter_all_paragraphs(doc):
    yield from doc.paragraphs
    yield from _iter_table_paragraphs(doc.tables)
    for section in doc.sections:
        for part in (section.header, section.first_page_header, section.even_page_header):
            yield from part.paragraphs
            yield from _iter_table_paragraphs(part.tables)
        for part in (section.footer, section.first_page_footer, section.even_page_footer):
            yield from part.paragraphs
            yield from _iter_table_paragraphs(part.tables)


def _iter_all_tables(doc):
    yield from _iter_nested_tables(doc.tables)
    for section in doc.sections:
        for part in (section.header, section.first_page_header, section.even_page_header):
            yield from _iter_nested_tables(part.tables)
        for part in (section.footer, section.first_page_footer, section.even_page_footer):
            yield from _iter_nested_tables(part.tables)


def cleanup_remaining_paragraphs(docx_path: Path) -> Counter:
    doc = Document(str(docx_path))
    counts: Counter[str] = Counter()
    changed = False
    for paragraph in _iter_all_paragraphs(doc):
        source = paragraph.text
        if not source or not LATIN_RE.search(source):
            continue
        target = translate_text(source)
        if target == source:
            continue
        _set_paragraph_text_minimal(paragraph, target)
        counts[source] += 1
        changed = True
    if changed:
        doc.save(str(docx_path))
    return counts


def patch_header_footer_textboxes(docx_path: Path) -> int:
    replacements = {
        "<w:t>Page</w:t>": "<w:t>Стр.</w:t>",
        "Page ": "Стр. ",
        "No.": "№",
        "Jan": "янв.",
        "Feb": "февр.",
        "Mar": "мар.",
        "Apr": "апр.",
        "May": "мая",
        "Jun": "июн.",
        "Jul": "июл.",
        "Aug": "авг.",
        "Sep": "сент.",
        "Oct": "окт.",
        "Nov": "нояб.",
        "Dec": "дек.",
    }
    changed_files = 0
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    try:
        with ZipFile(docx_path, "r") as zin, ZipFile(tmp_path, "w", compression=ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if re.fullmatch(r"word/(header|footer)\d+\.xml", info.filename):
                    text = data.decode("utf-8")
                    updated = text
                    updated = _translate_month_date_fragments(updated)
                    for source, target in replacements.items():
                        updated = updated.replace(source, target)
                    updated = _normalize_numeric_ranges(updated)
                    if updated != text:
                        data = updated.encode("utf-8")
                        changed_files += 1
                zout.writestr(info, data)
        shutil.move(str(tmp_path), str(docx_path))
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except OSError:
            pass
    return changed_files


def patch_document_xml_fragments(docx_path: Path) -> int:
    replacements = {
        "Steel to 300M": "Сталь, 300M",
        "Сталь to 300M": "Сталь, 300M",
        ">DIAMETER<": ">ДИАМЕТР<",
        "ДИАМЕТР ПОСЛЕ ХРОМИРОВАНИЯ PLATE 135,860-135,923mm": "ДИАМЕТР ПОСЛЕ ХРОМИРОВАНИЯ 135,860-135,923mm",
        "ДИАМЕТР ДО ХРОМИРОВАНИЯ PLATE": "ДИАМЕТР ДО ХРОМИРОВАНИЯ",
        "ДИАМЕТР AFTER GRINDING OF ХРОМОВОЕ ПОКРЫТИЕ": "ДИАМЕТР ПОСЛЕ ШЛИФОВАНИЯ ХРОМОВОГО ПОКРЫТИЯ",
        "ДИАМЕТР AFTER GRINDING ХРОМОВОЕ ПОКРЫТИЕ": "ДИАМЕТР ПОСЛЕ ШЛИФОВАНИЯ ХРОМОВОГО ПОКРЫТИЯ",
        "СЕЧЕНИЕ Z-Z (WITHOUT ВТУЛКИ)": "СЕЧЕНИЕ Z-Z (БЕЗ ВТУЛКИ)",
        "TERMINATION ПО M-DLPS1031-6. IRREGULAR LINE IS PERMITTED": "ГРАНИЦА ПО M-DLPS1031-6. НЕРОВНАЯ ЛИНИЯ ДОПУСКАЕТСЯ",
    }
    changed_files = 0
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    try:
        with ZipFile(docx_path, "r") as zin, ZipFile(tmp_path, "w", compression=ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/document.xml":
                    text = data.decode("utf-8")
                    updated = text
                    for source, target in replacements.items():
                        updated = updated.replace(source, target)
                    if updated != text:
                        data = updated.encode("utf-8")
                        changed_files += 1
                zout.writestr(info, data)
        shutil.move(str(tmp_path), str(docx_path))
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except OSError:
            pass
    return changed_files


def enforce_table_readability(docx_path: Path, *, min_font_size_pt: float = 9.0) -> int:
    doc = Document(str(docx_path))
    changed_runs = 0
    for table in _iter_all_tables(doc):
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fmt = paragraph.paragraph_format
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)
                    fmt.line_spacing = 0.9
                    for run in paragraph.runs:
                        size = run.font.size.pt if run.font.size else None
                        if size is None or size < min_font_size_pt:
                            run.font.size = Pt(min_font_size_pt)
                            changed_runs += 1
    if changed_runs:
        doc.save(str(docx_path))
    return changed_runs


def auto_fix_layout(source_docx: Path, translated_docx: Path) -> dict[str, int]:
    source_doc = Document(str(source_docx))
    translated_doc = Document(str(translated_docx))
    source_segments = collect_segments(source_doc, include_headers=True, include_footers=True)
    translated_segments = collect_segments(translated_doc, include_headers=True, include_footers=True)

    source_by_loc = {seg.location: seg for seg in source_segments}
    aligned: list = []
    for seg in translated_segments:
        source_seg = source_by_loc.get(seg.location)
        if source_seg is None:
            continue
        seg.source_plain = source_seg.source_plain
        seg.target_tagged = seg.paragraph_ref.text
        aligned.append(seg)

    cfg = PipelineConfig(
        include_headers=True,
        include_footers=True,
        layout_check=True,
        layout_auto_fix=True,
        layout_auto_fix_passes=3,
        layout_font_reduction_pt=0.6,
        layout_spacing_factor=0.83,
    )
    total_issues = 0
    total_fixes = 0
    for _ in range(cfg.layout_auto_fix_passes):
        issues = validate_layout(translated_doc, aligned, cfg)
        total_issues += len(issues)
        fixed = fix_expansion_issues(aligned, issues, cfg)
        total_fixes += fixed
        if fixed == 0:
            break
    translated_doc.save(str(translated_docx))
    try:
        com_stats = update_fields_via_com(
            translated_docx,
            autofit_textboxes=True,
            min_font_size_pt=7.5,
            max_shrink_steps=6,
            expand_overflowing=False,
            max_height_growth=1.3,
        )
    except Exception as exc:
        com_stats = {
            "fields_updated": 0,
            "tocs_updated": 0,
            "textboxes_seen": 0,
            "textboxes_autofit": 0,
            "textboxes_shrunk": 0,
            "textboxes_expanded": 0,
            "com_autofit_error": str(exc),
        }
    return {
        "issues": total_issues,
        "auto_fixed_segments": total_fixes,
        **com_stats,
    }


def run_translation(source: Path, output: Path, report: Path) -> dict[str, object]:
    output.parent.mkdir(parents=True, exist_ok=True)
    report.parent.mkdir(parents=True, exist_ok=True)

    unique_strings = collect_unique_english_strings(source)
    mapping = build_translation_map(unique_strings)
    unresolved = find_unresolved(unique_strings, mapping)

    shutil.copyfile(source, output)
    direct_counts = replace_segments_direct(output, mapping)
    cleanup_counts = cleanup_remaining_paragraphs(output)
    layout_stats = auto_fix_layout(source, output)
    table_font_floor_runs = enforce_table_readability(output)
    header_footer_patch_files = patch_header_footer_textboxes(output)
    document_xml_patch_files = patch_document_xml_fragments(output)

    report_payload = {
        "source": str(source),
        "output": str(output),
        "translation_context": {
            "glossary": str(ROOT / "for_test" / "new_formating" / "glossary.md"),
            "general_prompt": str(ROOT / "for_test" / "new_formating" / "general_prompt.md"),
        },
        "unique_english_strings": len(unique_strings),
        "mapped_strings": len(mapping),
        "unresolved_strings": unresolved,
        "replace_hits": int(sum(direct_counts.values())),
        "replaced_keys": int(len(direct_counts)),
        "cleanup_hits": int(sum(cleanup_counts.values())),
        "table_font_floor_runs": int(table_font_floor_runs),
        "header_footer_patch_files": int(header_footer_patch_files),
        "document_xml_patch_files": int(document_xml_patch_files),
        "layout": layout_stats,
    }
    report.write_text(json.dumps(report_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return report_payload


def main() -> int:
    parser = argparse.ArgumentParser(description="Translate original_new_part1.docx into Russian with formatting preserved.")
    parser.add_argument(
        "--source",
        default=str(ROOT / "for_test" / "new_formating" / "section" / "original_new_part1.docx"),
    )
    parser.add_argument(
        "--output",
        default=str(ROOT / "for_test" / "new_formating" / "section_translate" / "original_new_part1_ru.docx"),
    )
    parser.add_argument(
        "--report",
        default=str(ROOT / "tmp" / "docs" / "part1_manual_translation_report.json"),
    )
    args = parser.parse_args()

    source = Path(args.source).resolve()
    output = Path(args.output).resolve()
    report = Path(args.report).resolve()
    report_payload = run_translation(source, output, report)
    print(json.dumps(report_payload, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

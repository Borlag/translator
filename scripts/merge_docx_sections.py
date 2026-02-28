from __future__ import annotations

import argparse
import html
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document


def _import_renderer():
    scripts_dir = Path(__file__).resolve().parent
    sys.path.insert(0, str(scripts_dir))
    try:
        from render_docx_pages import render_docx_pages  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(f"Cannot import renderer from {scripts_dir}: {exc}") from exc
    return render_docx_pages


def _import_section_merge_module():
    root = Path(__file__).resolve().parents[1]
    src_dir = root / "src"
    sys.path.insert(0, str(src_dir))
    try:
        import docxru.section_merge as section_merge  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(f"Cannot import docxru.section_merge from {src_dir}: {exc}") from exc
    return section_merge


def _write_compare_index_html(
    output_dir: Path,
    *,
    left_label: str,
    right_label: str,
    left_dir: Path,
    right_dir: Path,
) -> Path:
    left_pages = sorted(left_dir.glob("page-*.png"))
    right_pages = sorted(right_dir.glob("page-*.png"))
    page_count = max(len(left_pages), len(right_pages))

    def rel(path: Path) -> str:
        return path.relative_to(output_dir).as_posix()

    rows: list[str] = []
    for idx in range(page_count):
        left_img = rel(left_pages[idx]) if idx < len(left_pages) else ""
        right_img = rel(right_pages[idx]) if idx < len(right_pages) else ""
        page_no = idx + 1
        rows.append(
            "<div class='row'>"
            f"<div class='cell'><div class='cap'>{html.escape(left_label)} p.{page_no}</div>"
            + (f"<img loading='lazy' src='{html.escape(left_img)}' />" if left_img else "<div class='missing'>missing</div>")
            + "</div>"
            f"<div class='cell'><div class='cap'>{html.escape(right_label)} p.{page_no}</div>"
            + (
                f"<img loading='lazy' src='{html.escape(right_img)}' />"
                if right_img
                else "<div class='missing'>missing</div>"
            )
            + "</div>"
            + "</div>"
        )

    css = """
    :root { --gap: 16px; --bg: #111; --fg: #eee; --muted: #999; }
    body { margin: 0; font-family: Arial, sans-serif; background: var(--bg); color: var(--fg); }
    header { position: sticky; top: 0; background: rgba(17,17,17,0.95); padding: 12px 16px; border-bottom: 1px solid #2a2a2a; }
    header .path { color: var(--muted); font-size: 12px; margin-top: 4px; word-break: break-all; }
    .wrap { padding: 16px; }
    .row { display: grid; grid-template-columns: 1fr 1fr; gap: var(--gap); margin-bottom: 18px; }
    .cell { background: #171717; border: 1px solid #2a2a2a; border-radius: 8px; padding: 10px; }
    .cap { font-size: 12px; color: var(--muted); margin-bottom: 8px; }
    img { width: 100%; height: auto; display: block; border-radius: 4px; }
    .missing { padding: 24px; text-align: center; color: var(--muted); border: 1px dashed #333; border-radius: 6px; }
    @media (max-width: 1000px) { .row { grid-template-columns: 1fr; } }
    """

    html_out = (
        "<!doctype html><html><head><meta charset='utf-8' />"
        "<meta name='viewport' content='width=device-width, initial-scale=1' />"
        f"<title>DOCX Compare: {html.escape(left_label)} vs {html.escape(right_label)}</title>"
        f"<style>{css}</style></head><body>"
        "<header>"
        f"<div><strong>{html.escape(left_label)}</strong> vs <strong>{html.escape(right_label)}</strong></div>"
        f"<div class='path'>{html.escape(str(output_dir.resolve()))}</div>"
        "</header>"
        "<div class='wrap'>"
        + "\n".join(rows)
        + "</div></body></html>"
    )
    out_path = output_dir / "index.html"
    out_path.write_text(html_out, encoding="utf-8")
    return out_path


def _preview(args: argparse.Namespace) -> int:
    render_docx_pages = _import_renderer()
    sm = _import_section_merge_module()

    base_path = Path(args.base).resolve()
    overlay_path = Path(args.overlay).resolve()
    output_dir = Path(args.output_dir).resolve()

    text_pages_dir = output_dir / "pages_text"
    pic_pages_dir = output_dir / "pages_pics"
    text_pages_dir.mkdir(parents=True, exist_ok=True)
    pic_pages_dir.mkdir(parents=True, exist_ok=True)

    render_docx_pages(base_path, text_pages_dir, backend=str(args.backend), dpi=int(args.dpi), keep_pdf=True)
    render_docx_pages(overlay_path, pic_pages_dir, backend=str(args.backend), dpi=int(args.dpi), keep_pdf=True)

    index_html = _write_compare_index_html(
        output_dir,
        left_label=str(args.left_label),
        right_label=str(args.right_label),
        left_dir=text_pages_dir,
        right_dir=pic_pages_dir,
    )

    base_doc = Document(str(base_path))
    overlay_doc = Document(str(overlay_path))
    mapping_result = sm.build_page_mapping(base_doc, overlay_doc, mode=str(args.mapping_mode))
    mapping_json = (
        Path(args.mapping_json).resolve()
        if args.mapping_json
        else (output_dir / "page_mapping.auto.json").resolve()
    )
    sm.save_page_mapping_json(
        mapping_json,
        mapping_result,
        base_path=base_path,
        overlay_path=overlay_path,
    )

    print(f"Preview HTML: {index_html}")
    print(f"Auto mapping JSON: {mapping_json}")
    print(
        "Sections/pages: "
        f"base={mapping_result.base_sections}, overlay={mapping_result.overlay_pages}"
    )
    print(f"Mapping mode: {mapping_result.mode}")
    return 0


def _merge(args: argparse.Namespace) -> int:
    sm = _import_section_merge_module()

    base_path = Path(args.base).resolve()
    overlay_path = Path(args.overlay).resolve()
    output_path = Path(args.output).resolve()

    pages = sm.parse_pages_spec(str(args.pages))
    if not pages:
        raise SystemExit("No pages selected. Use --pages, e.g. 5,12,45-50")

    base_doc = Document(str(base_path))
    overlay_doc = Document(str(overlay_path))

    if args.mapping_json:
        mapping_json_path = Path(args.mapping_json).resolve()
        mapping = sm.load_page_mapping_json(mapping_json_path)
        mapping_result = None
    else:
        mapping_result = sm.build_page_mapping(base_doc, overlay_doc, mode=str(args.mapping_mode))
        mapping = mapping_result.mapping

    if args.mapping_out:
        mapping_out_path = Path(args.mapping_out).resolve()
        if mapping_result is not None:
            sm.save_page_mapping_json(
                mapping_out_path,
                mapping_result,
                base_path=base_path,
                overlay_path=overlay_path,
            )
        else:
            payload: dict[str, Any] = {
                "mapping": {str(k): v for k, v in sorted(mapping.items())},
            }
            mapping_out_path.parent.mkdir(parents=True, exist_ok=True)
            mapping_out_path.write_text(
                json.dumps(payload, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
        print(f"Mapping JSON written: {mapping_out_path}")

    if args.dry_run:
        overlay_sections = len(sm.get_section_ranges(overlay_doc))
        base_sections = len(sm.get_section_ranges(base_doc))
        selected = [page for page in pages if 1 <= page <= overlay_sections]
        covered = sum(len(mapping.get(page, [])) for page in selected)
        print("Dry run only. No file changes.")
        print(f"Selected pages: {selected}")
        print(f"Base sections total: {base_sections}")
        print(f"Overlay pages total: {overlay_sections}")
        print(f"Mapped base sections (selected pages): {covered}")
        return 0

    stats = sm.replace_pages(
        base_doc,
        overlay_doc,
        pages=pages,
        page_mapping=mapping,
        keep_page_size=str(args.keep_page_size),
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    base_doc.save(str(output_path))

    print(f"Merged DOCX: {output_path}")
    print(f"Replaced pages: {stats['replaced_pages']}")
    print(f"Replaced pages count: {stats['replaced_pages_count']}")
    print(f"Replaced base sections count: {stats['replaced_base_sections_count']}")
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Preview and merge section-level content between two DOCX files."
    )
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_preview = sub.add_parser("preview", help="Render side-by-side preview and generate auto mapping JSON.")
    p_preview.add_argument("--base", required=True, help="Path to base DOCX (e.g. translated_1.docx)")
    p_preview.add_argument("--overlay", required=True, help="Path to overlay DOCX (e.g. TRANSLATED_PICTURES_1.docx)")
    p_preview.add_argument("--output-dir", required=True, help="Output directory for rendered pages and HTML")
    p_preview.add_argument("--backend", default="auto", choices=["auto", "soffice", "word"], help="DOCX->PDF backend")
    p_preview.add_argument("--dpi", type=int, default=150, help="Render DPI for PNG pages")
    p_preview.add_argument("--mapping-json", default=None, help="Optional output path for mapping JSON")
    p_preview.add_argument(
        "--mapping-mode",
        default="dp_one_to_one",
        choices=["dp_one_to_one", "greedy_group"],
        help="Auto mapping strategy",
    )
    p_preview.add_argument("--left-label", default="translated_1", help="Left label in HTML compare")
    p_preview.add_argument("--right-label", default="TRANSLATED_PICTURES_1", help="Right label in HTML compare")

    p_merge = sub.add_parser("merge", help="Replace selected base pages using overlay sections.")
    p_merge.add_argument("--base", required=True, help="Path to base DOCX (editable structure source)")
    p_merge.add_argument("--overlay", required=True, help="Path to overlay DOCX (images/annotations source)")
    p_merge.add_argument("--pages", required=True, help="Page list, e.g. 5,12,45-50,100")
    p_merge.add_argument("--output", required=True, help="Output merged DOCX path")
    p_merge.add_argument("--mapping-json", default=None, help="Optional manual mapping JSON path")
    p_merge.add_argument(
        "--mapping-mode",
        default="dp_one_to_one",
        choices=["dp_one_to_one", "greedy_group"],
        help="Auto mapping strategy (used when --mapping-json is not provided)",
    )
    p_merge.add_argument("--mapping-out", default=None, help="Optional path to write mapping used by merge")
    p_merge.add_argument("--keep-page-size", default="target", choices=["target", "source"], help="Keep page size/margins from target or source sectPr")
    p_merge.add_argument("--dry-run", action="store_true", help="Validate inputs and print plan, do not save DOCX")

    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    if args.cmd == "preview":
        return _preview(args)
    if args.cmd == "merge":
        return _merge(args)
    raise SystemExit(f"Unknown command: {args.cmd}")


if __name__ == "__main__":
    raise SystemExit(main())

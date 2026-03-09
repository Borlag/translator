from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.render_docx_pages import render_docx_pages


LATIN_TOKEN_RE = re.compile(r"[A-Za-z][A-Za-z0-9./()_-]*")
MONTH_RE = re.compile(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\.?")
CODE_RE = re.compile(r"(?:[A-Z]{1,4}|[A-Z]{1,8}[0-9][0-9A-Z./-]*|[A-Z]{1,8}[./-][0-9A-Z./-]*)")
SINGLE_LETTER_SUFFIX_RE = re.compile(r"^[A-Z]\)$")
SINGLE_UPPER_RE = re.compile(r"^[A-Z]$")
PARA_REF_RE = re.compile(r"^[A-Z]\.\(\d+\)\.?$")
LOWER_PREFIX_CODE_RE = re.compile(r"^[a-z]\)[A-Z0-9.-]+$")
LOWER_LIST_MARKER_RE = re.compile(r"^[a-z]$")
MAGNIFICATION_RE = re.compile(r"^\d+[xX]\.?$")
TEMP_RANGE_RE = re.compile(r"^(?:\d+)?o[CF]-(?:\d+)?o[CF]\.?$")
ALLOW = {
    "SAFRAN",
    "Safran",
    "Landing",
    "LANDING",
    "Systems",
    "SYSTEMS",
    "UK",
    "LTD",
    "Ltd",
    "LIMITED",
    "CAGE",
    "K0654",
    "Airbus",
    "MLG",
    "ECCN",
    "EAR",
    "Cheltenham",
    "Road",
    "Gloucester",
    "England",
    "Messier-Dowty",
    "Messier-",
    "Limited",
    "www.safran-landing-systems.com",
    "mm",
    "in",
    "oC",
    "oF",
    "Accomet",
    "Araldite",
    "Alocrom",
    "Fibreslip",
    "Loctite",
    "Almen",
    "Mastinox",
    "Messier",
    "Dowty",
    "Molykote",
    "MOLYKOTE",
    "Sermetel",
    "SERMETEL",
    "LOCTITE",
    "MPa",
    "ksi",
}


def _normalize_token(token: str) -> str:
    normalized = token.strip(".,;:")
    if normalized.startswith("("):
        normalized = normalized[1:]
    if normalized.endswith(")") and "(" not in normalized:
        normalized = normalized[:-1]
    return normalized


def _suspicious_tokens(text: str) -> list[str]:
    tokens = sorted(set(LATIN_TOKEN_RE.findall(text)))
    out: list[str] = []
    for token in tokens:
        normalized = _normalize_token(token)
        if not normalized:
            continue
        if token in ALLOW or normalized in ALLOW:
            continue
        if MONTH_RE.fullmatch(normalized):
            continue
        if CODE_RE.fullmatch(normalized):
            continue
        if SINGLE_UPPER_RE.fullmatch(normalized):
            continue
        if SINGLE_LETTER_SUFFIX_RE.fullmatch(token):
            continue
        if PARA_REF_RE.fullmatch(normalized):
            continue
        if LOWER_PREFIX_CODE_RE.fullmatch(normalized):
            continue
        if LOWER_LIST_MARKER_RE.fullmatch(normalized):
            continue
        if MAGNIFICATION_RE.fullmatch(normalized):
            continue
        if TEMP_RANGE_RE.fullmatch(normalized):
            continue
        out.append(token)
    return out


def _collect_docx_hits(docx_path: Path) -> list[dict[str, object]]:
    doc = Document(str(docx_path))
    hits: list[dict[str, object]] = []
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        tokens = _suspicious_tokens(text)
        if tokens:
            hits.append({"location": f"P{i}", "tokens": tokens, "text": text})
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = "\n".join(p.text for p in cell.paragraphs).strip()
                if not text:
                    continue
                tokens = _suspicious_tokens(text)
                if tokens:
                    hits.append({"location": f"T{ti}R{ri}C{ci}", "tokens": tokens, "text": text})
    return hits


def _collect_pdf_hits(pdf_path: Path) -> list[dict[str, object]]:
    import fitz

    doc = fitz.open(pdf_path)
    hits: list[dict[str, object]] = []
    for i in range(doc.page_count):
        text = doc.load_page(i).get_text("text")
        tokens = _suspicious_tokens(text)
        if tokens:
            hits.append({"page": i + 1, "tokens": tokens})
    return hits


def main() -> int:
    parser = argparse.ArgumentParser(description="Review a translated CMM DOCX for residual English and render issues.")
    parser.add_argument("--docx", required=True, help="Translated DOCX to inspect")
    parser.add_argument("--render-dir", required=True, help="Directory for rendered pages and review artifacts")
    parser.add_argument("--report", required=True, help="Path to the JSON review report")
    parser.add_argument("--backend", default="word", choices=["auto", "soffice", "word"])
    parser.add_argument("--dpi", type=int, default=120)
    args = parser.parse_args()

    docx_path = Path(args.docx).resolve()
    render_dir = Path(args.render_dir).resolve()
    report_path = Path(args.report).resolve()
    render_dir.mkdir(parents=True, exist_ok=True)
    report_path.parent.mkdir(parents=True, exist_ok=True)

    render_docx_pages(docx_path, render_dir, backend=args.backend, dpi=args.dpi, keep_pdf=True)
    pdf_path = render_dir / f"{docx_path.stem}.pdf"

    docx_hits = _collect_docx_hits(docx_path)
    pdf_hits = _collect_pdf_hits(pdf_path)

    payload = {
        "docx": str(docx_path),
        "pdf": str(pdf_path),
        "docx_hits": docx_hits,
        "pdf_hits": pdf_hits,
    }
    report_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W = f"{{{NS['w']}}}"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

MONTHS = {
    "Jan": "янв",
    "Feb": "фев",
    "Mar": "мар",
    "Apr": "апр",
    "May": "мая",
    "Jun": "июн",
    "Jul": "июл",
    "Aug": "авг",
    "Sep": "сен",
    "Oct": "окт",
    "Nov": "ноя",
    "Dec": "дек",
}

ALLOWED_LATIN_RE = re.compile(
    r"https?://|www\.|"
    r"\b(?:SAFRAN|Safran|CAGE|ECCN|EAR|CAA|MLG|NLG|UK|Ltd|GL2|QH|IFC|PCS|PR|"
    r"Airbus|Messier-Dowty|Landing|Systems|Cheltenham|Road|Gloucester|England|"
    r"Export|Administration|Regulations|Act|Airworthiness|Limitations|Section|Document|ALS)\b|"
    r"\b(?:EDES\d|MAF\d)[A-Z0-9-]*\b|"
    r"\b[A-Z]{1,4}-\d{2,}[A-Z0-9.-]*\b|"
    r"\bK\d{4}\b"
)


def collapse_ws(text: str) -> str:
    return " ".join((text or "").replace("\r", " ").replace("\n", " ").replace("\t", " ").split()).strip()


def normalize_lines(text: str) -> str:
    lines = [" ".join(line.split()) for line in (text or "").replace("\r\n", "\n").split("\n")]
    return "\n".join(lines).strip()


def paragraph_text(paragraph: etree._Element) -> str:
    parts: list[str] = []
    for node in paragraph.iter():
        if node is paragraph:
            continue

        parent = node.getparent()
        nested = False
        while parent is not None and parent is not paragraph:
            if parent.tag == W + "p":
                nested = True
                break
            parent = parent.getparent()
        if nested:
            continue

        if node.tag == W + "t":
            parts.append(node.text or "")
        elif node.tag == W + "tab":
            parts.append("\t")
        elif node.tag == W + "br":
            br_type = node.get(W + "type")
            if br_type in {None, "", "textWrapping"}:
                parts.append("\n")
    return "".join(parts)


def own_inline_nodes(paragraph: etree._Element) -> list[etree._Element]:
    nodes: list[etree._Element] = []
    for node in paragraph.iter():
        if node is paragraph:
            continue
        if node.tag not in {W + "t", W + "tab", W + "br"}:
            continue
        parent = node.getparent()
        nested = False
        while parent is not None and parent is not paragraph:
            if parent.tag == W + "p":
                nested = True
                break
            parent = parent.getparent()
        if not nested:
            nodes.append(node)
    return nodes


def first_text_run(paragraph: etree._Element) -> etree._Element | None:
    for node in own_inline_nodes(paragraph):
        run = node.getparent()
        if run is not None and run.tag == W + "r":
            return run
    return None


def set_run_text(run: etree._Element, text: str) -> None:
    insert_at = 0
    for idx, child in enumerate(list(run)):
        if child.tag == W + "rPr":
            insert_at = idx + 1

    while len(run) > insert_at:
        run.remove(run[insert_at])

    chunks: list[tuple[str, str]] = []
    buf: list[str] = []
    for ch in text:
        if ch in {"\n", "\t"}:
            if buf:
                chunks.append(("text", "".join(buf)))
                buf = []
            chunks.append(("br" if ch == "\n" else "tab", ch))
        else:
            buf.append(ch)
    if buf:
        chunks.append(("text", "".join(buf)))

    if not chunks:
        t = etree.Element(W + "t")
        run.insert(insert_at, t)
        return

    for kind, value in chunks:
        if kind == "text":
            t = etree.Element(W + "t")
            t.text = value
            if value[:1].isspace() or value[-1:].isspace():
                t.set(XML_SPACE, "preserve")
            run.insert(insert_at, t)
        elif kind == "tab":
            run.insert(insert_at, etree.Element(W + "tab"))
        else:
            run.insert(insert_at, etree.Element(W + "br"))
        insert_at += 1


def set_paragraph_text(paragraph: etree._Element, text: str) -> bool:
    run = first_text_run(paragraph)
    if run is None:
        return False

    nodes = own_inline_nodes(paragraph)
    keep = set(run.iterfind(".//w:rPr", namespaces=NS))
    for node in nodes:
        if node in keep:
            continue
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)

    set_run_text(run, text)
    return True


def parse_glossary(glossary_path: Path) -> list[tuple[str, str]]:
    text = glossary_path.read_text(encoding="utf-8")
    pairs: list[tuple[str, str]] = []
    for line in text.splitlines():
        if not line.startswith("|"):
            continue
        parts = [part.strip() for part in line.strip().strip("|").split("|")]
        if len(parts) < 2:
            continue
        en, ru = parts[0], parts[1]
        if en in {"English", "Standard"} or ru in {"Русский термин", "Назначение"}:
            continue
        if set(en) == {"-"}:
            continue
        if en and ru:
            pairs.append((en, ru))
    return pairs


def build_exact_map() -> dict[str, str]:
    samples_dir = Path(r"C:\Users\Urdul\Desktop\project\translator\samples")
    output_dir = Path(r"C:\Users\Urdul\Desktop\project\translator\output\doc")
    sample_src = next(samples_dir.glob("*abby_short.docx"))
    sample_ru = next(output_dir.glob("*consistency_v5.docx"))

    src_doc = Document(str(sample_src))
    ru_doc = Document(str(sample_ru))

    source_items: list[str] = []
    target_items: list[str] = []

    for paragraph in src_doc.paragraphs:
        source_items.append(paragraph.text)
    for table in src_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                source_items.append(cell.text)

    for paragraph in ru_doc.paragraphs:
        target_items.append(paragraph.text)
    for table in ru_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                target_items.append(cell.text)

    exact: dict[str, str] = {}
    for source_text, target_text in zip(source_items, target_items):
        source_key = collapse_ws(source_text)
        target_key = collapse_ws(target_text)
        if not source_key or source_key == target_key:
            continue
        if re.search(r"[A-Za-z]{3,}", source_key):
            exact[source_key] = normalize_lines(target_text)
    return exact


MANUAL_EXACT = {
    collapse_ws("COMPONENT MAINTENANCE MANUAL WITH"): "РУКОВОДСТВО ПО ТЕХНИЧЕСКОМУ ОБСЛУЖИВАНИЮ КОМПОНЕНТА С",
    collapse_ws(
        "NOTE: The above certification does not apply to revisions or amendments made after the date of initial certification by other Approved Organisations. Revisions or Amendments made by other Approved Organisations must each be separately certified and recorded on separate record sheets."
    ): (
        "ПРИМЕЧАНИЕ. Указанная выше сертификация не распространяется на редакции или поправки, "
        "внесенные после даты первоначальной сертификации другими утвержденными организациями. "
        "Каждая редакция или поправка, внесенная другой утвержденной организацией, должна быть "
        "сертифицирована отдельно и зарегистрирована на отдельном листе учета."
    ),
    collapse_ws(
        "The technical data in this document (or file) may contain US data and be controlled for export under the Export Administration Regulations (EAR), 15 CFR Parts 730-774, ECCN: 9E991. Violations of these laws may be subject to fines and penalties under the Export Administration Act."
    ): (
        "Технические данные в этом документе (или файле) могут содержать данные США и подпадать под "
        "экспортный контроль в соответствии с Export Administration Regulations (EAR), 15 CFR Parts "
        "730-774, ECCN: 9E991. Нарушение этих требований может повлечь штрафы и санкции в "
        "соответствии с Export Administration Act."
    ),
    collapse_ws(
        "No intellectual property rights are granted by the delivery of this document or the disclosure of its content. This document shall not be reproduced to a third party without the express written consent of Safran Landing Systems (and/or the appropriate affiliated company)."
    ): (
        "Никакие права интеллектуальной собственности не предоставляются путем передачи этого "
        "документа или раскрытия его содержания. Настоящий документ не подлежит воспроизведению "
        "для третьих лиц без прямого письменного согласия Safran Landing Systems "
        "(и/или соответствующей аффилированной компании)."
    ),
    collapse_ws(
        "PART No. 201587001 AND 201587002 COMPONENT MAINTENANCE MANUAL MAIN LANDING GEAR LEG"
    ): (
        "Номер детали 201587001 И 201587002 РУКОВОДСТВО ПО ТЕХНИЧЕСКОМУ ОБСЛУЖИВАНИЮ "
        "КОМПОНЕНТА ОСНОВНАЯ СТОЙКА ШАССИ"
    ),
    collapse_ws(
        "COMPONENT MAINTENANCE MANUAL 32-12-22 MAIN LANDING GEAR LEG"
    ): "РУКОВОДСТВО ПО ТЕХНИЧЕСКОМУ ОБСЛУЖИВАНИЮ КОМПОНЕНТА 32-12-22 ОСНОВНАЯ СТОЙКА ШАССИ",
    collapse_ws(
        "Record the issue date and insertion date of this revision in the Record of Revisions and retain this Letter of Transmittal."
    ): (
        "Запишите дату выпуска и дату внесения этой редакции в запись изменений и сохраните это "
        "сопроводительное письмо."
    ),
    collapse_ws("LIST OF EFFECTIVE PAGES"): "ПЕРЕЧЕНЬ СТРАНИЦ",
    collapse_ws("Remove and Destroy Pages"): "Удалить и уничтожить страницы",
    collapse_ws("Unit Identification Chart"): "ТАБЛИЦА ИДЕНТИФИКАЦИИ АГРЕГАТА",
    collapse_ws("Added Ref. Codes 2253 and 2255 details"): "Добавлены сведения по кодам 2253 и 2255",
    collapse_ws("Added content. Updated page numbers. Updated figure numbers"): (
        "Добавлено содержание. Обновлены номера страниц и рисунков."
    ),
    collapse_ws("Updated tables 501 and 502"): "Обновлены таблицы 501 и 502",
    collapse_ws("Updated table 601. Updated caution at para"): "Обновлена таблица 601. Обновлено предупреждение в пункте",
    collapse_ws("Updated figure titles. Deleted figures 626, 627, 649, 650, 653"): (
        "Обновлены заголовки рисунков. Удалены рисунки 626, 627, 649, 650, 653"
    ),
    collapse_ws("and 654. Updated figure 626. Added figures 642"): (
        "и 654. Обновлен рисунок 626. Добавлены рисунки 642"
    ),
    collapse_ws("648. Updated table 602. Updated figure numbers"): (
        "648. Обновлена таблица 602. Обновлены номера рисунков"
    ),
    collapse_ws("Safran Landing Systems UK Ltd"): "Safran Landing Systems UK Ltd",
    collapse_ws("Telephone: +44 (0) 1452 712424 Fax: +44 (0) 1452 713821"): (
        "Телефон: +44 (0) 1452 712424 Факс: +44 (0) 1452 713821"
    ),
    collapse_ws("Cheltenham Road, Gloucester, GL2 9QH, England Telephone: +44 (0) 1452 712424 Fax: +44 (0) 1452 713821"): (
        "Cheltenham Road, Gloucester, GL2 9QH, England Телефон: +44 (0) 1452 712424 Факс: +44 (0) 1452 713821"
    ),
    collapse_ws("CAGE: K0654"): "CAGE: K0654",
    collapse_ws("REV NO."): "РЕД. №",
    collapse_ws("ISSUE DATE"): "ДАТА ВЫПУСКА",
    collapse_ws("INSERTION DATE"): "ДАТА ВНЕСЕНИЯ",
    collapse_ws("INCORPORATED"): "ВНЕСЕНА",
    collapse_ws("REV."): "РЕД.",
    collapse_ws("DATE INSERTED"): "ДАТА ВНЕСЕНИЯ",
    collapse_ws("PAGE NUMBER"): "НОМЕР СТРАНИЦЫ",
    collapse_ws("DATE REMOVED"): "ДАТА УДАЛЕНИЯ",
    collapse_ws("DASH NO."): "№ ИСПОЛНЕНИЯ",
    collapse_ws("SERVICE BULLETIN NUMBER"): "НОМЕР СЕРВИСНОГО БЮЛЛЕТЕНЯ",
    collapse_ws("MOD. STRIKE NO."): "№ ОТМЕТКИ О МОД.",
    collapse_ws("SAFRAN LANDING SYSTEMS MODIFICATION NUMBER"): "НОМЕР МОДИФИКАЦИИ SAFRAN LANDING SYSTEMS",
    collapse_ws("SAFRAN LANDING SYSTEMS SERVICE BULLETIN NUMBER"): (
        "НОМЕР СЕРВИСНОГО БЮЛЛЕТЕНЯ SAFRAN LANDING SYSTEMS"
    ),
    collapse_ws("TITLE PAGE"): "ТИТУЛЬНАЯ СТРАНИЦА",
    collapse_ws("Subject"): "Тема",
    collapse_ws("Page"): "Страница",
    collapse_ws("Date"): "Дата",
    collapse_ws("Blank"): "Пусто",
    collapse_ws("Record of Temporary"): "Запись временных",
    collapse_ws("Revisions"): "изменений",
    collapse_ws("List of Service Bulletins"): "Список сервисных бюллетеней",
    collapse_ws("List of Effective"): "Перечень",
    collapse_ws("Pages (Continued)"): "страниц (продолжение)",
    collapse_ws("Description and"): "Описание и",
    collapse_ws("Operation"): "работа",
    collapse_ws("Testing and Fault"): "Испытания и локализация",
    collapse_ws("Isolation"): "неисправностей",
    collapse_ws("Isolation (Continued)"): "локализация неисправностей (продолжение)",
    collapse_ws("(Continued)"): "(продолжение)",
    collapse_ws("Cleaning"): "Очистка",
}


MANUAL_PHRASES = [
    ("COMPONENT MAINTENANCE MANUAL", "РУКОВОДСТВО ПО ТЕХНИЧЕСКОМУ ОБСЛУЖИВАНИЮ КОМПОНЕНТА"),
    ("ILLUSTRATED PARTS LIST", "ИЛЛЮСТРИРОВАННЫЙ ПЕРЕЧЕНЬ ДЕТАЛЕЙ"),
    ("STATEMENT OF INITIAL CERTIFICATION", "ЗАЯВЛЕНИЕ О ПЕРВОНАЧАЛЬНОЙ СЕРТИФИКАЦИИ"),
    ("LIST OF EFFECTIVE PAGES", "ПЕРЕЧЕНЬ СТРАНИЦ"),
    ("TABLE OF CONTENTS", "СОДЕРЖАНИЕ"),
    ("RECORD OF REVISIONS", "ЗАПИСЬ ИЗМЕНЕНИЙ"),
    ("RECORD OF TEMPORARY REVISIONS", "ЗАПИСЬ ВРЕМЕННЫХ ИЗМЕНЕНИЙ"),
    ("LIST OF SERVICE BULLETINS", "СПИСОК СЕРВИСНЫХ БЮЛЛЕТЕНЕЙ"),
    ("UNIT IDENTIFICATION CHART", "ТАБЛИЦА ИДЕНТИФИКАЦИИ АГРЕГАТА"),
    ("DATE INCORPORATED INTO MANUAL", "ДАТА ВНЕСЕНИЯ В РУКОВОДСТВО"),
    ("INCORPORATED INTO MANUAL", "ВНЕСЕНА В РУКОВОДСТВО"),
    ("UNIT IDENTIFICATION CHART (Continued)", "ТАБЛИЦА ИДЕНТИФИКАЦИИ АГРЕГАТА (продолжение)"),
    ("LIST OF EFFECTIVE PAGES (Continued)", "ПЕРЕЧЕНЬ СТРАНИЦ (продолжение)"),
    ("TABLE OF CONTENTS (Continued)", "СОДЕРЖАНИЕ (продолжение)"),
    ("ILLUSTRATIONS (Continued)", "ИЛЛЮСТРАЦИИ (продолжение)"),
    ("(Continued)", "(продолжение)"),
    ("Description and Operation", "Описание и работа"),
    ("Description", "Описание"),
    ("Operation", "Работа"),
    ("Testing and Fault Isolation", "Испытания и локализация неисправностей"),
    ("Assembly (Including Storage)", "Сборка (включая хранение)"),
    ("Including Storage", "включая хранение"),
    ("Special Tools, Fixtures and Equipment", "Специальные инструменты, приспособления и оборудование"),
    ("Special Tools, Fixtures", "Специальные инструменты, приспособления"),
    ("Detailed Parts List", "Подробный перечень деталей"),
    ("Vendor Codes, Names and Addresses", "Коды поставщиков, наименования и адреса"),
    ("How To Use This Illustrated Parts List", "Как пользоваться этим иллюстрированным перечнем деталей"),
    ("Numerical Index", "Числовой указатель"),
    ("Detailed Inspection", "Детальный осмотр"),
    ("Special Detailed Inspection", "Специальный детальный осмотр"),
    ("Repair Procedure Conditions", "Условия выполнения процедуры ремонта"),
    ("Protective Treatment", "Защитная обработка"),
    ("Approved Repairs", "Утвержденные ремонты"),
    ("Approved Repairs - Key Diagram", "Утвержденные ремонты - ключевая схема"),
    ("Main Fitting Repairs - Key Diagram", "Ремонты корпуса стойки - ключевая схема"),
    ("Torque Link Repairs - Key Diagram", "Ремонты рычага крутящего момента - ключевая схема"),
    ("Sliding Tube Repairs - Key Diagram", "Ремонты скользящей трубы - ключевая схема"),
    ("Upper Diaphragm Tube Repairs - Key Diagram", "Ремонты верхней трубы диафрагмы - ключевая схема"),
    ("Cylinder Repairs - Key Diagram", "Ремонты цилиндра - ключевая схема"),
    ("Transfer Block Repairs - Key Diagram", "Ремонты переходного блока - ключевая схема"),
    ("Harness Support Bracket Repairs - Key Diagram", "Ремонты кронштейна опоры жгута - ключевая схема"),
    ("Upper Pivot Bracket Repairs - Key Diagram", "Ремонты верхнего кронштейна шарнира - ключевая схема"),
    ("Key Diagram", "ключевая схема"),
    ("Diagram of Operation", "Схема работы"),
    ("Reason for Change", "Причина изменения"),
    ("Subject Reference", "Тема/ссылка"),
    ("Remove and Destroy Pages", "Удалить и уничтожить страницы"),
    ("Insert New/Revised", "Вставить новые/пересмотренные"),
    ("Record of Revisions", "Запись изменений"),
    ("Issued by", "Выдано"),
    ("Letter of Transmittal No.", "Сопроводительное письмо №"),
    ("Repair No.", "Ремонт №"),
    ("Page", "Страница"),
    ("Fig.", "Рис."),
    ("Sheet", "Лист"),
    ("Title", "Наименование"),
    ("Unit Identification", "Идентификация агрегата"),
    ("Chart", "Таблица"),
    ("General", "Общие сведения"),
    ("Procedure", "Процедура"),
    ("Check", "Контроль"),
    ("Repair", "Ремонт"),
    ("Introduction", "Введение"),
    ("Data", "Данные"),
    ("Torque Data", "Данные по рычагу крутящего момента"),
    ("Definitions", "Определения"),
    ("Remarks", "Примечания"),
    ("ILLUSTRATIONS", "ИЛЛЮСТРАЦИИ"),
    ("Illustrations", "Иллюстрации"),
    ("Equipment and Materials", "Оборудование и материалы"),
    ("Equipment", "Оборудование"),
    ("Test Conditions", "Условия испытания"),
    ("Fault Isolation", "Локализация неисправностей"),
    ("Cleaning", "Очистка"),
    ("Storage", "хранение"),
    ("Subassembly", "подсборка"),
    ("Main Fitting Subassembly", "подсборка корпуса стойки"),
    ("Repair to Main Fitting", "Ремонт корпуса стойки"),
    ("Lower Bearing Subassembly", "узел нижнего подшипника"),
    ("Main Fitting", "Корпус стойки"),
    ("Pivot Pin", "Штифт шарнира"),
    ("Uplock Pin Main Fitting", "Штифт замка убранного положения корпуса стойки"),
    ("Uplock Pin", "Штифт замка убранного положения"),
    ("Sliding Tube", "Скользящая труба"),
    ("Upper Slave Link", "Верхнее ведомое звено"),
    ("Lower Slave Link", "Нижнее ведомое звено"),
    ("Upper Diaphragm Tube", "Верхняя труба диафрагмы"),
    ("Cylinder", "Цилиндр"),
    ("Lock Stay Cardan", "Кардан фиксатора"),
    ("Transfer Block", "Переходной блок"),
    ("Spherical Bearing", "Сферический подшипник"),
    ("Harness Support Bracket", "Кронштейн опоры жгута"),
    ("Retaining Pin", "Стопорный штифт"),
    ("Upper Pivot Bracket", "Верхний кронштейн шарнира"),
    ("Pivot Bracket", "Кронштейн шарнира"),
    ("Valve Stem", "Шток клапана"),
    ("Inflation Valve", "Клапан зарядки"),
    ("Forward Pintle Pin", "Передний шкворневой штифт"),
    ("Slave Link", "Ведомое звено"),
    ("Torque Link", "Рычаг крутящего момента"),
    ("Machining and Installation", "Механическая обработка и установка"),
    ("Liner Installation", "Установка вкладыша"),
    ("Inner Liner Installation", "Установка внутреннего вкладыша"),
    ("Installation", "Установка"),
    ("Oversize Bushes", "Ремонтные втулки"),
    ("Oversize Bush(es)", "Ремонтная втулка увеличенного размера"),
    ("Bushes", "Втулки"),
    ("Bush", "Втулка"),
    ("Lubrication adapter", "смазочный адаптер"),
    ("Transfer Dowel", "передаточный штифт"),
    ("Chromium Plate Termination", "окончание хромового покрытия"),
    ("Electrical Bonding Resistance Test Points", "Точки проверки сопротивления электрического соединения"),
    ("Tables", "таблицы"),
    ("Superseded", "заменен"),
    ("Withdrawn", "изъят"),
    ("Only", "только"),
    ("or", "или"),
    ("Added fig-item", "Добавлен элемент рисунка"),
    ("Updated fig-item", "Обновлен элемент рисунка"),
    ("Updated fig-items", "Обновлены элементы рисунка"),
    ("Updated material specification", "Обновлена спецификация материала"),
    ("Updated Messier-Dowty Limited - Safran Landing Systems", "Наименование Messier-Dowty Limited изменено на Safran Landing Systems"),
    ("Updated Messier-Dowty Limited - Safran Landing System", "Наименование Messier-Dowty Limited изменено на Safran Landing Systems"),
    ("Updated Messier-Dowty Limited to Safran Landing Systems", "Наименование Messier-Dowty Limited изменено на Safran Landing Systems"),
    ("Updated Messier-Dowty Limited to Safran Landing System", "Наименование Messier-Dowty Limited изменено на Safran Landing Systems"),
    ("Updated Messier-Dowty Limited to", "Наименование Messier-Dowty Limited изменено на"),
    ("Updated figure titles", "Обновлены заголовки рисунков"),
    ("Updated figure numbers", "Обновлены номера рисунков"),
    ("Updated figure", "Обновлен рисунок"),
    ("Updated figures", "Обновлены рисунки"),
    ("Deleted figures", "Удалены рисунки"),
    ("Added figures", "Добавлены рисунки"),
    ("Added figure", "Добавлен рисунок"),
    ("Added figure 713. Updated figures 705, 706, 707", "Добавлен рисунок 713. Обновлены рисунки 705, 706, 707"),
    ("Updated table", "Обновлена таблица"),
    ("Updated tables", "Обновлены таблицы"),
    ("Updated para", "Обновлен пункт"),
    ("Updated paras", "Обновлены пункты"),
    ("Added para", "Добавлен пункт"),
    ("Added paras", "Добавлены пункты"),
    ("Updated IPL figs", "Обновлены рисунки IPL"),
    ("Updated IPL fig", "Обновлен рисунок IPL"),
    ("include Ref. Codes", "включая коды ссылок"),
    ("to include Ref. Codes", "с включением кодов ссылок"),
    ("in para", "в пункте"),
    ("in paras", "в пунктах"),
    ("only in para", "только в пункте"),
]


def compile_source_pattern(source: str) -> re.Pattern[str]:
    pattern = re.escape(source).replace(r"\ ", r"\s+")
    return re.compile(rf"(?<![A-Za-z]){pattern}(?![A-Za-z])", re.IGNORECASE)


def compile_phrase_patterns(glossary_pairs: list[tuple[str, str]]) -> list[tuple[re.Pattern[str], str]]:
    pairs = [pair for pair in glossary_pairs + MANUAL_PHRASES if pair[0] and pair[1]]
    pairs.sort(key=lambda item: len(item[0]), reverse=True)
    return [(compile_source_pattern(source), target) for source, target in pairs]


def convert_dates(text: str) -> str:
    def repl(match: re.Match[str]) -> str:
        month = MONTHS[match.group(1)]
        day = int(match.group(2))
        year = match.group(3)
        return f"{day} {month} {year}"

    return re.sub(r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+(\d{1,2})/(\d{4})\b", repl, text)


def convert_ranges(text: str) -> str:
    text = re.sub(r"\b(\d+)\s+to\s+(\d+(?:\.\d+)?)\b", r"\1-\2", text, flags=re.IGNORECASE)
    text = re.sub(r"\b(\d+)\s+and\s+(\d+)\b", r"\1 и \2", text, flags=re.IGNORECASE)
    return text


def cleanup_translation(text: str) -> str:
    text = re.sub(r"\s+\n", "\n", text)
    text = re.sub(r"\n\s+", "\n", text)
    text = re.sub(r" {2,}", " ", text)
    text = text.replace("Lиing", "Landing")
    text = text.replace("LРёing", "Landing")
    text = text.replace("SubСборка", "подсборка")
    text = text.replace("Repairs", "ремонты")
    text = text.replace("S-rage", "хранение")
    text = text.replace("Special -ols", "Специальные инструменты")
    text = text.replace("MMain Fitting", "Main Fitting")
    text = text.replace("MMMain Fitting", "Main Fitting")
    text = text.replace("www.safran-lиing-systems.com", "www.safran-landing-systems.com")
    text = re.sub(r"\(\s*продолжение\s*\)", "(продолжение)", text, flags=re.IGNORECASE)
    text = re.sub(r"\bPage PAGE\b", "Страница", text, flags=re.IGNORECASE)
    text = re.sub(r"\bTITLE PAGE\b", "ТИТУЛЬНАЯ СТРАНИЦА", text, flags=re.IGNORECASE)
    text = re.sub(r"\bBlank\b", "Пусто", text, flags=re.IGNORECASE)
    text = re.sub(r"\bDATE INCORPORATED INTO MANUAL\b", "ДАТА ВНЕСЕНИЯ В РУКОВОДСТВО", text, flags=re.IGNORECASE)
    text = re.sub(r"\bTelephone\b", "Телефон", text, flags=re.IGNORECASE)
    text = re.sub(r"\bTITLE\s+Страница\b", "Наименование Страница", text, flags=re.IGNORECASE)
    text = re.sub(r"\bUNIT IDENTIFICATION CHART\b", "ТАБЛИЦА ИДЕНТИФИКАЦИИ АГРЕГАТА", text, flags=re.IGNORECASE)
    text = re.sub(r"\bILLUSTRATIONS\s+\(продолжение\)\b", "ИЛЛЮСТРАЦИИ (продолжение)", text, flags=re.IGNORECASE)
    text = re.sub(r"\bINCORPORATED INTO MANUAL\b", "ВНЕСЕНИЯ В РУКОВОДСТВО", text, flags=re.IGNORECASE)
    text = re.sub(r"\bRepair to Main Fitting\b", "Ремонт корпуса стойки", text, flags=re.IGNORECASE)
    text = re.sub(r"\bRepair to\s+Корпус стойки\b", "Ремонт корпуса стойки", text, flags=re.IGNORECASE)
    text = re.sub(r"\bMain Fitting\b", "Корпус стойки", text, flags=re.IGNORECASE)
    text = re.sub(r"\bIncluding\b", "включая", text, flags=re.IGNORECASE)
    text = re.sub(r"\bUpdated figure\b", "Обновлен рисунок", text, flags=re.IGNORECASE)
    text = re.sub(r"\bUpdated figures\b", "Обновлены рисунки", text, flags=re.IGNORECASE)
    text = re.sub(r"\bAdded figure\b", "Добавлен рисунок", text, flags=re.IGNORECASE)
    text = re.sub(r"\bfigure\b", "рисунок", text, flags=re.IGNORECASE)
    text = re.sub(r"\bfigures\b", "рисунки", text, flags=re.IGNORECASE)
    text = re.sub(r"\bpara\b", "пункт", text, flags=re.IGNORECASE)
    text = re.sub(r"\bRef\.\s+Codes\b", "коды ссылок", text, flags=re.IGNORECASE)
    text = re.sub(r"\bTables\b", "таблицы", text, flags=re.IGNORECASE)
    text = re.sub(r"\bContinued\b", "продолжение", text, flags=re.IGNORECASE)
    text = re.sub(r"\bInner Liner Installation\b", "Установка внутреннего вкладыша", text, flags=re.IGNORECASE)
    text = re.sub(r"\bRepair No\.\b", "Ремонт №", text, flags=re.IGNORECASE)
    return text.strip()
    text = re.sub(r"\(\s*продолжение\s*\)", "(продолжение)", text, flags=re.IGNORECASE)
    text = re.sub(r"\bPage PAGE\b", "Страница", text, flags=re.IGNORECASE)
    text = re.sub(r"\bTITLE PAGE\b", "ТИТУЛЬНАЯ СТРАНИЦА", text, flags=re.IGNORECASE)
    text = re.sub(r"\bBlank\b", "Пусто", text, flags=re.IGNORECASE)
    return text.strip()


def translate_text(
    text: str,
    *,
    exact_map: dict[str, str],
    phrase_patterns: list[tuple[re.Pattern[str], str]],
) -> str:
    key = collapse_ws(text)
    if key in exact_map:
        return exact_map[key]

    translated = normalize_lines(text)
    translated = translated.replace("Lиing", "Landing")
    translated = translated.replace("LРёing", "Landing")
    translated = translated.replace("S-rage", "Storage")
    translated = translated.replace("Special -ols", "Special Tools")
    translated = translated.replace("-rque Link", "Torque Link")
    translated = translated.replace("SubСборка", "Subassembly")
    translated = translated.replace("MMain Fitting", "Main Fitting")
    translated = convert_dates(translated)
    translated = convert_ranges(translated)

    for pattern, replacement in phrase_patterns:
        translated = pattern.sub(replacement, translated)

    translated = re.sub(r"\bPART No\.\b", "Номер детали", translated, flags=re.IGNORECASE)
    translated = re.sub(r"\bPART NUMBER\b", "НОМЕР ДЕТАЛИ", translated, flags=re.IGNORECASE)
    translated = re.sub(r"\bDASH NO\.\b", "№ ИСПОЛНЕНИЯ", translated, flags=re.IGNORECASE)
    translated = re.sub(r"\bWITH\b", "С", translated, flags=re.IGNORECASE)
    translated = re.sub(r"\band\b", "и", translated, flags=re.IGNORECASE)
    translated = re.sub(r"\bor\b", "или", translated, flags=re.IGNORECASE)
    translated = cleanup_translation(translated)
    return translated


def remaining_latin(text: str) -> str:
    return ALLOWED_LATIN_RE.sub("", text)


def process_docx(
    input_path: Path,
    output_path: Path,
    *,
    glossary_path: Path,
    audit_path: Path | None,
) -> tuple[int, int, list[str]]:
    exact_map = build_exact_map()
    exact_map.update(MANUAL_EXACT)
    phrase_patterns = compile_phrase_patterns(parse_glossary(glossary_path))

    changed = 0
    scanned = 0
    leftovers: list[str] = []
    leftover_seen: set[str] = set()

    with zipfile.ZipFile(input_path, "r") as zin:
        with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    try:
                        root = etree.fromstring(data)
                    except Exception:
                        zout.writestr(item, data)
                        continue

                    xml_changed = False
                    for paragraph in root.xpath(".//w:p", namespaces=NS):
                        source_text = paragraph_text(paragraph)
                        collapsed = collapse_ws(source_text)
                        if not collapsed:
                            continue
                        if not re.search(r"[A-Za-z]{3,}", collapsed):
                            continue

                        scanned += 1
                        translated = translate_text(
                            source_text,
                            exact_map=exact_map,
                            phrase_patterns=phrase_patterns,
                        )
                        if collapse_ws(translated) != collapsed:
                            if set_paragraph_text(paragraph, translated):
                                changed += 1
                                xml_changed = True

                        remainder = remaining_latin(translated)
                        if re.search(r"[A-Za-z]{3,}", remainder):
                            key = collapse_ws(translated)
                            if key not in leftover_seen:
                                leftover_seen.add(key)
                                leftovers.append(key)

                    if xml_changed:
                        data = etree.tostring(root, encoding="UTF-8", xml_declaration=True, standalone="yes")

                zout.writestr(item, data)

    if audit_path is not None:
        audit_path.parent.mkdir(parents=True, exist_ok=True)
        audit_path.write_text("\n".join(leftovers), encoding="utf-8")

    return scanned, changed, leftovers


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Translate original_new_part1.docx preserving DOCX layout.")
    parser.add_argument("--input", required=True, help="Input DOCX path")
    parser.add_argument("--output", required=True, help="Output DOCX path")
    parser.add_argument("--glossary", required=True, help="Glossary markdown path")
    parser.add_argument("--audit-file", help="Optional path for remaining English audit output")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    scanned, changed, leftovers = process_docx(
        Path(args.input),
        Path(args.output),
        glossary_path=Path(args.glossary),
        audit_path=Path(args.audit_file) if args.audit_file else None,
    )
    print(f"paragraphs scanned: {scanned}")
    print(f"paragraphs changed: {changed}")
    print(f"remaining english items: {len(leftovers)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

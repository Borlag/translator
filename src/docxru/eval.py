from __future__ import annotations

import json
import re
from collections import Counter
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from time import perf_counter
from typing import Any

from docx import Document

from .config import PipelineConfig
from .docx_reader import collect_segments
from .pipeline import translate_docx


@dataclass(frozen=True)
class EvalResult:
    input_path: str
    output_path: str
    success: bool
    error_message: str | None
    duration_s: float
    segment_count: int
    translated_segments: int
    tm_hits: int
    error_count: int
    warn_count: int
    info_count: int
    glossary_mismatch_segments: int
    glossary_compliance_rate: float | None
    expansion_avg_ratio: float | None
    expansion_max_ratio: float | None
    expansion_warn_count: int
    layout_issue_count: int
    auto_fixes_applied: int
    issue_code_counts: dict[str, int]
    qa_jsonl_path: str
    log_path: str


def _safe_stem(value: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("._")
    return stem or "doc"


def _clone_eval_config(cfg: PipelineConfig, run_dir: Path) -> PipelineConfig:
    tm_cfg = cfg.tm.__class__(**{**cfg.tm.__dict__, "path": str(run_dir / "tm.sqlite")})
    return cfg.__class__(
        **{
            **cfg.__dict__,
            "tm": tm_cfg,
            "qa_report_path": str(run_dir / "qa_report.html"),
            "qa_jsonl_path": str(run_dir / "qa.jsonl"),
            "translation_history_path": str(run_dir / "history.jsonl"),
            "log_path": str(run_dir / "run.log"),
        }
    )


def _read_jsonl(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    rows: list[dict[str, Any]] = []
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
            except json.JSONDecodeError:
                continue
            if isinstance(obj, dict):
                rows.append(obj)
    return rows


def _count_history_records(path: Path) -> int:
    if not path.exists():
        return 0
    count = 0
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            if line.strip():
                count += 1
    return count


def _parse_tm_hits(log_path: Path) -> int:
    if not log_path.exists():
        return 0
    hits = 0
    tm_re = re.compile(r"\bTM hits:\s*(\d+)\b")
    with log_path.open("r", encoding="utf-8", errors="replace") as f:
        for line in f:
            m = tm_re.search(line)
            if m:
                hits = int(m.group(1))
    return hits


def _compute_expansion_stats(
    input_path: Path,
    output_path: Path,
    *,
    include_headers: bool,
    include_footers: bool,
    warn_ratio: float = 1.5,
) -> tuple[float | None, float | None, int]:
    in_doc = Document(str(input_path))
    out_doc = Document(str(output_path))
    src_segments = collect_segments(in_doc, include_headers=include_headers, include_footers=include_footers)
    tgt_segments = collect_segments(out_doc, include_headers=include_headers, include_footers=include_footers)

    ratios: list[float] = []
    warn_count = 0
    pair_count = min(len(src_segments), len(tgt_segments))
    for idx in range(pair_count):
        src = (src_segments[idx].source_plain or "").strip()
        tgt = (tgt_segments[idx].source_plain or "").strip()
        if not src:
            continue
        ratio = len(tgt) / max(1, len(src))
        ratios.append(ratio)
        if ratio > warn_ratio:
            warn_count += 1
    if not ratios:
        return None, None, 0
    avg_ratio = sum(ratios) / len(ratios)
    max_ratio = max(ratios)
    return avg_ratio, max_ratio, warn_count


def evaluate_single(
    input_path: Path,
    output_dir: Path,
    cfg: PipelineConfig,
    *,
    max_segments: int | None = None,
    resume: bool = False,
    output_name: str | None = None,
) -> EvalResult:
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    run_name = output_name or input_path.stem
    safe_name = _safe_stem(run_name)
    output_path = output_dir / f"{safe_name}.docx"
    run_dir = output_dir / "_eval" / safe_name
    run_dir.mkdir(parents=True, exist_ok=True)
    run_cfg = _clone_eval_config(cfg, run_dir)
    qa_jsonl_path = Path(run_cfg.qa_jsonl_path)
    log_path = Path(run_cfg.log_path)
    history_path = Path(run_cfg.translation_history_path or "")

    t0 = perf_counter()
    try:
        translate_docx(
            input_path=input_path,
            output_path=output_path,
            cfg=run_cfg,
            resume=resume,
            max_segments=max_segments,
        )
    except Exception as e:
        dt = perf_counter() - t0
        return EvalResult(
            input_path=str(input_path),
            output_path=str(output_path),
            success=False,
            error_message=str(e),
            duration_s=dt,
            segment_count=0,
            translated_segments=0,
            tm_hits=0,
            error_count=1,
            warn_count=0,
            info_count=0,
            glossary_mismatch_segments=0,
            glossary_compliance_rate=None,
            expansion_avg_ratio=None,
            expansion_max_ratio=None,
            expansion_warn_count=0,
            layout_issue_count=0,
            auto_fixes_applied=0,
            issue_code_counts={},
            qa_jsonl_path=str(qa_jsonl_path),
            log_path=str(log_path),
        )

    dt = perf_counter() - t0
    issue_rows = _read_jsonl(qa_jsonl_path)
    issue_codes = Counter(str(row.get("code", "")) for row in issue_rows if row.get("code"))
    severity_counts = Counter(str(row.get("severity", "")) for row in issue_rows if row.get("severity"))
    segment_count = len(
        collect_segments(
            Document(str(input_path)),
            include_headers=run_cfg.include_headers,
            include_footers=run_cfg.include_footers,
        )
    )

    glossary_mismatch_segments = len(
        {
            str(row.get("segment_id", ""))
            for row in issue_rows
            if str(row.get("code", "")) == "glossary_lemma_mismatch"
        }
    )
    glossary_compliance_rate: float | None = None
    if segment_count > 0 and run_cfg.glossary_lemma_check.strip().lower() != "off":
        glossary_compliance_rate = max(0.0, 1.0 - (glossary_mismatch_segments / segment_count))

    avg_ratio, max_ratio, expansion_warn_count = _compute_expansion_stats(
        input_path,
        output_path,
        include_headers=run_cfg.include_headers,
        include_footers=run_cfg.include_footers,
    )

    layout_issue_count = sum(
        count
        for code, count in issue_codes.items()
        if "layout" in code or "overflow" in code or code == "length_ratio_high"
    )
    auto_fixes_applied = int(issue_codes.get("layout_auto_fix_applied", 0))
    translated_segments = _count_history_records(history_path) if history_path else 0
    tm_hits = _parse_tm_hits(log_path)

    return EvalResult(
        input_path=str(input_path),
        output_path=str(output_path),
        success=True,
        error_message=None,
        duration_s=dt,
        segment_count=segment_count,
        translated_segments=translated_segments,
        tm_hits=tm_hits,
        error_count=int(severity_counts.get("error", 0)),
        warn_count=int(severity_counts.get("warn", 0)),
        info_count=int(severity_counts.get("info", 0)),
        glossary_mismatch_segments=glossary_mismatch_segments,
        glossary_compliance_rate=glossary_compliance_rate,
        expansion_avg_ratio=avg_ratio,
        expansion_max_ratio=max_ratio,
        expansion_warn_count=expansion_warn_count,
        layout_issue_count=layout_issue_count,
        auto_fixes_applied=auto_fixes_applied,
        issue_code_counts=dict(issue_codes),
        qa_jsonl_path=str(qa_jsonl_path),
        log_path=str(log_path),
    )


def evaluate_batch(
    input_dir: Path,
    output_dir: Path,
    cfg: PipelineConfig,
    *,
    max_segments: int | None = None,
    resume: bool = False,
) -> list[EvalResult]:
    input_dir = Path(input_dir)
    files = sorted(path for path in input_dir.glob("*.docx") if path.is_file() and not path.name.startswith("~$"))
    results: list[EvalResult] = []
    for input_path in files:
        results.append(
            evaluate_single(
                input_path=input_path,
                output_dir=output_dir,
                cfg=cfg,
                max_segments=max_segments,
                resume=resume,
                output_name=input_path.stem,
            )
        )
    return results


def summarize_results(results: list[EvalResult]) -> dict[str, Any]:
    success_count = sum(1 for r in results if r.success)
    failed_count = len(results) - success_count
    avg_duration = (sum(r.duration_s for r in results) / len(results)) if results else 0.0
    glossary_rates = [r.glossary_compliance_rate for r in results if r.glossary_compliance_rate is not None]
    avg_glossary_rate = (sum(glossary_rates) / len(glossary_rates)) if glossary_rates else None
    return {
        "files_total": len(results),
        "files_success": success_count,
        "files_failed": failed_count,
        "segment_count_total": sum(r.segment_count for r in results),
        "translated_segments_total": sum(r.translated_segments for r in results),
        "tm_hits_total": sum(r.tm_hits for r in results),
        "errors_total": sum(r.error_count for r in results),
        "warnings_total": sum(r.warn_count for r in results),
        "infos_total": sum(r.info_count for r in results),
        "layout_issues_total": sum(r.layout_issue_count for r in results),
        "auto_fixes_total": sum(r.auto_fixes_applied for r in results),
        "avg_duration_s": avg_duration,
        "avg_glossary_compliance_rate": avg_glossary_rate,
        "expansion_warn_total": sum(r.expansion_warn_count for r in results),
    }


def write_eval_report(results: list[EvalResult], report_path: Path) -> dict[str, Any]:
    summary = summarize_results(results)
    utc = getattr(datetime, "UTC", timezone.utc)  # noqa: UP017
    payload = {
        "generated_at": datetime.now(utc).isoformat(timespec="seconds").replace("+00:00", "Z"),
        "summary": summary,
        "results": [asdict(r) for r in results],
    }
    report_path = Path(report_path)
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return summary


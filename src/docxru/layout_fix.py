from __future__ import annotations

import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import _Cell

from .config import PipelineConfig
from .models import Issue, Segment, Severity

_TEXTBOX_INSET_DEFAULT_CAP_EMU = 25400  # ~2pt (1pt = 12700 EMU)
_DEFAULT_FRAME_HEIGHT_TWIPS = 240
_MIN_TEXTBOX_GROWTH_FACTOR = 1.03
_MIN_FRAME_GROWTH_FACTOR = 1.02
_MAX_TEXTBOX_GROWTH_FACTOR = 2.0
_MAX_FRAME_GROWTH_FACTOR = 2.0
_VML_HEIGHT_RE = re.compile(
    r"((?:^|;)\s*height\s*:\s*)([0-9]+(?:\.[0-9]+)?)(pt|in|cm|mm|px)?",
    re.IGNORECASE,
)


def reduce_font_size(paragraph, reduction_pt: float = 0.5, *, min_font_pt: float = 6.0) -> bool:
    changed = False
    step = max(0.1, float(reduction_pt))
    floor_pt = max(6.0, float(min_font_pt))
    for run in paragraph.runs:
        current = run.font.size.pt if run.font.size is not None else None
        if current is None:
            continue
        if float(current) <= floor_pt + 1e-6:
            continue
        updated = max(floor_pt, float(current) - step)
        if updated < float(current):
            run.font.size = Pt(updated)
            changed = True

    if changed:
        return True

    # If explicit run sizes are missing, apply a small fallback size to text runs.
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        current_size = _resolve_run_size_pt(run, paragraph)
        if current_size is not None:
            if float(current_size) <= floor_pt + 1e-6:
                continue
            fallback = max(floor_pt, float(current_size) - step)
        else:
            # Last-resort fallback if style font cannot be resolved.
            fallback = max(floor_pt, 10.0 - step)
        run.font.size = Pt(float(fallback))
        changed = True
    return changed


def _resolve_run_size_pt(run, paragraph) -> float | None:
    direct = run.font.size
    if direct is not None:
        return float(direct.pt)

    run_style = getattr(run, "style", None)
    run_style_font = getattr(run_style, "font", None)
    run_style_size = getattr(run_style_font, "size", None)
    if run_style_size is not None:
        return float(run_style_size.pt)

    para_style = getattr(paragraph, "style", None)
    para_style_font = getattr(para_style, "font", None)
    para_style_size = getattr(para_style_font, "size", None)
    if para_style_size is not None:
        return float(para_style_size.pt)

    return None


def _is_service_manual_title_paragraph(paragraph) -> bool:
    p_elm = getattr(paragraph, "_p", None)
    if p_elm is None:
        return False
    p_pr = getattr(p_elm, "pPr", None)
    if p_pr is None or p_pr.find(qn("w:framePr")) is None:
        return False
    text = " ".join("".join((run.text or "") for run in paragraph.runs).lower().split())
    if not text:
        return False
    return "руководство по техническому обслуживанию компонентов" in text


def apply_global_font_shrink(segments: list[Segment], cfg: PipelineConfig) -> int:
    """Unconditionally reduce font sizes after write-back for translated segments."""
    body_shrink = max(0.0, float(cfg.font_shrink_body_pt))
    table_shrink = max(0.0, float(cfg.font_shrink_table_pt))
    if body_shrink <= 0.0 and table_shrink <= 0.0:
        return 0

    min_font_pt = max(6.0, float(getattr(cfg, "font_shrink_min_font_pt", 6.0)))
    changed_segments = 0
    for seg in segments:
        paragraph = seg.paragraph_ref
        if paragraph is None or seg.target_tagged is None:
            continue
        if _is_service_manual_title_paragraph(paragraph):
            continue

        in_table_like = bool(seg.context.get("in_table") or seg.context.get("in_textbox"))
        shrink = table_shrink if in_table_like else body_shrink
        if shrink <= 0.0:
            continue

        changed = False
        for run in paragraph.runs:
            if not (run.text or "").strip():
                continue
            current_size = _resolve_run_size_pt(run, paragraph)
            if current_size is None:
                continue
            if current_size <= min_font_pt + 1e-6:
                continue
            new_size = max(min_font_pt, current_size - shrink)
            if new_size + 1e-6 >= current_size:
                continue
            run.font.size = Pt(new_size)
            changed = True

        if changed:
            changed_segments += 1
            seg.issues.append(
                Issue(
                    code="global_font_shrink_applied",
                    severity=Severity.INFO,
                    message="Applied unconditional post-writeback font shrink.",
                    details={
                        "segment_id": seg.segment_id,
                        "location": seg.location,
                        "shrink_pt": shrink,
                        "table_or_textbox": in_table_like,
                    },
                )
            )

    return changed_segments


def reduce_cell_spacing(cell: _Cell, factor: float = 0.8) -> bool:
    changed = False
    ratio = min(1.0, max(0.1, float(factor)))
    for paragraph in cell.paragraphs:
        changed = reduce_paragraph_spacing(paragraph, factor=ratio) or changed
    return changed


def reduce_character_spacing(paragraph, twips: int = -10) -> bool:
    changed = False
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        r_pr = run._r.get_or_add_rPr()
        spacing = r_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            r_pr.append(spacing)
        prev_val = spacing.get(qn("w:val"))
        spacing.set(qn("w:val"), str(int(twips)))
        if prev_val != str(int(twips)):
            changed = True
    return changed


def _local_name(node) -> str:
    tag = str(getattr(node, "tag", "") or "")
    if not tag:
        return ""
    if "}" in tag:
        return tag.split("}", 1)[1]
    if ":" in tag:
        return tag.split(":", 1)[1]
    return tag


def _iter_textbox_body_pr_nodes(paragraph):
    p_elm = getattr(paragraph, "_p", None)
    if p_elm is None:
        return
    # Keep scan scope tight to textbox-related containers only.
    container_hints = {"txbx", "textbox", "shape", "wsp", "drawing", "pict"}
    for ancestor in [p_elm, *list(p_elm.iterancestors())]:
        if _local_name(ancestor) not in container_hints:
            continue
        for node in ancestor.iter():
            if _local_name(node) == "bodyPr":
                yield node


def _iter_vml_textbox_nodes(paragraph):
    p_elm = getattr(paragraph, "_p", None)
    if p_elm is None:
        return
    for ancestor in [p_elm, *list(p_elm.iterancestors())]:
        if _local_name(ancestor) == "textbox":
            yield ancestor


def _set_body_pr_norm_autofit(body_pr) -> bool:
    no_autofit_nodes = []
    has_norm_autofit = False
    for child in list(body_pr):
        name = _local_name(child)
        if name == "noAutofit":
            no_autofit_nodes.append(child)
        elif name == "normAutofit":
            has_norm_autofit = True
    if not no_autofit_nodes:
        return False
    for node in no_autofit_nodes:
        body_pr.remove(node)
    if not has_norm_autofit:
        body_pr.append(OxmlElement("a:normAutofit"))
    return True


def _enable_textbox_norm_autofit(paragraph) -> bool:
    changed = False
    seen: set[int] = set()
    for body_pr in _iter_textbox_body_pr_nodes(paragraph):
        nid = id(body_pr)
        if nid in seen:
            continue
        seen.add(nid)
        changed = _set_body_pr_norm_autofit(body_pr) or changed
    return changed


def _tighten_textbox_insets(paragraph, *, cap_emu: int) -> bool:
    changed = False
    cap = max(0, int(cap_emu))
    seen: set[int] = set()
    for body_pr in _iter_textbox_body_pr_nodes(paragraph):
        nid = id(body_pr)
        if nid in seen:
            continue
        seen.add(nid)
        for attr in ("lIns", "rIns", "tIns", "bIns"):
            raw = body_pr.get(attr)
            if raw is None:
                continue
            try:
                val = int(str(raw))
            except (TypeError, ValueError):
                continue
            if val <= cap:
                continue
            body_pr.set(attr, str(cap))
            changed = True

    # For VML textboxes reclaim default internal paddings as well.
    for tb in _iter_vml_textbox_nodes(paragraph):
        raw_inset = str(tb.get("inset", "") or "").strip()
        if raw_inset == "0,0,0,0":
            continue
        tb.set("inset", "0,0,0,0")
        changed = True
    return changed


def _iter_vml_shape_nodes(paragraph):
    p_elm = getattr(paragraph, "_p", None)
    if p_elm is None:
        return
    for ancestor in [p_elm, *list(p_elm.iterancestors())]:
        if _local_name(ancestor) == "shape":
            yield ancestor


def _frame_pr_node(paragraph):
    p_elm = getattr(paragraph, "_p", None)
    if p_elm is None:
        return None
    p_pr = getattr(p_elm, "pPr", None)
    if p_pr is None:
        return None
    return p_pr.find(qn("w:framePr"))


def _expand_frame_height(paragraph, *, growth_factor: float) -> bool:
    frame_pr = _frame_pr_node(paragraph)
    if frame_pr is None:
        return False
    raw_h = frame_pr.get(qn("w:h"))
    try:
        current_h = int(str(raw_h)) if raw_h is not None else _DEFAULT_FRAME_HEIGHT_TWIPS
    except (TypeError, ValueError):
        current_h = _DEFAULT_FRAME_HEIGHT_TWIPS
    if current_h <= 0:
        current_h = _DEFAULT_FRAME_HEIGHT_TWIPS
    factor = max(_MIN_FRAME_GROWTH_FACTOR, min(_MAX_FRAME_GROWTH_FACTOR, float(growth_factor)))
    new_h = int(round(float(current_h) * factor))
    if new_h <= current_h:
        return False
    frame_pr.set(qn("w:h"), str(int(new_h)))
    frame_pr.set(qn("w:hRule"), "atLeast")
    return True


def _expand_textbox_extents(paragraph, *, growth_factor: float) -> bool:
    p_elm = getattr(paragraph, "_p", None)
    if p_elm is None:
        return False
    factor = max(_MIN_TEXTBOX_GROWTH_FACTOR, min(_MAX_TEXTBOX_GROWTH_FACTOR, float(growth_factor)))
    changed = False
    seen: set[int] = set()
    container_hints = {"drawing", "inline", "anchor", "shape", "wsp", "textbox", "txbx", "pict"}
    for ancestor in [p_elm, *list(p_elm.iterancestors())]:
        if ancestor is not p_elm and _local_name(ancestor) not in container_hints:
            continue
        for node in ancestor.iter():
            if _local_name(node) != "extent":
                continue
            nid = id(node)
            if nid in seen:
                continue
            seen.add(nid)
            raw_cy = node.get("cy")
            try:
                cy = int(str(raw_cy))
            except (TypeError, ValueError):
                continue
            if cy <= 0:
                continue
            new_cy = int(round(float(cy) * factor))
            if new_cy <= cy:
                continue
            node.set("cy", str(int(new_cy)))
            changed = True
    return changed


def _expand_vml_shape_style_height(paragraph, *, growth_factor: float) -> bool:
    factor = max(_MIN_TEXTBOX_GROWTH_FACTOR, min(_MAX_TEXTBOX_GROWTH_FACTOR, float(growth_factor)))
    changed = False
    seen: set[int] = set()
    for shape in _iter_vml_shape_nodes(paragraph):
        nid = id(shape)
        if nid in seen:
            continue
        seen.add(nid)
        style = str(shape.get("style") or "")
        if not style:
            continue
        match = _VML_HEIGHT_RE.search(style)
        if match is None:
            continue
        try:
            height_val = float(match.group(2))
        except (TypeError, ValueError):
            continue
        if height_val <= 0.0:
            continue
        new_val = float(height_val) * factor
        if new_val <= height_val + 1e-6:
            continue
        unit = match.group(3) or ""
        replacement = f"{match.group(1)}{new_val:.2f}{unit}"
        updated_style = style[: match.start()] + replacement + style[match.end() :]
        if updated_style == style:
            continue
        shape.set("style", updated_style)
        changed = True
    return changed


def _growth_factor_from_overflow(
    overflow_ratio: float,
    *,
    max_growth_factor: float,
    min_growth_factor: float,
) -> float:
    if overflow_ratio <= 1.01:
        return 1.0
    # Sublinear growth prevents huge jumps on noisy estimates.
    suggested = float(overflow_ratio) ** 0.68
    return min(float(max_growth_factor), max(float(min_growth_factor), suggested))


def _paragraph_cell(paragraph) -> _Cell | None:
    parent = getattr(paragraph, "_parent", None)
    if isinstance(parent, _Cell):
        return parent
    return None


def _relax_paragraph_frame_height_rule(paragraph) -> bool:
    p_elm = getattr(paragraph, "_p", None)
    if p_elm is None:
        return False
    p_pr = getattr(p_elm, "pPr", None)
    if p_pr is None:
        return False
    frame_pr = p_pr.find(qn("w:framePr"))
    if frame_pr is None:
        return False
    h_rule_attr = qn("w:hRule")
    h_rule = str(frame_pr.get(h_rule_attr, "")).strip().lower()
    if h_rule != "exact":
        return False
    frame_pr.set(h_rule_attr, "atLeast")
    return True


def _remove_paragraph_frame(paragraph) -> bool:
    p_elm = getattr(paragraph, "_p", None)
    if p_elm is None:
        return False
    p_pr = getattr(p_elm, "pPr", None)
    if p_pr is None:
        return False
    frame_pr = p_pr.find(qn("w:framePr"))
    if frame_pr is None:
        return False
    p_pr.remove(frame_pr)
    return True


def _relax_paragraph_exact_line_spacing(paragraph) -> bool:
    p_elm = getattr(paragraph, "_p", None)
    if p_elm is None:
        return False
    p_pr = getattr(p_elm, "pPr", None)
    if p_pr is None:
        return False
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        return False
    line_rule_attr = qn("w:lineRule")
    if str(spacing.get(line_rule_attr, "") or "").strip().lower() != "exact":
        return False
    spacing.set(line_rule_attr, "atLeast")
    return True


def _relax_table_row_exact_height(paragraph) -> bool:
    p_elm = getattr(paragraph, "_p", None)
    if p_elm is None:
        return False
    tr_height_tag = qn("w:trHeight")
    tr_pr_tag = qn("w:trPr")
    h_rule_attr = qn("w:hRule")

    for ancestor in p_elm.iterancestors():
        tag = str(getattr(ancestor, "tag", "")).lower()
        if not tag.endswith("}tr"):
            continue
        tr_pr = ancestor.find(tr_pr_tag)
        if tr_pr is None:
            return False
        changed = False
        for child in list(tr_pr):
            if child.tag != tr_height_tag:
                continue
            if str(child.get(h_rule_attr, "")).strip().lower() != "exact":
                continue
            tr_pr.remove(child)
            changed = True
        return changed
    return False


def reduce_paragraph_spacing(paragraph, factor: float = 0.8) -> bool:
    changed = False
    ratio = min(1.0, max(0.1, float(factor)))
    fmt = paragraph.paragraph_format
    before = fmt.space_before
    after = fmt.space_after
    if before is not None:
        new_before = max(0.0, before.pt * ratio)
        if new_before < before.pt:
            fmt.space_before = Pt(new_before)
            changed = True
    if after is not None:
        new_after = max(0.0, after.pt * ratio)
        if new_after < after.pt:
            fmt.space_after = Pt(new_after)
            changed = True
    if before is None and after is None:
        fmt.space_before = Pt(0.0)
        fmt.space_after = Pt(0.0)
        changed = True
    return changed


def set_single_line_spacing(paragraph) -> bool:
    fmt = paragraph.paragraph_format
    before = fmt.line_spacing
    fmt.line_spacing = 1.0
    return before != fmt.line_spacing


def insert_soft_wraps(paragraph, max_line_chars: int = 34) -> bool:
    if not paragraph.runs:
        return False
    non_empty_runs = [run for run in paragraph.runs if (run.text or "").strip()]
    # Rewriting run text can destroy inline formatting in rich paragraphs,
    # so soft wraps are applied only to simple single-run paragraphs.
    if len(non_empty_runs) != 1:
        return False
    text = "".join(run.text or "" for run in paragraph.runs)
    if len(text.strip()) <= max_line_chars:
        return False
    if "\n" in text:
        return False

    words = text.split()
    if len(words) < 4:
        return False

    chunks: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if len(candidate) <= max_line_chars:
            current = candidate
            continue
        chunks.append(current)
        current = word
    if current:
        chunks.append(current)
    if len(chunks) <= 1:
        return False

    # Minimal-touch rewrite: keep only first run text and clear rest.
    paragraph.runs[0].text = "\n".join(chunks)
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def _estimate_overflow_ratio(issue: Issue | None) -> float:
    if issue is None:
        return 1.0
    details = issue.details or {}
    # Preferred source: explicit capacity-based estimate from layout checks.
    try:
        target_len = float(details.get("target_len", 0) or 0)
        approx_capacity = float(details.get("approx_capacity_chars", 0) or 0)
        if approx_capacity > 0.0 and target_len > 0.0:
            return max(1.0, target_len / approx_capacity)
    except (TypeError, ValueError):
        pass
    # Fallback: ratio from expansion checker.
    try:
        ratio = float(details.get("ratio", 0) or 0)
        if ratio > 0.0:
            return max(1.0, ratio)
    except (TypeError, ValueError):
        pass
    # Last fallback: plain target/source lengths.
    try:
        source_len = float(details.get("source_len", 0) or 0)
        target_len = float(details.get("target_len", 0) or 0)
        if source_len > 0.0 and target_len > 0.0:
            return max(1.0, target_len / source_len)
    except (TypeError, ValueError):
        pass
    return 1.0


def _paragraph_average_font_pt(paragraph) -> float | None:
    sizes: list[float] = []
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        current_size = _resolve_run_size_pt(run, paragraph)
        if current_size is None:
            continue
        sizes.append(float(current_size))
    if not sizes:
        return None
    return sum(sizes) / len(sizes)


def _calculate_adaptive_reduction(
    overflow_ratio: float,
    current_font_pt: float,
    *,
    max_reduction: float = 3.0,
) -> float:
    if overflow_ratio <= 1.0:
        return 0.0
    needed = float(current_font_pt) * (1.0 - 1.0 / float(overflow_ratio))
    return min(float(max_reduction), max(0.2, needed))


def _fix_table_overflow(
    seg: Segment,
    cfg: PipelineConfig,
    *,
    issue: Issue | None = None,
    pass_number: int = 1,
) -> bool:
    if seg.paragraph_ref is None:
        return False

    changed = False
    overflow_ratio = _estimate_overflow_ratio(issue)
    cell = _paragraph_cell(seg.paragraph_ref)
    spacing_factor = float(cfg.layout_spacing_factor)
    if pass_number >= 2:
        spacing_factor *= 0.85
    if cell is not None:
        changed = _relax_table_row_exact_height(seg.paragraph_ref) or changed
        if len(cell.paragraphs) > 1:
            spacing_factor *= 0.85
        changed = reduce_cell_spacing(cell, factor=spacing_factor) or changed
        for paragraph in cell.paragraphs:
            changed = _relax_paragraph_exact_line_spacing(paragraph) or changed
            if pass_number >= 3 and overflow_ratio >= 1.8:
                char_spacing_twips = -12 if pass_number == 3 else -15
                changed = reduce_character_spacing(paragraph, twips=char_spacing_twips) or changed
            if pass_number >= 3 and overflow_ratio >= 1.8:
                changed = set_single_line_spacing(paragraph) or changed
    else:
        changed = _relax_paragraph_exact_line_spacing(seg.paragraph_ref) or changed
        changed = reduce_paragraph_spacing(seg.paragraph_ref, factor=spacing_factor) or changed
        if pass_number >= 3 and overflow_ratio >= 1.8:
            char_spacing_twips = -12 if pass_number == 3 else -15
            changed = reduce_character_spacing(seg.paragraph_ref, twips=char_spacing_twips) or changed
        if pass_number >= 3 and overflow_ratio >= 1.8:
            changed = set_single_line_spacing(seg.paragraph_ref) or changed

    avg_font_pt = _paragraph_average_font_pt(seg.paragraph_ref) or 10.0
    ratio_based_reduction = _calculate_adaptive_reduction(
        overflow_ratio,
        avg_font_pt,
        max_reduction=3.0 if pass_number <= 2 else 3.6,
    )
    table_font_reduction = max(float(cfg.layout_font_reduction_pt), ratio_based_reduction)
    if pass_number >= 3:
        table_font_reduction = max(table_font_reduction, 0.8)
    changed = reduce_font_size(
        seg.paragraph_ref,
        reduction_pt=table_font_reduction,
        min_font_pt=float(getattr(cfg, "layout_min_font_pt", 6.0)),
    ) or changed
    return changed


def _fix_textbox_overflow(
    seg: Segment,
    cfg: PipelineConfig,
    *,
    issue: Issue | None = None,
    allow_container_expand: bool = True,
    pass_number: int = 1,
) -> bool:
    if seg.paragraph_ref is None:
        return False

    changed = False
    overflow_ratio = _estimate_overflow_ratio(issue)
    inset_cap = int(getattr(cfg, "layout_textbox_inset_cap_emu", _TEXTBOX_INSET_DEFAULT_CAP_EMU))
    auto_expand = bool(getattr(cfg, "layout_overflow_box_auto_expand", True))
    max_growth = max(1.0, float(getattr(cfg, "layout_overflow_box_max_height_growth", 1.6)))
    changed = _relax_paragraph_exact_line_spacing(seg.paragraph_ref) or changed
    changed = _enable_textbox_norm_autofit(seg.paragraph_ref) or changed
    changed = _tighten_textbox_insets(seg.paragraph_ref, cap_emu=inset_cap) or changed
    if allow_container_expand and auto_expand and overflow_ratio > 1.05:
        growth_factor = _growth_factor_from_overflow(
            overflow_ratio,
            max_growth_factor=max_growth,
            min_growth_factor=_MIN_TEXTBOX_GROWTH_FACTOR,
        )
        changed = _expand_textbox_extents(seg.paragraph_ref, growth_factor=growth_factor) or changed
        changed = _expand_vml_shape_style_height(seg.paragraph_ref, growth_factor=growth_factor) or changed
    textbox_spacing_factor = min(0.9, max(0.4, float(cfg.layout_spacing_factor)))
    if pass_number >= 2:
        textbox_spacing_factor *= 0.85
    changed = reduce_paragraph_spacing(seg.paragraph_ref, factor=textbox_spacing_factor) or changed
    if pass_number >= 3 and overflow_ratio >= 1.6:
        char_spacing_twips = -10 if pass_number == 3 else -12
        changed = reduce_character_spacing(seg.paragraph_ref, twips=char_spacing_twips) or changed
    if pass_number >= 3:
        changed = set_single_line_spacing(seg.paragraph_ref) or changed
    avg_font_pt = _paragraph_average_font_pt(seg.paragraph_ref) or 10.0
    ratio_based_reduction = _calculate_adaptive_reduction(
        overflow_ratio,
        avg_font_pt,
        max_reduction=3.0 if pass_number <= 2 else 3.6,
    )
    textbox_font_reduction = max(float(cfg.layout_font_reduction_pt), ratio_based_reduction)
    forced_drop_pt = max(0.0, float(getattr(cfg, "layout_overflow_font_drop_pt", 0.0)))
    if pass_number == 1 and overflow_ratio > 1.05 and forced_drop_pt > 0.0:
        textbox_font_reduction = forced_drop_pt
    if pass_number >= 3:
        textbox_font_reduction = max(textbox_font_reduction, 0.7)
    changed = reduce_font_size(
        seg.paragraph_ref,
        reduction_pt=textbox_font_reduction,
        min_font_pt=float(getattr(cfg, "layout_min_font_pt", 6.0)),
    ) or changed
    return changed


def _fix_frame_overflow(
    seg: Segment,
    cfg: PipelineConfig,
    *,
    issue: Issue | None = None,
    pass_number: int = 1,
) -> bool:
    if seg.paragraph_ref is None:
        return False

    changed = False
    overflow_ratio = _estimate_overflow_ratio(issue)
    auto_expand = bool(getattr(cfg, "layout_overflow_box_auto_expand", True))
    max_growth = max(1.0, float(getattr(cfg, "layout_overflow_box_max_height_growth", 1.6)))
    changed = _relax_paragraph_exact_line_spacing(seg.paragraph_ref) or changed
    if auto_expand and overflow_ratio > 1.05:
        growth_factor = _growth_factor_from_overflow(
            overflow_ratio,
            max_growth_factor=max_growth,
            min_growth_factor=_MIN_FRAME_GROWTH_FACTOR,
        )
        changed = _expand_frame_height(seg.paragraph_ref, growth_factor=growth_factor) or changed
    if (overflow_ratio >= 1.35 and bool(seg.context.get("in_table"))) or pass_number >= 3:
        changed = _remove_paragraph_frame(seg.paragraph_ref) or changed
    changed = _relax_paragraph_frame_height_rule(seg.paragraph_ref) or changed
    frame_spacing_factor = min(0.92, max(0.5, float(cfg.layout_spacing_factor)))
    if pass_number >= 2:
        frame_spacing_factor *= 0.85
    changed = reduce_paragraph_spacing(seg.paragraph_ref, factor=frame_spacing_factor) or changed
    if pass_number >= 3 and overflow_ratio >= 1.5:
        char_spacing_twips = -10 if pass_number == 3 else -12
        changed = reduce_character_spacing(seg.paragraph_ref, twips=char_spacing_twips) or changed
    if pass_number >= 3 and overflow_ratio >= 1.5:
        changed = set_single_line_spacing(seg.paragraph_ref) or changed
    avg_font_pt = _paragraph_average_font_pt(seg.paragraph_ref) or 10.0
    ratio_based_reduction = _calculate_adaptive_reduction(
        overflow_ratio,
        avg_font_pt,
        max_reduction=2.6 if pass_number <= 2 else 3.2,
    )
    frame_font_reduction = max(float(cfg.layout_font_reduction_pt), max(0.4, ratio_based_reduction))
    forced_drop_pt = max(0.0, float(getattr(cfg, "layout_overflow_font_drop_pt", 0.0)))
    if pass_number == 1 and overflow_ratio > 1.05 and forced_drop_pt > 0.0:
        frame_font_reduction = forced_drop_pt
    if pass_number >= 3:
        frame_font_reduction = max(frame_font_reduction, 0.6)
    changed = reduce_font_size(
        seg.paragraph_ref,
        reduction_pt=frame_font_reduction,
        min_font_pt=float(getattr(cfg, "layout_min_font_pt", 6.0)),
    ) or changed
    return changed


def _fix_generic_overflow(seg: Segment, cfg: PipelineConfig, *, pass_number: int = 1) -> bool:
    if seg.paragraph_ref is None:
        return False

    changed = False
    changed = _relax_paragraph_exact_line_spacing(seg.paragraph_ref) or changed
    spacing_factor = float(cfg.layout_spacing_factor)
    if pass_number >= 2:
        spacing_factor *= 0.9
    changed = reduce_paragraph_spacing(seg.paragraph_ref, factor=spacing_factor) or changed
    if pass_number >= 3:
        changed = reduce_character_spacing(seg.paragraph_ref, twips=-10) or changed
    base_reduction = max(0.2, float(cfg.layout_font_reduction_pt))
    if pass_number >= 3:
        base_reduction = max(base_reduction, 0.6)
        changed = set_single_line_spacing(seg.paragraph_ref) or changed
    changed = reduce_font_size(
        seg.paragraph_ref,
        reduction_pt=base_reduction,
        min_font_pt=float(getattr(cfg, "layout_min_font_pt", 6.0)),
    ) or changed
    return changed


def fix_expansion_issues(
    segments: list[Segment],
    issues: list[Issue],
    cfg: PipelineConfig,
    *,
    pass_number: int = 1,
) -> int:
    if not segments or not issues:
        return 0

    seg_by_id = {seg.segment_id: seg for seg in segments}
    actionable_codes = {
        "length_ratio_high",
        "layout_table_overflow_risk",
        "layout_textbox_overflow_risk",
        "layout_frame_overflow_risk",
    }
    fixed_segment_ids: set[str] = set()
    expanded_textbox_keys: set[str] = set()

    for issue in issues:
        if issue.code not in actionable_codes:
            continue
        segment_id = str(issue.details.get("segment_id", "")).strip()
        if not segment_id or segment_id in fixed_segment_ids:
            continue
        seg = seg_by_id.get(segment_id)
        if seg is None or seg.paragraph_ref is None:
            continue

        strategy = "generic"
        if issue.code == "layout_table_overflow_risk" or (
            issue.code == "length_ratio_high" and seg.context.get("in_table")
        ):
            strategy = "table"
            changed = _fix_table_overflow(seg, cfg, issue=issue, pass_number=pass_number)
        elif issue.code == "layout_textbox_overflow_risk" or (
            issue.code == "length_ratio_high" and seg.context.get("in_textbox")
        ):
            strategy = "textbox"
            textbox_key = str(seg.context.get("textbox_id") or seg.location).strip()
            allow_expand = textbox_key not in expanded_textbox_keys
            changed = _fix_textbox_overflow(
                seg,
                cfg,
                issue=issue,
                allow_container_expand=allow_expand,
                pass_number=pass_number,
            )
            if allow_expand and textbox_key:
                expanded_textbox_keys.add(textbox_key)
        elif issue.code == "layout_frame_overflow_risk" or (
            issue.code == "length_ratio_high" and seg.context.get("in_frame")
        ):
            strategy = "frame"
            changed = _fix_frame_overflow(seg, cfg, issue=issue, pass_number=pass_number)
        elif issue.code == "length_ratio_high":
            # Do not apply generic destructive fixes to normal body paragraphs.
            continue
        else:
            changed = _fix_generic_overflow(seg, cfg, pass_number=pass_number)
        if not changed:
            continue

        fixed_segment_ids.add(segment_id)
        seg.issues.append(
            Issue(
                code="layout_auto_fix_applied",
                severity=Severity.INFO,
                message="Applied layout auto-fix (spacing/font reduction).",
                details={
                    "segment_id": seg.segment_id,
                    "location": seg.location,
                    "source_issue_code": issue.code,
                    "strategy": strategy,
                    "pass_number": int(pass_number),
                    "font_reduction_pt": float(cfg.layout_font_reduction_pt),
                    "spacing_factor": float(cfg.layout_spacing_factor),
                },
            )
        )

    return len(fixed_segment_ids)

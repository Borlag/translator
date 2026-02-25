from __future__ import annotations

import json
import logging
import re
from collections import Counter, deque
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from tqdm import tqdm

from .checker import (
    CHECKER_SYSTEM_PROMPT,
    attach_checker_edits_to_segments,
    apply_checker_suggestions_to_segments,
    filter_checker_suggestions,
    read_checker_suggestions,
    read_checker_trace_resume_state,
    run_llm_checker,
    write_checker_safe_suggestions,
    write_checker_suggestions,
)
from .config import PipelineConfig
from .consistency import report_consistency
from .dashboard_server import ensure_dashboard_html
from .docx_reader import collect_segments
from .layout_check import validate_layout
from .layout_fix import apply_global_font_shrink, fix_expansion_issues
from .llm import (
    apply_glossary_replacements,
    build_glossary_matchers,
    build_hard_glossary_replacements,
    build_llm_client,
    select_matched_glossary_terms,
    supports_repair,
)
from .logging_utils import setup_logging
from .model_sizing import (
    _median_source_chars,
    recommend_checker_timeout_s,
    recommend_grouped_timeout_s,
    recommend_runtime_model_sizing,
)
from .models import Issue, Segment, Severity
from .oxml_table_fix import normalize_abbyy_oxml
from .pricing import PricingTable, load_pricing_table
from .qa_report import write_qa_jsonl, write_qa_report
from .run_context import resolve_run_paths
from .run_status import RunStatusWriter
from .tagging import is_supported_paragraph, paragraph_to_tagged, tagged_to_runs
from .tm import FuzzyTMHit, TMStore, normalize_text, sha256_hex
from .token_shield import BRACKET_TOKEN_RE, shield, shield_terms, strip_bracket_tokens, unshield
from .usage import UsageTotals
from .validator import (
    is_glossary_lemma_check_available,
    validate_all,
    validate_glossary_lemmas,
    validate_numbers,
    validate_placeholders,
)


def _build_repair_payload(source_shielded: str, bad_output: str) -> str:
    return (
        "TASK: REPAIR_MARKERS\n\n"
        f"SOURCE:\n{source_shielded}\n\n"
        f"OUTPUT:\n{bad_output}"
    )


def _build_glossary_retry_payload(
    source_shielded: str,
    bad_output: str,
    missing_terms: list[dict[str, str]],
) -> str:
    terms = "\n".join(
        f"- {item['source']} -> {item['target']}"
        for item in missing_terms
        if item.get("source") and item.get("target")
    )
    if not terms:
        terms = "- (no terms)"
    return (
        "TASK: REWRITE_FOR_GLOSSARY\n\n"
        "Keep marker tokens exactly as in SOURCE.\n"
        "Do not add comments or explanations.\n\n"
        f"SOURCE:\n{source_shielded}\n\n"
        f"OUTPUT:\n{bad_output}\n\n"
        "REQUIRED_TERMS (EN -> RU):\n"
        f"{terms}\n\n"
        "Return only corrected translated text."
    )


def _read_optional_text(path_value: str | None, logger: logging.Logger, label: str) -> str | None:
    if not path_value:
        return None
    path = Path(path_value)
    try:
        text = path.read_text(encoding="utf-8-sig").strip()
    except OSError as e:
        raise RuntimeError(f"Cannot read {label} from '{path}': {e}") from e
    if not text:
        logger.warning(f"{label} is configured but empty: {path}")
        return None
    logger.info(f"Loaded {label}: {path}")
    return text


_STYLE_START_RE = re.compile(r"⟦S_(\d+)(?:\|[^⟧]*)?⟧")
_PLACEHOLDER_RE = re.compile(r"⟦(?!/?S_)[A-Z][A-Z0-9]*_\d+⟧")
_BRLINE_RE = re.compile(r"⟦BRLINE_\d+⟧")
_LATIN_RE = re.compile(r"[A-Za-z]")
_CYRILLIC_RE = re.compile(r"[А-Яа-яЁё]")
_CYRILLIC_WORD_RE = re.compile(r"[А-Яа-яЁё]{2,}")
_LAYOUT_SPLIT_RE = re.compile(r"((?:\.\s*){3,}|\t+)")
_FINAL_CLEANUP_RULES: tuple[tuple[re.Pattern[str], str], ...] = (
    # Remove zero-width spaces sometimes produced by machine translation (Google).
    (re.compile("\u200b"), ""),
    # Converted manuals often keep standalone English joiners on the title page.
    (re.compile(r"^\s*WITH\s*$", flags=re.IGNORECASE), "С"),
    (re.compile(r"\bNEW/REVISED\b"), "НОВЫЕ/ПЕРЕСМОТРЕННЫЕ"),
    # Legal small-print: tighten wording to avoid cover-page overflow in converted OEM manuals.
    (re.compile(r"\bНастоящий документ\b", flags=re.IGNORECASE), "Документ"),
    (re.compile(r"\bЭтот документ и вся информация\b", flags=re.IGNORECASE), "Документ и информация"),
    (re.compile(r"\bЭтот документ и информация, содержащаяся в нем,\b", flags=re.IGNORECASE), "Документ и его содержание"),
    (re.compile(r"\bявляются\s+исключительной\s+собственностью\b", flags=re.IGNORECASE), "являются собственностью"),
    (re.compile(r"\bсоответствующей\s+дочерней\s+компании\b", flags=re.IGNORECASE), "соответствующей компании"),
    (re.compile(r"^\s*Права на интеллектуальную собственность\b.*", flags=re.IGNORECASE), "Права ИС не предоставляются. Воспроизведение третьим лицам - только с письменного согласия Safran Landing Systems."),
    # Cover title phrasing cleanup.
    (
        re.compile(r"\bС\s+ИЛЛЮСТРИРОВАННЫЙ\s+(?:ПЕРЕЧЕНЬ|СПИСОК)\s+ДЕТАЛЕЙ\b", flags=re.IGNORECASE),
        "С ИЛЛЮСТРИРОВАННЫМ СПИСКОМ ДЕТАЛЕЙ",
    ),
)
_TABLE_REF_WORD_RE = re.compile(
    r"\btable\b(?=\s+(?:\d+[A-Z]?|[IVXLC]+(?:-[IVXLC]+)?|[A-Z]\d+))",
    flags=re.IGNORECASE,
)
_W_T_TAG = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
_W_HYPERLINK_TAG = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
_W_RUN_TAG = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
_TM_RULESET_VERSION = "2026-02-22-consistency-v5-toc-and-morphology"
_BATCH_PROVIDER_ALLOWLIST = {"openai", "ollama"}
_BATCH_MAX_PLACEHOLDER_TOKENS = 12
_JSON_FENCE_RE = re.compile(r"```(?:json)?\s*([\s\S]*?)\s*```", flags=re.IGNORECASE)
_LATIN_WORD_RE = re.compile(r"[A-Za-z]+")


def _style_start_tag(span_id: int, flags: tuple[str, ...]) -> str:
    if flags:
        return f"⟦S_{span_id}|{'|'.join(flags)}⟧"
    return f"⟦S_{span_id}⟧"


def _style_end_tag(span_id: int) -> str:
    return f"⟦/S_{span_id}⟧"


def _extract_style_inner_by_id(tagged_text: str) -> dict[int, str]:
    inner_by_id: dict[int, str] = {}
    pos = 0
    while True:
        m = _STYLE_START_RE.search(tagged_text, pos)
        if not m:
            break
        span_id = int(m.group(1))
        end_pat = re.compile(rf"⟦/S_{span_id}⟧")
        m_end = end_pat.search(tagged_text, m.end())
        if not m_end:
            break
        inner_by_id[span_id] = tagged_text[m.end() : m_end.start()]
        pos = m_end.end()
    return inner_by_id


def _protect_bracket_tokens(text: str) -> tuple[str, dict[str, str]]:
    token_map: dict[str, str] = {}

    def _repl(m: re.Match[str]) -> str:
        ph = f"__DOCXRU_BR_{len(token_map) + 1}__"
        token_map[ph] = m.group(0)
        return ph

    protected = BRACKET_TOKEN_RE.sub(_repl, text)
    return protected, token_map


def _restore_bracket_tokens(text: str, token_map: dict[str, str]) -> str:
    out = text
    for ph, token in token_map.items():
        out = out.replace(ph, token).replace(ph.lower(), token).replace(ph.upper(), token)
    return out


def _translate_shielded_fragment(
    text: str,
    llm_client,
    context: dict[str, Any],
) -> tuple[str, list[Issue]]:
    if not text:
        return text, []
    protected, token_map = _protect_bracket_tokens(text)
    try:
        translated = llm_client.translate(protected, context)
    except Exception as e:
        return text, [
            Issue(
                code="llm_error",
                severity=Severity.WARN,
                message=f"LLM ошибка на фрагменте: {e}",
                details={},
            )
        ]
    restored = _restore_bracket_tokens(translated, token_map)
    placeholder_issues = validate_placeholders(text, restored)
    if any(i.severity == Severity.ERROR for i in placeholder_issues):
        return text, placeholder_issues
    return restored, placeholder_issues


def _translate_plain_chunk(
    chunk: str,
    cfg: PipelineConfig,
    llm_client,
    context: dict[str, Any],
    cache: dict[str, str],
    glossary_terms: tuple[tuple[re.Pattern[str], str], ...] = (),
) -> tuple[str, list[Issue]]:
    if not chunk:
        return chunk, []

    if chunk in cache:
        return cache[chunk], []

    shielded = chunk
    token_map: dict[str, str] = {}
    # Apply hard glossary before generic shielding so phrase-level rules can still match
    # raw text fragments that would otherwise be split into DIM/PN placeholders.
    if glossary_terms and _LATIN_RE.search(shielded):
        shielded, glossary_map = shield_terms(
            shielded,
            glossary_terms,
            token_prefix="GLS",
            bridge_break_tokens=False,
        )
        if glossary_map:
            token_map = {**token_map, **glossary_map}
    shielded, pattern_map = shield(shielded, cfg.pattern_set)
    if pattern_map:
        token_map = {**pattern_map, **token_map}
    translated_shielded, issues = _translate_shielded_fragment(shielded, llm_client, context)
    translated = unshield(translated_shielded, token_map)
    # Run deterministic cleanup rules for Latin-bearing chunks to reduce EN leftovers
    # in TOC/label fragments (including fallback-to-source cases).
    if _LATIN_RE.search(translated):
        cleaned = apply_glossary_replacements(translated, ())
        if cleaned != translated:
            translated = cleaned
            issues = [
                *issues,
                Issue(
                    code="plain_chunk_cleanup_applied",
                    severity=Severity.INFO,
                    message="Applied deterministic cleanup for Latin chunk.",
                    details={},
                ),
            ]

    number_issues = validate_numbers(chunk, translated)
    all_issues = issues + number_issues
    cache[chunk] = translated
    return translated, all_issues


def _translate_toc_like_text(
    text: str,
    cfg: PipelineConfig,
    llm_client,
    context: dict[str, Any],
    cache: dict[str, str],
    glossary_terms: tuple[tuple[re.Pattern[str], str], ...] = (),
) -> tuple[str, list[Issue]]:
    parts = _LAYOUT_SPLIT_RE.split(text)
    out_parts: list[str] = []
    issues: list[Issue] = []

    for part in parts:
        if part == "":
            continue
        if _LAYOUT_SPLIT_RE.fullmatch(part):
            out_parts.append(part)
            continue
        if not _LATIN_RE.search(part):
            out_parts.append(part)
            continue

        m_prefix = re.match(r"^\s*", part)
        m_suffix = re.search(r"\s*$", part)
        prefix = m_prefix.group(0) if m_prefix else ""
        suffix = m_suffix.group(0) if m_suffix else ""
        core_end = len(part) - len(suffix)
        core = part[len(prefix) : core_end]
        if not core:
            out_parts.append(part)
            continue

        translated_core, tr_issues = _translate_plain_chunk(core, cfg, llm_client, context, cache, glossary_terms)
        issues.extend(tr_issues)
        out_parts.append(prefix + translated_core + suffix)

    return "".join(out_parts), issues


def _collect_complex_text_groups(paragraph) -> list[list[Any]]:
    groups: list[list[Any]] = []
    for child in paragraph._p.iterchildren():
        if child.tag == _W_RUN_TAG:
            run_nodes = list(child.iter(_W_T_TAG))
            if run_nodes:
                # Keep run-level boundaries to avoid shifting tabs/page-number alignment.
                groups.append(run_nodes)
            continue

        if child.tag == _W_HYPERLINK_TAG:
            hyperlink_nodes = list(child.iter(_W_T_TAG))
            if hyperlink_nodes:
                groups.append(hyperlink_nodes)
    return groups


def _run_is_safe_for_text_replace(run) -> bool:
    for child in run._r.iterchildren():
        local = child.tag.split("}")[-1]
        if local in {"rPr", "t", "tab", "br", "cr", "noBreakHyphen", "softHyphen", "lastRenderedPageBreak"}:
            continue
        return False
    return True


def _should_translate_segment_text(text: str) -> bool:
    """Heuristic: skip segments that are already in RU or contain no Latin text.

    This avoids:
    - retranslating already-Russian headings (often present in partially translated manuals)
    - rewriting purely numeric / symbol-only paragraphs (page labels, etc.)
    """
    if not text or not text.strip():
        return False
    latin_chars = len(_LATIN_RE.findall(text))
    if latin_chars == 0:
        return False
    cyr_chars = len(_CYRILLIC_RE.findall(text))
    if cyr_chars == 0:
        return True

    latin_words = _LATIN_WORD_RE.findall(text)
    cyr_words = _CYRILLIC_WORD_RE.findall(text)

    # Strong RU dominance with sparse Latin tokens (IDs, acronyms) -> keep as-is.
    if cyr_chars >= latin_chars * 2 and len(latin_words) <= 14:
        return False
    if len(cyr_words) >= max(3, len(latin_words) * 2):
        return False

    # Backward-compatible conservative rule for short mixed labels.
    return not (cyr_chars >= latin_chars * 3 and latin_chars <= 12)


def _compact_context_text(text: str, *, max_chars: int = 220) -> str:
    flat = re.sub(r"\s+", " ", text or "").strip()
    if not flat:
        return ""
    if len(flat) <= max_chars:
        return flat
    return flat[: max_chars - 3].rstrip() + "..."


def _build_matched_glossary_context(
    text: str,
    glossary_matchers,
    *,
    limit: int,
) -> list[dict[str, str]]:
    matched_pairs = select_matched_glossary_terms(text, glossary_matchers, limit=limit)
    if not matched_pairs:
        return []
    return [{"source": source, "target": target} for source, target in matched_pairs]


def _build_tm_references_context(
    hits: list[FuzzyTMHit],
    *,
    max_chars: int,
) -> list[dict[str, Any]]:
    if not hits:
        return []

    refs: list[dict[str, Any]] = []
    consumed = 0
    budget = max(0, int(max_chars))
    for hit in hits:
        source = _compact_context_text(strip_bracket_tokens(hit.source_norm), max_chars=180)
        target = _compact_context_text(strip_bracket_tokens(hit.target_text), max_chars=180)
        if not source or not target:
            continue
        line = f"{source} => {target}"
        line_cost = len(line) + 1
        if budget > 0 and consumed + line_cost > budget:
            break
        refs.append(
            {
                "source": source,
                "target": target,
                "similarity": round(float(hit.similarity), 4),
            }
        )
        consumed += line_cost
    return refs


def _estimate_manual_batch_output_tokens(
    *,
    batch_segments: int,
    batch_max_chars: int,
    source_char_lengths: list[int],
) -> int:
    median_chars = _median_source_chars(source_char_lengths)
    assumed_batch_chars = max(
        max(1, int(batch_max_chars)),
        max(1, int(batch_segments)) * max(1, int(median_chars)),
    )
    estimated_batch_output_chars = int(assumed_batch_chars * 1.35)
    return int(estimated_batch_output_chars / 2.2) + 420


def _effective_manual_max_output_tokens(
    *,
    auto_model_sizing: bool,
    batch_segments: int,
    batch_max_chars: int,
    max_output_tokens: int,
    source_char_lengths: list[int],
) -> int:
    if auto_model_sizing or int(batch_segments) <= 1:
        return int(max_output_tokens)
    estimated_output_tokens = _estimate_manual_batch_output_tokens(
        batch_segments=int(batch_segments),
        batch_max_chars=int(batch_max_chars),
        source_char_lengths=source_char_lengths,
    )
    return max(int(max_output_tokens), int(estimated_output_tokens))


def _recommended_grouped_batch_workers(
    *,
    concurrency: int,
    grouped_jobs_count: int,
    batch_max_chars: int,
) -> int:
    workers = min(max(1, int(concurrency)), max(1, int(grouped_jobs_count)))
    chars = max(0, int(batch_max_chars))
    if chars >= 100_000:
        return min(workers, 2)
    if chars >= 60_000:
        return min(workers, 2)
    if chars >= 36_000:
        return min(workers, 3)
    return workers


def _should_attach_neighbor_context(seg: Segment) -> bool:
    source = (seg.source_plain or "").strip()
    if not source:
        return False
    return bool(_LATIN_RE.search(source))


def _neighbor_context_max_chars(seg: Segment, cfg: PipelineConfig) -> int:
    base = int(cfg.llm.context_window_chars) if cfg.llm.context_window_chars > 0 else 220
    if seg.context.get("in_table"):
        return min(base, 100)
    return base


def _attach_neighbor_snippets(segments: list[Segment], cfg: PipelineConfig) -> None:
    for idx, seg in enumerate(segments):
        if not _should_attach_neighbor_context(seg):
            continue
        max_chars = _neighbor_context_max_chars(seg, cfg)
        prev_text = _compact_context_text(segments[idx - 1].source_plain, max_chars=max_chars) if idx > 0 else ""
        next_text = (
            _compact_context_text(segments[idx + 1].source_plain, max_chars=max_chars)
            if idx + 1 < len(segments)
            else ""
        )
        if prev_text:
            seg.context["prev_text"] = prev_text
        if next_text:
            seg.context["next_text"] = next_text


def _build_document_glossary_context(
    glossary_map: dict[str, str],
    *,
    limit: int,
) -> list[dict[str, str]]:
    if not glossary_map:
        return []
    capped = max(0, int(limit))
    items = list(glossary_map.items())
    if capped > 0:
        items = items[-capped:]
    return [{"source": source, "target": target} for source, target in items]


def _looks_sentence_like_english(text: str) -> bool:
    words = _LATIN_WORD_RE.findall(text or "")
    if len(words) < 8:
        return False
    return bool(re.search(r"[.!?;:]", text or ""))


def _should_apply_hard_glossary_to_segment(seg: Segment) -> bool:
    """Keep hard glossary on compact labels/TOC, relax for long narrative sentences.

    Hard glossary placeholders improve strict term stability, but can lock terms into
    nominative forms inside long RU sentences. We keep locking for TOC/table labels
    and short heading-like text, and relax it for sentence-like body prose.
    """
    if seg.context.get("is_toc_entry"):
        return True
    if seg.context.get("in_table"):
        return True

    source = (seg.source_plain or "").strip()
    if not source:
        return False

    latin_words = _LATIN_WORD_RE.findall(source)
    if len(latin_words) <= 6:
        return True
    if len(latin_words) >= 14:
        return False
    return not _looks_sentence_like_english(source)


def _segment_glossary_terms(
    seg: Segment,
    glossary_terms: tuple[tuple[re.Pattern[str], str], ...],
) -> tuple[tuple[re.Pattern[str], str], ...]:
    if not glossary_terms:
        return ()
    if _should_apply_hard_glossary_to_segment(seg):
        return glossary_terms
    return ()


def _append_recent_translation(
    ring: deque[tuple[str, str]],
    *,
    source_plain: str,
    target_plain: str,
) -> None:
    source = _compact_context_text(strip_bracket_tokens(source_plain), max_chars=220)
    target = _compact_context_text(strip_bracket_tokens(target_plain), max_chars=220)
    if not source or not target:
        return
    ring.append((source, target))


def _normalize_table_reference_labels(text: str) -> str:
    # Translate table references like "Table 1" but keep phrases such as "Table of Contents" untouched.
    def _replace(match: re.Match[str]) -> str:
        token = match.group(0)
        return "Таблица" if token and token[0].isupper() else "таблица"

    return _TABLE_REF_WORD_RE.sub(_replace, text)


def _apply_final_run_level_cleanup(segments: list[Segment]) -> int:
    changed_runs = 0
    seen_paragraphs: set[int] = set()
    for seg in segments:
        if seg.target_shielded_tagged is None or seg.target_tagged is None:
            continue
        para = seg.paragraph_ref
        key = id(para)
        if key in seen_paragraphs:
            continue
        seen_paragraphs.add(key)

        for run in para.runs:
            if not _run_is_safe_for_text_replace(run):
                continue
            original = run.text or ""
            updated = original
            for pattern, replacement in _FINAL_CLEANUP_RULES:
                updated = pattern.sub(replacement, updated)
            updated = _normalize_table_reference_labels(updated)
            if updated != original:
                run.text = updated
                changed_runs += 1
    return changed_runs


def _write_text_nodes(nodes: list[Any], text: str) -> None:
    if not nodes:
        return
    first = nodes[0]
    first.text = text
    xml_space = qn("xml:space")
    if text.startswith(" ") or text.endswith(" "):
        first.set(xml_space, "preserve")
    elif xml_space in first.attrib:
        del first.attrib[xml_space]
    for node in nodes[1:]:
        node.text = ""


def _translate_complex_paragraph_in_place(
    seg: Segment,
    cfg: PipelineConfig,
    llm_client,
    cache: dict[str, str],
    glossary_terms: tuple[tuple[re.Pattern[str], str], ...] = (),
) -> tuple[bool, list[Issue]]:
    groups = _collect_complex_text_groups(seg.paragraph_ref)
    if not groups:
        return False, []

    changed = False
    issues: list[Issue] = []
    for group_i, nodes in enumerate(groups):
        source_text = "".join(node.text or "" for node in nodes)
        if not source_text.strip():
            continue
        if not _LATIN_RE.search(source_text):
            continue
        ctx = dict(seg.context)
        ctx["complex_group"] = group_i
        translated, tr_issues = _translate_toc_like_text(source_text, cfg, llm_client, ctx, cache, glossary_terms)
        issues.extend(tr_issues)
        if translated != source_text:
            _write_text_nodes(nodes, translated)
            changed = True
    return changed, issues


def _fallback_translate_by_spans(
    seg: Segment,
    cfg: PipelineConfig,
    llm_client,
) -> tuple[str | None, list[Issue]]:
    if not seg.spans or not seg.shielded_tagged:
        return None, []

    source_inner = _extract_style_inner_by_id(seg.shielded_tagged)
    if not source_inner:
        return None, []

    rebuilt: list[str] = []
    issues: list[Issue] = []
    span_count = len(seg.spans)
    hard_glossary_terms = tuple(seg.context.get("_hard_glossary_terms") or ())
    for span in seg.spans:
        inner = source_inner.get(span.span_id)
        if inner is None:
            return None, issues

        ctx = dict(seg.context)
        ctx["span_id"] = span.span_id
        ctx["span_total"] = span_count
        shielded_inner = inner
        span_glossary_map: dict[str, str] = {}
        if hard_glossary_terms and _LATIN_RE.search(inner):
            shielded_inner, span_glossary_map = shield_terms(
                inner,
                hard_glossary_terms,
                token_prefix="GLS",
                bridge_break_tokens=False,
            )
        translated_inner, tr_issues = _translate_shielded_fragment(shielded_inner, llm_client, ctx)
        if span_glossary_map:
            translated_inner = unshield(translated_inner, span_glossary_map)
        issues.extend(tr_issues)

        rebuilt.append(_style_start_tag(span.span_id, span.flags))
        rebuilt.append(translated_inner)
        rebuilt.append(_style_end_tag(span.span_id))

    candidate = "".join(rebuilt)
    src_unshielded = unshield(seg.shielded_tagged, seg.token_map or {})
    tgt_unshielded = unshield(candidate, seg.token_map or {})
    src_plain = strip_bracket_tokens(src_unshielded)
    tgt_plain = strip_bracket_tokens(tgt_unshielded)
    final_issues = validate_all(
        source_shielded_tagged=seg.shielded_tagged,
        target_shielded_tagged=candidate,
        source_unshielded_plain=src_plain,
        target_unshielded_plain=tgt_plain,
        short_translation_min_ratio=cfg.short_translation_min_ratio,
        short_translation_min_source_chars=cfg.short_translation_min_source_chars,
        untranslated_latin_warn_ratio=cfg.untranslated_latin_warn_ratio,
        untranslated_latin_min_len=cfg.untranslated_latin_min_len,
        untranslated_latin_allowlist_path=cfg.untranslated_latin_allowlist_path,
        repeated_words_check=cfg.repeated_words_check,
        repeated_phrase_ngram_max=cfg.repeated_phrase_ngram_max,
        context_leakage_check=cfg.context_leakage_check,
        context_leakage_allowlist_path=cfg.context_leakage_allowlist_path,
    )
    issues.extend(final_issues)
    if any(i.severity == Severity.ERROR for i in final_issues):
        return None, issues
    issues.append(
        Issue(
            code="style_fallback_used",
            severity=Severity.INFO,
            message="Применён резервный перевод по span, чтобы сохранить маркеры форматирования.",
            details={},
        )
    )
    return candidate, issues


def _validate_segment_candidate(seg: Segment, cfg: PipelineConfig, out: str) -> list[Issue]:
    src_unshielded = unshield(seg.shielded_tagged or "", seg.token_map or {})
    tgt_unshielded = unshield(out, seg.token_map or {})
    src_plain = strip_bracket_tokens(src_unshielded)
    tgt_plain = strip_bracket_tokens(tgt_unshielded)
    issues = validate_all(
        source_shielded_tagged=seg.shielded_tagged or "",
        target_shielded_tagged=out,
        source_unshielded_plain=src_plain,
        target_unshielded_plain=tgt_plain,
        short_translation_min_ratio=cfg.short_translation_min_ratio,
        short_translation_min_source_chars=cfg.short_translation_min_source_chars,
        untranslated_latin_warn_ratio=cfg.untranslated_latin_warn_ratio,
        untranslated_latin_min_len=cfg.untranslated_latin_min_len,
        untranslated_latin_allowlist_path=cfg.untranslated_latin_allowlist_path,
        repeated_words_check=cfg.repeated_words_check,
        repeated_phrase_ngram_max=cfg.repeated_phrase_ngram_max,
        context_leakage_check=cfg.context_leakage_check,
        context_leakage_allowlist_path=cfg.context_leakage_allowlist_path,
    )
    leak_upper = out.upper()
    if "TASK: REPAIR_MARKERS" in leak_upper or ("SOURCE:" in leak_upper and "OUTPUT:" in leak_upper):
        issues.append(
            Issue(
                code="repair_payload_leak",
                severity=Severity.ERROR,
                message="Repair prompt payload leaked into segment output.",
                details={},
            )
        )
    return issues


def _target_plain_from_candidate(seg: Segment, out: str) -> str:
    target_unshielded = unshield(out, seg.token_map or {})
    return strip_bracket_tokens(target_unshielded)


def _extract_missing_glossary_terms(issues: list[Issue]) -> list[dict[str, str]]:
    missing_terms: list[dict[str, str]] = []
    for issue in issues:
        if issue.code != "glossary_lemma_mismatch":
            continue
        raw_missing = issue.details.get("missing")
        if not isinstance(raw_missing, list):
            continue
        for item in raw_missing:
            if not isinstance(item, dict):
                continue
            source = item.get("source")
            target = item.get("target")
            if isinstance(source, str) and isinstance(target, str):
                missing_terms.append({"source": source, "target": target})
    return missing_terms


def _count_missing_glossary_terms(issues: list[Issue]) -> int:
    return len(_extract_missing_glossary_terms(issues))


def _utc_now_iso() -> str:
    utc = getattr(datetime, "UTC", timezone.utc)  # noqa: UP017
    return datetime.now(utc).isoformat(timespec="seconds").replace("+00:00", "Z")


def _tm_text_fingerprint(text: str | None) -> str:
    if not text:
        return "none"
    return sha256_hex(normalize_text(text))


def _build_tm_profile_key(
    cfg: PipelineConfig,
    *,
    custom_system_prompt: str | None,
    glossary_text: str | None,
) -> str:
    return "|".join(
        (
            f"provider={cfg.llm.provider.strip().lower()}",
            f"model={cfg.llm.model.strip()}",
            f"source_lang={cfg.llm.source_lang.strip().lower()}",
            f"target_lang={cfg.llm.target_lang.strip().lower()}",
            f"hard_glossary={int(bool(cfg.llm.hard_glossary))}",
            f"glossary_in_prompt={int(bool(cfg.llm.glossary_in_prompt))}",
            f"glossary_prompt_mode={cfg.llm.glossary_prompt_mode.strip().lower()}",
            f"glossary_match_limit={int(cfg.llm.glossary_match_limit)}",
            f"structured_output_mode={cfg.llm.structured_output_mode.strip().lower()}",
            f"batch_segments={int(cfg.llm.batch_segments)}",
            f"batch_skip_on_brline={int(bool(cfg.llm.batch_skip_on_brline))}",
            f"batch_max_style_tokens={int(cfg.llm.batch_max_style_tokens)}",
            f"context_window_chars={int(cfg.llm.context_window_chars)}",
            f"fuzzy_enabled={int(bool(cfg.tm.fuzzy_enabled))}",
            f"fuzzy_top_k={int(cfg.tm.fuzzy_top_k)}",
            f"fuzzy_min_similarity={cfg.tm.fuzzy_min_similarity:.4f}",
            f"fuzzy_prompt_max_chars={int(cfg.tm.fuzzy_prompt_max_chars)}",
            f"fuzzy_token_regex={cfg.tm.fuzzy_token_regex}",
            f"fuzzy_rank_mode={cfg.tm.fuzzy_rank_mode.strip().lower()}",
            f"abbyy_profile={cfg.abbyy_profile.strip().lower()}",
            f"glossary_lemma_check={cfg.glossary_lemma_check.strip().lower()}",
            f"layout_check={int(bool(cfg.layout_check))}",
            f"layout_expansion_warn_ratio={float(cfg.layout_expansion_warn_ratio):.3f}",
            f"layout_auto_fix={int(bool(cfg.layout_auto_fix))}",
            f"layout_font_reduction_pt={float(cfg.layout_font_reduction_pt):.3f}",
            f"layout_spacing_factor={float(cfg.layout_spacing_factor):.3f}",
            f"reasoning_effort={(cfg.llm.reasoning_effort or '').strip().lower() or 'default'}",
            f"system_prompt_sha={_tm_text_fingerprint(custom_system_prompt)}",
            f"glossary_sha={_tm_text_fingerprint(glossary_text)}",
        )
    )


def _build_tm_meta(seg: Segment, cfg: PipelineConfig, source_hash: str, *, origin: str) -> dict[str, Any]:
    return {
        "provider": cfg.llm.provider,
        "model": cfg.llm.model,
        "origin": origin,
        "segment_id": seg.segment_id,
        "location": seg.location,
        "part": seg.context.get("part"),
        "section_header": seg.context.get("section_header"),
        "in_table": bool(seg.context.get("in_table")),
        "source_hash": source_hash,
    }


def _build_history_record(
    seg: Segment,
    cfg: PipelineConfig,
    *,
    source_hash: str | None,
    origin: str,
) -> dict[str, Any]:
    source_shielded = seg.shielded_tagged or ""
    target_shielded = seg.target_shielded_tagged or ""
    source_unshielded = unshield(source_shielded, seg.token_map or {})
    target_unshielded = unshield(target_shielded, seg.token_map or {})
    source_plain = strip_bracket_tokens(source_unshielded)
    target_plain = strip_bracket_tokens(target_unshielded)
    return {
        "timestamp_utc": _utc_now_iso(),
        "provider": cfg.llm.provider,
        "model": cfg.llm.model,
        "origin": origin,
        "source_hash": source_hash,
        "segment_id": seg.segment_id,
        "location": seg.location,
        "part": seg.context.get("part"),
        "section_header": seg.context.get("section_header"),
        "in_table": bool(seg.context.get("in_table")),
        "source_text": source_plain,
        "target_text": target_plain,
    }


def _as_warn_issues(issues: list[Issue]) -> list[Issue]:
    normalized: list[Issue] = []
    for issue in issues:
        if issue.severity == Severity.ERROR:
            normalized.append(
                Issue(
                    code=f"{issue.code}_downgraded",
                    severity=Severity.WARN,
                    message=issue.message,
                    details=issue.details,
                )
            )
        else:
            normalized.append(issue)
    return normalized


def _append_history_jsonl(path: Path, records: list[dict[str, Any]]) -> None:
    if not records:
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as f:
        for rec in records:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")


def _finalize_translation_result(
    *,
    seg: Segment,
    source_hash: str,
    source_norm: str,
    out: str,
    issues: list[Issue],
    cfg: PipelineConfig,
    llm_client,
    tm: TMStore,
    complex_chunk_cache: dict[str, str],
    glossary_terms: tuple[tuple[re.Pattern[str], str], ...],
    llm_translated_segments: set[str],
    recent_translations: deque[tuple[str, str]] | None = None,
) -> int:
    hard_errors = any(i.severity == Severity.ERROR for i in issues)
    if hard_errors:
        changed, fallback_issues = _translate_complex_paragraph_in_place(
            seg,
            cfg,
            llm_client,
            complex_chunk_cache,
            glossary_terms,
        )
        if changed:
            seg.target_shielded_tagged = None
            seg.issues.extend(
                [
                    Issue(
                        code="complex_fallback_after_hard_errors",
                        severity=Severity.WARN,
                        message=(
                            "Segment translated with complex in-place fallback after strict marker "
                            "validation failure."
                        ),
                        details={},
                    ),
                    *_as_warn_issues(fallback_issues),
                ]
            )
            return 1

    seg.target_shielded_tagged = out
    seg.issues.extend(issues)

    # Store to TM only if no hard errors (avoid caching broken markers).
    if not hard_errors:
        llm_translated_segments.add(seg.segment_id)
        tm.put_exact(
            source_hash=source_hash,
            source_norm=source_norm,
            target_text=out,
            meta=_build_tm_meta(seg, cfg, source_hash, origin="llm"),
        )
        tm.set_progress(seg.segment_id, "ok", source_hash=source_hash)
        if recent_translations is not None:
            _append_recent_translation(
                recent_translations,
                source_plain=seg.source_plain,
                target_plain=_target_plain_from_candidate(seg, out),
            )
    else:
        tm.set_progress(
            seg.segment_id,
            "error",
            source_hash=source_hash,
            error="; ".join(i.code for i in issues),
        )
    return 0


def _chunk_translation_jobs(
    jobs: list[tuple[Segment, str, str]],
    *,
    max_segments: int,
    max_chars: int,
) -> list[list[tuple[Segment, str, str]]]:
    if not jobs:
        return []

    seg_limit = max(1, int(max_segments))
    char_limit = max(1, int(max_chars))
    groups: list[list[tuple[Segment, str, str]]] = []
    current: list[tuple[Segment, str, str]] = []
    current_chars = 0

    for job in jobs:
        seg = job[0]
        estimated = len(seg.shielded_tagged or "") + 128
        if current and (len(current) >= seg_limit or current_chars + estimated > char_limit):
            groups.append(current)
            current = []
            current_chars = 0
        current.append(job)
        current_chars += estimated

    if current:
        groups.append(current)
    return groups


def _batch_ineligibility_reasons(seg: Segment, cfg: PipelineConfig) -> list[str]:
    text = str(seg.context.get("_batch_eligibility_text") or seg.shielded_tagged or "")
    reasons: list[str] = []
    if seg.context.get("is_toc_entry"):
        reasons.append("toc_entry")
    if cfg.llm.batch_skip_on_brline and _BRLINE_RE.search(text):
        reasons.append("contains_brline")
    if cfg.llm.batch_max_style_tokens >= 0 and len(_STYLE_START_RE.findall(text)) > cfg.llm.batch_max_style_tokens:
        reasons.append("too_many_style_tokens")
    if len(_PLACEHOLDER_RE.findall(text)) > _BATCH_MAX_PLACEHOLDER_TOKENS:
        reasons.append("too_many_placeholders")
    return reasons


def _collect_issue_code_counts(segments: list[Segment]) -> Counter[str]:
    counts: Counter[str] = Counter()
    for seg in segments:
        for issue in seg.issues:
            counts[issue.code] += 1
    return counts


def _attach_issues_to_segments(segments: list[Segment], issues: list[Issue]) -> int:
    if not segments or not issues:
        return 0
    seg_by_id = {seg.segment_id: seg for seg in segments}
    attached = 0
    for issue in issues:
        segment_id = str(issue.details.get("segment_id", "")).strip()
        if segment_id:
            seg = seg_by_id.get(segment_id)
            if seg is None:
                continue
            seg.issues.append(issue)
            attached += 1
            continue
        segments[0].issues.append(issue)
        attached += 1
    return attached


def _coerce_batch_glossary_context(value: Any, *, limit: int = 12) -> list[dict[str, str]]:
    if value is None:
        return []

    items = value if isinstance(value, (list, tuple)) else [value]
    max_items = max(0, int(limit))
    out: list[dict[str, str]] = []
    for item in items:
        if not isinstance(item, dict):
            continue
        source = str(item.get("source") or "").strip()
        target = str(item.get("target") or "").strip()
        if not source or not target:
            continue
        out.append({"source": source, "target": target})
        if max_items and len(out) >= max_items:
            break
    return out


def _coerce_batch_tm_hints(value: Any, *, limit: int = 1) -> list[dict[str, str]]:
    if value is None:
        return []

    items = value if isinstance(value, (list, tuple)) else [value]
    max_items = max(0, int(limit))
    out: list[dict[str, str]] = []
    for item in items:
        if not isinstance(item, dict):
            continue
        source = str(item.get("source") or "").strip()
        target = str(item.get("target") or "").strip()
        if not source or not target:
            continue
        out.append({"source": source, "target": target})
        if max_items and len(out) >= max_items:
            break
    return out


def _coerce_batch_recent_translations(value: Any, *, limit: int = 3) -> list[dict[str, str]]:
    if value is None:
        return []

    items = value if isinstance(value, (list, tuple)) else [value]
    max_items = max(0, int(limit))
    out: list[dict[str, str]] = []
    for item in items:
        if not isinstance(item, dict):
            continue
        source = str(item.get("source") or "").strip()
        target = str(item.get("target") or "").strip()
        if not source or not target:
            continue
        out.append({"source": source, "target": target})
        if max_items and len(out) >= max_items:
            break
    return out


def _build_batch_item_context(seg: Segment) -> str:
    parts: list[str] = []
    section_header = _compact_context_text(seg.context.get("section_header"), max_chars=120)
    if section_header:
        parts.append(f"SECTION={section_header}")

    part = str(seg.context.get("part") or "").strip().lower()
    if part and part != "body":
        parts.append(f"DOC_SECTION={part}")
    if seg.context.get("in_table"):
        parts.append("TABLE_CELL")
    if seg.context.get("in_textbox"):
        parts.append("TEXTBOX")
    if seg.context.get("is_toc_entry"):
        parts.append("TOC_ENTRY")

    prev_text = _compact_context_text(seg.context.get("prev_text"), max_chars=80)
    next_text = _compact_context_text(seg.context.get("next_text"), max_chars=80)
    if prev_text:
        parts.append(f"PREV={prev_text}")
    if next_text:
        parts.append(f"NEXT={next_text}")

    return " | ".join(parts) if parts else "(no context)"


def _build_batch_translation_prompt(items: list[dict[str, Any]]) -> str:
    input_json = json.dumps(items, ensure_ascii=False)
    return (
        "TASK: BATCH_TRANSLATE_SEGMENTS\n"
        "Translate each item.text from English to Russian.\n"
        "Use item.context for disambiguation between homonyms and nearby segment intent.\n"
        "If item.glossary is provided, prefer those EN->RU term mappings.\n"
        "If item.tm_hints/recent_translations are provided, keep terminology consistent with them.\n"
        "Do not translate item.id/item.context/item.glossary/item.tm_hints/item.recent_translations.\n"
        "Keep marker placeholders and style tags unchanged.\n"
        "Return ONLY valid JSON object (no markdown, no prose) where key=id and value=translated text.\n"
        "Preferred shape:\n"
        '{"<id>":"<translated_text>","<id2>":"<translated_text2>"}\n'
        "Also accepted shape:\n"
        '{"translations":[{"id":"<id>","text":"<translated_text>"}]}\n'
        "Rules:\n"
        "- Keep all IDs exactly as provided.\n"
        "- Do not merge or split items.\n"
        "- Include every item exactly once.\n"
        "INPUT_JSON:\n"
        f"{input_json}"
    )


def _iter_balanced_json_chunks(text: str, *, open_char: str, close_char: str) -> list[str]:
    chunks: list[str] = []
    in_string = False
    escaped = False
    depth = 0
    start_idx = -1

    for idx, ch in enumerate(text):
        if in_string:
            if escaped:
                escaped = False
                continue
            if ch == "\\":
                escaped = True
                continue
            if ch == '"':
                in_string = False
            continue

        if ch == '"':
            in_string = True
            continue
        if ch == open_char:
            if depth == 0:
                start_idx = idx
            depth += 1
            continue
        if ch == close_char and depth > 0:
            depth -= 1
            if depth == 0 and start_idx >= 0:
                candidate = text[start_idx : idx + 1].strip()
                if candidate:
                    chunks.append(candidate)
                start_idx = -1
    return chunks


def _extract_json_payload(raw: str) -> Any:
    text = (raw or "").strip()
    if not text:
        raise RuntimeError("Empty batch response")

    candidates: list[str] = []
    seen: set[str] = set()

    def _add_candidate(value: str) -> None:
        candidate = (value or "").strip()
        if not candidate or candidate in seen:
            return
        seen.add(candidate)
        candidates.append(candidate)

    _add_candidate(text)

    for m in _JSON_FENCE_RE.finditer(text):
        fenced = (m.group(1) or "").strip()
        if fenced:
            _add_candidate(fenced)

    obj_start = text.find("{")
    obj_end = text.rfind("}")
    if obj_start >= 0 and obj_end > obj_start:
        _add_candidate(text[obj_start : obj_end + 1])

    arr_start = text.find("[")
    arr_end = text.rfind("]")
    if arr_start >= 0 and arr_end > arr_start:
        _add_candidate(text[arr_start : arr_end + 1])

    for chunk in _iter_balanced_json_chunks(text, open_char="{", close_char="}"):
        _add_candidate(chunk)
    for chunk in _iter_balanced_json_chunks(text, open_char="[", close_char="]"):
        _add_candidate(chunk)

    for candidate in candidates:
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            continue
    raise RuntimeError("Batch response is not valid JSON")


def _coerce_batch_translation_map(payload: Any) -> dict[str, str]:
    entries: list[dict[str, Any]] | None = None
    if isinstance(payload, dict):
        if "translations" in payload and isinstance(payload["translations"], list):
            entries = [item for item in payload["translations"] if isinstance(item, dict)]
        elif all(isinstance(k, str) and isinstance(v, str) for k, v in payload.items()):
            return {str(k): str(v) for k, v in payload.items()}
    elif isinstance(payload, list):
        entries = [item for item in payload if isinstance(item, dict)]

    if entries is None:
        raise RuntimeError("Batch response has unsupported JSON schema")

    out: dict[str, str] = {}
    for item in entries:
        raw_id = item.get("id")
        raw_text = item.get("text")
        if not isinstance(raw_id, str) or not isinstance(raw_text, str):
            continue
        if raw_id in out:
            raise RuntimeError(f"Batch response has duplicate id: {raw_id}")
        out[raw_id] = raw_text
    return out


def _parse_batch_translation_output(raw: str, expected_ids: list[str]) -> dict[str, str]:
    payload = _extract_json_payload(raw)
    out = _coerce_batch_translation_map(payload)
    expected = set(expected_ids)
    got = set(out.keys())
    missing = sorted(expected - got)
    if missing:
        raise RuntimeError(f"Batch response missing ids: {missing[:5]}")
    extra = sorted(got - expected)
    if extra:
        raise RuntimeError(f"Batch response has unexpected ids: {extra[:5]}")
    return {seg_id: out[seg_id] for seg_id in expected_ids}


def _translate_batch_once(
    llm_client,
    jobs: list[tuple[Segment, str, str]],
    cfg: PipelineConfig,
) -> dict[str, str]:
    items: list[dict[str, Any]] = []
    for seg, _, _ in jobs:
        item: dict[str, Any] = {
            "id": seg.segment_id,
            "text": seg.shielded_tagged or "",
            "context": _build_batch_item_context(seg),
        }
        matched_glossary = _coerce_batch_glossary_context(seg.context.get("matched_glossary_terms"), limit=10)
        if matched_glossary:
            item["glossary"] = matched_glossary
        tm_hints = _coerce_batch_tm_hints(
            seg.context.get("tm_references"),
            limit=cfg.llm.batch_tm_hints_per_item,
        )
        if tm_hints:
            item["tm_hints"] = tm_hints
        recent_translations = _coerce_batch_recent_translations(
            seg.context.get("recent_translations"),
            limit=cfg.llm.batch_recent_translations_per_item,
        )
        if recent_translations:
            item["recent_translations"] = recent_translations
        items.append(item)

    prompt = _build_batch_translation_prompt(items)
    first_seg = jobs[0][0]
    context = {
        "task": "batch_translate",
        "part": first_seg.context.get("part"),
        "batch_size": len(jobs),
    }
    raw = llm_client.translate(prompt, context)
    return _parse_batch_translation_output(raw, [item["id"] for item in items])


def _is_batch_json_contract_error(error: Exception) -> bool:
    message = str(error).strip().lower()
    if not message:
        return False
    return any(
        token in message
        for token in (
            "json",
            "schema",
            "missing ids",
            "unexpected ids",
            "duplicate id",
        )
    )


def _is_timeout_error(error: Exception) -> bool:
    message = str(error).strip().lower()
    if not message:
        return False
    return any(
        token in message
        for token in (
            "timed out",
            "timeout",
            "time out",
            "deadline exceeded",
            "read operation timed out",
        )
    )


def _translate_batch_group(
    jobs: list[tuple[Segment, str, str]],
    cfg: PipelineConfig,
    llm_client,
    logger: logging.Logger,
    *,
    _root_batch_size: int | None = None,
    _from_timeout_bisect: bool = False,
) -> list[tuple[Segment, str, str, str, list[Issue]]]:
    fail_fast = bool(cfg.run.fail_fast_on_translate_error)
    root_batch_size = max(1, int(_root_batch_size or len(jobs)))
    if len(jobs) == 1:
        seg, source_hash, source_norm = jobs[0]
        out, issues = _translate_one(seg, cfg, llm_client, source_hash, source_norm, logger)
        if _from_timeout_bisect:
            issues = [
                Issue(
                    code="batch_fallback_single",
                    severity=Severity.WARN,
                    message="Grouped batch timed out and was bisected to single-segment fallback.",
                    details={"batch_size": root_batch_size, "reason": "timeout_bisect"},
                ),
                *issues,
            ]
        return [(seg, source_hash, source_norm, out, issues)]

    batch_map: dict[str, str] | None = None
    batch_error: Exception | None = None
    try:
        batch_map = _translate_batch_once(llm_client, jobs, cfg)
    except Exception as e:
        batch_error = e
        timeout_like = _is_timeout_error(e)
        if timeout_like and bool(cfg.run.batch_timeout_bisect):
            split_at = max(1, len(jobs) // 2)
            left_jobs = jobs[:split_at]
            right_jobs = jobs[split_at:]
            logger.warning(
                "Batch translate timed out (size=%d), bisecting into %d + %d.",
                len(jobs),
                len(left_jobs),
                len(right_jobs),
            )
            left_results = _translate_batch_group(
                left_jobs,
                cfg,
                llm_client,
                logger,
                _root_batch_size=root_batch_size,
                _from_timeout_bisect=True,
            )
            right_results = _translate_batch_group(
                right_jobs,
                cfg,
                llm_client,
                logger,
                _root_batch_size=root_batch_size,
                _from_timeout_bisect=True,
            )
            split_issue = Issue(
                code="batch_timeout_bisect",
                severity=Severity.INFO,
                message="Grouped batch timed out and was bisected.",
                details={
                    "batch_size": len(jobs),
                    "root_batch_size": root_batch_size,
                    "left_size": len(left_jobs),
                    "right_size": len(right_jobs),
                },
            )
            return [
                (seg, source_hash, source_norm, out, [split_issue, *issues])
                for seg, source_hash, source_norm, out, issues in (*left_results, *right_results)
            ]
        if timeout_like:
            logger.warning("Batch translate timed out (size=%d), retrying grouped request once.", len(jobs))
            try:
                batch_map = _translate_batch_once(llm_client, jobs, cfg)
                batch_error = None
            except Exception as retry_err:
                batch_error = retry_err

    if batch_map is None:
        e = batch_error if batch_error is not None else RuntimeError("unknown batch error")
        if fail_fast:
            raise RuntimeError(f"Batch translate failed (size={len(jobs)}): {e}") from e
        logger.warning(f"Batch translate failed (size={len(jobs)}), fallback to single-segment: {e}")
        results: list[tuple[Segment, str, str, str, list[Issue]]] = []
        has_json_contract_error = _is_batch_json_contract_error(e)
        for seg, source_hash, source_norm in jobs:
            out, issues = _translate_one(seg, cfg, llm_client, source_hash, source_norm, logger)
            diagnostics: list[Issue] = []
            if has_json_contract_error:
                diagnostics.append(
                    Issue(
                        code="batch_json_schema_violation",
                        severity=Severity.WARN,
                        message=f"Batch JSON/schema validation failed: {e}",
                        details={"batch_size": len(jobs)},
                    )
                )
            issues = diagnostics + [
                Issue(
                    code="batch_fallback_single",
                    severity=Severity.WARN,
                    message=f"Batch translation failed, fallback to single segment: {e}",
                    details={"batch_size": len(jobs)},
                ),
                *issues,
            ]
            results.append((seg, source_hash, source_norm, out, issues))
        return results

    results: list[tuple[Segment, str, str, str, list[Issue]]] = []
    for seg, source_hash, source_norm in jobs:
        candidate = batch_map.get(seg.segment_id)
        if candidate is None:
            if fail_fast:
                raise RuntimeError(
                    "Batch response missed segment id "
                    f"'{seg.segment_id}' (batch size={len(jobs)})."
                )
            out, fallback_issues = _translate_one(seg, cfg, llm_client, source_hash, source_norm, logger)
            issues = [
                Issue(
                    code="batch_missing_segment",
                    severity=Severity.WARN,
                    message="Batch response missed segment id, fallback to single segment translation.",
                    details={"batch_size": len(jobs)},
                ),
                *fallback_issues,
            ]
            results.append((seg, source_hash, source_norm, out, issues))
            continue

        batch_issues = _validate_segment_candidate(seg, cfg, candidate)
        if any(i.severity == Severity.ERROR for i in batch_issues):
            if fail_fast:
                raise RuntimeError(
                    "Batch output failed validation for segment "
                    f"'{seg.segment_id}' (batch size={len(jobs)})."
                )
            out, fallback_issues = _translate_one(seg, cfg, llm_client, source_hash, source_norm, logger)
            issues = [
                Issue(
                    code="batch_validation_fallback",
                    severity=Severity.WARN,
                    message="Batch output failed validation, fallback to single segment translation.",
                    details={"batch_size": len(jobs)},
                ),
                *batch_issues,
                *fallback_issues,
            ]
            results.append((seg, source_hash, source_norm, out, issues))
            continue

        batch_issues.append(
            Issue(
                code="batch_ok",
                severity=Severity.INFO,
                message="Segment translated in grouped batch mode.",
                details={"batch_size": len(jobs)},
            )
        )
        results.append((seg, source_hash, source_norm, candidate, batch_issues))
    return results


def _translate_one(
    seg: Segment,
    cfg: PipelineConfig,
    llm_client,
    source_hash: str,
    source_norm: str,
    logger: logging.Logger,
) -> tuple[str, list[Issue]]:
    """Translate a single segment with retries and marker validation."""
    last_output: str | None = None
    issues: list[Issue] = []
    can_repair = supports_repair(llm_client)
    max_attempts = cfg.llm.retries if can_repair else 1
    glossary_mode = cfg.glossary_lemma_check.strip().lower()
    if glossary_mode not in {"off", "warn", "retry"}:
        glossary_mode = "off"
    glossary_retry_enabled = glossary_mode == "retry" and can_repair
    glossary_retry_done = False

    for attempt in range(1, max_attempts + 1):
        try:
            if attempt == 1:
                out = llm_client.translate(seg.shielded_tagged or "", seg.context)
            else:
                # Repair attempt: do not retranslate, only fix markers.
                ctx = dict(seg.context)
                ctx["task"] = "repair"
                out = llm_client.translate(
                    _build_repair_payload(seg.shielded_tagged or "", last_output or ""),
                    ctx,
                )
        except Exception as e:
            issues = [
                Issue(
                    code="llm_error",
                    severity=Severity.ERROR,
                    message=f"LLM ошибка: {e}",
                    details={"attempt": attempt},
                )
            ]
            last_output = last_output or ""
            continue

        last_output = out

        # Validate markers and numbers
        issues = _validate_segment_candidate(seg, cfg, out)

        hard_errors = [i for i in issues if i.severity == Severity.ERROR]
        if not hard_errors:
            glossary_issues = validate_glossary_lemmas(
                _target_plain_from_candidate(seg, out),
                seg.context.get("matched_glossary_terms"),
                mode=glossary_mode,
            )
            if not glossary_issues:
                return out, issues

            base_issues = [*issues, *glossary_issues]
            if not glossary_retry_enabled or glossary_retry_done:
                return out, base_issues

            missing_terms = _extract_missing_glossary_terms(glossary_issues)
            if not missing_terms:
                return out, base_issues

            glossary_retry_done = True
            retry_ctx = dict(seg.context)
            retry_ctx["task"] = "repair"
            retry_ctx["glossary_retry"] = True
            try:
                retry_out = llm_client.translate(
                    _build_glossary_retry_payload(seg.shielded_tagged or "", out, missing_terms),
                    retry_ctx,
                )
            except Exception as e:
                return out, [
                    *base_issues,
                    Issue(
                        code="glossary_retry_llm_error",
                        severity=Severity.WARN,
                        message=f"Glossary-focused rewrite failed: {e}",
                        details={"attempt": attempt},
                    ),
                ]

            retry_marker_issues = _validate_segment_candidate(seg, cfg, retry_out)
            retry_hard_errors = [i for i in retry_marker_issues if i.severity == Severity.ERROR]
            if retry_hard_errors:
                return out, [
                    *base_issues,
                    Issue(
                        code="glossary_retry_rejected",
                        severity=Severity.WARN,
                        message="Glossary-focused rewrite failed marker validation; previous output kept.",
                        details={"attempt": attempt},
                    ),
                    *_as_warn_issues(retry_marker_issues),
                ]

            retry_glossary_issues = validate_glossary_lemmas(
                _target_plain_from_candidate(seg, retry_out),
                seg.context.get("matched_glossary_terms"),
                mode=glossary_mode,
            )
            if _count_missing_glossary_terms(retry_glossary_issues) < _count_missing_glossary_terms(glossary_issues):
                return retry_out, [
                    Issue(
                        code="glossary_retry_applied",
                        severity=Severity.INFO,
                        message="Glossary-focused rewrite improved term coverage.",
                        details={"attempt": attempt},
                    ),
                    *retry_marker_issues,
                    *retry_glossary_issues,
                ]

            return out, [
                *base_issues,
                Issue(
                    code="glossary_retry_no_improvement",
                    severity=Severity.INFO,
                    message="Glossary-focused rewrite did not improve term coverage; previous output kept.",
                    details={"attempt": attempt},
                ),
                *retry_marker_issues,
                *_as_warn_issues(retry_glossary_issues),
            ]

        # If we still have marker errors and there is room to retry, go to repair attempt.
        if attempt < max_attempts:
            continue

        # Last-resort recovery: translate span-by-span to salvage marker fidelity.
        fallback_out, fallback_issues = _fallback_translate_by_spans(seg, cfg, llm_client)
        if fallback_out is not None:
            return fallback_out, [
                Issue(
                    code="span_fallback_after_hard_errors",
                    severity=Severity.WARN,
                    message="Segment recovered via span-level fallback after marker validation errors.",
                    details={"attempts": attempt, "supports_repair": can_repair},
                ),
                *fallback_issues,
            ]
        return out, issues

    return last_output or (seg.shielded_tagged or ""), issues


def translate_docx(
    input_path: Path,
    output_path: Path,
    cfg: PipelineConfig,
    *,
    resume: bool = False,
    max_segments: int | None = None,
) -> None:
    logger = setup_logging(Path(cfg.log_path))
    run_paths = resolve_run_paths(cfg, output_path=output_path)
    run_paths.run_dir.mkdir(parents=True, exist_ok=True)
    ensure_dashboard_html(run_paths.dashboard_html_path.parent, filename=run_paths.dashboard_html_path.name)

    logger.info(f"Input: {input_path}")
    logger.info(f"Output: {output_path}")
    logger.info(f"Run dir: {run_paths.run_dir}")
    logger.info(f"Run status: {run_paths.status_path}")
    logger.info(f"Dashboard HTML: {run_paths.dashboard_html_path}")
    logger.info(
        "Mode: "
        f"{cfg.mode}; "
        f"concurrency={cfg.concurrency}; "
        f"headers={cfg.include_headers}; "
        f"footers={cfg.include_footers}; "
        f"structured_output_mode={cfg.llm.structured_output_mode}; "
        f"glossary_in_prompt={cfg.llm.glossary_in_prompt}; "
        f"glossary_prompt_mode={cfg.llm.glossary_prompt_mode}; "
        f"glossary_match_limit={cfg.llm.glossary_match_limit}; "
        f"hard_glossary={cfg.llm.hard_glossary}; "
        f"batch_skip_on_brline={cfg.llm.batch_skip_on_brline}; "
        f"batch_max_style_tokens={cfg.llm.batch_max_style_tokens}; "
        f"context_window_chars={cfg.llm.context_window_chars}; "
        f"reasoning_effort={cfg.llm.reasoning_effort or '(default)'}; "
        f"batch_segments={cfg.llm.batch_segments}; "
        f"batch_max_chars={cfg.llm.batch_max_chars}; "
        f"fuzzy_enabled={cfg.tm.fuzzy_enabled}; "
        f"abbyy_profile={cfg.abbyy_profile}; "
        f"glossary_lemma_check={cfg.glossary_lemma_check}; "
        f"layout_check={cfg.layout_check}; "
        f"layout_expansion_warn_ratio={cfg.layout_expansion_warn_ratio}; "
        f"layout_auto_fix={cfg.layout_auto_fix}; "
        f"layout_font_reduction_pt={cfg.layout_font_reduction_pt}; "
        f"layout_spacing_factor={cfg.layout_spacing_factor}; "
        f"history_jsonl={cfg.translation_history_path or '(off)'}"
    )
    logger.info(f"TM ruleset version: {_TM_RULESET_VERSION}")

    doc = Document(str(input_path))
    segments = collect_segments(doc, include_headers=cfg.include_headers, include_footers=cfg.include_footers)
    if max_segments is not None:
        if max_segments < 0:
            raise ValueError(f"max_segments must be >= 0, got {max_segments}")
        segments = segments[:max_segments]
        logger.info(f"Segment limit enabled: {len(segments)} segments")
    logger.info(f"Segments найдено: {len(segments)}")

    usage_totals = UsageTotals()
    pricing_table = PricingTable.empty(currency=cfg.pricing.currency)
    if cfg.pricing.enabled:
        if cfg.pricing.pricing_path:
            try:
                pricing_table = load_pricing_table(cfg.pricing.pricing_path, currency=cfg.pricing.currency)
                logger.info("Pricing table loaded: %s", cfg.pricing.pricing_path)
            except Exception as exc:
                logger.warning("Failed to load pricing table (%s): %s", cfg.pricing.pricing_path, exc)
        else:
            logger.warning("pricing.enabled=true but pricing.pricing_path is empty")

    status_writer = RunStatusWriter(
        path=run_paths.status_path,
        run_id=run_paths.run_id,
        total_segments=len(segments),
        flush_every_n_updates=cfg.run.status_flush_every_n_segments,
    )

    def _to_dashboard_link(path: Path) -> str:
        try:
            rel = path.resolve().relative_to(run_paths.run_dir.resolve())
            return rel.as_posix()
        except Exception:
            return str(path)

    status_writer.merge_paths(
        {
            "run_dir": str(run_paths.run_dir),
            "output": str(output_path),
            "qa_report": _to_dashboard_link(run_paths.qa_report_path),
            "qa_jsonl": _to_dashboard_link(run_paths.qa_jsonl_path),
            "dashboard_html": _to_dashboard_link(run_paths.dashboard_html_path),
            "checker_suggestions": _to_dashboard_link(run_paths.checker_suggestions_path),
            "checker_suggestions_safe": _to_dashboard_link(run_paths.checker_suggestions_safe_path),
            "checker_trace": _to_dashboard_link(run_paths.checker_trace_path),
        }
    )
    status_writer.set_phase("prepare")
    status_writer.set_done(0)
    status_writer.set_usage(usage_totals.snapshot())
    status_writer.write(force=True)

    # Attach neighbor snippets to improve local consistency.
    _attach_neighbor_snippets(segments, cfg)

    tm = TMStore(
        cfg.tm.path,
        fuzzy_token_regex=cfg.tm.fuzzy_token_regex,
        fuzzy_rank_mode=cfg.tm.fuzzy_rank_mode,
    )
    if cfg.tm.fuzzy_enabled and not tm.fts_enabled:
        logger.info("Fuzzy TM requested, but SQLite FTS5 is unavailable; continuing with exact-only TM behavior.")
    custom_system_prompt = _read_optional_text(cfg.llm.system_prompt_path, logger, "custom system prompt")
    glossary_text = _read_optional_text(cfg.llm.glossary_path, logger, "glossary")
    if cfg.glossary_lemma_check != "off" and not is_glossary_lemma_check_available():
        logger.info(
            "glossary_lemma_check=%s requested, but pymorphy3 is unavailable; fallback exact-term check is used.",
            cfg.glossary_lemma_check,
        )
    if cfg.llm.provider.strip().lower() == "google" and custom_system_prompt:
        logger.info("Provider 'google' does not support system prompts; custom prompt is ignored for this run.")
    glossary_terms: tuple[tuple[re.Pattern[str], str], ...] = ()
    glossary_matchers = build_glossary_matchers(glossary_text) if glossary_text else ()
    provider_norm = cfg.llm.provider.strip().lower()
    glossary_in_prompt = bool(cfg.llm.glossary_in_prompt)
    glossary_prompt_mode = cfg.llm.glossary_prompt_mode.strip().lower()
    hard_glossary = bool(cfg.llm.hard_glossary)
    if glossary_text and hard_glossary:
        glossary_terms = build_hard_glossary_replacements(glossary_text)
        logger.info(
            f"Hard glossary enforcement enabled ({len(glossary_terms)} terms); provider={cfg.llm.provider}; "
            f"glossary_in_prompt={glossary_in_prompt}; hard_glossary={hard_glossary}; "
            "scope=adaptive(toc/table/short-labels)"
        )
    if glossary_prompt_mode == "off" or not glossary_in_prompt:
        effective_glossary_text = None
    elif glossary_prompt_mode == "matched":
        # Segment-level term selection: only matched glossary entries are injected into prompts.
        effective_glossary_text = None
        if glossary_text:
            logger.info("Glossary prompt mode 'matched' enabled: prompts will include only matched terms.")
    else:
        effective_glossary_text = glossary_text
    if glossary_text and effective_glossary_text is None:
        if glossary_prompt_mode == "matched" and glossary_in_prompt:
            logger.info(
                "Global glossary prompt injection is disabled; matched segment-level glossary hints are enabled."
            )
        else:
            logger.info(
                "Glossary prompt injection is disabled; only post-translation glossary replacements are active "
                "(unless hard_glossary=true)."
            )
    if glossary_text and provider_norm == "google" and not hard_glossary:
        logger.info(
            "Provider 'google' is running without hard glossary locking; terminology can be less stable "
            "but wording is usually more natural."
        )
    effective_batch_segments = cfg.llm.batch_segments
    effective_batch_max_chars = cfg.llm.batch_max_chars
    effective_llm_max_output_tokens = cfg.llm.max_output_tokens
    effective_llm_timeout_s = float(cfg.llm.timeout_s)
    effective_checker_cfg = cfg.checker
    effective_checker_timeout_s = float(cfg.checker.timeout_s)
    if cfg.llm.auto_model_sizing:
        checker_provider_for_sizing = (cfg.checker.provider or cfg.llm.provider).strip()
        checker_model_for_sizing = (cfg.checker.model or cfg.llm.model).strip()
        source_lengths = [len(seg.source_plain or "") for seg in segments]
        prompt_chars = len(custom_system_prompt or "") + len(effective_glossary_text or "")
        if glossary_prompt_mode == "matched" and glossary_text:
            # Matched mode sends only term snippets, but reserve budget for those snippets.
            prompt_chars += min(1200, len(glossary_text))
        sizing = recommend_runtime_model_sizing(
            provider=cfg.llm.provider,
            model=cfg.llm.model,
            checker_provider=checker_provider_for_sizing,
            checker_model=checker_model_for_sizing,
            source_char_lengths=source_lengths,
            prompt_chars=prompt_chars,
            batch_segments=cfg.llm.batch_segments,
            batch_max_chars=cfg.llm.batch_max_chars,
            max_output_tokens=cfg.llm.max_output_tokens,
            context_window_chars=cfg.llm.context_window_chars,
            checker_pages_per_chunk=cfg.checker.pages_per_chunk,
            checker_fallback_segments_per_chunk=cfg.checker.fallback_segments_per_chunk,
            checker_max_output_tokens=cfg.checker.max_output_tokens,
        )
        effective_batch_segments = sizing.batch_segments
        effective_batch_max_chars = sizing.batch_max_chars
        effective_llm_max_output_tokens = sizing.max_output_tokens
        effective_checker_cfg = cfg.checker.__class__(
            **{
                **cfg.checker.__dict__,
                "pages_per_chunk": sizing.checker_pages_per_chunk,
                "fallback_segments_per_chunk": sizing.checker_fallback_segments_per_chunk,
                "max_output_tokens": sizing.checker_max_output_tokens,
            }
        )
        for note in sizing.notes:
            logger.info("Model auto-sizing: %s", note)
        logger.info(
            "Model auto-sizing effective values: batch_segments=%d; batch_max_chars=%d; "
            "llm_max_output_tokens=%d; checker_pages_per_chunk=%d; checker_fallback_segments_per_chunk=%d; "
            "checker_max_output_tokens=%d",
            int(effective_batch_segments),
            int(effective_batch_max_chars),
            int(effective_llm_max_output_tokens),
            int(effective_checker_cfg.pages_per_chunk),
            int(effective_checker_cfg.fallback_segments_per_chunk),
            int(effective_checker_cfg.max_output_tokens),
        )
    if not cfg.llm.auto_model_sizing and int(effective_batch_segments) > 1:
        source_lengths = [len(seg.source_plain or "") for seg in segments]
        estimated_output_tokens = _estimate_manual_batch_output_tokens(
            batch_segments=int(effective_batch_segments),
            batch_max_chars=int(effective_batch_max_chars),
            source_char_lengths=source_lengths,
        )
        effective_manual_output_tokens = _effective_manual_max_output_tokens(
            auto_model_sizing=False,
            batch_segments=int(effective_batch_segments),
            batch_max_chars=int(effective_batch_max_chars),
            max_output_tokens=int(effective_llm_max_output_tokens),
            source_char_lengths=source_lengths,
        )
        if int(effective_manual_output_tokens) > int(effective_llm_max_output_tokens):
            logger.info(
                "Auto-raised max_output_tokens to %d (batch content requires more than configured %d); "
                "auto_model_sizing=false, batch=%dx%dchars, estimated_min=%d",
                int(effective_manual_output_tokens),
                int(cfg.llm.max_output_tokens),
                int(effective_batch_segments),
                int(effective_batch_max_chars),
                int(estimated_output_tokens),
            )
            effective_llm_max_output_tokens = int(effective_manual_output_tokens)
    timeout_before = float(effective_llm_timeout_s)
    effective_llm_timeout_s = recommend_grouped_timeout_s(
        timeout_s=timeout_before,
        batch_segments=int(effective_batch_segments),
        batch_max_chars=int(effective_batch_max_chars),
    )
    if float(effective_llm_timeout_s) > timeout_before:
        logger.info(
            "Auto-raised llm timeout_s to %.1f for grouped batches (prev=%.1f; batch=%dx%dchars)",
            float(effective_llm_timeout_s),
            float(timeout_before),
            int(effective_batch_segments),
            int(effective_batch_max_chars),
        )
    checker_timeout_before = float(effective_checker_cfg.timeout_s)
    effective_checker_timeout_s = recommend_checker_timeout_s(
        timeout_s=checker_timeout_before,
        fallback_segments_per_chunk=int(effective_checker_cfg.fallback_segments_per_chunk),
    )
    if float(effective_checker_timeout_s) > checker_timeout_before:
        logger.info(
            "Auto-raised checker timeout_s to %.1f (prev=%.1f; fallback_segments_per_chunk=%d)",
            float(effective_checker_timeout_s),
            float(checker_timeout_before),
            int(effective_checker_cfg.fallback_segments_per_chunk),
        )
    tm_profile_key = _build_tm_profile_key(
        cfg,
        custom_system_prompt=custom_system_prompt,
        glossary_text=glossary_text,
    )
    logger.info(f"TM profile key: {tm_profile_key}")
    llm_client = build_llm_client(
        provider=cfg.llm.provider,
        model=cfg.llm.model,
        temperature=cfg.llm.temperature,
        timeout_s=effective_llm_timeout_s,
        max_output_tokens=effective_llm_max_output_tokens,
        source_lang=cfg.llm.source_lang,
        target_lang=cfg.llm.target_lang,
        base_url=cfg.llm.base_url,
        custom_system_prompt=custom_system_prompt,
        glossary_text=glossary_text,
        glossary_prompt_text=effective_glossary_text,
        prompt_examples_mode=cfg.llm.prompt_examples_mode,
        reasoning_effort=cfg.llm.reasoning_effort,
        prompt_cache_key=cfg.llm.prompt_cache_key,
        prompt_cache_retention=cfg.llm.prompt_cache_retention,
        structured_output_mode=cfg.llm.structured_output_mode,
        on_usage=usage_totals.add,
        estimate_cost=(pricing_table.estimate_cost if cfg.pricing.enabled else None),
        pricing_currency=pricing_table.currency,
    )

    # Stage 1: tagging + shielding + TM lookup
    to_translate: list[tuple[Segment, str, str]] = []  # (seg, source_hash, source_norm)
    tm_hits = 0
    resume_hits = 0
    tagging_errors = 0
    complex_translated = 0
    complex_chunk_cache: dict[str, str] = {}
    segment_source_hash: dict[str, str] = {}
    tm_hit_segments: set[str] = set()
    llm_translated_segments: set[str] = set()
    matched_glossary_segments = 0
    fuzzy_reference_segments = 0
    toc_inplace_translated = 0
    checker_suggestions_count = 0
    checker_safe_suggestions_count = 0
    checker_applied_suggestions = 0
    checker_requests_total = 0
    checker_requests_succeeded = 0
    checker_requests_failed = 0
    checker_chunks_failed = 0
    checker_chunks_total = 0
    checker_translatable_segments = 0
    checker_segments_checked = 0
    checker_split_events = 0
    checker_active_chunk_id = ""
    checker_active_attempt = 0
    checker_last_event = ""
    checker_last_error = ""
    checker_started_at_utc = ""
    checker_updated_at_utc = ""
    processed_segments = 0
    document_glossary: dict[str, str] = {}
    progress_cache: dict[str, dict[str, Any]] = {}
    batch_attempted_segments = 0
    batch_fallback_segments = 0
    batch_json_schema_violations = 0
    fail_fast_on_translate_error = bool(cfg.run.fail_fast_on_translate_error)

    def _flush_status(phase: str, *, force: bool = False) -> None:
        nonlocal processed_segments
        nonlocal checker_suggestions_count
        nonlocal checker_safe_suggestions_count
        nonlocal checker_applied_suggestions
        nonlocal checker_requests_total
        nonlocal checker_requests_succeeded
        nonlocal checker_requests_failed
        nonlocal checker_chunks_failed
        nonlocal checker_chunks_total
        nonlocal checker_translatable_segments
        nonlocal checker_segments_checked
        nonlocal checker_split_events
        nonlocal checker_active_chunk_id
        nonlocal checker_active_attempt
        nonlocal checker_last_event
        nonlocal checker_last_error
        nonlocal checker_started_at_utc
        nonlocal checker_updated_at_utc
        nonlocal batch_attempted_segments
        nonlocal batch_fallback_segments
        nonlocal batch_json_schema_violations
        step = max(1, int(cfg.run.status_flush_every_n_segments))
        if not force and processed_segments % step != 0:
            return
        batch_fallback_ratio = (
            float(batch_fallback_segments / batch_attempted_segments) if batch_attempted_segments > 0 else 0.0
        )
        checker_progress_pct = (
            float(100.0 * checker_segments_checked / checker_translatable_segments)
            if checker_translatable_segments > 0
            else 0.0
        )
        status_writer.set_phase(phase)
        status_writer.set_done(processed_segments)
        status_writer.set_usage(usage_totals.snapshot())
        status_writer.merge_metrics(
            {
                "tm_hits": tm_hits,
                "resume_hits": resume_hits,
                "tagging_errors": tagging_errors,
                "llm_queue": len(to_translate),
                "llm_translated": len(llm_translated_segments),
                "complex_in_place": complex_translated,
                "toc_in_place": toc_inplace_translated,
                "checker_suggestions": checker_suggestions_count,
                "checker_safe_suggestions": checker_safe_suggestions_count,
                "checker_applied_suggestions": checker_applied_suggestions,
                "checker_requests_total": checker_requests_total,
                "checker_requests_succeeded": checker_requests_succeeded,
                "checker_requests_failed": checker_requests_failed,
                "checker_chunks_failed": checker_chunks_failed,
                "checker_chunks_total": checker_chunks_total,
                "checker_translatable_segments": checker_translatable_segments,
                "checker_segments_checked": checker_segments_checked,
                "checker_split_events": checker_split_events,
                "checker_progress_pct": checker_progress_pct,
                "checker_active_chunk": checker_active_chunk_id,
                "checker_active_attempt": checker_active_attempt,
                "checker_last_event": checker_last_event,
                "checker_last_error": checker_last_error,
                "checker_started_at": checker_started_at_utc,
                "checker_updated_at": checker_updated_at_utc,
                "batch_attempted_segments": batch_attempted_segments,
                "batch_fallback_segments": batch_fallback_segments,
                "batch_json_schema_violations": batch_json_schema_violations,
                "batch_fallback_ratio": batch_fallback_ratio,
                "batch_fallback_warn_ratio": float(cfg.run.batch_fallback_warn_ratio),
                "issues_total": sum(len(seg.issues) for seg in segments),
            }
        )
        status_writer.write(force=force)

    def _ingest_batch_issue_metrics(issues: list[Issue]) -> None:
        nonlocal batch_attempted_segments
        nonlocal batch_fallback_segments
        nonlocal batch_json_schema_violations
        if not issues:
            return
        codes = {issue.code for issue in issues}
        attempted_codes = {
            "batch_ok",
            "batch_fallback_single",
            "batch_missing_segment",
            "batch_validation_fallback",
            "batch_worker_crash_fallback",
        }
        fallback_codes = {
            "batch_fallback_single",
            "batch_missing_segment",
            "batch_validation_fallback",
            "batch_worker_crash_fallback",
        }
        if codes.intersection(attempted_codes):
            batch_attempted_segments += 1
        if codes.intersection(fallback_codes):
            batch_fallback_segments += 1
        if "batch_json_schema_violation" in codes:
            batch_json_schema_violations += 1

    def _record_translate_crash(seg: Segment, source_hash: str, err: Exception) -> None:
        nonlocal processed_segments
        seg.issues.append(
            Issue(
                code="translate_crash",
                severity=Severity.ERROR,
                message=f"Translate crash: {err}",
                details={},
            )
        )
        tm.set_progress(seg.segment_id, "error", source_hash=source_hash, error=str(err))
        processed_segments += 1
        _flush_status("translate")
        if fail_fast_on_translate_error:
            raise RuntimeError(
                f"Translation aborted at segment '{seg.segment_id}' due to error: {err}"
            ) from err
    if resume and segments:
        try:
            progress_cache = tm.get_progress_bulk([seg.segment_id for seg in segments])
            logger.info(f"Resume progress cache loaded: {len(progress_cache)} records")
        except Exception as e:
            logger.warning(f"Resume progress bulk load failed; fallback to per-segment lookup: {e}")

    prepare_inplace_jobs: list[tuple[str, Segment, tuple[tuple[re.Pattern[str], str], ...]]] = []
    for seg in tqdm(segments, desc="Prepare", unit="seg"):
        prev_progress = progress_cache.get(seg.segment_id) if resume else None

        # Fast-path: if the segment contains no English (Latin) text, do not touch the paragraph at all.
        # This preserves layout and avoids degrading already-RU content in partially translated manuals.
        if not _should_translate_segment_text(seg.source_plain):
            seg.issues.append(
                Issue(
                    code="skip_no_latin",
                    severity=Severity.INFO,
                    message="Сегмент пропущен: нет английского текста (латиницы) или уже в RU — оставлено как в исходнике",
                    details={},
                )
            )
            processed_segments += 1
            _flush_status("prepare")
            continue

        if glossary_prompt_mode == "matched" and glossary_in_prompt and document_glossary:
            seg.context["document_glossary"] = _build_document_glossary_context(
                document_glossary,
                limit=cfg.llm.glossary_match_limit,
            )

        if glossary_prompt_mode == "matched" and glossary_matchers:
            matched_terms = _build_matched_glossary_context(
                seg.source_plain,
                glossary_matchers,
                limit=cfg.llm.glossary_match_limit,
            )
            if matched_terms:
                seg.context["matched_glossary_terms"] = matched_terms
                matched_glossary_segments += 1
                for pair in matched_terms:
                    source_term = str(pair.get("source") or "").strip()
                    target_term = str(pair.get("target") or "").strip()
                    if source_term and target_term:
                        if source_term in document_glossary:
                            document_glossary.pop(source_term, None)
                        document_glossary[source_term] = target_term
        elif glossary_matchers and not seg.context.get("matched_glossary_terms"):
            # Keep compact term evidence for checker pass even when prompt mode is not "matched".
            checker_terms = _build_matched_glossary_context(
                seg.source_plain,
                glossary_matchers,
                limit=min(12, cfg.llm.glossary_match_limit),
            )
            if checker_terms:
                seg.context["matched_glossary_terms"] = checker_terms

        seg_glossary_terms = _segment_glossary_terms(seg, glossary_terms)
        seg.context["_hard_glossary_terms"] = seg_glossary_terms

        # Dedicated TOC flow: preserve tab/page layout and translate column chunks in-place.
        if seg.context.get("is_toc_entry"):
            prepare_inplace_jobs.append(("toc", seg, seg_glossary_terms))
            continue

        # Safety gate: skip paragraphs that contain complex inline XML (hyperlinks, content controls, etc.)
        # to avoid reordering/corruption when rebuilding runs.
        if not is_supported_paragraph(seg.paragraph_ref):
            prepare_inplace_jobs.append(("complex", seg, seg_glossary_terms))
            continue
        try:
            tagged, spans, inline_map = paragraph_to_tagged(seg.paragraph_ref)
        except Exception as e:
            seg.issues.append(
                Issue(
                    code="tagging_error",
                    severity=Severity.ERROR,
                    message=f"Tagging ошибка: {e}",
                    details={},
                )
            )
            tagging_errors += 1
            processed_segments += 1
            _flush_status("prepare")
            continue

        seg.source_tagged = tagged
        seg.spans = spans
        seg.inline_run_map = inline_map

        # Compute batch-eligibility view before hard glossary shielding so BRLINE decisions
        # are based on original OCR line-break markers.
        batch_eligibility_text, _ = shield(tagged, cfg.pattern_set)
        seg.context["_batch_eligibility_text"] = batch_eligibility_text

        shielded_text = tagged
        token_map: dict[str, str] = {}
        # Apply hard glossary before generic shielding so long phrases with dimensions
        # are not broken by DIM placeholders before glossary matching.
        if seg_glossary_terms and _LATIN_RE.search(shielded_text):
            shielded_text, glossary_map = shield_terms(
                shielded_text,
                seg_glossary_terms,
                token_prefix="GLS",
                bridge_break_tokens=False,
            )
            if glossary_map:
                token_map = {**token_map, **glossary_map}
        shielded_text, pattern_map = shield(shielded_text, cfg.pattern_set)
        if pattern_map:
            token_map = {**pattern_map, **token_map}
        seg.shielded_tagged = shielded_text
        seg.token_map = token_map

        source_norm = normalize_text(shielded_text)
        source_norm_for_hash = f"{_TM_RULESET_VERSION}\n{tm_profile_key}\n{source_norm}"
        source_hash = sha256_hex(source_norm_for_hash)
        segment_source_hash[seg.segment_id] = source_hash

        if (
            resume
            and prev_progress is not None
            and prev_progress.get("status") in {"ok", "tm"}
            and prev_progress.get("source_hash") == source_hash
        ):
            hit = tm.get_exact(source_hash)
            if hit is not None:
                seg.target_shielded_tagged = hit.target_text
                tm_hits += 1
                resume_hits += 1
                tm_hit_segments.add(seg.segment_id)
                processed_segments += 1
                _flush_status("prepare")
                continue

        hit = tm.get_exact(source_hash)
        if hit is not None:
            seg.target_shielded_tagged = hit.target_text
            tm_hits += 1
            tm_hit_segments.add(seg.segment_id)
            tm.set_progress(seg.segment_id, "tm", source_hash=source_hash)
            processed_segments += 1
            _flush_status("prepare")
        else:
            if cfg.tm.fuzzy_enabled:
                fuzzy_hits = tm.get_fuzzy(
                    source_norm,
                    top_k=cfg.tm.fuzzy_top_k,
                    min_similarity=cfg.tm.fuzzy_min_similarity,
                )
                tm_refs = _build_tm_references_context(
                    fuzzy_hits,
                    max_chars=cfg.tm.fuzzy_prompt_max_chars,
                )
                if tm_refs:
                    seg.context["tm_references"] = tm_refs
                    seg.context["tm_references_max_chars"] = int(cfg.tm.fuzzy_prompt_max_chars)
                    fuzzy_reference_segments += 1
            to_translate.append((seg, source_hash, source_norm))

    if prepare_inplace_jobs:
        toc_jobs = sum(1 for kind, _, _ in prepare_inplace_jobs if kind == "toc")
        complex_jobs = len(prepare_inplace_jobs) - toc_jobs
        workers = min(max(1, int(cfg.concurrency)), len(prepare_inplace_jobs))
        logger.info(
            "Prepare queued in-place segments: total=%d; toc=%d; complex=%d; workers=%d",
            len(prepare_inplace_jobs),
            int(toc_jobs),
            int(complex_jobs),
            int(workers),
        )
        with ThreadPoolExecutor(max_workers=workers) as ex:
            futures = {
                ex.submit(
                    _translate_complex_paragraph_in_place,
                    seg,
                    cfg,
                    llm_client,
                    complex_chunk_cache,
                    seg_glossary_terms,
                ): (kind, seg)
                for kind, seg, seg_glossary_terms in prepare_inplace_jobs
            }
            for fut in tqdm(as_completed(futures), total=len(futures), desc="Prepare (in-place)", unit="seg"):
                kind, seg = futures[fut]
                try:
                    changed, inplace_issues = fut.result()
                except Exception as e:
                    seg.issues.append(
                        Issue(
                            code="prepare_in_place_crash",
                            severity=Severity.ERROR,
                            message=f"Prepare in-place crash: {e}",
                            details={"kind": kind},
                        )
                    )
                    processed_segments += 1
                    _flush_status("prepare")
                    continue
                seg.issues.extend(inplace_issues)
                if kind == "toc":
                    if changed:
                        toc_inplace_translated += 1
                else:
                    if changed:
                        complex_translated += 1
                    else:
                        seg.issues.append(
                            Issue(
                                code="skip_complex_paragraph",
                                severity=Severity.INFO,
                                message="Сегмент пропущен: сложная структура абзаца (не только runs) — оставлено как в исходнике",
                                details={},
                            )
                        )
                processed_segments += 1
                _flush_status("prepare")

    logger.info(f"TM hits: {tm_hits}")
    if resume:
        logger.info(f"Resume hits: {resume_hits}")
    logger.info(f"Tagging errors: {tagging_errors}")
    logger.info(f"Complex paragraphs translated in-place: {complex_translated}")
    logger.info(f"TOC segments translated in-place: {toc_inplace_translated}")
    logger.info(f"Segments with matched glossary prompt hints: {matched_glossary_segments}")
    logger.info(f"Segments with fuzzy TM prompt hints: {fuzzy_reference_segments}")
    logger.info(f"LLM segments: {len(to_translate)}")
    _flush_status("prepare", force=True)

    # Stage 2: LLM translate (single-segment or grouped translation requests)
    _flush_status("translate", force=True)
    if to_translate:
        provider_norm = cfg.llm.provider.strip().lower()
        context_window_chars = max(0, int(cfg.llm.context_window_chars))
        context_window_enabled = context_window_chars > 0
        if context_window_enabled:
            logger.info(
                "Context window is enabled (context_window_chars=%d): grouped translation requests disabled; "
                "running sequential translation with recent translations context.",
                context_window_chars,
            )
            recent_translations: deque[tuple[str, str]] = deque(maxlen=6)
            for seg, source_hash, source_norm in tqdm(to_translate, desc="Translate (sequential)", unit="seg"):
                if recent_translations:
                    seg.context["recent_translations"] = [
                        {"source": source, "target": target}
                        for source, target in recent_translations
                    ]
                else:
                    seg.context.pop("recent_translations", None)
                seg.context["recent_translations_max_chars"] = context_window_chars
                try:
                    out, issues = _translate_one(seg, cfg, llm_client, source_hash, source_norm, logger)
                except Exception as e:
                    _record_translate_crash(seg, source_hash, e)
                    continue
                _ingest_batch_issue_metrics(issues)
                complex_translated += _finalize_translation_result(
                    seg=seg,
                    source_hash=source_hash,
                    source_norm=source_norm,
                    out=out,
                    issues=issues,
                    cfg=cfg,
                    llm_client=llm_client,
                    tm=tm,
                    complex_chunk_cache=complex_chunk_cache,
                    glossary_terms=_segment_glossary_terms(seg, glossary_terms),
                    llm_translated_segments=llm_translated_segments,
                    recent_translations=recent_translations,
                )
                processed_segments += 1
                _flush_status("translate")
        else:
            grouped_mode = effective_batch_segments > 1 and provider_norm in _BATCH_PROVIDER_ALLOWLIST
            if effective_batch_segments > 1 and not grouped_mode:
                logger.info(
                    f"Grouped batch mode requested, but provider '{cfg.llm.provider}' is not supported; using per-segment mode."
                )

            if grouped_mode:
                batch_eligible: list[tuple[Segment, str, str]] = []
                single_only: list[tuple[tuple[Segment, str, str], list[str]]] = []
                for job in to_translate:
                    reasons = _batch_ineligibility_reasons(job[0], cfg)
                    if reasons:
                        single_only.append((job, reasons))
                    else:
                        batch_eligible.append(job)

                logger.info(
                    f"Grouped-request eligible segments: {len(batch_eligible)}; "
                    f"forced single by eligibility filter: {len(single_only)}"
                )
                if single_only:
                    reason_counts = Counter(reason for _, reasons in single_only for reason in reasons)
                    logger.info(
                        "Grouped-request ineligibility reasons: "
                        + ", ".join(f"{reason}={count}" for reason, count in reason_counts.most_common())
                    )

                if batch_eligible:
                    grouped_jobs = _chunk_translation_jobs(
                        batch_eligible,
                        max_segments=effective_batch_segments,
                        max_chars=effective_batch_max_chars,
                    )
                    logger.info(f"Grouped batches prepared: {len(grouped_jobs)}")
                else:
                    grouped_jobs = []
                    logger.info("No grouped batches prepared after eligibility filtering.")

                if grouped_jobs:
                    default_batch_workers = min(max(1, cfg.concurrency), len(grouped_jobs))
                    batch_workers = _recommended_grouped_batch_workers(
                        concurrency=cfg.concurrency,
                        grouped_jobs_count=len(grouped_jobs),
                        batch_max_chars=effective_batch_max_chars,
                    )
                    if batch_workers < default_batch_workers:
                        logger.info(
                            "Grouped batch workers auto-capped: %d -> %d (batch_max_chars=%d)",
                            int(default_batch_workers),
                            int(batch_workers),
                            int(effective_batch_max_chars),
                        )
                    else:
                        logger.info(f"Grouped batch workers: {batch_workers}")
                    with ThreadPoolExecutor(max_workers=batch_workers) as ex:
                        futures = {
                            ex.submit(_translate_batch_group, jobs, cfg, llm_client, logger): jobs
                            for jobs in grouped_jobs
                        }
                        for fut in tqdm(as_completed(futures), total=len(futures), desc="Translate (batch)", unit="batch"):
                            jobs = futures[fut]
                            try:
                                batch_results = fut.result()
                            except Exception as e:
                                if fail_fast_on_translate_error:
                                    raise RuntimeError(
                                        f"Batch worker crashed (size={len(jobs)}): {e}"
                                    ) from e
                                logger.warning(
                                    f"Batch worker crashed (size={len(jobs)}), fallback to single-segment: {e}"
                                )
                                for seg, source_hash, source_norm in jobs:
                                    try:
                                        out, issues = _translate_one(seg, cfg, llm_client, source_hash, source_norm, logger)
                                    except Exception as single_err:
                                        _record_translate_crash(seg, source_hash, single_err)
                                        continue
                                    issues = [
                                        Issue(
                                            code="batch_worker_crash_fallback",
                                            severity=Severity.WARN,
                                            message=f"Batch worker crashed, fallback to single segment translation: {e}",
                                            details={"batch_size": len(jobs)},
                                        ),
                                        *issues,
                                    ]
                                    _ingest_batch_issue_metrics(issues)
                                    complex_translated += _finalize_translation_result(
                                        seg=seg,
                                        source_hash=source_hash,
                                        source_norm=source_norm,
                                        out=out,
                                        issues=issues,
                                        cfg=cfg,
                                        llm_client=llm_client,
                                        tm=tm,
                                        complex_chunk_cache=complex_chunk_cache,
                                        glossary_terms=_segment_glossary_terms(seg, glossary_terms),
                                        llm_translated_segments=llm_translated_segments,
                                    )
                                    processed_segments += 1
                                    _flush_status("translate")
                                continue

                            for seg, source_hash, source_norm, out, issues in batch_results:
                                _ingest_batch_issue_metrics(issues)
                                complex_translated += _finalize_translation_result(
                                    seg=seg,
                                    source_hash=source_hash,
                                    source_norm=source_norm,
                                    out=out,
                                    issues=issues,
                                    cfg=cfg,
                                    llm_client=llm_client,
                                    tm=tm,
                                    complex_chunk_cache=complex_chunk_cache,
                                    glossary_terms=_segment_glossary_terms(seg, glossary_terms),
                                    llm_translated_segments=llm_translated_segments,
                                )
                                processed_segments += 1
                                _flush_status("translate")

                if single_only:
                    single_workers = min(max(1, cfg.concurrency), len(single_only))
                    logger.info(f"Single-filter workers: {single_workers}")
                    with ThreadPoolExecutor(max_workers=single_workers) as ex:
                        futures = {
                            ex.submit(_translate_one, seg, cfg, llm_client, source_hash, source_norm, logger): (
                                seg,
                                source_hash,
                                source_norm,
                                reasons,
                            )
                            for (seg, source_hash, source_norm), reasons in single_only
                        }
                        for fut in tqdm(
                            as_completed(futures),
                            total=len(futures),
                            desc="Translate (single-filter)",
                            unit="seg",
                        ):
                            seg, source_hash, source_norm, reasons = futures[fut]
                            try:
                                out, issues = fut.result()
                            except Exception as e:
                                _record_translate_crash(seg, source_hash, e)
                                continue

                            issues = [
                                Issue(
                                    code="batch_ineligible_single",
                                    severity=Severity.INFO,
                                    message=(
                                        "Segment excluded from grouped batch by eligibility filter; translated via single-pass mode."
                                    ),
                                    details={"reasons": reasons},
                                ),
                                *issues,
                            ]
                            _ingest_batch_issue_metrics(issues)
                            complex_translated += _finalize_translation_result(
                                seg=seg,
                                source_hash=source_hash,
                                source_norm=source_norm,
                                out=out,
                                issues=issues,
                                cfg=cfg,
                                llm_client=llm_client,
                                tm=tm,
                                complex_chunk_cache=complex_chunk_cache,
                                glossary_terms=_segment_glossary_terms(seg, glossary_terms),
                                llm_translated_segments=llm_translated_segments,
                            )
                            processed_segments += 1
                            _flush_status("translate")
            else:
                with ThreadPoolExecutor(max_workers=max(1, cfg.concurrency)) as ex:
                    futures = {
                        ex.submit(_translate_one, seg, cfg, llm_client, sh, sn, logger): (seg, sh, sn)
                        for (seg, sh, sn) in to_translate
                    }
                    for fut in tqdm(as_completed(futures), total=len(futures), desc="Translate", unit="seg"):
                        seg, source_hash, source_norm = futures[fut]
                        try:
                            out, issues = fut.result()
                        except Exception as e:
                            _record_translate_crash(seg, source_hash, e)
                            continue

                        _ingest_batch_issue_metrics(issues)
                        complex_translated += _finalize_translation_result(
                            seg=seg,
                            source_hash=source_hash,
                            source_norm=source_norm,
                            out=out,
                            issues=issues,
                            cfg=cfg,
                            llm_client=llm_client,
                            tm=tm,
                            complex_chunk_cache=complex_chunk_cache,
                            glossary_terms=_segment_glossary_terms(seg, glossary_terms),
                            llm_translated_segments=llm_translated_segments,
                        )
                        processed_segments += 1
                        _flush_status("translate")

    _flush_status("translate", force=True)
    if batch_attempted_segments > 0:
        batch_fallback_ratio = float(batch_fallback_segments / batch_attempted_segments)
        logger.info(
            "Batch fallback monitoring: attempted=%d; fallback=%d; ratio=%.2f%%; warn_threshold=%.2f%%",
            int(batch_attempted_segments),
            int(batch_fallback_segments),
            float(batch_fallback_ratio * 100.0),
            float(cfg.run.batch_fallback_warn_ratio * 100.0),
        )
        if batch_attempted_segments >= 20 and batch_fallback_ratio > float(cfg.run.batch_fallback_warn_ratio):
            logger.warning(
                "Batch fallback ratio exceeded threshold (%.2f%% > %.2f%%). "
                "Consider switching to grouped_fast/sequential_context or lowering batch density.",
                float(batch_fallback_ratio * 100.0),
                float(cfg.run.batch_fallback_warn_ratio * 100.0),
            )

    # Stage 3: Unshield + write back to DOCX
    status_writer.set_phase("writeback")
    status_writer.write(force=True)
    written = 0
    for seg in tqdm(segments, desc="Write", unit="seg"):
        if not seg.spans or seg.shielded_tagged is None:
            continue

        if seg.target_shielded_tagged is None:
            # No translation (tagging error etc) — keep source
            continue

        if any(i.severity == Severity.ERROR for i in seg.issues):
            seg.issues.append(
                Issue(
                    code="writeback_skipped_due_to_errors",
                    severity=Severity.INFO,
                    message="Write-back skipped due to hard validation/translation errors; source text preserved.",
                    details={},
                )
            )
            continue

        # Unshield placeholders back to original protected strings
        target_tagged_unshielded = unshield(seg.target_shielded_tagged, seg.token_map or {})
        seg.target_tagged = target_tagged_unshielded

        try:
            tagged_to_runs(
                seg.paragraph_ref,
                target_tagged_unshielded,
                seg.spans,
                inline_run_map=seg.inline_run_map,
            )
            written += 1
        except Exception as e:
            seg.issues.append(
                Issue(
                    code="write_error",
                    severity=Severity.ERROR,
                    message=f"Write-back ошибка: {e}",
                    details={},
                )
            )

    logger.info(f"Written segments: {written}/{len(segments)}; complex in-place: {complex_translated}")

    if cfg.font_shrink_body_pt > 0 or cfg.font_shrink_table_pt > 0:
        shrunk = apply_global_font_shrink(segments, cfg)
        logger.info(
            "Global font shrink applied: %d segments (body=%.2fpt, table/textbox=%.2fpt)",
            shrunk,
            float(cfg.font_shrink_body_pt),
            float(cfg.font_shrink_table_pt),
        )

    if glossary_matchers:
        consistency_issues = report_consistency(segments, glossary_matchers)
        attached = _attach_issues_to_segments(segments, consistency_issues)
        if consistency_issues:
            logger.info(f"Consistency check issues: {len(consistency_issues)} (attached={attached})")

    if cfg.layout_check or cfg.layout_auto_fix:
        layout_issues = validate_layout(doc, segments, cfg)
        attached = _attach_issues_to_segments(segments, layout_issues)
        if layout_issues:
            logger.info(f"Layout validation issues: {len(layout_issues)} (attached={attached})")
        if cfg.layout_auto_fix and layout_issues:
            applied_fixes = fix_expansion_issues(segments, layout_issues, cfg)
            logger.info(f"Layout auto-fixes applied: {applied_fixes}")

    cleaned_runs = _apply_final_run_level_cleanup(segments)
    logger.info(f"Final run-level cleanup changes: {cleaned_runs}")
    issue_counts = _collect_issue_code_counts(segments)
    for code in (
        "placeholders_mismatch",
        "style_tags_mismatch",
        "batch_fallback_single",
        "batch_validation_fallback",
    ):
        logger.info(f"Metric {code}: {issue_counts.get(code, 0)}")
    if issue_counts:
        logger.info(
            "Issue codes summary (top 12): "
            + ", ".join(f"{code}={count}" for code, count in issue_counts.most_common(12))
        )

    if cfg.abbyy_profile != "off":
        try:
            oxml_stats = normalize_abbyy_oxml(doc, profile=cfg.abbyy_profile)
            logger.info(
                "ABBYY OXML normalization (%s): trHeight_exact_removed=%d; framePr_removed=%d; "
                "lineSpacing_exact_relaxed=%d",
                cfg.abbyy_profile,
                int(oxml_stats.get("tr_height_exact_removed", 0)),
                int(oxml_stats.get("frame_pr_removed", 0)),
                int(oxml_stats.get("line_spacing_exact_relaxed", 0)),
            )
        except Exception as e:
            logger.warning(f"ABBYY OXML normalization failed: {e}")

    # Save outputs
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("DOCX сохранён")

    if cfg.mode.lower() == "com":
        try:
            from .com_word import update_fields_via_com

            com_stats = update_fields_via_com(output_path)
            logger.info(
                "Word COM post-process: fields_updated=%d; tocs_updated=%d; "
                "textboxes_seen=%d; textboxes_autofit=%d; textboxes_shrunk=%d",
                int(com_stats.get("fields_updated", 0)),
                int(com_stats.get("tocs_updated", 0)),
                int(com_stats.get("textboxes_seen", 0)),
                int(com_stats.get("textboxes_autofit", 0)),
                int(com_stats.get("textboxes_shrunk", 0)),
            )
        except Exception as e:
            logger.warning(f"Word COM post-process failed: {e}")

    checker_edits: list[dict[str, Any]] = []
    if effective_checker_cfg.enabled:
        status_writer.set_phase("checker")
        checker_started_at_utc = datetime.now(timezone.utc).isoformat()
        checker_updated_at_utc = checker_started_at_utc
        checker_last_event = "start"
        checker_last_error = ""
        checker_active_chunk_id = ""
        checker_active_attempt = 0
        checker_resume_skip_chunks: set[str] = set()
        checker_resume_base_requests_total = 0
        checker_resume_base_requests_succeeded = 0
        checker_resume_base_requests_failed = 0
        checker_resume_base_chunks_failed = 0
        checker_resume_base_split_events = 0
        checker_resume_base_suggestions = 0
        checker_resume_base_chunks_total = 0
        checker_resume_base_segments_checked = 0
        checker_resume_state: dict[str, Any] = {}
        existing_checker_edits: list[dict[str, Any]] = []

        if resume:
            if run_paths.checker_suggestions_path.exists():
                try:
                    existing_checker_edits = read_checker_suggestions(run_paths.checker_suggestions_path)
                except Exception as exc:
                    logger.warning("Checker resume: failed to read existing suggestions: %s", exc)
                    existing_checker_edits = []
            if run_paths.checker_trace_path.exists():
                try:
                    checker_resume_state = read_checker_trace_resume_state(run_paths.checker_trace_path)
                except Exception as exc:
                    logger.warning("Checker resume: failed to read checker trace state: %s", exc)
                    checker_resume_state = {}
                completed_chunks_raw = checker_resume_state.get("completed_chunk_ids") or []
                if existing_checker_edits and isinstance(completed_chunks_raw, list):
                    checker_resume_skip_chunks = {
                        str(item).strip()
                        for item in completed_chunks_raw
                        if str(item).strip()
                    }
                    if checker_resume_skip_chunks:
                        checker_resume_base_chunks_total = len(checker_resume_skip_chunks)
                        logger.info(
                            "Checker resume: will skip %d completed chunks from previous attempt.",
                            len(checker_resume_skip_chunks),
                        )
                elif completed_chunks_raw:
                    logger.warning(
                        "Checker resume: trace has completed chunks, but suggestions checkpoint is missing; "
                        "rerunning checker from the beginning to avoid losing prior edits."
                    )

                try:
                    checker_translatable_segments = max(
                        checker_translatable_segments,
                        int(checker_resume_state.get("translatable_segments") or 0),
                    )
                    checker_resume_base_requests_total = max(0, int(checker_resume_state.get("requests_total") or 0))
                    checker_resume_base_requests_succeeded = max(
                        0,
                        int(checker_resume_state.get("requests_succeeded") or 0),
                    )
                    checker_resume_base_requests_failed = max(
                        0,
                        int(checker_resume_state.get("requests_failed") or 0),
                    )
                    checker_resume_base_chunks_failed = max(0, int(checker_resume_state.get("chunks_failed") or 0))
                    checker_resume_base_split_events = max(0, int(checker_resume_state.get("split_events") or 0))
                    if existing_checker_edits:
                        checker_resume_base_suggestions = max(
                            checker_resume_base_suggestions,
                            int(checker_resume_state.get("suggestions_total") or 0),
                        )
                except Exception:
                    logger.debug("Checker resume: failed to normalize base trace counters", exc_info=True)
            if existing_checker_edits and checker_resume_skip_chunks:
                checker_edits = attach_checker_edits_to_segments(
                    segments=segments,
                    edits=existing_checker_edits,
                    logger=logger,
                )
                checker_suggestions_count = len(checker_edits)
                checker_resume_base_suggestions = checker_suggestions_count
                logger.info(
                    "Checker resume: loaded %d suggestions checkpoint entries.",
                    checker_suggestions_count,
                )
            elif existing_checker_edits:
                logger.info(
                    "Checker resume: suggestions checkpoint exists, but no completed chunk map; "
                    "checker will rerun from the beginning."
                )

        _flush_status("checker", force=True)

        def _on_checker_progress(event: dict[str, Any]) -> None:
            nonlocal checker_suggestions_count
            nonlocal checker_requests_total
            nonlocal checker_requests_succeeded
            nonlocal checker_requests_failed
            nonlocal checker_chunks_failed
            nonlocal checker_chunks_total
            nonlocal checker_translatable_segments
            nonlocal checker_segments_checked
            nonlocal checker_split_events
            nonlocal checker_active_chunk_id
            nonlocal checker_active_attempt
            nonlocal checker_last_event
            nonlocal checker_last_error
            nonlocal checker_started_at_utc
            nonlocal checker_updated_at_utc

            def _event_int(key: str) -> int | None:
                value = event.get(key)
                if value is None:
                    return None
                try:
                    return max(0, int(value))
                except Exception:
                    return None

            event_name = str(event.get("event") or "").strip()
            now_iso = datetime.now(timezone.utc).isoformat()
            if not checker_started_at_utc:
                checker_started_at_utc = now_iso
            checker_updated_at_utc = now_iso
            if event_name:
                checker_last_event = event_name

            chunk_id = str(event.get("chunk_id") or "").strip()
            if chunk_id:
                checker_active_chunk_id = chunk_id
            attempt = _event_int("attempt")
            if attempt is not None:
                checker_active_attempt = attempt
            if event_name in {"response", "failed_chunk", "summary"}:
                checker_active_attempt = 0
            if event_name == "summary":
                checker_active_chunk_id = ""

            err_text = str(event.get("error") or "").strip()
            if err_text:
                checker_last_error = err_text

            chunks_total_live = _event_int("chunks_total")
            if chunks_total_live is not None:
                checker_chunks_total = checker_resume_base_chunks_total + chunks_total_live
            translatable_live = _event_int("translatable_segments")
            if translatable_live is not None:
                checker_translatable_segments = translatable_live
            segments_checked_live = _event_int("segments_checked")
            if segments_checked_live is not None:
                checker_segments_checked = checker_resume_base_segments_checked + segments_checked_live
            requests_total_live = _event_int("requests_total")
            if requests_total_live is not None:
                checker_requests_total = checker_resume_base_requests_total + requests_total_live
            requests_ok_live = _event_int("requests_succeeded")
            if requests_ok_live is not None:
                checker_requests_succeeded = checker_resume_base_requests_succeeded + requests_ok_live
            requests_failed_live = _event_int("requests_failed")
            if requests_failed_live is not None:
                checker_requests_failed = checker_resume_base_requests_failed + requests_failed_live
            chunks_failed_live = _event_int("chunks_failed")
            if chunks_failed_live is not None:
                checker_chunks_failed = checker_resume_base_chunks_failed + chunks_failed_live
            split_live = _event_int("split_events")
            if split_live is not None:
                checker_split_events = checker_resume_base_split_events + split_live
            suggestions_live = _event_int("suggestions_total")
            if suggestions_live is not None:
                checker_suggestions_count = checker_resume_base_suggestions + suggestions_live
            _flush_status("checker", force=True)

        status_writer.write(force=True)
        checker_trace_stats: dict[str, Any] = {}
        run_paths.checker_trace_path.parent.mkdir(parents=True, exist_ok=True)
        if not (resume and run_paths.checker_trace_path.exists()):
            run_paths.checker_trace_path.write_text("", encoding="utf-8")
        write_checker_suggestions(run_paths.checker_suggestions_path, checker_edits)
        checker_provider = (effective_checker_cfg.provider or cfg.llm.provider).strip()
        checker_model = (effective_checker_cfg.model or cfg.llm.model).strip()
        checker_custom_prompt = _read_optional_text(
            effective_checker_cfg.system_prompt_path,
            logger,
            "checker system prompt",
        )
        checker_glossary_text = (
            _read_optional_text(
                effective_checker_cfg.glossary_path,
                logger,
                "checker glossary",
            )
            if effective_checker_cfg.glossary_path
            else glossary_text
        )
        checker_client = build_llm_client(
            provider=checker_provider,
            model=checker_model,
            temperature=effective_checker_cfg.temperature,
            timeout_s=effective_checker_timeout_s,
            max_output_tokens=effective_checker_cfg.max_output_tokens,
            source_lang=cfg.llm.source_lang,
            target_lang=cfg.llm.target_lang,
            base_url=cfg.llm.base_url,
            custom_system_prompt=checker_custom_prompt,
            glossary_text=checker_glossary_text,
            glossary_prompt_text=checker_glossary_text,
            reasoning_effort=cfg.llm.reasoning_effort,
            prompt_cache_key=cfg.llm.prompt_cache_key,
            prompt_cache_retention=cfg.llm.prompt_cache_retention,
            structured_output_mode="strict",
            base_system_prompt=CHECKER_SYSTEM_PROMPT,
            on_usage=usage_totals.add,
            estimate_cost=(pricing_table.estimate_cost if cfg.pricing.enabled else None),
            pricing_currency=pricing_table.currency,
        )
        for seg in segments:
            if seg.target_tagged is None and seg.target_shielded_tagged:
                seg.context["checker_target_text"] = unshield(seg.target_shielded_tagged, seg.token_map or {})
        checker_checkpoint_seen_keys: set[tuple[str, str, str, str]] = set()
        for item in checker_edits:
            checker_checkpoint_seen_keys.add(
                (
                    str(item.get("chunk_id") or "").strip(),
                    str(item.get("segment_id") or "").strip(),
                    str(item.get("suggested_target") or "").strip(),
                    str(item.get("instruction") or "").strip(),
                )
            )

        def _on_checker_chunk_complete(chunk_id: str, chunk_edits: list[dict[str, Any]]) -> None:
            nonlocal checker_edits
            if chunk_edits:
                for item in chunk_edits:
                    key = (
                        str(chunk_id).strip(),
                        str(item.get("segment_id") or "").strip(),
                        str(item.get("suggested_target") or "").strip(),
                        str(item.get("instruction") or "").strip(),
                    )
                    if key in checker_checkpoint_seen_keys:
                        continue
                    checker_checkpoint_seen_keys.add(key)
                    checker_edits.append(item)
            # Keep a durable checkpoint even when chunk has 0 edits.
            write_checker_suggestions(run_paths.checker_suggestions_path, checker_edits)

        checker_new_edits = run_llm_checker(
            segments=segments,
            checker_cfg=effective_checker_cfg,
            checker_client=checker_client,
            logger=logger,
            trace_path=run_paths.checker_trace_path,
            stats_out=checker_trace_stats,
            progress_callback=_on_checker_progress,
            skip_chunk_ids=checker_resume_skip_chunks,
            chunk_complete_callback=_on_checker_chunk_complete,
        )

        # Safety merge in case callback was interrupted between response and checkpoint write.
        for item in checker_new_edits:
            key = (
                str(item.get("chunk_id") or "").strip(),
                str(item.get("segment_id") or "").strip(),
                str(item.get("suggested_target") or "").strip(),
                str(item.get("instruction") or "").strip(),
            )
            if key in checker_checkpoint_seen_keys:
                continue
            checker_checkpoint_seen_keys.add(key)
            checker_edits.append(item)

        checker_chunks_total = checker_resume_base_chunks_total + int(
            checker_trace_stats.get("chunks_total", checker_chunks_total)
        )
        checker_segments_checked = checker_resume_base_segments_checked + int(
            checker_trace_stats.get("segments_checked", checker_segments_checked)
        )
        checker_requests_total = checker_resume_base_requests_total + int(
            checker_trace_stats.get("requests_total", 0)
        )
        checker_requests_succeeded = checker_resume_base_requests_succeeded + int(
            checker_trace_stats.get("requests_succeeded", 0)
        )
        checker_requests_failed = checker_resume_base_requests_failed + int(
            checker_trace_stats.get("requests_failed", 0)
        )
        checker_chunks_failed = checker_resume_base_chunks_failed + int(
            checker_trace_stats.get("chunks_failed", 0)
        )
        checker_split_events = checker_resume_base_split_events + int(
            checker_trace_stats.get("split_events", checker_split_events)
        )
        checker_suggestions_count = len(checker_edits)
        if checker_segments_checked > checker_translatable_segments:
            checker_translatable_segments = checker_segments_checked
        checker_last_event = "summary"
        checker_updated_at_utc = datetime.now(timezone.utc).isoformat()
        checker_active_chunk_id = ""
        checker_active_attempt = 0
        checker_safe_edits, checker_safe_skipped = filter_checker_suggestions(
            checker_edits,
            safe_only=True,
            min_confidence=float(effective_checker_cfg.auto_apply_min_confidence),
        )
        checker_safe_suggestions_count = len(checker_safe_edits)
        write_checker_suggestions(run_paths.checker_suggestions_path, checker_edits)
        write_checker_safe_suggestions(
            run_paths.checker_suggestions_safe_path,
            source_edits=checker_edits,
            safe_edits=checker_safe_edits,
            skipped=checker_safe_skipped,
        )
        logger.info("Checker suggestions: %d (%s)", checker_suggestions_count, run_paths.checker_suggestions_path)
        logger.info(
            "Checker safe suggestions: %d/%d (%s)",
            checker_safe_suggestions_count,
            checker_suggestions_count,
            run_paths.checker_suggestions_safe_path,
        )
        if effective_checker_cfg.auto_apply_safe and checker_safe_edits:
            apply_summary = apply_checker_suggestions_to_segments(
                segments=segments,
                edits=checker_safe_edits,
                safe_only=True,
                min_confidence=float(effective_checker_cfg.auto_apply_min_confidence),
                require_current_match=True,
                logger=logger,
            )
            checker_applied_suggestions = int(apply_summary.get("applied", 0))
            if checker_applied_suggestions > 0:
                output_path.parent.mkdir(parents=True, exist_ok=True)
                doc.save(str(output_path))
                logger.info(
                    "Checker safe auto-apply wrote updated DOCX: applied=%d; output=%s",
                    checker_applied_suggestions,
                    output_path,
                )
        logger.info(
            "Checker requests: total=%d; ok=%d; failed=%d; chunks_failed=%d",
            checker_requests_total,
            checker_requests_succeeded,
            checker_requests_failed,
            checker_chunks_failed,
        )
        _flush_status("checker", force=True)
    else:
        if not run_paths.checker_trace_path.exists():
            run_paths.checker_trace_path.write_text("", encoding="utf-8")
        if not run_paths.checker_suggestions_path.exists():
            write_checker_suggestions(run_paths.checker_suggestions_path, [])
        if not run_paths.checker_suggestions_safe_path.exists():
            write_checker_safe_suggestions(
                run_paths.checker_suggestions_safe_path,
                source_edits=[],
                safe_edits=[],
                skipped=[],
            )

    if not run_paths.checker_suggestions_path.exists():
        write_checker_suggestions(run_paths.checker_suggestions_path, [])
    if not run_paths.checker_suggestions_safe_path.exists():
        write_checker_safe_suggestions(
            run_paths.checker_suggestions_safe_path,
            source_edits=[],
            safe_edits=[],
            skipped=[],
        )

    # QA outputs
    status_writer.set_phase("qa")
    status_writer.write(force=True)
    qa_html = run_paths.qa_report_path
    qa_jsonl = run_paths.qa_jsonl_path
    write_qa_report(segments, qa_html)
    write_qa_jsonl(segments, qa_jsonl)
    logger.info(f"QA report: {qa_html}")
    logger.info(f"QA jsonl: {qa_jsonl}")

    if cfg.translation_history_path:
        history_path = Path(cfg.translation_history_path)
        history_records: list[dict[str, Any]] = []
        for seg in segments:
            if seg.target_shielded_tagged is None or seg.shielded_tagged is None:
                continue
            if any(i.severity == Severity.ERROR for i in seg.issues):
                continue
            source_hash = segment_source_hash.get(seg.segment_id)
            if seg.segment_id in llm_translated_segments:
                origin = "llm"
            elif seg.segment_id in tm_hit_segments:
                origin = "tm"
            else:
                origin = "reuse"
            history_records.append(
                _build_history_record(seg, cfg, source_hash=source_hash, origin=origin)
            )
        _append_history_jsonl(history_path, history_records)
        logger.info(f"Translation history appended: {history_path} ({len(history_records)} records)")

    status_writer.set_phase("done")
    status_writer.set_done(len(segments))
    status_writer.set_usage(usage_totals.snapshot())
    status_writer.merge_metrics(
        {
            "checker_suggestions": checker_suggestions_count,
            "checker_safe_suggestions": checker_safe_suggestions_count,
            "checker_applied_suggestions": checker_applied_suggestions,
            "checker_requests_total": checker_requests_total,
            "checker_requests_succeeded": checker_requests_succeeded,
            "checker_requests_failed": checker_requests_failed,
            "checker_chunks_failed": checker_chunks_failed,
            "issues_total": sum(len(seg.issues) for seg in segments),
            "written_segments": written,
            "llm_translated": len(llm_translated_segments),
            "tm_hits": tm_hits,
            "resume_hits": resume_hits,
            "complex_in_place": complex_translated,
        }
    )
    status_writer.write(force=True)

    tm.close()


def _build_checker_only_docx_segments(
    *,
    input_path: Path,
    output_path: Path,
    include_headers: bool,
    include_footers: bool,
    logger: logging.Logger,
) -> tuple[Document, list[Segment], dict[str, int]]:
    source_doc = Document(str(input_path))
    source_segments = collect_segments(
        source_doc,
        include_headers=include_headers,
        include_footers=include_footers,
    )
    output_doc = Document(str(output_path))
    output_segments = collect_segments(
        output_doc,
        include_headers=include_headers,
        include_footers=include_footers,
    )

    source_by_location = {seg.location: seg for seg in source_segments if seg.location}
    source_locations = set(source_by_location.keys())
    output_locations = {seg.location for seg in output_segments if seg.location}
    missing_source = 0
    skipped_no_latin = 0
    checker_candidates = 0
    untranslated_equal = 0

    for seg in output_segments:
        target_text = str(seg.source_plain or "")
        source_seg = source_by_location.get(seg.location)
        if source_seg is None:
            missing_source += 1
            seg.issues.append(
                Issue(
                    code="checker_source_missing",
                    severity=Severity.WARN,
                    message=(
                        "Checker-only alignment: source segment is missing by location; "
                        "segment skipped in checker pass."
                    ),
                    details={"location": seg.location},
                )
            )
            seg.target_tagged = ""
            seg.context["checker_target_text"] = ""
            continue

        seg.source_plain = source_seg.source_plain
        seg.context["part"] = source_seg.context.get("part")
        seg.context["section"] = source_seg.context.get("section")
        seg.context["section_header"] = source_seg.context.get("section_header")
        seg.context["paragraph_style"] = source_seg.context.get("paragraph_style")
        seg.context["in_table"] = bool(source_seg.context.get("in_table"))
        seg.context["in_textbox"] = bool(source_seg.context.get("in_textbox"))
        if source_seg.context.get("row_index") is not None:
            seg.context["row_index"] = source_seg.context.get("row_index")
        if source_seg.context.get("col_index") is not None:
            seg.context["col_index"] = source_seg.context.get("col_index")
        if source_seg.context.get("table_index") is not None:
            seg.context["table_index"] = source_seg.context.get("table_index")
        if source_seg.context.get("is_toc_entry") is not None:
            seg.context["is_toc_entry"] = bool(source_seg.context.get("is_toc_entry"))

        if _should_translate_segment_text(seg.source_plain):
            seg.target_tagged = target_text
            seg.context["checker_target_text"] = target_text
            if target_text.strip():
                checker_candidates += 1
            if normalize_text(seg.source_plain) == normalize_text(target_text):
                untranslated_equal += 1
        else:
            seg.target_tagged = ""
            skipped_no_latin += 1
            seg.context["checker_target_text"] = ""

    missing_target = max(0, len(source_locations - output_locations))
    if missing_source > 0 or missing_target > 0:
        logger.warning(
            "Checker-only alignment gaps: missing_source=%d; missing_target=%d",
            int(missing_source),
            int(missing_target),
        )
    logger.info(
        "Checker-only alignment: source_segments=%d; output_segments=%d; checker_candidates=%d; "
        "skipped_no_latin=%d; untranslated_equal=%d",
        len(source_segments),
        len(output_segments),
        int(checker_candidates),
        int(skipped_no_latin),
        int(untranslated_equal),
    )
    return (
        output_doc,
        output_segments,
        {
            "source_segments": len(source_segments),
            "output_segments": len(output_segments),
            "checker_candidates": int(checker_candidates),
            "skipped_no_latin": int(skipped_no_latin),
            "untranslated_equal": int(untranslated_equal),
            "missing_source": int(missing_source),
            "missing_target": int(missing_target),
        },
    )


def run_docx_checker_only(
    input_path: Path,
    output_path: Path,
    cfg: PipelineConfig,
    *,
    resume: bool = False,
    max_segments: int | None = None,
) -> None:
    logger = setup_logging(Path(cfg.log_path))
    run_paths = resolve_run_paths(cfg, output_path=output_path)
    run_paths.run_dir.mkdir(parents=True, exist_ok=True)
    ensure_dashboard_html(run_paths.dashboard_html_path.parent, filename=run_paths.dashboard_html_path.name)

    logger.info("Input (source): %s", input_path)
    logger.info("Output (translated DOCX): %s", output_path)
    logger.info("Run dir: %s", run_paths.run_dir)
    logger.info("Run status: %s", run_paths.status_path)
    logger.info("Dashboard HTML: %s", run_paths.dashboard_html_path)

    if output_path.suffix.lower() != ".docx":
        raise RuntimeError("Checker-only mode is currently supported for DOCX output only.")
    if not output_path.exists():
        raise FileNotFoundError(f"Checker-only mode requires existing translated output: {output_path}")

    effective_checker_cfg = cfg.checker
    if not effective_checker_cfg.enabled:
        effective_checker_cfg = cfg.checker.__class__(**{**cfg.checker.__dict__, "enabled": True})
        logger.info("Checker-only mode: checker.enabled was false; enabled in-memory for this run.")

    doc, segments, alignment_stats = _build_checker_only_docx_segments(
        input_path=input_path,
        output_path=output_path,
        include_headers=cfg.include_headers,
        include_footers=cfg.include_footers,
        logger=logger,
    )
    if max_segments is not None:
        if max_segments < 0:
            raise ValueError(f"max_segments must be >= 0, got {max_segments}")
        segments = segments[:max_segments]
        logger.info("Segment limit enabled for checker-only mode: %d segments", len(segments))
    logger.info("Checker-only segments loaded: %d", len(segments))

    usage_totals = UsageTotals()
    pricing_table = PricingTable.empty(currency=cfg.pricing.currency)
    if cfg.pricing.enabled:
        if cfg.pricing.pricing_path:
            try:
                pricing_table = load_pricing_table(cfg.pricing.pricing_path, currency=cfg.pricing.currency)
                logger.info("Pricing table loaded: %s", cfg.pricing.pricing_path)
            except Exception as exc:
                logger.warning("Failed to load pricing table (%s): %s", cfg.pricing.pricing_path, exc)
        else:
            logger.warning("pricing.enabled=true but pricing.pricing_path is empty")

    status_writer = RunStatusWriter(
        path=run_paths.status_path,
        run_id=run_paths.run_id,
        total_segments=len(segments),
        flush_every_n_updates=cfg.run.status_flush_every_n_segments,
    )

    def _to_dashboard_link(path: Path) -> str:
        try:
            rel = path.resolve().relative_to(run_paths.run_dir.resolve())
            return rel.as_posix()
        except Exception:
            return str(path)

    status_writer.merge_paths(
        {
            "run_dir": str(run_paths.run_dir),
            "output": str(output_path),
            "qa_report": _to_dashboard_link(run_paths.qa_report_path),
            "qa_jsonl": _to_dashboard_link(run_paths.qa_jsonl_path),
            "dashboard_html": _to_dashboard_link(run_paths.dashboard_html_path),
            "checker_suggestions": _to_dashboard_link(run_paths.checker_suggestions_path),
            "checker_suggestions_safe": _to_dashboard_link(run_paths.checker_suggestions_safe_path),
            "checker_trace": _to_dashboard_link(run_paths.checker_trace_path),
        }
    )
    status_writer.set_phase("prepare")
    status_writer.set_done(0)
    status_writer.set_usage(usage_totals.snapshot())
    status_writer.write(force=True)

    _attach_neighbor_snippets(segments, cfg)
    glossary_text = _read_optional_text(cfg.llm.glossary_path, logger, "glossary")
    glossary_matchers = build_glossary_matchers(glossary_text) if glossary_text else ()
    matched_glossary_segments = 0
    if glossary_matchers:
        for seg in segments:
            if not _should_translate_segment_text(seg.source_plain):
                continue
            checker_terms = _build_matched_glossary_context(
                seg.source_plain,
                glossary_matchers,
                limit=min(12, cfg.llm.glossary_match_limit),
            )
            if checker_terms:
                seg.context["matched_glossary_terms"] = checker_terms
                matched_glossary_segments += 1
        logger.info("Checker-only glossary hints attached to %d segments", int(matched_glossary_segments))

    checker_custom_prompt = _read_optional_text(
        effective_checker_cfg.system_prompt_path,
        logger,
        "checker system prompt",
    )
    checker_glossary_text = (
        _read_optional_text(
            effective_checker_cfg.glossary_path,
            logger,
            "checker glossary",
        )
        if effective_checker_cfg.glossary_path
        else glossary_text
    )

    effective_checker_timeout_s = float(effective_checker_cfg.timeout_s)
    if cfg.llm.auto_model_sizing:
        checker_provider_for_sizing = (effective_checker_cfg.provider or cfg.llm.provider).strip()
        checker_model_for_sizing = (effective_checker_cfg.model or cfg.llm.model).strip()
        source_lengths = [len(seg.source_plain or "") for seg in segments]
        prompt_chars = len(checker_custom_prompt or "") + len(checker_glossary_text or "")
        sizing = recommend_runtime_model_sizing(
            provider=cfg.llm.provider,
            model=cfg.llm.model,
            checker_provider=checker_provider_for_sizing,
            checker_model=checker_model_for_sizing,
            source_char_lengths=source_lengths,
            prompt_chars=prompt_chars,
            batch_segments=1,
            batch_max_chars=max(4000, int(cfg.llm.batch_max_chars)),
            max_output_tokens=cfg.llm.max_output_tokens,
            context_window_chars=0,
            checker_pages_per_chunk=effective_checker_cfg.pages_per_chunk,
            checker_fallback_segments_per_chunk=effective_checker_cfg.fallback_segments_per_chunk,
            checker_max_output_tokens=effective_checker_cfg.max_output_tokens,
        )
        effective_checker_cfg = effective_checker_cfg.__class__(
            **{
                **effective_checker_cfg.__dict__,
                "pages_per_chunk": sizing.checker_pages_per_chunk,
                "fallback_segments_per_chunk": sizing.checker_fallback_segments_per_chunk,
                "max_output_tokens": sizing.checker_max_output_tokens,
            }
        )
        for note in sizing.notes:
            logger.info("Checker-only model auto-sizing: %s", note)

    checker_timeout_before = float(effective_checker_cfg.timeout_s)
    effective_checker_timeout_s = recommend_checker_timeout_s(
        timeout_s=checker_timeout_before,
        fallback_segments_per_chunk=int(effective_checker_cfg.fallback_segments_per_chunk),
    )
    if float(effective_checker_timeout_s) > checker_timeout_before:
        logger.info(
            "Checker-only mode auto-raised checker timeout_s to %.1f (prev=%.1f; fallback_segments_per_chunk=%d)",
            float(effective_checker_timeout_s),
            float(checker_timeout_before),
            int(effective_checker_cfg.fallback_segments_per_chunk),
        )

    checker_suggestions_count = 0
    checker_safe_suggestions_count = 0
    checker_applied_suggestions = 0
    checker_requests_total = 0
    checker_requests_succeeded = 0
    checker_requests_failed = 0
    checker_chunks_failed = 0
    checker_chunks_total = 0
    checker_translatable_segments = 0
    checker_segments_checked = 0
    checker_split_events = 0
    checker_active_chunk_id = ""
    checker_active_attempt = 0
    checker_last_event = "start"
    checker_last_error = ""
    checker_started_at_utc = datetime.now(timezone.utc).isoformat()
    checker_updated_at_utc = checker_started_at_utc
    processed_segments = 0

    def _flush_status(phase: str, *, force: bool = False) -> None:
        step = max(1, int(cfg.run.status_flush_every_n_segments))
        if not force and processed_segments % step != 0:
            return
        checker_progress_pct = (
            float(100.0 * checker_segments_checked / checker_translatable_segments)
            if checker_translatable_segments > 0
            else 0.0
        )
        status_writer.set_phase(phase)
        status_writer.set_done(processed_segments)
        status_writer.set_usage(usage_totals.snapshot())
        status_writer.merge_metrics(
            {
                "tm_hits": 0,
                "resume_hits": 0,
                "tagging_errors": 0,
                "llm_queue": 0,
                "llm_translated": 0,
                "complex_in_place": 0,
                "toc_in_place": 0,
                "checker_suggestions": checker_suggestions_count,
                "checker_safe_suggestions": checker_safe_suggestions_count,
                "checker_applied_suggestions": checker_applied_suggestions,
                "checker_requests_total": checker_requests_total,
                "checker_requests_succeeded": checker_requests_succeeded,
                "checker_requests_failed": checker_requests_failed,
                "checker_chunks_failed": checker_chunks_failed,
                "checker_chunks_total": checker_chunks_total,
                "checker_translatable_segments": checker_translatable_segments,
                "checker_segments_checked": checker_segments_checked,
                "checker_split_events": checker_split_events,
                "checker_progress_pct": checker_progress_pct,
                "checker_active_chunk": checker_active_chunk_id,
                "checker_active_attempt": checker_active_attempt,
                "checker_last_event": checker_last_event,
                "checker_last_error": checker_last_error,
                "checker_started_at": checker_started_at_utc,
                "checker_updated_at": checker_updated_at_utc,
                "issues_total": sum(len(seg.issues) for seg in segments),
                "checker_alignment_source_segments": int(alignment_stats.get("source_segments", 0)),
                "checker_alignment_output_segments": int(alignment_stats.get("output_segments", 0)),
                "checker_alignment_candidates": int(alignment_stats.get("checker_candidates", 0)),
                "checker_alignment_skipped_no_latin": int(alignment_stats.get("skipped_no_latin", 0)),
                "checker_alignment_untranslated_equal": int(alignment_stats.get("untranslated_equal", 0)),
                "checker_alignment_missing_source": int(alignment_stats.get("missing_source", 0)),
                "checker_alignment_missing_target": int(alignment_stats.get("missing_target", 0)),
            }
        )
        status_writer.write(force=force)

    _flush_status("prepare", force=True)

    checker_resume_skip_chunks: set[str] = set()
    checker_resume_base_requests_total = 0
    checker_resume_base_requests_succeeded = 0
    checker_resume_base_requests_failed = 0
    checker_resume_base_chunks_failed = 0
    checker_resume_base_split_events = 0
    checker_resume_base_suggestions = 0
    checker_resume_base_chunks_total = 0
    checker_resume_base_segments_checked = 0
    checker_resume_state: dict[str, Any] = {}
    existing_checker_edits: list[dict[str, Any]] = []
    checker_edits: list[dict[str, Any]] = []

    if resume:
        if run_paths.checker_suggestions_path.exists():
            try:
                existing_checker_edits = read_checker_suggestions(run_paths.checker_suggestions_path)
            except Exception as exc:
                logger.warning("Checker-only resume: failed to read existing suggestions: %s", exc)
                existing_checker_edits = []
        if run_paths.checker_trace_path.exists():
            try:
                checker_resume_state = read_checker_trace_resume_state(run_paths.checker_trace_path)
            except Exception as exc:
                logger.warning("Checker-only resume: failed to read checker trace state: %s", exc)
                checker_resume_state = {}
            completed_chunks_raw = checker_resume_state.get("completed_chunk_ids") or []
            if existing_checker_edits and isinstance(completed_chunks_raw, list):
                checker_resume_skip_chunks = {str(item).strip() for item in completed_chunks_raw if str(item).strip()}
                if checker_resume_skip_chunks:
                    checker_resume_base_chunks_total = len(checker_resume_skip_chunks)
                    logger.info(
                        "Checker-only resume: will skip %d completed chunks from previous attempt.",
                        len(checker_resume_skip_chunks),
                    )
            elif completed_chunks_raw:
                logger.warning(
                    "Checker-only resume: trace has completed chunks, but suggestions checkpoint is missing; "
                    "rerunning checker from the beginning to avoid losing prior edits."
                )

            try:
                checker_translatable_segments = max(
                    checker_translatable_segments,
                    int(checker_resume_state.get("translatable_segments") or 0),
                )
                checker_resume_base_requests_total = max(0, int(checker_resume_state.get("requests_total") or 0))
                checker_resume_base_requests_succeeded = max(
                    0,
                    int(checker_resume_state.get("requests_succeeded") or 0),
                )
                checker_resume_base_requests_failed = max(
                    0,
                    int(checker_resume_state.get("requests_failed") or 0),
                )
                checker_resume_base_chunks_failed = max(0, int(checker_resume_state.get("chunks_failed") or 0))
                checker_resume_base_split_events = max(0, int(checker_resume_state.get("split_events") or 0))
                if existing_checker_edits:
                    checker_resume_base_suggestions = max(
                        checker_resume_base_suggestions,
                        int(checker_resume_state.get("suggestions_total") or 0),
                    )
            except Exception:
                logger.debug("Checker-only resume: failed to normalize base trace counters", exc_info=True)

        if existing_checker_edits and checker_resume_skip_chunks:
            checker_edits = attach_checker_edits_to_segments(
                segments=segments,
                edits=existing_checker_edits,
                logger=logger,
            )
            checker_suggestions_count = len(checker_edits)
            checker_resume_base_suggestions = checker_suggestions_count
            logger.info(
                "Checker-only resume: loaded %d suggestions checkpoint entries.",
                checker_suggestions_count,
            )
        elif existing_checker_edits:
            logger.info(
                "Checker-only resume: suggestions checkpoint exists, but no completed chunk map; "
                "checker will rerun from the beginning."
            )

    status_writer.set_phase("checker")
    _flush_status("checker", force=True)
    status_writer.write(force=True)

    def _on_checker_progress(event: dict[str, Any]) -> None:
        nonlocal checker_suggestions_count
        nonlocal checker_requests_total
        nonlocal checker_requests_succeeded
        nonlocal checker_requests_failed
        nonlocal checker_chunks_failed
        nonlocal checker_chunks_total
        nonlocal checker_translatable_segments
        nonlocal checker_segments_checked
        nonlocal checker_split_events
        nonlocal checker_active_chunk_id
        nonlocal checker_active_attempt
        nonlocal checker_last_event
        nonlocal checker_last_error
        nonlocal checker_started_at_utc
        nonlocal checker_updated_at_utc
        nonlocal processed_segments

        def _event_int(key: str) -> int | None:
            value = event.get(key)
            if value is None:
                return None
            try:
                return max(0, int(value))
            except Exception:
                return None

        event_name = str(event.get("event") or "").strip()
        now_iso = datetime.now(timezone.utc).isoformat()
        if not checker_started_at_utc:
            checker_started_at_utc = now_iso
        checker_updated_at_utc = now_iso
        if event_name:
            checker_last_event = event_name

        chunk_id = str(event.get("chunk_id") or "").strip()
        if chunk_id:
            checker_active_chunk_id = chunk_id
        attempt = _event_int("attempt")
        if attempt is not None:
            checker_active_attempt = attempt
        if event_name in {"response", "failed_chunk", "summary"}:
            checker_active_attempt = 0
        if event_name == "summary":
            checker_active_chunk_id = ""

        err_text = str(event.get("error") or "").strip()
        if err_text:
            checker_last_error = err_text

        chunks_total_live = _event_int("chunks_total")
        if chunks_total_live is not None:
            checker_chunks_total = checker_resume_base_chunks_total + chunks_total_live
        translatable_live = _event_int("translatable_segments")
        if translatable_live is not None:
            checker_translatable_segments = translatable_live
        segments_checked_live = _event_int("segments_checked")
        if segments_checked_live is not None:
            checker_segments_checked = checker_resume_base_segments_checked + segments_checked_live
            processed_segments = min(len(segments), checker_segments_checked)
        requests_total_live = _event_int("requests_total")
        if requests_total_live is not None:
            checker_requests_total = checker_resume_base_requests_total + requests_total_live
        requests_ok_live = _event_int("requests_succeeded")
        if requests_ok_live is not None:
            checker_requests_succeeded = checker_resume_base_requests_succeeded + requests_ok_live
        requests_failed_live = _event_int("requests_failed")
        if requests_failed_live is not None:
            checker_requests_failed = checker_resume_base_requests_failed + requests_failed_live
        chunks_failed_live = _event_int("chunks_failed")
        if chunks_failed_live is not None:
            checker_chunks_failed = checker_resume_base_chunks_failed + chunks_failed_live
        split_live = _event_int("split_events")
        if split_live is not None:
            checker_split_events = checker_resume_base_split_events + split_live
        suggestions_live = _event_int("suggestions_total")
        if suggestions_live is not None:
            checker_suggestions_count = checker_resume_base_suggestions + suggestions_live
        _flush_status("checker", force=True)

    run_paths.checker_trace_path.parent.mkdir(parents=True, exist_ok=True)
    if not (resume and run_paths.checker_trace_path.exists()):
        run_paths.checker_trace_path.write_text("", encoding="utf-8")
    write_checker_suggestions(run_paths.checker_suggestions_path, checker_edits)

    checker_provider = (effective_checker_cfg.provider or cfg.llm.provider).strip()
    checker_model = (effective_checker_cfg.model or cfg.llm.model).strip()
    checker_client = build_llm_client(
        provider=checker_provider,
        model=checker_model,
        temperature=effective_checker_cfg.temperature,
        timeout_s=effective_checker_timeout_s,
        max_output_tokens=effective_checker_cfg.max_output_tokens,
        source_lang=cfg.llm.source_lang,
        target_lang=cfg.llm.target_lang,
        base_url=cfg.llm.base_url,
        custom_system_prompt=checker_custom_prompt,
        glossary_text=checker_glossary_text,
        glossary_prompt_text=checker_glossary_text,
        reasoning_effort=cfg.llm.reasoning_effort,
        prompt_cache_key=cfg.llm.prompt_cache_key,
        prompt_cache_retention=cfg.llm.prompt_cache_retention,
        structured_output_mode="strict",
        base_system_prompt=CHECKER_SYSTEM_PROMPT,
        on_usage=usage_totals.add,
        estimate_cost=(pricing_table.estimate_cost if cfg.pricing.enabled else None),
        pricing_currency=pricing_table.currency,
    )

    checker_checkpoint_seen_keys: set[tuple[str, str, str, str]] = set()
    for item in checker_edits:
        checker_checkpoint_seen_keys.add(
            (
                str(item.get("chunk_id") or "").strip(),
                str(item.get("segment_id") or "").strip(),
                str(item.get("suggested_target") or "").strip(),
                str(item.get("instruction") or "").strip(),
            )
        )

    def _on_checker_chunk_complete(chunk_id: str, chunk_edits: list[dict[str, Any]]) -> None:
        nonlocal checker_edits
        if chunk_edits:
            for item in chunk_edits:
                key = (
                    str(chunk_id).strip(),
                    str(item.get("segment_id") or "").strip(),
                    str(item.get("suggested_target") or "").strip(),
                    str(item.get("instruction") or "").strip(),
                )
                if key in checker_checkpoint_seen_keys:
                    continue
                checker_checkpoint_seen_keys.add(key)
                checker_edits.append(item)
        write_checker_suggestions(run_paths.checker_suggestions_path, checker_edits)

    checker_trace_stats: dict[str, Any] = {}
    checker_new_edits = run_llm_checker(
        segments=segments,
        checker_cfg=effective_checker_cfg,
        checker_client=checker_client,
        logger=logger,
        trace_path=run_paths.checker_trace_path,
        stats_out=checker_trace_stats,
        progress_callback=_on_checker_progress,
        skip_chunk_ids=checker_resume_skip_chunks,
        chunk_complete_callback=_on_checker_chunk_complete,
    )

    for item in checker_new_edits:
        key = (
            str(item.get("chunk_id") or "").strip(),
            str(item.get("segment_id") or "").strip(),
            str(item.get("suggested_target") or "").strip(),
            str(item.get("instruction") or "").strip(),
        )
        if key in checker_checkpoint_seen_keys:
            continue
        checker_checkpoint_seen_keys.add(key)
        checker_edits.append(item)

    checker_chunks_total = checker_resume_base_chunks_total + int(checker_trace_stats.get("chunks_total", checker_chunks_total))
    checker_segments_checked = checker_resume_base_segments_checked + int(
        checker_trace_stats.get("segments_checked", checker_segments_checked)
    )
    checker_requests_total = checker_resume_base_requests_total + int(checker_trace_stats.get("requests_total", 0))
    checker_requests_succeeded = checker_resume_base_requests_succeeded + int(
        checker_trace_stats.get("requests_succeeded", 0)
    )
    checker_requests_failed = checker_resume_base_requests_failed + int(checker_trace_stats.get("requests_failed", 0))
    checker_chunks_failed = checker_resume_base_chunks_failed + int(checker_trace_stats.get("chunks_failed", 0))
    checker_split_events = checker_resume_base_split_events + int(
        checker_trace_stats.get("split_events", checker_split_events)
    )
    checker_suggestions_count = len(checker_edits)
    if checker_segments_checked > checker_translatable_segments:
        checker_translatable_segments = checker_segments_checked
    checker_last_event = "summary"
    checker_updated_at_utc = datetime.now(timezone.utc).isoformat()
    checker_active_chunk_id = ""
    checker_active_attempt = 0
    processed_segments = min(len(segments), checker_segments_checked)

    checker_safe_edits, checker_safe_skipped = filter_checker_suggestions(
        checker_edits,
        safe_only=True,
        min_confidence=float(effective_checker_cfg.auto_apply_min_confidence),
    )
    checker_safe_suggestions_count = len(checker_safe_edits)
    write_checker_suggestions(run_paths.checker_suggestions_path, checker_edits)
    write_checker_safe_suggestions(
        run_paths.checker_suggestions_safe_path,
        source_edits=checker_edits,
        safe_edits=checker_safe_edits,
        skipped=checker_safe_skipped,
    )
    logger.info("Checker suggestions: %d (%s)", checker_suggestions_count, run_paths.checker_suggestions_path)
    logger.info(
        "Checker safe suggestions: %d/%d (%s)",
        checker_safe_suggestions_count,
        checker_suggestions_count,
        run_paths.checker_suggestions_safe_path,
    )

    if effective_checker_cfg.auto_apply_safe and checker_safe_edits:
        apply_summary = apply_checker_suggestions_to_segments(
            segments=segments,
            edits=checker_safe_edits,
            safe_only=True,
            min_confidence=float(effective_checker_cfg.auto_apply_min_confidence),
            require_current_match=True,
            logger=logger,
        )
        checker_applied_suggestions = int(apply_summary.get("applied", 0))
        if checker_applied_suggestions > 0:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            logger.info(
                "Checker-only safe auto-apply wrote updated DOCX: applied=%d; output=%s",
                checker_applied_suggestions,
                output_path,
            )

    logger.info(
        "Checker-only requests: total=%d; ok=%d; failed=%d; chunks_failed=%d",
        checker_requests_total,
        checker_requests_succeeded,
        checker_requests_failed,
        checker_chunks_failed,
    )
    _flush_status("checker", force=True)

    if not run_paths.checker_suggestions_path.exists():
        write_checker_suggestions(run_paths.checker_suggestions_path, [])
    if not run_paths.checker_suggestions_safe_path.exists():
        write_checker_safe_suggestions(
            run_paths.checker_suggestions_safe_path,
            source_edits=[],
            safe_edits=[],
            skipped=[],
        )

    status_writer.set_phase("qa")
    status_writer.write(force=True)
    qa_html = run_paths.qa_report_path
    qa_jsonl = run_paths.qa_jsonl_path
    write_qa_report(segments, qa_html)
    write_qa_jsonl(segments, qa_jsonl)
    logger.info("QA report: %s", qa_html)
    logger.info("QA jsonl: %s", qa_jsonl)

    processed_segments = len(segments)
    status_writer.set_phase("done")
    status_writer.set_done(processed_segments)
    status_writer.set_usage(usage_totals.snapshot())
    status_writer.merge_metrics(
        {
            "checker_suggestions": checker_suggestions_count,
            "checker_safe_suggestions": checker_safe_suggestions_count,
            "checker_applied_suggestions": checker_applied_suggestions,
            "checker_requests_total": checker_requests_total,
            "checker_requests_succeeded": checker_requests_succeeded,
            "checker_requests_failed": checker_requests_failed,
            "checker_chunks_failed": checker_chunks_failed,
            "checker_chunks_total": checker_chunks_total,
            "checker_translatable_segments": checker_translatable_segments,
            "checker_segments_checked": checker_segments_checked,
            "checker_split_events": checker_split_events,
            "issues_total": sum(len(seg.issues) for seg in segments),
            "written_segments": 0,
            "llm_translated": 0,
            "tm_hits": 0,
            "resume_hits": 0,
            "complex_in_place": 0,
            "checker_alignment_source_segments": int(alignment_stats.get("source_segments", 0)),
            "checker_alignment_output_segments": int(alignment_stats.get("output_segments", 0)),
            "checker_alignment_candidates": int(alignment_stats.get("checker_candidates", 0)),
            "checker_alignment_skipped_no_latin": int(alignment_stats.get("skipped_no_latin", 0)),
            "checker_alignment_untranslated_equal": int(alignment_stats.get("untranslated_equal", 0)),
            "checker_alignment_missing_source": int(alignment_stats.get("missing_source", 0)),
            "checker_alignment_missing_target": int(alignment_stats.get("missing_target", 0)),
        }
    )
    status_writer.write(force=True)

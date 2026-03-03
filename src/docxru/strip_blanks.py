"""Strip excessive empty paragraphs and clean up soft line breaks in DOCX documents.

Two independent passes:
  Pass 1 – strip_blank_paragraphs(): remove consecutive empty <w:p> elements.
  Pass 2 – cleanup_soft_linebreaks(): replace mid-sentence <w:br> with spaces.

Both passes operate on the lxml tree directly and are safe for formatting.
"""

from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

# Tags that indicate non-text visible content (paragraph must NOT be removed)
_MC_ALTERNATE_CONTENT = "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent"

_VISUAL_CONTENT_TAGS = frozenset(
    {
        qn("w:drawing"),
        qn("w:pict"),
        _MC_ALTERNATE_CONTENT,
        qn("w:fldChar"),
        qn("w:fldSimple"),
    }
)

# Tags in pPr that make a paragraph structurally important
_PROTECTED_PPR_TAGS = frozenset(
    {
        qn("w:sectPr"),  # section break – removing kills page layout
    }
)

# Tags at paragraph level that make it protected
_PROTECTED_P_TAGS = frozenset(
    {
        qn("w:bookmarkStart"),  # bookmark anchor – orphans paired bookmarkEnd
    }
)


def _is_empty_paragraph(paragraph) -> bool:
    """Return True if the paragraph has no visible text content."""
    if (paragraph.text or "").strip():
        return False
    # Check all runs for any text
    p_elem = paragraph._p
    for t_elem in p_elem.iter(qn("w:t")):
        if (t_elem.text or "").strip():
            return False
    return True


def _is_protected_paragraph(paragraph) -> bool:
    """Return True if the paragraph must NOT be removed (structural/visual content)."""
    p_elem = paragraph._p

    # Check pPr children (e.g. w:sectPr)
    p_pr = p_elem.find(qn("w:pPr"))
    if p_pr is not None:
        for tag in _PROTECTED_PPR_TAGS:
            if p_pr.find(tag) is not None:
                return True

    # Check direct children of <w:p> (bookmarks, etc.)
    for tag in _PROTECTED_P_TAGS:
        if p_elem.find(tag) is not None:
            return True

    # Check for visual content inside runs
    for child in p_elem.iter():
        if child.tag in _VISUAL_CONTENT_TAGS:
            return True

    # Check for page/column breaks inside runs
    for br_elem in p_elem.iter(qn("w:br")):
        br_type = br_elem.get(qn("w:type"), "")
        if br_type in ("page", "column"):
            return True

    return False


def _can_remove(paragraph) -> bool:
    """Return True if this paragraph is empty AND safe to remove."""
    return _is_empty_paragraph(paragraph) and not _is_protected_paragraph(paragraph)


def _remove_paragraph_element(paragraph) -> None:
    """Remove a paragraph's XML element from its parent."""
    p_elem = paragraph._p
    parent = p_elem.getparent()
    if parent is not None:
        parent.remove(p_elem)


# ---------------------------------------------------------------------------
# Pass 1: Strip blank paragraphs
# ---------------------------------------------------------------------------


def _process_paragraph_sequence(
    paragraphs: list,
    *,
    keep_singles: bool,
    max_consecutive: int,
    min_remaining: int = 0,
) -> dict[str, int]:
    """Process a list of paragraphs, removing excessive empty ones.

    Args:
        paragraphs: list of python-docx Paragraph objects.
        keep_singles: preserve isolated single empty paragraphs.
        max_consecutive: for groups of 2+, keep at most this many.
        min_remaining: minimum paragraphs that must survive (1 for table cells).

    Returns:
        Stats dict with removal counts.
    """
    stats = {
        "total": len(paragraphs),
        "total_empty": 0,
        "removed": 0,
        "groups_processed": 0,
        "groups_kept_single": 0,
        "groups_reduced": 0,
        "groups_removed_entirely": 0,
    }

    # Phase 1: identify groups of consecutive removable-empty paragraphs
    groups: list[list] = []
    current_group: list = []

    for p in paragraphs:
        if _can_remove(p):
            stats["total_empty"] += 1
            current_group.append(p)
        else:
            if current_group:
                groups.append(current_group)
                current_group = []
    if current_group:
        groups.append(current_group)

    # Phase 2: decide which to remove
    to_remove: list = []
    for group in groups:
        stats["groups_processed"] += 1
        group_size = len(group)

        if group_size == 1 and keep_singles:
            stats["groups_kept_single"] += 1
            continue

        if group_size == 1 and not keep_singles:
            to_remove.extend(group)
            stats["groups_removed_entirely"] += 1
            continue

        # Group of 2+
        keep_count = min(max_consecutive, group_size)
        removable = group[keep_count:]
        to_remove.extend(removable)
        if keep_count == 0:
            stats["groups_removed_entirely"] += 1
        else:
            stats["groups_reduced"] += 1

    # Phase 3: enforce min_remaining constraint
    if min_remaining > 0:
        non_empty_count = stats["total"] - stats["total_empty"]
        kept_empty = stats["total_empty"] - len(to_remove)
        surviving = non_empty_count + kept_empty
        while surviving < min_remaining and to_remove:
            to_remove.pop()  # keep one more
            surviving += 1

    # Phase 4: remove (reverse order for safety, though lxml handles it fine)
    for p in reversed(to_remove):
        _remove_paragraph_element(p)
    stats["removed"] = len(to_remove)

    return stats


def strip_blank_paragraphs(
    doc: Document,
    *,
    keep_singles: bool = True,
    max_consecutive: int = 1,
    process_tables: bool = True,
) -> dict[str, int]:
    """Remove excessive empty paragraphs from a DOCX document.

    Args:
        doc: python-docx Document object (modified in place).
        keep_singles: preserve isolated single empty paragraphs.
        max_consecutive: for groups of 2+, keep at most this many (0 = remove all).
        process_tables: also strip blanks inside table cells.

    Returns:
        Statistics dictionary.
    """
    result: dict[str, int] = {}

    # --- Body ---
    body_stats = _process_paragraph_sequence(
        list(doc.paragraphs),
        keep_singles=keep_singles,
        max_consecutive=max_consecutive,
        min_remaining=0,
    )
    result["total_body_paragraphs"] = body_stats["total"]
    result["total_empty_body"] = body_stats["total_empty"]
    result["removed_body"] = body_stats["removed"]
    result["groups_processed"] = body_stats["groups_processed"]
    result["groups_kept_single"] = body_stats["groups_kept_single"]
    result["groups_reduced"] = body_stats["groups_reduced"]
    result["groups_removed_entirely"] = body_stats["groups_removed_entirely"]

    # --- Tables ---
    result["total_table_paragraphs"] = 0
    result["total_empty_table"] = 0
    result["removed_table"] = 0

    if process_tables:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_paras = list(cell.paragraphs)
                    if len(cell_paras) <= 1:
                        continue  # nothing to strip
                    cell_stats = _process_paragraph_sequence(
                        cell_paras,
                        keep_singles=keep_singles,
                        max_consecutive=max_consecutive,
                        min_remaining=1,  # OOXML: at least 1 paragraph per cell
                    )
                    result["total_table_paragraphs"] += cell_stats["total"]
                    result["total_empty_table"] += cell_stats["total_empty"]
                    result["removed_table"] += cell_stats["removed"]

    return result


# ---------------------------------------------------------------------------
# Pass 2: Cleanup soft line breaks
# ---------------------------------------------------------------------------

_W_BR = qn("w:br")
_W_T = qn("w:t")
_W_R = qn("w:r")
_W_TYPE = qn("w:type")
_W_CLEAR = qn("w:clear")
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _get_paragraph_alignment(paragraph) -> str | None:
    """Return paragraph alignment string or None."""
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is not None:
        jc = p_pr.find(qn("w:jc"))
        if jc is not None:
            return jc.get(qn("w:val"), "")
    return None


def _get_paragraph_style_name(paragraph) -> str:
    """Return the style name of a paragraph (lowercase)."""
    try:
        return (paragraph.style.name or "").lower()
    except Exception:
        return ""


def _is_short_frame_paragraph(paragraph, *, max_text_len: int = 120) -> bool:
    """Return True if the paragraph has w:framePr and short text content.

    Short framed paragraphs are typically titles, headers, or labels where
    soft line breaks serve as intentional visual line separators and should
    not be cleaned up. Content-heavy framed paragraphs (>max_text_len chars)
    benefit from break cleanup (e.g., translated text that wraps badly).
    """
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return False
    if p_pr.find(qn("w:framePr")) is None:
        return False
    # It's a frame — check text length
    text = ""
    for t_elem in paragraph._p.iter(qn("w:t")):
        text += (t_elem.text or "")
        if len(text) > max_text_len:
            return False  # Long content frame — don't skip
    return True  # Short text in frame — skip


def _should_skip_paragraph(
    paragraph, *, skip_toc: bool, skip_centered: bool, skip_short_frames: bool
) -> bool:
    """Return True if this paragraph should be skipped for break cleanup."""
    if skip_toc:
        style = _get_paragraph_style_name(paragraph)
        if "toc" in style:
            return True
    if skip_centered:
        align = _get_paragraph_alignment(paragraph)
        if align in ("center", "right"):
            return True
    if skip_short_frames:
        if _is_short_frame_paragraph(paragraph):
            return True
    return False


def _text_before_element(elem) -> str:
    """Get the last character of text content before this element within the same run."""
    # Check text of preceding siblings within the same run
    prev = elem.getprevious()
    while prev is not None:
        if prev.tag == _W_T and prev.text:
            return prev.text[-1]
        prev = prev.getprevious()
    return ""


def _text_after_element(elem) -> str:
    """Get the first character of text content after this element within the same run."""
    nxt = elem.getnext()
    while nxt is not None:
        if nxt.tag == _W_T and nxt.text:
            return nxt.text[0]
        nxt = nxt.getnext()
    return ""


def _text_in_next_run(run_elem) -> str:
    """Get the first text character from the next run in the paragraph."""
    nxt = run_elem.getnext()
    while nxt is not None:
        if nxt.tag == _W_R:
            for t_elem in nxt.iter(_W_T):
                if t_elem.text:
                    return t_elem.text[0]
        nxt = nxt.getnext()
    return ""


def _text_in_prev_run(run_elem) -> str:
    """Get the last text character from the previous run in the paragraph."""
    prev = run_elem.getprevious()
    while prev is not None:
        if prev.tag == _W_R:
            t_elems = list(prev.iter(_W_T))
            for t_elem in reversed(t_elems):
                if t_elem.text:
                    return t_elem.text[-1]
        prev = prev.getprevious()
    return ""


def cleanup_soft_linebreaks(
    doc: Document,
    *,
    process_tables: bool = True,
    skip_toc: bool = True,
    skip_centered: bool = True,
    skip_short_frames: bool = True,
) -> dict[str, int]:
    """Replace mid-sentence soft line breaks with spaces.

    Operates at the OXML level: removes <w:br> elements and inserts <w:t> spaces
    where needed, preserving run formatting.

    Args:
        doc: python-docx Document object (modified in place).
        process_tables: also process paragraphs inside table cells.
        skip_toc: skip paragraphs with TOC-related styles.
        skip_centered: skip center/right-aligned paragraphs.
        skip_short_frames: skip short-text framed paragraphs (titles/headers
            where breaks serve as visual line separators).

    Returns:
        Statistics dictionary.
    """
    stats = {
        "total_breaks_found": 0,
        "breaks_replaced_with_space": 0,
        "breaks_removed": 0,
        "breaks_skipped_page_column": 0,
        "breaks_skipped_clear": 0,
        "paragraphs_skipped_toc": 0,
        "paragraphs_skipped_aligned": 0,
        "paragraphs_skipped_short_frames": 0,
    }

    def _process_paragraphs(paragraphs):
        for para in paragraphs:
            if _should_skip_paragraph(
                para,
                skip_toc=skip_toc,
                skip_centered=skip_centered,
                skip_short_frames=skip_short_frames,
            ):
                style = _get_paragraph_style_name(para)
                if skip_toc and "toc" in style:
                    stats["paragraphs_skipped_toc"] += 1
                elif skip_short_frames and _is_short_frame_paragraph(para):
                    stats["paragraphs_skipped_short_frames"] += 1
                else:
                    stats["paragraphs_skipped_aligned"] += 1
                continue

            # Collect all <w:br> elements in this paragraph
            p_elem = para._p
            br_elements = list(p_elem.iter(_W_BR))
            if not br_elements:
                continue

            for br_elem in br_elements:
                stats["total_breaks_found"] += 1

                # Skip page/column breaks
                br_type = br_elem.get(_W_TYPE, "")
                if br_type in ("page", "column"):
                    stats["breaks_skipped_page_column"] += 1
                    continue

                # Skip breaks with w:clear (text wrapping around objects)
                if br_elem.get(_W_CLEAR) is not None:
                    stats["breaks_skipped_clear"] += 1
                    continue

                # This is a soft/textWrapping break — process it
                run_elem = br_elem.getparent()
                if run_elem is None or run_elem.tag != _W_R:
                    continue

                # Determine characters before and after
                char_before = _text_before_element(br_elem) or _text_in_prev_run(run_elem)
                char_after = _text_after_element(br_elem) or _text_in_next_run(run_elem)

                # Decide: replace with space or just remove
                needs_space = (
                    char_before
                    and char_after
                    and not char_before[-1:].isspace()
                    and not char_after[0:1].isspace()
                )

                if needs_space:
                    # Replace <w:br> with <w:t> containing a space
                    t_elem = run_elem.makeelement(_W_T, {_XML_SPACE: "preserve"})
                    t_elem.text = " "
                    br_idx = list(run_elem).index(br_elem)
                    run_elem.remove(br_elem)
                    run_elem.insert(br_idx, t_elem)
                    stats["breaks_replaced_with_space"] += 1
                else:
                    # Just remove the break (space already exists nearby)
                    run_elem.remove(br_elem)
                    stats["breaks_removed"] += 1

    # Process body paragraphs
    _process_paragraphs(doc.paragraphs)

    # Process table paragraphs
    if process_tables:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _process_paragraphs(cell.paragraphs)

    return stats


# ---------------------------------------------------------------------------
# Combined entry point
# ---------------------------------------------------------------------------


def strip_blanks_and_breaks(
    doc: Document,
    *,
    # Pass 1 params
    strip_empty: bool = True,
    keep_singles: bool = True,
    max_consecutive: int = 1,
    # Pass 2 params
    cleanup_breaks: bool = True,
    skip_toc: bool = True,
    skip_centered: bool = True,
    skip_short_frames: bool = True,
    # Common
    process_tables: bool = True,
) -> dict[str, int]:
    """Run both cleanup passes on a DOCX document.

    Pass 1: strip excessive empty paragraphs.
    Pass 2: replace mid-sentence soft line breaks with spaces.

    Returns combined statistics dictionary.
    """
    result: dict[str, int] = {}

    if strip_empty:
        p1 = strip_blank_paragraphs(
            doc,
            keep_singles=keep_singles,
            max_consecutive=max_consecutive,
            process_tables=process_tables,
        )
        result.update(p1)

    if cleanup_breaks:
        p2 = cleanup_soft_linebreaks(
            doc,
            process_tables=process_tables,
            skip_toc=skip_toc,
            skip_centered=skip_centered,
            skip_short_frames=skip_short_frames,
        )
        result.update(p2)

    return result

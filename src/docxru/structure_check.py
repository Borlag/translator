from __future__ import annotations

import json
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class TextMismatch:
    location: str
    source_snippet: str
    target_snippet: str


def iter_all_paragraphs(doc: Document) -> Iterator[tuple[str, Paragraph]]:
    """Yield (location, paragraph) for body and nested tables."""
    for i, p in enumerate(doc.paragraphs):
        yield f"body/p{i}", p

    def walk_table(tbl: Table, base: str) -> Iterator[tuple[str, Paragraph]]:
        for r_i, row in enumerate(tbl.rows):
            for c_i, cell in enumerate(row.cells):
                for p_i, p in enumerate(cell.paragraphs):
                    yield f"{base}/r{r_i}/c{c_i}/p{p_i}", p
                for t_i, t in enumerate(cell.tables):
                    yield from walk_table(t, f"{base}/r{r_i}/c{c_i}/t{t_i}")

    for t_i, t in enumerate(doc.tables):
        yield from walk_table(t, f"body/t{t_i}")


def doc_basic_stats(doc: Document) -> dict[str, int]:
    """Return basic doc-level counts."""
    # Count all paragraphs including nested tables
    total_paragraphs = sum(1 for _ in iter_all_paragraphs(doc))
    return {
        "body_paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "all_paragraphs_including_tables": total_paragraphs,
    }


def count_word_xml_patterns(docx_path: Path, patterns: dict[str, str]) -> dict[str, int]:
    """Count regex patterns across all word/*.xml parts in a DOCX."""
    counts = {k: 0 for k in patterns}
    with zipfile.ZipFile(docx_path) as z:
        xml_names = [n for n in z.namelist() if n.startswith("word/") and n.endswith(".xml")]
        for name in xml_names:
            try:
                xml = z.read(name).decode("utf-8", errors="ignore")
            except Exception:
                continue
            for key, pat in patterns.items():
                counts[key] += len(re.findall(pat, xml))
    return counts


def compare_docx_structure(
    input_path: Path,
    output_path: Path,
    *,
    check_text_equality: bool = False,
    max_text_mismatches: int = 20,
) -> dict:
    """Compare structural invariants between two DOCX files.

    Intended usage:
      - check_text_equality=False for real translations (text will differ)
      - check_text_equality=True for mock/dry runs (must be identical)
    """

    inp_doc = Document(str(input_path))
    out_doc = Document(str(output_path))

    basic_in = doc_basic_stats(inp_doc)
    basic_out = doc_basic_stats(out_doc)

    xml_patterns = {
        "drawing": r"<w:drawing",
        "pict": r"<w:pict",
        "br_any": r"<w:br",
        "br_column": r"<w:br[^>]*w:type=\"column\"",
        "hyperlink": r"<w:hyperlink",
        "fldChar": r"<w:fldChar",
        "instrText": r"<w:instrText",
        "sectPr": r"<w:sectPr",
    }
    xml_in = count_word_xml_patterns(input_path, xml_patterns)
    xml_out = count_word_xml_patterns(output_path, xml_patterns)

    # Marker leakage check: any remaining ⟦...⟧ tokens indicate broken write-back/unshield.
    # We keep this lightweight and scan visible paragraph text.
    marker_hits = 0
    for _, p in iter_all_paragraphs(out_doc):
        if "⟦" in (p.text or "") or "⟧" in (p.text or ""):
            marker_hits += 1
            if marker_hits >= max_text_mismatches:
                break

    text_mismatches: list[TextMismatch] = []
    if check_text_equality:
        for (loc_in, p_in), (loc_out, p_out) in zip(iter_all_paragraphs(inp_doc), iter_all_paragraphs(out_doc)):
            if loc_in != loc_out:
                text_mismatches.append(
                    TextMismatch(location=f"{loc_in} != {loc_out}", source_snippet="", target_snippet="")
                )
                break
            if (p_in.text or "") != (p_out.text or ""):
                text_mismatches.append(
                    TextMismatch(
                        location=loc_in,
                        source_snippet=(p_in.text or "")[:300],
                        target_snippet=(p_out.text or "")[:300],
                    )
                )
                if len(text_mismatches) >= max_text_mismatches:
                    break

    return {
        "input_path": str(input_path),
        "output_path": str(output_path),
        "basic": {"input": basic_in, "output": basic_out},
        "basic_diff": {k: basic_out.get(k, 0) - basic_in.get(k, 0) for k in basic_in},
        "word_xml_counts": {"input": xml_in, "output": xml_out},
        "word_xml_diff": {k: xml_out.get(k, 0) - xml_in.get(k, 0) for k in xml_in},
        "marker_leak_paragraphs": marker_hits,
        "text_mismatches": [m.__dict__ for m in text_mismatches],
    }


def write_structure_report(report: dict, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

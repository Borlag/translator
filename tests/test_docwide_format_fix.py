"""Tests for document-wide formatting fixes in oxml_table_fix.py."""

from __future__ import annotations

from lxml import etree
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docxru.oxml_table_fix import (
    apply_document_wide_formatting_fixes,
    reduce_font_sizes_in_frames,
    reduce_font_sizes_in_textboxes,
    relax_frame_exact_heights_aggressive,
)

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
NS_WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"


def _make_elem(ns: str, local: str, **attribs) -> etree._Element:
    """Make a namespaced element using raw lxml (bypasses python-docx ns registry)."""
    elem = etree.SubElement(etree.Element("_dummy"), f"{{{ns}}}{local}")
    elem.getparent().remove(elem)
    for k, v in attribs.items():
        elem.set(k, v)
    return elem


def _make_doc_with_textbox(font_pt: float | None = None, text: str = "Hello") -> Document:
    """Create a minimal DOCX with a DrawingML textbox containing text."""
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run()

    # Build DrawingML textbox using raw lxml elements
    drawing = OxmlElement("w:drawing")
    inline = _make_elem(NS_WP, "inline")
    extent = _make_elem(NS_WP, "extent")
    extent.set("cx", "1000000")
    extent.set("cy", "500000")
    inline.append(extent)

    graphic = _make_elem(NS_A, "graphic")
    graphic_data = _make_elem(NS_A, "graphicData")
    wsp = _make_elem(NS_WPS, "wsp")
    body_pr = _make_elem(NS_A, "bodyPr")
    wsp.append(body_pr)

    txbx = _make_elem(NS_WPS, "txbx")
    txbx_content = OxmlElement("w:txbxContent")
    tb_p = OxmlElement("w:p")
    tb_r = OxmlElement("w:r")
    if font_pt is not None:
        tb_rPr = OxmlElement("w:rPr")
        tb_sz = OxmlElement("w:sz")
        tb_sz.set(qn("w:val"), str(int(font_pt * 2)))
        tb_rPr.append(tb_sz)
        tb_r.append(tb_rPr)
    tb_t = OxmlElement("w:t")
    tb_t.text = text
    tb_r.append(tb_t)
    tb_p.append(tb_r)
    txbx_content.append(tb_p)
    txbx.append(txbx_content)
    wsp.append(txbx)
    graphic_data.append(wsp)
    graphic.append(graphic_data)
    inline.append(graphic)
    drawing.append(inline)
    r._r.append(drawing)

    return doc


def _make_doc_with_frame(font_pt: float, h_rule: str = "exact") -> Document:
    """Create a minimal DOCX with a framed paragraph."""
    doc = Document()
    p = doc.add_paragraph("Frame text")
    for run in p.runs:
        run.font.size = Pt(font_pt)

    p_pr = p._p.get_or_add_pPr()
    frame_pr = OxmlElement("w:framePr")
    frame_pr.set(qn("w:w"), "5000")
    frame_pr.set(qn("w:h"), "2000")
    frame_pr.set(qn("w:hRule"), h_rule)
    p_pr.append(frame_pr)

    return doc


class TestReduceFontSizesInTextboxes:
    def test_caps_large_font(self):
        doc = _make_doc_with_textbox(font_pt=12.0)
        stats = reduce_font_sizes_in_textboxes(doc, max_font_pt=9.0)
        assert stats["textbox_fonts_capped"] == 1

    def test_skips_small_font(self):
        doc = _make_doc_with_textbox(font_pt=8.0)
        stats = reduce_font_sizes_in_textboxes(doc, max_font_pt=9.0)
        assert stats["textbox_fonts_capped"] == 0

    def test_sets_inherited_font(self):
        doc = _make_doc_with_textbox(font_pt=None)
        stats = reduce_font_sizes_in_textboxes(
            doc, max_font_pt=9.0, set_inherited_font_pt=8.5
        )
        assert stats["textbox_fonts_set_inherited"] == 1

    def test_no_inherited_when_disabled(self):
        doc = _make_doc_with_textbox(font_pt=None)
        stats = reduce_font_sizes_in_textboxes(
            doc, max_font_pt=9.0, set_inherited_font_pt=None
        )
        assert stats["textbox_fonts_set_inherited"] == 0


class TestReduceFontSizesInFrames:
    def test_caps_large_font(self):
        doc = _make_doc_with_frame(font_pt=12.0)
        stats = reduce_font_sizes_in_frames(doc, max_font_pt=9.0)
        assert stats["frame_fonts_capped"] >= 1

    def test_skips_small_font(self):
        doc = _make_doc_with_frame(font_pt=8.0)
        stats = reduce_font_sizes_in_frames(doc, max_font_pt=9.0)
        assert stats["frame_fonts_capped"] == 0


class TestRelaxFrameExactHeights:
    def test_relaxes_exact_to_atleast(self):
        doc = _make_doc_with_frame(font_pt=10.0, h_rule="exact")
        stats = relax_frame_exact_heights_aggressive(doc)
        assert stats["frame_exact_relaxed_all"] == 1

    def test_skips_non_exact(self):
        doc = _make_doc_with_frame(font_pt=10.0, h_rule="atLeast")
        stats = relax_frame_exact_heights_aggressive(doc)
        assert stats["frame_exact_relaxed_all"] == 0


class TestApplyDocumentWideFormattingFixes:
    def test_combined_pass(self):
        doc = _make_doc_with_textbox(font_pt=12.0)
        stats = apply_document_wide_formatting_fixes(
            doc,
            max_textbox_font_pt=9.0,
            set_inherited_textbox_font_pt=8.5,
        )
        assert stats["textbox_fonts_capped"] >= 1

    def test_no_changes_when_all_ok(self):
        doc = _make_doc_with_textbox(font_pt=8.0)
        stats = apply_document_wide_formatting_fixes(
            doc,
            max_textbox_font_pt=9.0,
            set_inherited_textbox_font_pt=None,
        )
        total = sum(stats.values())
        assert total == 0

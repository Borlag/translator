from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement

from docxru.docx_reader import collect_segments


def _append_textbox_paragraph(container_paragraph, text: str) -> None:
    run = container_paragraph.add_run("")
    txbx_content = OxmlElement("w:txbxContent")
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    txbx_content.append(p)
    run._r.append(txbx_content)


def test_collect_segments_includes_body_textbox_paragraphs():
    doc = Document()
    doc.add_paragraph("Regular body line")
    host = doc.add_paragraph("Host")
    _append_textbox_paragraph(host, "Text from body textbox")

    segments = collect_segments(doc, include_headers=False, include_footers=False)
    texts = [seg.source_plain for seg in segments]

    assert "Regular body line" in texts
    assert "Host" in texts
    assert "Text from body textbox" in texts
    textbox_seg = next(seg for seg in segments if seg.source_plain == "Text from body textbox")
    assert textbox_seg.context.get("in_textbox") is True
    assert textbox_seg.location.startswith("body/textbox")


def test_collect_segments_textbox_in_header_respects_include_headers_flag():
    doc = Document()
    header = doc.sections[0].header
    host = header.add_paragraph("Header host")
    _append_textbox_paragraph(host, "Header textbox text")

    no_header_segments = collect_segments(doc, include_headers=False, include_footers=False)
    assert all(seg.source_plain != "Header textbox text" for seg in no_header_segments)

    header_segments = collect_segments(doc, include_headers=True, include_footers=False)
    header_textbox_seg = next(seg for seg in header_segments if seg.source_plain == "Header textbox text")
    assert header_textbox_seg.context.get("part") == "header"
    assert header_textbox_seg.context.get("in_textbox") is True
    assert header_textbox_seg.location.startswith("header0/textbox")


def test_collect_segments_marks_toc_like_entries():
    doc = Document()
    doc.add_paragraph("TABLE OF CONTENTS")
    doc.add_paragraph("Repair No. 1-1 Lower Bearing\tRepair No.\t1-1\t601")

    segments = collect_segments(doc, include_headers=False, include_footers=False)

    title = next(seg for seg in segments if seg.source_plain == "TABLE OF CONTENTS")
    entry = next(seg for seg in segments if "Repair No. 1-1 Lower Bearing" in seg.source_plain)
    assert title.context.get("is_toc_entry") is True
    assert entry.context.get("is_toc_entry") is True

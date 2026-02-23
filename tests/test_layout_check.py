from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement

from docxru.config import PipelineConfig
from docxru.layout_check import (
    check_table_cell_overflow,
    check_text_expansion,
    check_textbox_overflow,
    validate_layout,
)
from docxru.models import Segment
from docxru.oxml_table_fix import set_cell_width_twips


def _make_segment(
    *,
    seg_id: str,
    source: str,
    target: str,
    paragraph_ref,
    in_table: bool = False,
    in_textbox: bool = False,
) -> Segment:
    context = {"part": "body"}
    if in_table:
        context["in_table"] = True
    if in_textbox:
        context["in_textbox"] = True
    return Segment(
        segment_id=seg_id,
        location=f"body/p{seg_id}",
        context=context,
        source_plain=source,
        paragraph_ref=paragraph_ref,
        target_tagged=target,
        target_shielded_tagged=target,
    )


def _attach_extent(paragraph, *, width_twips: int, height_twips: int) -> None:
    run = paragraph.add_run("")
    extent = OxmlElement("wp:extent")
    extent.set("cx", str(int(width_twips) * 635))
    extent.set("cy", str(int(height_twips) * 635))
    run._r.append(extent)


def test_check_text_expansion_warns_on_high_ratio():
    seg = _make_segment(
        seg_id="1",
        source="Install bolt",
        target="Установите крепежный болт с контролем момента затяжки",
        paragraph_ref=None,
    )

    issues = check_text_expansion([seg], warn_ratio=1.4)
    assert len(issues) == 1
    assert issues[0].code == "length_ratio_high"
    assert issues[0].details["segment_id"] == "1"


def test_check_table_cell_overflow_uses_cell_width():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_width_twips(cell, 600)  # very narrow
    paragraph = cell.paragraphs[0]
    paragraph.text = "Main fitting"

    seg = _make_segment(
        seg_id="2",
        source="Main fitting",
        target="Основной фитинг в сборе со вспомогательными элементами",
        paragraph_ref=paragraph,
        in_table=True,
    )

    issues = check_table_cell_overflow(doc, [seg])
    assert len(issues) == 1
    assert issues[0].code == "layout_table_overflow_risk"
    assert issues[0].details["segment_id"] == "2"


def test_check_textbox_overflow_uses_extent_dimensions():
    doc = Document()
    paragraph = doc.add_paragraph("Textbox host")
    _attach_extent(paragraph, width_twips=960, height_twips=300)

    seg = _make_segment(
        seg_id="2a",
        source="Bearing",
        target="Very long translated text that should overflow the textbox area",
        paragraph_ref=paragraph,
        in_textbox=True,
    )

    issues = check_textbox_overflow(doc, [seg])
    assert len(issues) == 1
    details = issues[0].details
    assert details["width_twips"] == 960
    assert details["height_twips"] == 300
    assert details["approx_capacity_chars"] < len(seg.target_tagged or "")


def test_check_textbox_overflow_fallback_by_ratio():
    seg = _make_segment(
        seg_id="3",
        source="Bearing",
        target="Подшипник в составе узла со значительным увеличением текста",
        paragraph_ref=None,
        in_textbox=True,
    )

    issues = check_textbox_overflow(None, [seg])
    assert len(issues) == 1
    assert issues[0].code == "layout_textbox_overflow_risk"


def test_validate_layout_combines_checks():
    doc = Document()
    paragraph = doc.add_paragraph("Bolt")
    seg = _make_segment(
        seg_id="4",
        source="Bolt",
        target="Крепежный элемент болтового типа",
        paragraph_ref=paragraph,
    )
    cfg = PipelineConfig(layout_expansion_warn_ratio=1.1)
    issues = validate_layout(doc, [seg], cfg)
    assert any(issue.code == "length_ratio_high" for issue in issues)

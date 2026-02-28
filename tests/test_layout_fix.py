from __future__ import annotations

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docxru.config import PipelineConfig
from docxru.layout_fix import apply_global_font_shrink, fix_expansion_issues
from docxru.models import Issue, Segment, Severity


def _issue(seg_id: str, code: str = "length_ratio_high") -> Issue:
    return Issue(
        code=code,
        severity=Severity.WARN,
        message="overflow risk",
        details={"segment_id": seg_id},
    )


def test_fix_expansion_issues_reduces_frame_font_and_emits_issue():
    doc = Document()
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    frame_pr = OxmlElement("w:framePr")
    frame_pr.set(qn("w:hRule"), "exact")
    p_pr.append(frame_pr)
    run = paragraph.add_run("Long translated text")
    run.font.size = Pt(12)

    seg = Segment(
        segment_id="s1",
        location="body/p1",
        context={"part": "body", "in_frame": True},
        source_plain="text",
        paragraph_ref=paragraph,
        target_tagged="Long translated text",
    )
    cfg = PipelineConfig(layout_auto_fix=True, layout_font_reduction_pt=0.5)

    applied = fix_expansion_issues([seg], [_issue("s1", code="layout_frame_overflow_risk")], cfg)
    assert applied == 1
    assert run.font.size is not None
    assert run.font.size.pt == pytest.approx(11.5)
    frame_after = paragraph._p.get_or_add_pPr().find(qn("w:framePr"))
    assert frame_after is not None
    assert frame_after.get(qn("w:hRule")) == "atLeast"
    assert any(
        issue.code == "layout_auto_fix_applied" and issue.details.get("strategy") == "frame" for issue in seg.issues
    )


def test_fix_expansion_issues_can_detach_high_overflow_frame():
    doc = Document()
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    frame_pr = OxmlElement("w:framePr")
    frame_pr.set(qn("w:hRule"), "exact")
    p_pr.append(frame_pr)
    paragraph.add_run("Long translated text")

    seg = Segment(
        segment_id="s1-frame-detach",
        location="body/p1",
        context={"part": "body", "in_frame": True, "in_table": True},
        source_plain="text",
        paragraph_ref=paragraph,
        target_tagged="Long translated text",
    )
    issue = _issue("s1-frame-detach", code="layout_frame_overflow_risk")
    issue.details.update({"approx_capacity_chars": 40, "target_len": 80})
    cfg = PipelineConfig(layout_auto_fix=True, layout_font_reduction_pt=0.5)

    applied = fix_expansion_issues([seg], [issue], cfg)
    assert applied == 1
    assert paragraph._p.get_or_add_pPr().find(qn("w:framePr")) is None


def test_fix_expansion_issues_skips_generic_length_ratio_for_body_paragraphs():
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Long translated text")
    run.font.size = Pt(12)

    seg = Segment(
        segment_id="s1-generic",
        location="body/p1",
        context={"part": "body"},
        source_plain="text",
        paragraph_ref=paragraph,
        target_tagged="Long translated text",
    )
    cfg = PipelineConfig(layout_auto_fix=True, layout_font_reduction_pt=0.5)

    applied = fix_expansion_issues([seg], [_issue("s1-generic", code="length_ratio_high")], cfg)

    assert applied == 0
    assert run.font.size is not None
    assert run.font.size.pt == pytest.approx(12.0)
    assert not any(issue.code == "layout_auto_fix_applied" for issue in seg.issues)


def test_fix_expansion_issues_reduces_non_table_spacing():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("Long translated text")
    run.font.size = Pt(11)

    seg = Segment(
        segment_id="s1a",
        location="body/p2",
        context={"part": "body", "in_textbox": True},
        source_plain="text",
        paragraph_ref=paragraph,
        target_tagged="Long translated text",
    )
    cfg = PipelineConfig(layout_auto_fix=True, layout_spacing_factor=0.5, layout_font_reduction_pt=0.5)

    applied = fix_expansion_issues([seg], [_issue("s1a", code="layout_textbox_overflow_risk")], cfg)
    assert applied == 1
    assert paragraph.paragraph_format.space_before is not None
    assert paragraph.paragraph_format.space_before.pt < 14
    assert paragraph.paragraph_format.space_after is not None
    assert paragraph.paragraph_format.space_after.pt < 8
    assert run.font.size is not None
    assert run.font.size.pt == pytest.approx(10.0)
    assert any(
        issue.code == "layout_auto_fix_applied" and issue.details.get("strategy") == "textbox"
        for issue in seg.issues
    )


def test_fix_expansion_issues_reduces_table_spacing():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:hRule"), "exact")
    tr_height.set(qn("w:val"), "240")
    tr_pr.append(tr_height)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.text = "Translated"
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(10)

    seg = Segment(
        segment_id="s2",
        location="body/t0/r0/c0/p0",
        context={"part": "body", "in_table": True},
        source_plain="text",
        paragraph_ref=paragraph,
        target_tagged="Translated",
    )
    cfg = PipelineConfig(layout_auto_fix=True, layout_spacing_factor=0.5, layout_font_reduction_pt=0.5)

    applied = fix_expansion_issues([seg], [_issue("s2", code="layout_table_overflow_risk")], cfg)
    assert applied == 1
    before = paragraph.paragraph_format.space_before
    after = paragraph.paragraph_format.space_after
    assert before is not None and before.pt < 12
    assert after is not None and after.pt < 10
    assert paragraph.paragraph_format.line_spacing == pytest.approx(1.0)
    spacing = paragraph.runs[0]._r.get_or_add_rPr().find(qn("w:spacing"))
    assert spacing is not None
    assert spacing.get(qn("w:val")) == "-15"
    assert table.rows[0]._tr.get_or_add_trPr().find(qn("w:trHeight")) is None
    assert any(
        issue.code == "layout_auto_fix_applied" and issue.details.get("strategy") == "table"
        for issue in seg.issues
    )


def test_fix_expansion_issues_table_shrink_scales_by_overflow_ratio():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    run = paragraph.add_run("Translated")
    run.font.size = Pt(12)

    seg = Segment(
        segment_id="s2b",
        location="body/t0/r0/c0/p0",
        context={"part": "body", "in_table": True},
        source_plain="source",
        paragraph_ref=paragraph,
        target_tagged="translated",
    )
    issue = _issue("s2b", code="layout_table_overflow_risk")
    issue.details.update({"approx_capacity_chars": 40, "target_len": 160})
    cfg = PipelineConfig(layout_auto_fix=True, layout_font_reduction_pt=0.5)

    applied = fix_expansion_issues([seg], [issue], cfg)

    assert applied == 1
    assert run.font.size is not None
    # ratio=4.0 should trigger near-max reduction (~2pt)
    assert run.font.size.pt == pytest.approx(10.0)


def test_apply_global_font_shrink_uses_different_body_and_table_steps():
    doc = Document()
    body_p = doc.add_paragraph()
    body_run = body_p.add_run("Body translated")
    body_run.font.size = Pt(12)

    table = doc.add_table(rows=1, cols=1)
    table_p = table.cell(0, 0).paragraphs[0]
    table_run = table_p.add_run("Table translated")
    table_run.font.size = Pt(11)

    body_seg = Segment(
        segment_id="b1",
        location="body/p1",
        context={"part": "body"},
        source_plain="Body source",
        paragraph_ref=body_p,
        target_tagged="Body translated",
    )
    table_seg = Segment(
        segment_id="t1",
        location="body/t0/r0/c0/p0",
        context={"part": "body", "in_table": True},
        source_plain="Table source",
        paragraph_ref=table_p,
        target_tagged="Table translated",
    )
    cfg = PipelineConfig(font_shrink_body_pt=2.0, font_shrink_table_pt=3.0)

    changed = apply_global_font_shrink([body_seg, table_seg], cfg)

    assert changed == 2
    assert body_run.font.size is not None and body_run.font.size.pt == pytest.approx(10.0)
    assert table_run.font.size is not None and table_run.font.size.pt == pytest.approx(8.0)
    assert any(issue.code == "global_font_shrink_applied" for issue in body_seg.issues)
    assert any(issue.code == "global_font_shrink_applied" for issue in table_seg.issues)


def test_apply_global_font_shrink_skips_non_translated_segments():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Unchanged")
    run.font.size = Pt(12)

    seg = Segment(
        segment_id="s-no-target",
        location="body/p1",
        context={"part": "body"},
        source_plain="Source",
        paragraph_ref=p,
        target_tagged=None,
    )
    cfg = PipelineConfig(font_shrink_body_pt=2.0, font_shrink_table_pt=3.0)

    changed = apply_global_font_shrink([seg], cfg)

    assert changed == 0
    assert run.font.size is not None and run.font.size.pt == pytest.approx(12.0)

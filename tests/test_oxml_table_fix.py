from __future__ import annotations

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docxru.oxml_table_fix import (
    normalize_abbyy_oxml,
    normalize_table_cell_margins,
    normalize_textbox_insets,
    set_textbox_autofit,
)


def _append_exact_tr_height(doc: Document):
    table = doc.add_table(rows=1, cols=1)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), "240")
    tr_height.set(qn("w:hRule"), "exact")
    tr_pr.append(tr_height)
    return tr_pr


def _append_frame_pr(doc: Document):
    p = doc.add_paragraph("Sample")
    p_pr = p._p.get_or_add_pPr()
    frame_pr = OxmlElement("w:framePr")
    frame_pr.set(qn("w:w"), "100")
    frame_pr.set(qn("w:h"), "240")
    frame_pr.set(qn("w:hRule"), "exact")
    p_pr.append(frame_pr)
    return p_pr


def _append_framed_paragraph_with_spacing(
    doc: Document,
    *,
    y_twips: int,
    x_twips: int = 0,
    text: str = "Framed paragraph",
):
    p = doc.add_paragraph(text)
    p_pr = p._p.get_or_add_pPr()
    frame_pr = OxmlElement("w:framePr")
    frame_pr.set(qn("w:w"), "100")
    frame_pr.set(qn("w:h"), "240")
    frame_pr.set(qn("w:hRule"), "exact")
    frame_pr.set(qn("w:x"), str(int(x_twips)))
    frame_pr.set(qn("w:y"), str(int(y_twips)))
    frame_pr.set(qn("w:vAnchor"), "page")
    frame_pr.set(qn("w:hAnchor"), "page")
    p_pr.append(frame_pr)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "exact")
    p_pr.append(spacing)
    return frame_pr, spacing


def _append_exact_line_spacing(doc: Document):
    p = doc.add_paragraph("Line spacing sample")
    p_pr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "exact")
    p_pr.append(spacing)
    return spacing


def _append_textbox(doc: Document, *, text: str) -> tuple[object, object]:
    host = doc.add_paragraph("Host")
    run = host.add_run("")

    shape = OxmlElement("w:shape")
    body_pr = OxmlElement("a:bodyPr")
    body_pr.append(OxmlElement("a:noAutofit"))

    txbx_content = OxmlElement("w:txbxContent")
    if text:
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        p.append(r)
        txbx_content.append(p)

    shape.append(body_pr)
    shape.append(txbx_content)
    run._r.append(shape)
    return body_pr, txbx_content


def _has_child(node, local_name: str) -> bool:
    for child in list(node):
        tag = str(getattr(child, "tag", ""))
        if tag.endswith("}" + local_name) or tag.endswith(":" + local_name) or tag == local_name:
            return True
    return False


def test_set_textbox_autofit_replaces_noautofit_for_non_empty_textbox():
    doc = Document()
    body_pr, _ = _append_textbox(doc, text="Textbox text")

    changed = set_textbox_autofit(doc)

    assert changed == 1
    assert _has_child(body_pr, "noAutofit") is False
    assert _has_child(body_pr, "normAutofit") is True


def test_set_textbox_autofit_skips_empty_textbox():
    doc = Document()
    body_pr, _ = _append_textbox(doc, text="")

    changed = set_textbox_autofit(doc)

    assert changed == 0
    assert _has_child(body_pr, "noAutofit") is True
    assert _has_child(body_pr, "normAutofit") is False


def test_normalize_abbyy_oxml_safe_removes_exact_trheight_only():
    doc = Document()
    tr_pr = _append_exact_tr_height(doc)
    p_pr = _append_frame_pr(doc)
    spacing = _append_exact_line_spacing(doc)

    stats = normalize_abbyy_oxml(doc, profile="safe")

    assert stats["tr_height_exact_removed"] == 1
    assert stats["frame_pr_removed"] == 0
    assert stats["frame_pr_exact_relaxed"] == 0
    assert stats["line_spacing_exact_relaxed"] == 0
    assert stats["textbox_autofit_updated"] == 0
    assert stats["table_cell_margins_normalized"] == 0
    assert tr_pr.find(qn("w:trHeight")) is None
    assert p_pr.find(qn("w:framePr")) is not None
    assert spacing.get(qn("w:lineRule")) == "exact"


def test_normalize_abbyy_oxml_aggressive_relaxes_framepr_height_rule():
    doc = Document()
    tr_pr = _append_exact_tr_height(doc)
    p_pr = _append_frame_pr(doc)
    spacing = _append_exact_line_spacing(doc)

    stats = normalize_abbyy_oxml(doc, profile="aggressive")

    assert stats["tr_height_exact_removed"] == 1
    assert stats["frame_pr_removed"] == 0
    assert stats["frame_pr_exact_relaxed"] == 1
    assert stats["line_spacing_exact_relaxed"] == 1
    assert stats["textbox_autofit_updated"] == 0
    assert stats["table_cell_margins_normalized"] == 0
    assert tr_pr.find(qn("w:trHeight")) is None
    frame_pr = p_pr.find(qn("w:framePr"))
    assert frame_pr is not None
    assert frame_pr.get(qn("w:hRule")) == "auto"
    assert spacing.get(qn("w:lineRule")) == "atLeast"


def test_normalize_abbyy_oxml_aggressive_keeps_header_frame_vertical_anchor():
    doc = Document()
    top_frame, top_spacing = _append_framed_paragraph_with_spacing(
        doc,
        y_twips=1717,
        text="РУКОВОДСТВО ПО ТЕХНИЧЕСКОМУ ОБСЛУЖИВАНИЮ КОМПОНЕНТОВ ДЛЯ ДЕТАЛЕЙ № 201587001 И 201587002",
    )
    body_frame, body_spacing = _append_framed_paragraph_with_spacing(doc, y_twips=5000, text="Body frame")

    stats = normalize_abbyy_oxml(doc, profile="aggressive")

    assert stats["frame_pr_exact_relaxed"] == 0
    assert stats["line_spacing_exact_relaxed"] == 0
    assert stats["frame_auto_height_set"] >= 2
    assert stats["frame_vertical_anchor_margin_set"] >= 1
    assert stats["frame_vertical_text_distance_set"] >= 1
    assert stats["frame_header_blocks_detected"] >= 1
    assert stats["frame_header_y_nudged"] >= 1
    assert top_frame.get(qn("w:hRule")) == "auto"
    assert int(top_frame.get(qn("w:y"))) >= 2100
    assert top_spacing.get(qn("w:lineRule")) == "exact"
    assert top_frame.get(qn("w:vAnchor")) == "page"
    assert top_frame.get(qn("w:vSpace")) is None
    assert body_frame.get(qn("w:hRule")) == "auto"
    assert body_spacing.get(qn("w:lineRule")) == "exact"
    assert body_frame.get(qn("w:vAnchor")) == "margin"
    assert body_frame.get(qn("w:vSpace")) == "142"


def test_normalize_abbyy_oxml_full_applies_textbox_autofit():
    doc = Document()
    tr_pr = _append_exact_tr_height(doc)
    p_pr = _append_frame_pr(doc)
    spacing = _append_exact_line_spacing(doc)
    body_pr, _ = _append_textbox(doc, text="Overflowing textbox")

    stats = normalize_abbyy_oxml(doc, profile="full")

    assert stats["tr_height_exact_removed"] == 1
    assert stats["frame_pr_removed"] == 0
    assert stats["frame_pr_exact_relaxed"] == 1
    assert stats["line_spacing_exact_relaxed"] == 1
    assert stats["textbox_autofit_updated"] == 1
    assert stats["textbox_insets_normalized"] == 0
    assert stats["frame_width_expanded"] == 0
    assert stats["frame_height_expanded"] == 0
    assert stats["frame_vertical_reflowed"] == 0
    assert stats["table_cell_margins_normalized"] == 0
    assert tr_pr.find(qn("w:trHeight")) is None
    frame_pr = p_pr.find(qn("w:framePr"))
    assert frame_pr is not None
    assert frame_pr.get(qn("w:hRule")) == "auto"
    assert spacing.get(qn("w:lineRule")) == "atLeast"
    assert _has_child(body_pr, "noAutofit") is False
    assert _has_child(body_pr, "normAutofit") is True


def test_normalize_textbox_insets_caps_excessive_bodypr_insets():
    doc = Document()
    body_pr, _ = _append_textbox(doc, text="Textbox text")
    body_pr.set("lIns", "91440")
    body_pr.set("rIns", "91440")
    body_pr.set("tIns", "91440")
    body_pr.set("bIns", "91440")

    changed = normalize_textbox_insets(doc, max_inset_emu=25400)

    assert changed == 1
    assert body_pr.get("lIns") == "25400"
    assert body_pr.get("bIns") == "25400"


def test_normalize_abbyy_oxml_full_expands_and_spaces_overlapping_page_frames():
    doc = Document()
    frame_a, _ = _append_framed_paragraph_with_spacing(doc, y_twips=4000, x_twips=1200)
    frame_b, _ = _append_framed_paragraph_with_spacing(doc, y_twips=4100, x_twips=1200)
    frame_a.set(qn("w:w"), "3000")
    frame_b.set(qn("w:w"), "3000")
    frame_a.set(qn("w:h"), "900")
    frame_b.set(qn("w:h"), "900")

    stats = normalize_abbyy_oxml(
        doc,
        profile="full",
        frame_expand_width_factor=1.1,
        frame_expand_height_factor=1.1,
        frame_vertical_gap_twips=283,
    )

    assert stats["frame_width_expanded"] >= 2
    assert stats["frame_height_expanded"] >= 2
    assert stats["frame_vertical_reflowed"] >= 1
    y_a = int(frame_a.get(qn("w:y")))
    h_a = int(frame_a.get(qn("w:h")))
    y_b = int(frame_b.get(qn("w:y")))
    assert y_b >= y_a + h_a + 283


def test_normalize_table_cell_margins_caps_excessive_values():
    doc = Document()
    cell = doc.add_table(rows=1, cols=1).cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    left = OxmlElement("w:left")
    left.set(qn("w:type"), "dxa")
    left.set(qn("w:w"), "480")
    tc_mar.append(left)
    tc_pr.append(tc_mar)

    changed = normalize_table_cell_margins(doc, max_margin_twips=108)

    assert changed == 1
    assert left.get(qn("w:w")) == "108"


def test_normalize_abbyy_oxml_rejects_invalid_profile():
    with pytest.raises(ValueError, match="Unsupported ABBYY profile"):
        normalize_abbyy_oxml(Document(), profile="unknown")

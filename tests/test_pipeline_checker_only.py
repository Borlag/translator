from __future__ import annotations

import logging

from docx import Document

from docxru.pipeline import _build_checker_only_docx_segments


def test_build_checker_only_docx_segments_aligns_source_and_target(tmp_path):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "translated.docx"

    source_doc = Document()
    source_doc.add_paragraph("Install actuator")
    source_doc.add_paragraph("12345")
    source_doc.save(str(source_path))

    output_doc = Document()
    output_doc.add_paragraph("Установите привод")
    output_doc.add_paragraph("12345")
    output_doc.save(str(output_path))

    _, segments, stats = _build_checker_only_docx_segments(
        input_path=source_path,
        output_path=output_path,
        include_headers=False,
        include_footers=False,
        logger=logging.getLogger("test_checker_only_alignment"),
    )

    assert len(segments) == 2
    assert segments[0].source_plain == "Install actuator"
    assert segments[0].context.get("checker_target_text") == "Установите привод"
    assert segments[1].source_plain == "12345"
    assert segments[1].context.get("checker_target_text") == ""
    assert stats["checker_candidates"] == 1
    assert stats["skipped_no_latin"] == 1
    assert stats["missing_source"] == 0
    assert stats["missing_target"] == 0

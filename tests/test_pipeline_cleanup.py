from __future__ import annotations

from docx import Document

from docxru.models import Segment
from docxru.pipeline import _apply_final_run_level_cleanup, _should_translate_segment_text


def test_final_cleanup_skips_untranslated_segments():
    doc = Document()
    paragraph = doc.add_paragraph("Table")
    seg = Segment(
        segment_id="s1",
        location="body/p1",
        context={"part": "body"},
        source_plain="Table",
        paragraph_ref=paragraph,
    )

    changed = _apply_final_run_level_cleanup([seg])

    assert changed == 0
    assert paragraph.text == "Table"


def test_final_cleanup_applies_to_translated_segments():
    doc = Document()
    paragraph = doc.add_paragraph("Table")
    seg = Segment(
        segment_id="s1",
        location="body/p1",
        context={"part": "body"},
        source_plain="Table",
        paragraph_ref=paragraph,
        target_shielded_tagged="table-ru",
        target_tagged="table-ru",
    )

    changed = _apply_final_run_level_cleanup([seg])

    assert changed == 0
    assert paragraph.text == "Table"


def test_final_cleanup_translates_table_references_with_numbers():
    doc = Document()
    paragraph = doc.add_paragraph("Table 5A")
    seg = Segment(
        segment_id="s1",
        location="body/p1",
        context={"part": "body"},
        source_plain="Table 5A",
        paragraph_ref=paragraph,
        target_shielded_tagged="table-ref-ru",
        target_tagged="table-ref-ru",
    )

    changed = _apply_final_run_level_cleanup([seg])

    assert changed == 1
    assert paragraph.text == "Таблица 5A"


def test_final_cleanup_keeps_table_of_contents_phrase():
    doc = Document()
    paragraph = doc.add_paragraph("Table of Contents")
    seg = Segment(
        segment_id="s1",
        location="body/p1",
        context={"part": "body"},
        source_plain="Table of Contents",
        paragraph_ref=paragraph,
        target_shielded_tagged="toc-ru",
        target_tagged="toc-ru",
    )

    changed = _apply_final_run_level_cleanup([seg])

    assert changed == 0
    assert paragraph.text == "Table of Contents"


def test_should_translate_segment_text_skips_russian_only():
    assert _should_translate_segment_text("Руководство по техническому обслуживанию компонентов") is False


def test_should_translate_segment_text_skips_russian_dominant_with_sparse_latin():
    text = "Руководство по техническому обслуживанию PN 201587001 ATA 32-12-22"
    assert _should_translate_segment_text(text) is False


def test_should_translate_segment_text_translates_english_text():
    text = "Install the charging valve and measure the nitrogen pressure."
    assert _should_translate_segment_text(text) is True

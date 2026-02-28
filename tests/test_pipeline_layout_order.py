from __future__ import annotations

import logging

from docx import Document
from docx.oxml import OxmlElement

import docxru.pipeline as pipeline
from docxru.config import PipelineConfig
from docxru.models import Issue, Segment, Severity


def test_apply_abbyy_and_layout_passes_runs_normalization_before_layout(monkeypatch):
    call_order: list[str] = []

    def _fake_normalize(doc, *, profile):  # noqa: ANN001, ANN202
        del doc
        call_order.append("normalize")
        assert profile == "full"
        return {
            "tr_height_exact_removed": 1,
            "frame_pr_removed": 1,
            "line_spacing_exact_relaxed": 1,
            "textbox_autofit_updated": 1,
        }

    def _fake_validate(doc, segments, cfg):  # noqa: ANN001, ANN202
        del doc, cfg
        call_order.append("validate")
        return [
            Issue(
                code="layout_textbox_overflow_risk",
                severity=Severity.WARN,
                message="overflow",
                details={"segment_id": segments[0].segment_id},
            )
        ]

    def _fake_attach(segments, issues):  # noqa: ANN001, ANN202
        del segments, issues
        call_order.append("attach")
        return 1

    def _fake_fix(segments, issues, cfg, *, pass_number=1):  # noqa: ANN001, ANN202
        del segments, issues, cfg, pass_number
        call_order.append("fix")
        return 1

    monkeypatch.setattr(pipeline, "normalize_abbyy_oxml", _fake_normalize)
    monkeypatch.setattr(pipeline, "validate_layout", _fake_validate)
    monkeypatch.setattr(pipeline, "_attach_issues_to_segments", _fake_attach)
    monkeypatch.setattr(pipeline, "fix_expansion_issues", _fake_fix)

    doc = Document()
    paragraph = doc.add_paragraph("Text")
    segments = [
        Segment(
            segment_id="s1",
            location="body/p1",
            context={"part": "body", "in_textbox": True},
            source_plain="Text",
            paragraph_ref=paragraph,
            target_tagged="Long translated text",
        )
    ]
    cfg = PipelineConfig(abbyy_profile="full", layout_check=True, layout_auto_fix=True)

    pipeline._apply_abbyy_and_layout_passes(
        doc,
        segments,
        cfg,
        logging.getLogger("test_pipeline_layout_order"),
    )

    assert call_order == ["normalize", "validate", "attach", "fix"]


def test_apply_abbyy_and_layout_passes_can_recheck_multiple_fix_passes(monkeypatch):
    validate_calls = {"count": 0}
    fix_calls = {"count": 0}

    def _fake_validate(doc, segments, cfg):  # noqa: ANN001, ANN202
        del doc, cfg
        validate_calls["count"] += 1
        if validate_calls["count"] == 1:
            return [
                Issue(
                    code="layout_table_overflow_risk",
                    severity=Severity.WARN,
                    message="overflow",
                    details={"segment_id": segments[0].segment_id},
                )
            ]
        return []

    def _fake_fix(segments, issues, cfg, *, pass_number=1):  # noqa: ANN001, ANN202
        del segments, issues, cfg, pass_number
        fix_calls["count"] += 1
        return 1

    monkeypatch.setattr(pipeline, "validate_layout", _fake_validate)
    monkeypatch.setattr(pipeline, "fix_expansion_issues", _fake_fix)

    doc = Document()
    paragraph = doc.add_paragraph("Text")
    segments = [
        Segment(
            segment_id="s2",
            location="body/p2",
            context={"part": "body", "in_table": True},
            source_plain="Text",
            paragraph_ref=paragraph,
            target_tagged="Long translated text",
        )
    ]
    cfg = PipelineConfig(layout_check=True, layout_auto_fix=True, layout_auto_fix_passes=2)

    pipeline._apply_abbyy_and_layout_passes(
        doc,
        segments,
        cfg,
        logging.getLogger("test_pipeline_layout_order"),
    )

    assert fix_calls["count"] == 1
    assert validate_calls["count"] == 2


def test_attach_container_constraints_sets_space_limit_for_tight_textboxes():
    doc = Document()
    paragraph = doc.add_paragraph("Textbox host")
    run = paragraph.add_run("")
    extent = OxmlElement("wp:extent")
    extent.set("cx", str(900 * 635))
    extent.set("cy", str(260 * 635))
    run._r.append(extent)

    seg = Segment(
        segment_id="s-space",
        location="body/textbox0/p0",
        context={"part": "body", "in_textbox": True},
        source_plain="Main fitting assembly text for limited shape",
        paragraph_ref=paragraph,
    )

    constrained = pipeline._attach_container_constraints([seg], PipelineConfig(layout_expansion_warn_ratio=1.2))
    assert constrained == 1
    assert int(seg.context.get("max_target_chars", 0)) > 0


def test_resolve_runtime_formatting_preset_auto_detects_abbyy():
    doc = Document()
    doc.core_properties.author = "ABBYY FineReader"
    cfg = PipelineConfig(formatting_preset="auto")

    resolved = pipeline._resolve_runtime_formatting_preset(cfg, doc, logging.getLogger("test_auto_preset"))
    assert resolved.formatting_preset == "abbyy_standard"
    assert resolved.translate_enable_formatting_fixes is True
    assert resolved.abbyy_profile == "aggressive"

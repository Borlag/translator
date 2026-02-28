from __future__ import annotations

from docx import Document

import docxru.pipeline as pipeline
from docxru.config import PipelineConfig


def test_postformat_docx_runs_layout_passes_on_existing_output(monkeypatch, tmp_path):
    input_path = tmp_path / "translated.docx"
    output_path = tmp_path / "final.docx"

    doc = Document()
    p = doc.add_paragraph("Translated paragraph")
    p.paragraph_format.space_before = None
    p.paragraph_format.space_after = None
    doc.save(str(input_path))

    call_order: list[str] = []

    def _fake_abbyy_and_layout(doc_obj, segments, cfg, logger):  # noqa: ANN001, ANN202
        del doc_obj, cfg, logger
        call_order.append("layout")
        assert segments
        assert all(seg.target_tagged for seg in segments)

    def _fake_final_cleanup(segments):  # noqa: ANN001, ANN202
        call_order.append("cleanup")
        assert segments
        return 0

    def _fake_com(output, cfg, logger):  # noqa: ANN001, ANN202
        del output, cfg, logger
        call_order.append("com")

    monkeypatch.setattr(pipeline, "_apply_abbyy_and_layout_passes", _fake_abbyy_and_layout)
    monkeypatch.setattr(pipeline, "_apply_final_run_level_cleanup", _fake_final_cleanup)
    monkeypatch.setattr(pipeline, "_run_word_com_postprocess", _fake_com)

    cfg = PipelineConfig(mode="reflow", abbyy_profile="full", layout_check=True, layout_auto_fix=True)
    pipeline.postformat_docx(input_path=input_path, output_path=output_path, cfg=cfg)

    assert output_path.exists()
    assert call_order == ["layout", "cleanup", "com"]


def test_postformat_docx_runs_layout_before_global_font_shrink(monkeypatch, tmp_path):
    input_path = tmp_path / "translated.docx"
    output_path = tmp_path / "final.docx"

    doc = Document()
    doc.add_paragraph("Translated paragraph")
    doc.save(str(input_path))

    call_order: list[str] = []

    def _fake_abbyy_and_layout(doc_obj, segments, cfg, logger):  # noqa: ANN001, ANN202
        del doc_obj, segments, cfg, logger
        call_order.append("layout")

    def _fake_shrink(segments, cfg):  # noqa: ANN001, ANN202
        del segments, cfg
        call_order.append("shrink")
        return 1

    def _fake_final_cleanup(segments):  # noqa: ANN001, ANN202
        del segments
        call_order.append("cleanup")
        return 0

    def _fake_com(output, cfg, logger):  # noqa: ANN001, ANN202
        del output, cfg, logger
        call_order.append("com")

    monkeypatch.setattr(pipeline, "_apply_abbyy_and_layout_passes", _fake_abbyy_and_layout)
    monkeypatch.setattr(pipeline, "apply_global_font_shrink", _fake_shrink)
    monkeypatch.setattr(pipeline, "_apply_final_run_level_cleanup", _fake_final_cleanup)
    monkeypatch.setattr(pipeline, "_run_word_com_postprocess", _fake_com)

    cfg = PipelineConfig(
        mode="reflow",
        abbyy_profile="full",
        layout_check=True,
        layout_auto_fix=True,
        font_shrink_body_pt=0.3,
    )
    pipeline.postformat_docx(input_path=input_path, output_path=output_path, cfg=cfg)

    assert output_path.exists()
    assert call_order == ["layout", "shrink", "cleanup", "com"]


def test_postformat_docx_rejects_non_docx_input(tmp_path):
    bad_input = tmp_path / "translated.txt"
    bad_input.write_text("not a docx", encoding="utf-8")
    output_path = tmp_path / "final.docx"

    cfg = PipelineConfig()
    try:
        pipeline.postformat_docx(input_path=bad_input, output_path=output_path, cfg=cfg)
    except RuntimeError as exc:
        assert "DOCX" in str(exc)
    else:
        raise AssertionError("postformat_docx should reject non-DOCX input")

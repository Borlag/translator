from __future__ import annotations

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docxru.section_merge import (
    NS,
    build_page_mapping,
    get_section_ranges,
    parse_pages_spec,
    replace_pages,
)


def _append_textbox_text(paragraph, text: str) -> None:
    run = paragraph.add_run("")
    txbx_content = OxmlElement("w:txbxContent")
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    txbx_content.append(p)
    run._r.append(txbx_content)


def _build_base_doc() -> Document:
    doc = Document()
    doc.add_paragraph("alpha one")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("alpha two")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("bravo one")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("charlie")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("delta")
    return doc


def _build_overlay_doc() -> Document:
    doc = Document()
    p1 = doc.add_paragraph("")
    _append_textbox_text(p1, "alpha one alpha two")
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    p2 = doc.add_paragraph("")
    _append_textbox_text(p2, "bravo one")
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    p3 = doc.add_paragraph("")
    _append_textbox_text(p3, "charlie delta")
    return doc


def test_parse_pages_spec():
    assert parse_pages_spec("5,12,45-50,100") == [5, 12, 45, 46, 47, 48, 49, 50, 100]


def test_build_page_mapping_and_replace_pages(tmp_path):
    base_doc = _build_base_doc()
    overlay_doc = _build_overlay_doc()

    mapping = build_page_mapping(base_doc, overlay_doc, mode="greedy_group")
    assert mapping.mode == "greedy_group"
    assert mapping.mapping[1] == [1, 2]
    assert mapping.mapping[2] == [3]
    assert mapping.mapping[3] == [4, 5]

    stats = replace_pages(
        base_doc,
        overlay_doc,
        pages=[1, 3],
        page_mapping=mapping.mapping,
        keep_page_size="target",
    )
    assert stats["replaced_pages"] == [1, 3]
    assert stats["replaced_pages_count"] == 2
    assert stats["replaced_base_sections_count"] == 4

    txbx_count = int(base_doc._element.body.xpath("count(.//w:txbxContent//w:t)"))
    assert txbx_count > 0

    out_path = tmp_path / "merged.docx"
    base_doc.save(str(out_path))
    reopened = Document(str(out_path))
    assert len(get_section_ranges(reopened)) == 3


def test_build_page_mapping_dp_one_to_one_is_monotonic():
    base_doc = _build_base_doc()
    overlay_doc = _build_overlay_doc()

    mapping = build_page_mapping(base_doc, overlay_doc, mode="dp_one_to_one")
    assert mapping.mode == "dp_one_to_one"
    assert len(mapping.mapping) == 3

    targets = [mapping.mapping[i][0] for i in (1, 2, 3)]
    assert all(len(mapping.mapping[i]) == 1 for i in (1, 2, 3))
    assert targets == sorted(targets)
    assert all(1 <= target <= 5 for target in targets)


def test_replace_pages_preserves_target_section_break_type():
    base_doc = _build_base_doc()
    overlay_doc = _build_overlay_doc()

    # Force source first section to continuous and ensure merge keeps target section break semantics.
    source_first = get_section_ranges(overlay_doc)[0].sect_pr
    for existing in source_first.findall("./w:type", namespaces=NS):
        source_first.remove(existing)
    src_type = OxmlElement("w:type")
    src_type.set(qn("w:val"), "continuous")
    source_first.append(src_type)

    replace_pages(
        base_doc,
        overlay_doc,
        pages=[1],
        page_mapping={1: [2]},
        keep_page_size="source",
    )

    target_second = get_section_ranges(base_doc)[1].sect_pr
    assert target_second.find("./w:type", namespaces=NS) is None


def test_replace_pages_copies_overlay_styles():
    base_doc = _build_base_doc()
    overlay_doc = _build_overlay_doc()

    overlay_style = overlay_doc.styles.add_style("OverlayBody", WD_STYLE_TYPE.PARAGRAPH)
    overlay_para = overlay_doc.paragraphs[0]
    overlay_para.style = overlay_style

    base_style_ids_before = {
        style.get(qn("w:styleId"))
        for style in base_doc.part._styles_part.element.xpath("./w:style")
        if style.get(qn("w:styleId"))
    }
    assert overlay_style.style_id not in base_style_ids_before

    replace_pages(
        base_doc,
        overlay_doc,
        pages=[1],
        page_mapping={1: [2]},
        keep_page_size="source",
    )

    base_style_ids_after = {
        style.get(qn("w:styleId"))
        for style in base_doc.part._styles_part.element.xpath("./w:style")
        if style.get(qn("w:styleId"))
    }
    assert overlay_style.style_id in base_style_ids_after

    replaced_section = get_section_ranges(base_doc)[1]
    used_pstyles = {
        style_id
        for block in replaced_section.blocks
        for style_id in block.xpath(".//w:pStyle/@w:val")
    }
    assert overlay_style.style_id in used_pstyles

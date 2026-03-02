"""Tests for strip_blanks module: empty paragraph removal + soft break cleanup."""

from __future__ import annotations

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docxru.strip_blanks import (
    _can_remove,
    _is_empty_paragraph,
    _is_protected_paragraph,
    cleanup_soft_linebreaks,
    strip_blank_paragraphs,
    strip_blanks_and_breaks,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_doc_with_paragraphs(*texts: str) -> Document:
    """Create a Document with the given paragraph texts ('' = empty)."""
    doc = Document()
    # python-docx creates a default empty paragraph; remove it
    for child in list(doc.element.body):
        doc.element.body.remove(child)
    for t in texts:
        doc.add_paragraph(t)
    return doc


def _add_sect_pr_to_paragraph(paragraph) -> None:
    """Add a w:sectPr element to the paragraph's pPr (simulates section break)."""
    p_elem = paragraph._p
    p_pr = p_elem.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p_elem.insert(0, p_pr)
    sect_pr = OxmlElement("w:sectPr")
    p_pr.append(sect_pr)


def _add_page_break_to_paragraph(paragraph) -> None:
    """Add a w:br type=page inside the paragraph."""
    run = paragraph.add_run("")
    br_elem = OxmlElement("w:br")
    br_elem.set(qn("w:type"), "page")
    run._r.append(br_elem)


def _add_bookmark_to_paragraph(paragraph) -> None:
    """Add a w:bookmarkStart element to the paragraph."""
    bm = OxmlElement("w:bookmarkStart")
    bm.set(qn("w:id"), "0")
    bm.set(qn("w:name"), "test_bookmark")
    paragraph._p.append(bm)


def _add_drawing_to_paragraph(paragraph) -> None:
    """Add a w:drawing element inside a run."""
    run = paragraph.add_run("")
    drawing = OxmlElement("w:drawing")
    run._r.append(drawing)


def _add_soft_break_to_run(run) -> None:
    """Add a <w:br/> (soft line break) to a run."""
    br_elem = OxmlElement("w:br")
    run._r.append(br_elem)


def _paragraph_texts(doc: Document) -> list[str]:
    """Return list of paragraph texts from the document body."""
    return [p.text for p in doc.paragraphs]


# ===========================================================================
# Pass 1: _is_empty_paragraph / _is_protected_paragraph
# ===========================================================================


class TestIsEmptyParagraph:
    def test_empty_no_runs(self):
        doc = Document()
        p = doc.add_paragraph("")
        assert _is_empty_paragraph(p) is True

    def test_whitespace_only(self):
        doc = Document()
        p = doc.add_paragraph("   \t  ")
        assert _is_empty_paragraph(p) is True

    def test_with_text(self):
        doc = Document()
        p = doc.add_paragraph("Hello")
        assert _is_empty_paragraph(p) is False

    def test_with_drawing_still_looks_empty_text(self):
        doc = Document()
        p = doc.add_paragraph("")
        _add_drawing_to_paragraph(p)
        # Text is empty, but drawing is present — _is_empty_paragraph returns True
        # (the protection check is in _is_protected_paragraph)
        assert _is_empty_paragraph(p) is True


class TestIsProtectedParagraph:
    def test_normal_empty_not_protected(self):
        doc = Document()
        p = doc.add_paragraph("")
        assert _is_protected_paragraph(p) is False

    def test_sect_pr_is_protected(self):
        doc = Document()
        p = doc.add_paragraph("")
        _add_sect_pr_to_paragraph(p)
        assert _is_protected_paragraph(p) is True

    def test_page_break_is_protected(self):
        doc = Document()
        p = doc.add_paragraph("")
        _add_page_break_to_paragraph(p)
        assert _is_protected_paragraph(p) is True

    def test_bookmark_is_protected(self):
        doc = Document()
        p = doc.add_paragraph("")
        _add_bookmark_to_paragraph(p)
        assert _is_protected_paragraph(p) is True

    def test_drawing_is_protected(self):
        doc = Document()
        p = doc.add_paragraph("")
        _add_drawing_to_paragraph(p)
        assert _is_protected_paragraph(p) is True


class TestCanRemove:
    def test_empty_unprotected(self):
        doc = Document()
        p = doc.add_paragraph("")
        assert _can_remove(p) is True

    def test_empty_with_sect_pr(self):
        doc = Document()
        p = doc.add_paragraph("")
        _add_sect_pr_to_paragraph(p)
        assert _can_remove(p) is False

    def test_non_empty(self):
        doc = Document()
        p = doc.add_paragraph("text")
        assert _can_remove(p) is False


# ===========================================================================
# Pass 1: strip_blank_paragraphs
# ===========================================================================


class TestStripBlankParagraphsBody:
    def test_single_empty_kept_with_keep_singles(self):
        doc = _make_doc_with_paragraphs("A", "", "B")
        stats = strip_blank_paragraphs(doc, keep_singles=True, process_tables=False)
        assert stats["removed_body"] == 0
        assert _paragraph_texts(doc) == ["A", "", "B"]

    def test_single_empty_removed_without_keep_singles(self):
        doc = _make_doc_with_paragraphs("A", "", "B")
        stats = strip_blank_paragraphs(doc, keep_singles=False, process_tables=False)
        assert stats["removed_body"] == 1
        assert _paragraph_texts(doc) == ["A", "B"]

    def test_pair_reduced_to_one(self):
        doc = _make_doc_with_paragraphs("A", "", "", "B")
        stats = strip_blank_paragraphs(doc, max_consecutive=1, process_tables=False)
        assert stats["removed_body"] == 1
        assert _paragraph_texts(doc) == ["A", "", "B"]

    def test_pair_removed_entirely(self):
        doc = _make_doc_with_paragraphs("A", "", "", "B")
        stats = strip_blank_paragraphs(doc, max_consecutive=0, process_tables=False)
        assert stats["removed_body"] == 2
        assert _paragraph_texts(doc) == ["A", "B"]

    def test_group_of_four_reduced_to_one(self):
        doc = _make_doc_with_paragraphs("A", "", "", "", "", "B")
        stats = strip_blank_paragraphs(doc, max_consecutive=1, process_tables=False)
        assert stats["removed_body"] == 3
        assert _paragraph_texts(doc) == ["A", "", "B"]

    def test_large_block_reduced(self):
        texts = ["Before"] + [""] * 20 + ["After"]
        doc = _make_doc_with_paragraphs(*texts)
        stats = strip_blank_paragraphs(doc, max_consecutive=1, process_tables=False)
        assert stats["removed_body"] == 19
        assert _paragraph_texts(doc) == ["Before", "", "After"]

    def test_no_empty_noop(self):
        doc = _make_doc_with_paragraphs("A", "B", "C")
        stats = strip_blank_paragraphs(doc, process_tables=False)
        assert stats["removed_body"] == 0

    def test_preserves_order(self):
        doc = _make_doc_with_paragraphs("A", "", "", "B", "", "C")
        strip_blank_paragraphs(doc, keep_singles=True, max_consecutive=0, process_tables=False)
        # Pair between A-B removed entirely, single between B-C kept
        assert _paragraph_texts(doc) == ["A", "B", "", "C"]

    def test_sect_pr_paragraph_not_removed(self):
        doc = _make_doc_with_paragraphs("A", "", "", "B")
        # Make the first empty paragraph have a sectPr
        paras = list(doc.paragraphs)
        sect_p_elem = paras[1]._p
        _add_sect_pr_to_paragraph(paras[1])
        stats = strip_blank_paragraphs(doc, max_consecutive=0, process_tables=False)
        # The sectPr paragraph XML element should survive
        body_p_elems = [p._p for p in doc.paragraphs]
        assert sect_p_elem in body_p_elems
        texts = _paragraph_texts(doc)
        assert texts[0] == "A"
        assert texts[-1] == "B"

    def test_page_break_paragraph_not_removed(self):
        doc = _make_doc_with_paragraphs("A", "", "B")
        paras = list(doc.paragraphs)
        _add_page_break_to_paragraph(paras[1])
        stats = strip_blank_paragraphs(doc, keep_singles=False, max_consecutive=0, process_tables=False)
        # Protected — should not be removed even with aggressive settings
        assert stats["removed_body"] == 0


class TestStripBlankParagraphsTable:
    def test_table_cell_keeps_min_one_paragraph(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = ""  # single empty paragraph
        stats = strip_blank_paragraphs(doc, keep_singles=False, max_consecutive=0, process_tables=True)
        # Cell has only 1 paragraph — nothing removed
        assert len(cell.paragraphs) >= 1

    def test_table_cell_pair_removed(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].text = "Content"
        cell.add_paragraph("")
        cell.add_paragraph("")
        cell.add_paragraph("More")
        stats = strip_blank_paragraphs(doc, max_consecutive=0, process_tables=True)
        assert stats["removed_table"] == 2
        texts = [p.text for p in cell.paragraphs]
        assert texts == ["Content", "More"]

    def test_table_all_empty_keeps_one(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].text = ""
        cell.add_paragraph("")
        cell.add_paragraph("")
        stats = strip_blank_paragraphs(doc, keep_singles=False, max_consecutive=0, process_tables=True)
        # Must keep at least 1
        assert len(cell.paragraphs) >= 1

    def test_process_tables_false(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].text = ""
        cell.add_paragraph("")
        stats = strip_blank_paragraphs(doc, process_tables=False)
        assert stats["removed_table"] == 0


# ===========================================================================
# Pass 2: cleanup_soft_linebreaks
# ===========================================================================


class TestCleanupSoftLinebreaks:
    def test_soft_break_between_words_replaced_with_space(self):
        doc = Document()
        p = doc.add_paragraph("")
        run = p.add_run("hello")
        _add_soft_break_to_run(run)
        t2 = OxmlElement("w:t")
        t2.text = "world"
        run._r.append(t2)

        stats = cleanup_soft_linebreaks(doc, process_tables=False)
        assert stats["breaks_replaced_with_space"] == 1
        # Verify the break is gone and text contains space
        br_count = len(list(p._p.iter(qn("w:br"))))
        assert br_count == 0

    def test_soft_break_next_to_space_just_removed(self):
        doc = Document()
        p = doc.add_paragraph("")
        run = p.add_run("hello ")
        _add_soft_break_to_run(run)
        t2 = OxmlElement("w:t")
        t2.text = "world"
        run._r.append(t2)

        stats = cleanup_soft_linebreaks(doc, process_tables=False)
        assert stats["breaks_removed"] == 1
        assert stats["breaks_replaced_with_space"] == 0

    def test_page_break_not_touched(self):
        doc = Document()
        p = doc.add_paragraph("")
        run = p.add_run("text")
        br_elem = OxmlElement("w:br")
        br_elem.set(qn("w:type"), "page")
        run._r.append(br_elem)

        stats = cleanup_soft_linebreaks(doc, process_tables=False)
        assert stats["breaks_skipped_page_column"] == 1
        assert stats["breaks_replaced_with_space"] == 0
        assert stats["breaks_removed"] == 0
        # Verify the break is still there
        assert len(list(p._p.iter(qn("w:br")))) == 1

    def test_column_break_not_touched(self):
        doc = Document()
        p = doc.add_paragraph("")
        run = p.add_run("text")
        br_elem = OxmlElement("w:br")
        br_elem.set(qn("w:type"), "column")
        run._r.append(br_elem)

        stats = cleanup_soft_linebreaks(doc, process_tables=False)
        assert stats["breaks_skipped_page_column"] == 1

    def test_toc_paragraph_skipped(self):
        doc = Document()
        # Create a TOC-style paragraph
        p = doc.add_paragraph("", style="TOC Heading")
        run = p.add_run("Entry")
        _add_soft_break_to_run(run)
        t2 = OxmlElement("w:t")
        t2.text = "text"
        run._r.append(t2)

        stats = cleanup_soft_linebreaks(doc, skip_toc=True, process_tables=False)
        assert stats["paragraphs_skipped_toc"] == 1
        # Break should still be there
        assert len(list(p._p.iter(qn("w:br")))) == 1

    def test_cleanup_breaks_false_noop(self):
        doc = Document()
        p = doc.add_paragraph("")
        run = p.add_run("hello")
        _add_soft_break_to_run(run)
        t2 = OxmlElement("w:t")
        t2.text = "world"
        run._r.append(t2)

        stats = strip_blanks_and_breaks(doc, cleanup_breaks=False, strip_empty=False, process_tables=False)
        # No break stats should exist
        assert stats.get("breaks_replaced_with_space", 0) == 0

    def test_break_between_runs(self):
        """Soft break at end of one run, text starts in next run."""
        doc = Document()
        p = doc.add_paragraph("")
        run1 = p.add_run("hello")
        _add_soft_break_to_run(run1)
        run2 = p.add_run("world")

        stats = cleanup_soft_linebreaks(doc, process_tables=False)
        assert stats["breaks_replaced_with_space"] == 1


# ===========================================================================
# Combined function
# ===========================================================================


class TestStripBlanksAndBreaks:
    def test_both_passes_run(self):
        doc = _make_doc_with_paragraphs("A", "", "", "B")
        # Also add a soft break in paragraph "B"
        b_para = doc.paragraphs[3] if len(doc.paragraphs) > 3 else doc.paragraphs[-1]
        # Find the "B" paragraph
        for p in doc.paragraphs:
            if p.text == "B":
                b_para = p
                break
        run = b_para.runs[0] if b_para.runs else b_para.add_run("B")
        _add_soft_break_to_run(run)
        t2 = OxmlElement("w:t")
        t2.text = "extra"
        run._r.append(t2)

        stats = strip_blanks_and_breaks(doc, max_consecutive=0, process_tables=False)
        assert stats["removed_body"] == 2  # pair of empties
        assert stats["breaks_replaced_with_space"] >= 1

    def test_strip_empty_false(self):
        doc = _make_doc_with_paragraphs("A", "", "", "B")
        stats = strip_blanks_and_breaks(doc, strip_empty=False, cleanup_breaks=False, process_tables=False)
        assert "removed_body" not in stats
        assert _paragraph_texts(doc) == ["A", "", "", "B"]

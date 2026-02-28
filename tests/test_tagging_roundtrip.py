from __future__ import annotations

from docx import Document

from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docxru.tagging import paragraph_to_tagged, tagged_to_runs, tagged_to_runs_inplace


def _flags_from_run(run):
    flags = []
    if run.bold:
        flags.append("B")
    if run.italic:
        flags.append("I")
    if run.underline:
        flags.append("U")
    if run.font.superscript:
        flags.append("SUP")
    if run.font.subscript:
        flags.append("SUB")
    return tuple(flags)


def test_tagging_roundtrip_char_level_flags():
    doc = Document()
    p = doc.add_paragraph()

    r1 = p.add_run("Remove the ")
    r2 = p.add_run("bolt")
    r2.bold = True
    r3 = p.add_run(" and ")
    r4 = p.add_run("washer")
    r4.italic = True
    r5 = p.add_run(" PN ")
    r6 = p.add_run("201587001")
    r6.underline = True

    tagged, spans, inline_map = paragraph_to_tagged(p)

    orig_text = "".join(s.source_text for s in spans)
    orig_flags = []
    for s in spans:
        orig_flags.extend([s.flags] * len(s.source_text))

    # Apply roundtrip
    tagged_to_runs(p, tagged, spans, inline_run_map=inline_map)

    new_text = "".join(r.text for r in p.runs)
    new_flags = []
    for r in p.runs:
        new_flags.extend([_flags_from_run(r)] * len(r.text))

    assert new_text == orig_text
    assert new_flags == orig_flags


def test_tagging_roundtrip_preserves_column_break():
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Left")
    r1.add_break(WD_BREAK.COLUMN)
    p.add_run("Right")

    tagged, spans, inline_map = paragraph_to_tagged(p)
    # Must contain a column-break token
    assert "BRCOL" in tagged

    tagged_to_runs(p, tagged, spans, inline_run_map=inline_map)
    xml = p._p.xml
    assert "w:br" in xml and "w:type=\"column\"" in xml


def test_tagging_roundtrip_preserves_run_rpr_xml():
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Main Fitting")
    r.bold = True
    r.font.name = "Arial"
    orig_rpr = r._r.rPr.xml

    tagged, spans, inline_map = paragraph_to_tagged(p)
    assert spans and spans[0].rpr_xml is not None

    tagged_to_runs(p, tagged, spans, inline_run_map=inline_map)
    assert p.runs and p.runs[0]._r.rPr is not None
    assert p.runs[0]._r.rPr.xml == orig_rpr


def test_paragraph_to_tagged_merges_runs_with_same_visible_style():
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Main ")
    r2 = p.add_run("Fitting")
    r1.bold = True
    r2.bold = True

    # Make rPr XML different in a non-visual way.
    rpr2 = r2._r.get_or_add_rPr()
    no_proof = OxmlElement("w:noProof")
    rpr2.append(no_proof)

    tagged, spans, _ = paragraph_to_tagged(p)
    assert len(spans) == 1
    assert spans[0].source_text == "Main Fitting"
    assert "⟦S_1|B⟧Main Fitting⟦/S_1⟧" == tagged


def test_paragraph_to_tagged_allows_last_rendered_page_break():
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Subject")
    # Insert non-visible pagination cache marker used by Word.
    r._r.insert(0, OxmlElement("w:lastRenderedPageBreak"))

    tagged, spans, _ = paragraph_to_tagged(p)
    assert "OBJ_" not in tagged
    assert len(spans) == 1
    assert spans[0].source_text == "Subject"


def test_tagging_roundtrip_preserves_hyperlink_structure():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("See ")

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), "rId42")
    h_run = OxmlElement("w:r")
    h_text = OxmlElement("w:t")
    h_text.text = "manual"
    h_run.append(h_text)
    hyperlink.append(h_run)
    p._p.append(hyperlink)

    p.add_run(" now")

    tagged, spans, inline_map = paragraph_to_tagged(p)
    assert any(any(flag.startswith("HREF_") for flag in span.flags) for span in spans)

    tagged_to_runs(p, tagged, spans, inline_run_map=inline_map)
    xml = p._p.xml
    assert "<w:hyperlink" in xml
    assert "manual" in p.text


def test_tagged_to_runs_inplace_preserves_run_structure():
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Main ")
    r1.bold = True
    r2 = p.add_run("Fitting")
    r2.bold = True

    tagged, spans, _ = paragraph_to_tagged(p)
    assert spans and spans[0].original_run_lengths == (5, 7)

    translated_inner = "ОсновнойФитингУзла"
    translated_tagged = tagged.replace("Main Fitting", translated_inner)
    inplace_ok = tagged_to_runs_inplace(p, translated_tagged, spans)
    assert inplace_ok is True
    assert len(p.runs) == 2
    assert "".join(run.text for run in p.runs) == translated_inner
    assert p.runs[0].bold is True and p.runs[1].bold is True


def test_tagged_to_runs_inplace_returns_false_on_structure_changes():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Install")
    tagged, spans, _ = paragraph_to_tagged(p)

    # Prefix text outside style tags should force rebuild path.
    bad_tagged = "PREFIX " + tagged
    assert tagged_to_runs_inplace(p, bad_tagged, spans) is False
